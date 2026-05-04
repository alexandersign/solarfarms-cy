"""
Generate PV Installation Contract (Net Metering) for Tamara Morozova.
Based on generate_pv_residential.py (canonical template, 27/04/2026).
Changes vs. template:
  - All client details filled in
  - §1.2 subsidy bullet removed
  - §2 timeline updated to reflect EAC application already submitted + connection terms issued
  - §3.1 price filled in (€9,000 incl. VAT)
  - §3.2 payment schedule: 40% / 50% / 10%
  - §3.3 Government Subsidies removed entirely
  - §4 warranties filled with Jinko Tiger Neo + Huawei SUN2000 OEM specs
  - §8 changed from Net Billing to Net Metering (kWh-for-kWh, 15yr EAC contract)
  - Contractor capacity: Business Development Manager (not Director)
"""
import io
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand colours ──────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1A, 0x36, 0x5D)
GOLD  = RGBColor(0xC9, 0xA4, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x40, 0x40, 0x40)
LIGHT_GREY_BG = "F0F4F8"
NAVY_HEX  = "1A365D"
GOLD_HEX  = "C9A432"

OUTPUT = os.path.join(
    os.path.dirname(__file__),
    "..", "..", "docs", "clients", "Individual_Tamara_Morozova",
    "pv-installation-contract-dec2025.docx"
)
LOGO_SVG = os.path.join(os.path.dirname(__file__),
                        "..", "..", "public", "favicon.svg")

# ── Client & project data ──────────────────────────────────────────────────────
CLIENT = {
    "name":    "Tamara Morozova",
    "id":      "—",
    "address": "Lombardi Gardens Villas, Flat 3, 531A,\nProtara Cavo Greco, Cyprus",
    "email":   "—",
    "tel":     "+357 99933225",
    "rep":     "Tamara Morozova",
    "eac_no":  "3959063626",
}

CONTRACTOR = {
    "company": "LIGHTHIEF CYPRUS LTD",
    "reg":     "HE 477423",
    "tin":     "60187188Q",
    "address": "28 October Ave 249, Lophitis Business Center 1,\nOffice 201, 3035 Limassol, Cyprus",
    "email":   "office@lighthief.com",
    "tel":     "+357 77 77 00 50",
    "rep":     "Alexander Papacosta",
    "cap":     "Business Development Manager",
}

CONTRACT = {
    "ref":        "LCY-PV-TMR-001",
    "date":       "19 December 2025",
    "price_vat":  "€9,000.00",
    "price_net":  "€7,563.03",
    "price_tax":  "€1,436.97",
    "eac_fee":    "€365.00 + VAT 19% = €434.35",
}

SYSTEM = {
    "address":    "Protara, Kavo Greco, 531, Cyprus",
    "panels":     "16 × JinkoSolar Tiger Neo N-Type 590W",
    "kwp":        "9.44 kWp",
    "inverter":   "1 × Huawei SUN2000-10KTL-M1 (10 kW)",
    "battery":    "Not applicable",
    "bat_cap":    "Not applicable",
    "ats":        "☐ Included   ☑ Not included",
    "monitoring": "☑ Included — Huawei FusionSolar app",
    "yield":      "Approximately 14,500 kWh/year (indicative — see §1.4)",
}

WARRANTY = {
    "panels_product":  "15 years from delivery date",
    "panels_perf":     "30 years — linear, ≥87.4% rated power at year 30",
    "inverter":        "10 years from commissioning (extendable to 20 yr)",
    "mounting":        "10 years from installation date",
    "workmanship":     "5 years from commissioning date",
    "statutory":       "2 years from delivery (Cyprus Law 154(I)/2021)",
}


# ── Helpers (identical to generate_pv_residential.py) ─────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=True, bottom=True, left=True, right=True,
                     color="CCCCCC", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, active in [("top", top), ("bottom", bottom),
                         ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if active:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def para_spacing(para, before=0, after=60, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(13)
        run.font.color.rgb = GOLD
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), NAVY_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
    para_spacing(p, before=160, after=60)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    para_spacing(p, before=100, after=40)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    para_spacing(p, before=0, after=50)
    return p


def add_bullet(doc, text, indent=0.4):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    p.paragraph_format.left_indent = Cm(indent)
    para_spacing(p, before=0, after=30)
    return p


def add_notice_box(doc, text, bg=LIGHT_GREY_BG, border_hex=NAVY_HEX):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_borders(cell, color=border_hex, size="12")
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = BLACK
    para_spacing(p, before=60, after=60)
    doc.add_paragraph()
    return tbl


def add_two_col_table(doc, rows_data, header_row=None, col_widths=None):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    if header_row:
        row = tbl.add_row()
        for i, txt in enumerate(header_row):
            c = row.cells[i]
            set_cell_bg(c, NAVY_HEX)
            p = c.paragraphs[0]
            run = p.add_run(txt)
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = WHITE
            para_spacing(p, before=40, after=40)
    for label, value in rows_data:
        row = tbl.add_row()
        for i, txt in enumerate([label, value]):
            c = row.cells[i]
            p = c.paragraphs[0]
            if i == 0:
                run = p.add_run(txt)
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = NAVY
                set_cell_bg(c, "F5F7FA")
            else:
                run = p.add_run(txt)
                run.font.size = Pt(9.5)
                run.font.color.rgb = BLACK
            para_spacing(p, before=40, after=40)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return tbl


def add_payment_table(doc, rows):
    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = "Table Grid"
    hrow = tbl.add_row()
    for i, txt in enumerate(["Instalment", "Trigger / Milestone", "Amount"]):
        c = hrow.cells[i]
        set_cell_bg(c, NAVY_HEX)
        p = c.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, before=40, after=40)
    for inst, trigger, pct in rows:
        row = tbl.add_row()
        for i, (txt, bold) in enumerate([(inst, True), (trigger, False), (pct, True)]):
            c = row.cells[i]
            p = c.paragraphs[0]
            run = p.add_run(txt)
            run.bold = bold
            run.font.size = Pt(9.5)
            run.font.color.rgb = NAVY if bold else BLACK
            if i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_spacing(p, before=40, after=40)
    for row in tbl.rows:
        row.cells[0].width = Inches(1.4)
        row.cells[1].width = Inches(4.0)
        row.cells[2].width = Inches(1.1)
    doc.add_paragraph()
    return tbl


def add_horizontal_rule(doc, color_hex=NAVY_HEX):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    para_spacing(p, before=0, after=80)
    return p


def logo_png_bytes():
    from svglib.svglib import svg2rlg
    from reportlab.graphics import renderPM
    with open(LOGO_SVG, "r") as f:
        svg_text = f.read()
    svg_text = svg_text.replace(
        '<rect width="512" height="512" fill="#0B0B0B"/>',
        '<rect width="512" height="512" fill="#1A365D"/>'
    )
    svg_text = svg_text.replace('<rect width="256" height="256" fill="#0B0B0B"/>', '')
    svg_text = svg_text.replace('fill="white"', 'fill="#C9A432"')
    tmp = os.path.join(os.path.dirname(__file__), "_tmp_logo.svg")
    with open(tmp, "w") as f:
        f.write(svg_text)
    drawing = svg2rlg(tmp)
    buf = io.BytesIO()
    renderPM.drawToFile(drawing, buf, fmt="PNG")
    buf.seek(0)
    os.remove(tmp)
    return buf


def add_header(doc):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    logo_cell = tbl.cell(0, 0)
    text_cell = tbl.cell(0, 1)
    set_cell_bg(logo_cell, NAVY_HEX)
    set_cell_bg(text_cell, NAVY_HEX)
    logo_cell.width = Inches(1.1)
    text_cell.width = Inches(5.4)

    try:
        img_buf = logo_png_bytes()
        p_logo = logo_cell.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(img_buf, width=Inches(0.75))
        para_spacing(p_logo, before=60, after=60)
    except Exception:
        p_logo = logo_cell.paragraphs[0]
        run = p_logo.add_run("☀")
        run.font.size = Pt(32)
        run.font.color.rgb = GOLD
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1 = text_cell.paragraphs[0]
    r1 = p1.add_run("LIGHTHIEF CYPRUS LTD")
    r1.bold = True
    r1.font.size = Pt(16)
    r1.font.color.rgb = GOLD
    para_spacing(p1, before=40, after=0)

    p2 = text_cell.add_paragraph()
    r2 = p2.add_run("Solar Photovoltaic System Installation Contract — Net Metering")
    r2.font.size = Pt(10)
    r2.font.color.rgb = WHITE
    para_spacing(p2, before=0, after=0)

    p3 = text_cell.add_paragraph()
    r3 = p3.add_run(f"Ref: {CONTRACT['ref']}  ·  Contract Date: {CONTRACT['date']}")
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = RGBColor(0xB0, 0xBE, 0xD4)
    para_spacing(p3, before=0, after=60)

    doc.add_paragraph()
    return tbl


# ── Main document ──────────────────────────────────────────────────────────────

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = BLACK

    # ════════════════════════════════════════════════════════════════════════
    # HEADER
    # ════════════════════════════════════════════════════════════════════════
    add_header(doc)
    add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════════════
    # PARTIES
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "PARTIES")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c_contractor = tbl.cell(0, 0)
    c_client     = tbl.cell(0, 1)
    set_cell_bg(c_contractor, "F0F4F8")
    set_cell_bg(c_client,     "FAFBFC")

    def party_block(cell, title, rows):
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY
        para_spacing(p, before=40, after=20)
        for label, val in rows:
            p2 = cell.add_paragraph()
            r_lbl = p2.add_run(label + ": ")
            r_lbl.bold = True; r_lbl.font.size = Pt(9); r_lbl.font.color.rgb = GREY
            r_val = p2.add_run(val)
            r_val.font.size = Pt(9); r_val.font.color.rgb = BLACK
            para_spacing(p2, before=0, after=20)

    party_block(c_contractor, "CONTRACTOR", [
        ("Company",        CONTRACTOR["company"]),
        ("Reg. No.",       CONTRACTOR["reg"]),
        ("TIN",            CONTRACTOR["tin"]),
        ("Address",        CONTRACTOR["address"]),
        ("Email",          CONTRACTOR["email"]),
        ("Tel",            CONTRACTOR["tel"]),
        ("Representative", CONTRACTOR["rep"] + ", " + CONTRACTOR["cap"]),
    ])
    party_block(c_client, "CLIENT", [
        ("Full Name",      CLIENT["name"]),
        ("Address",        CLIENT["address"]),
        ("Tel",            CLIENT["tel"]),
        ("EAC Customer No.", CLIENT["eac_no"]),
        ("Representative", CLIENT["rep"]),
    ])
    for row in tbl.rows:
        row.cells[0].width = Inches(3.25)
        row.cells[1].width = Inches(3.25)
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run(
        "Contractor and Client are referred to individually as a \"Party\" and collectively as "
        "the \"Parties.\" This Contract is governed by the laws of the Republic of Cyprus."
    )
    r.font.size = Pt(9.5); r.font.color.rgb = GREY; r.italic = True
    para_spacing(p, before=0, after=80)
    add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 1. SCOPE OF WORK
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1.  SCOPE OF WORK")

    add_subheading(doc, "1.1  System Specifications")
    add_body(doc, (
        "The Contractor shall supply, deliver, install, and commission a grid-connected photovoltaic "
        "system (the \"System\") at the premises specified below. The scope includes all labour, "
        "materials, cabling, protection devices, programming, single-line diagram, and EAC Net "
        "Metering application pack."
    ))

    # Spec table (filled in for Tamara)
    tbl_spec = doc.add_table(rows=0, cols=2)
    tbl_spec.style = "Table Grid"
    hrow = tbl_spec.add_row()
    for i, txt in enumerate(["Parameter", "Specification"]):
        c = hrow.cells[i]
        set_cell_bg(c, NAVY_HEX)
        p = c.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True; run.font.size = Pt(9.5); run.font.color.rgb = WHITE
        para_spacing(p, before=40, after=40)
    specs = [
        ("Installation Address",               SYSTEM["address"]),
        ("Solar Panels",                       SYSTEM["panels"]),
        ("Total PV Capacity",                  SYSTEM["kwp"]),
        ("Inverter",                           SYSTEM["inverter"]),
        ("Battery Storage",                    SYSTEM["battery"]),
        ("Mounting System",                    "Flat roof"),
        ("Remote Monitoring",                  SYSTEM["monitoring"]),
        ("Net Metering Application",           "☑ Included"),
        ("Estimated Annual Yield (indicative)", SYSTEM["yield"]),
    ]
    for label, val in specs:
        row = tbl_spec.add_row()
        c0 = row.cells[0]; c1 = row.cells[1]
        set_cell_bg(c0, "F5F7FA")
        p0 = c0.paragraphs[0]; p1 = c1.paragraphs[0]
        r0 = p0.add_run(label); r0.bold = True; r0.font.size = Pt(9.5); r0.font.color.rgb = NAVY
        r1 = p1.add_run(val);   r1.font.size = Pt(9.5);   r1.font.color.rgb = BLACK
        para_spacing(p0, before=40, after=40); para_spacing(p1, before=40, after=40)
    for row in tbl_spec.rows:
        row.cells[0].width = Inches(2.2); row.cells[1].width = Inches(4.3)
    doc.add_paragraph()

    add_subheading(doc, "1.2  Services Included")
    add_body(doc, "The Contractor shall provide:")
    for item in [
        "Site assessment and system design;",
        "Submission of Net Metering application to EAC (Electricity Authority of Cyprus);",
        "Handling of all necessary permits and approvals;",
        "Supply of all equipment and materials;",
        "Professional installation in accordance with Cyprus standards;",
        "System commissioning and testing;",
        "Training on system operation and monitoring.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "1.3  Installation Standards")
    add_body(doc, "All work shall comply with:")
    for item in [
        "Cyprus Electricity Authority (EAC) regulations;",
        "Cyprus Energy Regulatory Authority (CERA) requirements;",
        "EU safety and quality standards;",
        "Local building and planning regulations;",
        "All applicable Cyprus laws and regulations.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "1.4  Yield Disclaimer")
    add_notice_box(doc, (
        "IMPORTANT: The estimated annual yield figure above is indicative only. Actual energy "
        "generation depends on solar irradiation, shading, weather conditions, panel orientation, "
        "system degradation, and grid availability. The Contractor makes no warranty as to actual "
        "energy output unless a written performance guarantee is separately executed."
    ), bg="FFF8E1", border_hex=GOLD_HEX)

    # ════════════════════════════════════════════════════════════════════════
    # 2. TIMELINE AND MILESTONES
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2.  TIMELINE AND MILESTONES")

    add_subheading(doc, "2.1  Project Timeline")
    rows_tl = [
        ("Contract Signing Date",          CONTRACT["date"]),
        ("EAC Application Submission",
         "Submitted prior to application deadline — COMPLETED"),
        ("Connection Terms (Preliminary)", "Issued by EAC/DSO — ISSUED"),
        ("Equipment Delivery",             "Within agreed programme; Client notified ≥ 3 days in advance"),
        ("Installation",                   "Within 4 months of contract signing; duration 2–3 working days"),
        ("EAC Inspection",                 "To be scheduled by EAC following installation"),
        ("System Commissioning",           "Upon successful EAC inspection and grid connection approval"),
    ]
    add_two_col_table(doc, rows_tl, header_row=["Milestone", "Status / Target"],
                      col_widths=[2.4, 4.1])

    add_subheading(doc, "2.2  Delays")
    add_body(doc, "The Contractor shall not be liable for delays caused by:")
    for item in [
        "EAC processing times;",
        "Weather conditions or force majeure events (see §13.6);",
        "Client-requested changes;",
        "Delays in obtaining permits beyond the Contractor's reasonable control;",
        "Client's failure to provide required documents or access.",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # 3. PAYMENT TERMS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3.  PAYMENT TERMS")

    add_subheading(doc, "3.1  Total Contract Price (VAT Inclusive)")
    add_two_col_table(doc, [
        ("Total Contract Price (VAT incl. 19%)", CONTRACT["price_vat"]),
        ("Net Amount (excl. VAT)",               CONTRACT["price_net"]),
        ("VAT (19%)",                            CONTRACT["price_tax"]),
    ], col_widths=[3.5, 3.0])

    add_subheading(doc, "3.2  Payment Schedule")
    add_body(doc, (
        "Payments shall be made by bank transfer to the Contractor's account. "
        "No works shall commence or materials ordered until the Deposit (Instalment 1) is received. "
        "All amounts are VAT-inclusive."
    ))
    add_payment_table(doc, [
        ("#1 — Deposit (40%)",
         "Due upon contract signing",
         "40%\n€3,600.00"),
        ("#2 — Installation (50%)",
         "Due upon completion of physical installation works.",
         "50%\n€4,500.00"),
        ("#3 — Final (10%)",
         "Due upon EAC inspection and grid connection approval.",
         "10%\n€900.00"),
    ])

    add_subheading(doc, "3.2.1  Late Payment Interest")
    add_body(doc, (
        "If the Client fails to make any payment when due, interest shall accrue on the overdue "
        "amount from the due date at the rate prescribed by EU Directive 2011/7/EU — being the "
        "European Central Bank main refinancing rate plus eight percentage points (8%) per annum. "
        "The Contractor may suspend works upon seven (7) days' written notice of non-payment."
    ))

    add_subheading(doc, "3.2.2  EAC/DSO Connection Application Fee")
    add_notice_box(doc, (
        "Not included in the Contract Price: The EAC/DSO one-time connection application fee "
        "(covering DSO processing, inspection, system connection to the distribution network, "
        "and ripple control connection) is payable directly by the Client to EAC upon issuance "
        "of connection terms. Current fee: " + CONTRACT["eac_fee"] + ". "
        "This fee is set by EAC/CERA and may change without notice."
    ), bg="FFF8E1", border_hex=GOLD_HEX)

    # §3.3 Government Subsidies intentionally omitted — client cannot apply

    # ════════════════════════════════════════════════════════════════════════
    # 4. WARRANTIES AND GUARANTEES
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4.  WARRANTIES AND GUARANTEES")

    add_subheading(doc, "4.1  Warranty Summary")
    tbl_w = doc.add_table(rows=0, cols=3)
    tbl_w.style = "Table Grid"
    hrow = tbl_w.add_row()
    for i, txt in enumerate(["Item", "Coverage", "Basis"]):
        c = hrow.cells[i]
        set_cell_bg(c, NAVY_HEX)
        p = c.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True; run.font.size = Pt(9.5); run.font.color.rgb = WHITE
        para_spacing(p, before=40, after=40)
    warranty_data = [
        ("Workmanship",
         WARRANTY["workmanship"],
         "Contractual (above statutory minimum)"),
        ("JinkoSolar Tiger Neo 590W\n— Product (defect)",
         WARRANTY["panels_product"],
         "JinkoSolar OEM warranty"),
        ("JinkoSolar Tiger Neo 590W\n— Power output",
         WARRANTY["panels_perf"],
         "JinkoSolar linear performance guarantee"),
        ("Huawei SUN2000-10KTL-M1\n— Inverter",
         WARRANTY["inverter"],
         "Huawei OEM warranty"),
        ("Mounting System",
         WARRANTY["mounting"],
         "Manufacturer/installer warranty"),
        ("All equipment (goods)",
         WARRANTY["statutory"],
         "Cyprus Law 154(I)/2021 — statutory minimum"),
    ]
    for item, cov, basis in warranty_data:
        row = tbl_w.add_row()
        for i, txt in enumerate([item, cov, basis]):
            c = row.cells[i]
            p = c.paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9.5)
            run.font.color.rgb = NAVY if i == 0 else BLACK
            if i == 0:
                run.bold = True
            para_spacing(p, before=40, after=40)
    for row in tbl_w.rows:
        row.cells[0].width = Inches(1.7)
        row.cells[1].width = Inches(2.8)
        row.cells[2].width = Inches(2.0)
    doc.add_paragraph()

    add_subheading(doc, "4.2  Statutory Legal Guarantee (Cyprus Law 154(I)/2021)")
    add_body(doc, (
        "In accordance with Cyprus Law 154(I)/2021 implementing EU Directive 2019/771 on the "
        "sale of goods, the Client is entitled to a minimum two-year legal guarantee from the "
        "date of delivery of each item of equipment, covering any lack of conformity. This "
        "statutory guarantee is in addition to any commercial OEM warranty and cannot be excluded "
        "or reduced by contract."
    ))

    add_subheading(doc, "4.3  Workmanship Warranty")
    add_body(doc, (
        "The Contractor warrants all installation workmanship for five (5) years from the "
        "commissioning date. This covers defects directly attributable to the quality of "
        "installation performed by the Contractor's personnel. During this period, the "
        "Contractor shall remedy any covered defect within a reasonable time following written "
        "notice from the Client, at no additional cost."
    ))

    add_subheading(doc, "4.4  Warranty Conditions and Exclusions")
    add_body(doc, "Warranties are void or limited if:")
    for item in [
        "The system is modified or repaired without the Contractor's prior written consent;",
        "Damage is caused by Client negligence or misuse;",
        "Damage is caused by acts of God, fire, flood, or other events beyond both Parties' control;",
        "The Client has failed to fulfil payment obligations under §3.",
    ]:
        add_bullet(doc, item)
    add_body(doc, (
        "Warranties do not cover routine consumables, panel soiling, shading from new structures "
        "erected after installation, or performance variations within the OEM's published tolerance range."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # 5. CLIENT OBLIGATIONS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5.  CLIENT OBLIGATIONS")
    add_body(doc, "The Client shall:")
    for item in [
        "Provide accurate electricity consumption data (last 6 EAC bills);",
        "Ensure clear, safe, and unobstructed access to the installation site;",
        "Obtain and provide to the Contractor all necessary consents from co-owners, neighbours, "
        "or property management authorities before installation commences;",
        "Provide all necessary documentation required for the EAC Net Metering application in a "
        "timely manner;",
        "Make payments in accordance with the agreed schedule in §3.2;",
        "Not modify or repair the system without the Contractor's prior written approval;",
        "Notify the Contractor of any distribution board extension requirements before "
        "installation commences — such works, if required, are not included in the Contract "
        "Price and will be quoted separately. Any modifications required to the electrical "
        "distribution board or the siting and shading mitigation of the inverter are the "
        "Client's sole responsibility; the Contractor is available to advise upon request;",
        "Maintain appropriate property insurance covering the installed system.",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # 6. CONTRACTOR OBLIGATIONS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6.  CONTRACTOR OBLIGATIONS")
    add_body(doc, "The Contractor shall:")
    for item in [
        "Perform all work with professional competence and in accordance with applicable standards;",
        "Use only certified, high-quality equipment conforming to EU CE marking requirements;",
        "Comply with all applicable safety regulations;",
        "Obtain all necessary permits and regulatory approvals within its scope;",
        "Provide the Client with all required documentation and certificates upon commissioning;",
        "Train the Client on system operation and monitoring;",
        "Respond to warranty claims in writing within ten (10) Business Days of notification;",
        "Replace any roof tiles broken or directly damaged during installation with tiles of the "
        "same style and colour (see §14.1).",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # 7. MAINTENANCE AND SUPPORT
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7.  MAINTENANCE AND SUPPORT")

    add_subheading(doc, "7.1  Optional Maintenance Services")
    add_body(doc, "The Contractor offers optional maintenance packages, including:")
    for item in [
        "Annual inspection and cleaning;",
        "Performance monitoring and reporting;",
        "Priority response for service calls.",
    ]:
        add_bullet(doc, item)
    add_body(doc, "Maintenance packages are quoted separately and are not included in the Contract Price.")

    add_subheading(doc, "7.2  Client Self-Maintenance")
    add_body(doc, "The Client should:")
    for item in [
        "Keep panels reasonably clean and free from obstruction;",
        "Monitor system performance via the Huawei FusionSolar monitoring application;",
        "Report any faults, alarms, or unusual readings promptly;",
        "Not install new structures that shade the panels.",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # 8. NET METERING TERMS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8.  NET METERING TERMS")

    add_subheading(doc, "8.1  EAC Net Metering Contract")
    add_body(doc, (
        "Upon grid connection approval, the Client will enter into a Net Metering contract "
        "directly with EAC for a period of fifteen (15) years (residential category), in "
        "accordance with the EAC/CERA Net Metering framework applicable at the time of connection."
    ))

    add_subheading(doc, "8.2  How Net Metering Works")
    add_body(doc, (
        "Under the EAC Net Metering framework, electricity produced by the photovoltaic system "
        "is first consumed on-site. Surplus electricity exported to the EAC distribution network "
        "is credited against the Client's import consumption on a kilowatt-hour-for-kilowatt-hour "
        "basis at the full retail tariff applicable to the Client's supply category:"
    ))
    for item in [
        "Electricity imported from the grid is charged at the standard EAC retail tariff;",
        "Electricity exported to the grid offsets import consumption kWh-for-kWh at the same "
        "retail tariff rate — this is the key distinction of Net Metering from Net Billing;",
        "Net Metering credits are applied against the Client's EAC bill at the end of each "
        "two-month billing period;",
        "Unused credits carry forward and are valid for thirty-six (36) months;",
        "Credits are cleared every three (3) years; any excess credits not offset against "
        "consumption are not compensated in cash after the clearing period.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "8.3  Annual Net Metering Fee")
    add_body(doc, (
        "The Client shall pay EAC an annual Net Metering administrative fee of "
        "€47.23 per kWp installed + VAT (19%) per year. For this system (9.44 kWp), the "
        "indicative annual fee is approximately €447.93 + VAT = €532.03 per year. "
        "This fee is set by EAC/CERA and is subject to regulatory adjustment."
    ))

    add_subheading(doc, "8.4  EAC Tariff Changes")
    add_body(doc, (
        "EAC/CERA tariffs, fees, and Net Metering scheme conditions are subject to regulatory "
        "change. The Contractor has no control over and accepts no liability for any future "
        "adjustments to Net Metering tariffs, credit rates, scheme duration, or eligibility "
        "conditions enacted by EAC or CERA after the contract date."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # 9. INSURANCE AND LIABILITY
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9.  INSURANCE AND LIABILITY")

    add_subheading(doc, "9.1  Contractor Insurance")
    add_body(doc, "The Contractor maintains the following insurance policies:")
    for item in [
        "Professional Liability (Errors & Omissions) Insurance;",
        "Public / General Liability Insurance;",
        "Workers' Compensation / Employer's Liability Insurance.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "9.2  Limitation of Liability")
    add_body(doc, (
        "Except for fraud, wilful misconduct, or death or personal injury caused by the "
        "Contractor's negligence, the Contractor's total aggregate liability under this Contract "
        "shall not exceed the total Contract Price paid by the Client. Neither Party shall be "
        "liable for indirect, consequential, or loss-of-profit damages, except as required by "
        "mandatory consumer protection law."
    ))

    add_subheading(doc, "9.3  Indemnification")
    add_body(doc, (
        "Each Party shall indemnify the other against third-party claims arising directly from "
        "their own negligence or wilful misconduct in the performance of their respective "
        "obligations under this Contract."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # 10. RIGHT OF WITHDRAWAL AND CANCELLATION
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10.  RIGHT OF WITHDRAWAL AND CANCELLATION")

    add_subheading(doc, "10.1  14-Day Right of Withdrawal (Statutory — Off-Premises Contracts)")
    add_notice_box(doc, (
        "STATUTORY RIGHT: Under Cyprus Law 133(I)/2013 implementing EU Consumer Rights Directive "
        "2011/83/EU, if this Contract is concluded at the Client's premises or any location other "
        "than the Contractor's business premises (an off-premises contract), the Client has a "
        "statutory right to withdraw from this Contract without giving any reason within FOURTEEN "
        "(14) CALENDAR DAYS from the date of signing.\n\n"
        "To exercise this right, the Client must notify the Contractor before the 14-day period "
        "expires (email to office@lighthief.com is acceptable).\n\n"
        "Effect of withdrawal: The Contractor shall refund all payments received within fourteen "
        "(14) days of receiving the withdrawal notice, subject to §10.1.1 below."
    ), bg="E8F4FD", border_hex=NAVY_HEX)

    add_subheading(doc, "10.1.1  Early Start and Partial Withdrawal Refund")
    add_body(doc, (
        "If the Client expressly requests that works commence before the 14-day withdrawal period "
        "expires, and the Client subsequently withdraws, the Client shall pay a proportionate "
        "amount for services actually performed up to the point of withdrawal."
    ))

    add_subheading(doc, "10.2  Cancellation After the 14-Day Period")
    add_body(doc, (
        "If the Client cancels this Contract after the 14-day withdrawal period has expired, "
        "the following cancellation charges apply:"
    ))
    add_two_col_table(doc, [
        ("After 14-day period, before installation commences",
         "10% of Contract Price (administration and design costs)"),
        ("After installation commences, before completion",
         "30% of Contract Price (engineering, material procurement costs)"),
        ("After installation completed",
         "No refund — full Contract Price is due"),
    ], header_row=["Stage of Cancellation", "Cancellation Charge"],
       col_widths=[3.3, 3.2])

    add_subheading(doc, "10.3  Contractor Termination Rights")
    add_body(doc, "The Contractor may terminate this Contract by written notice for:")
    for item in [
        "Client non-payment following thirty (30) days' written notice;",
        "Material breach of Client obligations that remains unremedied after fourteen (14) days' "
        "written notice;",
        "Discovery of unsafe site conditions that cannot be resolved within a reasonable timeframe.",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # 11. DISPUTE RESOLUTION
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11.  DISPUTE RESOLUTION")

    add_subheading(doc, "11.1  Good-Faith Negotiation")
    add_body(doc, (
        "The Parties shall first attempt to resolve any dispute through good-faith negotiation "
        "within thirty (30) days of one Party giving written notice of the dispute to the other."
    ))

    add_subheading(doc, "11.2  Mediation and Consumer Protection")
    add_body(doc, (
        "If negotiation fails, the Parties shall attempt mediation before commencing legal "
        "proceedings. The Client may also contact the following authorities for assistance:"
    ))
    for item in [
        "Cyprus Consumer Protection Service: tel. 1429 (free consumer helpline);",
        "Cyprus Energy Regulatory Authority (CERA): +357 22 666 363;",
        "EU Online Dispute Resolution: https://ec.europa.eu/consumers/odr/.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "11.3  Legal Proceedings")
    add_body(doc, (
        "Any dispute not resolved by negotiation or mediation shall be subject to the exclusive "
        "jurisdiction of the courts of the Republic of Cyprus."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # 12. DATA PROTECTION
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "12.  DATA PROTECTION")
    add_body(doc, (
        "The Contractor shall comply with the General Data Protection Regulation (EU) 2016/679 "
        "(GDPR) and applicable Cyprus data protection laws regarding any personal data collected "
        "in connection with this Contract. Personal data will be processed solely for the purpose "
        "of performing this Contract and will not be disclosed to third parties without the "
        "Client's consent, except as required by law. The Client may request access to, "
        "correction, or deletion of their personal data at any time by contacting "
        "office@lighthief.com."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # 13. GENERAL PROVISIONS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "13.  GENERAL PROVISIONS")
    clauses = [
        ("13.1  Entire Agreement",
         "This Contract constitutes the entire agreement between the Parties regarding its "
         "subject matter and supersedes all prior negotiations, representations, or agreements."),
        ("13.2  Amendments",
         "Any amendment must be in writing and signed by authorised representatives of both Parties."),
        ("13.3  Governing Law",
         "This Contract is governed by the laws of the Republic of Cyprus."),
        ("13.4  Language",
         "This Contract is executed in English. In the event of any inconsistency with a "
         "translated version, the English version shall prevail."),
        ("13.5  Severability",
         "If any provision is found invalid or unenforceable, the remaining provisions continue "
         "in full force and effect."),
        ("13.6  Force Majeure",
         "Neither Party shall be liable for delays or failure to perform obligations due to "
         "circumstances beyond their reasonable control, including acts of God, war, terrorism, "
         "pandemic, governmental action, or grid authority decisions, provided the affected Party "
         "gives prompt written notice and uses reasonable efforts to mitigate the impact."),
        ("13.7  Assignment",
         "Neither Party may assign this Contract without the prior written consent of the other."),
        ("13.8  Notices",
         "All formal notices under this Contract shall be in writing and sent by email (with "
         "delivery confirmation) or by registered post to the addresses in the Parties section. "
         "Email notices are deemed received on the next Business Day."),
    ]
    for title, text in clauses:
        add_subheading(doc, title)
        add_body(doc, text)

    # ════════════════════════════════════════════════════════════════════════
    # 14. SPECIAL CONDITIONS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "14.  SPECIAL CONDITIONS")

    add_subheading(doc, "14.1  Roof Tile Guarantee")
    add_body(doc, (
        "The Contractor guarantees to replace any roof tiles broken or directly damaged during "
        "panel installation with tiles of the same style and colour. Where exactly matching tiles "
        "are no longer commercially available, the Contractor will replace with tiles of the "
        "closest available colour and grade. This guarantee covers only tiles directly damaged "
        "during installation; it does not extend to pre-existing roof membrane, waterproofing, "
        "or structural conditions."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # PAYMENT BANK DETAILS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "PAYMENT DETAILS")
    add_two_col_table(doc, [
        ("Account Name",   "LIGHTHIEF CYPRUS LTD"),
        ("Account Number", "357044102353"),
        ("IBAN",           "CY86002001950000357044102353"),
        ("SWIFT/BIC",      "BCYPCY2N"),
        ("Bank",           "Bank of Cyprus"),
    ], col_widths=[2.2, 4.3])

    add_notice_box(doc, (
        "⚠  FRAUD PREVENTION NOTICE: Lighthief Cyprus Ltd will NEVER change bank account details "
        "by email or SMS. Before making any payment, verify the account details above are correct "
        "by calling +357 77 77 00 50. Do not transfer funds to any account other than the one "
        "above without first confirming directly with a Lighthief representative by phone."
    ), bg="FFF3CD", border_hex="FF8C00")

    # ════════════════════════════════════════════════════════════════════════
    # SIGNATURES
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "SIGNATURES")
    add_body(doc, (
        "By signing below, both Parties confirm they have read, understood, and agreed to all "
        "terms of this Contract, including the Client's right of withdrawal in §10.1."
    ))
    doc.add_paragraph()

    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (party, name, cap) in enumerate([
        ("FOR THE CLIENT", CLIENT["name"], ""),
        ("FOR THE CONTRACTOR — LIGHTHIEF CYPRUS LTD",
         CONTRACTOR["rep"], CONTRACTOR["cap"]),
    ]):
        cell = sig_tbl.cell(0, i)
        p_title = cell.paragraphs[0]
        r = p_title.add_run(party)
        r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY
        para_spacing(p_title, before=0, after=80)
        rows_sig = [("Name:", name), ("Capacity:", cap)] if cap else [("Name:", name)]
        rows_sig += [
            ("Signature:", "___________________________"),
            ("Date:",       "___________________________"),
        ]
        for label, val in rows_sig:
            p = cell.add_paragraph()
            r_lbl = p.add_run(label + "  ")
            r_lbl.bold = True; r_lbl.font.size = Pt(9.5); r_lbl.font.color.rgb = GREY
            r_val = p.add_run(val)
            r_val.font.size = Pt(9.5); r_val.font.color.rgb = BLACK
            para_spacing(p, before=0, after=50)
        cell.width = Inches(3.25)

    doc.add_paragraph()
    p_note = doc.add_paragraph()
    r_note = p_note.add_run("* Each Party shall retain one signed original of this Contract.")
    r_note.font.size = Pt(8.5); r_note.font.color.rgb = GREY; r_note.italic = True

    # ════════════════════════════════════════════════════════════════════════
    # CLIENT NOTICE (page 2)
    # ════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "IMPORTANT NOTICE TO CLIENT")
    add_notice_box(doc, (
        "This Contract includes important terms regarding your rights and obligations.\n\n"
        "You are advised to:\n"
        "  •  Read all terms carefully before signing.\n"
        "  •  Keep a copy for your records.\n"
        "  •  Note your 14-day right of withdrawal (§10.1).\n"
        "  •  Understand the payment schedule and milestones.\n"
        "  •  Note the warranty periods for panels (15yr / 30yr), inverter (10yr), "
        "and workmanship (5yr).\n"
        "  •  Note the annual EAC Net Metering administrative fee (€47.23/kWp + VAT).\n"
        "  •  Be aware of your consumer rights under Cyprus law.\n\n"
        "For more information on your consumer rights, contact:\n"
        "  •  Cyprus Consumer Protection Service: 1429\n"
        "  •  Cyprus Energy Regulatory Authority (CERA): +357 22 666 363\n"
        "  •  EU Online Dispute Resolution: https://ec.europa.eu/consumers/odr/"
    ), bg="E8F4FD", border_hex=NAVY_HEX)

    # ════════════════════════════════════════════════════════════════════════
    # ANNEX A — WITHDRAWAL FORM
    # ════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "ANNEX A — STANDARD WITHDRAWAL FORM")
    add_body(doc, "(Complete and return this form only if you wish to withdraw from the Contract)")
    add_notice_box(doc, (
        "TO: Lighthief Cyprus Ltd, 28 October Ave 249, Lophitis Business Center 1, Office 201, "
        "3035 Limassol, Cyprus\n"
        "Email: office@lighthief.com\n\n"
        "I / We (*) hereby give notice that I / We (*) withdraw from my / our (*) contract for "
        "the supply and installation of the following:\n\n"
        "System address: _______________________________________________\n\n"
        "Contract date: " + CONTRACT["date"] + "\n\n"
        "Client name(s): " + CLIENT["name"] + "\n\n"
        "Client address: _______________________________________________\n\n"
        "Client signature(s) (only if this form is notified on paper):\n\n"
        "_______________________________________________\n\n"
        "Date: _______________________________________________\n\n"
        "(*) Delete as appropriate."
    ), bg="FAFBFC", border_hex=NAVY_HEX)

    out = os.path.normpath(OUTPUT)
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
