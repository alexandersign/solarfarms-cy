"""Generate General Introducer Agreement — IPN International Property Network (Cyprus) Limited."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand palette ─────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1A, 0x36, 0x5D)
GOLD  = RGBColor(0xC9, 0xA4, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x40, 0x40, 0x40)

COMPANY = {
    "name":    "LIGHTHIEF CYPRUS LTD",
    "reg":     "HE 477423",
    "tin":     "60187188Q",
    "address": "28 October Ave 249, Lophitis Business Center I, Office 201, 3035 Limassol, Cyprus",
    "email":   "office@lighthief.com",
    "phone":   "+357 77 77 00 50",
    "rep":     "Alexander Papacosta",
    "title":   "Cyprus Director",
    "website": "solarfarms.cy",
}

PARTNER = {
    "name":    "IPN INTERNATIONAL PROPERTY NETWORK (CYPRUS) LIMITED",
    "reg":     "HE 351537",
    "address": "Cyprus",
    "rep":     "Vasos Kolonas",
    "title":   "Authorised Representative",
}

REF_NO = "LCY-INT-IPN-2026-001"


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def setup(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(3.0)
        s.right_margin  = Cm(3.0)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = GOLD
        run.font.bold = True
        run.font.size = Pt(14)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = GOLD
        run.font.bold = True
        run.font.size = Pt(12)
    return p


def body(doc, text, bold=False, italic=False, colour=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = colour if colour else BLACK
    return p


def spacer(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def two_col_table(doc, rows, header=None):
    total = len(rows) + (1 if header else 0)
    t = doc.add_table(rows=total, cols=2)
    t.style = "Table Grid"
    idx = 0
    if header:
        for ci, val in enumerate(header):
            cell = t.cell(0, ci)
            cell.text = val
            set_cell_bg(cell, "1A365D")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(10)
        idx = 1
    for label, value in rows:
        t.cell(idx, 0).text = label
        t.cell(idx, 1).text = value
        for para in t.cell(idx, 0).paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        for para in t.cell(idx, 1).paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = BLACK
        idx += 1
    return t


def clause(doc, number, title, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.0)
    run_num = p.add_run(f"{number}.  ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_num.font.color.rgb = BLACK
    run_title = p.add_run(f"{title}.  ")
    run_title.bold = True
    run_title.underline = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = BLACK
    run_text = p.add_run(text)
    run_text.font.size = Pt(11)
    run_text.font.color.rgb = BLACK


def step(doc, title, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r_title = p.add_run(f"{title}:  ")
    r_title.bold = True
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = NAVY
    p.add_run(text).font.size = Pt(11)


def divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("\u2500" * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY


# ── Document build ────────────────────────────────────────────────────────────

def build(output_path):
    doc = Document()
    setup(doc)

    # ── Cover header ─────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("GENERAL INTRODUCER AGREEMENT")
    r.font.name = "Calibri"
    r.font.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle_p.add_run(
        "Renewable Energy & Clean Technology Solutions — Worldwide\n"
        "Lighthief Cyprus Ltd  \u00b7  General Introducer Programme"
    )
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    r2.font.color.rgb = GREY
    r2.italic = True

    spacer(doc)

    two_col_table(doc, [
        ("Agreement Reference",  REF_NO),
        ("Date of Agreement",    "_______ / _______ / _______"),
        ("Classification",       "CONFIDENTIAL \u2014 Between Parties Only"),
    ])

    spacer(doc)
    divider(doc)
    spacer(doc)

    # ── 1. Parties ────────────────────────────────────────────────────────────
    h1(doc, "1.  PARTIES")

    h2(doc, "1.1  Principal")
    two_col_table(doc, [
        ("Company",          COMPANY["name"]),
        ("Reg. No.",         COMPANY["reg"]),
        ("TIN",              COMPANY["tin"]),
        ("Address",          COMPANY["address"]),
        ("Email",            COMPANY["email"]),
        ("Tel",              COMPANY["phone"]),
        ("Website",          COMPANY["website"]),
        ("Representative",   f"{COMPANY['rep']}, {COMPANY['title']}"),
    ])

    spacer(doc)
    h2(doc, "1.2  Introducer")
    two_col_table(doc, [
        ("Company Name",              PARTNER["name"]),
        ("Registration No.",          PARTNER["reg"]),
        ("Registered Address",        "______________________________________"),
        ("Email",                     "______________________________________"),
        ("Tel",                       "______________________________________"),
        ("Authorised Representative", PARTNER["rep"]),
        ("Professional Capacity",     PARTNER["title"]),
    ])

    spacer(doc)
    body(
        doc,
        "The Principal and the Introducer are referred to individually as a \u201cParty\u201d and "
        "collectively as the \u201cParties\u201d.",
        italic=True, colour=GREY,
    )
    spacer(doc)

    # ── 2. Recitals ───────────────────────────────────────────────────────────
    h1(doc, "2.  RECITALS")
    body(
        doc,
        "(A)  The Principal is a Cyprus-registered company providing renewable energy products "
        "and services including, without limitation, solar photovoltaic (PV) systems, battery "
        "energy storage systems (BESS), operations and maintenance (O&M) services, and related "
        "clean-technology solutions, to clients globally.",
    )
    spacer(doc)
    body(
        doc,
        "(B)  The Introducer has established relationships and networks which may be valuable in "
        "identifying prospective clients for the Principal's services.",
    )
    spacer(doc)
    body(
        doc,
        "(C)  The Parties wish to formalise the terms on which the Introducer will introduce "
        "prospective clients to the Principal and on which the Principal will pay an introducer "
        "commission in the event a contract is concluded with an Introduced Client.",
    )
    spacer(doc)

    # ── 3. Definitions ────────────────────────────────────────────────────────
    h1(doc, "3.  DEFINITIONS")
    definitions = [
        ("\u201cAgreement\u201d",
         "this General Introducer Agreement, including all schedules."),
        ("\u201cServices\u201d",
         "any and all products and services provided by the Principal from time to time, "
         "including solar PV systems, BESS installations, O&M contracts, energy management "
         "solutions, and any other clean-technology or renewable energy offering, in any "
         "jurisdiction worldwide."),
        ("\u201cClient Contract\u201d",
         "any contract, order, or agreement executed between the Principal and an Introduced "
         "Client for the supply of Services, howsoever structured."),
        ("\u201cIntroduction\u201d",
         "a written notification by the Introducer to the Principal, sent by email to "
         f"{COMPANY['email']}, identifying a prospective client by full name and a "
         "partial telephone number (country code plus at least the first digits of the "
         "subscriber number, sufficient to verify identity), in accordance with Clause 4."),
        ("\u201cIntroduced Client\u201d",
         "a legal entity or natural person identified in a valid, confirmed Introduction "
         "who subsequently signs a Client Contract with the Principal within the "
         "Protection Period."),
        ("\u201cRegistration Confirmation\u201d",
         "written confirmation issued by the Principal within five (5) Working Days of receipt "
         "of an Introduction, confirming either: (i) the prospective client is a new contact "
         "not already in the Principal\u2019s active pipeline (\u201cNew Registration\u201d); or "
         "(ii) the prospective client is an existing contact already engaged by the Principal "
         "(\u201cExisting Contact\u201d)."),
        ("\u201cProtection Period\u201d",
         "the twelve (12) calendar month period following a New Registration Confirmation, "
         "during which the Introducer\u2019s Commission entitlement is preserved."),
        ("\u201cContract Value\u201d",
         "the net value of the Client Contract excluding VAT, as stated in the signed contract."),
        ("\u201cCommission\u201d",
         "the fee payable by the Principal to the Introducer on the terms of this Agreement."),
        ("\u201cWorking Days\u201d",
         "Monday to Friday, excluding public holidays in Cyprus."),
    ]
    for term, definition in definitions:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        run_term = p.add_run(f"{term}  ")
        run_term.bold = True
        run_term.font.size = Pt(11)
        p.add_run(f"means {definition}").font.size = Pt(11)

    spacer(doc)

    # ── 4. Appointment ────────────────────────────────────────────────────────
    h1(doc, "4.  APPOINTMENT")
    body(
        doc,
        "4.1  The Principal hereby appoints the Introducer as a non-exclusive general introducer "
        "for the purpose of identifying and introducing prospective clients for the Principal\u2019s "
        "Services, in any country or territory worldwide.",
    )
    spacer(doc)
    body(
        doc,
        "4.2  The Introducer has no authority to negotiate pricing, make representations, accept "
        "orders, or enter into any contractual commitment on behalf of the Principal. All client "
        "negotiations and agreements are conducted solely by the Principal.",
    )
    spacer(doc)
    body(
        doc,
        "4.3  The Introducer acts as an independent contractor. Nothing in this Agreement creates "
        "an employment, partnership, joint venture, or agency relationship beyond the limited "
        "introducer scope described herein.",
    )
    spacer(doc)

    # ── 5. Introduction Process ───────────────────────────────────────────────
    h1(doc, "5.  INTRODUCTION & REGISTRATION PROCESS")
    body(
        doc,
        "All Introductions must follow the steps below. An Introduction that does not comply "
        "with these steps shall not give rise to any Commission entitlement.",
    )
    spacer(doc)

    step(doc,
         "Step 1 \u2014 Submit Introduction by Email",
         f"The Introducer sends an email to {COMPANY['email']} with the subject line "
         "\u201cNew Introducer Registration\u201d. The email must include: "
         "(a) the full legal name of the prospective client or company; "
         "(b) a partial telephone number (country code + at minimum the first 4\u20136 digits "
         "of the subscriber number); and "
         "(c) any other contextual information available (country, sector, project type). "
         "No telephone number format or email address alone is sufficient \u2014 both name "
         "and partial number are required to register.")
    spacer(doc)

    step(doc,
         "Step 2 \u2014 Registration Confirmation (5 Working Days)",
         "The Principal shall issue a Registration Confirmation within five (5) Working Days "
         "of receiving the Introduction email. The confirmation will state either: "
         "(i) \u201cNew Registration Confirmed\u201d \u2014 the prospective client is not in the "
         "Principal\u2019s active pipeline and the Protection Period commences from the "
         "confirmation date; or "
         "(ii) \u201cExisting Contact \u2014 No Commission\u201d \u2014 the prospective client is "
         "already engaged by the Principal, and no Commission entitlement arises. "
         "If the Principal fails to issue a Registration Confirmation within five (5) Working "
         "Days, the Introduction shall be deemed confirmed as a New Registration.")
    spacer(doc)

    step(doc,
         "Step 3 \u2014 No Direct Contact Before Confirmation",
         "The Principal shall not initiate or establish any direct contact with the introduced "
         "prospective client until a Registration Confirmation has been issued under Step 2. "
         "If the Principal independently contacts the same prospective client within five (5) "
         "Working Days of receipt of the Introduction email and before issuing the Registration "
         "Confirmation, the Introduction shall be deemed a New Registration.")
    spacer(doc)

    step(doc,
         "Step 4 \u2014 Facilitated Introduction",
         "Following a New Registration Confirmation, the Introducer may, at the Principal\u2019s "
         "request, facilitate an introductory meeting, call, or written introduction between "
         "the prospective client and the Principal. The Introducer is not required to manage "
         "or participate in subsequent commercial negotiations.")
    spacer(doc)

    step(doc,
         "Step 5 \u2014 Commission Trigger",
         "Commission becomes payable upon execution of a Client Contract between the Principal "
         "and the Introduced Client within the Protection Period, in accordance with Clause 6.")
    spacer(doc)

    # ── 6. Commission ─────────────────────────────────────────────────────────
    h1(doc, "6.  COMMISSION")
    two_col_table(doc, [
        ("Commission Rate",    "3.0% (three per cent) of Contract Value"),
        ("Basis",              "Net Contract Value excluding VAT, as stated in the signed Client Contract"),
        ("Scope",              "Applies to all Services, all jurisdictions, with no geographic or "
                               "product-category limitation"),
        ("Multiple Contracts", "Where the Introduced Client signs Client Contracts for multiple "
                               "projects, a separate Commission is calculated on each individual "
                               "Contract Value"),
        ("Variations",         "If the final Contract Value increases via a signed variation order, "
                               "Commission on the additional amount follows the same 3.0% rate"),
        ("Large Contracts",    "For any single Client Contract Value exceeding \u20ac10,000,000, "
                               "the Commission on the portion above \u20ac10,000,000 shall be agreed "
                               "in writing between the Parties before execution of that contract"),
    ], header=["Commission Term", "Detail"])

    spacer(doc)

    # ── 7. Payment Terms ──────────────────────────────────────────────────────
    h1(doc, "7.  PAYMENT TERMS")
    two_col_table(doc, [
        ("Instalment 1 \u2014 30% of Commission",
         "Invoiced and payable within 30 days of the Principal receiving the Introduced "
         "Client\u2019s advance payment under the Client Contract"),
        ("Instalment 2 \u2014 70% of Commission",
         "Invoiced and payable within 30 days of the Principal receiving the Introduced "
         "Client\u2019s final payment under the Client Contract"),
        ("Payment Method",
         "Bank transfer to the Introducer\u2019s designated account (details in Schedule B)"),
        ("Currency",
         "Euro (\u20ac). Where the Client Contract is denominated in another currency, "
         "Commission is converted at the ECB mid-rate on the date the triggering payment "
         "is received by the Principal."),
        ("Late Payment Interest",
         "2% per month on overdue amounts, accruing from the due date"),
    ], header=["Milestone", "Detail"])

    spacer(doc)
    body(
        doc,
        "No Commission is payable if the Introduced Client cancels the Client Contract before "
        "final payment, or if any partial payment already received from the Introduced Client "
        "is refunded by the Principal. Instalment 1 already paid shall be credited against "
        "future Commissions from the same Introducer.",
        italic=True, colour=GREY,
    )
    spacer(doc)

    # ── 8. Protection Period ──────────────────────────────────────────────────
    h1(doc, "8.  PROTECTION PERIOD & COMMISSION ENTITLEMENT")
    clause(doc, "8.1", "Duration",
           "A New Registration Confirmation triggers a Protection Period of twelve (12) calendar "
           "months from the confirmation date. If the Introduced Client signs a Client Contract "
           "within this period, Commission is payable at the rate in Clause 6.")
    clause(doc, "8.2", "Extension",
           "The Protection Period may be extended by mutual written agreement if the Introduced "
           "Client is actively engaged in contract negotiation at the expiry date.")
    clause(doc, "8.3", "Existing Contacts",
           "No Commission is payable where the Registration Confirmation states \u201cExisting "
           "Contact\u201d, provided such confirmation is issued within the five (5) Working Day "
           "window under Clause 5 Step 2.")
    clause(doc, "8.4", "Failure to Confirm",
           "If the Principal fails to issue any Registration Confirmation within five (5) Working "
           "Days of receiving a valid Introduction email, the introduced prospective client shall "
           "be deemed a New Registration and the Protection Period shall commence from the sixth "
           "Working Day after receipt of the Introduction.")
    clause(doc, "8.5", "No Contact Before Confirmation",
           "Any direct commercial contact established by the Principal with the introduced "
           "prospective client before a Registration Confirmation is issued shall not extinguish "
           "the Introducer\u2019s Commission entitlement. Such contact shall be treated as if "
           "a New Registration had been confirmed.")
    spacer(doc)

    # ── 9. Obligations of the Introducer ─────────────────────────────────────
    h1(doc, "9.  OBLIGATIONS OF THE INTRODUCER")
    clause(doc, "9.1", "Accuracy",
           "The Introducer shall provide accurate and complete information in each Introduction "
           "email, including the required name and partial telephone number.")
    clause(doc, "9.2", "No Misrepresentation",
           "The Introducer shall not make any false, misleading, or unauthorised representations "
           "about the Principal, its products, pricing, or capabilities.")
    clause(doc, "9.3", "No Pricing Disclosure",
           "The Introducer shall not disclose or discuss specific pricing, margins, or commercial "
           "terms with any prospective client unless expressly authorised in writing by the Principal.")
    clause(doc, "9.4", "Compliance",
           "The Introducer shall at all times comply with all applicable laws and professional "
           "regulations, including anti-bribery, anti-corruption, and data protection legislation "
           "in all relevant jurisdictions.")
    clause(doc, "9.5", "Conflict of Interest",
           "The Introducer shall immediately disclose to the Principal any actual or potential "
           "conflict of interest in relation to any introduced prospective client.")
    spacer(doc)

    # ── 10. Obligations of the Principal ─────────────────────────────────────
    h1(doc, "10.  OBLIGATIONS OF THE PRINCIPAL")
    clause(doc, "10.1", "Timely Response",
           f"The Principal shall respond to each Introduction email sent to {COMPANY['email']} "
           "within five (5) Working Days with a Registration Confirmation.")
    clause(doc, "10.2", "No Direct Contact Before Confirmation",
           "The Principal shall not initiate direct commercial contact with a prospective client "
           "identified in an Introduction email until a Registration Confirmation has been issued.")
    clause(doc, "10.3", "Commission Payment",
           "The Principal shall pay Commission in accordance with Clause 7 and shall provide "
           "the Introducer with written notification of each triggered Commission event within "
           "ten (10) Working Days of the triggering payment being received.")
    clause(doc, "10.4", "Records",
           "The Principal shall maintain accurate records of all Introductions, Registration "
           "Confirmations, Client Contracts, and Commission payments for the duration of this "
           "Agreement plus three (3) years, and shall make such records available to the "
           "Introducer on reasonable written request.")
    spacer(doc)

    # ── 11. Confidentiality ───────────────────────────────────────────────────
    h1(doc, "11.  CONFIDENTIALITY")
    body(
        doc,
        "11.1  Each Party shall treat as strictly confidential all information received from "
        "the other Party in connection with this Agreement, including commission rates, client "
        "identities, Introduction details, technical specifications, and business strategies "
        "(\u201cConfidential Information\u201d).",
    )
    spacer(doc)
    body(
        doc,
        "11.2  Neither Party shall disclose Confidential Information to any third party without "
        "the prior written consent of the disclosing Party, except to professional advisers "
        "bound by equivalent confidentiality obligations.",
    )
    spacer(doc)
    body(
        doc,
        "11.3  This obligation of confidentiality survives termination of this Agreement for "
        "a period of three (3) years.",
    )
    spacer(doc)

    # ── 12. Term & Termination ────────────────────────────────────────────────
    h1(doc, "12.  TERM & TERMINATION")
    clause(doc, "12.1", "Term",
           "This Agreement commences on the date of last signature and continues for a period "
           "of twenty-four (24) months, unless terminated earlier in accordance with this clause.")
    clause(doc, "12.2", "Renewal",
           "The Agreement shall automatically renew for successive twelve (12) month periods "
           "unless either Party gives thirty (30) days written notice of non-renewal before "
           "the end of the then-current term.")
    clause(doc, "12.3", "Termination for Convenience",
           "Either Party may terminate this Agreement for convenience on thirty (30) days "
           "written notice. Commission entitlement for Introductions already confirmed as New "
           "Registrations before the termination notice date is preserved for the remainder "
           "of the relevant Protection Period.")
    clause(doc, "12.4", "Termination for Cause",
           "The Principal may terminate immediately on written notice if the Introducer: "
           "(a) breaches Clause 9 or 11; (b) is found to have engaged in fraudulent, corrupt, "
           "or illegal conduct; or (c) becomes insolvent. No Commission shall be payable "
           "following termination for cause in respect of any Introduction made after the "
           "breach.")
    spacer(doc)

    # ── 13. Taxes & Withholding ───────────────────────────────────────────────
    h1(doc, "13.  TAXES & WITHHOLDING")
    body(
        doc,
        "13.1  Each Party is solely responsible for its own tax obligations arising from "
        "this Agreement. Commission amounts stated herein are exclusive of VAT or any "
        "equivalent indirect tax.",
    )
    spacer(doc)
    body(
        doc,
        "13.2  Where applicable under Cyprus law, the Principal may withhold tax at source "
        "from Commission payments and shall issue the requisite withholding tax certificates. "
        "The Introducer is responsible for declaring Commission income to the relevant tax "
        "authorities in its jurisdiction.",
    )
    spacer(doc)

    # ── 14. Limitation of Liability ───────────────────────────────────────────
    h1(doc, "14.  LIMITATION OF LIABILITY")
    body(
        doc,
        "14.1  The Principal\u2019s total liability to the Introducer under this Agreement "
        "shall not exceed the total Commission actually paid in the twelve (12) months "
        "preceding the relevant claim.",
    )
    spacer(doc)
    body(
        doc,
        "14.2  Neither Party shall be liable for indirect, consequential, or special loss, "
        "including lost profits or business opportunities, howsoever arising.",
    )
    spacer(doc)
    body(
        doc,
        "14.3  The Principal makes no warranty as to the commercial outcome of any proposal "
        "submitted to an Introduced Client, nor as to the acceptance of any quotation.",
    )
    spacer(doc)

    # ── 15. Governing Law ─────────────────────────────────────────────────────
    h1(doc, "15.  GOVERNING LAW & DISPUTE RESOLUTION")
    body(
        doc,
        "15.1  This Agreement is governed by, and construed in accordance with, the laws "
        "of the Republic of Cyprus.",
    )
    spacer(doc)
    body(
        doc,
        "15.2  The Parties shall endeavour to resolve any dispute amicably within twenty-one "
        "(21) days of written notice. Failing amicable resolution, disputes shall be subject "
        "to the exclusive jurisdiction of the courts of Limassol, Cyprus.",
    )
    spacer(doc)
    body(
        doc,
        "15.3  Any notice under this Agreement shall be in writing and delivered by email "
        "(with read receipt) or registered post to the addresses stated in Clause 1. "
        "Notices are effective on receipt.",
    )
    spacer(doc)

    # ── 16. General ───────────────────────────────────────────────────────────
    h1(doc, "16.  GENERAL")
    clause(doc, "16.1", "Entire Agreement",
           "This Agreement, together with its Schedules, constitutes the entire agreement "
           "between the Parties relating to its subject matter and supersedes all prior "
           "discussions, representations, and arrangements.")
    clause(doc, "16.2", "Amendments",
           "No amendment to this Agreement is valid unless made in writing and signed by "
           "authorised representatives of both Parties.")
    clause(doc, "16.3", "Waiver",
           "Failure to exercise, or delay in exercising, any right under this Agreement "
           "shall not constitute a waiver of that right.")
    clause(doc, "16.4", "Severability",
           "If any provision of this Agreement is found invalid or unenforceable, the "
           "remaining provisions shall continue in full force and effect.")
    clause(doc, "16.5", "Assignment",
           "The Introducer may not assign or transfer any rights or obligations under "
           "this Agreement without the prior written consent of the Principal.")
    clause(doc, "16.6", "Programme Amendments",
           "The Principal reserves the right to amend Commission rates or programme terms "
           "with thirty (30) days written notice. Such amendments shall not affect Commission "
           "entitlements arising from Introductions confirmed as New Registrations before "
           "the notice date.")
    spacer(doc)

    # ── Signatures ────────────────────────────────────────────────────────────
    divider(doc)
    spacer(doc)
    h1(doc, "EXECUTION")
    body(
        doc,
        "By signing below, the Parties confirm they have read, understood, and agreed to "
        "all terms and conditions of this General Introducer Agreement.",
    )
    spacer(doc)

    sig = doc.add_table(rows=6, cols=2)
    sig.style = "Table Grid"

    sig_headers = [
        f"FOR THE INTRODUCER\n{PARTNER['name']}",
        f"FOR THE PRINCIPAL\n{COMPANY['name']}",
    ]
    for ci, val in enumerate(sig_headers):
        cell = sig.cell(0, ci)
        cell.text = val
        set_cell_bg(cell, "1A365D")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)

    left_rows = [
        f"Name:  {PARTNER['rep']}",
        f"Capacity:  {PARTNER['title']}",
        "Signature:  _________________________",
        "Date:  ______________________________",
        "Stamp (if applicable):  ______________",
    ]
    right_rows = [
        f"Name:  {COMPANY['rep']}",
        f"Capacity:  {COMPANY['title']}",
        "Signature:  _________________________",
        "Date:  ______________________________",
        "Company Seal",
    ]

    for i, (l, r) in enumerate(zip(left_rows, right_rows), start=1):
        sig.cell(i, 0).text = l
        sig.cell(i, 1).text = r
        for ci in range(2):
            for para in sig.cell(i, ci).paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = BLACK

    spacer(doc)
    body(
        doc,
        "* Each Party should initial every page and retain one signed original. "
        "An electronically signed copy (PDF) shall have equal legal validity.",
        italic=True, colour=GREY,
    )

    spacer(doc)
    divider(doc)
    spacer(doc)

    # ── Schedule A — Introduction Log ─────────────────────────────────────────
    h1(doc, "SCHEDULE A \u2014 INTRODUCTION LOG & REGISTRATION RECORD")
    body(
        doc,
        "To be completed by the Principal upon issuing each Registration Confirmation. "
        "A copy shall be emailed to the Introducer within five (5) Working Days of receipt "
        "of the Introduction email.",
    )
    spacer(doc)
    two_col_table(doc, [
        ("Agreement Reference",          REF_NO),
        ("Introduction Email Date",      "_______ / _______ / _______"),
        ("Prospective Client Name",      "______________________________________"),
        ("Partial Telephone Number",     "______________________________________"),
        ("Country / Territory",          "______________________________________"),
        ("Sector / Project Type",        "______________________________________"),
        ("Registration Status",          "\u25a1  New Registration \u25a1  Existing Contact \u2014 No Commission"),
        ("Protection Period Expiry",     "_______ / _______ / _______  (12 months from confirmation)"),
        ("Commission Rate",              "3.0% of Contract Value"),
        ("Confirmation Issued By",       f"{COMPANY['rep']}, {COMPANY['title']}"),
        ("Confirmation Date",            "_______ / _______ / _______"),
        ("Confirmation Sent to",         "______________________________________  (Introducer email)"),
    ], header=["Field", "Details"])

    spacer(doc)
    body(
        doc,
        "Note: A separate Schedule A shall be completed for each Introduction submitted "
        "under this Agreement.",
        italic=True, colour=GREY,
    )

    spacer(doc)
    divider(doc)
    spacer(doc)

    # ── Schedule B — Bank Details ──────────────────────────────────────────────
    h1(doc, "SCHEDULE B \u2014 INTRODUCER BANK DETAILS")
    body(
        doc,
        "Bank account details for Commission payments. "
        "The Introducer is responsible for notifying the Principal of any changes.",
    )
    spacer(doc)
    two_col_table(doc, [
        ("Account Holder Name",     "______________________________________"),
        ("Bank Name",               "______________________________________"),
        ("IBAN",                    "______________________________________"),
        ("SWIFT / BIC",             "______________________________________"),
        ("Bank Address",            "______________________________________"),
        ("Currency",                "Euro (\u20ac)"),
        ("VAT Registration No.",    "______________________________________"),
    ])

    spacer(doc)
    divider(doc)
    spacer(doc)

    # ── Footer ────────────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = footer_p.add_run(
        f"\u00a9 2026 {COMPANY['name']}  |  Reg. HE 477423  |  {COMPANY['website']}  |  {COMPANY['email']}\n"
        "This agreement is confidential. Introducer Commission is subject to execution of this Agreement "
        "and the registration process in Clause 5. All commission rates are as stated herein."
    )
    r_f.font.size = Pt(8)
    r_f.font.color.rgb = GREY
    r_f.italic = True

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    build("referral-agreement-ipn-jun2026.docx")
