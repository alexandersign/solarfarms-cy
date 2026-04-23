"""Generate off-grid PV installation contract for Lighthief Cyprus Ltd."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x1A, 0x6B, 0x3C)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x40, 0x40, 0x40)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

COMPANY = {
    "name":    "LIGHTHIEF CYPRUS LTD",
    "reg":     "HE 477423",
    "tin":     "60187188Q",
    "address": "28 October Ave 249, Lophitis Business Center 1, Office 201, 3035 Limassol, Cyprus",
    "email":   "office@lighthief.com",
    "phone":   "+357 77 77 00 50",
    "rep":     "Alexander Papacosta",
}


# ── helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def setup(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    for s in doc.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(3.0)
        s.right_margin  = Cm(3.0)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = GREEN
        run.font.bold = True
        run.font.size = Pt(13)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = GREEN
        run.font.size = Pt(11)


def body(doc, text, bold=False, italic=False, colour=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.bold = bold
    run.italic = italic
    if colour:
        run.font.color.rgb = colour
    return p


def spacer(doc):
    doc.add_paragraph()


def two_col_table(doc, rows, header=None):
    """Simple two-column table with optional header row."""
    total = len(rows) + (1 if header else 0)
    t = doc.add_table(rows=total, cols=2)
    t.style = "Table Grid"
    idx = 0
    if header:
        for ci, val in enumerate(header):
            cell = t.cell(0, ci)
            cell.text = val
            set_cell_bg(cell, "1A6B3C")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)
        idx = 1
    for label, value in rows:
        t.cell(idx, 0).text = label
        t.cell(idx, 1).text = value
        for para in t.cell(idx, 0).paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        for para in t.cell(idx, 1).paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
        idx += 1
    return t


def numbered_clause(doc, number, title, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(f"{number}. {title}. ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)


# ── document build ───────────────────────────────────────────────────────────

def build(output_path):
    doc = Document()
    setup(doc)

    # ── Title block ──────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("OFF-GRID PHOTOVOLTAIC & ENERGY STORAGE\nINSTALLATION CONTRACT")
    r.font.name = "Calibri"
    r.font.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = GREEN

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Lighthief Cyprus Ltd  ·  Supply, Installation & Commissioning")
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    r2.font.color.rgb = GREY

    doc.add_paragraph()

    # ── 1. Parties ───────────────────────────────────────────────────────────
    h1(doc, "1. PARTIES")

    h2(doc, "Contractor")
    two_col_table(doc, [
        ("Company",             COMPANY["name"]),
        ("Reg. No.",            COMPANY["reg"]),
        ("TIN",                 COMPANY["tin"]),
        ("Address",             COMPANY["address"]),
        ("Email",               COMPANY["email"]),
        ("Tel",                 COMPANY["phone"]),
        ("Representative",      COMPANY["rep"]),
    ])

    spacer(doc)
    h2(doc, "Client")
    two_col_table(doc, [
        ("Full Name / Company",  "______________________________________"),
        ("ID / Reg. No.",        "______________________________________"),
        ("Address",              "______________________________________"),
        ("Email",                "______________________________________"),
        ("Tel",                  "______________________________________"),
        ("Representative",       "______________________________________"),
    ])

    spacer(doc)

    # ── 2. Scope ─────────────────────────────────────────────────────────────
    h1(doc, "2. SCOPE OF WORK")
    body(doc,
         "The Contractor shall supply, deliver, install, and commission a fully off-grid photovoltaic "
         "and battery energy storage system (the \"System\") at the premises specified below. "
         "The scope includes all labour, materials, cabling, protection devices, and programming "
         "required for a fully operational, isolated power system.")

    spacer(doc)
    two_col_table(doc, [
        ("Installation Address", "______________________________________"),
        ("PV Capacity (kWp)",    "______________________________________"),
        ("Battery Capacity",     "______________________________________"),
        ("Inverter / Charger",   "______________________________________"),
        ("Generator Interface",  "□ ATS included  □ Not included"),
        ("Monitoring",           "□ Included  □ Not included"),
    ], header=["Parameter", "Specification"])

    spacer(doc)

    # ── 3. Contract Price ────────────────────────────────────────────────────
    h1(doc, "3. CONTRACT PRICE")
    two_col_table(doc, [
        ("Total Contract Price (VAT incl. 19%)",  "€ ________________"),
        ("Net Amount (excl. VAT)",                "€ ________________"),
        ("VAT (19%)",                             "€ ________________"),
    ])

    spacer(doc)

    # ── 4. Payment Terms ─────────────────────────────────────────────────────
    h1(doc, "4. PAYMENT TERMS")
    body(doc, "Payments shall be made by bank transfer to the Contractor's account. "
              "No works shall commence or materials ordered until the deposit (Instalment 1) is received.")

    spacer(doc)
    two_col_table(doc, [
        ("#1 — Deposit (50%)",              "Due upon contract signing"),
        ("#2 — Pre-delivery (30%)",         "Due 3 days before equipment delivery"),
        ("#3 — Final (20%)",                "Due on commissioning date"),
    ], header=["Instalment", "Trigger"])

    spacer(doc)
    body(doc, "All amounts are VAT-inclusive. Late payments accrue interest at 2% per month on the outstanding balance.",
         italic=True, colour=GREY)

    spacer(doc)

    # ── 5. Timeline ──────────────────────────────────────────────────────────
    h1(doc, "5. TIMELINE")
    two_col_table(doc, [
        ("Contract Signing / PID",            "_______ / _______ / _______"),
        ("Estimated Installation Start",      "Within ____ weeks of deposit receipt"),
        ("Estimated Completion",              "Within ____ working days of start"),
    ])

    spacer(doc)
    body(doc,
         "Timelines are estimates. The Contractor shall not be liable for delays caused by weather, "
         "supply chain disruptions, permitting, or other circumstances beyond its reasonable control.",
         italic=True, colour=GREY)

    spacer(doc)

    # ── 6. Client Obligations ────────────────────────────────────────────────
    h1(doc, "6. CLIENT OBLIGATIONS")
    obligations = [
        ("6.1", "Site Access",
         "Provide unobstructed access to the installation site during working hours for the duration of the works."),
        ("6.2", "Existing Infrastructure",
         "Disclose all known underground cables, pipes, and structural constraints before works begin. "
         "The Contractor is not liable for damage to undisclosed infrastructure."),
        ("6.3", "No Modification",
         "Not to modify, reprogram, or tamper with any component of the System without prior written consent "
         "from the Contractor. Unauthorized modifications void all warranties."),
        ("6.4", "Off-Grid Confirmation",
         "Acknowledge that the System is designed exclusively for isolated (off-grid) operation and must "
         "not be connected to any public distribution network without a separate written agreement and "
         "all required regulatory approvals."),
        ("6.5", "Load Management",
         "Manage electrical loads within the system's designed capacity. Overloading the System releases "
         "the Contractor from liability for equipment damage or loss of supply."),
        ("6.6", "Generator Fuel",
         "Ensure adequate fuel supply for any backup generator. Generator maintenance remains the Client's responsibility."),
    ]
    for num, title, text in obligations:
        numbered_clause(doc, num, title, text)

    spacer(doc)

    # ── 7. Warranties ────────────────────────────────────────────────────────
    h1(doc, "7. WARRANTIES")
    two_col_table(doc, [
        ("Workmanship",      "1 year from commissioning date"),
        ("Inverter / Charger", "As per OEM — typically 2–5 years"),
        ("Battery Modules",  "As per OEM — typically 5–10 years"),
        ("PV Modules",       "Product: 10 years / Performance: 25 years (OEM)"),
    ], header=["Item", "Coverage"])

    spacer(doc)
    body(doc,
         "Warranties are void if: (a) the system is modified without the Contractor's consent; "
         "(b) damage is caused by Client negligence, misuse, or acts of God; "
         "(c) payment obligations are not fulfilled.",
         italic=True, colour=GREY)

    spacer(doc)

    # ── 8. Limitation of Liability ───────────────────────────────────────────
    h1(doc, "8. LIMITATION OF LIABILITY")
    body(doc,
         "The Contractor's total aggregate liability under or in connection with this Contract shall not "
         "exceed the total Contract Price actually paid. The Contractor shall not be liable for any "
         "indirect, consequential, or special loss, including loss of income, business interruption, "
         "or data loss, howsoever arising.")

    spacer(doc)

    # ── 9. Cancellation ──────────────────────────────────────────────────────
    h1(doc, "9. CANCELLATION")
    body(doc,
         "If the Client cancels after signing: all payments received are non-refundable. "
         "Equipment already ordered or works already completed shall be invoiced and are payable in full. "
         "The Contractor may cancel this Contract in writing if the Client is in material breach "
         "and fails to remedy within 14 days of written notice.")

    spacer(doc)

    # ── 10. Governing Law ────────────────────────────────────────────────────
    h1(doc, "10. GOVERNING LAW")
    body(doc,
         "This Contract is governed by the laws of the Republic of Cyprus. "
         "The courts of Cyprus shall have exclusive jurisdiction over any dispute arising hereunder. "
         "Any notice under this Contract shall be in writing and sent by email or registered post to "
         "the addresses above.")

    spacer(doc)
    doc.add_paragraph("─" * 72)

    # ── Signatures ───────────────────────────────────────────────────────────
    h1(doc, "SIGNATURES")
    body(doc,
         "By signing below, both parties confirm they have read, understood, and agreed to all "
         "terms of this Contract.")

    spacer(doc)

    sig = doc.add_table(rows=5, cols=2)
    sig.style = "Table Grid"

    headers = ["FOR THE CLIENT", "FOR THE CONTRACTOR — LIGHTHIEF CYPRUS LTD"]
    for ci, val in enumerate(headers):
        cell = sig.cell(0, ci)
        cell.text = val
        set_cell_bg(cell, "1A6B3C")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)

    left_rows  = ["Name: ______________________", "Capacity: ___________________", "Signature: __________________", "Date: _______________________"]
    right_rows = [f"Name: {COMPANY['rep']}", "Capacity: Director", "Signature: __________________", "Date: _______________________"]

    for i, (l, r) in enumerate(zip(left_rows, right_rows), start=1):
        sig.cell(i, 0).text = l
        sig.cell(i, 1).text = r
        for ci in range(2):
            for para in sig.cell(i, ci).paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    spacer(doc)
    body(doc,
         "* Both parties should initial each page and retain one signed original each.",
         italic=True, colour=GREY)

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    build("offgrid-installation-contract-template.docx")
