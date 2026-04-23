"""Generate PV O&M contract template for Lighthief Cyprus Ltd.

Utility-scale ground-mount / large rooftop PV plants (500 kWp+).
Three service tiers: A — Basic, B — Standard, C — Premium.

Usage:
    python generate_om_contract.py
Outputs:
    pv-om-contract-template.docx  (in the current working directory)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Lighthief brand palette ───────────────────────────────────────────────────
NAVY  = RGBColor(0x1A, 0x36, 0x5D)   # primary — table headers, bars
GOLD  = RGBColor(0xC9, 0xA4, 0x32)   # headings, reference numbers
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x40, 0x40, 0x40)   # sub-body, footnotes

COMPANY = {
    "name":    "LIGHTHIEF CYPRUS LTD",
    "reg":     "HE 477423",
    "tin":     "60187188Q",
    "address": "28 October Ave 249, Lophitis Business Center 1, Office 201, 3035 Limassol, Cyprus",
    "email":   "office@lighthief.com",
    "phone":   "+357 77 77 00 50",
    "rep":     "Alexander Papacosta",
    "title":   "Director",
}


# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def setup(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    for s in doc.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(3.0)
        s.right_margin  = Cm(3.0)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = GOLD
        run.font.bold = True
        run.font.size = Pt(13)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(11)


def body(doc, text, bold=False, italic=False, colour=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.bold = bold
    run.italic = italic
    if colour:
        run.font.color.rgb = colour
    return p


def spacer(doc):
    doc.add_paragraph()


def navy_table(doc, rows, header=None, col_widths=None):
    """General-purpose table with NAVY header row and alternating body rows."""
    total = len(rows) + (1 if header else 0)
    ncols = len(header) if header else len(rows[0])
    t = doc.add_table(rows=total, cols=ncols)
    t.style = "Table Grid"
    idx = 0
    if header:
        for ci, val in enumerate(header):
            cell = t.cell(0, ci)
            cell.text = val
            set_cell_bg(cell, "1A365D")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9)
        idx = 1
    for row_data in rows:
        for ci, val in enumerate(row_data):
            cell = t.cell(idx, ci)
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if ci == 0:
                        run.font.bold = True
        idx += 1
    return t


def two_col_table(doc, rows, header=None):
    """Convenience wrapper for a 2-column label/value table."""
    formatted = [(label, value) for label, value in rows]
    return navy_table(doc, formatted, header=header)


def numbered_clause(doc, number, title, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(f"{number}. {title}. ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)


def divider(doc):
    doc.add_paragraph("─" * 72)


# ── document build ────────────────────────────────────────────────────────────

def build(output_path):
    doc = Document()
    setup(doc)

    # ── Title block ──────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PHOTOVOLTAIC OPERATIONS & MAINTENANCE AGREEMENT")
    r.font.name  = "Calibri"
    r.font.bold  = True
    r.font.size  = Pt(15)
    r.font.color.rgb = GOLD

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Lighthief Cyprus Ltd  ·  Monitoring, Maintenance & Performance Management")
    r2.font.name  = "Calibri"
    r2.font.size  = Pt(10)
    r2.font.color.rgb = GREY

    spacer(doc)
    two_col_table(doc, [
        ("Contract Ref.",    "LC-OM-_____ / _____"),
        ("Commencement Date", "_______ / _______ / _______"),
        ("Contract Year 1 End", "_______ / _______ / _______"),
        ("Selected Tier",    "☐ Tier A — Basic   ☐ Tier B — Standard   ☐ Tier C — Premium"),
    ])
    spacer(doc)

    # ── 1. Parties ───────────────────────────────────────────────────────────
    h1(doc, "1. PARTIES")

    h2(doc, "Service Provider")
    two_col_table(doc, [
        ("Company",         COMPANY["name"]),
        ("Reg. No.",        COMPANY["reg"]),
        ("TIN",             COMPANY["tin"]),
        ("Address",         COMPANY["address"]),
        ("Email",           COMPANY["email"]),
        ("Tel",             COMPANY["phone"]),
        ("Representative",  COMPANY["rep"]),
    ])

    spacer(doc)
    h2(doc, "Plant Owner / Operator (Client)")
    two_col_table(doc, [
        ("Full Name / Company",    "______________________________________"),
        ("Registration No.",       "______________________________________"),
        ("Address",                "______________________________________"),
        ("Email",                  "______________________________________"),
        ("Tel",                    "______________________________________"),
        ("Authorized Representative", "______________________________________"),
    ])
    spacer(doc)

    # ── 2. Definitions ───────────────────────────────────────────────────────
    h1(doc, "2. DEFINITIONS")
    body(doc,
         "In this Agreement the following terms shall have the meanings set out below:")
    spacer(doc)
    navy_table(doc, [
        ("Agreement",             "This Photovoltaic Operations & Maintenance Agreement, including all schedules."),
        ("Plant",                 "The photovoltaic generating installation described in Section 3, including all associated electrical equipment, cabling, protection devices, monitoring hardware, and inverters."),
        ("PV System",             "The photovoltaic modules, inverter(s), mounting structure, DC/AC cabling, and protection equipment forming part of the Plant."),
        ("Contract Year",         "Each successive 12-month period commencing on the Commencement Date or any anniversary thereof."),
        ("Commencement Date",     "The date entered in the header table above, from which Service obligations begin."),
        ("Monitoring Period",     "Each calendar month during the term of this Agreement."),
        ("Scheduled Maintenance Visit (SMV)", "A planned on-site maintenance visit carried out in accordance with the schedule for the selected Service Tier."),
        ("Corrective Maintenance", "Unplanned remedial work carried out to restore the Plant to normal operating condition following a Fault."),
        ("Fault",                 "Any condition causing the Plant to operate outside normal parameters, including inverter trip, string failure, communication loss, or protection relay activation."),
        ("Response Time",         "The elapsed time from Service Provider's acknowledgment of a reported Fault to on-site attendance by a qualified technician."),
        ("Availability",          "The ratio of hours in a Contract Year during which all inverters are operational and producing power (subject to irradiance > 50 W/m²) to the total number of daylight hours in the same period, expressed as a percentage."),
        ("Performance Ratio (PR)", "The ratio of actual energy yield to reference yield (based on in-plane irradiance), as defined in IEC 61724-1, expressed as a percentage and corrected for temperature."),
        ("P50 Baseline",          "The median expected annual energy yield for the Plant as established by the most recent PVGIS simulation or independent energy yield assessment on file."),
        ("Point of Common Coupling (PCC)", "The electrical connection point between the Plant and the EAC distribution network."),
        ("Force Majeure",         "Any event beyond a party's reasonable control, including acts of God, lightning, flood, earthquake, grid network failure, vandalism, theft, or war."),
        ("EAC",                   "The Electricity Authority of Cyprus, acting as Distribution System Operator (DSO) for the purposes of this Agreement."),
        ("CERA",                  "The Cyprus Energy Regulatory Authority."),
        ("HICP",                  "The Harmonised Index of Consumer Prices for Cyprus published by Eurostat."),
        ("Parts",                 "Replacement components, consumables, or materials required for Corrective Maintenance. Parts are excluded from all Service Tier fees unless expressly stated."),
    ], header=["Term", "Meaning"])
    spacer(doc)

    # ── 3. Plant Description ─────────────────────────────────────────────────
    h1(doc, "3. PLANT DESCRIPTION")
    body(doc,
         "The Services under this Agreement apply exclusively to the Plant described in the table below. "
         "Any material change to the Plant — including capacity increase, inverter replacement, or "
         "structural modification — shall be notified in writing. The Service Provider may revise fees "
         "accordingly with 30 days' notice.")
    spacer(doc)
    two_col_table(doc, [
        ("Plant Site Address",           "______________________________________"),
        ("Installed PV Capacity (kWp)",  "______________________________________"),
        ("Number of PV Modules",         "______________________________________"),
        ("Inverter Make / Model",        "______________________________________"),
        ("Number of Inverters",          "______________________________________"),
        ("Monitoring Platform / SCADA",  "______________________________________"),
        ("Grid Connection Voltage",      "□ LV (0.4 kV)   □ MV (11 kV / 22 kV)"),
        ("EAC Connection Approval Ref.", "______________________________________"),
        ("Commissioning Date",           "______________________________________"),
        ("Prior O&M Provider (if any)",  "______________________________________"),
    ], header=["Parameter", "Details"])
    spacer(doc)

    # ── 4. Service Tiers ─────────────────────────────────────────────────────
    h1(doc, "4. SERVICE TIERS")
    body(doc,
         "The Client selects one Service Tier at signing (indicated in the header table). "
         "The services included in each tier are cumulative — Tier B includes all Tier A services; "
         "Tier C includes all Tier B services. The selected tier is recorded in the header and confirmed "
         "by both parties' signatures. Upgrading between tiers requires a written amendment signed by both parties.")
    spacer(doc)

    # Tier A
    h2(doc, "Tier A — Basic (Monitoring & Reporting)")
    navy_table(doc, [
        ("Remote Monitoring",    "Portal access to inverter / SCADA data with automated performance tracking throughout the Contract Year."),
        ("Fault Alerts",         "Automated email notification to Client within 4 hours of detection of any Fault visible via the monitoring platform."),
        ("Monthly Report",       "A report delivered within 10 business days of month-end covering: energy generated (kWh), estimated CO₂ saving, specific yield (kWh/kWp), and any faults logged."),
        ("Annual Summary",       "A year-end performance report comparing actual yield against P50 Baseline, with a deviation analysis and maintenance recommendations."),
        ("Site Visits",          "Not included in Tier A. Site visits may be requested at the applicable day-rate (see Section 5)."),
    ], header=["Service Element", "Description"])
    spacer(doc)

    # Tier B
    h2(doc, "Tier B — Standard (+ Preventive Maintenance)")
    navy_table(doc, [
        ("All Tier A Services",  "All services listed under Tier A are included."),
        ("Scheduled Maintenance Visits", "2 (two) SMVs per Contract Year, scheduled by mutual agreement and conducted during normal working hours (08:00–17:00, Mon–Fri, excluding public holidays)."),
        ("SMV Scope — each visit includes:",
         "• Visual inspection of all PV modules for soiling, shading, damage, and delamination\n"
         "• String IV curve and Voc testing to identify underperforming or failed strings\n"
         "• Thermal imaging of all string combiner boxes, DC isolators, and AC panels\n"
         "• Torque verification on all DC and AC electrical terminations\n"
         "• Inverter inspection: event log review, firmware version check, air filter cleaning\n"
         "• Earth continuity and insulation resistance (IR) testing on DC circuits\n"
         "• Visual inspection of mounting structure, cable trays, and conduits for mechanical integrity\n"
         "• Grounding and bonding continuity check"),
        ("SMV Report",           "A written maintenance report delivered within 10 business days of each SMV, including test results, photographs, and any corrective action recommendations."),
        ("Corrective Maintenance", "Not included in Tier B. Corrective maintenance following an SMV finding may be quoted separately or upgraded to Tier C."),
        ("Emergency Callouts",   "Not included. Available at the applicable emergency day-rate (see Section 5)."),
    ], header=["Service Element", "Description"])
    spacer(doc)

    # Tier C
    h2(doc, "Tier C — Premium (+ Corrective Maintenance & Performance Guarantee)")
    navy_table(doc, [
        ("All Tier B Services",  "All services listed under Tiers A and B are included."),
        ("Corrective Maintenance — Unlimited Callouts",
         "The Service Provider shall carry out all Corrective Maintenance required to restore the Plant to normal operation. "
         "Callouts are unlimited in number. Parts and consumables are excluded and invoiced separately at cost."),
        ("Response Time SLA",
         "P1 — Critical (total plant output loss): on-site within 4 hours of acknowledgment\n"
         "P2 — Major (>30% capacity loss): on-site within 24 hours of acknowledgment\n"
         "P3 — Minor (<30% capacity loss, no safety risk): on-site within 72 hours of acknowledgment\n"
         "Acknowledgment: within 2 hours of fault report during business hours; within 4 hours outside business hours."),
        ("Availability Guarantee",
         "The Service Provider guarantees an annual Availability of ≥95% per Contract Year, measured on an inverter-operating-hours basis "
         "and excluding Force Majeure events, grid outages at the PCC, and downtime attributable to Client obligations under Section 7."),
        ("Performance Ratio Guarantee",
         "The Service Provider guarantees an annual PR of ≥78%, irradiance-corrected using PVGIS TMY data for the plant location. "
         "Measured over the full Contract Year. Excludes periods of grid curtailment by EAC."),
        ("Liquidated Damages",
         "For each full percentage point by which Availability or PR falls below the guaranteed threshold, the Service Provider shall credit "
         "5% of the applicable annual fee against the next invoice. Total LD credit is capped at 20% of the annual fee per Contract Year. "
         "LD credit is the Client's sole remedy for Availability and PR shortfall."),
        ("Annual Technical Review",
         "An annual in-person technical review meeting to present performance data, maintenance findings, and the forward maintenance plan for the next Contract Year."),
    ], header=["Service Element", "Description"])
    spacer(doc)

    # ── 5. Fees & Payment ────────────────────────────────────────────────────
    h1(doc, "5. FEES & PAYMENT")
    body(doc,
         "Annual O&M fees are denominated in Euros per kWp of installed capacity per annum, as agreed "
         "and recorded in the table below at signing. Fees are VAT-exclusive; VAT at the prevailing Cyprus "
         "rate (currently 19%) is added to each invoice.")
    spacer(doc)
    two_col_table(doc, [
        ("Selected Tier",                    "☐ Tier A   ☐ Tier B   ☐ Tier C"),
        ("Annual Fee Rate (€/kWp/year)",     "€ _____________ / kWp / year"),
        ("Installed Capacity (kWp)",         "_____________ kWp"),
        ("Annual Fee — Year 1 (excl. VAT)",  "€ _____________"),
        ("VAT (19%)",                        "€ _____________"),
        ("Annual Fee — Year 1 (incl. VAT)",  "€ _____________"),
        ("Invoice Frequency",                "☐ Annually in advance   ☐ Quarterly in advance"),
        ("Payment Method",                   "Bank transfer to Service Provider's account (details on invoice)"),
    ], header=["Item", "Value"])
    spacer(doc)
    navy_table(doc, [
        ("Annual Escalation",
         "From Contract Year 2 onwards, fees escalate annually by the greater of: (a) 3%, or (b) the "
         "HICP (Cyprus) annual index as published by Eurostat for the preceding calendar year. "
         "The Service Provider shall notify the Client of the revised fee no later than 60 days before each Contract Year anniversary."),
        ("Emergency / Ad-Hoc Day Rate",
         "Works not covered by the selected tier (including parts-only Tier B callouts or additional SMVs) "
         "are charged at €_______ / technician-day plus parts at cost plus 15%."),
        ("Late Payment",
         "Invoices unpaid after 30 days from the due date accrue interest at 2% per month on the outstanding balance."),
        ("Parts",
         "Replacement parts and consumables are excluded from all tier fees and invoiced separately at cost plus 15%, "
         "regardless of selected tier. A written quotation is provided before any parts order exceeding €500."),
    ], header=["Fee Term", "Details"])
    spacer(doc)

    # ── 6. Service Schedule ──────────────────────────────────────────────────
    h1(doc, "6. SERVICE SCHEDULE")
    body(doc,
         "The Service Provider shall, no later than 30 days before the start of each Contract Year, "
         "provide the Client with a draft Service Schedule for the year, setting out proposed SMV dates "
         "(Tier B/C), reporting due dates, and the annual review date (Tier C). "
         "The Client shall confirm or propose alternative dates within 10 business days. "
         "If no response is received, the proposed dates shall be deemed accepted. "
         "SMVs shall be conducted during normal working hours unless otherwise agreed.")
    spacer(doc)
    body(doc,
         "The Client shall provide the Service Provider with a minimum of 48 hours' notice prior to any "
         "access restriction affecting a scheduled SMV. For Tier C emergency callouts, a minimum of "
         "24 hours' notice of access restriction is required; shorter notice may be given in genuine emergency situations.",
         italic=True, colour=GREY)
    spacer(doc)

    # ── 7. Client Obligations ────────────────────────────────────────────────
    h1(doc, "7. CLIENT OBLIGATIONS")
    obligations = [
        ("7.1", "Site Access",
         "Provide the Service Provider's personnel with safe, unobstructed access to all parts of the Plant — "
         "including rooftops, cable routes, inverter rooms, and switchgear — on reasonable notice (48 hours for "
         "scheduled visits; 24 hours for Tier C emergency callouts). Access shall be free of charge."),
        ("7.2", "Grid Connection & Meter",
         "Maintain the EAC grid connection, metering equipment, and account in good standing throughout the term. "
         "Notify the Service Provider within 24 hours of any EAC disconnection, tariff change, or meter fault."),
        ("7.3", "No Unauthorized Modification",
         "Not to modify, reprogram, bypass, or tamper with any component of the Plant — including inverter settings, "
         "protection relays, monitoring hardware, or CT clamps — without prior written consent from the Service Provider. "
         "Unauthorized modification voids all Service Tier guarantees and the Availability / PR commitments in Tier C."),
        ("7.4", "Regulatory Notifications",
         "Immediately notify the Service Provider in writing if contacted by EAC, CERA, the DSO, or any regulatory "
         "authority regarding the Plant, its grid connection, or its generating licence."),
        ("7.5", "Site Utilities",
         "Provide, at no charge to the Service Provider, access to a 230V/16A electrical supply and water connection "
         "at the Plant site for use during SMVs and corrective maintenance visits."),
        ("7.6", "Civil & Structural Integrity",
         "Maintain the civil and structural integrity of the plant boundary, access roads, mounting structures, "
         "cable trench covers, and any weather protection enclosures. The Service Provider is not liable for equipment "
         "damage resulting from structural failure caused by Client's neglect."),
        ("7.7", "Insurance",
         "Maintain in force throughout the term: (a) property all-risks insurance covering the Plant for its full "
         "replacement value; and (b) public liability insurance of no less than €2,000,000 per occurrence. "
         "Certificates shall be provided to the Service Provider upon written request."),
        ("7.8", "Data Sharing",
         "Provide the Service Provider with ongoing remote access to the monitoring platform and, on request, "
         "EAC billing data, irradiance sensor data, and any third-party SCADA credentials necessary to deliver "
         "the contracted Services."),
    ]
    for num, title, text in obligations:
        numbered_clause(doc, num, title, text)
    spacer(doc)

    # ── 8. Exclusions ────────────────────────────────────────────────────────
    h1(doc, "8. EXCLUSIONS")
    body(doc,
         "The following are expressly excluded from all Service Tiers and shall, where applicable, "
         "be quoted and charged separately:")
    spacer(doc)
    navy_table(doc, [
        ("PV Module Replacement",        "Replacement of physically damaged, degraded, or end-of-life PV modules. Covered by OEM product warranty where applicable — Service Provider can assist with claim management at day-rate."),
        ("Inverter Replacement",         "Full inverter unit replacement. Labour and parts charged separately. Service Provider can manage OEM warranty claims at day-rate."),
        ("Grid Infrastructure",          "Any infrastructure beyond the Plant's Point of Common Coupling (PCC), including EAC transformer, metering, and distribution network."),
        ("Civil & Structural Works",     "Repair or replacement of roofing, ground foundations, mounting structures, cable trenches, or any civil works."),
        ("Vandalism & Theft",            "Damage or loss caused by vandalism, theft, or malicious acts. Client's property insurance is the applicable remedy."),
        ("Force Majeure",                "Damage or performance loss caused by Force Majeure events as defined in Section 2."),
        ("EAC Grid Faults",              "Loss of generation or equipment damage caused by EAC network faults, voltage events, or grid curtailment orders originating at or beyond the PCC."),
        ("Third-Party Software Licences", "EMS software licences, SCADA platform subscription fees, or third-party data-management tools not supplied by the Service Provider."),
        ("CERA / EAC Regulatory Changes", "Any cost or work arising from changes in Net Billing tariff rates, generating licence conditions, or grid code requirements enacted after the Commencement Date."),
        ("Works Outside Normal Hours",   "SMVs or callouts outside 08:00–17:00 Mon–Fri (excluding public holidays) unless agreed in advance. Surcharge applies as per Section 5."),
    ], header=["Exclusion", "Notes"])
    spacer(doc)

    # ── 9. Term & Renewal ────────────────────────────────────────────────────
    h1(doc, "9. TERM & RENEWAL")
    body(doc,
         "This Agreement commences on the Commencement Date and continues for an initial term of one (1) "
         "Contract Year. It shall renew automatically for successive Contract Years unless either party "
         "gives written notice of termination no less than three (3) months before the end of the then-current "
         "Contract Year. The maximum aggregate term without a full re-signing is five (5) Contract Years.")
    spacer(doc)
    body(doc,
         "Either party may terminate this Agreement with immediate effect by written notice if the other party: "
         "(a) is in material breach and fails to remedy within 30 days of written notice; "
         "(b) becomes insolvent or enters into any form of liquidation or administration; or "
         "(c) ceases to hold any licence, permit, or authorisation material to the performance of its obligations.")
    spacer(doc)
    body(doc,
         "On termination, the Client shall pay all fees accrued and outstanding up to the effective date of "
         "termination. Fees paid in advance for periods beyond termination shall be refunded on a pro-rata basis "
         "within 30 days, except where termination is due to the Client's material breach.",
         italic=True, colour=GREY)
    spacer(doc)

    # ── 10. Limitation of Liability ──────────────────────────────────────────
    h1(doc, "10. LIMITATION OF LIABILITY")
    body(doc,
         "The Service Provider's total aggregate liability to the Client under or in connection with this "
         "Agreement — whether in contract, tort (including negligence), or otherwise — shall not exceed "
         "the total annual fees actually paid by the Client in the Contract Year in which the relevant "
         "claim arises.")
    spacer(doc)
    body(doc,
         "Neither party shall be liable to the other for any indirect, consequential, special, or punitive "
         "loss or damage — including loss of generation revenue, lost Net Billing credits, loss of profit, "
         "loss of use, or business interruption — howsoever arising.")
    spacer(doc)
    body(doc,
         "For Tier C, the Liquidated Damages mechanism in Section 4 (Tier C) is the Client's sole and "
         "exclusive remedy for any shortfall in Availability or Performance Ratio against the guaranteed thresholds. "
         "Nothing in this Agreement limits either party's liability for death or personal injury caused by negligence, "
         "or for fraudulent misrepresentation.")
    spacer(doc)

    # ── 11. Governing Law ────────────────────────────────────────────────────
    h1(doc, "11. GOVERNING LAW")
    body(doc,
         "This Agreement is governed by and shall be construed in accordance with the laws of the Republic "
         "of Cyprus. The courts of Cyprus shall have exclusive jurisdiction to settle any dispute or claim "
         "arising out of or in connection with this Agreement.")
    spacer(doc)
    body(doc,
         "All notices under this Agreement shall be in writing and delivered by email (with read receipt "
         "requested) or registered post to the addresses set out in Section 1. Notices take effect on the "
         "date of confirmed receipt.")
    spacer(doc)
    divider(doc)

    # ── Tier Selection & Signatures ──────────────────────────────────────────
    h1(doc, "TIER SELECTION & SIGNATURES")
    body(doc,
         "The Client confirms the selected Service Tier by initialling the corresponding box below. "
         "By signing, both parties confirm they have read and agree to all terms of this Agreement.")
    spacer(doc)

    # Tier selection table
    tier_sel = doc.add_table(rows=4, cols=3)
    tier_sel.style = "Table Grid"
    headers_ts = ["Service Tier", "Description", "Client Initials"]
    for ci, val in enumerate(headers_ts):
        cell = tier_sel.cell(0, ci)
        cell.text = val
        set_cell_bg(cell, "1A365D")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)

    tier_rows = [
        ("Tier A — Basic",    "Remote monitoring, fault alerts, monthly & annual reports."),
        ("Tier B — Standard", "All of Tier A + 2 scheduled preventive maintenance visits per year."),
        ("Tier C — Premium",  "All of Tier B + unlimited corrective maintenance, response time SLA, Availability ≥95% and PR ≥78% guarantees."),
    ]
    for i, (tier, desc) in enumerate(tier_rows, start=1):
        tier_sel.cell(i, 0).text = tier
        tier_sel.cell(i, 1).text = desc
        tier_sel.cell(i, 2).text = "______"
        for ci in range(3):
            for para in tier_sel.cell(i, ci).paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if ci == 0:
                        run.font.bold = True

    spacer(doc)

    # Signature table
    sig = doc.add_table(rows=5, cols=2)
    sig.style = "Table Grid"
    sig_headers = ["FOR THE CLIENT (PLANT OWNER / OPERATOR)", f"FOR THE SERVICE PROVIDER — {COMPANY['name']}"]
    for ci, val in enumerate(sig_headers):
        cell = sig.cell(0, ci)
        cell.text = val
        set_cell_bg(cell, "1A365D")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)

    left_rows  = [
        "Name: _______________________",
        "Capacity: ____________________",
        "Signature: ___________________",
        "Date: ________________________",
    ]
    right_rows = [
        f"Name: {COMPANY['rep']}",
        f"Capacity: {COMPANY['title']}",
        "Signature: ___________________",
        "Date: ________________________",
    ]
    for i, (l, r) in enumerate(zip(left_rows, right_rows), start=1):
        sig.cell(i, 0).text = l
        sig.cell(i, 1).text = r
        for ci in range(2):
            for para in sig.cell(i, ci).paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    spacer(doc)
    body(doc,
         "* Both parties should initial each page. Each party retains one signed original of this Agreement.",
         italic=True, colour=GREY)

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    build("pv-om-contract-template.docx")
