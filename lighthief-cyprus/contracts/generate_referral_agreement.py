"""Generate Referral Commission Agreement template for Lighthief Cyprus Ltd — BESS EPC Program."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand palette (Lighthief brand identity) ─────────────────────────────────
NAVY  = RGBColor(0x1A, 0x36, 0x5D)
GOLD  = RGBColor(0xC9, 0xA4, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x40, 0x40, 0x40)

COMPANY = {
    "name":    "LIGHTHIEF CYPRUS LTD",
    "reg":     "HE 477423",
    "tin":     "60187188Q",
    "address": "28 October Ave 249, Lophitis Business Center I, Office 201, 3035 Limassol, Cyprus",
    "email":   "office@lighthief.com",
    "phone":   "+357 77 77 00 50",
    "rep":     "Alexander Papacosta",
    "title":   "Managing Director",
}

REF_NO = "LCY-REF-2026-___"


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def setup(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(3.0)
        s.right_margin  = Cm(3.0)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = GOLD
        run.font.bold = True
        run.font.size = Pt(14)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = GOLD
        run.font.bold = True
        run.font.size = Pt(12)
    return p


def body(doc, text, bold=False, italic=False, colour=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = colour if colour else BLACK
    return p


def spacer(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def two_col_table(doc, rows, header=None, col_widths=None):
    total = len(rows) + (1 if header else 0)
    t = doc.add_table(rows=total, cols=2)
    t.style = "Table Grid"
    idx = 0
    if header:
        for ci, val in enumerate(header):
            cell = t.cell(0, ci)
            cell.text = val
            set_cell_bg(cell, "1A365D")  # NAVY
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(10)
        idx = 1
    for label, value in rows:
        t.cell(idx, 0).text = label
        t.cell(idx, 1).text = value
        for para in t.cell(idx, 0).paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        for para in t.cell(idx, 1).paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = BLACK
        idx += 1
    return t


def four_col_table(doc, rows, header):
    total = len(rows) + 1
    t = doc.add_table(rows=total, cols=4)
    t.style = "Table Grid"
    for ci, val in enumerate(header):
        cell = t.cell(0, ci)
        cell.text = val
        set_cell_bg(cell, "1A365D")  # NAVY
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
    for ri, row_vals in enumerate(rows, start=1):
        for ci, val in enumerate(row_vals):
            t.cell(ri, ci).text = val
            for para in t.cell(ri, ci).paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if ci == 0:
                        run.font.bold = True
    return t


def clause(doc, number, title, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.0)
    run_num = p.add_run(f"{number}.  ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_num.font.color.rgb = BLACK
    run_title = p.add_run(f"{title}.  ")
    run_title.bold = True
    run_title.underline = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = BLACK
    run_text = p.add_run(text)
    run_text.font.size = Pt(11)
    run_text.font.color.rgb = BLACK


def divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY


# ── Document build ────────────────────────────────────────────────────────────

def build(output_path):
    doc = Document()
    setup(doc)

    # ── Cover header ─────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("REFERRAL COMMISSION AGREEMENT")
    r.font.name = "Calibri"
    r.font.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle_p.add_run(
        "Battery Energy Storage System (BESS) EPC Turnkey Solutions — Cyprus\n"
        "Lighthief Cyprus Ltd  ·  Referral Partner Program"
    )
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    r2.font.color.rgb = GREY
    r2.italic = True

    spacer(doc)

    # Reference / meta table
    two_col_table(doc, [
        ("Agreement Reference",  REF_NO),
        ("Date of Agreement",    "_______ / _______ / _______"),
        ("Classification",       "CONFIDENTIAL — Between Parties Only"),
        ("Program Reference",    "LCY-REF-2026-001"),
    ])

    spacer(doc)
    divider(doc)
    spacer(doc)

    # ── 1. Parties ────────────────────────────────────────────────────────────
    h1(doc, "1.  PARTIES")

    h2(doc, "1.1  Principal")
    two_col_table(doc, [
        ("Company",         COMPANY["name"]),
        ("Reg. No.",        COMPANY["reg"]),
        ("TIN",             COMPANY["tin"]),
        ("Address",         COMPANY["address"]),
        ("Email",           COMPANY["email"]),
        ("Tel",             COMPANY["phone"]),
        ("Representative",  f"{COMPANY['rep']}, {COMPANY['title']}"),
    ])

    spacer(doc)
    h2(doc, "1.2  Referral Partner")
    two_col_table(doc, [
        ("Full Name / Company Name",  "______________________________________"),
        ("Registration / ID No.",     "______________________________________"),
        ("VAT / TIN No.",             "______________________________________"),
        ("Registered Address",        "______________________________________"),
        ("Email",                     "______________________________________"),
        ("Tel",                       "______________________________________"),
        ("Authorised Representative", "______________________________________"),
        ("Professional Capacity",     "______________________________________"),
    ])

    spacer(doc)
    body(
        doc,
        "The Principal and the Referral Partner are referred to individually as a \"Party\" and "
        "collectively as the \"Parties\".",
        italic=True, colour=GREY,
    )
    spacer(doc)

    # ── 2. Definitions ────────────────────────────────────────────────────────
    h1(doc, "2.  DEFINITIONS")
    definitions = [
        (u"\u201cAgreement\u201d",
         "this Referral Commission Agreement, including all schedules and annexes."),
        (u"\u201cBESS\u201d",
         "Battery Energy Storage System \u2014 a containerised, grid-connected energy storage "
         "installation forming part of a utility-scale renewable energy project."),
        (u"\u201cEPC Contract\u201d",
         "the Engineering, Procurement, and Construction (turnkey) contract signed between "
         "the Principal and a Referred Client for the supply and installation of a BESS system."),
        (u"\u201cReferred Client\u201d",
         "a legal entity or natural person introduced by the Referral Partner to the Principal "
         "pursuant to this Agreement, who subsequently signs an EPC Contract with the Principal."),
        (u"\u201cReferral\u201d",
         "a formal written notification by the Referral Partner to the Principal identifying a "
         "prospective client, their contact details, and the relevant PV park(s)."),
        (u"\u201cCommission\u201d",
         "the fee payable by the Principal to the Referral Partner upon the terms of this Agreement."),
        (u"\u201cContract Value\u201d",
         "the net value of the EPC Contract excluding VAT, as stated in the signed EPC Contract."),
        (u"\u201cExclusivity Period\u201d",
         "the twelve (12) month period following acceptance of a valid Referral during which "
         "the Commission entitlement is preserved."),
        (u"\u201cPipeline\u201d",
         "the Principal's internal register of active prospects already engaged prior to the Referral."),
    ]
    for term, definition in definitions:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        run_term = p.add_run(f"{term}  ")
        run_term.bold = True
        run_term.font.size = Pt(11)
        p.add_run(f"means {definition}").font.size = Pt(11)

    spacer(doc)

    # ── 3. Appointment ────────────────────────────────────────────────────────
    h1(doc, "3.  APPOINTMENT")
    body(
        doc,
        "3.1  The Principal hereby appoints the Referral Partner as a non-exclusive referral "
        "agent for the purpose of identifying and introducing potential clients for the "
        "Principal's BESS EPC turnkey services in Cyprus.",
    )
    spacer(doc)
    body(
        doc,
        "3.2  The Referral Partner has no authority to negotiate pricing, make representations, "
        "accept orders, or enter into any contractual commitment on behalf of the Principal. "
        "All client negotiations and agreements are conducted solely by the Principal.",
    )
    spacer(doc)
    body(
        doc,
        "3.3  The Referral Partner acts as an independent contractor. Nothing in this Agreement "
        "creates an employment, partnership, joint venture, or agency relationship beyond "
        "the limited referral scope described herein.",
    )
    spacer(doc)

    # ── 4. Referral Process ───────────────────────────────────────────────────
    h1(doc, "4.  REFERRAL PROCESS")
    steps = [
        ("Step 1 — Notify",
         "The Referral Partner submits a written Referral to the Principal (by email to "
         f"{COMPANY['email']}) providing: (a) client / company name; (b) contact person "
         "and role; (c) PV park location(s) and licensed capacity (MW); and "
         "(d) estimated project timeline."),
        ("Step 2 — Validation",
         "The Principal shall confirm in writing within three (3) business days whether the "
         "Referral is valid (i.e. the prospective client is not already in the Principal's "
         "active Pipeline). A valid Referral triggers the Exclusivity Period."),
        ("Step 3 — Introduction",
         "The Referral Partner facilitates an introductory meeting, call, or written "
         "introduction between the prospective client and the Principal."),
        ("Step 4 — Sales Process",
         "The Principal manages all subsequent sales activity, technical proposals, site "
         "assessments, pricing, and contract negotiations independently."),
        ("Step 5 — Commission",
         "Upon execution of an EPC Contract with the Referred Client, the Commission "
         "becomes payable in accordance with Schedule A hereto."),
    ]
    for step_title, step_text in steps:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r_title = p.add_run(f"{step_title}:  ")
        r_title.bold = True
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = NAVY
        p.add_run(step_text).font.size = Pt(11)

    spacer(doc)

    # ── 5. Commission ─────────────────────────────────────────────────────────
    h1(doc, "5.  COMMISSION STRUCTURE")
    body(
        doc,
        "The Commission rate applicable to this Agreement is confirmed in the table below and "
        "in Schedule A. The applicable Tier is determined at the time of Referral validation "
        "and recorded in Schedule A.",
    )
    spacer(doc)

    four_col_table(doc, [
        ("Tier A — Introduction",
         "Referral Partner provides client name, contact, and facilitates an introductory meeting",
         "2.0% of Contract Value",
         "Client signs EPC Contract"),
        ("Tier B — Qualified Lead",
         "Referred Client has confirmed intent to proceed (licensed PV park, budget confirmed)",
         "3.0% of Contract Value",
         "Client signs EPC Contract"),
        ("Tier C — Multi-Project",
         "Referred Client has two (2) or more PV parks to equip",
         "3.5% of Contract Value",
         "Per EPC Contract signed"),
    ], header=["Tier", "Qualification Criteria", "Commission Rate", "Payment Trigger"])

    spacer(doc)
    h2(doc, "5.1  Commission Cap")
    two_col_table(doc, [
        ("Up to €3,000,000 Contract Value",
         "Standard rate applies — no cap"),
        ("€3,000,001 — €10,000,000",
         "Standard rate on first €3,000,000; 1.5% on the balance above €3,000,000"),
        ("Above €10,000,000",
         "By separate written agreement between the Parties"),
    ], header=["Contract Value Band", "Commission Calculation"])

    spacer(doc)
    h2(doc, "5.2  Commission Basis")
    clause(doc, "5.2.1", "Net Value",
           "Commission is calculated on the net Contract Value, excluding VAT, as stated "
           "in the signed EPC Contract.")
    clause(doc, "5.2.2", "Multiple Parks",
           "Where the Referred Client signs EPC Contracts for multiple parks, a separate "
           "Commission is calculated on each individual Contract Value.")
    clause(doc, "5.2.3", "Variations",
           "If the final EPC Contract Value increases via a signed variation order, "
           "the Commission on the additional amount follows the same Tier rate.")
    spacer(doc)

    # ── 6. Payment Terms ──────────────────────────────────────────────────────
    h1(doc, "6.  PAYMENT TERMS")
    two_col_table(doc, [
        ("Instalment 1 — 30% of Commission",
         "Invoiced and payable within 30 days of the Principal receiving the Referred Client's "
         "advance payment under the EPC Contract"),
        ("Instalment 2 — 70% of Commission",
         "Invoiced and payable within 30 days of the Principal receiving the Referred Client's "
         "final payment under the EPC Contract"),
        ("Payment Method",
         "Bank transfer to the Referral Partner's designated account (details in Schedule B)"),
        ("Currency",
         "Euro (€)"),
        ("Late Payment Interest",
         "2% per month on overdue amounts, accruing from the due date"),
    ], header=["Milestone", "Detail"])

    spacer(doc)
    body(
        doc,
        "No Commission is payable if the Referred Client cancels the EPC Contract before "
        "final payment, or if any partial payment already made by the Referred Client is "
        "refunded by the Principal.",
        italic=True, colour=GREY,
    )
    spacer(doc)

    # ── 7. Exclusivity & Protection ───────────────────────────────────────────
    h1(doc, "7.  EXCLUSIVITY PERIOD & REFERRAL PROTECTION")
    clause(doc, "7.1", "Duration",
           "Once a Referral is validated in writing by the Principal, the Referral Partner "
           "is entitled to Commission if the Referred Client signs an EPC Contract within "
           "twelve (12) months of the validation date (the \"Exclusivity Period\").")
    clause(doc, "7.2", "Extension",
           "The Exclusivity Period may be extended by mutual written agreement if the "
           "Referred Client is actively engaged in contract negotiation at the expiry date.")
    clause(doc, "7.3", "Pre-existing Pipeline",
           "No Commission is payable if the prospective client was already in the Principal's "
           "active Pipeline at the time of the Referral, as notified to the Referral Partner "
           "within the three (3) business day validation window.")
    clause(doc, "7.4", "Direct Contact",
           "If a Referred Client contacts the Principal directly (independently of the "
           "Referral Partner) before the Referral is submitted, no Commission entitlement arises.")
    spacer(doc)

    # ── 8. Obligations of the Referral Partner ────────────────────────────────
    h1(doc, "8.  OBLIGATIONS OF THE REFERRAL PARTNER")
    clause(doc, "8.1", "Accuracy",
           "The Referral Partner shall provide accurate, complete, and up-to-date information "
           "about each Referred Client and their project.")
    clause(doc, "8.2", "No Misrepresentation",
           "The Referral Partner shall not make any false, misleading, or unauthorised "
           "representations about the Principal, its products, pricing, or capabilities.")
    clause(doc, "8.3", "No Pricing Disclosure",
           "The Referral Partner shall not disclose or discuss specific pricing, margins, "
           "or commercial terms with any Referred Client unless expressly authorised in "
           "writing by the Principal.")
    clause(doc, "8.4", "Compliance",
           "The Referral Partner shall at all times comply with all applicable laws and "
           "professional regulations, including anti-bribery, anti-corruption, and data "
           "protection legislation.")
    clause(doc, "8.5", "Notification of Conflict",
           "The Referral Partner shall immediately disclose to the Principal any actual or "
           "potential conflict of interest in relation to any Referred Client.")
    spacer(doc)

    # ── 9. Confidentiality ────────────────────────────────────────────────────
    h1(doc, "9.  CONFIDENTIALITY")
    body(
        doc,
        "9.1  Each Party shall treat as strictly confidential all information received from "
        "the other Party in connection with this Agreement, including pricing, commission "
        "rates, client identities, technical specifications, and business strategies "
        "(\"Confidential Information\").",
    )
    spacer(doc)
    body(
        doc,
        "9.2  Neither Party shall disclose Confidential Information to any third party "
        "without the prior written consent of the disclosing Party, except to professional "
        "advisers bound by equivalent confidentiality obligations.",
    )
    spacer(doc)
    body(
        doc,
        "9.3  This obligation of confidentiality survives termination of this Agreement for "
        "a period of three (3) years.",
    )
    spacer(doc)

    # ── 10. Term & Termination ────────────────────────────────────────────────
    h1(doc, "10.  TERM & TERMINATION")
    clause(doc, "10.1", "Term",
           "This Agreement commences on the date of last signature and continues for a period "
           "of twenty-four (24) months, unless terminated earlier in accordance with this clause.")
    clause(doc, "10.2", "Renewal",
           "The Agreement shall automatically renew for successive twelve (12) month periods "
           "unless either Party gives thirty (30) days written notice of non-renewal before "
           "the end of the then-current term.")
    clause(doc, "10.3", "Termination for Convenience",
           "Either Party may terminate this Agreement for convenience on thirty (30) days "
           "written notice. Commission entitlement for Referrals already validated before "
           "the termination notice date is preserved.")
    clause(doc, "10.4", "Termination for Cause",
           "The Principal may terminate immediately on written notice if the Referral Partner: "
           "(a) breaches Clause 8 or 9; (b) is found to have engaged in fraudulent, corrupt, "
           "or illegal conduct; or (c) becomes insolvent. No Commission shall be payable "
           "following termination for cause.")
    spacer(doc)

    # ── 11. Taxes & Withholding ───────────────────────────────────────────────
    h1(doc, "11.  TAXES & WITHHOLDING")
    body(
        doc,
        "11.1  Each Party is solely responsible for its own tax obligations arising from "
        "this Agreement. Commission amounts stated herein are exclusive of VAT.",
    )
    spacer(doc)
    body(
        doc,
        "11.2  Where applicable under Cyprus law, the Principal may withhold tax at source "
        "from Commission payments and shall issue the requisite withholding tax certificates. "
        "The Referral Partner is responsible for declaring Commission income to the relevant "
        "tax authorities in its jurisdiction.",
    )
    spacer(doc)

    # ── 12. Limitation of Liability ───────────────────────────────────────────
    h1(doc, "12.  LIMITATION OF LIABILITY")
    body(
        doc,
        "12.1  The Principal's total liability to the Referral Partner under this Agreement "
        "shall not exceed the total Commission actually paid in the twelve (12) months "
        "preceding the relevant claim.",
    )
    spacer(doc)
    body(
        doc,
        "12.2  Neither Party shall be liable for indirect, consequential, or special loss, "
        "including lost profits or business opportunities, howsoever arising.",
    )
    spacer(doc)
    body(
        doc,
        "12.3  The Principal makes no warranty as to the commercial outcome of any EPC "
        "proposal submitted to a Referred Client, nor as to the acceptance of any quotation.",
    )
    spacer(doc)

    # ── 13. Governing Law ─────────────────────────────────────────────────────
    h1(doc, "13.  GOVERNING LAW & DISPUTE RESOLUTION")
    body(
        doc,
        "13.1  This Agreement is governed by, and construed in accordance with, the laws "
        "of the Republic of Cyprus.",
    )
    spacer(doc)
    body(
        doc,
        "13.2  The Parties shall endeavour to resolve any dispute amicably within twenty-one "
        "(21) days of written notice. Failing amicable resolution, disputes shall be subject "
        "to the exclusive jurisdiction of the courts of Limassol, Cyprus.",
    )
    spacer(doc)
    body(
        doc,
        "13.3  Any notice under this Agreement shall be in writing and delivered by email "
        "(with read receipt) or registered post to the addresses stated in Clause 1. "
        "Notices are effective on receipt.",
    )
    spacer(doc)

    # ── 14. General ───────────────────────────────────────────────────────────
    h1(doc, "14.  GENERAL")
    clause(doc, "14.1", "Entire Agreement",
           "This Agreement, together with its Schedules, constitutes the entire agreement "
           "between the Parties relating to its subject matter and supersedes all prior "
           "discussions, representations, and arrangements.")
    clause(doc, "14.2", "Amendments",
           "No amendment to this Agreement is valid unless made in writing and signed by "
           "authorised representatives of both Parties.")
    clause(doc, "14.3", "Waiver",
           "Failure to exercise, or delay in exercising, any right under this Agreement "
           "shall not constitute a waiver of that right.")
    clause(doc, "14.4", "Severability",
           "If any provision of this Agreement is found invalid or unenforceable, the "
           "remaining provisions shall continue in full force and effect.")
    clause(doc, "14.5", "Assignment",
           "The Referral Partner may not assign or transfer any rights or obligations under "
           "this Agreement without the prior written consent of the Principal.")
    clause(doc, "14.6", "Program Amendments",
           "The Principal reserves the right to amend Commission rates or program terms with "
           "thirty (30) days written notice. Such amendments shall not affect Commission "
           "entitlements arising from Referrals validated before the notice date.")
    spacer(doc)

    # ── Signatures ────────────────────────────────────────────────────────────
    divider(doc)
    spacer(doc)
    h1(doc, "EXECUTION")
    body(
        doc,
        "By signing below, the Parties confirm they have read, understood, and agreed to "
        "all terms and conditions of this Referral Commission Agreement.",
    )
    spacer(doc)

    sig = doc.add_table(rows=6, cols=2)
    sig.style = "Table Grid"

    sig_headers = [
        "FOR THE REFERRAL PARTNER",
        f"FOR THE PRINCIPAL — {COMPANY['name']}",
    ]
    for ci, val in enumerate(sig_headers):
        cell = sig.cell(0, ci)
        cell.text = val
        set_cell_bg(cell, "1A365D")  # NAVY
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)

    left_rows = [
        "Name:  ______________________________",
        "Capacity:  __________________________",
        "Signature:  _________________________",
        "Date:  ______________________________",
        "Stamp (if applicable):  ______________",
    ]
    right_rows = [
        f"Name:  {COMPANY['rep']}",
        f"Capacity:  {COMPANY['title']}",
        "Signature:  _________________________",
        "Date:  ______________________________",
        "Company Seal",
    ]

    for i, (l, r) in enumerate(zip(left_rows, right_rows), start=1):
        sig.cell(i, 0).text = l
        sig.cell(i, 1).text = r
        for ci in range(2):
            for para in sig.cell(i, ci).paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = BLACK

    spacer(doc)
    body(
        doc,
        "* Each Party should initial every page and retain one signed original. "
        "An electronically signed copy (PDF) shall have equal legal validity.",
        italic=True, colour=GREY,
    )

    spacer(doc)
    divider(doc)
    spacer(doc)

    # ── Schedule A — Referral Details ─────────────────────────────────────────
    h1(doc, "SCHEDULE A — REFERRAL DETAILS & COMMISSION CONFIRMATION")
    body(
        doc,
        "To be completed and signed by both Parties upon validation of each Referral.",
    )
    spacer(doc)
    two_col_table(doc, [
        ("Agreement Reference",          REF_NO),
        ("Referral Date",                "_______ / _______ / _______"),
        ("Referred Client Name",         "______________________________________"),
        ("Referred Client Contact",      "______________________________________"),
        ("PV Park Location(s)",          "______________________________________"),
        ("Licensed Capacity (MW)",       "______________________________________"),
        ("Estimated Project Value (€)",  "______________________________________"),
        ("Applicable Commission Tier",   "□ Tier A — 2.0%   □ Tier B — 3.0%   □ Tier C — 3.5%"),
        ("Commission Rate Confirmed",    "_______ %  of Contract Value"),
        ("Exclusivity Period Expiry",    "_______ / _______ / _______"),
        ("Principal Confirmation By",    "______________________________________"),
        ("Principal Confirmation Date",  "_______ / _______ / _______"),
    ], header=["Field", "Details"])

    spacer(doc)
    divider(doc)
    spacer(doc)

    # ── Schedule B — Payment Details ──────────────────────────────────────────
    h1(doc, "SCHEDULE B — REFERRAL PARTNER BANK DETAILS")
    body(
        doc,
        "Bank account details for Commission payments. "
        "The Referral Partner is responsible for notifying the Principal of any changes.",
    )
    spacer(doc)
    two_col_table(doc, [
        ("Account Holder Name",     "______________________________________"),
        ("Bank Name",               "______________________________________"),
        ("IBAN",                    "______________________________________"),
        ("SWIFT / BIC",             "______________________________________"),
        ("Bank Address",            "______________________________________"),
        ("Currency",                "Euro (€)"),
        ("VAT Registration No.",    "______________________________________"),
    ])

    spacer(doc)
    divider(doc)
    spacer(doc)

    # ── Footer note ───────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = footer_p.add_run(
        f"© 2026 {COMPANY['name']}  |  {COMPANY['address']}  |  {COMPANY['email']}  |  www.lighthief.com\n"
        "This agreement is confidential. Referral commissions are subject to execution of this Agreement. "
        "All pricing and commission rates are indicative and subject to the terms herein."
    )
    r_f.font.size = Pt(8)
    r_f.font.color.rgb = GREY
    r_f.italic = True

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    build("referral-commission-agreement-template.docx")
