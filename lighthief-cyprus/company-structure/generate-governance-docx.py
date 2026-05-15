#!/usr/bin/env python3
"""
Generate Word (.docx) versions of governance Markdown files in this folder.
Run from repo root: python lighthief-cyprus/company-structure/generate-governance-docx.py
Or from this folder: python generate-governance-docx.py

Requires: python-docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Lighthief brand (see .cursor/rules/lighthief-brand-identity.mdc)
NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x40, 0x40, 0x40)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex6: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex6)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shd)


def apply_runs(paragraph, text: str, base_size_pt: float = 10) -> None:
    """Split on **bold** markers."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(base_size_pt)
            r.font.color.rgb = BLACK
        else:
            r = paragraph.add_run(part)
            r.font.size = Pt(base_size_pt)
            r.font.color.rgb = BLACK


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = GOLD
    if level <= 1:
        r.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(0)
    elif level == 2:
        r.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(14)
    else:
        r.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            cell = tbl.rows[ri].cells[ci]
            txt = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            apply_runs(p, txt, 9 if ri > 0 else 10)
            for run in p.runs:
                run.font.name = "Calibri"
                if ri == 0:
                    run.bold = True
                    run.font.color.rgb = WHITE
                else:
                    run.font.color.rgb = BLACK
            if ri == 0:
                set_cell_bg(cell, "1A365D")


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.4)
    apply_runs(p, text.strip().lstrip("- ").strip(), 10)


def add_code_block(doc: Document, body: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(body if body.strip() else "(empty)")
    r.font.name = "Consolas"
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def is_table_sep(line: str) -> bool:
    s = line.strip().replace(" ", "")
    if not s.startswith("|"):
        return False
    inner = s.strip("|")
    return bool(re.match(r"^[-:|]+$", inner))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and is_table_row(lines[i]):
        if is_table_sep(lines[i]):
            i += 1
            continue
        parts = lines[i].strip().split("|")
        if parts and parts[0].strip() == "":
            parts = parts[1:]
        if parts and parts[-1].strip() == "":
            parts = parts[:-1]
        raw = [c.strip() for c in parts]
        rows.append(raw)
        i += 1
    return rows, i


def md_file_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    sty = doc.styles["Normal"]
    sty.font.name = "Calibri"
    sty.font.size = Pt(11)

    # Title from filename
    title = md_path.stem.replace("-", " ").replace("_", " ")
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(title)
    r0.bold = True
    r0.font.size = Pt(14)
    r0.font.color.rgb = NAVY
    doc.add_paragraph()

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
                lang = stripped[3:].strip()
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_buf.append(lines[i])
                    i += 1
                i += 1  # closing fence
                body = "\n".join(code_buf)
                if lang.lower() == "mermaid":
                    p = doc.add_paragraph()
                    apply_runs(
                        p,
                        "**Diagram:** Mermaid source is in the Markdown file; flowchart not rendered in Word.",
                        10,
                    )
                else:
                    add_code_block(doc, body)
                in_code = False
                continue
            in_code = False
            i += 1
            continue

        if is_table_row(line):
            rows, ni = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
            i = ni
            continue

        if stripped.startswith("#"):
            hashes = len(line) - len(line.lstrip("#"))
            title_text = line.lstrip("#").strip()
            add_heading(doc, title_text, hashes)
            i += 1
            continue

        if re.match(r"^[\s]*[-*]\s+", line):
            level = (len(line) - len(line.lstrip())) // 2
            add_bullet(doc, re.sub(r"^[\s]*[-*]\s+", "", line), min(level, 2))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        apply_runs(p, stripped, 11)
        p.paragraph_format.space_after = Pt(4)
        i += 1

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "Lighthief Cyprus Ltd · HE 477423 · solarfarms.cy · office@lighthief.com"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = GREY

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


def main() -> None:
    here = Path(__file__).resolve().parent
    md_files = sorted(here.glob("*.md"))
    if not md_files:
        print("No .md files found.", file=sys.stderr)
        sys.exit(1)
    for md in md_files:
        out = md.with_suffix(".docx")
        md_file_to_docx(md, out)
    print(f"Done. {len(md_files)} document(s).")


if __name__ == "__main__":
    main()
