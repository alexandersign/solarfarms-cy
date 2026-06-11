#!/usr/bin/env python3
"""Generate INTERNAL meeting prep .docx — Amit (management) + Yossi (sales).
Run from repo root:
  python docs/internal/meetings/generate_meeting_prep_amit_yossi_jun2026.py
Requires: python-docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x40, 0x40, 0x40)
NAVY_HEX = "1A365D"


def set_cell_bg(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.color.rgb = GOLD
    r.font.size = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)


def para(doc: Document, text: str, *, grey: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = GREY if grey else BLACK


def bullets(doc: Document, items: list[str]) -> None:
    for t in items:
        p = doc.add_paragraph(t, style="List Bullet")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for i, hdr in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = hdr
        set_cell_bg(cell, NAVY_HEX)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.bold = True
                run.font.color.rgb = WHITE
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    run.font.color.rgb = BLACK
    doc.add_paragraph()


def main() -> None:
    out_dir = Path(__file__).resolve().parent
    out = out_dir / "LCY-MEETING-PREP-AMIT-YOSSI-JUN2026.docx"

    doc = Document()
    sect = doc.sections[0]
    sect.top_margin = sect.bottom_margin = Cm(2)
    sect.left_margin = sect.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(
        "INTERNAL — Meeting prep\n"
        "Amit: Management implementation  |  Yossi: Sales & pipeline"
    )
    tr.bold = True
    tr.font.name = "Calibri"
    tr.font.size = Pt(16)
    tr.font.color.rgb = NAVY
    title.paragraph_format.space_after = Pt(6)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("June 2026 · Lighthief Cyprus Ltd · HE 477423")
    sr.font.name = "Calibri"
    sr.font.size = Pt(11)
    sr.font.color.rgb = GREY
    sub.paragraph_format.space_after = Pt(14)

    para(
        doc,
        "Classification: INTERNAL — not for client circulation. Figures from lib/portfolio-data.ts "
        "(May 2026). Verify against executed contracts before reliance. Do not share CIF, margins, "
        "or quotation references externally.",
        grey=True,
    )

    h(doc, "Meeting structure", 1)
    table(
        doc,
        ["Segment", "With", "Focus", "Duration"],
        [
            ["Part A", "Amit", "Implementation & management process", "~30–40 min"],
            ["Part B", "Yossi", "Sales, pipeline, closes", "~30–40 min"],
        ],
    )
    para(doc, "Shared anchor: Galascope Batch 1 (2 parks, 7.5 MW / 30 MWh, €3.44M ex VAT) — first execution test and first signed revenue gate.", bold=True)

    # ─── PART A ───
    doc.add_page_break()
    h(doc, "Part A — Amit: Management & implementation", 1)

    para(doc, "Opening line:", bold=True)
    para(
        doc,
        "We are moving from deal packaging to repeatable EPC delivery: one confirmed batch (Galascope), "
        "then scale the same playbook across the group order and pipeline LOIs.",
    )

    h(doc, "A1. Management model — who does what", 2)
    bullets(
        doc,
        [
            "Alexander Papacosta — Cyprus Director / programme owner: client, lender, Linyang/Voltus, commissioning presence.",
            "Costas Hadjikyriacou — ETEK: DSO applications, witness test, planning, O&M team lead (Cyprus).",
            "Andreas Christoforou — BESS consultant (UK → Cyprus): energy systems, specs, stakeholder support.",
            "Field O&M — Dawid, Szymon, Kacper (Poland-trained, deploying to Cyprus).",
            "Linyang — Kamil (install/SCADA), Michael (product), China commissioning team.",
            "Subcontractors — civil, electrical install (TBD award), Soulis transport, Joha cables, DEHN LPS.",
            "Lighthief = EPC integrator + O&M manager; client SPVs own assets. Single point of responsibility to client.",
        ],
    )

    h(doc, "A2. Implementation phases (repeatable per batch)", 2)
    table(
        doc,
        ["Phase", "Duration", "Owner", "Management focus"],
        [
            ["0 — Contract & CPs", "Until EPC stamped", "Alexander + legal", "APG, DWU, CAR insurance, Disperon SHA rates"],
            ["1 — Pre-order", "6–8 weeks", "Costas", "DSO / connection terms, permits, civil design"],
            ["2 — Production", "~90 days", "Linyang + Alexander", "FAT, 55% pre-shipment milestone"],
            ["3 — Logistics", "~50 days", "Alexander + Soulis", "Permits 15 d; €2,500/container; Famagusta routes"],
            ["4 — Site EPC", "Overlap G1+G2", "Costas + subs", "Parallel installs — not sequential"],
            ["5 — Commissioning", "→ PAC", "Costas + field", "SCADA IEC-104, DSO witness"],
            ["6 — O&M handover", "From PAC", "Costas + engineers", "LTSA, Voltus NOC, quarterly visits"],
        ],
    )
    para(doc, "Galascope PAC target: 31 January 2027 (both parks). Production May–Jul 2026; ship Aug 2026; CIF Sep 2026.")

    h(doc, "A3. Galascope execution risks", 2)
    table(
        doc,
        ["Risk", "Management response"],
        [
            ["G1 and G2 in parallel (Famagusta)", "Budget +0.3–0.7 FTE or site coordinator (~€15–45k overlap)"],
            ["One supervisor, two crane windows", "Second site cover or extended hours"],
            ["PM load (Soulis, Linyang, DSO ×2)", "Logistics / site admin fraction"],
            ["Thin batch margin (~7.3% on €3.19M installed)", "~1% cost slip ≈ €32k; large overrun can wipe batch"],
        ],
    )
    para(doc, "Decision for Amit: approve parallel programme with incremental supervision budget, or accept schedule/quality risk.")

    h(doc, "A4. Subcontractor & procurement governance", 2)
    table(
        doc,
        ["Workstream", "Status", "Action"],
        [
            ["Civil (€2,000/MWh)", "Confirmed", "Monitor subcontractor delivery"],
            ["Transport (A. Soulis)", "Confirmed", "Per-container playbook"],
            ["Cables (Joha)", "Per-park quoted", "Cost control per park total"],
            ["Electrical install", "No RFP awarded", "PRIORITY — blocks execution"],
            ["CAR insurance", "Marsh in progress", "Bind before Linyang advance"],
            ["EMS/SCADA (Voltus)", "Quoted", "SHA Exhibit C rates — 14 days post-sign"],
        ],
    )

    h(doc, "A5. O&M management ramp", 2)
    bullets(
        doc,
        [
            "Portfolio (active model): 29 parks, ~517 MWh — Galascope adds 30 MWh.",
            "Monitoring: Voltus SCADA + Lighthief NOC 24/7.",
            "Recommended field staffing: 5 engineers (2 night / 2 day / 1 float); ~320 night shifts/year (batched).",
            "Planned maintenance: midnight–6am; critical response 1 hr remote / 4 hr on-site.",
            "Engineers completing 3-week Czestochowa training (BMS, PCS, EN 50549-2, DSO witness).",
            "LTSA ~€1,740/MWh/yr — separate revenue stream; do not under-resource because EPC margin is tight.",
        ],
    )

    h(doc, "A6. Governance & resourcing", 2)
    bullets(
        doc,
        [
            "Origination pool (up to 30% net margin): existing Cyprus team that built backlog — not new hires unless agreed.",
            "Galascope EPC sign triggers operational scale (MD salary step, parallel investor tracks).",
            "Investor-introduced staff: separate employment/advisory terms; no automatic commission on closed deals.",
        ],
    )

    h(doc, "A7. Questions for Amit", 2)
    bullets(
        doc,
        [
            "What management systems expected (PM tool, weekly cadence, milestone dashboard)?",
            "Staffing Cyprus PMO vs advisory only?",
            "Cap on parallel sites before dedicated programme manager?",
            "PV-O&M integration — shared field team or separate BESS unit?",
        ],
    )

    h(doc, "A8. Amit — next 30 days", 2)
    table(
        doc,
        ["#", "Action", "Owner"],
        [
            ["1", "Galascope EPC signed", "Alexander + Dino"],
            ["2", "Electrical subcontractor awarded", "Alexander"],
            ["3", "CAR insurance bound", "Alexander"],
            ["4", "Parallel site programme issued", "Costas / PM"],
            ["5", "O&M roster confirmed for PAC window", "Costas"],
        ],
    )

    # ─── PART B ───
    doc.add_page_break()
    h(doc, "Part B — Yossi: Sales & pipeline", 1)

    para(doc, "Opening line:", bold=True)
    para(
        doc,
        "Cyprus BESS sales are relationship-led group orders — Esperia/Galascope is the anchor; "
        "the rest of the ~€61M group backlog follows the same EPC + LTSA product.",
    )

    h(doc, "B1. Product (client language)", 2)
    table(
        doc,
        ["Layer", "What the client buys"],
        [
            ["Turnkey EPC", "Linyang BESS + install + commissioning + grid witness support"],
            ["LTSA", "15-year O&M, availability, performance management"],
            ["EMS/SCADA", "Voltus/Disperon — €400/MWh/yr (Galascope model)"],
            ["Extended warranty", "Yrs 6–15 — paid direct to Linyang (not Lighthief P&L)"],
        ],
    )
    para(
        doc,
        "Elevator pitch: Exclusive Linyang distributor; largest Cyprus utility-scale BESS rollout; "
        "single contractor from container to PAC; ETEK sign-off and DSO coordination included.",
    )

    h(doc, "B2. Pipeline funnel", 2)
    para(doc, "Active model: 29 parks · ~517 MWh · ~€64M revenue when fully contracted (28 group + Aeolian).")
    table(
        doc,
        ["Priority", "Client", "MWh", "Probability", "Revenue (approx)", "Sales action"],
        [
            ["1", "Galascope (Esperia)", "30", "Confirmed — legal", "€3.44M", "Support counsel; hold G1/G2 price"],
            ["2", "Spanercom (Anarita ×2)", "40", "~85%", "~€4.76M", "EPC day after Galascope sign"],
            ["3", "Timotheos (9 parks)", "81", "~50%", "~€11.8M", "EPC signature for B1 slot"],
            ["3", "Lampros (2 parks)", "15", "~50%", "~€2.1M", "Same"],
            ["4", "Esperia pipeline LOI", "285+", "Post-Galascope", "~€33M+", "Batches 2026–28"],
            ["5", "Aeolian (wind hybrid)", "20", "~85%", "€2.66M", "Separate track"],
            ["Park", "Kerasi / Karis", "45", "~30%", "~€6.3M", "After B1 committed"],
            ["Lost", "ABIO", "200", "~20%", "—", "Lost to CATL; do not forecast"],
        ],
    )

    h(doc, "B3. Galascope — closed vs still selling", 2)
    para(doc, "Closed commercially (May 2026):", bold=True)
    bullets(
        doc,
        [
            "G1: €2,238,000 (20 MWh effective); G2: €1,206,300 (10 MWh).",
            "Payments: 30 / 55 / 10 / 5; 3-month DLP; PAC 31 January 2027.",
        ],
    )
    para(doc, "Still with counsel (do not re-trade price):", bold=True)
    bullets(
        doc,
        [
            "APG transfer to Galascope / Alpha Bank.",
            "EMS addendum €400/MWh/yr for 5 years from PAC.",
            "Scope: Costas sign-off included; licensed site drawings excluded (€5–15k/site).",
        ],
    )
    para(doc, "Do NOT say in client meetings: CIF, margins, quotation numbers, internal pipeline totals.", grey=True)

    h(doc, "B4. Sales narrative by segment", 2)
    bullets(
        doc,
        [
            "Esperia / Dino: Galascope proves model for 9-park pipeline LOI; split LOIs reduce legal drag.",
            "Spanercom: Same hardware as Galascope G1; ~€119k/MWh — send EPC immediately after Galascope stamp.",
            "Timotheos / Lampros: Verbal intent — urgency: sign by Jun 2026 or lose B1 production slot.",
            "Aeolian: Wind + BESS flagship; €2.66M turnkey — marketing and standalone close.",
        ],
    )

    h(doc, "B5. Competitive positioning", 2)
    table(
        doc,
        ["Objection", "Response"],
        [
            ["CATL / €124k/MWh flat (ABIO)", "We sell turnkey EPC + grid compliance + LTSA, not container-only"],
            ["Cheaper container quote", "Client still needs civil, DSO, SCADA, install — our price is installed PAC"],
            ["6-month DLP ask", "3 months with stronger bank security package"],
        ],
    )

    h(doc, "B6. Revenue phasing (forecast)", 2)
    table(
        doc,
        ["Wave", "Timing", "MWh (approx)", "Note"],
        [
            ["B1 Galascope", "Sign → PAC Jan 2027", "30", "Advance after Connection Terms"],
            ["B1 extension", "If 7 parks sign", "90", "Same slot as Galascope"],
            ["Esperia 2026 pipeline", "Post-trust", "75", "LOI batch 2"],
            ["2027 / 2028", "LOI", "115 + 125", "Frenaros 100 MWh flagship 2028"],
        ],
    )

    h(doc, "B7. Origination & incentives", 2)
    bullets(
        doc,
        [
            "Up to 30% of net margin on executed projects → existing Cyprus origination team (pre-investment).",
            "New sales hires / investor-introduced personnel → outside pool unless expressly agreed.",
            "Clarify Yossi role: closing vs origination vs channel (Israel / EU) before promising commission.",
        ],
    )

    h(doc, "B8. Prospecting (beyond funnel)", 2)
    table(
        doc,
        ["Director / group", "BESS MW", "Angle"],
        [
            ["Timotheos Timotheou", "61", "Already in funnel — deepen"],
            ["HELIOPEX (Patssalis / Kallenos)", "13 each", "Post-B1 hybrid EPC"],
            ["GKA Georgiou", "6.4", "Smaller hybrid"],
            ["Santiam / Stephanidis", "0 BESS", "PV O&M cross-sell only"],
        ],
    )

    h(doc, "B9. Questions for Yossi", 2)
    bullets(
        doc,
        [
            "Which accounts does he personally own vs support?",
            "Israel / EU channel — exclusive or introducer?",
            "After Galascope: Spanercom first or Timotheos block for volume?",
            "Need one-page commercial summary for new meetings?",
        ],
    )

    h(doc, "B10. Yossi — sales targets", 2)
    table(
        doc,
        ["#", "Target", "By"],
        [
            ["1", "Galascope EPC + EMS addendum signed", "ASAP"],
            ["2", "Spanercom EPC issued", "48h after Galascope"],
            ["3", "Timotheos + Lampros chase meetings", "Jun 2026"],
            ["4", "Aeolian signature path", "Q2 2026"],
            ["5", "Esperia 2026 pipeline LOI paper", "Post-Galascope / Q3 2026"],
        ],
    )

    h(doc, "Handoff (30 seconds)", 1)
    para(
        doc,
        "To Yossi (after Amit): Management is building the delivery machine for Galascope; sales must stamp "
        "Galascope then clone the contract to Spanercom and B1 extension (~€10M+ ex VAT in 90 days if executed).",
    )
    para(
        doc,
        "To Amit (before Yossi): Sales holds €3.44M at counsel; execution prep cannot wait for ink — "
        "electrical sub and parallel site plan must run in parallel.",
    )

    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "Lighthief Cyprus Ltd · HE 477423 · office@lighthief.com · +357 99 164 158 · solarfarms.cy\n"
        "Prepared June 2026 · Source: lib/portfolio-data.ts · INTERNAL"
    )
    fr.font.name = "Calibri"
    fr.font.size = Pt(9)
    fr.font.color.rgb = GREY

    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
