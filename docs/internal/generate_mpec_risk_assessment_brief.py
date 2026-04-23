#!/usr/bin/env python3
"""
Generate MPEC Risk Assessment Brief (DOCX)
Briefing document for Michaelis at MPEC covering:
  - Legal basis (Κ.Δ.Π. 15/2026, Official Gazette 5992)
  - Certifications & documents already held (Linyang battery + Kehua PCS)
  - Service required from ΕΞΥΠΠ-certified engineer
Date: April 2026
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

# ── Brand colours ──────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1A, 0x36, 0x5D)
GOLD  = RGBColor(0xC9, 0xA4, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x40, 0x40, 0x40)

TODAY = datetime.date.today().strftime("%-d %B %Y")

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = BLACK
style.paragraph_format.space_after = Pt(5)


# ── Helper: cell background colour ────────────────────────────────────────────
def set_cell_bg(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour)
    tcPr.append(shd)


# ── Heading helpers ────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = GOLD
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)


def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)


def para(text, bold=False, italic=False, colour=BLACK):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(10.5)
    r.font.color.rgb = colour
    return p


def bullet(text, level=0, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLACK
    p.paragraph_format.left_indent   = Cm(1.0 + level * 0.5)
    p.paragraph_format.space_after   = Pt(3)


def add_table(headers, rows, col_widths=None):
    """Add a styled table with navy header row."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '1A365D')
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            if ri % 2 == 1:
                set_cell_bg(cell, 'F5F7FA')

    # Column widths
    if col_widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Cm(w)

    doc.add_paragraph()
    return t


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════════════════

# ── Title block ───────────────────────────────────────────────────────────────
doc.add_paragraph()
h1('BESS Risk Assessment — Briefing Note')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Prepared for: Michaelis — MPEC\nPrepared by: Lighthief Cyprus Ltd')
r.font.size = Pt(11)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = meta.add_run(f'Date: {TODAY}    |    Reference: LCY-EXIPP-BRIEF-001    |    Classification: Confidential')
r2.font.size = Pt(9)
r2.font.color.rgb = GREY
r2.italic = True

doc.add_paragraph()

# ── Section 1: Purpose ────────────────────────────────────────────────────────
h2('1.  Purpose of This Document')
para(
    'Lighthief Cyprus Ltd is the EPC contractor and system integrator for a portfolio of Battery Energy Storage '
    'System (BESS) installations on licensed solar parks across Cyprus. This briefing note is addressed to the '
    'certified ΕΞΥΠΠ engineer at MPEC and sets out:'
)
bullet('The legal obligation that triggers the requirement for a Written Risk Assessment;')
bullet('The manufacturer certifications and technical documentation already in our possession (Linyang battery + Kehua PCS);')
bullet('The documents and site-specific inputs that remain outstanding; and')
bullet('The scope of service we are requesting from MPEC.')


# ── Section 2: Legal Basis ────────────────────────────────────────────────────
h2('2.  Legal Basis')
para(
    'The requirement arises directly from Κ.Δ.Π. 15/2026, published in Official Gazette No. 5992 on '
    '16 January 2026, which establishes the conditions under which BESS installations within licensed '
    'RES (solar) stations are exempt from building and planning permits in Cyprus.'
)
para(
    'The relevant conditions under the Labour Inspection Department (Τμήμα Εργασίας) are:'
)

add_table(
    headers=['Condition No.', 'Requirement (English)', 'Greek Term'],
    rows=[
        ['32', 'A Written Risk Assessment must be prepared covering all battery-related hazards.', 'Γραπτή Εκτίμηση Κινδύνου'],
        ['33', 'The risk assessment must be based on the manufacturer\'s specifications and prepared by a certified ΕΞΥΠΠ or ΕΣΥΠΠ engineer.', 'Εκτίμηση βάσει προδιαγραφών κατασκευαστή — ΕΞΥΠΠ/ΕΣΥΠΠ'],
    ],
    col_widths=[3.0, 9.0, 4.5]
)

para(
    'The permit exemption (and therefore the legal ability to commence installation works) is '
    'conditional on this document being in place. It is not a post-installation formality.',
    bold=False, italic=True, colour=GREY
)


# ── Section 3: System Description ─────────────────────────────────────────────
h2('3.  System Description')
para(
    'Each BESS installation in the portfolio consists of the following major components, sourced '
    'from two manufacturers:'
)

add_table(
    headers=['Component', 'Manufacturer', 'Key Model(s)', 'Function'],
    rows=[
        ['Battery Container (ESS)', 'Linyang Energy (中国储能)', 'LY-Ocean Series (LFP, 20HC container)', 'Energy storage — Lithium Iron Phosphate cells'],
        ['Battery Management System', 'Linyang Energy', 'Integrated BMS + EMS', 'Cell monitoring, protection, SOC/SOH management'],
        ['Power Conversion System (PCS)', 'Kehua Digital Energy', 'BCS1250K-C-HUD (T1 skid)', 'AC/DC conversion, grid interface'],
        ['MV Skid / Transformer', 'Kehua Digital Energy', 'BCS T1/T2/T4/T8 series (0.69/22kV)', 'MV step-up and switchgear'],
        ['Thermal Management (BTMS)', 'Linyang (Kelvin BTMS)', 'Liquid cooling, 45 kW per container', 'Temperature control'],
    ],
    col_widths=[4.0, 4.0, 4.5, 4.0]
)

para(
    'Typical project size: 2.5 MW / 10 MWh per site (2× 1.25 MW PCS + 4–5× battery containers). '
    'The systems are Category C standalone storage facilities under the Cyprus DSO Technical Guide 2025.1.'
)


# ── Section 4: Documents Already in Our Possession ────────────────────────────
h2('4.  Manufacturer Documents & Certifications Already Available')
para(
    'The following documents are held by Lighthief Cyprus Ltd and will be provided to MPEC in full '
    'upon engagement. They form the manufacturer specification basis required by Condition 33.'
)

h3('4.1  Linyang Battery — Certifications')
add_table(
    headers=['Standard / Certification', 'Scope', 'Certificate / Reference', 'Risk Assessment Relevance'],
    rows=[
        ['UL 9540A (Unit Level)', 'No thermal runaway propagation between modules', '80239432', 'Core fire risk finding — confirms flame containment at unit level'],
        ['IEC 62619:2022', 'Secondary battery safety — CB Scheme', 'BE-51213', 'General battery safety baseline'],
        ['IEC 63056:2020', 'Safety requirements for secondary lithium cells', 'SG SGS-00632', 'Cell-level safety compliance'],
        ['UL 1973', 'Battery systems for stationary applications (cell level)', 'MH63503', 'Stationary storage safety'],
        ['UN 38.3', 'Transport safety classification for lithium batteries', '01112300005127', 'Hazardous goods classification'],
        ['EN 60204-1 / EN ISO 12100', 'Machinery safety — Thermal Management System (BTMS)', 'M7 122013 0012', 'Cooling system mechanical safety'],
        ['IEC 61000-6-2 / 61000-6-4', 'EMC immunity and emissions — container level', 'E6A 130105 0005', 'Electromagnetic compatibility'],
        ['IEC 62933-5-2', 'Grid integration safety — in progress / recently issued', '704082542702-00', 'Grid-level safety — status to be confirmed'],
    ],
    col_widths=[4.2, 4.5, 3.5, 4.3]
)

h3('4.2  Kehua PCS — Certifications')
add_table(
    headers=['Standard / Certification', 'Scope', 'Document', 'Risk Assessment Relevance'],
    rows=[
        ['EN 50549-2', 'Grid connection standard for generating plants >16A — full certificate + 41MB test report', 'BCS1250K-C-HUD EN50549-2 cert + report', 'LVRT/HVRT behaviour, grid fault response'],
        ['IEC 62477-1 (CB Scheme)', 'Safety of power electronics — TÜV SÜD cert, valid to 2030', 'BCS1000~1250K IEC62477-1 CB cert', 'PCS electrical safety'],
        ['CE-LVD (Low Voltage Directive)', 'EU market access — Low Voltage', 'BCS1000~1250K CE-LVD cert', 'LVD compliance'],
        ['CE-EMC', 'Electromagnetic compatibility', 'BCS1000~1250K CE-EMC cert', 'EMC compliance'],
        ['IEC 62116 + IEC 61727', 'Anti-islanding + grid connection characteristics', 'BCS1000~1250K cert', 'Anti-islanding protection'],
        ['IEC 62909-1/-2 (TÜV SÜD)', 'Bidirectional power conversion systems', 'BCS1000-1250K cert', 'PCS safety for storage applications'],
    ],
    col_widths=[4.2, 5.0, 4.0, 3.3]
)

h3('4.3  Technical Data Available for Risk Assessment Drafting')
add_table(
    headers=['Document', 'Content', 'Source'],
    rows=[
        ['Kehua BCS1250K efficiency curves', 'Operating efficiency 97.5–99%; operating temperature range -35°C to +60°C', 'Kehua (on file)'],
        ['Kehua PCS P-Q capability curves (single + parallel)', 'Reactive power capability at all operating points', 'Kehua (on file)'],
        ['FRT / Frequency Response curves', 'Behaviour under grid frequency disturbance events', 'Kehua (on file)'],
        ['THD + Flicker + DC Injection test results', 'Power quality; DC injection <0.5%; THDi <3%', 'Kehua (on file)'],
        ['IEC 60870-5-104 SCADA point list', 'Remote control and monitoring register map', 'Kehua (on file)'],
        ['Kehua MV Skid datasheets (T1/T2/T4/T8)', 'Physical dimensions, weight, IP ratings, MV transformer specs', 'Kehua (on file)'],
        ['System Single Line Diagrams (22kV)', 'Reference SLDs from comparable Cyprus BESS projects', 'Kehua (on file)'],
        ['Linyang User Manual V2.0', 'Installation, commissioning, maintenance, and safety procedures', 'Linyang (on file)'],
        ['DSO Technical Guide 2025.1 (EAC)', 'Cyprus DSO safety distances, fire, SCADA, and protection requirements', 'EAC (public document)'],
        ['Κ.Δ.Π. 15/2026 and Κ.Δ.Π. 17/2026', 'Full text of the permit exemption decree including all 33 conditions', 'Official Gazette 5992'],
    ],
    col_widths=[5.0, 6.5, 4.0]
)


# ── Section 5: What Is Still Missing ──────────────────────────────────────────
h2('5.  Documents Pending — To Be Provided Before / During Engagement')
para(
    'The following items are not yet in our possession. We are actively pursuing these from Kehua and Linyang '
    'and expect to receive them within the coming weeks. We will share them with MPEC as soon as available.'
)

add_table(
    headers=['Item', 'From', 'Status', 'Impact on Risk Assessment'],
    rows=[
        ['Kehua PCS Installation & Commissioning Manual', 'Kehua', 'Requested — pending receipt', 'Needed for installation hazard section (PPE, isolation procedures, live-work rules)'],
        ['Kehua Protection Settings Guide', 'Kehua', 'Requested — pending receipt', 'Relay coordination data for electrical fault risk section'],
        ['UL 9540A — Installation Level test report', 'Linyang', 'Unit level passed; installation level pending', 'Required to assert no fire propagation between containers at site level (TSO requirement)'],
        ['IEC 62933-5-2 certificate (final)', 'Linyang', 'Testing completed; certificate issuance pending', 'Grid integration safety — confirms system-level behaviour'],
        ['Site-specific layout plan per project', 'Lighthief / Civil engineer', 'Varies per site', 'Container spacing, boundary setbacks, fire access road dimensions'],
    ],
    col_widths=[4.5, 2.5, 3.5, 6.0]
)


# ── Section 6: Scope of Service Required ──────────────────────────────────────
h2('6.  Scope of Service Required from MPEC')
para(
    'We are requesting the following services from MPEC for the Cyprus BESS portfolio. '
    'The risk assessment must be site-specific; however, a standardised methodology based on '
    'the common Linyang/Kehua platform can be developed once and adapted per project.'
)

h3('6.1  Written Risk Assessment (Γραπτή Εκτίμηση Κινδύνου)')
para('One risk assessment document per project site, to be authored and signed by the ΕΞΥΠΠ-certified engineer at MPEC. Each document must cover, at minimum:')
bullet('Electrical hazards: DC arc flash (1000–1500 V DC), AC fault, short circuit at PCS and MV level')
bullet('Fire hazard: thermal runaway risk, fire suppression system (aerosol + sprinkler), container spacing (≥1 m general, ≥3 m at access doors), 6 m fire access perimeter')
bullet('Chemical hazard: LFP off-gas (H₂, CO, CO₂, HC) — emission scenarios, ventilation and detection requirements')
bullet('Mechanical/structural hazard: container weight (up to 25,000 kg for 5 MW skid), seismic considerations (3.55 m/s² horizontal per TSO specification), lifting and placement during installation')
bullet('Environmental hazard: electrolyte containment, battery recycling obligations, decommissioning procedures')
bullet('Operational hazard: maintenance access procedures, PPE requirements, emergency shutdown procedures, hot-work controls (no welding/grinding within 6 m of energised containers)')
bullet('Commissioning hazard: first energisation procedures, SAT conditions, isolation and earthing protocols')

h3('6.2  Basis of Assessment')
para('The risk assessment must explicitly reference the manufacturer specifications as required by Condition 33. The primary source documents are:')
bullet('Linyang User Manual V2.0 (battery system)')
bullet('Kehua BCS1250K-C-HUD technical datasheet and installation manual')
bullet('UL 9540A test report (Linyang, unit level)')
bullet('EN 50549-2 test report (Kehua PCS)')
bullet('Cyprus DSO Technical Guide 2025.1 (safety distances)')
bullet('Κ.Δ.Π. 15/2026, Conditions 18–31 (fire, environment, and labour requirements)')

h3('6.3  ETEK Registration & Professional Liability')
para(
    'The signing engineer must hold a valid ETEK professional licence and provide a copy of their '
    'professional liability insurance certificate. This is required under Section 18.6 (Appendix II) '
    'of Κ.Δ.Π. 15/2026 as part of the Υπεύθυνη Δήλωση Μελετητή submission.'
)

h3('6.4  Deliverable Format')
add_table(
    headers=['Deliverable', 'Format', 'Notes'],
    rows=[
        ['Written Risk Assessment per site', 'PDF (signed) + editable DOCX', 'One per project; reuse common methodology across portfolio'],
        ['Engineer\'s Declaration (ETEK number + insurance)', 'PDF scan', 'To be attached to each site file'],
        ['Summary compliance matrix', 'DOCX or PDF', 'Optional — useful for DSO/insurance submission packages'],
    ],
    col_widths=[5.5, 3.5, 7.5]
)

h3('6.5  Timeline')
para(
    'The first installations are expected to commence in Q3 2026. We require the first risk assessment '
    'to be completed at least 4 weeks prior to the scheduled start of civil works on each site, to allow '
    'time for the Υπεύθυνη Δήλωση Μελετητή to be submitted and any queries from the Labour Inspection '
    'Department to be resolved.'
)


# ── Section 7: Next Steps ──────────────────────────────────────────────────────
h2('7.  Proposed Next Steps')
add_table(
    headers=['Step', 'Action', 'Owner', 'Target'],
    rows=[
        ['1', 'Introductory call to review scope and methodology', 'Lighthief + MPEC', 'Within 1 week'],
        ['2', 'Share full document pack (certifications + manuals on file)', 'Lighthief', 'Upon engagement'],
        ['3', 'MPEC reviews documents and confirms any additional items required', 'MPEC', 'Within 2 weeks of step 2'],
        ['4', 'Receive Kehua installation manual + Linyang UL 9540A installation report', 'Lighthief (from OEMs)', 'Ongoing — expected within 4–6 weeks'],
        ['5', 'MPEC drafts standard risk assessment methodology for Linyang/Kehua platform', 'MPEC', 'Following step 3'],
        ['6', 'Site-specific risk assessments issued per project, starting with first installations', 'MPEC', 'Q2–Q3 2026'],
    ],
    col_widths=[1.5, 7.0, 3.5, 3.5]
)


# ── Section 8: Contact ─────────────────────────────────────────────────────────
h2('8.  Contact')
add_table(
    headers=['Company', 'Contact', 'Role', 'Email'],
    rows=[
        ['Lighthief Cyprus Ltd', 'Alexander Papacosta', 'Managing Director / EPC Lead', 'alex@lighthief.com'],
        ['Lighthief Cyprus Ltd', 'Costas Hadjikyriacou', 'Technical Manager', 'costas@lighthief.com'],
    ],
    col_widths=[4.0, 4.0, 4.5, 4.0]
)

doc.add_paragraph()
para(
    'This document is confidential and intended solely for the use of Michaelis and the team at MPEC. '
    'Please do not distribute without prior written consent from Lighthief Cyprus Ltd.',
    italic=True, colour=GREY
)

# ── Save ───────────────────────────────────────────────────────────────────────
out = 'docs/internal/mpec-risk-assessment-brief-apr2026.docx'
doc.save(out)
print(f'✅  Saved: {out}')
