#!/usr/bin/env python3
"""Build the Aeolian LTSA (03) from the Esperia/Galascope LTSA template.

Bases off LTSA-Galascope-Esperia-may2026.docx and updates it for the standalone
single-park Aeolian project:
  - Party / reference / version identity → Aeolian (HE 168239, LCY-LTSA-AEO-2026)
  - Removes group / portfolio availability framing (Aeolian is ONE park)
  - Strips the internal €661–€3,181/MWh "portfolio range" EMS pricing leak
  - EMS/SCADA subscription set to the Aeolian rate (EUR 400/MWh/yr = EUR 8,024/yr)
  - Fixes residual "bi-annual" → "quarterly" (maintenance is quarterly)
  - Schedule 1 site/equipment filled for Agia Anna / T8 config
  - Signatory Sotiris Shiacallis; companion doc ref → LCY-EPC-AEO-2026

Run: python docs/clients/Individual_Aeolian_Dynamics_Larnaca/contracts/v1/build-aeolian-ltsa-jun2026.py
"""
from __future__ import annotations

from pathlib import Path
from docx import Document

V1 = Path(__file__).resolve().parent
REPO = V1.parents[4]
SRC = (REPO / "docs" / "clients" / "group-order" / "Group2_Esperia_Energy"
       / "contracts" / "LTSA-Galascope-Esperia-may2026.docx")
PKG = V1 / "CLIENT-PACKAGE"
OUT = PKG / "03-LTSA-Aeolian-Agia-Anna.docx"

# Substring replacements (applied within each paragraph's runs)
SUBS = [
    # Identity / reference
    ("LCY-LTSA-GAL-2026", "LCY-LTSA-AEO-2026"),
    ("LCY-EPC-GAL-B1-2026 v4.0", "LCY-EPC-AEO-2026"),
    ("LCY-EPC-GAL-B1-2026", "LCY-EPC-AEO-2026"),
    # Client party
    ("Galascope Ltd,",
     "T.P. Aeolian Dynamics Ltd,"),
    ("with registered office at Karaiskaki 6, City House, 3032 Limassol, Cyprus, "
     "Registration No. HE 303759",
     "with registered office at Karaiskaki 13, 3032 Limassol, Cyprus, Company No. HE 168239"),
    ("Galascope Ltd", "T.P. Aeolian Dynamics Ltd"),
    # Maintenance frequency leftovers
    ("bi-annual servicing", "quarterly servicing"),
    ("BI-ANNUAL ON-SITE MAINTENANCE CHECKLIST", "QUARTERLY ON-SITE MAINTENANCE CHECKLIST"),
    # SCADA dashboard accounts — single client, not a group
    ("10 read-only SCADA dashboard accounts per client group",
     "10 read-only SCADA dashboard accounts"),
    # EMS/SCADA subscription header row + basis note
    ("| EMS/SCADA Software Subscription | €[●] | €[●] |",
     "| EMS/SCADA Software Subscription | €400 / MWh / Year | €8,024 |"),
    ("(20% of installed EMS/SCADA cost per year)",
     "(DISPERON software subscription — see EMS Subscription Addendum)"),
    # EMS subscription detail table (Aeolian values; no internal cost basis)
    ("| System Capacity | [●] MWh |", "| System Capacity | 20.06 MWh |"),
    ("| Installed EMS/SCADA Cost | €[●] |", "| EMS/SCADA Subscription Rate | €400 / MWh / Year |"),
    ("| Annual Subscription (20%) | €[●] |", "| Annual EMS/SCADA Subscription | €8,024 |"),
    ("| EUR/MWh/Year | €[●] |", "| Billing | Annually in advance from PAC |"),
    # Schedule 2 System Capacity line
    ("SYSTEM CAPACITY: [●] MWh", "SYSTEM CAPACITY: 20.06 MWh"),
    # Schedule 1 site
    ("Site Name: [●]", "Site Name: Agia Anna Wind Farm Hybrid (BESS)"),
    ("Site Address: [●]", "Site Address: Agia Anna, Larnaca, Cyprus"),
    # Schedule 1 equipment rows
    ("| Battery Container | Linyang Power Atlantic 5MWh | [●] | [●] |",
     "| Battery Container | Linyang Power Atlantic 5.015 MWh | 4 | [●] |"),
    ("| Power Conversion System | [●] MW PCS | [●] | [●] |",
     "| Power Conversion System | Kehua BCS1000K (1.0 MW), 8 units in one T8 MV skid | 8 | [●] |"),
    ("| EMS | [Third Party – specify] | [●] | [●] |",
     "| EMS | DISPERON (R&D Innovations Sp. z o.o.) | 1 | [●] |"),
    # Signature block
    ("For [End Customer]", "For T.P. Aeolian Dynamics Ltd"),
]

# Full-paragraph rewrites (group/portfolio framing → single park). Keyed by a
# distinctive substring; the whole paragraph text is replaced.
REWRITES = [
    ("VERSION HISTORY",
     "Document Reference: LCY-LTSA-AEO-2026\nVersion: 1.0\nDate: June 2026"),
    ("the BESS portfolio covered under this Agreement shall achieve a minimum annual Availability",
     "Where the Client has selected Tier C services, the Service Provider guarantees that the BESS "
     "at the Site shall achieve a minimum annual Availability of ninety-seven percent (97%), "
     "calculated in accordance with Schedule 4."),
    ("Availability shall be calculated on a Group Aggregate basis across all parks",
     "Availability shall be calculated for the BESS at the Site, as follows:"),
    ("Group Availability (%) = [SUM(Available Hours per Park)",
     "Availability (%) = [Available Hours / Total Hours] x 100"),
    ("Total Hours per Park = 8,760 hours per year",
     "(a) Total Hours = 8,760 hours per year (or 8,784 in a leap year), commencing from the PAC date;"),
    ("Unavailable Hours per Park = hours during which that park's BESS",
     "(b) Unavailable Hours = hours during which the BESS is unable to operate at its rated capacity "
     "due to equipment fault, failure, or Service Provider-caused downtime;"),
    ("Parks are added to the Group Availability calculation from their respective PAC dates",
     "(c) The Availability Guarantee commences from the PAC date for the Site;"),
    ("For partial years (the year of commissioning), only the period from PAC to year-end is included",
     "(d) For the partial year of commissioning, only the period from PAC to year-end is included in "
     "the calculation."),
    ("9.2A Group-Level Rationale", "9.2A Availability Basis"),
    ("The 97% Availability Guarantee applies at the group portfolio level",
     "The 97% Availability Guarantee applies to the BESS at the Site, measured annually in "
     "accordance with Schedule 4. An isolated component outage does not constitute a breach provided "
     "the annual Availability of the Site exceeds 97% after the excluded hours in Schedule 4."),
    ("Availability shall be calculated on a Group Aggregate basis",
     "Availability shall be calculated for the BESS at the Site."),
    ("Rate varies by park size and group allocation — portfolio range",
     "For this Project the EMS/SCADA software subscription is EUR 400/MWh/Year (EUR 8,024 per year "
     "for 20.06 MWh), billed annually in advance from PAC, as set out in the EMS Subscription "
     "Addendum. Calculation basis: DISPERON software subscription (cloud platform, updates, external "
     "data feeds, cybersecurity/NIS2, support)."),
]


def _replace_in_paragraph(p, old, new):
    joined = "".join(r.text for r in p.runs)
    if old in joined:
        full = joined.replace(old, new)
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = full
        else:
            p.add_run(full)


def _rewrite_paragraph(p, full_new):
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = full_new
    else:
        p.add_run(full_new)


def main():
    if not SRC.exists():
        print(f"!! SOURCE NOT FOUND: {SRC}")
        return
    doc = Document(str(SRC))

    for p in doc.paragraphs:
        # Full-paragraph rewrites first (group/portfolio framing)
        done = False
        for marker, new in REWRITES:
            if marker in p.text:
                _rewrite_paragraph(p, new)
                done = True
                break
        if done:
            continue
        for old, new in SUBS:
            _replace_in_paragraph(p, old, new)

    PKG.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"LTSA -> {OUT}")


if __name__ == "__main__":
    main()
