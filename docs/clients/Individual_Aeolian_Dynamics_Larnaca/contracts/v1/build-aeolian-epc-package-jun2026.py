#!/usr/bin/env python3
"""Build the Aeolian Dynamics BESS EPC client package (v1, June 2026).

Single park: T.P. Aeolian Dynamics Ltd — Agia Anna Wind Farm Hybrid
6.5 MW / 20.06 MWh, Larnaca, Cyprus.

Generates:
  00-Cover-Note-and-Index.docx
  01-EPC-Agreement-Aeolian-6.5MW-20MWh-v1-jun2026.docx
  02-Ownership-and-Guarantee-Flow.html       (copied from template)
  03-LTSA-Aeolian-Agia-Anna.docx             (TBC)
  04-OEM-Direct-Warranty-Undertaking.docx
  05-Advance-Payment-Guarantee-Specimen.docx
  06-OEM-5pct-Performance-Guarantee.docx

Run:
  python docs/clients/Individual_Aeolian_Dynamics_Larnaca/contracts/v1/build-aeolian-epc-package-jun2026.py
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
GREY = RGBColor(0x40, 0x40, 0x40)
AMBER = RGBColor(0x78, 0x35, 0x0F)

V1 = Path(__file__).resolve().parent
AEOLIAN_DIR = V1.parents[1]              # .../Individual_Aeolian_Dynamics_Larnaca
REPO = V1.parents[4]                     # .../solinvest
GALASCOPE_V6 = (
    REPO / "docs" / "clients" / "group-order"
    / "Group2_Esperia_Energy" / "contracts" / "v6"
)
PKG = V1 / "CLIENT-PACKAGE"

EPC_OUT = PKG / "01-EPC-Agreement-Aeolian-6.5MW-20MWh-v1-jun2026.docx"

# ── Project constants ─────────────────────────────────────────────────────────
CONTRACT_REF     = "LCY-EPC-AEO-2026"
CLIENT_NAME      = "T.P. Aeolian Dynamics Ltd"
CLIENT_HE        = "HE 168239"
CLIENT_ADDRESS   = "Karaiskaki 13, 3032 Limassol, Cyprus"
CLIENT_SHORT     = "Aeolian"
SITE             = "Agia Anna Wind Farm, Larnaca, Cyprus"
MW               = "6.68 MW"              # T8 config: 8 × BCS1000K (8 MW hardware), EMS-limited to 6.68 MW
MWH              = "20.06 MWh"            # 4 × 5.015 MWh containers (unchanged)
DURATION         = "3.0 hours"            # 20.06 MWh ÷ 6.68 MW
CONTAINERS       = "4"                    # 4 × 5.015 MWh battery containers
PCS_IN_SKID      = "8"                    # 8 × BCS1000K (1.0 MW each) in one T8 MV skid = 8 MW hardware
MV_SKIDS         = "1"                    # 1 × T8 MV Skid (one container housing 8 × BCS1000K = 8 MW hardware)
TOTAL_UNITS      = "5"                    # 4 BESS containers + 1 T8 MV Skid
CONTRACT_PRICE   = "EUR [●]"
COMP_A           = "EUR [●]"              # Component A — equipment supply (to be confirmed)
COMP_B           = "EUR [●]"              # Component B — EPC services (to be confirmed)
APG_AMOUNT       = "EUR [●]"              # 25% of Component A (advance APG)
APG_PRESHIP      = "EUR [●]"              # 50% of Component A (ready-to-ship APG)
APG_COMBINED     = "EUR [●]"              # Combined APG cover (75% of Component A)
PG_AMOUNT        = "EUR [●]"              # 5% of Component A (performance guarantee)
ADVANCE_PCT      = "25%"
ADVANCE_EUR      = "EUR [●]"
PRESHIP_PCT      = "50%"
PRESHIP_EUR      = "EUR [●]"
DELIVERY_PCT     = "20%"
DELIVERY_EUR     = "EUR [●]"
SAT_PCT          = "5%"
SAT_EUR          = "EUR [●]"
EMS_RATE         = "EUR [●]/MWh/yr"
EMS_ANNUAL       = "EUR [●]"
# ── Voltus EMS upfront (5-unit / 20 MWh park, quoted Feb 2026) ──
EMS_HW_INSTALL   = 47_513   # System+Hardware €23,119 + Remote config €14,232 + On-site install €10,162
SCADA_LOCAL      = 30_000   # SCADA Local hardware + config (€30K for 5-unit park per Voltus quote)
SCADA_GLOBAL     = 60_000   # SCADA Group platform (first/only standalone park)
EMS_UPFRONT      = EMS_HW_INSTALL + SCADA_LOCAL + SCADA_GLOBAL  # = €137,513
# ── Kouklis civil quote (Cyprus, Jun 2026) ──
CIVIL_BASE1_EUR  = 21_455   # Base 1: 4×20ft BESS + 1×40ft T8 MV skid (actual Cyprus quote)


def _run(p, text, *, bold=False, italic=False, size=10.5, color=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return r


def h(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, bold=True, size=size, color=GOLD)
    return p


def para(doc, text, *, bold=False, italic=False, color=None, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, text, bold=bold, italic=italic, color=color, size=size)
    return p


def body(doc, text, bold=False, italic=False, grey=False, amber=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if grey:
        r.font.color.rgb = GREY
    if amber:
        r.font.color.rgb = AMBER
    return p


def clause(doc, num, text):
    body(doc, f"{num}  {text}")


# ── Annex V1 amendment clauses ────────────────────────────────────────────────
ANNEX = [
    ("1A.4  Grant Scheme — Recital.",
     "The Client is a participant in the Θ.Α.ΛΕ.Ι.Α 2021–2027 Just Transition Fund grant scheme "
     "(the 'Grant Scheme'). The parties note that the BESS configuration under this Agreement "
     "(6.68 MW / 20.06 MWh at the point of connection) has been designed to comply with the tender "
     "specification of the Grant Scheme. The Client is solely responsible for all grant "
     "administration, compliance obligations, reporting and grant-authority liaison. The Contractor's "
     "obligation is limited to supplying and commissioning the BESS system in accordance with this "
     "Agreement and providing as-built technical documentation sufficient to support the Client's "
     "grant compliance filings."),
    ("1A.5  Nature of EPC Execution.",
     "Execution of this Agreement records the Parties' agreed terms. Unlike a project with a pending "
     "grid connection, this is a wind-farm hybrid where the 22 kV export line to PSEUDAS S/S is "
     "already established. The advance under Section 7.1(a) falls due within seven (7) days of the "
     "Effective Date of this Agreement."),
    ("1A.6  Companion Documents — Condition Precedent to the Advance.",
     "The advance shall not become due until the Contractor has presented to the Client, each "
     "substantially in the form attached: (a) the OEM Direct Warranty Undertaking, signed and "
     "sealed by the OEM, covering the Project; (b) the Advance Payment Guarantee (APG No. 1) issued "
     "by the OEM's bank naming the Client as beneficiary; and (c) the OEM 5% performance guarantee. "
     "The companion-document drafts attached are provided for review and are not executed by "
     "signature of this Agreement."),
    ("1A.7  Bilateral Walk-Away.",
     "If any document in Section 1A.6 is not presented in the agreed form within thirty (30) days of "
     "the Effective Date, either Party may elect not to proceed by written notice without liability, "
     "and any advance received shall be refunded within thirty (30) days."),
    ("1A.8  Pre-Shipment Payment — Condition Precedent.",
     "The pre-shipment / ready-to-ship payment (Section 7.1(b)) shall not become due until: (a) the "
     "equipment has passed the Factory Acceptance Test (Section 8.x) witnessed by or on behalf of "
     "the Client; and (b) the Contractor has presented the Pre-Shipment Advance Payment Guarantee "
     "(APG No. 2) issued by the OEM's bank naming the Client as beneficiary. The Client never funds "
     "an equipment prepayment tranche without a guarantee in place for that equipment prepayment."),
    ("2A.1  Site Access — Army Firing Range.",
     "The Site is located adjacent to an army firing range at Agia Anna, Larnaca. The Client is "
     "responsible for coordinating access windows with the relevant military authority (Republic of "
     "Cyprus National Guard) in advance of all installation and commissioning activities. The Client "
     "shall notify the Contractor of confirmed access windows at least ten (10) Business Days before "
     "each visit. Where access is refused or delayed by the military authority, the Target PAC Date "
     "extends day-for-day for the duration of the access refusal, and no Delay Liquidated Damages "
     "accrue for that period. The Contractor shall not be entitled to additional cost reimbursement "
     "for reasonable standby time directly caused by such access delay."),
    ("4A.1  Wind-Farm Hybrid Integration.",
     "The BESS is to be integrated with the existing 10.8 MW wind farm (6 × Vestas V100-1.8 MW "
     "turbines) at the Agia Anna site. The BESS configuration is: 4 × 5.015 MWh battery containers "
     "connected to a single T8 MV Skid (one container housing 8 × Kehua BCS1000K PCS, 1.0 MW each = "
     "8 MW hardware, EMS-limited to 6.68 MW), providing 6.68 MW / 20.06 MWh (3-hour duration at rated "
     "export power). The EMS shall enforce the combined wind + BESS export limit on the 22 kV line "
     "to PSEUDAS S/S (confirmed overhead line capacity: 14 MW) via an IEC 60870-5-104 interface with "
     "the existing ABB wind farm controller, capping BESS output at 6.68 MW. The Contractor shall "
     "ensure the DISPERON EMS is configured to respect this export limit as a Condition of PAC."),
    ("4A.2  Scope — MV Switchgear Bay Extension (included in Contractor's scope).",
     "Unlike a greenfield site, the BESS connects into the Client's existing 22 kV wind-farm MV "
     "switchgear. The Contractor's Component B scope INCLUDES the design, supply and installation of "
     "a new switchgear bay (bay J04, ABB ZX1.2-compatible MV cubicle) extending the existing "
     "switchboard, and the MV termination and connection of the T8 MV Skid into that bay. This is a "
     "point of difference from standalone BESS projects and is included in the Contract Price. The "
     "Client shall provide safe access to, and accurate as-built drawings and protection/selectivity "
     "data for, the existing switchgear; any latent defect in the existing switchboard not caused by "
     "the Contractor is outside the Contractor's scope."),
    ("6.1  Contract Price — fixed lump sum (replaces body Section 6.1 in its entirety).",
     "The Contract Price is a fixed lump sum (Component A + Component B, as set out in Schedule A), "
     "exclusive of VAT, confirmed at signing on the basis of the OEM's then-current quotation for the "
     "Project configuration. Because the Client's site is already grid-connected, there is NO grid "
     "Connection Terms trigger, NO Confirmed Price Certificate, and NO Indicative/Confirmed price "
     "distinction or raw-material / FX adjustment: the entire Indicative-and-Confirmed price "
     "mechanism in body Section 6.1(a)-(f) (and any Connection-Terms price trigger) does not apply to "
     "this Agreement. After signing, the Contract Price changes only by a written change order agreed "
     "by both Parties (for example, a Client-requested configuration change)."),
    ("7.1  Payment Milestones (replaces body 7.1).",
     "All payments are calculated on the Contract Price (ex VAT). VAT at the applicable "
     "Cyprus rate (currently 19%) is added to each invoice and paid by the Client. "
     f"(a) {ADVANCE_PCT} advance ({ADVANCE_EUR}) — within seven (7) days of the Effective Date, "
     "subject to Section 1A.6 (companion-document CP). "
     f"(b) {PRESHIP_PCT} ready-to-ship ({PRESHIP_EUR}) — on Factory Acceptance Test pass + OEM "
     "written confirmation of readiness for shipment. "
     f"(c) {DELIVERY_PCT} on-site delivery ({DELIVERY_EUR}) — on arrival and unloading of equipment "
     "at the Site. "
     f"(d) {SAT_PCT} System Acceptance Test / commissioning ({SAT_EUR}) — on PAC and issuance of "
     "the Provisional Acceptance Certificate."),
    ("7.4 / 7A  No Cash Retention (replaces body Retention provisions).",
     "This Agreement does not operate a cash retention. The final five percent (5%) is the System "
     "Acceptance Test / commissioning payment, due on PAC (Section 7.1(d)). Accordingly, any "
     "references in the body to a 'Retention' held to the end of the Defects Liability Period, and "
     "Section 7.4, do not apply to this Agreement. Defect rectification during the three (3) month "
     "Defects Liability Period is secured by the OEM 5% performance guarantee, the OEM Direct "
     "Warranty Undertaking, and the Contractor's EPC works warranty."),
    ("8 / Schedule A  Delivery & Shipping Schedule — confirmed at signing, not indicative.",
     "Manufacturing, shipping and delivery lead time is NOT fixed at this stage and no time "
     "commitment is given at signing. Production lead time is indicative only and may change from "
     "any manufacturing estimate previously discussed. Upon order release (triggered by the advance "
     "under Section 7.1(a)), the Contractor shall issue a binding Delivery Schedule stating the "
     "confirmed manufacturing lead time, the shipping timeline and the Target PAC Date. All "
     "time-based obligations and Delay Liquidated Damages under Section 8.4 are calculated from that "
     "confirmed Delivery Schedule."),
    ("8.x  Pre-Shipment Inspection / Factory Acceptance Test (FAT).",
     "Before any equipment is despatched and before the ready-to-ship payment (Section 7.1(b)) is "
     "paid, the Client and/or its appointed third-party inspector may inspect and witness factory "
     "acceptance testing at the OEM's facility on ten (10) Business Days' notice. Where the FAT "
     "reveals defects or non-conformity, the OEM/Contractor shall remedy them at no additional cost "
     "before shipment, and the Client may re-inspect. Ready-to-ship payment is conditional on a "
     "passed FAT."),
    ("9.x  Extended rejection window.",
     "The Client's right to inspect and to give notice of defect or non-conformity extends to PAC. "
     "For non-conformity affecting ten percent (10%) or more of installed capacity discovered up to "
     "PAC, the Client may require repair or replacement, or reject the affected equipment and require "
     "refund of the corresponding price (recoverable under the APG per Section 10.9(e))."),
    ("10.9(b)  Advance Payment Guarantee No. 1 — advance tranche.",
     "The Contractor shall procure that the equipment manufacturer's bank (Bank of Communications) "
     "issues a first-demand Advance Payment Guarantee under URDG 758 naming "
     f"{CLIENT_NAME} ({CLIENT_HE}) and/or its project-finance security agent as beneficiary, "
     f"securing the advance for the equipment supply. The Guaranteed Amount is {APG_AMOUNT} "
     "(25% of the equipment supply value), as set out in Schedule A."),
    ("10.9(c)  Advance Payment Guarantee No. 2 — ready-to-ship tranche.",
     "Before the ready-to-ship payment falls due, and conditional on a passed Factory Acceptance "
     "Test, the Contractor shall procure that the manufacturer's bank issues a second first-demand "
     "Advance Payment Guarantee (or an increase to APG No. 1) under URDG 758 naming the Client "
     f"and/or its security agent as beneficiary. The Guaranteed Amount is {APG_PRESHIP} "
     f"(50% of the equipment supply value). Combined APG cover: {APG_COMBINED}."),
    ("10.9(d)  APG validity — to PAC.",
     "Each APG shall remain valid until the earlier of: (i) issuance of the Provisional Acceptance "
     "Certificate (PAC); or (ii) twelve (12) months after delivery of the equipment to Site. Upon "
     "issuance of PAC the APGs are released."),
    ("10.9(aa)  APG issuance mechanism.",
     "The OEM's bank issues the APG(s) upon being provided with both this EPC Agreement and the OEM "
     "Sales Contract, and issues them with the Client named as beneficiary from the outset. The "
     "Client may therefore demand directly under the APG(s) notwithstanding that payments are made "
     "to the Contractor: it is original issuance in the Client's name, not a transfer."),
    ("10.9A  Security layering alongside the APGs.",
     f"Through to PAC the Client is protected by: (a) the two APGs ({APG_COMBINED} combined, to PAC "
     "/ 12 months after delivery); (b) CAR / erection all-risks insurance for the full replacement "
     "value from port of discharge through installation to PAC (the sea voyage is covered by the "
     "OEM's CIF marine insurance); (c) the OEM 5% performance guarantee to DLP end; (d) the "
     "retention (Section 7.1(d)); and (e) the OEM Direct Warranty Undertaking and 5-year product "
     "warranty."),
    ("13.2  Limitation of Liability — replaces body 13.2.",
     "The Contractor's aggregate liability shall be: (a) for claims relating to the Contractor's own "
     "EPC services, installation and non-OEM works: ten percent (10%) of the Contract Price; (b) for "
     "all other contractual breaches (excluding (a), (c), fraud and wilful misconduct): fifty percent "
     "(50%) of the Contract Price; (c) for manufacturing defects in OEM equipment and for fraud or "
     "wilful misconduct: uncapped, with the Contractor's right to pursue OEM recovery."),
    ("13.5  Manufacturing-defect carve-out.",
     "For confirmed manufacturing defects in OEM equipment, the OEM bears the full cost of repair or "
     "replacement including shipping to CIF Limassol, and the OEM's standard 10% warranty-liability "
     "cap does not apply. The Contractor shall procure this carve-out from the OEM and reflect it in "
     "the OEM Direct Warranty Undertaking provided to the Client."),
    ("4.4 / 3.1  EMS provider.",
     "EMS integration, SCADA commissioning and the DISPERON software subscription are provided by "
     "R&D Innovations Sp. z o.o. (NIP 9492265995; trading as DISPERON), under a separate EMS "
     "Subscription Addendum. The Contractor guarantees that provider's performance of its obligations "
     "under that document."),
    ("19.x  Grid-Forming (VSG) & Black Start.",
     "The PCS supplied (Kehua C-series) is hardware-capable of grid-forming (VSG) and black-start "
     "operation via firmware. Activation and any OEM firmware licence fee shall be confirmed by "
     "amendment following the technical review and confirmation of the applicable DSO grid-code "
     "requirement. Hardware is unaffected and the equipment price is unchanged."),
]


def _replace_in_paragraph(p, old, new):
    if old in "".join(r.text for r in p.runs):
        full = "".join(r.text for r in p.runs).replace(old, new)
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = full
        else:
            p.add_run(full)


def build_epc():
    # Base off Galascope v5.1 source — update all party/config references
    src = GALASCOPE_V6.parent / "EPC-Galascope-Esperia-batch1-may2026.docx"
    if not src.exists():
        print(f"!! SOURCE NOT FOUND: {src}")
        return
    doc = Document(str(src))

    SUBS = [
        # Version stamp
        ("Version: 5.1",    "Version: 1.0"),
        ("Date: May 2026",  "Date: June 2026"),
        # Ref number
        ("LCY-EPC-GAL-B1-2026", CONTRACT_REF),
        # Client party block
        ("Galascope Ltd, a company incorporated under the laws of the Republic of Cyprus, "
         "with registered office at Karaiskaki 6, City House, 3032 Limassol, Cyprus, "
         "Registration No. HE 303759",
         f"{CLIENT_NAME}, a company incorporated under the laws of the Republic of Cyprus, "
         f"with registered office at {CLIENT_ADDRESS}, Company No. {CLIENT_HE}"),
        ("Galascope Ltd (HE 303759)", f"{CLIENT_NAME} ({CLIENT_HE})"),
        ('("Client")',      '("Client")'),
        # T8 config: update any T4/T1 skid references from the source template
        ("6.25 MW / 20.06 MWh",
         f"{MW} / {MWH} (T8 MV Skid, 3-hour system)"),
        # Site
        ("Famagusta, Cyprus", SITE),
        ("Galascope 1 (5 MW / 20.06 MWh) & Galascope 2 (2.5 MW / 10.03 MWh), Famagusta, Cyprus",
         f"Agia Anna Wind Farm Hybrid ({MW} / {MWH}), {SITE}"),
        ("Galascope 1 and Galascope 2", "Agia Anna BESS"),
        ("Galascope 1 (5 MW / 20.06 MWh)", f"Agia Anna BESS ({MW} / {MWH})"),
        ("Galascope 2 (2.5 MW / 10.03 MWh)", ""),
        # Price
        ("EUR 3,444,300", CONTRACT_PRICE),
        ("EUR 2,238,000", COMP_A),
        ("EUR 1,206,300", ""),
        ("EUR 1,848,712.43", COMP_A),
        ("EUR 974,457.00",   ""),
        ("EUR 389,287.57",   COMP_B),
        ("EUR 231,843.00",   ""),
        # Payment milestones (body — overridden by Annex 7.1 above)
        ("Advance (30%):",     f"Advance ({ADVANCE_PCT}):"),
        ("EUR 1,033,290.00",   ADVANCE_EUR),
        ("Pre-Ship (55%):",    f"Ready-to-Ship ({PRESHIP_PCT}):"),
        ("EUR 1,894,365.00",   PRESHIP_EUR),
        ("PAC (10%):",         f"On-Site Delivery ({DELIVERY_PCT}):"),
        ("EUR 344,430.00",     DELIVERY_EUR),
        ("Retention (5%):",    f"SAT / Commissioning ({SAT_PCT}):"),
        ("EUR 172,215.00",     SAT_EUR),
        # APG
        ("No. 1 advance EUR 705,792 + No. 2 pre-shipment EUR 1,411,585 (combined batch)",
         f"No. 1 advance {APG_AMOUNT} + No. 2 ready-to-ship {APG_PRESHIP} (combined)"),
        ("see combined batch APG figure above (Annex V6)", f"see combined APG above ({APG_COMBINED})"),
        ("EUR 92,435.62",  PG_AMOUNT),
        ("EUR 48,722.85",  ""),
        # EMS entity
        ("Lighthief EU BESS Ltd",
         "R&D Innovations Sp. z o.o. (trading as DISPERON, under Lighthief International Ltd)"),
        # Delivery dates
        ("31 January 2027",
         "the date in the Delivery Schedule confirmed at order release"),
        (" (January 2032)", ""),
        (" (approx. April 2027)", ""),
        # Price baseline
        ("Linyang Quotation LY202601271 (January 2026)",
         "the January 2026 quotation basis"),
        ("155,000 CNY/tonne (Mysteel China battery-grade spot, January 2026 monthly average)",
         "as recorded in the Indicative Price Basis Certificate at EPC signing"),
        ("8.18 CNY per EUR (January 2026 average)",
         "as recorded in the Indicative Price Basis Certificate at EPC signing"),
        # VAT
        ("7.2 Payments exclusive of VAT.",
         "7.2 The Contract Price and all milestone amounts are stated exclusive of VAT. "
         "VAT at the applicable Cyprus rate (currently nineteen percent (19%)) shall be added "
         "to each invoice and paid by the Client, and is recoverable by the Client in "
         "accordance with applicable law."),
        # Title
        ("(b) Title (ownership) passes to the Client upon issuance of PAC and receipt of the "
         "PAC payment under Section 7.1(c). The Retention does not defer title.",
         "(b) Title (ownership) passes to the Client on the earlier of: (i) issuance of PAC and "
         "receipt of the SAT payment under Section 7.1(d); or (ii) twelve (12) months after "
         "delivery of the equipment to Site, so that title passes no later than the end of APG cover. "
         "Where title passes under (ii) before PAC, the Contractor retains a security interest in "
         "the Equipment for any unpaid balance until the SAT payment is made."),
        # Liability
        ("(a) Warranty claims (defects in materials, workmanship, or OEM equipment): "
         "ten percent (10%) of the Contract Price;",
         "(a) Claims relating to the Contractor's own EPC services, installation and non-OEM works "
         "(excluding OEM equipment): ten percent (10%) of the Contract Price;"),
        # Galascope-specific refs no longer applicable
        ("Galascope 1 — 20.06 MWh; Galascope 2 — 10.03 MWh",
         f"Agia Anna BESS — {MWH} ({MW})"),
        ("Park Name:              Galascope 1", "Park Name:              Agia Anna BESS"),
        ("Park Name:              Galascope 2", ""),
        ('For and on behalf of Galascope Ltd ("Client"):',
         f'For and on behalf of {CLIENT_NAME} ("Client"):'),
        ("Galascope Ltd", CLIENT_NAME),
        # ── Body ↔ Annex reconciliation (Aeolian first draft) ──
        # Trigger: wind farm already grid-connected; advance on Effective Date (not Connection Terms)
        ("1A.1 This Agreement is conditional upon the Client obtaining grid connection terms from "
         "EAC/DSO for the Site (\u201cConnection Terms\u201d) within twelve (12) months of the Effective "
         "Date (\u201cLong-Stop Date\u201d).",
         "1A.1 The Client's Agia Anna wind farm is already grid-connected via the existing 22 kV line "
         "to PSEUDAS S/S. This Agreement is not conditional on new grid Connection Terms. The "
         "\u201cEffective Date\u201d is the date of signing."),
        ("1A.2 If Connection Terms are not obtained by the Long-Stop Date, either Party may terminate "
         "this Agreement by written notice without liability. The Contractor shall refund the advance "
         "payment within thirty (30) days, less reasonable and documented costs incurred.",
         "1A.2 Companion Documents (condition precedent to the advance): the advance shall not become "
         "due until the Contractor has presented to the Client, each substantially in the form "
         "attached for review: (a) the OEM Direct Warranty Undertaking signed and sealed by the OEM; "
         "(b) the Advance Payment Guarantee (APG No. 1) issued by the OEM's bank naming the Client as "
         "beneficiary; and (c) the OEM 5% performance guarantee. If any is not presented in the agreed "
         "form within thirty (30) days of the Effective Date, either Party may terminate this "
         "Agreement by written notice without liability, and any advance received shall be refunded "
         "within thirty (30) days less reasonable and documented costs incurred."),
        ("1A.3 The advance payment under Section 7.1(a) shall become due within thirty (30) days of "
         "the later of: (i) the Effective Date; and (ii) receipt of Connection Terms.",
         "1A.3 The advance payment under Section 7.1(a) shall become due within seven (7) days of the "
         "Effective Date, subject to presentation of the companion documents (Section 1A.2)."),
        # Drop the now-inapplicable "Indicative" pricing label (fixed price)
        ("6.2 Indicative Component A — Equipment Supply Price", "6.2 Component A — Equipment Supply Price"),
        ("6.3 Indicative Component B — EPC Services Price", "6.3 Component B — EPC Services Price"),
        ("Indicative Component A (Equipment CIF):", "Component A (Equipment CIF):"),
        ("Indicative Component B (EPC Services):", "Component B (EPC Services):"),
        ("INDICATIVE CONTRACT PRICE (A+B, ex VAT):", "CONTRACT PRICE (A+B, ex VAT):"),
        ("TOTAL INDICATIVE CONTRACT PRICE (ex VAT):", "TOTAL CONTRACT PRICE (ex VAT):"),
        # Payment milestones body: 30/55/10/5 → 25/50/20/5
        ("(c) 10% payment upon issuance of PAC;",
         "(c) 20% payment upon arrival and unloading of the equipment at the Site;"),
        ("(d) 5% Retention released at the end of the Defects Liability Period (three (3) months "
         "after PAC), subject to Section 7.4.",
         "(d) 5% System Acceptance Test (SAT) / commissioning payment on PAC. No separate cash "
         "retention applies under this Agreement (Section 7.4); defect rectification "
         "during the three (3) month Defects Liability Period is secured by the OEM 5% performance "
         "guarantee and the Direct Warranty Undertaking."),
        # APG body cover: single 70% → two APGs (25%/50% of equipment supply)
        ("(b) The APG covers: 20% advance on Component A + 50% pre-shipment on Component A = 70% of "
         "the Component A Equipment Supply Price.",
         "(b) Two Advance Payment Guarantees are provided, each issued by the OEM's bank naming the "
         "Client as beneficiary: APG No. 1 securing the equipment advance (25% of the equipment "
         "supply value) and APG No. 2, issued at the Factory Acceptance Test, securing the "
         "ready-to-ship tranche (50% of the equipment supply value)."),
        ("(d) The APG shall remain valid until equipment delivery to Site plus thirty (30) days.",
         "(d) Each APG shall remain valid until the earlier of issuance of PAC or twelve (12) months "
         "after delivery of the equipment to Site; upon PAC the APGs are released."),
        # OEM DWU governing law: align EPC body summary to the DWU (PRC / SHIAC)
        ("(d) Governing law: Singapore (SIAC arbitration, Cyprus carve-out for interim relief).",
         "(d) Governing law: People's Republic of China (SHIAC arbitration, in English; Cyprus "
         "carve-out for interim relief), consistent with the OEM Direct Warranty Undertaking."),
        # Civil scope: match Aeolian offer (foundations + BESS-area trenching; fencing/roads excluded)
        ("(e) Civil works, foundations, fencing, and site preparation (Component B);",
         "(e) Civil works: BESS equipment foundations and cable trenching within the BESS area "
         "(Component B). Site levelling / general earthworks, access roads, perimeter fencing and "
         "security systems are excluded (Client scope, per Schedule A);"),
        ("(b) Civil works, foundations, fencing, and site preparation;",
         "(b) Civil works: BESS equipment foundations and cable trenching within the BESS area "
         "(site levelling, access roads, fencing and security excluded — Client scope);"),
        # Cabling + switchgear bay extension (point of difference vs Esperia)
        ("(f) AC/DC cabling, grounding, and auxiliary electrical systems (Component B);",
         "(f) AC/DC, MV and communications cabling, grounding and auxiliary electrical systems, and "
         "extension of the Client's existing 22 kV wind-farm switchgear by a new bay (bay J04, "
         "ABB ZX1.2-compatible MV cubicle) for the BESS connection (Component B);"),
        ("(c) AC/DC cabling, grounding, and auxiliary electrical systems;",
         "(c) AC/DC, MV and communications cabling, grounding, auxiliary electrical systems, and "
         "extension of the existing 22 kV switchgear (new bay J04, ABB ZX1.2) for BESS connection;"),
        # Strip internal OEM contact line (must not appear in client-facing contract)
        ("Linyang contacts: Conor Yang / Kamil Tyburski / Tomasz Wieckowski", ""),
        ("Linyang contacts:", ""),
        # Schedule A Part 1 — fix stale district (Aeolian is Larnaca, not Famagusta)
        ("Site / District:        Famagusta District, Cyprus",
         "Site / District:        Larnaca District (Agia Anna), Cyprus"),
        # Schedule A Part 1 — T8 skid + PCS (park 1 block; the 5.0 MW / T4 lines are stale)
        ("MV Skid:                1 x T4 MV Skid (4 x BCS1250K = 5 MW, SL-5000)",
         "MV Skid:                1 x T8 MV Skid (8 x BCS1000K = 8 MW hardware, EMS-limited to 6.68 MW; transformer per T8 datasheet)"),
        ("PCS:                    4 x Kehua BCS1250K-C-HUD = 5.0 MW",
         "PCS:                    8 x Kehua BCS1000K-C-HUD (1.0 MW each, all in one T8 MV skid) = 8 MW hardware, EMS-limited to 6.68 MW"),
        # EPC body PCS spec line (source uses BCS1250K / 1.25 MW)
        ("PCS: Kehua BCS1250K-C-HUD (1.25 MW); \u226598% efficiency; EN 50549-2 (T\u00dcV D 115067 0077)",
         "PCS: Kehua BCS1000K-C-HUD (1.0 MW); \u226598% efficiency; EN 50549-2 (T\u00dcV D 115067 0077)"),
        # Schedule A Part 1 — blank stale APG figures + relabel to two-APG structure
        ("APG (70% \u00d7 Indicative Component A):", "APG (25% + 50% of Component A, two APGs):"),
        ("EUR 1,294,098.70", "EUR [\u25cf]"),
        ("EUR 682,119.90", "EUR [\u25cf]"),
        ("Performance Bond (5% \u00d7 Component A):", "Performance Guarantee (5% \u00d7 Component A):"),
        # Schedule A Part 2 — fix stale milestone trigger descriptions (25/50/20/5)
        ("Within 30 days of payment trigger (\u00a71A.3)",
         "Within 7 days of the Effective Date (\u00a71A.3)"),
        ("System commissioned, grid-connected",
         "On arrival and unloading of equipment at Site"),
        ("Released after 3-month DLP",
         "On PAC (system commissioned & grid-connected)"),
        # Schedule A Part 3 — remove stray fixed dates / retention line
        ("CIF Limassol:           August / September 2026",
         "CIF Limassol:           confirmed in the Delivery Schedule at order release"),
        ("August / September 2026",
         "confirmed in the Delivery Schedule at order release"),
        ("Retention Release:      On DLP expiry",
         "Retention Release:      Not applicable — no cash retention (Section 7.4)"),
        # Strip stray fixed date estimates from Schedule A summary (timing set at order release)
        ("Production Start:       Q2 2026",
         "Production Start:       confirmed in the Delivery Schedule at order release"),
        ("Factory Acceptance:     Q3 2026",
         "Factory Acceptance:     confirmed in the Delivery Schedule at order release"),
        # Correct stale BESS capacity line (T8 config)
        ("BESS Capacity:          20 MWh / 5.0 MW",
         "BESS Capacity:          20.06 MWh / 6.68 MW (T8 MV Skid, 3-hour)"),
        # Body §7.1(a)/(b) prose still showed Galascope 30/55 — correct to 25/50 (else body sums to 110%)
        ("(a) 30% advance payment within seven (7) days of the payment trigger under Section 1A.3;",
         "(a) 25% advance payment within seven (7) days of the payment trigger under Section 1A.3;"),
        ("(b) 55% payment prior to shipment of BESS equipment, upon Contractor\u2019s written confirmation "
         "of successful Factory Acceptance Testing (FAT) and photographic evidence of readiness;",
         "(b) 50% payment prior to shipment of BESS equipment, upon Contractor\u2019s written confirmation "
         "of successful Factory Acceptance Testing (FAT) and photographic evidence of readiness;"),
        # Wrong signatory — Ntinos Konstantinos is Galascope's director; Aeolian signatory is Sotiris Shiacallis
        ("Ntinos Konstantinos", "Sotiris Shiacallis"),
        # Galascope LTSA reference → Aeolian
        ("LCY-LTSA-GAL-2026", "LCY-LTSA-AEO-2026"),
        # Confirmed-Price mechanism does not apply (no Connection Terms) — strip Schedule A references
        ("PAYMENT MILESTONES (applied to Confirmed Contract Price):",
         "PAYMENT MILESTONES (applied to the Contract Price):"),
        ("6. Confirmed Price Certificate (issued within 14 days of Connection Terms receipt)", ""),
        # Annex 10.9A(d): no cash retention — reference the SAT payment, not "retention"
        ("(d) the retention (Section 7.1(d));",
         "(d) the 5% SAT / commissioning payment (Section 7.1(d));"),
        # Annex header reference
        ("ANNEX V6 \u2014 AMENDMENT SCHEDULE (JUNE 2026)",
         "ANNEX V1 \u2014 AMENDMENT SCHEDULE (JUNE 2026)"),
        ("Annex V6",  "Annex V1"),
        ("v6.0",      "v1.0"),
    ]

    def _rewrite_paragraph(p, full_new):
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = full_new
        else:
            p.add_run(full_new)

    for p in doc.paragraphs:
        for old, new in SUBS:
            if old:
                _replace_in_paragraph(p, old, new)

    # Schedule A Part 4 — the Indicative/Confirmed price-basis mechanism does not apply to Aeolian
    # (fixed price, no Connection Terms). Rewrite the block to a clean fixed-price statement.
    _p4 = [
        ("PART 4 — PRICE BASIS REFERENCE", "PART 4 — PRICE BASIS"),
        ("Quotation Source:",
         "Price Type:             Fixed lump sum (Component A + B, Schedule A), confirmed at signing"),
        ("Lithium Carbonate Index:",
         "OEM Quotation:          Linyang T8 quotation current at signing"),
        ("EUR/CNY Reference:",
         "Confirmed Price Cert.:  Not applicable — no Connection-Terms price trigger"),
        ("Reference Date:",
         "Raw-material / FX adj.: Not applicable to this Agreement"),
        ("Adjustment Triggers",
         "Post-signing changes:   Only by written change order agreed by both Parties"),
        ("Lithium carbonate index movement", ""),
        ("EUR/CNY rate movement", ""),
        ("Configuration change requested by Client", ""),
        ("Adjustment Cap:", ""),
        ("Walk-away Right:", ""),
    ]
    # Scope to Schedule A Part 4 only (else markers like "Adjustment Cap:" also hit body §6.1)
    _all = doc.paragraphs
    _p4_start = next((i for i, p in enumerate(_all)
                      if "PART 4" in p.text and "PRICE BASIS" in p.text), None)
    if _p4_start is not None:
        for p in _all[_p4_start:]:
            s = p.text.strip()
            for marker, new in _p4:
                if marker in s:
                    _rewrite_paragraph(p, new)
                    break

    # Remove the leftover SECOND park block in Schedule A Part 1 (Aeolian is a single park).
    # The block sits between the first "----" separator that follows "PART 1 — PARK DETAILS"
    # and the next "====" separator (PART 2). Detection uses the separators, which SUBS leave intact.
    def _dashes(s):  # a separator line of dashes
        s = s.strip()
        return len(s) >= 8 and set(s) <= set("-\u2014\u2013")
    def _equals(s):
        s = s.strip()
        return len(s) >= 8 and set(s) <= set("=")
    paras = doc.paragraphs
    part1_idx = next((i for i, p in enumerate(paras) if "PART 1 — PARK DETAILS" in p.text), None)
    if part1_idx is not None:
        # first dash-separator after PART 1 = end of park-1 block
        first_dash = next((i for i in range(part1_idx + 1, len(paras)) if _dashes(paras[i].text)), None)
        # next equals-separator after that = start of PART 2
        part2_sep = next((i for i in range(first_dash + 1, len(paras)) if _equals(paras[i].text)), None) if first_dash else None
        if first_dash is not None and part2_sep is not None:
            for p in paras[first_dash + 1:part2_sep]:
                p._element.getparent().remove(p._element)

    # ── Integrate all terms into the body as one clean contract (NO amendment schedule) ──
    def _find(sub):
        return next((p for p in doc.paragraphs if sub in p.text), None)
    def _find_start(sub):
        return next((p for p in doc.paragraphs if p.text.strip().startswith(sub)), None)
    def _insert_after(anchor, text):
        new_p = OxmlElement("w:p")
        anchor._p.addnext(new_p)
        np = Paragraph(new_p, anchor._parent)
        np.paragraph_format.space_after = Pt(5)
        np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _run(np, text, size=10.5)
        return np

    # Clean the version history (Galascope carryover; Client has not seen prior versions)
    vh = _find("VERSION HISTORY")
    if vh is not None:
        _rewrite_paragraph(vh, f"Document Reference: {CONTRACT_REF}\nVersion: 1.0\nDate: June 2026")

    # §6.1 — fixed price (remove Indicative/Confirmed mechanism)
    p61 = _find_start("6.1 Contract Price")
    if p61 is not None:
        _rewrite_paragraph(p61,
            "6.1 Contract Price: the Contract Price is a fixed lump sum (Component A + Component B, as "
            "set out in Schedule A), exclusive of VAT, confirmed at signing on the basis of the OEM's "
            "then-current quotation for the Project configuration. The Client's site is already "
            "grid-connected, so there is no grid connection-terms trigger and no price-adjustment "
            "mechanism: the Contract Price does not change with raw-material or exchange-rate "
            "movements. After signing, the Contract Price changes only by a written change order "
            "agreed by both Parties (for example, a Client-requested configuration change).")

    # Remove the "Retention" definition (no cash retention in this Agreement)
    defr = _find_start('"Retention" means')
    if defr is not None:
        defr._element.getparent().remove(defr._element)

    # Remove the now-unused "Connection Terms" definition (site already grid-connected)
    cdef = _find_start('"Connection Terms" means')
    if cdef is not None:
        cdef._element.getparent().remove(cdef._element)

    # "Annexed:" → "Attached:" (Schedule C intro — avoid amendment-adjacent wording)
    anx = _find("[Annexed: OEM-Direct-Warranty-Undertaking")
    if anx is not None:
        _replace_in_paragraph(anx, "[Annexed:", "[Attached:")

    # §7.4 — No Cash Retention (rewrite header; delete sub-paragraphs (a)-(d))
    p74 = _find_start("7.4 Retention")
    if p74 is not None:
        _rewrite_paragraph(p74,
            "7.4 No Cash Retention: this Agreement does not operate a cash retention. The final five "
            "percent (5%) is the System Acceptance Test / commissioning payment, due on PAC "
            "(Section 7.1(d)). Defect rectification during the three (3) month Defects Liability "
            "Period is secured by the OEM 5% performance guarantee, the OEM Direct Warranty "
            "Undertaking, and the Contractor's works warranty.")
        for lead in ["(a) The Client shall retain five percent",
                     "(b) Retention shall be released within thirty",
                     "(c) Where punch-list items remain outstanding",
                     "(d) The Client shall not withhold, set off"]:
            q = _find_start(lead)
            if q is not None:
                q._element.getparent().remove(q._element)

    # §7.5(c) — drop "consideration for the Retention"
    p75c = _find("This obligation is the consideration for the Retention.")
    if p75c is not None:
        _rewrite_paragraph(p75c,
            "(c) Defects Liability: The Contractor undertakes to rectify all defects notified during "
            "the Defects Liability Period at its own cost.")

    # §9.2 — remove retention references
    p92c = _find("the Retention is released per Section 7.4.")
    if p92c is not None:
        _rewrite_paragraph(p92c,
            "(c) The Defects Liability Period ends three (3) months after PAC; no cash retention is "
            "held (Section 7.4).")
    p92e = _find("There is no Final Acceptance Certificate.")
    if p92e is not None:
        _rewrite_paragraph(p92e,
            "There is no Final Acceptance Certificate and no cash retention; the final five percent "
            "(5%) is the SAT / commissioning payment due on PAC (Section 7.1(d)).")

    # Insert project-specific clauses into their natural sections
    a = _find("The Parties wish to set out the terms under which the Contractor")
    if a is not None:
        _insert_after(a,
            "2.4 The Project is a participant in the \u0398.\u0391.\u039b\u0395.\u0399.\u0391 2021\u20132027 "
            "Just Transition Fund grant scheme (the \u201cGrant Scheme\u201d). The BESS configuration under "
            "this Agreement (6.68 MW / 20.06 MWh at the point of connection) is designed to meet the "
            "tender specification of the Grant Scheme. The Client is solely responsible for all grant "
            "administration, compliance, reporting and grant-authority liaison; the Contractor's "
            "obligation is limited to supplying and commissioning the BESS in accordance with this "
            "Agreement and providing as-built technical documentation sufficient to support the "
            "Client's grant filings.")

    a = _find("subject to presentation of the companion documents (Section 1A.2).")
    if a is not None:
        _insert_after(a,
            "1A.4 Pre-shipment payment: the ready-to-ship payment under Section 7.1(b) shall not "
            "become due until (a) the equipment has passed the Factory Acceptance Test (Section 8.6) "
            "witnessed by or on behalf of the Client, and (b) the Contractor has presented the "
            "pre-shipment Advance Payment Guarantee (APG No. 2) issued by the OEM's bank naming the "
            "Client as beneficiary. The Client never funds an equipment prepayment tranche without a "
            "guarantee in place for that tranche.")

    a = _find("Any works not expressly listed are excluded unless agreed in writing.")
    if a is not None:
        a = _insert_after(a,
            "4.6 Wind-Farm Hybrid Integration and Export Limit: the BESS is integrated with the "
            "Client's existing 10.8 MW wind farm (6 \u00d7 Vestas V100-1.8 MW turbines) at the Agia Anna "
            "site. The configuration is 4 \u00d7 5.015 MWh battery containers (20.06 MWh) and a single T8 "
            "MV Skid \u2014 one container housing 8 \u00d7 Kehua BCS1000K PCS (1.0 MW each) = 8 MW hardware, "
            "EMS-limited to 6.68 MW \u2014 providing 6.68 MW / 20.06 MWh (3-hour duration). The DISPERON EMS shall enforce the "
            "combined wind + BESS export limit on the 22 kV line to PSEUDAS S/S (confirmed overhead "
            "line capacity 14 MW) via an IEC 60870-5-104 interface with the existing ABB wind-farm "
            "controller, capping BESS export at 6.68 MW. Configuration of this export limit is a "
            "condition of PAC.")
        a = _insert_after(a,
            "4.7 MV Switchgear Bay Extension: because the BESS connects into the Client's existing "
            "22 kV wind-farm switchgear, the Contractor's scope includes the design, supply and "
            "installation of a new switchgear bay (bay J04, ABB ZX1.2-compatible MV cubicle) "
            "extending the existing switchboard, and the MV termination of the T8 MV Skid into that "
            "bay. The Client shall provide safe access to, and accurate as-built drawings and "
            "protection/selectivity data for, the existing switchgear; any latent defect in the "
            "existing switchboard not caused by the Contractor is outside the Contractor's scope.")
        _insert_after(a,
            "4.8 Grid-Forming (VSG) and Black Start: the PCS supplied is hardware-capable of "
            "grid-forming (VSG) and black-start operation via firmware. Activation, the test point, "
            "and any OEM firmware licence fee shall be confirmed by amendment to the Technical "
            "Agreement following the technical review and confirmation of the applicable DSO "
            "grid-code requirement. Hardware is unaffected and the equipment price is unchanged.")

    a = _find("No Contractor Liability for delays or losses arising from Client")
    if a is not None:
        _insert_after(a,
            "5A.5 Site Access — Army Firing Range: the Site is adjacent to an army firing range at "
            "Agia Anna. The Client is responsible for coordinating access windows with the Republic "
            "of Cyprus National Guard in advance of all installation and commissioning activities, "
            "and shall notify the Contractor of confirmed access windows at least ten (10) Business "
            "Days before each visit. Where access is refused or delayed by the military authority, "
            "the Target PAC Date extends day-for-day for the duration of the refusal and no Delay "
            "Liquidated Damages accrue for that period; the Contractor is not entitled to additional "
            "cost reimbursement beyond reasonable standby costs directly caused by such access delay.")

    a = _find("Contractor notifies within seven (7) Business Days.")
    if a is not None:
        a = _insert_after(a,
            "8.5 Delivery & Shipping Schedule: manufacturing, shipping and delivery lead time is not "
            "fixed at signing and no time commitment is given at signing; production lead time is "
            "indicative only and may change. Upon order release (triggered by the advance under "
            "Section 7.1(a)), the Contractor shall issue a binding Delivery Schedule stating the "
            "confirmed manufacturing lead time, the shipping timeline and the Target PAC Date. All "
            "time-based obligations and Delay Liquidated Damages under Section 8.4 are calculated "
            "from that confirmed Delivery Schedule.")
        _insert_after(a,
            "8.6 Factory Acceptance Test (FAT): before any equipment is despatched and before the "
            "ready-to-ship payment (Section 7.1(b)) is paid, the Client and/or its appointed "
            "third-party inspector may inspect and witness factory acceptance testing at the OEM's "
            "facility on ten (10) Business Days' notice. Where the FAT reveals defects or "
            "non-conformity, the OEM/Contractor shall remedy them at no additional cost before "
            "shipment, and the Client may re-inspect. The ready-to-ship payment is conditional on a "
            "passed FAT.")

    if p92e is not None:
        _insert_after(p92e,
            "9.3 Extended Rejection Window: the Client's right to inspect and to give notice of "
            "defect or non-conformity, and the associated repair/replace/refund remedy, extends to "
            "commissioning / PAC and is not limited to thirty (30) days after delivery. For "
            "non-conformity affecting ten percent (10%) or more of installed capacity discovered up "
            "to PAC, the Client may require repair or replacement, or reject the affected equipment "
            "and require refund of the corresponding price (recoverable, where unpaid, under the "
            "Advance Payment Guarantees per Section 10.9).")

    a = _find("The Contractor shall also procure that the OEM provides a performance bond equal to five percent (5%)")
    if a is not None:
        a = _insert_after(a,
            "(g) The OEM's bank issues the Advance Payment Guarantee(s) upon being provided with both "
            "this Agreement and the OEM supply contract, and issues them with the Client named as "
            "beneficiary from the outset. The Client may therefore demand directly under the APG(s) "
            "notwithstanding that payments are made to the Contractor: it is original issuance in the "
            "Client's name, not a transfer or assignment.")
        _insert_after(a,
            "(h) Through to PAC the Client is protected, in combination, by: the two APGs (advance and "
            "pre-shipment refund, valid to the earlier of PAC or twelve (12) months after delivery); "
            "CAR / erection all-risks insurance for the full replacement value from port of discharge "
            "through installation to PAC (the sea voyage being covered by the OEM's CIF marine "
            "insurance); the OEM 5% performance guarantee to the end of the Defects Liability Period; "
            "and the OEM Direct Warranty Undertaking and 5-year product warranty.")

    # Clean version banner (no amendment/annex language)
    first = doc.paragraphs[0]
    banner = first.insert_paragraph_before()
    _run(banner, "EPC AGREEMENT — VERSION 1.0 (June 2026)", bold=True, size=13, color=NAVY)
    first.insert_paragraph_before()

    PKG.mkdir(parents=True, exist_ok=True)
    doc.save(str(EPC_OUT))
    print(f"EPC v1 -> {EPC_OUT}")


def build_cover_note():
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(10.5)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(t, "Agia Anna Wind Farm Hybrid \u2014 BESS EPC Package", bold=True, size=15, color=NAVY)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(st,
         f"T.P. Aeolian Dynamics Ltd \u2014 6.68\u202fMW / 20.06\u202fMWh (T8 MV Skid, 3-hour), Agia Anna, Larnaca, Cyprus  "
         f"\u00b7  Ref. {CONTRACT_REF} (v1, June 2026)",
         italic=True, size=10, color=GOLD)
    doc.add_paragraph()

    para(doc, f"To: {CLIENT_NAME} \u2014 Sotiris Shiacallis (Director), and advisers / board.")
    para(doc, "From: Lighthief Cyprus Ltd \u2014 Alexander Papacosta (Director).")

    h(doc, "How this package works")
    for tx in [
        "The EPC Agreement is signed now. Signing triggers the advance payment within seven (7) days, "
        "conditional on Lighthief presenting the Direct Warranty Undertaking, the bank APG (with "
        "Aeolian as beneficiary) and the OEM performance guarantee.",
        "Unlike some projects, there is no separate grid-connection trigger: the 22\u202fkV line to "
        "PSEUDAS S/S is already in place. The advance is due from EPC Effective Date.",
        "The companion documents below travel with this pack for review now; they are executed at "
        "the order trigger (same moment as the advance), not before.",
        "No advance is payable until the signed DWU, bank APG and performance guarantee are all "
        "presented. If any is not provided in the agreed form, either party may withdraw with no "
        "liability.",
    ]:
        b = doc.add_paragraph(style="List Bullet")
        _run(b, tx)

    h(doc, "Documents in this package")
    rows = [
        ("01", "EPC Agreement — Agia Anna Wind-Farm Hybrid BESS", "For signature"),
        ("03", "LTSA — Long-Term Service Agreement (Ref. LCY-LTSA-AEO-2026)", "For review"),
        ("04", "OEM Direct Warranty Undertaking",                        "Draft — signed by OEM at order"),
        ("05", f"Advance Payment Guarantee — specimen (APG, Aeolian beneficiary)", "Form for review; bank-issued before advance"),
        ("06", "OEM 5% Performance Guarantee",                          "Form for review; issued at order"),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for c, txt in zip(hdr, ["#", "Document", "Status"]):
        _run(c.paragraphs[0], txt, bold=True, size=10)
    for num, name, status in rows:
        cells = tbl.add_row().cells
        _run(cells[0].paragraphs[0], num)
        _run(cells[1].paragraphs[0], name)
        _run(cells[2].paragraphs[0], status)

    para(doc,
         "The EMS Subscription Addendum (DISPERON) and the one-page Ownership & Guarantee Flow "
         "explainer will follow under separate cover.",
         italic=True, size=9.5, color=GREY)

    h(doc, "Key protections built in for the Client")
    for tx in [
        f"The manufacturer's equipment supply is secured by two first-demand bank guarantees "
        f"totalling {APG_COMBINED}: APG No.\u20091 ({APG_AMOUNT}, 25% of the equipment supply value) "
        "issued before the advance, and APG No.\u20092 "
        f"({APG_PRESHIP}, 50% of the equipment supply value) issued at the Factory Acceptance Test "
        "before the ready-to-ship payment. These secure the equipment refund from the manufacturer. "
        f"Your milestone payments ({ADVANCE_PCT} / {PRESHIP_PCT} of the total contract price) also "
        "include EPC services; that portion is protected by FAT before payment, the 5% performance "
        "guarantee, and your possession and title of the equipment once delivered.",
        "Both APGs name Aeolian directly as beneficiary and are kept valid to PAC (or 12\u202fmonths "
        "after delivery), so faults found at commissioning remain covered.",
        "A Factory Acceptance Test (FAT) lets you or your inspector witness factory testing and have "
        "defects fixed before the ready-to-ship payment and before shipment.",
        "Your right to reject and require remedy runs to commissioning / PAC, not just 30\u202fdays "
        "after delivery.",
        "Title passes to you on the earlier of PAC or 12\u202fmonths after delivery \u2014 so ownership "
        "never lags the end of APG cover (Lighthief retaining a security interest for any unpaid "
        "balance if title passes before PAC).",
        "The Contract Price is a fixed lump sum, confirmed at signing on the manufacturer's "
        "then-current quotation; it does not change with raw-material or exchange-rate movements. "
        "Any change after signing is only by written change order agreed by both parties.",
        "Delivery timeline and production lead time are confirmed at order release; production "
        "lead time is indicative only and may change from any figure discussed. If the confirmed "
        "timeline is unacceptable, you may withdraw without liability (bilateral walk-away right, Section 1A.2).",
        "Continuous layered cover: APGs to PAC \u2192 CAR / erection all-risks (port of discharge \u2192 "
        "PAC) \u2192 5% performance guarantee \u2192 5-year OEM warranty & DWU.",
        "The OEM's standard warranty caps aggregate liability at 10% of the payment received for "
        "defective products. The EPC and DWU include a manufacturing-defect carve-out (OEM bears "
        "full repair/replacement cost for confirmed manufacturing defects). FAT is the primary "
        "protection against a large proportion of equipment being faulty.",
    ]:
        b = doc.add_paragraph(style="List Bullet")
        _run(b, tx)

    h(doc, "Note \u2014 Kehua PCS and OEM scope")
    para(doc,
         "The PCS is Kehua C-series (BCS1000K-C-HUD), supplied within Linyang's commercial scope "
         "as part of Linyang's integrated BESS offering. The Kehua PCS is an OEM-supplied product "
         "and is covered under the Linyang product warranty and Direct Warranty Undertaking.")

    h(doc, "Note \u2014 BESS configuration: T8 MV Skid / 3-hour system")
    para(doc,
         "The BESS comprises four (4) battery containers (5.015\u202fMWh each = 20.06\u202fMWh) and a single "
         "T8 MV Skid \u2014 one container housing 8 \u00d7 Kehua BCS1000K PCS (1.0\u202fMW each) = 8\u202fMW hardware, "
         "EMS-limited to 6.68\u202fMW. This provides "
         "20.06\u202fMWh \u00f7 6.68\u202fMW = exactly 3.0\u202fhours duration at full rated export power \u2014 "
         "meeting the grant-scheme requirement with headroom. The DISPERON EMS enforces the "
         "combined wind\u202f+\u202fBESS export limit on the 22\u202fkV line to PSEUDAS S/S via "
         "IEC\u202f60870-5-104 interface with the existing ABB wind farm controller (Section\u202f4.6). "
         "This is a PAC condition and is included in the EPC scope at no additional cost.")

    h(doc, "Note \u2014 Switchgear bay extension (included in our scope)")
    para(doc,
         "Because the BESS connects into your existing 22\u202fkV wind-farm switchgear, our Component B "
         "scope includes the design, supply and installation of a new switchgear bay (bay J04, "
         "ABB ZX1.2-compatible MV cubicle) extending the existing switchboard, and the MV termination "
         "of the T8 MV Skid into that bay (Section 4.7). This is included in the Contract "
         "Price. We ask Aeolian to provide safe access and accurate as-built / protection data for "
         "the existing switchboard; any latent defect in the existing switchgear not caused by us is "
         "outside our scope.")

    h(doc, "Note \u2014 Army firing range site access")
    para(doc,
         "The Agia Anna site is adjacent to an army firing range. Aeolian is responsible for "
         "coordinating access windows with the National Guard (Section 2A.1). Access delays caused "
         "by the military authority extend the Target PAC Date day-for-day with no Delay LDs "
         "accruing for that period.")

    h(doc, "Note \u2014 Grant scheme")
    para(doc,
         "This project is a participant in the \u0398.\u0391.\u039b\u0395.\u0399.\u0391 2021\u20132027 "
         "Just Transition Fund grant scheme. The BESS configuration (6.68\u202fMW / 20.06\u202fMWh at POC) "
         "meets the tender specification (\u22655.4\u202fMW, \u226516.2\u202fMWh, \u22657,000 cycles) with headroom. "
         "Grant administration and compliance are Aeolian's responsibility; Lighthief provides "
         "as-built technical documentation to support grant filings.")

    h(doc, "Note \u2014 Governing law")
    para(doc,
         "The EPC, LTSA and EMS Addendum are governed by Cyprus law. The OEM Direct Warranty "
         "Undertaking follows the manufacturer's standard warranty terms under PRC law, with "
         "disputes resolved at SHIAC in English \u2014 consistent with international BESS procurement "
         "practice.")

    doc.add_paragraph()
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(f, "Lighthief Cyprus Ltd \u00b7 HE\u202f477423 \u00b7 28 October Ave 249, Lophitis Business Center 1, "
            "Office 201, 3035 Limassol, Cyprus \u00b7 office@lighthief.com \u00b7 +357\u202f99\u202f164\u202f158 \u00b7 solarfarms.cy",
         size=8, color=GREY)

    out = PKG / "00-Cover-Note-and-Index.docx"
    doc.save(str(out))
    print(f"Cover note -> {out}")


def build_dwu():
    """OEM DWU for Aeolian — same structure as Galascope, updated for single park + export limit."""
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(10.5)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("OEM PRODUCT WARRANTY CONFIRMATION & DIRECT UNDERTAKING")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{CLIENT_NAME} \u2014 Draft for client review")
    r2.font.size = Pt(11); r2.font.color.rgb = GOLD
    doc.add_paragraph()

    body(doc,
         "DRAFT FOR CLIENT REVIEW \u2014 to be signed and sealed by Linyang. This Undertaking "
         "reflects the manufacturer's product warranty as set out in its Overseas Energy Storage "
         "System Product Warranty Manual (LYCN/WI-3410, v2) and the manufacturer's confirmed "
         "specification.",
         bold=True, amber=True)
    doc.add_paragraph()

    body(doc, "FROM: Jiangsu Linyang Energy Storage Technology Co., Ltd (\u201cLinyang\u201d / \u201cOEM\u201d)")
    body(doc, f"TO: {CLIENT_NAME}, {CLIENT_ADDRESS} (\u201cEnd-Customer\u201d)")
    body(doc, "THROUGH: Lighthief Cyprus Ltd (HE\u202f477423), authorised exclusive distributor (\u201cDistributor\u201d)")
    body(doc, f"PROJECT: Agia Anna Wind Farm Hybrid BESS ({MW} / {MWH}), {SITE}")
    doc.add_paragraph()

    def h2(t): body(doc, t, bold=True)

    h2("1. WARRANTY CONFIRMATION")
    clause(doc, "1.1", "Linyang confirms its standard product warranty of five (5) years applies to the "
           "Products supplied for the Project (PCS, battery container, BMS, HVAC and fire-suppression, "
           "transformer and MV switchgear, DC/AC enclosure, UPS, and OEM-supplied ancillaries), per "
           "Linyang Warranty Manual v2. The PCS (Kehua C-series, BCS1000K-C-HUD) is supplied within "
           "Linyang's commercial scope as part of Linyang's integrated BESS offering and is covered "
           "under this Undertaking as an OEM-supplied product.")
    clause(doc, "1.2", "Warranty Start Date. The warranty period commences on the earlier of: (a) the "
           "start date of commissioning; or (b) six (6) months after delivery of the Products to the "
           "Project Site, preventing warranty time from burning during sea transit. This is consistent "
           "with EPC Section 3.1.")
    clause(doc, "1.3", "Coastal / corrosion rating. All battery containers and MV skids supplied are "
           "IP55 / C5 corrosion-rated, consistent with the Agia Anna coastal proximity. The warranty "
           "is unaffected by coastal distance for C5-rated equipment (Warranty Manual v2, \u00a7II.3(2)).")
    clause(doc, "1.4", "PCS sub-supply (Kehua) \u2014 Linyang responsibility. The PCS is the Kehua "
           "BCS1000K-C-HUD, sourced by Linyang from Kehua and supplied as an integral part of "
           "Linyang's BESS system. Linyang confirms the PCS is a \u2018Product\u2019 under this Undertaking "
           "and that Linyang is fully and directly responsible to the End-Customer for the PCS \u2014 "
           "including the five (5) year warranty, performance guarantees, remedy and the "
           "direct-enforcement right \u2014 to the same extent as for products Linyang manufactures "
           "itself. The End-Customer is not required to pursue Kehua directly. Linyang shall pass "
           "through the benefit of the applicable Kehua product warranty and grid-code certification "
           "(EN\u202f50549 / T\u00dcV certified for this PCS model).")

    h2("2. PERFORMANCE GUARANTEES")
    clause(doc, "2.1", "SOH (State of Health): End of Year\u202f5 \u226585%; End of Year\u202f10 \u226579.58%; "
           "End of Year\u202f15 \u226570% (1 cycle/day). Note: Year\u202f10 capacity of 16.0\u202fMWh (20\u202fMWh "
           "\u00d7 80%) meets the grant-scheme 16.2\u202fMWh tender target within tolerance.")
    clause(doc, "2.2", "Extended warranty: Year\u202f5 SOH guaranteed within the base warranty. Year\u202f10 "
           "and Year\u202f15 SOH guaranteed only where the extended warranty has been purchased "
           "(BESS \u20ac913.92/MWh/yr for years 6\u201310; \u20ac1,157.62/MWh/yr for years 11\u201315), or under LTSA.")
    clause(doc, "2.3", "Round-Trip Efficiency (RTE): the OEM guarantees \u226586.32% at the equipment "
           "measurement boundary (OEM/PCS terminal output, 0.5C, standard test conditions). The EPC "
           "site PAC floor of 84% RTE reflects system-level performance inclusive of BOP losses and "
           "is the Contractor's responsibility.")
    clause(doc, "2.4", "Cycle life: 7,000 equivalent full cycles to 70% end-of-life at 0.5C, 90% DoD, "
           "25\u00b0C. Note: 7,000 cycles is nearly double the grant-scheme requirement of 3,650 cycles "
           "(1\u202fcycle/day \u00d7 10\u202fyears).")

    h2("3. REMEDY (parts only \u2014 service and labour excluded)")
    clause(doc, "3.1", "On a valid warranty claim, Linyang shall repair or replace the defective Product "
           "or supply replacement modules/components delivered FOB or CIF Limassol at Linyang's cost.")
    clause(doc, "3.2", "On-site labour and commissioning are provided by the Contractor under the EPC "
           "works warranty and the LTSA, so the End-Customer is made whole.")
    clause(doc, "3.3", "Linyang reserves the exclusive right to determine the cause and nature of any "
           "defect. Disputed liability: joint independent third-party testing agency on agreed "
           "protocol; test costs borne by the non-liable party.")

    h2("4. DIRECT ENFORCEMENT RIGHT")
    clause(doc, "4.1", "If the Distributor is unable to fulfil warranty obligations by reason of: "
           "(a)\u202finsolvency or dissolution; (b)\u202fcessation of business; (c)\u202ftermination of the "
           "distribution arrangement; or (d)\u202ffailure to process a valid claim within thirty (30) "
           "days, the End-Customer may enforce this warranty directly against Linyang.")
    clause(doc, "4.2", "Linyang shall acknowledge a claim within forty-eight (48) hours and respond "
           "substantively within fourteen (14) days.")

    h2("5. CONDITIONS, EXCLUSIONS & LIABILITY")
    clause(doc, "5.1", "This warranty is subject to the conditions and exclusions of Linyang Warranty "
           "Manual v2. Installation and commissioning must follow Linyang's manuals and be performed "
           "by Linyang-authorised personnel; the Distributor's installation is covered by Linyang's "
           "written authorisation of the Distributor as installer for the Project.")
    clause(doc, "5.2", "Payment condition. The warranty applies only where all amounts due have been "
           "paid. Contractual retention held under the EPC is not an \u2018amount due and unpaid\u2019.")
    clause(doc, "5.3", "Liability limitation and manufacturing-defect carve-out. Linyang's aggregate "
           "liability for general warranty services and related costs is limited to ten percent (10%) "
           "of the payment received for the defective Products (Warranty Manual v2, \u00a7XI). "
           "Carve-out overriding \u00a7XI for this Project: notwithstanding \u00a7XI, Linyang agrees that "
           "the ten percent (10%) aggregate liability cap shall NOT apply to confirmed manufacturing "
           "defects. For any confirmed manufacturing defect, Linyang shall bear the full cost of "
           "repair or replacement including shipping to CIF Limassol without limitation. Linyang's "
           "execution of this Undertaking constitutes its binding agreement to this carve-out as a "
           "term that overrides Warranty Manual v2 \u00a7XI for this Project.")
    clause(doc, "5.4", "Insurance. Linyang maintains product liability insurance of EUR\u202f5,000,000 per "
           "occurrence (AXA), evidence available on request.")

    h2("6. DURATION, TRANSFER & GOVERNING LAW")
    clause(doc, "6.1", "This Undertaking is effective from delivery of the Products to Site and remains "
           "in force for the warranty period (5\u202fyears from the Warranty Start Date), and survives "
           "termination of the distribution arrangement and the Sales Contract.")
    clause(doc, "6.2", "Warranty rights transfer to a new owner where the Products remain at the "
           "original site and ownership is legally transferred.")
    clause(doc, "6.3", "This OEM Undertaking is governed by the laws of the People's Republic of China, "
           "with disputes to the Shanghai International Arbitration Center (SHIAC), in English. "
           "The EPC and Cyprus-side documents are governed by Cyprus law.")

    doc.add_paragraph()
    body(doc, "SIGNED for Jiangsu Linyang Energy Storage Technology Co., Ltd:  "
              "Name / Title / Signature / Date / Company Seal")
    body(doc, f"ACKNOWLEDGED by Lighthief Cyprus Ltd (Distributor) and {CLIENT_NAME} (End-Customer).")
    doc.add_paragraph()
    body(doc, "Lighthief Cyprus Ltd \u00b7 HE\u202f477423 \u00b7 office@lighthief.com \u00b7 "
              "+357\u202f99\u202f164\u202f158 \u00b7 solarfarms.cy",
         grey=True, size=8)

    out = PKG / "04-OEM-Direct-Warranty-Undertaking.docx"
    doc.save(str(out))
    print(f"DWU -> {out}")


def build_apg():
    """Single APG specimen (advance) — Aeolian direct beneficiary."""
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SPECIMEN \u2014 ADVANCE PAYMENT GUARANTEE")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Beneficiary: {CLIENT_NAME} \u00b7 for bank pre-confirmation")
    r2.font.size = Pt(10.5); r2.font.color.rgb = GOLD
    doc.add_paragraph()

    body(doc, "SPECIMEN FOR CLIENT REVIEW \u2014 not an issued instrument. The issuing bank will "
              "release the final guarantee on its own letterhead.", bold=True)
    doc.add_paragraph()

    body(doc, "From:  [Issuing bank \u2014 full name, registered office, SWIFT BIC] (the \u201cBank\u201d)")
    body(doc, f"To:    {CLIENT_NAME} ({CLIENT_HE}), {CLIENT_ADDRESS}, "
              "and/or [Security Agent] as project-finance security agent (the \u201cBeneficiary\u201d)")
    body(doc, "Issuing Date:  [\u25cf]      Guarantee No.:  [\u25cf]")
    doc.add_paragraph()
    body(doc, "ADVANCE PAYMENT GUARANTEE", bold=True)
    doc.add_paragraph()

    clause(doc, "1", f"Whereas the Beneficiary is procuring the supply, delivery and installation of "
           f"a grid-connected BESS for the Agia Anna Wind Farm Hybrid ({MW} / {MWH}), "
           f"{SITE} (the \u201cProject\u201d), and has agreed to make the advance payment for the "
           "equipment supply for the Project.")
    clause(doc, "2", f"We irrevocably guarantee to pay you, on your first written demand, any sum or "
           f"sums up to a maximum aggregate amount of {APG_AMOUNT} (the \u201cGuaranteed Amount\u201d), "
           "being twenty-five percent (25%) of the equipment supply value for the Project.")
    clause(doc, "3", "We shall pay within five (5) calendar days of receipt of your complying written "
           "demand stating that the equipment has not been delivered to Site and/or that the advance "
           "payment has not been refunded when due, without set-off or deduction.")
    clause(doc, "4", "A second guarantee (or increase hereto) of an equivalent fifty percent (50%) of "
           "the equipment supply value will be issued at the Factory Acceptance Test, before the "
           "ready-to-ship payment falls due. Combined equipment supply cover at that point: "
           f"{APG_COMBINED}.")
    clause(doc, "5", "This guarantee comes into force on receipt by the applicant of the advance "
           "payment and remains valid until the earlier of: (i) issuance of the Provisional "
           "Acceptance Certificate (PAC); or (ii) twelve (12) months after delivery of the "
           "equipment to Site.")
    clause(doc, "6", "Any demand must be presented in writing or by authenticated SWIFT, your bank "
           "confirming the signatories are authorised to bind the Beneficiary.")
    clause(doc, "7", "The Beneficiary may assign this guarantee or its proceeds to a project-finance "
           "lender or security agent without the Bank's further consent.")
    clause(doc, "8", "This guarantee is subject to URDG\u202f758. [Governing law / jurisdiction: as "
           "required by the Bank \u2014 to be confirmed.]")
    doc.add_paragraph()
    body(doc, "For and on behalf of the issuing Bank:  [authorised signatures / SWIFT authentication]")
    doc.add_paragraph()
    body(doc, "Specimen prepared by Lighthief Cyprus Ltd \u00b7 HE\u202f477423 \u00b7 "
              "office@lighthief.com \u00b7 +357\u202f99\u202f164\u202f158 \u00b7 solarfarms.cy \u2014 "
              "for confirmation by the issuing bank.",
         grey=True, size=8)

    out = PKG / "05-Advance-Payment-Guarantee-Specimen.docx"
    doc.save(str(out))
    print(f"APG specimen -> {out}")


def build_performance_guarantee():
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CORPORATE PERFORMANCE GUARANTEE")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"T.P. Aeolian Dynamics Ltd \u2014 Draft for client review")
    r2.font.size = Pt(11); r2.font.color.rgb = GOLD
    doc.add_paragraph()

    body(doc,
         "DRAFT FOR CLIENT REVIEW \u2014 Not valid until signed and sealed by Linyang. This is a 5% "
         "corporate performance guarantee from the manufacturer \u2014 a direct corporate obligation, "
         "not a bank guarantee; the bank instruments are the separate Advance Payment Guarantees.",
         bold=True, amber=True)
    doc.add_paragraph()

    body(doc, "Place and date:  [\u25cf]")
    body(doc, "Reference No.:  [\u25cf]")
    body(doc, f"Beneficiary:  {CLIENT_NAME} ({CLIENT_HE}), {CLIENT_ADDRESS}, "
              "and/or its project-finance security agent (\u201cBeneficiary\u201d)")
    body(doc, "Guarantor:  Jiangsu Linyang Energy Storage Technology Co., Ltd (\u201cGuarantor\u201d / \u201cLinyang\u201d)")
    doc.add_paragraph()
    body(doc, "CORPORATE PERFORMANCE GUARANTEE", bold=True)
    doc.add_paragraph()

    clause(doc, "(1)", f"Background. The Beneficiary and the Guarantor have entered into the Sales "
           f"Contract for supply of BESS equipment for the Agia Anna Wind Farm Hybrid ({MW} / "
           f"{MWH}), {SITE}. The Component A (equipment) price is approximately {COMP_A}.")
    clause(doc, "(2)", f"Guarantee amount. The Guarantor irrevocably undertakes to pay the Beneficiary "
           f"any sum or sums not exceeding in the aggregate {PG_AMOUNT} (five percent (5%) of the "
           "Component A equipment price) (\u201cGuaranteed Amount\u201d).")
    clause(doc, "(3)", "Nature. A direct corporate obligation of the Guarantor \u2014 not a bank guarantee, "
           "letter of credit, or URDG\u202f758 demand guarantee.")
    clause(doc, "(4)", "Purpose. Security for the Guarantor's performance: (a) delivery of conforming "
           "equipment CIF Limassol; (b) defects in materials and workmanship during the Defects "
           "Liability Period; (c) related OEM obligations under the supply terms.")
    clause(doc, "(5)", "Issuance. Delivered no later than one (1) month before first shipment, signed "
           "by an authorised officer with company seal.")
    clause(doc, "(6)", "Validity. From execution until the end of the Defects Liability Period, being "
           "three (3) months after PAC. Actual dates set by the confirmed Delivery Schedule; no "
           "fixed calendar date applies at signing.")
    clause(doc, "(7)", "Demand. The Beneficiary may demand after written notice of default and a "
           "thirty (30) day cure period (immediately on insolvency); payment within thirty (30) days.")
    clause(doc, "(8)", "Backing. Supported by the Guarantor's warranty reserve and product liability "
           "insurance of EUR\u202f5,000,000 per occurrence (AXA).")
    clause(doc, "(9)", "Relationship to APG. Separate from, and additional to, the Advance Payment "
           "Guarantee. Failure to deliver this guarantee does not reduce the APG requirement.")
    clause(doc, "(10)", "Assignment. Aeolian is a direct named Beneficiary. Aeolian may further assign "
           "this guarantee to a project-finance lender on written notice (effective immediately on "
           "the Distributor's insolvency or cessation of business).")
    clause(doc, "(11)", "Governing law. Governed by the laws of the People's Republic of China, with "
           "disputes to SHIAC in English, consistent with the OEM Direct Warranty Undertaking.")

    doc.add_paragraph()
    body(doc, "For and on behalf of Jiangsu Linyang Energy Storage Technology Co., Ltd")
    body(doc, "Name / Title / Signature / Company seal / Date:  [\u25cf]")
    doc.add_paragraph()
    body(doc, "Lighthief Cyprus Ltd \u00b7 HE\u202f477423 \u00b7 office@lighthief.com \u00b7 "
              "+357\u202f99\u202f164\u202f158 \u00b7 solarfarms.cy",
         grey=True, size=8)

    out = PKG / "06-OEM-5pct-Performance-Guarantee.docx"
    doc.save(str(out))
    print(f"Performance Guarantee -> {out}")


if __name__ == "__main__":
    PKG.mkdir(parents=True, exist_ok=True)
    build_epc()
    build_cover_note()
    build_dwu()
    build_apg()
    build_performance_guarantee()
    print(f"\nDONE. Package folder: {PKG}")
