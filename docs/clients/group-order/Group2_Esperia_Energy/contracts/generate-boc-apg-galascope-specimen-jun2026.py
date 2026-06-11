#!/usr/bin/env python3
"""SPECIMEN Advance Payment Guarantee — Bank of Communications form, Galascope beneficiary.

Recast of the Bank of Communications template Linyang already sent ('APG draft.docx'),
adapted so the END-CUSTOMER (Galascope Ltd) is the beneficiary from issuance — the option
Linyang confirmed the bank can do once it sees the EPC + Sales Contract.

Purpose: give Dino / Galascope a concrete DRAFT to review before EPC signature, and to send
to Linyang's bank for pre-confirmation of wording. Attach as 'Form of APG' schedule to the EPC.

NOT the issued instrument — the bank issues the final on its own letterhead after it receives
the signed EPC + Sales Contract.

Run: python docs/clients/group-order/Group2_Esperia_Energy/contracts/generate-boc-apg-galascope-specimen-jun2026.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
GREY = RGBColor(0x40, 0x40, 0x40)
OUT = Path(__file__).resolve().parent

# SSOT amounts
G1_COMP_A = 1_848_712.43
G2_COMP_A = 974_457.00
TOTAL_COMP_A = G1_COMP_A + G2_COMP_A
APG_AMOUNT = round(TOTAL_COMP_A * 0.30)  # 846,951 — advance on Component A (equipment) only


def fmt_eur(n: int) -> str:
    return f"EUR {n:,}"


def title(doc, text, sub):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sub)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GOLD
    doc.add_paragraph()


def body(doc, text, bold=False, italic=False, grey=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.bold = bold
    r.italic = italic
    if grey:
        r.font.color.rgb = GREY
    return p


def clause(doc, num, text):
    body(doc, f"({num})  {text}")


def note(doc, text):
    body(doc, "Note: " + text, italic=True, grey=True)


def build():
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(10.5)

    title(
        doc,
        "SPECIMEN — ADVANCE PAYMENT GUARANTEE",
        "Bank of Communications form · Beneficiary: Galascope Ltd · for bank pre-confirmation",
    )

    body(doc, "SPECIMEN FOR REVIEW — NOT AN ISSUED INSTRUMENT.", bold=True)
    body(
        doc,
        "Prepared by Lighthief from the issuing bank's own template, showing the terms on which "
        "the end-customer (Galascope Ltd) is named beneficiary. To be confirmed by Bank of "
        "Communications and issued on its letterhead after it receives the signed EPC and Sales "
        "Contract. Attach to the EPC as 'Form of Advance Payment Guarantee'.",
        italic=True, grey=True,
    )
    doc.add_paragraph()

    body(doc, "From:  Bank of Communications Co., Ltd., Jiangsu Provincial Branch")
    body(doc, "       11 Floor, No.218 Lushan Road, Jianye District, Nanjing, Jiangsu, China (the \u201cBank\u201d)")
    body(doc, "To:    Galascope Ltd (HE 303759), Famagusta / Cyprus, and/or [Security Agent / Alpha Bank "
              "Cyprus] as its project finance security agent (the \u201cBeneficiary\u201d)")
    body(doc, "Issuing Date:  [\u25cf]      Guarantee No.:  [\u25cf]      Advised through:  [\u25cf]")
    doc.add_paragraph()
    body(doc, "ADVANCE PAYMENT GUARANTEE", bold=True)
    doc.add_paragraph()

    clause(
        doc, "1",
        "Whereas: (i) our client, Jiangsu Linyang Energy Storage Technology Co., Ltd "
        "(the \u201cApplicant\u201d), has entered into a Sales Contract dated 17 March 2026 with "
        "Lighthief Cyprus Ltd (HE 477423) for the supply of grid-connected BESS equipment "
        "(Component A) for the Galascope 1 (5 MW / 20.06 MWh effective) and Galascope 2 "
        "(2.5 MW / 10.03 MWh effective) projects, Famagusta, Cyprus (the \u201cSales Contract\u201d); "
        "and (ii) Lighthief Cyprus Ltd has entered into an EPC contract ref. LCY-EPC-GAL-B1-2026 "
        "with you, the Beneficiary, for those projects (the \u201cEPC Contract\u201d), under which the "
        "advance payment in respect of the equipment is funded by the Beneficiary. At the request "
        "of the Applicant, we, the Bank, hereby issue this independent guarantee in your favour.",
    )
    note(
        doc,
        "The recital states the payment chain (Galascope \u2192 Lighthief \u2192 Linyang) so the bank can "
        "name Galascope as beneficiary directly. This is original issuance, not a transfer or "
        "assignment.",
    )

    clause(
        doc, "2",
        f"We irrevocably guarantee to pay you, on first written demand, any sum up to a maximum "
        f"aggregate amount of {fmt_eur(APG_AMOUNT)} (say: [\u25cf] only) (the \u201cGuaranteed Amount\u201d), "
        f"being one hundred percent (100%) of the advance payment made in respect of the equipment "
        f"(Component A) under the Sales Contract.",
    )
    note(
        doc,
        f"Guaranteed Amount = {fmt_eur(APG_AMOUNT)} (advance on Component A / equipment only). EPC "
        "services (Component B) are outside this guarantee \u2014 the Beneficiary has accepted this.",
    )

    clause(
        doc, "3",
        "We shall pay upon receipt of your complying written demand stating that the Applicant has "
        "failed to deliver the equipment under the Sales Contract and/or has failed to refund the "
        "advance payment when obliged to do so, and that the equipment has not been delivered to "
        "Site, provided the aggregate amount claimed does not exceed the Guaranteed Amount. Payment "
        "shall be made in EUR within five (5) Business Days, without set-off or deduction.",
    )

    clause(
        doc, "4",
        "Our liability shall be reduced by each amount paid under this guarantee. The Guaranteed "
        "Amount shall not be reduced merely on presentation of shipping documents or a bill of "
        "lading.",
    )

    clause(
        doc, "5",
        "This guarantee comes into force upon the Applicant's receipt of the advance payment and "
        "remains valid until, and any demand must be received by the Bank on or before, the earlier "
        "of: (i) the date of issuance of the Provisional Acceptance Certificate (PAC) for the "
        "projects; and (ii) twelve (12) months after delivery of the equipment to Site, but in any "
        "event not later than [\u25cf] (the \u201cExpiry Date\u201d). After the Expiry Date this guarantee "
        "becomes null and void whether or not returned to us.",
    )
    note(
        doc,
        "Validity aligned to EPC v5.1 \u00a710.9(d) / Anastasis board item #4: earlier of PAC or 12 "
        "months after delivery, auto-release on PAC.",
    )

    clause(
        doc, "6",
        "Any demand must be presented through your bank by authenticated SWIFT or in writing, your "
        "bank confirming that the signatories are authorised to bind the Beneficiary.",
    )

    clause(
        doc, "7",
        "This guarantee is subject to the Uniform Rules for Demand Guarantees (URDG), ICC "
        "Publication No. 758. [Governing law / jurisdiction: as required by the Bank \u2014 to be "
        "confirmed; the Beneficiary's lender may request a stated governing law.]",
    )

    clause(
        doc, "8",
        "All charges of the issuing Bank are for the account of the Applicant; charges of other "
        "banks are for the account of the Beneficiary.",
    )

    doc.add_paragraph()
    body(doc, "For and on behalf of Bank of Communications Co., Ltd., Jiangsu Provincial Branch")
    body(doc, "[Authorised signatures / SWIFT authentication]")
    doc.add_paragraph()
    body(
        doc,
        "Specimen prepared by Lighthief Cyprus Ltd \u00b7 HE 477423 \u00b7 office@lighthief.com \u00b7 "
        "+357 99 164 158 \u00b7 solarfarms.cy \u2014 for confirmation by the issuing bank.",
        grey=True,
    )

    path = OUT / "BoC-APG-Galascope-beneficiary-SPECIMEN-jun2026.docx"
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    build()
