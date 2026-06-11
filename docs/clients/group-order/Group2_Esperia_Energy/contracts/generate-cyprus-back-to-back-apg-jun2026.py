#!/usr/bin/env python3
"""Generate the DOWNSTREAM Cyprus back-to-back Advance Payment Guarantee draft.

Lighthief (Applicant) -> Cyprus/EU bank -> Galascope Ltd / security agent (Beneficiary).
First-demand bank guarantee under URDG 758, governed by Cyprus law. This solves the
transferability gap created because the upstream Linyang / Bank of Communications APG
is NON-TRANSFERABLE (beneficiary remains Lighthief).

Run: python docs/clients/group-order/Group2_Esperia_Energy/contracts/generate-cyprus-back-to-back-apg-jun2026.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Lighthief brand tokens
NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
GREY = RGBColor(0x40, 0x40, 0x40)
OUT = Path(__file__).resolve().parent

# Amounts (SSOT: Schedule A / Component A split)
G1_COMP_A = 1_848_712.43
G2_COMP_A = 974_457.00
TOTAL_COMP_A = G1_COMP_A + G2_COMP_A           # 2,823,169.43
UPSTREAM_ADVANCE = round(TOTAL_COMP_A * 0.30)  # 846,951 (upstream BoCom cover)

CLIENT_CONTRACT_PRICE = 3_444_300              # agreed Galascope batch (A3)
CLIENT_ADVANCE = round(CLIENT_CONTRACT_PRICE * 0.30)  # 1,033,290 (downstream cover)
COLLATERAL_GAP = CLIENT_ADVANCE - UPSTREAM_ADVANCE    # 186,339 services-leg


def fmt_eur(n: int) -> str:
    return f"EUR {n:,}"


def title(doc, text, sub):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sub)
    r2.font.size = Pt(11)
    r2.font.color.rgb = GOLD
    doc.add_paragraph()


def body(doc, text, bold=False, italic=False, grey=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.bold = bold
    r.italic = italic
    if grey:
        r.font.color.rgb = GREY
    return p


def clause(doc, num, text):
    body(doc, f"({num})  {text}")


def note(doc, text):
    body(doc, "Drafting note: " + text, italic=True, grey=True)


def build():
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(10.5)

    title(
        doc,
        "DRAFT — ADVANCE PAYMENT GUARANTEE (CYPRUS / BACK-TO-BACK)",
        "First-demand bank guarantee under URDG 758 — Lighthief to Galascope / lender",
    )

    body(doc, "DRAFT FOR DISCUSSION — Not valid until issued by a bank", bold=True)
    body(
        doc,
        "Downstream instrument issued by the Applicant's Cyprus/EU bank in favour of the "
        "end-customer and/or its security agent. Back-to-back with the upstream Linyang / "
        "Bank of Communications APG, which is NON-TRANSFERABLE and remains in favour of "
        "Lighthief Cyprus Ltd.",
        italic=True, grey=True,
    )
    doc.add_paragraph()

    body(doc, "Place and date of issue:  [\u25cf]")
    body(doc, "Guarantee No.:  [\u25cf]")
    body(doc, "Issuing Bank:  [Bank of Cyprus Public Company Ltd / EU bank — full name, "
              "registered office, SWIFT BIC] (the \u201cBank\u201d)")
    body(doc, "Applicant:  Lighthief Cyprus Ltd (HE 477423), 28 October Ave 249, Lophitis "
              "Business Center 1, Office 201, 3035 Limassol, Cyprus (the \u201cApplicant\u201d)")
    body(doc, "Beneficiary:  Galascope Ltd (HE 303759) and/or [Security Agent / Alpha Bank "
              "Cyprus] as its project finance security agent (the \u201cBeneficiary\u201d)")
    doc.add_paragraph()
    body(doc, "ADVANCE PAYMENT GUARANTEE", bold=True)
    doc.add_paragraph()

    clause(
        doc, "1",
        "We have been informed by the Applicant that the Beneficiary and the Applicant have "
        "entered into an EPC agreement ref. LCY-EPC-GAL-B1-2026 (the \u201cContract\u201d) for the "
        "engineering, procurement and construction of grid-connected BESS at Galascope 1 "
        "(5 MW / 20.06 MWh effective) and Galascope 2 (2.5 MW / 10.03 MWh effective), Famagusta, "
        f"Cyprus, with an indicative contract value of {fmt_eur(CLIENT_CONTRACT_PRICE)} (excl. VAT). "
        "Under the Contract the Beneficiary is required to pay the Applicant an advance payment "
        f"(the \u201cAdvance Payment\u201d) of {fmt_eur(CLIENT_ADVANCE)} against delivery of this guarantee.",
    )

    clause(
        doc, "2",
        "We, the Bank, acting on the instructions of the Applicant, hereby irrevocably and "
        "unconditionally guarantee, irrespective of the validity and legal effect of the Contract, "
        "to pay to the Beneficiary on first written demand any sum up to a maximum aggregate amount of:",
    )
    body(doc, f"      {fmt_eur(CLIENT_ADVANCE)}  (in words: [\u25cf])", bold=True)
    body(
        doc,
        "being one hundred percent (100%) of the Advance Payment, upon receipt of the Beneficiary\u2019s "
        "first written demand accompanied by the Beneficiary\u2019s statement that the Applicant has "
        "failed to refund the Advance Payment, or any part of it, when obliged to do so under the "
        "Contract, and stating the amount claimed.",
    )
    note(
        doc,
        f"Amount set to the full client advance ({fmt_eur(CLIENT_ADVANCE)}). The upstream Linyang / "
        f"Bank of Communications APG covers {fmt_eur(UPSTREAM_ADVANCE)} (equipment leg only); the "
        f"balance of {fmt_eur(COLLATERAL_GAP)} (services-leg advance) is supported by the Applicant\u2019s "
        "cash margin / counter-indemnity to the Bank. If the lender accepts equipment-leg cover only, "
        f"reduce the guaranteed amount to {fmt_eur(UPSTREAM_ADVANCE)}.",
    )

    clause(
        doc, "3",
        "Payment under this guarantee shall be made within five (5) calendar days of the Bank\u2019s "
        "receipt of a complying demand, in EUR, without set-off and without deduction for any costs, "
        "fees, taxes or charges. All banking charges of the Bank are for the account of the Applicant.",
    )

    clause(
        doc, "4",
        "The demand must be presented in writing to the Bank at its address in paragraph (8) or by "
        "authenticated SWIFT. Where presented through the Beneficiary\u2019s bank, that bank shall "
        "confirm the signatures are those of persons authorised to bind the Beneficiary.",
    )

    clause(
        doc, "5",
        "Our liability under this guarantee shall be reduced by each amount paid by us under it.",
    )

    clause(
        doc, "5A",
        "Without prejudice to paragraph (5), the maximum amount of this guarantee shall reduce only "
        "as and when the Advance Payment is contractually amortised against accepted equipment under "
        "the Contract, such reduction taking effect on delivery of the relevant equipment to Site and "
        "acceptance by the Beneficiary, evidenced by a delivery/acceptance certificate. Presentation "
        "of shipping documents or a bill of lading alone shall not reduce the amount of this guarantee.",
    )

    clause(
        doc, "6",
        "The Beneficiary may submit a demand only after the Advance Payment has been credited to the "
        "Applicant\u2019s account No. [\u25cf] maintained with [name of Applicant\u2019s bank] (SWIFT: [\u25cf]).",
    )

    clause(
        doc, "7",
        "This guarantee is valid until, and any demand must be received by the Bank on or before, the "
        "earlier of: (i) the date falling thirty (30) days after delivery of the equipment to Site "
        "(CIF Limassol, unloaded, Incoterms 2020); and (ii) [\u25cf] (a hard longstop date) "
        "(the \u201cExpiry Date\u201d).",
    )
    note(
        doc,
        "Expiry keyed to delivery-to-Site + 30 days, matching the upstream APG so the back-to-back "
        "chain runs off and on at the same time. The gap between delivery and PAC is covered by the "
        "separate 5% performance security and the OEM Direct Warranty Undertaking, not by extending "
        "this advance-refund guarantee.",
    )

    clause(
        doc, "8",
        "Any demand must be delivered to the Bank at [\u25cf full address of issuing branch / trade "
        "finance unit] no later than [\u25cf] hours (local time) on the Expiry Date.",
    )

    clause(
        doc, "9",
        "This guarantee is subject to the Uniform Rules for Demand Guarantees, ICC Publication "
        "No. 758 (URDG 758), and, to the extent not inconsistent therewith, is governed by the laws "
        "of Cyprus. Any dispute between the Bank and the Beneficiary shall be subject to the "
        "exclusive jurisdiction of the courts of Cyprus.",
    )

    clause(
        doc, "10",
        "Assignment of proceeds. The Beneficiary may assign the proceeds of this guarantee to a "
        "project finance lender or security agent without the Bank\u2019s further consent, against "
        "written notice to the Bank. This guarantee is issued directly in favour of the Beneficiary "
        "(and/or its security agent), so no transfer of beneficiary status from a third party is "
        "required.",
    )
    note(
        doc,
        "This is the key advantage of the downstream Cyprus instrument: because it is issued directly "
        "to Galascope / the security agent, the lender does not depend on transferring the upstream "
        "Chinese guarantee (which Linyang\u2019s bank will not make transferable).",
    )

    doc.add_paragraph()
    body(doc, "Issued by the Bank:")
    body(doc, "[Issuing Bank \u2014 authorised signatures]")
    body(doc, "_______________________________        _______________________________")
    body(doc, "Authorised signatory                                   Authorised signatory")
    doc.add_paragraph()
    body(doc, "Lighthief Cyprus Ltd \u00b7 HE 477423 \u00b7 office@lighthief.com \u00b7 +357 99 164 158 \u00b7 solarfarms.cy",
         grey=True)

    path = OUT / "Cyprus-Back-to-Back-APG-Galascope-DRAFT-jun2026.docx"
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    build()
