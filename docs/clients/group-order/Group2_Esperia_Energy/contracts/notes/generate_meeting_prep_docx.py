#!/usr/bin/env python3
"""One-off: build INTERNAL meeting prep .docx for Esperia client meeting.
Run from repo root: python docs/clients/group-order/Group2_Esperia_Energy/contracts/notes/generate_meeting_prep_docx.py
Requires: python-docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor

# Lighthief brand (lighthief-brand-identity.mdc)
NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x40, 0x40, 0x40)


def h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = GOLD
    r.font.size = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)


def para(doc: Document, text: str, *, grey: bool = False) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = GREY if grey else BLACK


def bullets(doc: Document, items: list[str]) -> None:
    for t in items:
        p = doc.add_paragraph(t, style="List Bullet")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK


def main() -> None:
    root = Path(__file__).resolve().parent
    out = root / "LCY-ESPERIA-INTERNAL-meeting-prep-20-May-2026.docx"

    doc = Document()
    sect = doc.sections[0]
    sect.top_margin = sect.bottom_margin = Cm(2)
    sect.left_margin = sect.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    tr = title.add_run(
        "INTERNAL — Esperia / Galascope — Pre–client meeting notes\n"
        "Wednesday 20 May 2026"
    )
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.color.rgb = NAVY
    title.paragraph_format.space_after = Pt(12)

    para(
        doc,
        "Lighthief Cyprus Ltd · HE 477423 · Classification: INTERNAL — not for circulation to "
        "client without review. Summarises internal discussions (commercial, bank, guarantees, "
        "insurance, contract stack). Not legal or financial advice.",
        grey=True,
    )

    h(doc, "1. Contract & programme status (Galascope / Esperia)")
    bullets(
        doc,
        [
            "Galascope Ltd — two BESS sites (Galascope 1: 5 MW / 20 MWh; Galascope 2: 2.5 MW / 10 MWh), "
            "Famagusta, hybrid Category B.",
            "Full EPC package (draft for signature) with Esperia side: EPC (LCY-EPC-001 v5.0 baseline; "
            "v5.1 commercial/schedule mechanics where applicable), LTSA, EMS addendum, OEM Direct Warranty "
            "Undertaking (Linyang). Status discussed: with client (Dino / Esperia) for legal review — "
            "confirm actual status on the call.",
            "Upstream: Linyang sales contract (17 Mar 2026 baseline) recorded as executed in internal SSOT; "
            "operational gates still include APG before relevant advances to Linyang and construction-phase "
            "insurance per risk checklist — confirm with counsel what is already satisfied.",
            "Internal bankability brief (May 2026) is NOT client-facing; client-facing materials are the "
            "executed/issued contract pack and curated summaries without internal repo paths or sensitive terms.",
        ],
    )

    h(doc, "2. Payment waterfall (headline model)")
    bullets(
        doc,
        [
            "Client → Lighthief (internal model): 30% advance; 55% pre-shipment; 10% PAC; 5% retention "
            "(SSOT text — align to signed EPC §7.1).",
            "Lighthief → Linyang (internal model): 20% advance; 50% pre-shipment; 20% DAP; 10% SAT — "
            "align to executed sales contract.",
            "Cash timing: after client 85% and Linyang 70%, the next Linyang 20% (DAP) can create a short "
            "window vs client PAC 10% — model per project; consider WC or milestone alignment.",
            "Reconcile SSOT “24-month DLP” retention wording with EPC v5 bankability narrative (3-month DLP) "
            "before any client-facing schedule cites retention release.",
        ],
    )

    h(doc, "3. Advance Payment Guarantee (APG) — how it works")
    bullets(
        doc,
        [
            "APG is issued by Linyang’s bank (applicant Linyang) for the benefit of Lighthief as buyer — "
            "secures refund/reperformance if equipment prepayments are taken and supply fails.",
            "Operational rule: do not remit Linyang equipment prepayments until an acceptable APG is in hand "
            "(condition precedent in the sales/EPC chain — exact clause per executed documents).",
            "This does NOT replace possible separate security that the client’s bank (e.g. Alpha) may want "
            "from Lighthief for the client’s advance (different direction: contractor / bank to employer or "
            "security agent).",
        ],
    )

    h(doc, "4. Bank of Cyprus & Alpha / lender context")
    bullets(
        doc,
        [
            "BoC relationship is young (from 1 July); modest cumulative flows (~€200k discussed internally). "
            "Expect limited appetite for large unsecured on-demand bonds — early discussion with Trade "
            "Finance / Guarantees and RM (Marianna) to size realistic products.",
            "Alpha (or other lender) will drive CPs for the borrower; Lighthief must map what BoC can actually "
            "issue vs what Alpha asks (advance payment bond, performance bond, escrow, LC).",
            "If corporate bonds are not available unsecured, fall back to: escrow / controlled account, "
            "client-side LC, lower advance %, smaller tranches, cash-covered surety, or (if acceptable) JV "
            "with established EPC balance sheet.",
        ],
    )

    h(doc, "5. Newco / no completed BESS EPC under Lighthief Cyprus")
    bullets(
        doc,
        [
            "Treat lack of Cyprus EPC completion history as normal: credit story for thin SPV is usually "
            "structural (escrow, LC, OEM APG, insurance, delivery chain bios) plus any group experience "
            "presented accurately — not as blanket cross-guarantee unless counsel structures it.",
        ],
    )

    h(doc, "6. Sinosure / export credit (Linyang side)")
    bullets(
        doc,
        [
            "Sinosure-style cover is typically for commercial credit risk; if milestones are truly "
            "client-funded and APG gates upstream prepay, clarify with Linyang counsel exactly which "
            "product and premium they require and whether it applies to fully prepaid tranches.",
        ],
    )

    h(doc, "7. Poland bank guarantee samples")
    bullets(
        doc,
        [
            "Useful as non-binding format / negotiation examples only; not legally portable. Each deal "
            "needs a new instrument matching applicant, beneficiary, contract reference, amount, currency, "
            "law, and EPC/Linyang definitions.",
        ],
    )

    h(doc, "8. Escrow option (client cash + Linyang payment)")
    bullets(
        doc,
        [
            "Client funds escrow; release to Lighthief when objective CPs are met (e.g. APG + "
            "documentation); Lighthief pays Linyang same day / back-to-back from cleared funds.",
            "Direct escrow-to-Linyang splits are possible but need Linyang acceptance, bank KYC, and tax/invoice alignment.",
        ],
    )

    h(doc, "9. Holland (NBI) insurance supplementary pack")
    bullets(
        doc,
        [
            "Internal exposure schedule to Holland Verzekeringsmakelaars: strong numerics for CAR/DSU/LD, "
            "but treat as underwriting appendix; add explicit broker asks, deadlines, and market minimums "
            "if used as formal RFP.",
            "Fix internal inconsistencies before re-issue (park/MWh totals vs subtitle; total EPC sum "
            "reconciliation; EPC version cite should match current v5 / v5.1 stack).",
        ],
    )

    h(doc, "10. Suggested talking points for Wednesday")
    bullets(
        doc,
        [
            "Confirm signing timeline and any Alpha CPs already known.",
            "Confirm milestone dates vs batch CIF/PAC assumptions.",
            "Agree communication path for bond/escrow/LC if lender is in the loop.",
            "No internal quotation refs or margin language in client-visible follow-ups.",
        ],
    )

    doc.add_paragraph()
    fp = doc.add_paragraph("Prepared from internal working notes. Verify all figures against executed contracts and lib/portfolio-data.ts before reliance.")
    for run in fp.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = GREY

    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
