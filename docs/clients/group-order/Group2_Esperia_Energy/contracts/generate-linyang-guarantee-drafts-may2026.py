#!/usr/bin/env python3
"""Generate Linyang APG + Performance Bond draft DOCX (English).

Aligned to Polish source templates in financial/:
  - Załącznik NR 9 — Wzór Gwarancji Bankowej Zwrotu Zaliczki (PL).pdf
  - PKO BP Wzór NR 6 — Gwarancja Bankowa Zwrotu Zaliczki (URDG 758 option)

Run: python docs/clients/group-order/Group2_Esperia_Energy/contracts/generate-linyang-guarantee-drafts-may2026.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
OUT = Path(__file__).resolve().parent

# Galascope Component A (equipment CIF) — EPC v5.1 Schedule A / portfolio-data
G1_COMP_A = 1_848_712
G2_COMP_A = 974_457
TOTAL_COMP_A = G1_COMP_A + G2_COMP_A
CONTRACT_REF = "Sales Contract dated 17 March 2026"
CONTRACT_NO = "[●]"
# APG = 100% of Advance Payment (per Linyang §9B.1 / Załącznik 9 — kwota gwarancji = zaliczka)
ADVANCE_PCT = 0.30  # indicative prepay % on equipment leg — confirm vs signed sales contract
ADVANCE_AMOUNT = round(TOTAL_COMP_A * ADVANCE_PCT)
APG_AMOUNT = ADVANCE_AMOUNT  # guarantee covers 100% of the advance paid
PB_AMOUNT = round(0.05 * TOTAL_COMP_A)


def fmt_eur(n: int) -> str:
    return f"EUR {n:,}".replace(",", " ")


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(11)
    r2.font.color.rgb = GOLD
    doc.add_paragraph()


def body(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.bold = bold
    r.italic = italic


def clause(doc: Document, num: str, text: str) -> None:
    body(doc, f"({num})  {text}")


def build_apg() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    add_title(
        doc,
        "DRAFT — ADVANCE PAYMENT GUARANTEE",
        "(Refund of Advance / Gwarancja Zwrotu Zaliczki)",
    )

    body(doc, "DRAFT FOR DISCUSSION — Not valid until issued by a bank", bold=True)
    body(
        doc,
        "English draft based on Linyang/Poland templates: Załącznik nr 9 (advance guarantee) "
        "and PKO BP Wzór nr 6, adapted for Galascope Cyprus batch.",
        italic=True,
    )
    doc.add_paragraph()

    body(doc, "Place and date of issue:  [●]")
    body(doc, "Guarantee No.:  [●]")
    body(doc, "Beneficiary:  Lighthief Cyprus Ltd (HE 477423), 28 October Ave 249, "
         "Lophitis Business Center 1, Office 201, 3035 Limassol, Cyprus (\"Beneficiary\")")
    doc.add_paragraph()
    body(doc, "ADVANCE PAYMENT GUARANTEE", bold=True)
    doc.add_paragraph()

    clause(
        doc,
        "1",
        "We have been informed by our client: Jiangsu Linyang Energy Storage Technology Co., Ltd "
        f"(\"Applicant\") that you have entered into {CONTRACT_NO} {CONTRACT_REF} with the Applicant "
        f"for the supply of grid-connected BESS equipment for Galascope 1 (5 MW / 20.06 MWh effective) "
        f"and Galascope 2 (2.5 MW / 10.03 MWh effective), Famagusta, Cyprus, in the total contract value of "
        f"approximately {fmt_eur(TOTAL_COMP_A)} for Component A (equipment) (\"Contract\"). "
        f"Under the Contract you are required to pay the Applicant an Advance Payment of "
        f"{fmt_eur(ADVANCE_AMOUNT)} (\"Advance Payment\") after receipt of this bank guarantee securing "
        f"refund of the Advance Payment.",
    )

    clause(
        doc,
        "2",
        f"We, [full name of Issuing Bank, registered office, SWIFT BIC, company registration numbers] "
        f"(\"Bank\"), acting on the instructions of the Applicant, hereby irrevocably and unconditionally "
        f"guarantee, irrespective of the validity and legal effect of the Contract, payment to you of any "
        f"amount up to:  {fmt_eur(APG_AMOUNT)}  (in words: [●])  being one hundred percent (100%) of the "
        f"Advance Payment, on your first written demand for payment (\"Demand\") containing a statement that "
        f"the Applicant has failed to refund the Advance Payment to you when obliged to do so under the Contract.",
    )

    clause(
        doc,
        "3",
        "Payment under this guarantee shall be made within five (5) calendar days of receipt of your "
        "complying Demand, in the full amount claimed, without deduction for any costs, fees, or charges. "
        "All transfer and banking charges shall be for the account of the Applicant.",
    )

    clause(
        doc,
        "4",
        "Your Demand must be presented through the bank maintaining your account. For the purpose of "
        "verifying signatures on your Demand, that bank shall confirm that the signatures on the Demand "
        "are those of persons authorised to bind the Beneficiary. If the Demand is presented by SWIFT, "
        "the message shall contain the full text of the Demand and the statement referred to in paragraph (2).",
    )

    clause(
        doc,
        "5",
        "Our obligation under this guarantee shall be reduced by each amount paid by us to you under this guarantee.",
    )

    clause(
        doc,
        "5A",
        "Without prejudice to paragraph (5), the maximum amount of this guarantee shall be automatically "
        "reduced pro-rata by the CIF value of equipment delivered to Site (Incoterms 2020 CIF Limassol) "
        "upon presentation to the Bank of shipping documents acceptable to the Beneficiary.",
    )

    clause(
        doc,
        "6",
        "Under this guarantee, the Beneficiary is entitled to submit a Demand to the Bank only after "
        "the Advance Payment has been credited to the Applicant's account No. [●] maintained with "
        "[name of Applicant's bank] (SWIFT: [●]).",
    )

    clause(
        doc,
        "7",
        "This guarantee is valid until the earlier of: (i) the date of issuance of the Provisional "
        "Acceptance Certificate (PAC) for the Project; or (ii) twelve (12) months after delivery of "
        "equipment to Site, but in any event not later than [●] (\"Expiry Date\"). If the Expiry Date "
        "falls on a day on which the Bank is closed for business, this guarantee remains valid until the "
        "first banking day of the Bank following that day.",
    )

    clause(
        doc,
        "8",
        "Your Demand for payment must be delivered to the Bank's office at [● full address of bank branch "
        "or trade finance unit] no later than [●] hours on the last day of the Expiry Date.",
    )

    clause(
        doc,
        "9",
        "This guarantee shall be returned to the Bank after: (a) expiry of the Expiry Date; or (b) payment "
        "by the Bank of the full guarantee amount under this guarantee; or (c) release of the Applicant and "
        "the Bank from their obligations under this guarantee by the Beneficiary.",
    )

    clause(
        doc,
        "10",
        "The Bank's obligation under this guarantee expires upon expiry of the Expiry Date, notwithstanding "
        "that the original guarantee has not been returned to the Bank. The guarantee shall also expire "
        "automatically and in full if a complying Demand is not received by the Bank before the Expiry Date, "
        "whether or not the original guarantee is returned to the Bank.",
    )

    clause(
        doc,
        "11",
        "This guarantee is subject to the Uniform Rules for Demand Guarantees, ICC Publication No. 758 "
        "(URDG 758), and, to the extent not inconsistent therewith, to the laws of [●]. "
        "Disputes between the Bank and the Beneficiary shall be submitted to the courts of [●] / as required "
        "by the Beneficiary's project lender.",
    )

    clause(
        doc,
        "12",
        "Assignment (passable guarantee): The Beneficiary may assign or transfer this guarantee, in whole "
        "or in part, to Galascope Ltd (HE 303759), Esperia Energy Group, or a project finance lender / "
        "security agent, upon prior written notice to the Bank. Assignment shall not increase the guaranteed "
        "amount or extend the Expiry Date without the Bank's written consent (such consent not to be "
        "unreasonably withheld where assignment is for project financing purposes). "
        "(Note: Polish standard form nr 6 provides that the guarantee may not be transferred without the "
        "Bank's express written consent — this paragraph reflects the agreed Cyprus passable structure.)",
    )

    doc.add_paragraph()
    body(doc, "[Issuing Bank — corporate stamp and authorised signatures]")
    body(doc, "Source templates: financial/Załącznik NR 9; financial/PKO BP Wzór NR 6", italic=True)

    path = OUT / "Linyang-APG-draft-Galascope-G1-G2-may2026.docx"
    doc.save(path)
    print(f"Wrote {path}")


def build_corporate_performance_guarantee() -> None:
    """Corporate performance guarantee — NOT a bank guarantee (per Linyang position / RFI V3 §9B.2)."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    add_title(
        doc,
        "DRAFT — CORPORATE PERFORMANCE GUARANTEE",
        "Galascope Ltd Batch · Issued by Linyang (not a bank instrument)",
    )

    body(doc, "DRAFT FOR DISCUSSION — Not valid until signed and sealed by Linyang", bold=True)
    body(
        doc,
        "Form aligned to Lighthief RFI V3 / proposed Sales Contract §9B.2: five percent (5%) "
        "corporate performance guarantee from the Seller — not an on-demand bank guarantee. "
        "Advance payment security remains a separate bank APG (Załącznik nr 9 template).",
        italic=True,
    )
    doc.add_paragraph()

    body(doc, "Place and date:  [●]")
    body(doc, "Reference No.:  [●]")
    body(doc, "Beneficiary:  Lighthief Cyprus Ltd (HE 477423), Limassol, Cyprus (\"Beneficiary\")")
    body(
        doc,
        "Guarantor:  Jiangsu Linyang Energy Storage Technology Co., Ltd "
        "(\"Guarantor\" / \"Linyang\")",
    )
    doc.add_paragraph()
    body(doc, "CORPORATE PERFORMANCE GUARANTEE", bold=True)
    doc.add_paragraph()

    clause(
        doc,
        "1",
        f"Background. The Beneficiary and the Guarantor have entered into {CONTRACT_NO} "
        f"{CONTRACT_REF} for supply of BESS equipment for Galascope 1 (5 MW / 20.06 MWh effective) "
        f"and Galascope 2 (2.5 MW / 10.03 MWh effective), Famagusta, Cyprus. The Component A "
        f"(equipment) price is approximately {fmt_eur(TOTAL_COMP_A)} (\"Contract\").",
    )

    clause(
        doc,
        "2",
        f"Guarantee amount. The Guarantor irrevocably undertakes to pay the Beneficiary, on the terms "
        f"of this instrument, any sum or sums not exceeding in the aggregate {fmt_eur(PB_AMOUNT)} "
        f"(five percent (5%) of Component A equipment price) (\"Guaranteed Amount\").",
    )

    clause(
        doc,
        "3",
        "Nature of instrument. This is a direct corporate obligation of the Guarantor. It is not a "
        "bank guarantee, letter of credit, or demand guarantee under URDG 758. No bank is primary "
        "obligor unless the Parties later agree a separate bank instrument in writing.",
    )

    clause(
        doc,
        "4",
        "Purpose. Security for the Guarantor's performance under the Contract, including: "
        "(a) delivery of conforming equipment to CIF Limassol; "
        "(b) defects in materials and workmanship during the Defects Liability Period; and "
        "(c) related obligations that are the Guarantor's responsibility under the Contract and "
        "OEM warranty framework (RFI V3).",
    )

    clause(
        doc,
        "5",
        "Issuance. Delivered to the Beneficiary no later than one (1) month before the first "
        "shipment of equipment for the Project, signed by an authorised officer of the Guarantor "
        "and affixed with the Guarantor's company seal (chop), with a certificate of incumbency "
        "or board resolution if requested by the Beneficiary.",
    )

    clause(
        doc,
        "6",
        "Validity. From the date of execution until the end of the Defects Liability Period being "
        "three (3) months after Provisional Acceptance Certificate (PAC) for the Project "
        "(indicative: April 2027) (\"Expiry Date\").",
    )

    clause(
        doc,
        "7",
        "Demand procedure. The Beneficiary may make a written demand if the Guarantor fails to "
        "perform a material obligation under the Contract after the Beneficiary has given written "
        "notice of default and the Guarantor has failed to cure within thirty (30) days (or "
        "immediately in the event of insolvency or abandonment). The demand shall state the "
        "amount claimed (not exceeding the Guaranteed Amount) and the basis of the claim.",
    )

    clause(
        doc,
        "8",
        "Payment. The Guarantor shall pay the demanded amount within thirty (30) calendar days of "
        "receipt of a valid demand at office@lighthief.com (or such address as the Beneficiary "
        "notifies), by wire transfer to the Beneficiary's designated account.",
    )

    clause(
        doc,
        "9",
        "Reduction and release. The Guaranteed Amount reduces by any amount paid under this "
        "guarantee. This guarantee is released upon the Expiry Date if no outstanding demand has "
        "been made, or upon written release by the Beneficiary.",
    )

    clause(
        doc,
        "10",
        "Backing (RFI V3 — no parent guarantee). The Beneficiary acknowledges this corporate "
        "guarantee is supported by: (a) the Guarantor's OEM warranty reserve (approximately 1.9% "
        "per RFI V3); (b) product liability insurance of EUR 5,000,000 per occurrence (AXA); and "
        "(c) the Direct Warranty Undertaking to end customers. The Guarantor confirms it does not "
        "provide a parent company guarantee from Jiangsu Linyang Energy Co., Ltd (601222) unless "
        "separately agreed in writing.",
    )

    clause(
        doc,
        "11",
        "Relationship to APG. This instrument is separate from the Advance Payment Guarantee "
        "(bank guarantee per Załącznik nr 9 / PKO wzór). Failure to deliver this corporate "
        "guarantee does not replace or reduce the APG requirement.",
    )

    clause(
        doc,
        "12",
        "Assignment. The Beneficiary may assign this guarantee to Galascope Ltd (HE 303759) or a "
        "project finance lender / security agent upon thirty (30) days' prior written notice to "
        "the Guarantor. Assignment shall not increase the Guaranteed Amount or extend the Expiry Date.",
    )

    clause(
        doc,
        "13",
        "Governing law. Governed by the laws applicable to the Contract [● PRC / Singapore / as "
        "stated in Sales Contract]. Disputes subject to the dispute resolution clause in the Contract.",
    )

    doc.add_paragraph()
    body(doc, "For and on behalf of")
    body(doc, "Jiangsu Linyang Energy Storage Technology Co., Ltd")
    body(doc, "Name:  [●]")
    body(doc, "Title:  [●]")
    body(doc, "Signature:  _________________________     Company seal:  [●]")
    body(doc, "Date:  [●]")

    path = OUT / "Linyang-Corporate-Performance-Guarantee-Galascope-G1-G2-may2026.docx"
    doc.save(path)
    print(f"Wrote {path}")

    # Legacy filename — same content for anyone with old link
    legacy = OUT / "Linyang-Performance-Bond-draft-Galascope-G1-G2-may2026.docx"
    doc.save(legacy)
    print(f"Wrote {legacy} (legacy name — corporate form, not bank bond)")


if __name__ == "__main__":
    build_apg()
    build_corporate_performance_guarantee()
