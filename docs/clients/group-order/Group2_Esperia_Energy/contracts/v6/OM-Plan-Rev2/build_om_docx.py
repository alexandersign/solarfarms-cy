# -*- coding: utf-8 -*-
"""
Regenerates the Galascope O&M Plan (Rev.2 PDF) as a Word document,
in the Lighthief brand palette (NAVY structure / GOLD headings / WHITE-on-navy /
BLACK body / GREY sub-labels), with corrected GPS + PCS/Transformer/MVS data.
"""
import copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX

from build_blocks import flatten, classify_and_merge, merge_split_tables

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
GOLD_DARK = RGBColor(0x9C, 0x7D, 0x22)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x40, 0x40, 0x40)
GREY_LIGHT = RGBColor(0x6B, 0x6B, 0x6B)

NAVY_HEX = "1A365D"
GOLD_HEX = "C9A432"
ZEBRA_HEX = "F2F2F6"
CALLOUT_BG_HEX = "F7EFD6"
CALLOUT_BORDER_HEX = "C9A432"

FONT = "Calibri"


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color=GOLD_HEX, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)


def set_table_borders(table, color="BFBFBF", sz=4):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)


def set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)
    for idx, w in enumerate(widths_cm):
        if idx < len(table.columns):
            table.columns[idx].width = Cm(w)


def para_set(p, text, size=11, bold=False, italic=False, color=BLACK, font=FONT,
             align=None, highlight=False, space_after=4, space_before=0, line_spacing=None):
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if line_spacing:
        p.paragraph_format.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    if highlight:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return r


def cell_text(cell, text, size=10, bold=False, color=BLACK, align=None, valign=WD_ALIGN_VERTICAL.CENTER,
              highlight=False):
    cell.vertical_alignment = valign
    cell.text = ''
    lines = text.split('\n')
    p = cell.paragraphs[0]
    for r_old in list(p.runs):
        r_old._r.getparent().remove(r_old._r)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if align is not None:
        p.alignment = align
    for i, line in enumerate(lines):
        pp = p if i == 0 else cell.add_paragraph()
        if i > 0:
            pp.paragraph_format.space_after = Pt(2)
            pp.paragraph_format.space_before = Pt(0)
            if align is not None:
                pp.alignment = align
        r = pp.add_run(line)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = color
        if highlight:
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


# ---------------------------------------------------------------------------
# Header / footer
# ---------------------------------------------------------------------------

def build_header(section):
    header = section.header
    header.is_linked_to_previous = False
    # clear default paragraph
    for p in list(header.paragraphs):
        p.text = ''
    tbl = header.add_table(rows=1, cols=2, width=Cm(17.2))
    tbl.autofit = False
    set_col_widths(tbl, [8.6, 8.6])
    left, right = tbl.rows[0].cells
    shade_cell(left, NAVY_HEX)
    shade_cell(right, NAVY_HEX)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run('\u25cf LIGHTHIEF')
    r.font.name = FONT
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = GOLD
    p2 = right.paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run('O&M Services | PV Plant Construction | BESS Systems | Due Diligence | PV Panel Recycling')
    r2.font.name = FONT
    r2.font.size = Pt(7.5)
    r2.font.color.rgb = GOLD
    # remove table borders
    set_table_borders(tbl, color="1A365D", sz=0)
    # small spacer paragraph after header table
    header.add_paragraph()


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    t = OxmlElement('w:r')
    tt = OxmlElement('w:t')
    tt.text = '1'
    t.append(tt)
    fld.append(t)
    run._r.addnext(fld)


def build_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p.text = ''
    tbl = footer.add_table(rows=3, cols=2, width=Cm(17.2))
    tbl.autofit = False
    set_col_widths(tbl, [8.6, 8.6])
    set_table_borders(tbl, color="FFFFFF", sz=0)
    rows_left = ['Lighthief Cyprus Ltd', 'Company No. HE 477423 | TIN 60187188Q', 'Limassol, Cyprus']
    rows_right = ['Plan O&M - Galascope 1 / Galascope 2', 'www.lighthief.energy | www.lighthief.com', None]
    for i in range(3):
        lc, rc = tbl.rows[i].cells
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        rl = lp.add_run(rows_left[i])
        rl.font.name = FONT
        rl.font.size = Pt(7.5)
        rl.font.color.rgb = GREY
        rp = rc.paragraphs[0]
        rp.paragraph_format.space_after = Pt(0)
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if rows_right[i] is not None:
            rr = rp.add_run(rows_right[i])
            rr.font.name = FONT
            rr.font.size = Pt(7.5)
            rr.font.color.rgb = GREY
        else:
            rr = rp.add_run('Page ')
            rr.font.name = FONT
            rr.font.size = Pt(7.5)
            rr.font.color.rgb = GREY
            add_page_field(rp)


def setup_page(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    build_header(section)
    build_footer(section)
    return section


# ---------------------------------------------------------------------------
# Data corrections (GPS + PCS/Transformer/MVS) applied to the extracted blocks
# ---------------------------------------------------------------------------

LOC_G1 = "Foinitzi, Avgorou, Famagusta\n35\u00b003'53.8\"N 33\u00b051'30.0\"E"
LOC_G2 = "Kerimis, Avgorou, Famagusta\n35\u00b003'52.0\"N 33\u00b051'47.9\"E"
LOC_REMARK = ("GPS per site coordinates provided Aug 2026. Distance to nearest coastline "
              "(Famagusta Bay, NE, straight-line \u2014 not surveyed): Galascope 1 \u2248 10.0 km, "
              "Galascope 2 \u2248 9.7 km.")

PCS_G1 = ("4\u00d7 Kehua PCS BCS1250K-C-HUD (skid: BCS5000K-C-HUD/T4), 5,000 kW\n"
          "1\u00d7 step-up transformer 5,000 kVA 22/0.69 kV, oil-immersed, Dy11\n"
          "1\u00d7 auxiliary transformer 250 kVA 0.69/0.38 kV, dry-type, Dyn11\n"
          "1\u00d7 RMU: Mingyang, 24 kV, SF6-free, 630 A")
PCS_G2 = ("2\u00d7 Kehua PCS BCS1250K-C-HUD (skid: BCS2500K-C-HUD/T2), 2,500 kW\n"
          "1\u00d7 step-up transformer 2,500 kVA 22/0.69 kV, oil-immersed, Dy11\n"
          "1\u00d7 auxiliary transformer 125 kVA 0.69/0.38 kV, dry-type, Dyn11\n"
          "1\u00d7 RMU: Mingyang, 24 kV, SF6-free, 630 A")
PCS_REMARK = ("Per SA-13/SA-14 Technical Agreements v6.5 and OEM.rmu in portfolio-data.ts. RMU brand "
              "confirmed Mingyang (Aug 2026, Linyang/Costas). Serial numbers from FAT reports/asset register.")

REV3_ROW = ["Rev.3",
            "GPS coordinates (Galascope 1 & 2) and PCS/Transformer/MVS technical data completed "
            "from SA-13/SA-14; RMU brand confirmed Mingyang (Linyang/Costas, Aug 2026)",
            "Lighthief", "[TO BE COMPLETED]"]


def apply_corrections(blocks):
    for b in blocks:
        if b['type'] != 'table' or not b['rows']:
            continue
        header = b['rows'][0]
        if header == ['Item', 'Galascope 1', 'Galascope 2', 'Remarks']:
            hl = b.setdefault('highlight', set())
            for ridx, row in enumerate(b['rows']):
                if row[0] == 'Location':
                    row[1], row[2], row[3] = LOC_G1, LOC_G2, LOC_REMARK
                    hl.update({(ridx, 1), (ridx, 2), (ridx, 3)})
                elif row[0] == 'PCS / Transformer / MVS':
                    row[1], row[2], row[3] = PCS_G1, PCS_G2, PCS_REMARK
                    hl.update({(ridx, 1), (ridx, 2), (ridx, 3)})
        elif header == ['Field', 'Value', 'Field', 'Value']:
            hl = b.setdefault('highlight', set())
            for ridx, row in enumerate(b['rows']):
                for cidx, cell in enumerate(row):
                    if cell == 'REV.2 - DRAFT FOR REVIEW':
                        row[cidx] = 'REV.3 - DRAFT FOR REVIEW'
                        hl.add((ridx, cidx))
        elif header == ['Revision', 'Change Description', 'Prepared by', 'Approved by']:
            b['rows'].append(list(REV3_ROW))
            hl = b.setdefault('highlight', set())
            new_r = len(b['rows']) - 1
            hl.update({(new_r, c) for c in range(4)})


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------

def render_cover(doc, blocks):
    idx = 0
    # cover_kicker, cover_title, cover_subtitle, cover_desc
    assert blocks[0]['type'] == 'cover_kicker'
    p = doc.add_paragraph()
    para_set(p, blocks[0]['text'], size=10, bold=True, color=GOLD_DARK, space_after=2, space_before=6)
    p = doc.add_paragraph()
    para_set(p, blocks[1]['text'], size=22, bold=True, color=GOLD, space_after=2)
    p = doc.add_paragraph()
    para_set(p, blocks[2]['text'], size=14, bold=True, color=NAVY, space_after=6)
    p = doc.add_paragraph()
    para_set(p, blocks[3]['text'], size=10.5, italic=True, color=GREY, space_after=10)
    return 4


def render_table(doc, rows, highlight=None, col_widths=None, header=True, font_size=9.5):
    highlight = highlight or set()
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="BFBFBF", sz=4)
    if col_widths:
        set_col_widths(tbl, col_widths)
    for ridx, row in enumerate(rows):
        is_header = header and ridx == 0
        for cidx in range(ncols):
            cell = tbl.rows[ridx].cells[cidx]
            text = row[cidx] if cidx < len(row) else ''
            if is_header:
                shade_cell(cell, NAVY_HEX)
                cell_text(cell, text, size=font_size + 0.5, bold=True, color=WHITE)
            else:
                hl = (ridx, cidx) in highlight
                if not hl and ridx % 2 == 0:
                    shade_cell(cell, ZEBRA_HEX)
                cell_text(cell, text, size=font_size, bold=False, color=BLACK, highlight=hl)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def render_callout(doc, label, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, CALLOUT_BG_HEX)
    set_cell_borders(cell, color=CALLOUT_BORDER_HEX, sz=6)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r1 = p.add_run(label + ': ')
    r1.font.name = FONT
    r1.font.size = Pt(9.5)
    r1.font.bold = True
    r1.font.color.rgb = GOLD_DARK
    r2 = p.add_run(text)
    r2.font.name = FONT
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = BLACK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def render_blocks(doc, blocks):
    for b in blocks:
        t = b['type']
        if t in ('cover_kicker', 'cover_title', 'cover_subtitle', 'cover_desc'):
            continue  # handled by render_cover
        elif t == 'h1':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '8')
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), NAVY_HEX)
            pbdr.append(bottom)
            pPr.append(pbdr)
            num = b.get('num', '')
            prefix = (str(num) + '.  ') if num != '' else ''
            r = p.add_run(prefix + b['text'].upper())
            r.font.name = FONT
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = GOLD
        elif t == 'h2':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            para_set(p, b['text'], size=12, bold=True, color=GOLD_DARK, space_after=4, space_before=0)
        elif t == 'label':
            p = doc.add_paragraph()
            para_set(p, b['text'], size=11.5, bold=True, color=NAVY, space_after=4, space_before=8)
        elif t == 'paragraph':
            p = doc.add_paragraph()
            para_set(p, b['text'], size=11, color=BLACK, space_after=6)
        elif t == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            para_set(p, b['text'], size=11, color=BLACK, space_after=3)
        elif t == 'numbered':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.6)
            para_set(p, f"{b['num']}.\t{b['text']}", size=11, color=BLACK, space_after=3)
        elif t == 'callout':
            render_callout(doc, b['label'], b['text'])
        elif t == 'table':
            render_table(doc, b['rows'], highlight=b.get('highlight'))
        elif t == 'endmark':
            p = doc.add_paragraph()
            para_set(p, b['text'], size=10, bold=True, color=GREY_LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=20)
        else:
            raise ValueError('unknown block type ' + t)


def main():
    stream = flatten()
    blocks = classify_and_merge(stream)
    blocks = merge_split_tables(blocks)
    apply_corrections(blocks)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(11)

    setup_page(doc)

    n_cover = render_cover(doc, blocks)
    render_blocks(doc, blocks[n_cover:])

    out_path = 'Plan_OM_Galascope_G1_G2_Rev3_EN.docx'
    doc.save(out_path)
    print('saved', out_path)


if __name__ == '__main__':
    main()
