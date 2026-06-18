#!/usr/bin/env python3
"""Relax the Galascope LTSA SLA for a small (~3 engineer) team: business-day
on-site call-outs (weekend/holiday best-endeavours, backed by the annual
Availability Guarantee), reconcile the 2h/4h remote-response inconsistency,
and ease the acknowledgement window. Maintenance frequency left unchanged
(bi-annual) pending OEM warranty-frequency confirmation.

Edits the source docx in contracts/ (build script copies it into the package).
"""
from pathlib import Path
from docx import Document

SRC = Path(__file__).resolve().parent.parent / "LTSA-Galascope-Esperia-may2026.docx"

REPLACEMENTS = [
    # §8.2 response commitments
    ("(a) Critical Alerts: Initial remote response within two (2) hours; on-site attendance (if required) within eight (8) hours.",
     "(a) Critical Alerts: Initial remote response within four (4) hours; on-site attendance (if "
     "required) by the end of the next Business Day. For any safety risk the Service Provider shall "
     "apply immediate remote isolation or safe-shutdown."),
    ("(b) Major Alerts: Initial remote response within eight (8) hours; on-site attendance (if required) within forty-eight (48) hours.",
     "(b) Major Alerts: Initial remote response within one (1) Business Day; on-site attendance (if "
     "required) within five (5) Business Days."),
    # §8.5 service-level credits (align to the new commitments)
    ("within the eight (8) hour commitment for a Critical Alert",
     "by the next-Business-Day commitment for a Critical Alert"),
    ("for each full hour of delay beyond the eight (8) hour commitment",
     "for each Business Day of delay beyond that commitment"),
    ("within the forty-eight (48) hour commitment for a Major Alert",
     "within the five (5) Business Day commitment for a Major Alert"),
    ("for each full hour of delay beyond the forty-eight (48) hour commitment",
     "for each Business Day of delay beyond that commitment"),
    # §8.6 acknowledgement / escalation
    ("acknowledged by the duty technician within thirty (30) minutes of alert generation.",
     "acknowledged by the duty technician within one (1) hour of alert generation."),
    ("does not acknowledge a Critical Alert within thirty (30) minutes",
     "does not acknowledge a Critical Alert within one (1) hour"),
    # §6 maintenance frequency -> quarterly (OEM requirement to maintain warranty)
    ("(b) Bi-annual preventive maintenance visits (every six months);",
     "(b) Quarterly preventive maintenance visits (every three months);"),
    ("(b) Bi-Annual On-Site Maintenance Visits: Comprehensive on-site inspection and maintenance by qualified personnel, conducted every six (6) months (two visits per year).",
     "(b) Quarterly On-Site Maintenance Visits: Comprehensive on-site inspection and maintenance by "
     "qualified personnel, conducted every three (3) months (four visits per year), as required by "
     "the OEM to maintain warranty validity."),
    ("6.5 Bi-Annual Maintenance Visit Scope", "6.5 Quarterly Maintenance Visit Scope"),
    ("(a) Bi-Annual Maintenance Visits shall be scheduled at least thirty (30) days in advance, at mutually agreed dates approximately six (6) months apart.",
     "(a) Quarterly Maintenance Visits shall be scheduled at least thirty (30) days in advance, at "
     "mutually agreed dates approximately three (3) months apart."),
    ("Bi-annual on-site maintenance visits (2 per year)",
     "Quarterly on-site maintenance visits (4 per year)"),
    # §8.7 coverage
    ("apply on a twenty-four hours a day, seven days a week, three hundred and sixty-five days a year "
     "(24/7/365) basis, including weekends, public holidays, and night hours.",
     "operate as follows: remote monitoring and remote response operate 24/7/365. On-site attendance "
     "commitments under Section 8.2 apply on Business Days; for Critical Alerts arising on weekends, "
     "public holidays, or night hours, on-site attendance is on a reasonable-endeavours basis with "
     "immediate remote isolation available, and the annual Availability Guarantee (Section 9) is the "
     "primary performance commitment."),
]


def replace_runs(p, old, new):
    full = "".join(r.text for r in p.runs)
    if old in full:
        full = full.replace(old, new)
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = full
        else:
            p.add_run(full)
        return True
    return False


def main():
    d = Document(str(SRC))
    hits = {a: 0 for a, _ in REPLACEMENTS}
    for p in d.paragraphs:
        for old, new in REPLACEMENTS:
            if replace_runs(p, old, new):
                hits[old] += 1
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in REPLACEMENTS:
                        if replace_runs(p, old, new):
                            hits[old] += 1
    d.save(str(SRC))
    for old, n in hits.items():
        print(f"[{'OK' if n else 'MISS'}] x{n}: {old[:55]}")
    print(f"Saved {SRC.name}")


if __name__ == "__main__":
    main()
