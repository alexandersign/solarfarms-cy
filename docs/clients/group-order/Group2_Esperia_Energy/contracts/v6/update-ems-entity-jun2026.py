#!/usr/bin/env python3
"""Update the Galascope EMS Subscription Addendum: change the EMS provider entity
from 'Lighthief EU BESS Ltd' to 'R&D Innovations Sp. z o.o.' (trading as DISPERON),
NIP 9492265995, under Lighthief International Ltd.

KRS and registered office left as [●] placeholders pending confirmation.
Edits the source docx in contracts/ (the build script then copies it into the package).
"""
from pathlib import Path
from docx import Document

SRC = Path(__file__).resolve().parent.parent / "EMS-Subscription-Galascope-may2026.docx"

OLD = "Lighthief EU BESS Ltd"
NEW = "R&D Innovations Sp. z o.o."


def replace_runs(p, old, new):
    full = "".join(r.text for r in p.runs)
    if old in full:
        full = full.replace(old, new)
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = full
        else:
            p.add_run(full)


def set_cell_lines(cell, lines):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    p0.add_run(lines[0])
    for ln in lines[1:]:
        cell.add_paragraph(ln)


def main():
    d = Document(str(SRC))

    # Body paragraphs
    for p in d.paragraphs:
        replace_runs(p, "Lighthief EU BESS Ltd (trading as DISPERON)",
                     "R&D Innovations Sp. z o.o. (trading as DISPERON, under Lighthief International Ltd)")
        replace_runs(p, OLD, NEW)
        replace_runs(p, "Cyprus Managing Director", "Director")
        replace_runs(p, "Authorised Representative", "Director")
        # No fixed PAC date — Connection Terms timing is unknown
        replace_runs(p, "from PAC date (31 January 2027, target)",
                     "from the PAC date (set per the EPC Delivery Schedule confirmed at Connection Terms)")

    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                txt = cell.text
                if "Service Provider:" in txt and "Lighthief EU BESS Ltd" in txt and ("Reg. No." in txt or "Limassol" in txt):
                    # full party block cell -> rebuild for the Polish entity
                    set_cell_lines(cell, [
                        "Service Provider:",
                        "R\u0026D Innovations Sp. z o.o.",
                        "(trading as DISPERON)",
                        "NIP 9492265995 \u00b7 KRS [\u25cf]",
                        "Registered office: [\u25cf], Poland",
                        "a company under Lighthief International Ltd",
                    ])
                    continue
                for p in cell.paragraphs:
                    replace_runs(p, "Lighthief EU BESS Ltd (trading as DISPERON)",
                                 "R&D Innovations Sp. z o.o. (trading as DISPERON), under Lighthief International Ltd")
                    replace_runs(p, "For Lighthief EU BESS Ltd (DISPERON)",
                                 "For R&D Innovations Sp. z o.o. (DISPERON)")
                    replace_runs(p, OLD, NEW)
                    replace_runs(p, "Cyprus Managing Director", "Director")
                    replace_runs(p, "Authorised Representative", "Director")

    d.save(str(SRC))
    print(f"Updated EMS entity in {SRC.name} -> R&D Innovations Sp. z o.o. (DISPERON)")


if __name__ == "__main__":
    main()
