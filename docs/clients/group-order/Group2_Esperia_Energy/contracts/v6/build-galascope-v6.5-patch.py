#!/usr/bin/env python3
"""Patch the Galascope EPC v6.3 (READY-TO-SEND) body into v6.4, inline, in place.

Source of the changes:
  - Stelios Constantinou's tracked-change redline from the SCC call
    ("01-EPC-Agreement-Galascope-G1-G2-v6.3 SCC CALL COMMNETS.docx"), and
  - Dr. Arkadius Sybaris's (Lighthief International) push-back memo on that
    redline ("EPC_Amendments_Sybaris.docx", "arkadiusz push back" folder).

Every change is applied directly into the numbered clause at its correct
place in the body (no separate annex), colour-coded per the existing legend:
  RED     included as requested
  PURPLE  included -- our changes
  ORANGE  included -- our counter (Sybaris push-back)
  NAVY    base drafting / structural fix (ours)

Run:
  python3 build-galascope-v6.4-patch.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

AMD_RED = RGBColor(0xC0, 0x00, 0x00)
AMD_PURPLE = RGBColor(0x70, 0x20, 0x9C)
AMD_ORANGE = RGBColor(0xC0, 0x60, 0x00)
AMD_BLUE = RGBColor(0x1F, 0x49, 0xC0)
NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)

V6 = Path(__file__).resolve().parent
SRC = V6 / "Galascope-EPC-Package-v6.3-READY-TO-SEND" / "01-EPC-Agreement-Galascope-G1-G2-v6.3.docx"
DEST_DIR = V6 / "Galascope-EPC-Package-v6.5-READY-TO-SEND"
DEST = DEST_DIR / "01-EPC-Agreement-Galascope-G1-G2-v6.5.docx"


def _find_para(doc, prefix, occurrence=1):
    n = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith(prefix):
            rest = t[len(prefix):]
            if rest == "" or rest[0] in " \n.:,;()":
                n += 1
                if n == occurrence:
                    return p
    return None


def _run(p, text, *, bold=False, italic=False, size=10.5, color=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return r


def _clear_para(p):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)


def _write_clause(p, headline, text, color, *, justify=True):
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if headline:
        _run(p, headline.strip() + "  ", bold=True, size=10.5, color=color)
    if text:
        _run(p, text, size=10.5, color=color)


def _replace_clause(p, headline, text, color):
    _clear_para(p)
    _write_clause(p, headline, text, color)
    return p


def _new_para_after(anchor_para):
    new_p = OxmlElement("w:p")
    anchor_para._p.addnext(new_p)
    return Paragraph(new_p, anchor_para._parent)


def _insert_clause_after(anchor_para, headline, text, color, *, justify=True):
    new_para = _new_para_after(anchor_para)
    _write_clause(new_para, headline, text, color, justify=justify)
    return new_para


def _insert_many_after(anchor_para, items):
    """items: list of (headline, text, color) tuples inserted in order."""
    cur = anchor_para
    for headline, text, color in items:
        cur = _insert_clause_after(cur, headline, text, color)
    return cur


def _delete_para(p):
    p._p.getparent().remove(p._p)


def _move_after(p, anchor_para):
    """Detach p and reinsert it immediately after anchor_para."""
    el = p._p
    el.getparent().remove(el)
    anchor_para._p.addnext(el)


def build():
    assert SRC.exists(), f"source not found: {SRC}"
    DEST_DIR.mkdir(exist_ok=True)
    doc = Document(str(SRC))

    # ── header / version banner ─────────────────────────────────────────
    p0 = doc.paragraphs[0]
    _clear_para(p0)
    _run(p0, "EPC AGREEMENT — VERSION 6.5 (August 2026)", bold=True, size=12, color=NAVY)

    for p in doc.paragraphs[:3]:
        for r in p.runs:
            if "v6.3" in r.text:
                r.text = r.text.replace("v6.3", "v6.5")

    # version-history block (paragraph 8) mixes "Document Reference / Version /
    # Date" with a multi-line VERSION HISTORY list inside ONE paragraph — patch
    # by rewriting the whole run rather than a substring match, and append a
    # new v6.5 entry (do NOT reuse "v6.4": that label was already consumed by
    # earlier changelog entries folded into the v6.3 file, even though the top
    # banner was never bumped at the time).
    p_hist = None
    for p in doc.paragraphs:
        if p.text.startswith("Document Reference:"):
            p_hist = p
            break
    assert p_hist is not None, "version-history paragraph not found"
    for r in p_hist.runs:
        if "Version: 6.3" in r.text:
            r.text = r.text.replace("Version: 6.3", "Version: 6.5")
        if "Date: July 2026" in r.text:
            r.text = r.text.replace("Date: July 2026", "Date: August 2026")
    _insert_clause_after(p_hist, "",
        "v6.5 — August 2026: Dr. Arkadius Sybaris (Lighthief International) push-back on "
        "Stelios's SCC-call redline, plus the remaining Stelios SCC-call items \u2014 new "
        "\u00a71A.8 condition-subsequent structure for Companion Documents (Deficiency Notice / "
        "30-Business-Day cure, per-document independent termination \u2014 Sybaris counter to the "
        "client's global void-ab-initio draft); \u00a75A.5 burden-of-proof broadened to all \u00a75A.3 "
        "Client obligations; \u00a76.1C deleted entirely; \u00a77.4(b) retention released in 10 days "
        "(was 30), 12-month corporate PG; \u00a77.4A set-off capped at 15% of the milestone and "
        "excludes the PAC payment, improper set-off = Payment Default; \u00a711.5 replaced the "
        "most-favourable-to-Client interpretation rule with a fixed document order of "
        "precedence (EPC first); \u00a713.3A latent-defect carve-out time-limited to 2 years from "
        "PAC and safety/grid-code carve-out narrowed to Contractor-attributable, pre-PAC "
        "non-conformity, with a new \u201clatent defect\u201d definition; \u00a713.5 now expressly lists "
        "the itemised manufacturing-defect categories (Timotheos's original ask, previously "
        "only logged as \u2018included\u2019 in the response matrix but never drafted); \u00a714.4 body "
        "text moved back under its own heading (was mis-ordered after \u00a714.5) plus a new note "
        "that OEM (Linyang) provides marine cargo insurance to CIF Limassol; \u00a716.3B and new "
        "Schedule D (Lighthief International parent-company Performance Undertaking Letter) "
        "narrowed to an insolvency-only trigger and capped to the same \u00a713 limits as the "
        "Contractor (Sybaris counter to the client's broader due-performance draft); \u00a78.1 "
        "Incoterm changed to DDP to the Client's Site (OEM\u2192Contractor leg remains CIF "
        "Limassol); \u00a76.1/\u00a77.2/\u00a77C consolidated into one VAT clause keeping the Article 11B "
        "reverse-charge position while reserving the Contractor's right to invoice VAT if the "
        "Cyprus Tax Department later rules against reverse-charge; \u00a77.1 Component A gained a "
        "5% retention tranche released at FAC (Delivery tranche reduced 15%\u219210%, cash-neutral); "
        "\u00a79.1A PAC acceptance matrix gained an explicit SAT-pass item; new \u201cOrder Date\u201d "
        "definition (= Effective Date, per the Parties' decision to sign all Companion "
        "Documents together, no delay) and \u00a71A.4 updated accordingly. LTSA, DWU, EMS "
        "Addendum and the Redline Response Matrix are NOT yet updated for this round \u2014 "
        "mirroring those documents to \u00a74.4B, the EMS carve-out, SOH conditionality and the "
        "new \u00a71A.8 structure remains an open task.",
        NAVY)

    # ── 1A.8 Condition Subsequent — Companion Documents (new clause) ────
    anchor = _find_para(doc, "1A.7A")
    assert anchor is not None, "1A.7A anchor not found"
    _insert_many_after(anchor, [
        ("1A.8  Condition subsequent — agreement and execution of Companion Documents.",
         "", AMD_BLUE),
        ("(a)", "Notwithstanding signature of this Agreement, the Parties acknowledge that "
         "this Agreement is executed before the following documents (together the "
         "\u201cCompanion Documents\u201d) have been agreed and executed: (i) the Long-Term Service "
         "Agreement (Ref. LCY-LTSA-GAL-2026); (ii) the OEM Direct Warranty Undertaking for "
         "Galascope 1 and Galascope 2; (iii) Advance Payment Guarantees No. 1 and No. 2; "
         "(iv) the OEM Performance Guarantee; (v) the EMS Subscription Addendum; "
         "(vi) the Technical Agreements for Galascope 1 and Galascope 2; and (vii) the "
         "Lighthief International Ltd performance undertaking under Section 16.3B "
         "(Schedule D).", AMD_BLUE),
        ("(b)", "This Agreement is conditional upon each Companion Document being: "
         "(i) agreed in form and substance satisfactory to the Client, acting reasonably; "
         "and (ii) executed, dated and, where applicable, sealed by all parties to it, in "
         "each case within thirty (30) days after the Effective Date (the \u201cDocuments "
         "Long-Stop Date\u201d).", AMD_PURPLE),
        ("(c)", "If, in respect of any particular Companion Document, that Companion "
         "Document has not been so agreed and executed by the Documents Long-Stop Date, "
         "the Party wishing to invoke this Section 1A.8 in respect of that Companion "
         "Document shall give the other Party written notice identifying the Companion "
         "Document in question (a \u201cDeficiency Notice\u201d). If the relevant Companion "
         "Document has still not been agreed and executed within thirty (30) Business "
         "Days following a Deficiency Notice (the \u201cCure Period\u201d), either Party may "
         "terminate this Agreement by written notice, without liability to the other, save "
         "that: (i) any amounts paid by the Client shall be refunded within thirty (30) "
         "days, less reasonable and documented costs incurred by the Contractor prior to "
         "termination; and (ii) Sections 15 (Governing Law), 16.6 (Notices) and any "
         "confidentiality obligations shall survive. For the avoidance of doubt, the Cure "
         "Period and this right to terminate apply separately and independently in respect "
         "of each Companion Document listed in Section 1A.8(a).", AMD_ORANGE),
        ("(d)", "Neither Party is obliged to agree to any term of a Companion Document "
         "proposed by the other. No course of dealing, part performance of this Agreement, "
         "or payment made under it shall, of itself, be construed as either Party's "
         "acceptance that a Companion Document is agreed in form and substance "
         "satisfactory to it for the purposes of this Section 1A.8.", AMD_ORANGE),
        ("(e)", "The Parties may extend the Documents Long-Stop Date only by written "
         "agreement signed by both Parties.", AMD_BLUE),
    ])

    # ── 1A.4 — anchor Order Date to the Effective Date, no delay ────────
    p_1a4 = _find_para(doc, "1A.4")
    assert p_1a4 is not None
    _replace_clause(p_1a4, "1A.4  Nature of EPC Execution.",
        "Execution of this Agreement on the Effective Date constitutes placement of the "
        "order with the OEM (the \u201cOrder Date\u201d), consistent with the Parties' agreement to "
        "execute the Companion Documents referred to in Section 1A.8 together with this "
        "Agreement, without delay. For the avoidance of doubt, placement of the order under "
        "this Section creates the Contractor's obligation to the OEM but does not itself "
        "trigger any payment obligation of the Client: the advance payment under Section "
        "7.1(a) falls due only once the conditions in Sections 1A.1 to 1A.3 and 1A.5 are "
        "satisfied.", AMD_PURPLE)

    # ── 3.1 Definitions — add "Order Date"; tighten "Retention" ──────────
    p_dwu_def = _find_para(doc, '"OEM Direct Warranty Undertaking"')
    assert p_dwu_def is not None
    _insert_clause_after(p_dwu_def, "",
        "\u201cOrder Date\u201d means the Effective Date; see Section 1A.4.", AMD_PURPLE)

    p_retention_def = _find_para(doc, '"Retention"')
    if p_retention_def is not None:
        _replace_clause(p_retention_def, "",
            "\u201cRetention\u201d means the five percent (5%) retention amounts withheld on "
            "Component A and Component B under Sections 7.1 and 7.4, and released as "
            "provided therein.", NAVY)

    # ── 5A.5 — broaden the Client-breach burden of proof (Sybaris) ──────
    p_5a5 = _find_para(doc, "5A.5")
    assert p_5a5 is not None
    _replace_clause(p_5a5, "5A.5  Client-breach causation and mitigation.",
        "No extension of time, cost reimbursement, warranty limitation or liability "
        "exclusion arising from a Client act or omission shall apply unless the Contractor "
        "demonstrates that the Client breach directly and materially caused the relevant "
        "delay, defect or loss and that the Contractor took reasonable mitigation "
        "measures. Conversely, in respect of any obligation of the Client under Section "
        "5A.3 (and not only operational non-compliance such as low state-of-charge, denial "
        "of site access or loss of connectivity), the Client bears the burden of showing "
        "that it complied with the relevant obligation, to the same standard of "
        "directness, materiality and mitigation as set out above.", AMD_ORANGE)

    # ── 6.1(a) — flag Section 11B note ───────────────────────────────────
    p_61 = _find_para(doc, "6.1 Contract Price")
    assert p_61 is not None
    for r in p_61.runs:
        if "exclusive of VAT." in r.text:
            r.text = r.text.replace(
                "exclusive of VAT.",
                "exclusive of VAT (Section 11B rules apply \u2014 see Section 7.2).")
            r.font.color.rgb = AMD_PURPLE
            break

    # ── 6.1C — delete entirely (Sybaris) ─────────────────────────────────
    p_61c = _find_para(doc, "6.1C")
    assert p_61c is not None
    _delete_para(p_61c)

    # ── 7.1 — Component A retention tranche restructuring ───────────────
    p_71 = _find_para(doc, "7.1  Payment schedule")
    assert p_71 is not None
    _clear_para(p_71)
    _write_clause(p_71, "",
        "7.1  Payment schedule \u2014 decoupled streams.  The Contract Price (EUR 3,462,849.40, "
        "ex VAT) is structured in two independent payment streams with separate milestone "
        "triggers:\n\n"
        "COMPONENT A \u2014 Equipment Supply Price (EUR 2,823,169.43, CIF Limassol):\n"
        "A1. Equipment Advance (25%, EUR 705,792.36): due within seven (7) days of the "
        "payment trigger (\u00a71A.3); secured by APG No. 1; Contractor simultaneously pays the "
        "OEM advance.\n"
        "A2. Pre-Shipment (50%, EUR 1,411,584.72): due on joint Contractor/Client written "
        "confirmation of FAT and issuance of APG No. 2, before shipment.\n"
        "A3. Delivery (10%, EUR 282,316.94): due within seven (7) days of equipment "
        "delivery to Site, completion of joint inspection and quantity/serial-number "
        "verification.\n"
        "A4. PAC (10%, EUR 282,316.94): due within seven (7) days of PAC issuance.\n"
        "Retention (5%, EUR 141,158.47): released upon issuance of FAC at end of DLP "
        "(\u00a79.2A), subject to Section 7.4.\n\n"
        "COMPONENT B \u2014 EPC Services Price (EUR 639,679.97):\n"
        "M1a. Mobilisation + Engineering (20%, EUR 127,935.99): due at the payment trigger "
        "\u2014 professional services commence on this date: protection engineering, DSO/EAC "
        "application preparation, procurement management, ETEK design and project "
        "programme.\n"
        "M1b. Civil/Platforms complete (20%, EUR 127,935.99): due on inspection and "
        "written approval of concrete platforms by the Client or its representative.\n"
        "M2. Delivery + Placement (15%, EUR 95,951.99): due on completion of joint "
        "equipment inspection and placement of containers and MV skid on platforms.\n"
        "M3. Cabling, Earthing and LPS complete (15%, EUR 95,951.99): due on completion of "
        "all MV/LV cabling, earthing grid and DEHN LPS/SPD, with test records submitted.\n"
        "M4. EMS/SCADA + Cold Commissioning (15%, EUR 95,951.99): due on EMS/SCADA "
        "installation and cold-commissioning report.\n"
        "M5. PAC (10%, EUR 63,967.99): due within seven (7) days of PAC issuance.\n"
        "Retention (5%, EUR 31,984.00): released upon issuance of FAC at end of DLP "
        "(\u00a79.2A).\n\n"
        "APG No. 1 secures A1 (EUR 705,792.36); APG No. 2 secures A2 (EUR 1,411,584.72). "
        "Component B payments are secured by the Contractor\u2019s performance obligations and "
        "the performance guarantee. The first combined payment due at trigger = A1 + M1a = "
        "EUR 833,728.35. Payments are due within seven (7) Business Days of the relevant "
        "trigger. Cross-reference key: elsewhere in this Agreement, \u201cSection 7.1(a)\u201d means "
        "the advance payment (A1 above); \u201cSection 7.1(b)\u201d means the pre-shipment payment "
        "(A2); and \u201cSection 7.1(c)\u201d means the PAC payment (A4 and M5 together, i.e. the "
        "Component A and Component B amounts both due on PAC).",
        AMD_RED)

    # ── 7.2 / 7C — consolidate VAT clauses, add reservation, keep import note ──
    p_72 = _find_para(doc, "7.2 The Contract Price")
    p_7c = _find_para(doc, "7C  Import VAT")
    assert p_72 is not None and p_7c is not None
    _replace_clause(p_72, "7.2  VAT treatment.",
        "The Contract Price and all milestone amounts are stated exclusive of VAT. The "
        "Parties' present position is that the EPC services and equipment supplied under "
        "this Agreement are subject to Cyprus VAT on a domestic reverse-charge basis under "
        "Article 11B of the Value Added Tax Law (L. 95(I)/2000), and on that basis the "
        "Contractor shall not add or collect VAT on its invoices; the Client is responsible "
        "for self-accounting the VAT directly with the Cyprus Tax Department. No formal "
        "ruling on this treatment has yet been obtained from the Cyprus Tax Department in "
        "respect of this Agreement. If the Cyprus Tax Department at any time determines, "
        "whether by ruling, audit or otherwise, that the reverse-charge mechanism under "
        "Article 11B does not apply to any part of the Works or Equipment, the Contractor "
        "reserves the right to invoice, and the Client shall pay, VAT at the applicable "
        "rate on the amounts affected, in addition to amounts already invoiced or paid. "
        "Import duty on the BESS equipment is included within the Contract Price. Import "
        "VAT, where applicable at the point of import, shall be dealt with by whichever "
        "Party acts as importer of record for Cypriot customs purposes, as agreed in "
        "writing between the Parties before the equipment is dispatched; neither Party "
        "shall be required to pre-fund import VAT on the other's behalf.", AMD_PURPLE)
    _delete_para(p_7c)

    # ── 7.4(b) retention release + PG period (Stelios base + Sybaris 10-day) ──
    p_74b = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("(b) Retention shall be released within thirty"):
            p_74b = p
            break
    assert p_74b is not None
    _replace_clause(p_74b, "(b)",
        "Retention shall be released within ten (10) days of the end of the Defects "
        "Liability Period (three (3) months after PAC). The Contractor shall maintain a "
        "five percent (5%) performance guarantee valid for twelve (12) months after PAC in "
        "the form of a corporate performance guarantee issued by Lighthief Cyprus Ltd "
        "directly to the Client and agreed with the Client.", AMD_ORANGE)

    # ── 7.4A — set-off cap + Payment Default consequence (Sybaris) ──────
    p_74a = _find_para(doc, "7.4A")
    assert p_74a is not None
    _replace_clause(p_74a, "7.4A  Set-off rights.",
        "The Client may set off against any unpaid milestone payment (excluding, for the "
        "avoidance of doubt, any PAC payment under Section 7.1(c), which shall not be "
        "subject to set-off under this Section) any amount that is: (i) agreed in writing "
        "by the Parties as due from the Contractor; or (ii) determined as due by a binding "
        "adjudicator's or expert's decision under Section 15.3. Set-off shall not apply to "
        "amounts that are disputed in good faith and not yet ascertained. The amount set "
        "off under this Section shall not exceed fifteen percent (15%) of the milestone "
        "payment against which it is set off. The Client shall give the Contractor seven "
        "(7) Business Days' prior written notice before exercising any set-off right, "
        "identifying the amount and basis of the set-off. Any set-off exercised by the "
        "Client otherwise than in accordance with this Section 7.4A (including without the "
        "required notice, without satisfying condition (i) or (ii), or in excess of the "
        "applicable cap) shall be treated as a failure to pay by the Client under Section "
        "5A.4(b) (Payment Default), with the consequences set out therein.", AMD_ORANGE)

    # ── 8.1 — Incoterm: DDP to Client Site; OEM leg stays CIF ────────────
    p_81 = _find_para(doc, "8.1 Equipment supply")
    assert p_81 is not None
    _replace_clause(p_81, "8.1",
        "Equipment supply DDP (Delivered Duty Paid) to the Client's Site in Cyprus "
        "(Incoterms\u00ae 2020), with the Contractor bearing transport, customs clearance and "
        "import duty through to Site, consistent with Section 7.5(a) (risk transfer on "
        "delivery and unloading at Site) and Section 14.5 (insurance coverage structure). "
        "For the avoidance of doubt, the OEM's (Linyang's) supply of equipment to the "
        "Contractor remains on a CIF Limassol basis under the OEM's separate supply "
        "contract; the OEM shall provide the Contractor with its certificate of marine "
        "cargo insurance for that leg, which the Contractor shall forward to the Client per "
        "Section 14.4.", AMD_PURPLE)

    # ── 9.1A — add explicit SAT pass item to the PAC acceptance matrix ──
    p_91a = _find_para(doc, "9.1A")
    assert p_91a is not None
    for r in p_91a.runs:
        if "grid-forming/VSG and black-start readiness verified (where DSO activation is required)." in r.text:
            r.text = r.text.replace(
                "grid-forming/VSG and black-start readiness verified (where DSO activation is required).",
                "grid-forming/VSG and black-start readiness verified (where DSO activation is "
                "required); and (j) full pass of the SAT protocol set out in Schedule A.")
            r.font.color.rgb = AMD_RED
            break

    # ── 11.5 — order of precedence replaces most-favourable-to-Client (Sybaris) ──
    p_115 = _find_para(doc, "11.5")
    assert p_115 is not None
    _replace_clause(p_115, "11.5  Integrated document suite and interpretation.",
        "The EPC Agreement, the LTSA, the OEM Direct Warranty Undertaking, the Advance "
        "Payment Guarantees, the OEM Performance Guarantee, the EMS Subscription Addendum, "
        "the Technical Agreements and the OEM Warranty Terms shall be read together as a "
        "single, integrated project-document suite. In the event of any inconsistency or "
        "conflict between them, the following order of precedence shall apply (with the "
        "higher-ranked document prevailing to the extent of any inconsistency): (1) the "
        "EPC Agreement; (2) the LTSA; (3) the Advance Payment Guarantees and the OEM "
        "Performance Guarantee; (4) the OEM Direct Warranty Undertaking and the OEM "
        "Warranty Terms; (5) the EMS Subscription Addendum; (6) the Technical Agreements.",
        AMD_ORANGE)

    # ── 13.3A — time-limit + narrow the uncapped carve-outs (Sybaris) ───
    p_1333a = _find_para(doc, "13.3A")
    assert p_1333a is not None
    _replace_clause(p_1333a, "13.3A  Additional liability carve-outs (not subject to any cap).",
        "In addition to the carve-outs in Section 13.3, the liability caps in Section 13.2 "
        "shall not apply to: (a) latent defects in the Works or Equipment, provided that "
        "any claim in respect of a latent defect is brought within two (2) years of the "
        "date of PAC, after which the applicable cap under Section 13.2 shall apply; and "
        "(b) non-conformity with mandatory safety requirements or applicable grid-code "
        "requirements existing as of the date of PAC and attributable to the Contractor, "
        "excluding any non-conformity arising from a change in applicable law or "
        "regulation after the date of PAC, or from any act or omission of the Client. For "
        "the avoidance of doubt, general firmware, software and design defects that are "
        "the responsibility of the OEM remain subject to the OEM Direct Warranty "
        "Undertaking and the manufacturing-defect carve-out (Section 13.5), and are not "
        "additionally uncapped under this Section.", AMD_ORANGE)
    _insert_clause_after(p_1333a, "",
        "For the purposes of paragraph (a), a \u201clatent defect\u201d means a defect in the Works "
        "or Equipment that could not have been discovered through the exercise of "
        "reasonable due diligence by way of standard acceptance testing, visual inspection "
        "or commissioning procedures conducted as of the date of PAC or during the Defects "
        "Liability Period.", AMD_ORANGE)

    # ── 13.5 — insert Timotheos's itemised manufacturing-defect list ────
    p_135 = _find_para(doc, "13.5")
    assert p_135 is not None
    _insert_clause_after(p_135, "",
        "For the avoidance of doubt, \u201cmanufacturing defects\u201d for the purposes of this "
        "Section 13.5 expressly include systemic, latent, design, firmware, batch, "
        "safety-recall and serial defects across the PCS, BMS, HVAC, fire-suppression, "
        "transformer, MV switchgear, battery cells and ancillary equipment supplied by the "
        "OEM.", AMD_PURPLE)

    # ── 14.4 — fix mis-ordered body (currently after 14.5) + add CIF-insurance note ──
    p_144_heading = _find_para(doc, "14.4 Proof of Insurance")
    p_144_body = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("The Contractor shall provide insurance certificates or cover notes"):
            p_144_body = p
            break
    assert p_144_heading is not None and p_144_body is not None
    _move_after(p_144_body, p_144_heading)
    for r in p_144_body.runs:
        r.font.color.rgb = NAVY
    _insert_clause_after(p_144_body, "",
        "For the avoidance of doubt, marine cargo insurance for the sea voyage from the "
        "port of shipment to CIF Limassol (Section 14.5(a)) is procured and provided by "
        "the OEM (Linyang) under the CIF terms of the OEM's supply contract, and not by "
        "the Contractor. The Contractor shall forward the OEM's certificate of marine "
        "cargo insurance to the Client without delay upon its receipt from the OEM.",
        AMD_PURPLE)

    # ── 16.3B — narrow to insolvency-only trigger, cap to EPC caps, ref Schedule D ──
    p_1636 = _find_para(doc, "16.3B")
    assert p_1636 is not None
    _replace_clause(p_1636, "16.3B  Lighthief International Ltd \u2014 performance undertaking.",
        "Lighthief International Ltd (\u201cGuarantor\u201d) guarantees to the Client the due "
        "performance by Lighthief Cyprus Ltd of its core EPC obligations under this "
        "Agreement, being: (a) delivery of the BESS equipment to Site as EPC Contractor; "
        "(b) installation, commissioning and handover of the BESS to the technical "
        "specification in Schedule A; and (c) the Contractor's warranty obligations under "
        "Section 10.1 for the Warranty Period. This undertaking: (i) does not extend to "
        "financial liabilities beyond the Contract Price, to the LTSA or EMS services "
        "(which are contracted separately), or to OEM obligations guaranteed under the DWU "
        "and performance guarantee; (ii) is governed by Cyprus law; (iii) remains in force "
        "for the duration of the Warranty Period; (iv) is limited to the same caps, "
        "exclusions, carve-outs and defences available to the Contractor under Sections "
        "13.1, 13.2, 13.3 and 13.3A, as if the Guarantor were itself the Contractor; and "
        "(v) arises solely if the Contractor becomes insolvent, is placed into "
        "liquidation, administration, receivership or any analogous insolvency proceeding "
        "under the laws of the Republic of Cyprus, or otherwise formally ceases to trade "
        "or admits in writing its inability to perform, and is not triggered by mere "
        "non-performance, delay, or dispute while the Contractor remains solvent and "
        "operating. This undertaking is provided by way of a parent-company confirmation "
        "letter from Lighthief International Ltd substantially in the form of Schedule D, "
        "signed by its authorised officer (Dr. Arkadius Sybaris, Founder & CEO), delivered "
        "to the Client before the advance payment under Section 7.1(a) becomes due.",
        AMD_ORANGE)

    # ── Schedule D — Lighthief International Performance Undertaking Letter ──
    p_sched_c_last = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("(d) Governing law: Singapore"):
            p_sched_c_last = p
            break
    assert p_sched_c_last is not None

    letter_items = [
        ("", "", None),  # spacer
        ("SCHEDULE D \u2014 LIGHTHIEF INTERNATIONAL PERFORMANCE UNDERTAKING LETTER", "", GOLD),
        ("[Draft \u2014 to be issued on Lighthief International Ltd letterhead and delivered to "
         "the Client before the advance payment under Section 7.1(a) becomes due, per "
         "Section 16.3B.]", "", AMD_ORANGE),
        ("LIGHTHIEF INTERNATIONAL LTD", "", AMD_ORANGE),
        ("Date: [\u25cf] 2026", "", AMD_ORANGE),
        ("To: Galascope Ltd", "", AMD_ORANGE),
        ("[Client address]", "", AMD_ORANGE),
        ("Re: Parent Company Performance Undertaking \u2014 EPC Agreement Ref. "
         "LCY-EPC-GAL-B1-2026", "", AMD_ORANGE),
        ("Dear Sirs,", "", AMD_ORANGE),
        ("", "We refer to the Engineering, Procurement & Construction Agreement dated "
         "[\u25cf] 2026 between Lighthief Cyprus Ltd (\u201cContractor\u201d) and Galascope Ltd "
         "(\u201cClient\u201d), Reference LCY-EPC-GAL-B1-2026 (the \u201cEPC Agreement\u201d).", AMD_ORANGE),
        ("1.", "Lighthief International Ltd (\u201cGuarantor\u201d) confirms and guarantees to the "
         "Client the due performance by the Contractor of the following core obligations "
         "under the EPC Agreement (the \u201cGuaranteed Obligations\u201d):", AMD_ORANGE),
        ("(a)", "delivery of the BESS equipment to Site as EPC Contractor;", AMD_ORANGE),
        ("(b)", "installation, commissioning and handover of the BESS to the technical "
         "specification set out in Schedule A;", AMD_ORANGE),
        ("(c)", "the Contractor's warranty obligations under Section 10.1 of the EPC "
         "Agreement for the Warranty Period.", AMD_ORANGE),
        ("2.", "Exclusions. This undertaking does not extend to: (i) financial liabilities "
         "beyond the Contract Price; (ii) the Long-Term Service Agreement or EMS services, "
         "which are contracted separately; or (iii) OEM obligations guaranteed under the "
         "OEM Direct Warranty Undertaking or OEM Performance Guarantee.", AMD_ORANGE),
        ("3.", "Limitation of liability. The Guarantor's liability under this undertaking "
         "shall in no circumstances exceed, and shall be subject to, the same limitations "
         "of liability, caps, exclusions, carve-outs and defences available to the "
         "Contractor under Sections 13.1, 13.2, 13.3 and 13.3A of the EPC Agreement, as if "
         "the Guarantor were itself the Contractor.", AMD_ORANGE),
        ("4.", "Trigger of Guarantor's liability. The Guarantor's obligations under this "
         "undertaking arise solely in the event that the Contractor becomes insolvent, is "
         "placed into liquidation, administration, receivership or any analogous "
         "insolvency proceeding under the laws of the Republic of Cyprus, or otherwise "
         "formally ceases to trade or admits in writing its inability to perform the "
         "Guaranteed Obligations. For the avoidance of doubt, this undertaking shall not "
         "be triggered by mere non-performance, delay, or dispute concerning the "
         "Contractor's obligations under the EPC Agreement while the Contractor remains "
         "solvent and operating.", AMD_ORANGE),
        ("5.", "Governing law and jurisdiction. This undertaking is governed by the laws "
         "of the Republic of Cyprus, and the courts of Cyprus shall have exclusive "
         "jurisdiction.", AMD_ORANGE),
        ("6.", "Duration. This undertaking takes effect on the date of this letter and "
         "shall automatically terminate, without further notice or action, upon expiry of "
         "the Warranty Period defined in the EPC Agreement.", AMD_ORANGE),
        ("7.", "No assignment. The benefit of this undertaking may not be assigned by the "
         "Client without the Guarantor's prior written consent, save for assignment to the "
         "Client's project-finance lender or security agent in accordance with Section "
         "16.3A of the EPC Agreement.", AMD_ORANGE),
        ("Yours faithfully,", "", AMD_ORANGE),
        ("Dr. Arkadius Sybaris", "", AMD_ORANGE),
        ("Founder & CEO", "", AMD_ORANGE),
        ("Lighthief International Ltd", "", AMD_ORANGE),
    ]
    cur = p_sched_c_last
    for headline, text, color in letter_items:
        new_p = _new_para_after(cur)
        if headline == "SCHEDULE D \u2014 LIGHTHIEF INTERNATIONAL PERFORMANCE UNDERTAKING LETTER":
            new_p.paragraph_format.space_before = Pt(12)
            new_p.paragraph_format.space_after = Pt(4)
            _run(new_p, headline, bold=True, size=12, color=color)
        elif headline == "" and text == "":
            pass  # spacer paragraph, left empty
        else:
            _write_clause(new_p, headline + ("  " if headline and text else ""), text, color)
        cur = new_p

    doc.save(str(DEST))
    print(f"Saved: {DEST}")


if __name__ == "__main__":
    build()
