#!/usr/bin/env python3
"""Client-review drafts for Galascope — aligned to UPSTREAM-ONLY Linyang terms.

Two outputs, both DRAFT FOR CLIENT REVIEW (not yet signed by Linyang):

1. OEM-DWU-Galascope-CLIENT-REVIEW-jun2026.docx
   The Direct Warranty Undertaking rewritten to reflect ONLY what Linyang
   actually provides upstream — per 'Linyang Warranty Terms v2' (LYCN/WI-3410)
   and the RFI V3 responses (rfi-linyang-responses-feb2026.md). It does NOT
   include the enhanced asks we drafted in the v5.1 redlines (PAC-only warranty
   start, 60-day solvent-dispute direct claim, install-labour remedy) because
   Linyang has not agreed those. Honest pass-through, so the client is not
   shown more than the OEM will sign.

2. Performance-Guarantee-Galascope-CLIENT-REVIEW-jun2026.docx
   The 5% corporate performance guarantee, names corrected for the Galascope
   deal (replaces the mis-named Polish 'Solarfun' template Linyang returned).
   Beneficiary Lighthief, assignable to Galascope / lender. For client review.

Run: python docs/clients/group-order/Group2_Esperia_Energy/contracts/v6/generate-galascope-client-review-drafts-jun2026.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
GREY = RGBColor(0x40, 0x40, 0x40)
AMBER = RGBColor(0x78, 0x35, 0x0F)
OUT = Path(__file__).resolve().parent

# Component A (equipment) — EPC v5.1 Schedule A
G1_COMP_A = 1_848_712
G2_COMP_A = 974_457
TOTAL_COMP_A = G1_COMP_A + G2_COMP_A
PB_AMOUNT = round(0.05 * TOTAL_COMP_A)  # EUR 141,158 (5% of Component A)
CONTRACT_REF = "Sales Contract dated 17 March 2026"


def fmt_eur(n: int) -> str:
    return f"EUR {n:,}"


def title(doc, text, sub):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sub)
    r2.font.size = Pt(11)
    r2.font.color.rgb = GOLD
    doc.add_paragraph()


def body(doc, text, bold=False, italic=False, grey=False, amber=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if grey:
        r.font.color.rgb = GREY
    if amber:
        r.font.color.rgb = AMBER
    return p


def h(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = NAVY
    return p


def clause(doc, num, text):
    body(doc, f"{num}  {text}")


# ──────────────────────────────────────────────────────────────────────────
# 1. DWU — UPSTREAM-ONLY (client review)
# ──────────────────────────────────────────────────────────────────────────
def build_dwu():
    doc = Document()
    s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(10.5)

    title(doc, "OEM PRODUCT WARRANTY CONFIRMATION & DIRECT UNDERTAKING",
          "Galascope Ltd — Draft for client review")

    body(doc, "DRAFT FOR CLIENT REVIEW \u2014 to be signed and sealed by Linyang. This Undertaking reflects "
              "the manufacturer's product warranty as set out in its Overseas Energy Storage System "
              "Product Warranty Manual (LYCN/WI-3410, v2) and the manufacturer's confirmed specification.",
         bold=True, amber=True)
    doc.add_paragraph()

    body(doc, "FROM: Jiangsu Linyang Energy Storage Technology Co., Ltd (\u201cLinyang\u201d / \u201cOEM\u201d)")
    body(doc, "TO: Galascope Ltd, Karaiskaki 6, City House, 3032 Limassol, Cyprus (\u201cEnd-Customer\u201d)")
    body(doc, "THROUGH: Lighthief Cyprus Ltd (HE 477423), authorised exclusive distributor (\u201cDistributor\u201d)")
    body(doc, "PROJECT: Galascope 1 (5 MW / 20.06 MWh) & Galascope 2 (2.5 MW / 10.03 MWh), Famagusta, Cyprus")
    doc.add_paragraph()

    h(doc, "1. WARRANTY CONFIRMATION")
    clause(doc, "1.1", "Linyang confirms its standard product warranty of five (5) years applies to the "
           "Products supplied for the Project (PCS, battery container, BMS, HVAC and fire-suppression, "
           "transformer and MV switchgear, DC/AC enclosure, UPS, and OEM-supplied ancillaries), per "
           "Linyang Warranty Manual v2. The PCS (Kehua C-series, BCS1250K-C-HUD) is supplied within "
           "Linyang's commercial scope as part of Linyang's integrated BESS offering and is covered "
           "under this Undertaking as an OEM-supplied product.")
    clause(doc, "1.2", "Warranty Start Date. The warranty period commences on the earlier of: (a) the start "
           "date of commissioning; or (b) six (6) months after delivery of the Products to the Project "
           "Site (Warranty Manual v2, §II.4, modified by this Undertaking for this Project to run from "
           "delivery to Site rather than from factory shipment, preventing warranty time from burning "
           "during sea transit). This is consistent with EPC Section 3.1 Warranty Start Date.")
    clause(doc, "1.3", "Coastal installation. For systems without a C5-rated enclosure installed 2\u20135 km from "
           "the sea, the warranty for PCS, transformer and switchgear is reduced to two (2) years; within "
           "2 km it is void. For C5-rated systems within 500 m of the sea, that warranty is two (2) years. "
           "The Project's distance from the sea and enclosure rating shall be recorded in Schedule A. "
           "(Warranty Manual v2, §II.3(2).)")
    clause(doc, "1.4", "PCS sub-supply (Kehua) \u2014 Linyang responsibility. The power conversion system supplied "
           "for the Project is the Kehua BCS1250K-C-HUD, sourced by Linyang from Kehua and supplied as an "
           "integral part of Linyang's BESS system. Linyang confirms that the PCS is a \u201cProduct\u201d under this "
           "Undertaking and that Linyang is fully and directly responsible to the End-Customer for the PCS "
           "\u2014 including the five (5) year warranty (Section 1.1), the performance guarantees (Section 2), the "
           "remedy (Section 3) and the direct-enforcement right (Section 4) \u2014 to the same extent as for "
           "products Linyang manufactures itself. The End-Customer's recourse for the PCS is against Linyang "
           "under this Undertaking, and the End-Customer is not required to pursue Kehua directly. Linyang "
           "shall pass through to the End-Customer the benefit of the applicable Kehua product warranty and "
           "grid-code certification (including the EN 50549 / T\u00dcV certification held for the PCS).")

    h(doc, "2. PERFORMANCE GUARANTEES (capacity, efficiency, cycle life)")
    clause(doc, "2.1", "SOH (State of Health), per the product Technical Agreement, as confirmed by the manufacturer: "
           "End of Year 5 \u2265 85%; End of Year 10 \u2265 79.58%; End of Year 15 \u2265 70% (1 cycle/day).")
    clause(doc, "2.2", "Within the base five (5) year warranty, the guaranteed SOH point is Year 5 (\u2265 85%). "
           "The Year 10 and Year 15 SOH guarantees apply only where the End-Customer has purchased the "
           "extended warranty for years 6\u201315 (manufacturer rates: BESS \u20ac913.92/MWh/yr for years 6\u201310 and "
           "\u20ac1,157.62/MWh/yr for years 11\u201315; PCS+MVS \u20ac747.76 and \u20ac926.10 respectively), or under the LTSA. "
           "Maximum OEM SOH guarantee is fifteen (15) years.")
    clause(doc, "2.3", "Round-Trip Efficiency (RTE): the OEM guarantees \u2265 86.32% at the equipment measurement "
           "boundary (OEM/PCS terminal output, 0.5C, standard test conditions), consistent with the EPC "
           "technical specification. This measurement excludes losses from Distributor-supplied BOP cabling, "
           "transformer and auxiliary consumption, which are outside the OEM measurement scope. The EPC site "
           "PAC floor of 84% RTE reflects system-level performance inclusive of BOP losses and is the "
           "Contractor's responsibility under the EPC.")
    clause(doc, "2.4", "Cycle life: 7,000 equivalent full cycles to 70% end-of-life at 0.5C, 90% DoD, 25\u00b0C "
           "as confirmed by the manufacturer.")

    h(doc, "3. REMEDY (parts only — service and labour excluded)")
    clause(doc, "3.1", "On a valid warranty claim, Linyang shall repair or replace the defective Product, or "
           "supply replacement modules/components (new or equivalent refurbished parts), delivered FOB or "
           "CIF Limassol at Linyang's cost.")
    clause(doc, "3.2", "Linyang supplies replacement parts; under the OEM product warranty, freight beyond the "
           "stated Incoterm, travel, on-site removal/installation labour, commissioning, and third-party fees "
           "are not part of the OEM remedy (Warranty Manual v2, §IX). For the End-Customer, on-site "
           "replacement labour and commissioning are provided by the Contractor under the EPC works warranty "
           "and the LTSA, so the End-Customer is made whole.")
    clause(doc, "3.3", "Linyang reserves the exclusive right to determine the cause and nature of any defect. "
           "Where the Parties dispute liability, they shall jointly appoint an independent third-party testing "
           "agency on an agreed protocol; if the defect is found not attributable to the Products, the test "
           "costs are borne by the Warranty Right Holder (Warranty Manual v2, §VIII.4 / §X).")

    h(doc, "4. DIRECT ENFORCEMENT RIGHT")
    clause(doc, "4.1", "If the Distributor is unable to fulfil its warranty obligations to the End-Customer by "
           "reason of: (a) insolvency, liquidation, or dissolution; (b) cessation of business; (c) termination "
           "of the distribution arrangement; or (d) failure to process a valid warranty claim within thirty "
           "(30) days, the End-Customer may enforce this warranty directly against Linyang, providing evidence "
           "of purchase, the defect, and the Distributor's inability to act.")
    clause(doc, "4.2", "Linyang shall acknowledge a claim within forty-eight (48) hours (excluding weekends and "
           "public holidays) and respond substantively within fourteen (14) days (Warranty Manual v2, §VIII.2).")

    h(doc, "5. CONDITIONS, EXCLUSIONS & LIABILITY")
    clause(doc, "5.1", "This warranty is subject to the conditions and exclusions of Linyang Warranty Manual v2, "
           "including proper operation per the manuals, authorised maintenance, intact serial numbers, and no "
           "unauthorised modification/relocation. Installation and commissioning must follow Linyang's manuals "
           "and be performed by Linyang-authorised personnel; the Distributor's installation is covered by "
           "Linyang's written authorisation of the Distributor as installer for the Project.")
    clause(doc, "5.2", "Payment condition. The warranty applies only where the Warranty Right Holder has paid all "
           "amounts due (Warranty Manual v2, §III.2). Contractual retention held under the EPC is not an "
           "\u201camount due and unpaid\u201d for this purpose.")
    clause(doc, "5.3", "Liability limitation and manufacturing-defect carve-out. Linyang's aggregate liability "
           "for general warranty services and related costs is limited to ten percent (10%) of the payment "
           "received for the defective Products (Warranty Manual v2, §XI). Carve-out overriding §XI for this "
           "Project: notwithstanding §XI, Linyang agrees that the ten percent (10%) aggregate liability cap "
           "shall NOT apply to confirmed manufacturing defects in the Products. For any confirmed "
           "manufacturing defect, Linyang shall bear the full cost of repair or replacement of the affected "
           "Products, including shipping to CIF Limassol, without limitation by reference to the 10% cap. "
           "Linyang's execution of this Undertaking (authorised signature and company seal below) constitutes "
           "its binding agreement to this carve-out as a term that overrides Warranty Manual v2 §XI for this "
           "Project. The signed Undertaking incorporating this carve-out is a Condition Precedent to the "
           "advance under the EPC.")
    clause(doc, "5.4", "Insurance. Linyang maintains product liability insurance of EUR 5,000,000 per occurrence "
           "(AXA), evidence available on request.")

    h(doc, "6. DURATION, TRANSFER & GOVERNING LAW")
    clause(doc, "6.1", "This Undertaking is effective from delivery of the Products to Site and remains in force "
           "for the warranty period (5 years from the Warranty Start Date), and survives termination of the "
           "distribution arrangement and the Sales Contract.")
    clause(doc, "6.2", "Warranty rights transfer to a new owner where the Products remain at the original site and "
           "ownership is legally transferred (Warranty Manual v2, §VI) \u2014 enabling assignment to the End-Customer "
           "and its project-finance security agent.")
    clause(doc, "6.3", "This OEM Undertaking is governed by the laws of the People's Republic of China, with "
           "disputes to the Shanghai International Arbitration Center (SHIAC), in English, per the manufacturer's "
           "standard warranty terms (Warranty Manual v2, §XIII). The EPC and the Cyprus-side documents are "
           "governed by Cyprus law.")

    doc.add_paragraph()
    body(doc, "SIGNED for Jiangsu Linyang Energy Storage Technology Co., Ltd:  Name / Title / Signature / Date / Company Seal")
    body(doc, "ACKNOWLEDGED by Lighthief Cyprus Ltd (Distributor) and Galascope Ltd (End-Customer).")
    doc.add_paragraph()
    body(doc, "Lighthief Cyprus Ltd \u00b7 HE 477423 \u00b7 office@lighthief.com \u00b7 +357 99 164 158 \u00b7 solarfarms.cy",
         grey=True, size=8)

    path = OUT / "OEM-DWU-Galascope-CLIENT-REVIEW-jun2026.docx"
    doc.save(path)
    print(f"Wrote {path}")


# ──────────────────────────────────────────────────────────────────────────
# 2. Performance guarantee — names corrected (client review)
# ──────────────────────────────────────────────────────────────────────────
def build_performance_guarantee():
    doc = Document()
    s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(10.5)

    title(doc, "CORPORATE PERFORMANCE GUARANTEE",
          "Galascope Ltd batch \u2014 Draft for client review")

    body(doc, "DRAFT FOR CLIENT REVIEW \u2014 Not valid until signed and sealed by Linyang. This is a 5% "
              "corporate performance guarantee from the manufacturer \u2014 a direct corporate obligation, not "
              "a bank guarantee; the bank instruments are the separate Advance Payment Guarantees.",
         bold=True, amber=True)
    doc.add_paragraph()

    body(doc, "Place and date:  [\u25cf]")
    body(doc, "Reference No.:  [\u25cf]")
    body(doc, "Beneficiary:  Galascope Ltd (HE 303759), Karaiskaki 6, City House, 3032 Limassol, Cyprus, "
              "and/or its project-finance security agent; Lighthief Cyprus Ltd (HE 477423) is named as "
              "co-beneficiary for any obligation owed to the Distributor (\u201cBeneficiary\u201d)")
    body(doc, "Guarantor:  Jiangsu Linyang Energy Storage Technology Co., Ltd (\u201cGuarantor\u201d / \u201cLinyang\u201d)")
    doc.add_paragraph()
    body(doc, "CORPORATE PERFORMANCE GUARANTEE", bold=True)
    doc.add_paragraph()

    clause(doc, "(1)", f"Background. The Beneficiary and the Guarantor have entered into the {CONTRACT_REF} for "
           f"supply of BESS equipment for Galascope 1 (5 MW / 20.06 MWh) and Galascope 2 (2.5 MW / 10.03 MWh), "
           f"Famagusta, Cyprus. The Component A (equipment) price is approximately {fmt_eur(TOTAL_COMP_A)}.")
    clause(doc, "(2)", f"Guarantee amount. The Guarantor irrevocably undertakes to pay the Beneficiary any sum "
           f"or sums not exceeding in the aggregate {fmt_eur(PB_AMOUNT)} (five percent (5%) of the Component A "
           f"equipment price) (\u201cGuaranteed Amount\u201d).")
    clause(doc, "(3)", "Nature. A direct corporate obligation of the Guarantor \u2014 not a bank guarantee, letter of "
           "credit, or URDG 758 demand guarantee.")
    clause(doc, "(4)", "Purpose. Security for the Guarantor's performance under the Contract: (a) delivery of "
           "conforming equipment CIF Limassol; (b) defects in materials and workmanship during the Defects "
           "Liability Period; and (c) related OEM obligations under the supply terms.")
    clause(doc, "(5)", "Issuance. Delivered no later than one (1) month before first shipment, signed by an "
           "authorised officer and affixed with the Guarantor's company seal.")
    clause(doc, "(6)", "Validity. From execution until the end of the Defects Liability Period, being three (3) "
           "months after PAC. Actual dates are set by the Delivery Schedule confirmed at Connection Terms; "
           "no fixed calendar date applies at signing.")
    clause(doc, "(7)", "Demand. The Beneficiary may demand if the Guarantor fails to perform a material obligation "
           "after written notice of default and a thirty (30) day cure period (immediately on insolvency or "
           "abandonment); payment within thirty (30) days.")
    clause(doc, "(8)", "Backing. Supported by the Guarantor's warranty reserve, product "
           "liability insurance of EUR 5,000,000 per occurrence (AXA), and the Direct Warranty Undertaking. "
           "No parent-company guarantee is provided unless separately agreed.")
    clause(doc, "(9)", "Relationship to APG. Separate from, and additional to, the Advance Payment Guarantee. "
           "Failure to deliver this guarantee does not reduce the APG requirement.")
    clause(doc, "(10)", "Assignment. Galascope is a direct named Beneficiary, so no assignment step is required "
           "for Galascope to claim. Galascope may further assign this guarantee or its proceeds to a "
           "project-finance lender or security agent on written notice (effective immediately on the "
           "Distributor's insolvency, liquidation or cessation of business); assignment does not increase "
           "the Guaranteed Amount or extend the Expiry Date.")
    clause(doc, "(11)", "Governing law. Governed by the laws of the People's Republic of China, with disputes to "
           "the Shanghai International Arbitration Center (SHIAC) in English, consistent with the OEM Direct "
           "Warranty Undertaking and the manufacturer's standard supply terms.")

    doc.add_paragraph()
    body(doc, "For and on behalf of Jiangsu Linyang Energy Storage Technology Co., Ltd")
    body(doc, "Name / Title / Signature / Company seal / Date:  [\u25cf]")
    doc.add_paragraph()
    body(doc, "Lighthief Cyprus Ltd \u00b7 HE 477423 \u00b7 office@lighthief.com \u00b7 +357 99 164 158 \u00b7 solarfarms.cy",
         grey=True, size=8)

    path = OUT / "Performance-Guarantee-Galascope-CLIENT-REVIEW-jun2026.docx"
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    build_dwu()
    build_performance_guarantee()
