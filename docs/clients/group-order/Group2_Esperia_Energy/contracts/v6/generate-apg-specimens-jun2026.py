#!/usr/bin/env python3
"""Two client-clean APG specimens for Galascope (advance + pre-shipment).

Both name Galascope Ltd as beneficiary, use neutral descriptions (no reference
to the OEM Sales Contract, no CIF/25% derivation), state EUR amounts, URDG 758,
valid to PAC / 12 months after delivery. No internal margin notes.

Outputs into CLIENT-PACKAGE/:
  05a-Advance-Payment-Guarantee-No1-advance-specimen.docx       (EUR 705,792)
  05b-Advance-Payment-Guarantee-No2-preshipment-specimen.docx   (EUR 1,411,585)
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
GREY = RGBColor(0x40, 0x40, 0x40)

V6 = Path(__file__).resolve().parent
PKG = V6 / "CLIENT-PACKAGE"

ADVANCE_AMOUNT = "EUR 705,792"
PRESHIP_AMOUNT = "EUR 1,411,585"


def _run(p, text, *, bold=False, italic=False, size=10.5, color=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color


def title(doc, text, sub):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, bold=True, size=13, color=NAVY)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, sub, size=10.5, color=GOLD)
    doc.add_paragraph()


def body(doc, text, *, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    _run(p, text, bold=bold, italic=italic, color=color)


def clause(doc, num, text):
    body(doc, f"({num})  {text}")


def build(filename, no, label, amount, tranche_words):
    doc = Document()
    s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(10.5)

    title(doc, f"SPECIMEN \u2014 ADVANCE PAYMENT GUARANTEE No. {no}",
          f"Beneficiary: Galascope Ltd \u00b7 {label} \u00b7 for bank pre-confirmation")

    body(doc, "SPECIMEN FOR CLIENT REVIEW \u2014 not an issued instrument. The issuing bank will "
              "release the final guarantee on its own letterhead.", bold=True)
    doc.add_paragraph()

    body(doc, "From:  [Issuing bank \u2014 full name, registered office, SWIFT BIC] (the \u201cBank\u201d)")
    body(doc, "To:    Galascope Ltd (HE 303759), Karaiskaki 6, City House, 3032 Limassol, Cyprus (registered "
              "office), and/or [Security Agent / Alpha Bank Cyprus] as its project-finance security agent "
              "(the \u201cBeneficiary\u201d)")
    body(doc, "Issuing Date:  [\u25cf]      Guarantee No.:  [\u25cf]")
    doc.add_paragraph()
    body(doc, "ADVANCE PAYMENT GUARANTEE", bold=True)
    doc.add_paragraph()

    clause(doc, "1", "Whereas the Beneficiary is procuring the supply, delivery and installation of a "
           "grid-connected BESS for Galascope 1 (5 MW / 20.06 MWh) and Galascope 2 (2.5 MW / 10.03 MWh), "
           "Famagusta, Cyprus (the \u201cProject\u201d), and has agreed to make the "
           f"{tranche_words} for the equipment supply for the Project.")
    clause(doc, "2", "We irrevocably guarantee to pay you, on your first written demand, any sum or sums "
           f"up to a maximum aggregate amount of {amount} (the \u201cGuaranteed Amount\u201d), being one "
           f"hundred percent (100%) of the {label.lower()} for the equipment supply for the Project.")
    clause(doc, "3", "We shall pay within five (5) calendar days of receipt of your complying written "
           "demand stating that the equipment has not been delivered to Site and/or that the "
           f"corresponding {label.lower()} has not been refunded when due, without set-off or deduction. "
           "All issuing-bank charges are for the account of the applicant.")
    clause(doc, "4", "Our liability reduces by each amount paid under this guarantee. The Guaranteed "
           "Amount shall not be reduced merely on presentation of shipping documents or a bill of lading.")
    clause(doc, "5", "This guarantee comes into force on receipt by the applicant of the corresponding "
           "payment and remains valid until, and any demand must be received by the Bank on or before, "
           "the earlier of: (i) issuance of the Provisional Acceptance Certificate (PAC) for the Project; "
           "or (ii) twelve (12) months after delivery of the equipment to Site; upon issuance of PAC this "
           "guarantee is released (the \u201cExpiry Date\u201d).")
    clause(doc, "6", "Any demand must be presented in writing to the Bank, or by authenticated SWIFT, "
           "your bank confirming the signatories are authorised to bind the Beneficiary.")
    clause(doc, "7", "The Beneficiary may assign this guarantee or its proceeds to a project-finance "
           "lender or security agent without the Bank's further consent.")
    clause(doc, "8", "This guarantee is subject to the Uniform Rules for Demand Guarantees, ICC "
           "Publication No. 758 (URDG 758). [Governing law / jurisdiction: as required by the Bank \u2014 "
           "to be confirmed.]")
    doc.add_paragraph()
    body(doc, "For and on behalf of the issuing Bank:  [authorised signatures / SWIFT authentication]")
    doc.add_paragraph()
    body(doc, "Specimen prepared by Lighthief Cyprus Ltd \u00b7 HE 477423 \u00b7 office@lighthief.com \u00b7 "
              "+357 99 164 158 \u00b7 solarfarms.cy \u2014 for confirmation by the issuing bank.",
         italic=True, color=GREY)

    PKG.mkdir(parents=True, exist_ok=True)
    doc.save(str(PKG / filename))
    print(f"wrote {filename}")


if __name__ == "__main__":
    build("05a-Advance-Payment-Guarantee-No1-advance-specimen.docx",
          "1", "Advance payment", ADVANCE_AMOUNT,
          "advance payment")
    build("05b-Advance-Payment-Guarantee-No2-preshipment-specimen.docx",
          "2", "Pre-shipment payment", PRESHIP_AMOUNT,
          "pre-shipment payment (after a passed Factory Acceptance Test)")
