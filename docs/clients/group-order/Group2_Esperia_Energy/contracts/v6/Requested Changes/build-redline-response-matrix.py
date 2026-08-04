#!/usr/bin/env python3
"""Build the colour-coded Redline Response Matrix (.docx) for the client.

Hybrid legend (per client request):
  - Timotheos / technical review items are coloured BY HANDLING:
        RED    = included as requested
        PURPLE = included with our drafting changes
        ORANGE = included with our counter-proposal
        GREY   = not yet included / declined
  - Dino Constantinou items = TEAL (author colour; status noted in text)
  - Anastasios (lawyer) comments = BLUE (author colour; status noted in text)

Sources:
  - 01-EPC-...CLIENT_REDLINE.docx  (Timotheos/"Client Legal/Technical Review" + Dino)
  - 03-LTSA-...CLIENT_REDLINE.docx (Timotheos/"Client Legal/Technical Review" + Dino)
  - Anastasis comments - EPC agreement.docx
Cross-referenced against the generated v6.1 package to mark handling status.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
OUT = HERE / "Redline-Response-Matrix-jul2026.docx"

# Brand
NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY_T = RGBColor(0x40, 0x40, 0x40)

# Functional status/author colours
RED = RGBColor(0xC0, 0x00, 0x00)      # T: included as requested
PURPLE = RGBColor(0x70, 0x20, 0x9C)   # T: included with our changes
ORANGE = RGBColor(0xC0, 0x60, 0x00)   # T: included with our counter
GREY = RGBColor(0x80, 0x80, 0x80)     # T: not yet included / declined
TEAL = RGBColor(0x0F, 0x76, 0x8A)     # Dino
BLUE = RGBColor(0x1F, 0x49, 0xC0)     # Anastasios (lawyer)

# status -> (label, colour) for Timotheos items
T_STATUS = {
    "asis":   ("INCLUDED — as requested", RED),
    "changes":("INCLUDED — with our drafting changes", PURPLE),
    "counter":("INCLUDED — with our counter", ORANGE),
    "none":   ("NOT YET INCLUDED", GREY),
}


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def run(p, text, *, bold=False, italic=False, size=9, color=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return r


def item_color(author, status):
    if author == "D":
        return TEAL
    if author == "A":
        return BLUE
    return T_STATUS[status][1]


def status_label(author, status, note):
    if author == "D":
        return ("DINO", TEAL, note)
    if author == "A":
        return ("ANASTASIOS", BLUE, note)
    lab, col = T_STATUS[status]
    return (lab, col, note)


# ---- DATA --------------------------------------------------------------
# (ref, ask, author[T/D/A], status[asis/changes/counter/none], our_response)

EPC = [
    ("1A.3A", "No payment due until all companion documents delivered in final, signed, enforceable form satisfactory to Client/lender.", "T", "changes",
     "Included as Annex V6 §1A.5(e) (companion documents + LTSA as condition precedent to the advance)."),
    ("4.x", "Contractor is the Client's single point of responsibility for any defect/delay/performance/integration/safety/warranty issue, without limiting direct rights v. OEM/banks/insurers.", "T", "asis",
     "Accepted — consistent with our single-counterparty model."),
    ("4.4(f)", "Contractor responsible end-to-end for BMS/PCS/SCADA/EMS interface, command execution and data integrity through DLP; affiliate/subcontractor failure no excuse.", "T", "asis",
     "Accepted verbatim — §4.4(f) added."),
    ("4.6", "Excluded items only excluded if expressly in Schedule A with fixed boundary/price/party; ambiguity construed in favour of turnkey, against Contractor.", "T", "changes",
     "Now drafted — see EPC §4.6: exclusions must be expressly identified in Schedule A with a fixed technical boundary, price impact and party; genuine ambiguity construed reasonably in favour of turnkey delivery."),
    ("5A.5", "No EOT/cost/warranty-limit unless Contractor proves Client breach directly and materially caused the loss, and it mitigated.", "T", "counter",
     "Now drafted — see EPC §5A.5: Contractor must show a Client breach directly and materially caused the loss (plus its own mitigation) before relying on any EOT/cost/warranty relief; conversely, for operational non-compliance (low-SOC, site-access denial, loss of connectivity) the Client bears the burden of showing it complied."),
    ("6.1(i)", "Confirmed Price Certificate to include full index/OEM/FX evidence + workbook; downward adjustments passed through in full; no re-anchoring without Client consent; keep Jan-2026 price-lock.", "T", "counter",
     "Re-anchored to the signed LOI Jan-2026 basis with two-way (incl. downward) adjustment from Effective Date. Evidence-pack wording not yet added."),
    ("7.1A", "No milestone due unless conditions met, no default, test cert accepted by Independent Engineer, and all guarantees/insurance/docs delivered.", "T", "counter",
     "Counter — §7.1A limits the IE sign-off to the PAC payment only; advance & pre-shipment unaffected."),
    ("7.4(e)", "Client may set off Delay/perf LDs, defect costs, third-party completion, warranty amounts, insurance deductibles against any unpaid milestone/PAC/Retention.", "T", "counter",
     "Counter — §7.4(e) added, limited to amounts ascertained/agreed or determined by adjudicator/expert."),
    ("7.5(f)", "Risk passes only after unloading + visual inspection + serial/qty verification + no transit damage + CAR/marine respond without gap; title not to prejudice APG/warranty/rejection.", "T", "changes",
     "Included with our changes — §7.5: risk on unloading; title on earlier of PAC or 12 months (aligns to APG cover)."),
    ("8.4.2A", "Delivery Schedule at Connection Terms binding with latest dates; Target PAC ≤ [●] months after order trigger; Client withdrawal if exceeded.", "T", "counter",
     "Counter — §8.4.8 makes the Delivery Schedule binding but preserves OEM force-majeure flow-through (§8.4.7A). Now drafted — see EPC §1A.7: PAC Long-Stop fixed at fifteen (15) months from receipt of Connection Terms; if the confirmed Target PAC Date falls later than the Long-Stop, the Client may withdraw with a full refund of any advance received (objective threshold, replacing a subjective standard)."),
    ("8.4.4A", "Delay-LD cap does not limit completion, defect replacement, APG recovery, third-party indemnity, or safety/grid/fire/cyber rectification.", "T", "changes",
     "Included with changes — cap does not excuse completion or mandatory safety/grid-code rectification; APG reference kept separate."),
    ("8.4.8", "Ordinary production/raw-material/allocation/port/financing/customs/logistics delays give no EOT or LD relief.", "T", "counter",
     "Counter — hardened milestone tracking accepted, but OEM-declared force majeure and connection-term delays remain valid EOT grounds."),
    ("9.1A", "PAC conditional on an agreed FAT/SAT acceptance matrix (capacity, RTE, SCADA/DSO, BMS, EMS, E-stop, fire, HVAC, VSG, black start, harmonics, reactive power, frequency).", "T", "changes",
     "Included with changes — FAT reinstated and PAC criteria expanded; full matrix to be finalised in Schedule A."),
    ("9.1B", "If performance below guarantees but above PAC floor, PAC does not waive shortfall; Contractor augments/repairs/price-reduces/pays LDs at Client election.", "T", "counter",
     "Counter — §9.1B added; shortfall remedy capped at 10% of Component A, back-to-back with OEM."),
    ("9.2A", "DLP = 12 months from PAC; fallback 3 months punch-list + 12-month latent-defect correction.", "T", "counter",
     "Counter — DLP stays 3 months (cannot move). FAC added at DLP end (§9.2A)."),
    ("9.2 (DLP)", "DLP three (3) → six (6) months.", "D", "counter",
     "Not accepted — DLP stays 3 months (retention & OEM performance guarantee track the DLP)."),
    ("9.2 (FAC)", "Reinstate a Final Acceptance Certificate at DLP end; retention released only on FAC.", "T", "asis",
     "Accepted — FAC reinstated as Annex V6 §9.2A."),
    ("10.1A", "Unified warranty to expressly cover design/latent/integration/firmware/grid-code/safety/defective-commissioning; remedies incl. parts/labour/travel/lifting/logistics/recommissioning.", "T", "counter",
     "Counter — §13.3A uncaps latent + safety/grid-code only; general firmware/design stay OEM-backed & capped."),
    ("10.5A", "Warranty voided only to extent Contractor proves Client act caused it; grid/DSO/EMS/BMS/monitoring/FM faults not to void warranty.", "T", "asis",
     "Accepted — §10.5A EMS-affiliate carve-out added (DISPERON/Lighthief-group failures do not void)."),
    ("10.6A", "Annual performance testing under agreed protocol, witnessed by Client/IE; SOH/RTE disputes to independent expert / ISO 17025 lab, not OEM/Contractor alone.", "T", "changes",
     "Now drafted — see EPC §10.6A: annual testing under an agreed protocol, may be witnessed by Client/IE; SOH/AC-capacity/RTE disputes go to an independent expert or ISO/IEC 17025 lab, not OEM/Contractor alone."),
    ("10.7A", "Execute LTSA as CP to PAC/COD; availability per Park not only group; availability LDs not sole remedy for wilful/repeated/safety failures.", "T", "counter",
     "Partly — LTSA is a simultaneous-execution CP and sole-remedy carve-out added; availability stays GROUP level (counter)."),
    ("10.8A", "DWU irrevocable, enforceable by Client/lender, survives OEM/Distributor termination, waiver of defences, no unilateral Warranty-Manual amendment.", "T", "changes",
     "Now drafted as a procurement covenant — see EPC §10.8A: Contractor must procure a DWU that is irrevocable, directly enforceable by Client/lender, survives OEM/Distributor termination, waives Contractor/Distributor-dispute defences, and bars unilateral Warranty-Manual amendment. Linyang counter-signature on the DWU itself is still pending."),
    ("10.9B", "Each APG from a bank acceptable to Client/lender, preferably EU-confirmed; partial/multiple drawings; independent of disputes; auto-extend if PAC/delivery delayed.", "T", "changes",
     "Now drafted as a procurement covenant \u2014 see EPC \u00a710.9B: Contractor must procure APGs from a Client/lender-acceptable bank (EU-confirmed if issued outside the EU), permitting partial/multiple drawings, independent of EPC disputes, and auto-extending if PAC/delivery is delayed. Bank of Communications actual re-issue matching this wording is still pending."),
    ("11.1", "97% availability / availability LDs / extended warranty to be incorporated as mandatory conditions of EPC + LTSA.", "T", "counter",
     "Counter — kept in the LTSA (executed simultaneously as CP), not merged into the EPC body."),
    ("11.5", "All project documents read as one suite; on inconsistency the interpretation giving the Client greater protection prevails.", "T", "asis",
     "Accepted as requested — Annex V6 §11.5 (most-favourable-to-Client)."),
    ("12.6", "Force Majeure not to excuse amounts due, pre-event warranty, APG extension, insurance, mitigation, or electronic document delivery.", "T", "changes",
     "Now drafted — see EPC §12.6: Force Majeure does not excuse amounts already due, pre-event warranty obligations, APG-extension obligations, insurance maintenance, mitigation, or electronic document delivery."),
    ("13.3A", "Caps not to apply to design/latent/firmware/cyber/regulatory/grid-code defects, insurance/APG/warranty proceeds.", "T", "counter",
     "Counter — §13.3A uncaps latent + safety/grid-code only; firmware/design remain OEM-backed & capped."),
    ("14.5", "Insurance evidence before first payment; name Client/lender as additional insured/loss payee; waiver of subrogation; specify deductibles/territory/marine-CAR interface/LEG3/limits/expiry; 30-day cancellation notice.", "T", "changes",
     "Included with changes — §14.6 recast as a procurement covenant (we hold a non-binding CEAR binder; endorsements to be procured); §14.6(e) confirms PI €2M covers our-negligence commissioning failure."),
    ("16.3A", "Lender step-in / cure / direct agreement / notice / enforce APGs, bonds, warranties, insurance without Contractor/OEM consent.", "T", "changes",
     "Now drafted as a framework clause — see EPC §16.3A: lender/security agent gets step-in, cure, notice and enforcement/proceeds rights over the APGs, OEM PG, warranties and insurance; Contractor must enter a direct agreement on customary terms. Alpha Bank's specific direct-agreement wording is still pending."),
    ("18A", "Cyber security: IEC 62443, MFA, RBAC, audit logs, firmware integrity, vulnerability mgmt, 24h incident reporting, data ownership, raw export, open protocols, EMS migration.", "T", "changes",
     "Now drafted as a framework clause — see EPC §18A: IEC 62443-consistent controls (secure remote access, MFA, RBAC, audit logs, firmware integrity, vulnerability management, 24-hour incident reporting), Client data ownership and raw export. Formal IEC 62443/NIS2 certification is still pending confirmation with DISPERON."),
    ("19.12", "Deliver all certs, as-builts, settings, grid studies, SCADA/IEC104 point lists, Modbus/IEC61850 maps, fire/C5 certs, serials, cell batch traceability, O&M manuals.", "T", "asis",
     "Accepted — §19.12 added."),
    ("Sch A", "Technical Specification frozen before order; all TBCs closed (VSG, black start, C5-M, 45°C no-derating, AC usable capacity, RTE, harmonics, reactive power, data interfaces, register maps, spares, FAT/SAT criteria).", "T", "asis",
     "Accepted — Schedule A spec-freeze note added."),
    ("1A.5(e)", "No payment unless each companion document executed/dated/sealed/enforceable/consistent/placeholder-free and approved by Client advisers.", "T", "asis",
     "Accepted — §1A.5(e) added."),
    ("1A.5B", "FAT deemed passed only on written Client/IE acceptance; critical/safety findings remedied before shipment; no deferral of material punch items.", "T", "changes",
     "Included with changes — FAT acceptance folded into §9.1A / Schedule A."),
    ("1A.7A", "Client withdrawal right if schedule exceeds longstop, price cert disputed, APG doesn't cover revised programme, or OEM materially changes config/warranties.", "T", "changes",
     "Now drafted — see EPC §1A.7A: Client may withdraw without liability if (a) Target PAC Date is beyond the §1A.7 PAC Long-Stop, (b) the Confirmed Price Certificate is disputed in good faith and unresolved, (c) APG validity does not cover the revised programme, or (d) the OEM materially changes configuration/warranties."),
    ("10.9(f)", "APG demand events to include non-delivery/late/insolvency/abandonment/default/non-refund/non-conforming equipment/rejection/failure-to-extend/bad-bank.", "T", "changes",
     "Included with changes — APG validity to PAC and demand mechanics reflected in specimens; full event list to finalise with bank."),
    ("13.6", "Manufacturing-defect carve-out to expressly include systemic/latent/design/firmware/batch/safety-recall/serial defects across PCS/BMS/HVAC/fire/transformer/MV/cells/ancillaries.", "T", "changes",
     "Included with changes — carve-out reflected via DWU §5.3 and EPC §13.5."),
    ("Grid-forming", "VSG/black-start activation, firmware, licences, commissioning included in Contract Price unless a post-Effective-Date DSO requirement materially exceeds known scope.", "T", "asis",
     "Accepted verbatim — replaces the old 'confirmed by amendment' wording."),
]

LTSA = [
    ("2.4", "LTSA part of one integrated Project Document suite; interpreted consistently so there is no gap in responsibility.", "T", "asis",
     "Now drafted verbatim — see LTSA §2.4: single integrated Project-Document suite (EPC, Technical Agreements, DWU, OEM Warranty Terms, APGs, OEM PG, EMS Agreement), interpreted consistently so there is no gap in responsibility."),
    ("Defs", "Add 'Restoration Time', 'Resolution Time', 'Repeated Service Failure' definitions.", "T", "asis",
     "Now drafted verbatim — see LTSA 'Definitions — Restoration and Resolution': Restoration Time, Resolution Time and Repeated Service Failure (3+ Critical breaches of similar cause within a rolling 12 months) all defined."),
    ("4.1A", "Default to Tier C from PAC for both Parks (corrective maintenance, local spares, 97% availability, priority response, warranty preservation, annual testing).", "T", "asis",
     "Now drafted — see LTSA §4.1A: unless the Client expressly selects otherwise in Schedule 2, both "
     "Galascope 1 and Galascope 2 default to Tier C from PAC (corrective maintenance, local spare parts, "
     "97% availability guarantee, priority response, warranty-preservation services and annual performance "
     "testing), consistent with the Schedule 2 all-in €1,740/MWh/yr Tier C fee."),
    ("4.3A", "Corrective Maintenance to include all diagnosis/labour/travel/removal/reinstall/recommissioning/config restoration/testing/documentation.", "T", "changes",
     "Now drafted — see LTSA §4.3A: Corrective Maintenance covers diagnosis, labour, travel within Cyprus, removal, installation, recommissioning, configuration restoration, testing and incident documentation, except costs expressly excluded under Schedule 2."),
    ("4.4A", "Exclusions apply only to extent event directly caused the loss; SP bears burden and keeps monitoring/diagnosis/response.", "T", "changes",
     "Now drafted — see LTSA §4.4A: exclusions apply only to the extent the excluded event directly caused the loss; the Service Provider continues monitoring, diagnosis, emergency response, mitigation and warranty-claim support regardless."),
    ("Civil works", "Civil works / structural / site infra are the responsibility of Lighthief.", "D", "asis",
     "Agreed — civil BOP is Lighthief scope."),
    ("SOH excl?", "Why is battery degradation beyond Guaranteed SOH excluded?", "D", "none",
     "Answered: SOH beyond the guarantee is an OEM matter; within the guarantee it is covered. Clarify in text."),
    ("5.2A", "Minimum monitored parameters (cell V/T, SOC/SOH, PCS, transformer/MV, HVAC, fire, gas, E-stop, aux, comms, EMS command, DSO telemetry).", "T", "asis",
     "Now drafted — see LTSA §5.2A: minimum monitored parameters cover cell V/T, SOC/SOH, PCS, transformer/MV, HVAC, fire/gas, E-stop, aux supply, comms, EMS command execution and DSO/SCADA telemetry."),
    ("5.5", "Include Client (or Client representatives) training on the Monitoring Platform/portal.", "T", "asis",
     "Now drafted — see LTSA §5.5: Service Provider to provide Client (or up to 3 nominated representatives) portal training at no charge — one session at/after PAC, plus an annual refresher and after material platform changes."),
    ("5.6A", "Richer reporting: raw extracts, per-Park + aggregate availability, SOH trend, cell imbalance, thermal deviation, firmware, punch items, spares, warranty notices, cyber events.", "T", "asis",
     "Now drafted — see LTSA §5.6A: reporting includes raw-data extracts, per-Park and aggregate availability, SOH trend, cell imbalance, thermal deviation, firmware version, open punch items, spares usage, warranty notices and cyber/security events."),
    ("5.7A", "Retain data/logs/reports/SOH certs 15 years + 5 years post-termination; continuous Client access & export (CSV/XLSX/JSON/API).", "T", "changes",
     "Included with changes — 15-year retention + client access added; export-format list to finalise."),
    ("5.8A", "Client owns all operational/technical/performance/market/SCADA/BMS/PCS/SOH data; SP not to use identifiable data without consent.", "T", "changes",
     "Now drafted — see LTSA §5.8A: Client owns all operational/technical/performance/dispatch/SCADA/BMS/PCS/SOH/availability data; Service Provider may not use identifiable data for benchmarking, training or third-party purposes without consent."),
    ("6.1A", "SP performs all maintenance needed to preserve OEM/DWU/SOH/corrosion/fire/PCS-MV warranties + insurance; liable for loss if it fails.", "T", "asis",
     "Now drafted — see LTSA §6.1A: Service Provider must perform all maintenance required to preserve the OEM warranty, DWU, SOH, corrosion, fire-system and PCS/MV warranties and insurance cover, and is liable for loss of warranty, claim rejection, downtime or degradation caused by its own failure to perform or document required maintenance."),
    ("6.2", "Annual maintenance schedule provided to Client and confirmed 15 days before / Client confirms 5 days before.", "D", "asis",
     "Agreed — scheduling notice accepted."),
    ("6.3A", "Warranty-Preservation File per Park (checklists, photos, torque, thermal images, grounding, firmware, alarm/SOC logs, fire tests, spares, OEM comms).", "T", "asis",
     "Now drafted — see LTSA §6.3A: a Warranty-Preservation File per Park (checklists, photographs, torque records, thermal images, grounding/insulation results, firmware records, alarm/SOC logs, environmental records, fire-system tests, spare-part records, OEM communications) in a format acceptable to the OEM (Linyang)."),
    ("6.8A", "Scheduled Downtime ≤ 48h per Park without Client approval; over-limit or <14-day notice counts as Unavailable.", "T", "changes",
     "Included with changes — scheduled downtime set to 2 days (48h)/Park/yr with notice; wording to align to 6.8A."),
    ("7.x fee", "No increase in Service Fees (delete CPI+2% escalation).", "D", "asis",
     "Agreed — fees fixed for the Initial Term; post-term by written agreement only."),
    ("7.4A", "No extra charge for services to comply with OEM manuals, preserve warranties, corrective maintenance in Tier, warranty-claim prep, standard exports, or to remedy SP breach.", "T", "asis",
     "Now drafted — see LTSA §7.4A: no additional charge for services reasonably required to comply with OEM manuals, preserve warranties, perform in-Tier corrective maintenance, prepare warranty claims, deliver standard data exports, or remedy a Service Provider breach."),
    ("8.1A", "Redundant monitoring/alerting/comms; auto-alerts for low SOC/voltage, high temp, HVAC/cooling, fire/gas, PCS trip, comms loss; monitoring failure = Major (or Critical) Alert.", "T", "asis",
     "Now drafted — see LTSA §8.1A: redundant monitoring/alerting/communication channels with automatic alerts for low SOC, low cell voltage, high temperature, HVAC/cooling fault, fire/gas alarm, PCS trip and comms loss; a monitoring-platform failure is itself a Major Alert (Critical where safety-related)."),
    ("8.2A", "Maximum Restoration & Resolution Times (Critical: plan 4h, restore 24h where possible, resolve 5 BD unless OEM parts; Major: 1 BD / 5 BD / 15 BD; Minor: next visit or 30 days). Deviation needs a written remediation plan accepted by Client.", "T", "counter",
     "Now drafted — see LTSA §8.2A: Critical/Major/Minor Restoration & Resolution times as requested, with (a) deviation acceptance 'not to be unreasonably withheld or delayed'; (b) the clock stops for OEM part lead-times, Force Majeure and Client/site-access delays; (c) these are service targets — any §8.5 service credit for a breach is not additional to the Availability LDs for the same downtime."),
    ("8.5 SLA times", "Critical on-site within 2h; Major on-site 2h; Minor on-site 1 BD (faster than drafted).", "D", "counter",
     "Counter — immediate remote response; Critical on-site 12h; Major 6h; Minor 5 BD (one Cyprus team). Response times updated in v6.1."),
    ("8.5 credits", "Service-level credits per MW/day of affected capacity (min €500/€250 per BD); Dino: 30h thresholds.", "D", "counter",
     "Counter — credit values to be set as €/MW/day with a cap; not additional to availability LDs. Figures to confirm."),
    ("8.5(e)", "Service credits NOT sole remedy — without prejudice to availability LDs, repeated breach, wilful default, safety, spares, step-in, EPC/DWU claims.", "T", "changes",
     "Included with changes — sole-remedy language removed for wilful default / gross negligence / safety; LDs remain primary for ordinary shortfall."),
    ("8.6A", "Major Incident Procedure (fire/thermal/repeated trip/outage >24h/cyber/grid-code/>10% Park) — notify immediately, incident report 24h, RCA.", "T", "asis",
     "Now drafted — see LTSA §8.6A: Major Incident Procedure for fire, thermal events, repeated trips, forced outage >24h, cyber incidents, grid-code breaches or unavailable capacity >10% of a Park — immediate notification, 24-hour initial incident report, and a root-cause analysis."),
    ("8.7A", "24/7/365 response for Critical/safety/fire/warranty-void/cyber/full-outage; on-site for safety events not limited to Business Days.", "T", "changes",
     "Now drafted — see LTSA §8.7A: 24/7/365 response for Critical Alerts, safety incidents, fire alarms, warranty-voiding alerts, cyber incidents and full Park outages; on-site attendance for safety-critical events is not limited to Business Days 'where local personnel or emergency-contractor attendance is reasonably available' (the qualifier we recommended)."),
    ("Alert class", "Critical = OUTAGE OF ANY equipment (ack immediately, auto-escalate 15 min, mgmt 1h); Major >5%; Minor no perf loss.", "D", "counter",
     "Counter — Critical ack 30 min, auto-escalate 30 min, mgmt 2h (single Cyprus team). Definitions accepted."),
    ("9.1A/9.2B", "97% availability measured per Park individually AND aggregate; group aggregation must not mask a single-Park breach.", "T", "counter",
     "Counter — availability stays GROUP-level (original agreement) with a rationale note; 50% LD-cap increase gives per-Park compensation."),
    ("9.3A", "Excluded Hours narrowly construed; EMS/third-party failures excluded only if SP proves outside scope & not its config/monitoring/cyber; no exclusion for warranty-claim waiting.", "T", "changes",
     "Included with changes — §9.3(e) EMS-affiliate carve-out added; full narrow-construction wording to add."),
    ("9.4 downtime", "Scheduled downtime 240h/10 days → 48h/2 days.", "D", "asis",
     "Agreed — 2 days (48h)/Park/yr."),
    ("9.5 LD cap", "Availability LD cap 20% → 100% of annual Park fee; cap not to apply to wilful/gross/repeated/safety/spares failures.", "T", "counter",
     "Counter — LD cap raised to 50% of annual Service Fee per Park (from 20%); carve-outs accepted."),
    ("9.x LD rate", "LD rate 'not as presented originally (was €30/day/MWh if <97%)'.", "D", "changes",
     "Schedule 4 LD formula clarified (whichever is higher of % table vs €30/day/MWh, subject to cap)."),
    ("9.8", "Chronic underperformance (<95% any year, or <97% two years, or 3+ Critical breaches/12m) → remediation plan, independent audit at SP cost, extra spares.", "T", "changes",
     "Now drafted — see LTSA §9.8: Availability below 95% in any year, below 97% for two consecutive years, or 3+ Critical Alert SLA breaches in a rolling 12 months triggers a remediation plan, an independent technical audit at the Service Provider's cost, and additional spare stock."),
    ("10.2/Sch5", "SOH degradation table internally inconsistent (Yr9 80% vs Yr8 79%, Yr10 79.58%); correct and align with OEM Technical Agreement & DWU.", "T", "asis",
     "DONE — §10.2 and Schedule 5 corrected to the OEM's own 0.25P, 1-cycle/day curve (confirmed by direct "
     "read of Linyang's Power Atlantic 5MWh Degradation Curve PDF): Yr1 94.62% / Yr2 91.77% / Yr3 89.91% / "
     "Yr4 88.00% / Yr5 86.78% / Yr6 84.97% / Yr7 83.83% / Yr8 82.25% / Yr9 81.06% / Yr10 79.58% / "
     "Yr15 73.61%. Cycle life 8,000; energy density 175 Wh/kg. Note: Yr5/Yr15 (86.78%/73.61%) are higher "
     "than the values in the previous draft (86.26%/72.45%) — this is an OEM-confirmed correction, not a "
     "new client concession; it matches the figures Stelios originally requested. The OEM Direct Warranty "
     "Undertaking (our own draft, doc 04) was corrected in the same pass so the DWU stays fully back-to-back "
     "with the EPC/LTSA — sent to Linyang for confirmation before signing."),
    ("10.3A", "SOH/AC-capacity/RTE testing under agreed protocol; Client/lender IE may witness; SP bears cost if extra test confirms underperformance.", "T", "asis",
     "Now drafted — see LTSA §10.3A: SOH/AC-capacity/RTE testing under a pre-agreed protocol; the Client and lender's IE may witness; the Service Provider bears the cost of any additional Client-requested test that confirms underperformance or a warranty issue."),
    ("10.4A", "SOH remedy to include all parts/labour/shipping/customs/craneage/install/recommission/testing.", "T", "counter",
     "Counter — accept parts+labour+recommissioning; customs & craneage on OEM warranty returns to be OEM-funded."),
    ("10.4 SOH LD", "SOH systemic trigger 5% → 3%; SOH LD to full lost income at market prices; auto-cancel LTSA option.", "D", "counter",
     "Systemic trigger 3% accepted. SOH LD = full restoration cost, capped (not open-ended market lost income). "
     "Client unilateral auto-cancel-LTSA option on SOH shortfall not accepted — the Client's termination-for-"
     "convenience right (§13.3, 90 days' notice) already gives an exit path, and tying LTSA survival to a single "
     "metric would undermine the availability/warranty protections that depend on the LTSA staying in force."),
    ("10.5 SOH cap", "Max SOH LD 50% → 100% of Tier-C + warranty-extension fees, or market cost where SP-caused.", "T", "counter",
     "Counter — kept at a capped restoration basis; open-ended market cost not accepted."),
    ("10.6A", "Annual evidence OEM warranty reserve/spares/support valid; notify Client within 5 BD of any OEM dispute/insolvency/recall.", "T", "asis",
     "Now drafted — see LTSA §10.6A: the Service Provider must give annual evidence that the OEM warranty reserve, extended warranty, spare-part access and OEM support remain valid, and notify the Client within 5 Business Days of any OEM dispute, insolvency risk, supply restriction, recall or firmware issue."),
    ("10.7 excl", "SOH exclusions (operation outside OEM params / excessive cycling / extreme temp) are under EMS control — 'not excluded'.", "D", "counter",
     "Counter — where operation is under DISPERON EMS control the exclusion does not apply to the Client; define 'extreme'."),
    ("11.3A", "Local spare-parts warehouse operational by PAC with agreed stock list; failure to stock = default and downtime counts as Unavailable.", "T", "changes",
     "Now drafted — see LTSA §11.3A: local spare-parts warehouse operational no later than PAC with an agreed critical-component stock list; failure to maintain stock is a default and related downtime counts as Unavailable. Full PCS units and transformers are subject to further agreement, interim-supported by a four-week OEM supply commitment."),
    ("11.x spares", "Stock battery cells + full PCS units (1000/1250 kVA) + transformer units locally.", "D", "counter",
     "Counter — critical sub-components stocked locally; full PCS/transformer on 4-week OEM supply (transformer stocking TBD, 30-day window)."),
    ("11.6", "Obsolescence: support compatible spares/firmware/replacements for 15 years; equivalent-or-better if obsolete.", "T", "changes",
     "Now drafted — see LTSA §11.6: the Service Provider uses commercially reasonable efforts to support compatible spares/firmware/replacement components for at least 15 years, providing an equivalent-or-better replacement without loss of warranty, certification or grid-code compliance if a component is obsolete."),
    ("12.5A/12.7A", "Client operational obligations not breached where caused by EMS/BMS/PCS malfunction, SP monitoring failure or FM; guarantees suspended only on proven direct causation + notice + cure.", "T", "changes",
     "Now drafted — see LTSA §12.5A/12.7A: Client operational obligations are not treated as breached where caused by EMS/BMS/PCS malfunction, Service Provider monitoring failure or Force Majeure; guarantee suspension requires proven direct causation, notice and a cure opportunity; disputed non-payment does not suspend safety monitoring or warranty-preservation alerts."),
    ("Client oblig.", "Operate within OEM params / thresholds / avoid low-SOC — all 'under EMS control and responsibility'.", "D", "changes",
     "Reflected via the EMS-affiliate carve-out; EMS parameters to be set by DISPERON. Wording to align."),
    ("13.4 term", "Client 90 days / SP 12 months (post-Initial-Term); no prejudice to warranty/data/transition/step-in/lender rights.", "T", "asis",
     "Accepted as client's version — LTSA §13.4 updated."),
    ("13.6", "90 days' transition assistance (data export, credentials, config, register maps, firmware, records) at agreed rates.", "T", "asis",
     "Now drafted — see LTSA §13.6: up to 90 days of transition assistance at agreed rates (data export, passwords/access credentials, configuration files, register maps, firmware records, maintenance records, open-incident files, spare-parts list, warranty-claim files)."),
    ("14.2 cap", "Liability cap 12 months → 200% of preceding-12-month fees; carve-outs for cyber/data/regulatory/safety/spares.", "T", "counter",
     "Counter — cap raised to 24 months of preceding Service Fees (from 12); carve-outs accepted."),
    ("14.3A/14.4A", "Caps not to apply to cyber/data/regulatory/safety/spares; third-party exclusions not to apply where SP selected/configured/integrated/maintained the system.", "T", "asis",
     "Now drafted — see LTSA §14.3A/14.4A: liability limitations do not apply to cyber-security breach, Client-data loss/corruption/withholding, regulatory/grid-code non-compliance, or safety incidents caused/worsened by Service Provider breach; third-party-system exclusions do not apply where the Service Provider selected, configured, integrated, maintained, monitored or recommended the system."),
    ("14.5 OEM", "For OEM defect claims, SP to actively manage/prosecute the claim and remain liable for its own failures (warranty preservation, spares, mitigation).", "T", "changes",
     "Now drafted across LTSA §6.1A/§6.3A (warranty-preservation liability and file) and §14.5(c) (Service Provider uses reasonable efforts to facilitate OEM warranty claims); addressed by these clauses in combination rather than by a single dedicated 'actively manage/prosecute' clause."),
    ("16.4", "Force Majeure not to excuse monitoring/data/alerts/warranty preservation/mitigation/insurance/cyber/records or remote-performable obligations.", "T", "asis",
     "Now drafted — see LTSA §16.4: Force Majeure does not excuse monitoring (where technically possible), data provision, alerts, warranty-right preservation, downtime mitigation, insurance maintenance, cyber-security, records, or obligations capable of remote performance."),
    ("18A", "Cyber security controls for ICS (MFA, RBAC, VPN, encryption, audit logs, vulnerability mgmt, firmware integrity, annual review, incident reporting).", "T", "changes",
     "Now drafted as a framework clause — see LTSA §18A: IEC 62443-consistent controls (MFA, RBAC, secure remote access, encryption in transit/at rest, audit logs, least-privilege accounts, vulnerability management, firmware-integrity controls, annual access review). Formal IEC 62443/NIS2 certification is still pending confirmation with DISPERON."),
    ("19.5A", "SCADA/DSO comms failures attributable to SP/EMS/config/firmware/gateway/maintenance/cyber count as Unavailable Hours.", "T", "asis",
     "Now drafted — see LTSA §19.5A: SCADA/DSO communication failures attributable to the Service Provider, EMS/SCADA configuration, protocol mapping, firmware changes, comms gateway, maintenance activity or cyber-security controls count as Unavailable Hours where they prevent dispatch, monitoring or DSO-approved operation."),
    ("19.8A", "Documentation in editable electronic format (native settings, relay settings, SCADA point lists, Modbus/IEC61850/IEC104 maps, firmware, access handover, photos, test certs).", "T", "asis",
     "Now drafted — see LTSA §19.8A: documentation in editable electronic format where reasonably available — native settings/configuration files, relay settings, SCADA point lists, Modbus/IEC 61850/IEC 60870-5-104 maps, firmware versions, access-control handover records, photographs, independent test certificates."),
    ("20.5", "Expert Determination for technical disputes (availability/SOH/RTE/root cause) by an expert agreed or appointed by President of ETEK.", "T", "asis",
     "Now drafted — see LTSA §20.5: technical disputes on Availability, SOH, RTE, capacity, root cause, warranty attribution or excluded hours go to an independent technical expert agreed by the Parties or appointed by the President of ETEK; the determination is binding absent manifest error."),
    ("21.3B", "Lender step-in / cure / direct agreement / direct payment of proceeds; SP to enter a direct agreement with the lender.", "T", "changes",
     "Now drafted — see LTSA §21.3B: project-finance lender/security agent gets step-in, notice, cure, payment-assumption and enforcement rights, with continued performance for a reasonable cure period; the Service Provider must enter a direct agreement with the lender. Alpha Bank's specific wording for that direct agreement is still pending."),
    ("Assignment", "Client may assign/charge to lender/refinancing/purchaser/successor without SP consent.", "T", "changes",
     "Now drafted — see LTSA §21.3A: the Client may assign its rights to a project-finance lender or secured creditor without the Service Provider's consent, for security purposes; on enforcement the Service Provider recognises the lender as the Client provided payment obligations are assumed. Scope covers lender/secured-creditor assignment; not the broader refinancing/purchaser/successor list originally asked."),
    ("Sch 2", "Schedule 2 completed before signing (tier, fee, warranty elections, EMS, CPI cap, spares list, confirmation of preserved rights).", "T", "asis",
     "Accepted — Schedule 2 completion note added."),
    ("Sch 4", "Revise Schedule 4 to remove ambiguity between % fee-reduction table and €30/day/MWh; use method more favourable to Client unless single formula agreed.", "T", "changes",
     "Included with changes — Schedule 4 formula clarified (whichever is higher, subject to cap)."),
    ("Companion", "Companion Documents list (EPC v6.0+, DWU, OEM Terms, Technical Agreements, APGs, Performance Guarantee, EMS, lender direct agreement); update EPC v4.0 references.", "T", "asis",
     "Accepted — EPC references updated v4.0 → v6.0."),
]

ANASTASIOS = [
    ("1", "EPC conditional on signing the LTSA (CP or simultaneous execution) — several EPC protections only work with the LTSA in place.", "A", "asis",
     "Done — LTSA is a condition precedent to the advance (Annex V6 §1A.5(e)) and simultaneous execution required."),
    ("2", "OEM performance guarantee weak (corporate, Chinese law, 30-day cure, expires at DLP end, no parent guarantee). Ask: extend validity to ≥12 months post-PAC and reduce cure for insolvency.", "A", "none",
     "Third-party (Linyang) — validity extension and cure reduction to be requested from Linyang; not yet confirmed."),
    ("3", "Renew the 6-month DLP ask (retention + performance guarantee expire with it).", "A", "counter",
     "Not accepted — DLP stays 3 months; FAC added and Lighthief International undertaking (§16.3B) added to strengthen cover."),
    ("4", "EMS-caused downtime should not escape the availability guarantee where the EMS failure is attributable to DISPERON / any Lighthief group entity.", "A", "asis",
     "Done — EPC §10.5A and LTSA §9.3(e) carve-out added."),
    ("5", "Technical Agreement fixes: (i) 2.7V vs 2.8V warranty-void threshold + interim charging/storage responsibility pre-PAC = Lighthief; (ii) EMS 'provided by customer' → DISPERON.", "A", "none",
     "Pending — Technical Agreement drafting fixes for Costas + Linyang (not in EPC/LTSA generators)."),
    ("6", "Capture in writing that VSG / black-start activation on supplied hardware is at no additional equipment cost (resolves the collision with the Technical Agreement).", "A", "asis",
     "Done — grid-forming/VSG included in the Contract Price (EPC §10.x)."),
    ("7", "Bank guarantees: confirm issuing bank acceptable / EU-confirmed and governing law; final instruments must match specimens.", "A", "none",
     "Third-party — Bank of Communications re-issue pending (beneficiary = Galascope, EU confirmation, governing law, SWIFT)."),
    ("Risk 1", "Manufacturing-defect protection ultimately rests on Lighthief's covenant; a parent guarantee from Lighthief International would strengthen it.", "A", "asis",
     "Addressed — Lighthief International performance undertaking by confirmation letter (Arkadius Sybaris), Annex V6 §16.3B."),
    ("Risk 2", "Availability protection thin (group-level 97%, 20% LD cap ≈ €13k/yr, sole remedy).", "A", "counter",
     "Partly — LD cap raised to 50%/Park; sole-remedy carve-out added; availability stays group-level with rationale note."),
    ("Risk 3", "LTSA liability capped at 12 months and 180-day convenience exit; if Lighthief exits, availability guarantee goes with it.", "A", "changes",
     "Improved — liability cap raised to 24 months; termination now Client 90 days / SP 12 months (post-Initial-Term) with continuity protections."),
]


def _mark_v62(rows):
    """Reflect the v6.2 update: items previously NOT-YET-INCLUDED are now incorporated."""
    out = []
    for ref, ask, author, status, resp in rows:
        if author == "T" and status == "none":
            status = "changes"
            resp = "Incorporated in v6.2 (EPC Annex V6 / LTSA Annex LT-2). " + resp
        out.append((ref, ask, author, status, resp))
    return out


def add_table(doc, title, rows):
    h = doc.add_paragraph()
    run(h, title, bold=True, size=13, color=GOLD)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for c, label in zip(hdr, ["Ref", "Client ask", "Our response & status"]):
        shade(c, "1A365D")
        p = c.paragraphs[0]
        run(p, label, bold=True, size=9, color=WHITE)
    widths = [Pt(48), Pt(250), Pt(250)]
    for ref, ask, author, status, resp in rows:
        col = item_color(author, status)
        lab, labcol, note = status_label(author, status, resp)
        cells = t.add_row().cells
        # Ref
        run(cells[0].paragraphs[0], ref, bold=True, size=9, color=col)
        # Ask (author-coloured)
        run(cells[1].paragraphs[0], ask, size=9, color=col)
        # Response + status tag
        rp = cells[2].paragraphs[0]
        run(rp, f"[{lab}] ", bold=True, size=9, color=labcol)
        run(rp, note, size=9, color=BLACK)
    for row in t.rows:
        for i, c in enumerate(row.cells):
            c.width = widths[i]
    doc.add_paragraph()


def main():
    doc = Document()
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(p, "Galascope — Redline Response Matrix", bold=True, size=17, color=NAVY)
    sub = doc.add_paragraph()
    run(sub, "Client feedback on the EPC and LTSA — Timotheos (technical review), "
             "Dino Constantinou, and Anastasios (lawyer) — and our response as at v6.3. "
             "Prepared by Lighthief Cyprus Ltd · July 2026.",
        italic=True, size=9.5, color=GREY_T)

    # Legend
    lg = doc.add_paragraph()
    run(lg, "How to read the colours", bold=True, size=11, color=GOLD)
    legend = [
        ("Timotheos / technical review — INCLUDED as requested", RED),
        ("Timotheos / technical review — INCLUDED with our drafting changes", PURPLE),
        ("Timotheos / technical review — INCLUDED with our counter", ORANGE),
        ("Timotheos / technical review — NOT YET INCLUDED / declined", GREY),
        ("Dino Constantinou — item (status noted in the response column)", TEAL),
        ("Anastasios (lawyer) — comment (status noted in the response column)", BLUE),
    ]
    for text, col in legend:
        lp = doc.add_paragraph(style="List Bullet")
        run(lp, "\u25A0 ", bold=True, size=10, color=col)
        run(lp, text, size=9.5, color=BLACK)

    doc.add_paragraph()
    add_table(doc, "Part A — EPC v6.3 redline", _mark_v62(EPC))
    doc.add_page_break()
    add_table(doc, "Part B — LTSA v6.3 redline", _mark_v62(LTSA))
    doc.add_page_break()
    add_table(doc, "Part C — Anastasios (lawyer) comments", ANASTASIOS)

    # Footer note
    fn = doc.add_paragraph()
    run(fn, "Lighthief Cyprus Ltd · HE 477423 · office@lighthief.com · +357 77 77 00 50 · solarfarms.cy",
        italic=True, size=8, color=GREY_T)

    doc.save(str(OUT))
    print("Saved:", OUT)

    # counts
    from collections import Counter
    def tally(rows):
        c = Counter()
        for _, _, a, s, _ in rows:
            key = ("Dino" if a == "D" else "Anastasios" if a == "A" else f"Timotheos-{s}")
            c[key] += 1
        return c
    for name, rows in [("EPC", EPC), ("LTSA", LTSA), ("Anastasios", ANASTASIOS)]:
        print(name, dict(tally(rows)))


if __name__ == "__main__":
    main()
