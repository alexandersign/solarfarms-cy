#!/usr/bin/env python3
"""Build the FINAL Galascope client package (for Dino / lawyer / board).

1. Generates EPC v6 = the v5.1 body + a governing "Annex V6 — Amendment Schedule
   (June 2026)" that integrates every agreed v6 change (APG to PAC, RTE split,
   FAT, extended rejection window, price validity long-stop, VSG note, liability
   split, manufacturing-defect carve-out, order-trigger CP). A top banner marks
   the document as v6.0; the Annex prevails over the body where inconsistent.
2. Assembles docs/.../v6/CLIENT-PACKAGE/ with ONLY client-facing final files
   (renamed, numbered) + a clean client cover note (no internal content).

Run: python docs/clients/group-order/Group2_Esperia_Energy/contracts/v6/build-galascope-final-package-jun2026.py
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
GREY = RGBColor(0x40, 0x40, 0x40)

# Amendment colour legend (matches the Redline Response Matrix)
AMD_RED = RGBColor(0xC0, 0x00, 0x00)     # Timotheos / technical — included as requested
AMD_PURPLE = RGBColor(0x70, 0x20, 0x9C)  # Timotheos / technical — included with our changes
AMD_ORANGE = RGBColor(0xC0, 0x60, 0x00)  # Timotheos / technical — included with our counter
AMD_TEAL = RGBColor(0x0F, 0x76, 0x8A)    # Dino Constantinou
AMD_BLUE = RGBColor(0x1F, 0x49, 0xC0)    # Anastasios (lawyer)

# leading-token -> colour (clauses that answer a client redline item)
_AMD_MAP = {
    # EPC Annex V6
    "11.5": AMD_RED, "9.2A": AMD_RED, "4.4B": AMD_RED, "10.5A": AMD_RED,
    "1A.5A": AMD_RED, "19.12": AMD_RED, "Schedule": AMD_RED, "10.10": AMD_RED,
    "7.1B": AMD_ORANGE, "13.3A": AMD_ORANGE, "9.1B": AMD_ORANGE, "8.4.8": AMD_ORANGE,
    "7.4A": AMD_ORANGE, "5A.5": AMD_ORANGE,
    "16.3B": AMD_BLUE,
    "14.5": AMD_PURPLE, "4.6": AMD_PURPLE, "10.6A": AMD_PURPLE, "12.6": AMD_PURPLE,
    "1A.7A": AMD_PURPLE, "16.3A": AMD_PURPLE, "10.8A": AMD_PURPLE, "10.9B": AMD_PURPLE,
    "18A": AMD_PURPLE,
    # LTSA Annex LT-2 counters (rest default to purple in _append_annex)
    "8.2A": AMD_ORANGE, "10.4A": AMD_ORANGE,
}


def _amd_color(headline, default):
    tok = headline.split()[0] if headline.split() else ""
    return _AMD_MAP.get(tok, default)


def _annex_colour_legend(doc):
    p = doc.add_paragraph()
    _run(p, "Colour key:  ", bold=True, size=8.5, color=GREY)
    _run(p, "\u25A0 included as requested   ", size=8.5, color=AMD_RED)
    _run(p, "\u25A0 included \u2014 our changes   ", size=8.5, color=AMD_PURPLE)
    _run(p, "\u25A0 included \u2014 our counter   ", size=8.5, color=AMD_ORANGE)
    _run(p, "\u25A0 from Dino   ", size=8.5, color=AMD_TEAL)
    _run(p, "\u25A0 from Anastasios   ", size=8.5, color=AMD_BLUE)
    _run(p, "\u25A0 base v6 amendment", size=8.5, color=NAVY)


# ── Inline integration: relocate amendment-schedule clauses into the body ────
# The client (Timotheos/Stelios) could not find the agreed changes because they
# were appended as a separate Annex at the end of the document. This moves each
# clause to its correct numbered position in the body (colour-coding preserved),
# so there is exactly one document, in clause-number order, no side annexes.

def _find_para(doc, prefix, occurrence=1):
    """Return the Nth paragraph (1-indexed) whose stripped text starts with
    `prefix` followed by a word boundary (space, newline, '.', or end)."""
    n = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith(prefix):
            rest = t[len(prefix):]
            if rest == "" or rest[0] in " \n.:,":
                n += 1
                if n == occurrence:
                    return p
    return None


def _clear_para(p):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)


def _write_clause(p, headline, text, color):
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, headline.strip() + "  ", bold=True, size=10.5, color=color)
    _run(p, text, size=10.5, color=color)


def _insert_clause_after(anchor_para, headline, text, color):
    new_p = OxmlElement("w:p")
    anchor_para._p.addnext(new_p)
    new_para = Paragraph(new_p, anchor_para._parent)
    _write_clause(new_para, headline, text, color)
    return new_para


def _replace_clause(p, headline, text, color):
    _clear_para(p)
    _write_clause(p, headline, text, color)
    return p


def _delete_para(p):
    p._p.getparent().remove(p._p)


def _clean_annex_refs(doc, replacements):
    for p in doc.paragraphs:
        for old, new in replacements:
            for r in p.runs:
                if old in r.text:
                    r.text = r.text.replace(old, new)


V6 = Path(__file__).resolve().parent
CONTRACTS = V6.parent
REPO = CONTRACTS.parents[4]  # .../solinvest
PKG = V6 / "CLIENT-PACKAGE"

EPC_SRC = CONTRACTS / "EPC-Galascope-Esperia-batch1-may2026.docx"
EPC_OUT = PKG / "01-EPC-Agreement-Galascope-G1-G2-v6-jun2026.docx"


def _run(p, text, *, bold=False, italic=False, size=10.5, color=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return r


def h(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, bold=True, size=size, color=GOLD)
    return p


def para(doc, text, *, bold=False, italic=False, color=None, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, text, bold=bold, italic=italic, color=color, size=size)
    return p


# ── v6 amendment clauses (verbatim agreed text) ──────────────────────────────
ANNEX = [
    ("1A.4  Nature of EPC Execution.",
     "Execution of this Agreement records the Parties' agreed terms. It is not an order to "
     "manufacture and creates no payment obligation. The order is placed, and the advance under "
     "Section 7.1(a) falls due, only on satisfaction of the Conditions Precedent in Sections 1A.1 "
     "and 1A.5."),
    ("1A.5  Companion Documents \u2014 Condition Precedent to the Advance.",
     "The advance shall not become due until the Contractor has presented to the Client, and the "
     "Client has accepted or signed off (such acceptance not to be unreasonably withheld or delayed "
     "beyond ten (10) Business Days of presentation), each of the following, substantially in the "
     "form attached: (a) the OEM Direct Warranty Undertaking, signed and sealed by the OEM, covering "
     "the Project (Galascope 1 and Galascope 2); (b) the Advance Payment Guarantee (APG No. 1) issued "
     "by the OEM's bank naming the Client as beneficiary, and consistent with the agreed APG wording "
     "per Section 10.9-APG; (c) the OEM 5% performance guarantee; (d) the Confirmed Price Certificate "
     "per Section 6.1; (e) the LTSA (LCY-LTSA-GAL-2026), executed and delivered per Section 1A.5A; "
     "(f) the Technical Agreement(s) for Galascope 1 and Galascope 2, signed off by both Parties "
     "confirming the frozen technical specification (Schedule A); and (g) the EMS Subscription "
     "Addendum, executed by the Client and the EMS provider (Disperon Sp. z o.o., trading as "
     "DISPERON)."),
    ("1A.5B  Pre-Shipment Payment \u2014 Condition Precedent.",
     "The pre-shipment payment (Section 7.1(b)) shall not become due until: (a) the equipment has "
     "passed the Factory Acceptance Test (Section 8.3B) witnessed by or on behalf of the Client; and "
     "(b) the Contractor has presented the Pre-Shipment Advance Payment Guarantee (APG No. 2) issued "
     "by the OEM's bank naming the Client as beneficiary (Section 10.9(c)). The Client never funds an "
     "equipment prepayment tranche without a guarantee in place for that equipment prepayment."),
    ("1A.6  Bilateral Walk-Away.",
     "If any document in Section 1A.5 is not presented in the agreed form by the Long-Stop Date in "
     "Section 1A.1, either Party may elect not to proceed by written notice without liability, and "
     "any advance received shall be refunded within thirty (30) days less reasonable and documented "
     "costs. The companion-document drafts attached are provided for review and are not executed by "
     "signature of this Agreement."),
    ("6.1B  Price Basis Certificates.",
     "The Contractor shall issue: (i) an Indicative Price Basis Certificate at signing, stating the "
     "Indicative Contract Price, the reference indices (Mysteel lithium carbonate and EUR/CNY as at "
     "the Effective Date) and effective MWh per park; and (ii) a Delivery Price Basis Certificate "
     "upon delivery, stating whether the Confirmed Price equals the Indicative Price or the adjusted "
     "Confirmed Price under (c)-(d), with full index values and calculation shown. The Client may "
     "notify a dispute within fourteen (14) days of each Certificate; absent dispute the Certificate "
     "is deemed accepted for payment milestones following that date. Schedule A effective MWh: "
     "Galascope 1 \u2014 20.06 MWh; Galascope 2 \u2014 10.03 MWh."),
    ("6.1C  Price Validity Long-Stop.",
     "Where the Client receives Connection Terms within six (6) months of the Effective Date, the "
     "Confirmed Price is set under sub-clauses (c)-(d) (two-way adjustment; upward movement capped "
     "at five percent (5%); downward movement passed through in full on milestones not yet "
     "invoiced). Where Connection Terms are received more than six (6) months after the Effective "
     "Date, the Contractor shall obtain a refreshed OEM quotation and re-anchor the Indicative Price "
     "and reference indices to the month of that refreshed quotation; the same two-way adjustment and "
     "5% cap then apply from the re-anchored basis. Either Party may terminate under Section 6.1(d) "
     "if the resulting upward adjustment would exceed five percent (5%)."),
    ("6.1A  Price basis \u2014 consistent with the signed Letter of Intent (LOI Clause 4.4).",
     "Consistent with the signed LOI (Clause 4.4): Schedule A pricing is locked with the OEM on the "
     "January 2026 quotation basis, not spot index. The Indicative Price (Schedule A, or as updated "
     "in writing before signing) shall be the Confirmed Price for equipment-cost purposes, without "
     "upward raw-material or FX adjustment for index movement occurring before the EPC Effective "
     "Date. The two-way adjustment under Section 6.1(c)-(d) applies after signing only, to verified "
     "movement from the indices stated in the Indicative Price Basis Certificate at EPC signing "
     "(Section 6.1B). The OEM has confirmed it continues to hold this price at the current level. "
     "References in the body and Schedule A to the January 2026 quotation or reference values are "
     "read consistently with this paragraph and LOI Clause 4.4."),
    ("8 / Schedule A Part 3  Delivery & Shipping Schedule and Target PAC \u2014 confirmed at Connection Terms, not fixed at signing.",
     "Manufacturing, shipping and delivery lead time is NOT fixed at signing and no time commitment is "
     "given at signing. Production lead time is indicative only and may change from any manufacturing "
     "estimate previously discussed (including any three (3) month manufacturing figure), depending on "
     "the OEM's production capacity at the time of order. Upon receipt of Connection Terms (the order "
     "trigger under Section 1A), the Contractor shall issue a binding Delivery Schedule stating the "
     "confirmed manufacturing lead time, the shipping timeline and the Target PAC Date, together with the "
     "Confirmed Price Certificate under Section 6.1B; the Confirmed Price may differ from the "
     "Indicative Price only within the two-way adjustment and the five percent (5%) upward cap under "
     "Section 6.1(c)-(d). All time-based obligations and all Delay Liquidated Damages under Section 8.4 "
     "are calculated from that confirmed Delivery Schedule, and not from any duration or calendar date "
     "stated at or before signing."),
    ("1A.7  Client withdrawal on the delivery timeline.",
     "If the Target PAC Date in the Delivery Schedule confirmed at Connection Terms falls later than "
     "fifteen (15) months after the Client's receipt of Connection Terms (the \u201cPAC Long-Stop\u201d), "
     "the Client may withdraw from this Agreement by written notice within ten (10) Business Days of "
     "receiving the Delivery Schedule, without liability to either Party, and any advance received "
     "shall be refunded within thirty (30) days. This reflects that neither Party can predict the "
     "OEM's market or production-capacity movements at the time of order, while giving the Client "
     "certainty that the Contractor bears the risk of any Delivery Schedule beyond the PAC Long-Stop. "
     "The Client has no right of withdrawal under this Section where the confirmed Target PAC Date is "
     "on or before the PAC Long-Stop."),
    ("5A.1A  Planning.",
     "The Client shall use reasonable endeavours to obtain planning permissions and landowner "
     "consents within sixty (60) days of the Effective Date. Where delay by a planning or competent "
     "authority beyond the Client's reasonable control causes delay to the Target PAC Date, the "
     "Target PAC Date extends day-for-day. The Contractor shall not be entitled to additional cost "
     "reimbursement beyond reasonable and documented standby costs directly caused by such planning "
     "delay. The Client shall notify the Contractor promptly of material planning delay."),
    ("8.4.7A  Upstream Force Majeure (OEM Flow-Through) \u2014 replaces the body clause.",
     "(a) Where the OEM has formally notified the Contractor in writing of an event it claims is "
     "Force Majeure under the upstream supply contract, the Contractor may notify the Client within "
     "seven (7) Business Days, attaching a copy of the OEM notice. (b) The Target PAC Date extends "
     "only if the notified event constitutes Force Majeure under Section 12.1 (no extension for "
     "events within the §12.1 exclusions, including ordinary commercial or production delays, raw "
     "material sourcing, port congestion, or financing). (c) If the Client reasonably contends the "
     "event is not §12.1 Force Majeure, it shall notify the Contractor within ten (10) Business Days "
     "of the Contractor's notice; the Target PAC Date shall not extend during any dispute. (d) If "
     "the Parties do not agree within thirty (30) days, either Party may refer the matter to dispute "
     "resolution under Section 15.3; no extension is deemed agreed unless resolved in favour of FM "
     "qualification. (e) Where the event is agreed or determined to be Force Majeure, the Target PAC "
     "Date extends day-for-day and no Delay LDs accrue for that period. (f) The Contractor shall use "
     "reasonable endeavours to challenge any OEM notice that clearly does not meet Section 12.1."),
    ("10.9(b)  Advance Payment Guarantee No. 1 \u2014 advance tranche (replaces body 10.9(b)).",
     "The Contractor shall procure that the equipment manufacturer's bank (Bank of Communications) "
     "issues a first-demand Advance Payment Guarantee under URDG 758 naming Galascope Ltd (HE 303759) "
     "and/or its project-finance security agent as beneficiary, securing the advance payment for the "
     "supply of the equipment. The Guaranteed Amount is EUR 705,792 (the equipment-side advance), as "
     "set out in Schedule A and stated on the face of the APG. This and Section 10.9(c) replace the "
     "APG figures in the body and in Schedule A Parts 2 and 4."),
    ("10.9(a)  APG issuance & validity \u2014 confirmed mechanism (addresses the legality query).",
     "The OEM's bank (Bank of Communications) issues the Advance Payment Guarantee(s) upon being "
     "provided with both this EPC Agreement and the OEM Sales Contract, and issues them with "
     "Galascope Ltd named as beneficiary from the outset. The Client may therefore demand directly "
     "under the APG(s) notwithstanding that the Client pays the Contractor (and not the OEM): the "
     "APG is original issuance naming the Client as beneficiary, not a transfer or assignment. This "
     "confirms the validity of the APG for the Client and underpins the payment terms in Section 7.1."),
    ("10.9(c)  Advance Payment Guarantee No. 2 \u2014 pre-shipment tranche (new).",
     "Before the pre-shipment payment falls due, and conditional on a passed Factory Acceptance Test, "
     "the Contractor shall procure that the manufacturer's bank issues a second first-demand Advance "
     "Payment Guarantee (or, at the Contractor's election, an increase to APG No. 1) under URDG 758 "
     "naming the Client and/or its security agent as beneficiary, securing the pre-shipment payment "
     "for the supply of the equipment. The Guaranteed Amount is EUR 1,411,585 (the equipment-side "
     "pre-shipment tranche), as set out in Schedule A. Combined APG cover is therefore EUR 2,117,377 "
     "(the equipment payments made before delivery). The DAP and acceptance tranches, paid on or "
     "after delivery, are not prepayments and are not APG-secured (covered by possession of the "
     "equipment, the performance guarantee, and the Retention)."),
    ("10.9(d)  APG validity \u2014 to PAC (replaces body 10.9(d) 'delivery + 30 days').",
     "Each APG shall remain valid until the earlier of: (i) issuance of the Provisional Acceptance "
     "Certificate (PAC); or (ii) twelve (12) months after delivery of the equipment to Site; upon "
     "issuance of PAC the APGs are released. This keeps the advance-refund cover alive through "
     "commissioning, when latent equipment faults are typically discovered."),
    ("10.9(e)  APG trigger.",
     "A demand may be made under either APG where the manufacturer has failed to refund the relevant "
     "prepayment when due, or has failed to deliver conforming/functioning equipment, including on the "
     "Client's rejection of equipment at commissioning (see Section 9 rejection window and the "
     "manufacturing-defect carve-out below)."),
    ("10.9A  Security layering alongside the APGs.",
     "Through to PAC the Client is protected by, in combination: (a) the two APGs (advance + "
     "pre-shipment refund, to PAC / 12 months, EUR 2,117,377 combined); (b) CAR / erection all-risks "
     "insurance for the full replacement value from "
     "arrival at the discharge port through installation to PAC \u2014 the sea voyage is covered by the "
     "OEM's CIF marine insurance, so the Contractor's policy runs from port of discharge; (c) the "
     "OEM 5% performance guarantee to the end of the Defects Liability Period; (d) the Retention "
     "(Section 7.4); and (e) the OEM Direct Warranty Undertaking and 5-year product warranty "
     "(Section 10.8)."),
    ("8.x  Pre-Shipment Inspection / Factory Acceptance Test (FAT).",
     "Before any equipment is despatched and before the pre-shipment milestone (Section 7.1(b)) is "
     "paid, the Client and/or its appointed third-party inspector may inspect and witness factory "
     "acceptance testing at the OEM's facility on ten (10) Business Days' notice. Where the FAT "
     "reveals defects or non-conformity, the OEM/Contractor shall remedy them at no additional cost "
     "before shipment, and the Client may re-inspect. Pre-shipment payment is conditional on a "
     "passed FAT. The FAT is the primary protection against a material proportion of equipment being "
     "defective, as it prevents payment for and shipment of defective equipment."),
    ("9.2A  Final Acceptance Certificate (FAC).",
     "At the end of the Defects Liability Period (three (3) months after PAC), the Contractor shall "
     "issue a Final Acceptance Certificate once: (a) all defects notified during the Defects "
     "Liability Period have been rectified to the Client's reasonable satisfaction; (b) all "
     "punch-list items are closed or formally deferred by written agreement; and (c) the "
     "deliverables in Section 19.12 have been provided. The Retention shall be released within "
     "thirty (30) days of FAC issuance. If the Contractor fails to issue a FAC within thirty (30) "
     "days of the end of the Defects Liability Period despite the above conditions being met, the "
     "Client may issue the FAC itself and the Retention shall be released automatically."),
    ("9.1B  PAC with performance shortfall \u2014 remedy and cap.",
     "If measured capacity or Round-Trip Efficiency is below the Contract guarantees (Section 10.6) "
     "but above the PAC hard floor (Section 9.1(a)\u2013(b)), PAC shall be issued but the Contractor "
     "shall, at the Client's election, (i) perform augmentation, repair or recommissioning to "
     "restore performance to the guaranteed level within sixty (60) days of PAC; or (ii) pay "
     "Liquidated Damages calculated as: LD = (Guaranteed level \u2212 Measured level) \u00d7 System "
     "Capacity \u00d7 prevailing market rate per kWh for equivalent battery modules. The aggregate "
     "Liquidated Damages for performance shortfall under this Section shall not exceed ten percent "
     "(10%) of the Component A Equipment Supply Price. This cap shall not apply to manufacturing "
     "defects, which are subject to the uncapped remedy in Section 13.5 and the OEM Direct "
     "Warranty Undertaking. Performance shortfall LDs under this Section and Delay LDs under "
     "Section 8.4 shall not be aggregated beyond the respective caps for each."),
    ("8.4.8  Delivery Schedule hardening and FM preservation.",
     "Upon issue of the Delivery Schedule at Connection Terms, the milestones set out in that "
     "Schedule are binding on the Contractor. The Contractor shall provide monthly progress reports "
     "against each milestone. Where any milestone is at risk of delay, the Contractor shall notify "
     "the Client within seven (7) Business Days of that risk arising. Notwithstanding the foregoing, "
     "the extension-of-time regime in Section 8.4.7A (OEM Force Majeure flow-through), Section "
     "8.4.6 (Exclusions), and Section 8.4.7 (Extension of Time) remain in full force and effect "
     "and are not affected by the binding nature of the Delivery Schedule. Ordinary production-capacity "
     "constraints, raw-material sourcing, port congestion, and financing issues shall not entitle "
     "the Contractor to an extension of time under this Agreement (consistent with Section 8.4.7A's "
     "exclusion of such events from Force Majeure). The Client accepts that the Delivery Schedule "
     "is established at Connection Terms and may differ materially from indicative lead times "
     "discussed before signing."),
    ("7.4A  Set-off rights.",
     "The Client may set off against any unpaid milestone or PAC payment any amount that is: "
     "(i) agreed in writing by the Parties as due from the Contractor; or (ii) determined as due "
     "by a binding adjudicator's or expert's decision under Section 15.3. Set-off shall not apply "
     "to amounts that are disputed in good faith and not yet ascertained. The Client shall give "
     "the Contractor seven (7) Business Days' prior written notice before exercising any set-off "
     "right, identifying the amount and basis of the set-off."),
    ("14.5  Insurance \u2014 coverage structure and Contractor's procurement covenant.",
     "(a) Coverage structure. The sea voyage from the port of shipment to CIF Limassol is intended "
     "to be covered by the OEM's (Linyang's) CIF marine cargo insurance, and the Contractor's "
     "Construction All-Risks (CAR) policy is intended to attach from arrival at the port of "
     "discharge (CIF Limassol) and to run through inland transit and installation to PAC. "
     "(b) Procurement covenant. The Contractor shall use all reasonable endeavours to procure, and "
     "to maintain in force for the relevant periods, insurances that: (i) name Lighthief Cyprus Ltd "
     "as insured; (ii) name the Client (Galascope Ltd) and, where notified, its project-finance "
     "lender or security agent as additional insured on the CAR and Public Liability policies, with "
     "a waiver of subrogation in their favour and, in respect of the lender, a loss-payee / lenders' "
     "endorsement; (iii) provide inland-transit cover attaching from the port of discharge through "
     "to delivery at Site; and (iv) so far as reasonably obtainable from insurers, bridge the "
     "interface between the OEM's marine cargo cover and the CAR cover (including by way of a "
     "difference-in-conditions or 50/50 loss-sharing basis). (c) Evidence. The Contractor shall "
     "provide certificates of insurance and the relevant endorsements evidencing the above before "
     "work commences at Site (and, in respect of any advance payment, per Section 14.4), and shall "
     "notify the Client promptly of any material change to or cancellation of any policy. "
     "(d) Where any endorsement in (b) is not obtainable from insurers on commercially reasonable "
     "terms, the Contractor shall notify the Client and the Parties shall agree an alternative "
     "arrangement in good faith; nothing in this Section warrants cover that insurers decline to "
     "provide. (e) Commissioning failure due to Contractor negligence. For the avoidance of doubt, "
     "any failure to achieve commissioning, or any defect, that is attributable to the Contractor's "
     "professional negligence (including electrical-engineering or design negligence) is covered "
     "under the Contractor's Professional Indemnity insurance, which the Contractor shall procure "
     "and maintain in force no later than the Order Date defined in Section 1A.4 and throughout the "
     "Works thereafter, with a limit of not less than EUR 2,000,000, per Section 14.4; this is in "
     "addition to, and does not limit, the Contractor's rectification, delay-LD, performance-LD, "
     "retention and performance-guarantee obligations under this Agreement."),
    ("11.5  Integrated document suite and interpretation (client-agreed).",
     "The EPC Agreement, the LTSA, the OEM Direct Warranty Undertaking, the Advance Payment "
     "Guarantees, the OEM Performance Guarantee, the EMS Subscription Addendum, the Technical "
     "Agreements and the OEM Warranty Terms shall be read together as a single, integrated "
     "project-document suite. In the event of any inconsistency or conflict between them, the "
     "interpretation that is most favourable to the Client shall govern."),
    ("7.1B  Independent Engineer sign-off \u2014 PAC payment.",
     "The PAC payment under Section 7.1(c) shall not become due until the relevant PAC test "
     "certificate and PAC report (Section 9.1) have been accepted in writing by the Client or its "
     "appointed Independent Engineer, such acceptance not to be unreasonably withheld or delayed "
     "beyond ten (10) Business Days of the Client's receipt of a complying PAC certificate. This "
     "condition applies to the PAC payment only and does not affect the advance or pre-shipment "
     "payments under Sections 7.1(a)-(b)."),
    ("13.3A  Additional liability carve-outs (not subject to any cap).",
     "In addition to the carve-outs in Section 13.3, the liability caps in Section 13.2 shall not "
     "apply to: (a) latent defects in the Works or Equipment; and (b) non-conformity with mandatory "
     "safety requirements or applicable grid-code requirements. For the avoidance of doubt, general "
     "firmware, software and design defects that are the responsibility of the OEM remain subject "
     "to the OEM Direct Warranty Undertaking and the manufacturing-defect carve-out (Section 13.5), "
     "and are not additionally uncapped under this Section."),
    ("16.3B  Lighthief International Ltd \u2014 performance undertaking.",
     "Lighthief International Ltd guarantees to the Client the due performance by Lighthief Cyprus "
     "Ltd of its core EPC obligations under this Agreement, being: (a) delivery of the BESS "
     "equipment to Site as EPC Contractor; (b) installation, commissioning and handover of the "
     "BESS to the technical specification in Schedule A; and (c) the Contractor's warranty "
     "obligations under Section 10.1 for the Warranty Period. This undertaking: (i) does not "
     "extend to financial liabilities beyond the Contract Price, to the LTSA or EMS services "
     "(which are contracted separately), or to OEM obligations guaranteed under the DWU and "
     "performance guarantee; (ii) is governed by Cyprus law; and (iii) remains in force for the "
     "duration of the Warranty Period. This undertaking is provided by way of a parent-company "
     "confirmation letter from Lighthief International Ltd, signed by its authorised officer "
     "(Dr. Arkadius Sybaris, Founder & CEO), delivered to the Client before the advance payment "
     "under Section 7.1(a) becomes due."),
    ("9.x  Extended rejection window.",
     "The Client's right to inspect and to give notice of defect or non-conformity, and the "
     "associated repair/replace/refund remedy, extends until commissioning / PAC and is not limited "
     "to thirty (30) days after delivery. For non-conformity affecting ten percent (10%) or more of "
     "installed capacity discovered up to PAC, the Client may require repair or replacement, or "
     "reject the affected equipment and require refund of the corresponding price (recoverable, "
     "where unpaid, under the APG per Section 10.9(e))."),
    ("13.2  Limitation of Liability \u2014 replaces body 13.2.",
     "The Contractor's aggregate liability shall be: (a) for claims relating to the Contractor's own "
     "EPC services, installation and non-OEM works (excluding OEM equipment): ten percent (10%) of "
     "the Contract Price; (b) for all other contractual breaches (excluding (a), (c), fraud and "
     "wilful misconduct): fifty percent (50%) of the Contract Price; (c) for manufacturing defects "
     "in OEM equipment and for fraud or wilful misconduct: uncapped, with the Contractor's right to "
     "pursue OEM recovery per Section 13.4."),
    ("13.5  Manufacturing-defect carve-out.",
     "For confirmed manufacturing defects in OEM equipment, the OEM bears the full cost of repair or "
     "replacement, including shipping to CIF Limassol, and the OEM's standard 10% warranty-liability "
     "cap does not apply. The Contractor shall procure this carve-out from the OEM and reflect it in "
     "the OEM Direct Warranty Undertaking provided to the Client."),
    ("4.4 / 3.1  EMS provider (replaces references to the EMS integration affiliate).",
     "EMS integration, SCADA commissioning and the DISPERON software subscription are provided by "
     "Disperon Sp. z o.o. (NIP 9492265995; trading as DISPERON), a company under Lighthief "
     "International Ltd, under a separate EMS Integration Agreement and EMS Subscription Addendum. "
     "The Contractor guarantees that provider's performance of its obligations under those documents."),
    ("4.4B  End-to-end interface responsibility.",
     "The Contractor remains responsible for end-to-end BMS, PCS, SCADA and EMS interface "
     "compatibility, communications, command execution and data integrity at PAC and during the "
     "Defects Liability Period. Failure of an affiliate, software provider or subcontractor to "
     "deliver these functions shall not excuse the Contractor's responsibility to the Client."),
    ("10.5A  Warranty-void carve-out \u2014 EMS/monitoring failures.",
     "The warranty-voiding conditions set out in Section 10.5 shall not be invoked by the "
     "Contractor, and shall not restrict the Client's rights under the OEM Direct Warranty "
     "Undertaking or the LTSA availability guarantee, where the condition was caused by or "
     "attributable to: (a) DISPERON (Disperon Sp. z o.o.) in its capacity as EMS provider; "
     "(b) any other entity within the Lighthief group of companies acting in connection with the "
     "Project; or (c) the Contractor's own monitoring failure or failure to issue a timely alert "
     "that would have enabled the Client to prevent the condition."),
    ("1A.5A  LTSA \u2014 condition precedent to the advance payment.",
     "The Long-Term Service Agreement (LTSA, Ref. LCY-LTSA-GAL-2026) between the same Parties "
     "must be executed and in force before the advance payment under Section 7.1(a) falls due. "
     "The advance payment shall not become due unless the LTSA is already executed and delivered, "
     "together with the other companion documents in Sections 1A.5(a)\u2013(d). The EPC may be "
     "executed independently of the LTSA; the Parties intend that the LTSA will be executed before "
     "Connection Terms are received, but in any event it must be in place no later than the date "
     "the advance payment falls due under Section 1A.3. For the avoidance of doubt, the LTSA is "
     "not a document to be agreed at a later unspecified time \u2014 it must accompany the advance "
     "payment milestone."),
    ("19.12  Technical deliverables at handover.",
     "No later than thirty (30) days after PAC, the Contractor shall deliver to the Client: "
     "(a) as-built drawings and single-line diagrams (SLD); (b) all test reports, commissioning "
     "certificates and inspection certificates issued at FAT, SAT and PAC; (c) SCADA/EMS "
     "IEC\u202f60870-5-104 and Modbus register maps / point lists; (d) protection-relay settings files "
     "and test certificates; (e) firmware version records and access credentials for each device; "
     "(f) fire-system inspection and pressure test certificates; (g) C5 corrosion-coating "
     "certificates; (h) serial-number records for all containers, PCS units, transformers and "
     "MV switchgear; (i) DSO/EAC submission files and approval records; and (j) OEM operation "
     "and maintenance manuals. Documents shall be provided in editable electronic format where "
     "reasonably available. These deliverables are a condition of final retention release under "
     "Section 7.4."),
    ("Schedule A \u2014 Specification Freeze.",
     "The Technical Specification in Schedule A (and the companion Technical Agreements) shall be "
     "frozen and all placeholders, \u2018to be confirmed\u2019 items and open comments shall be closed before "
     "the order-placement advance becomes due under Section 7.1(a). The Contractor shall not place "
     "the order or invoke the advance-payment trigger against an open specification."),
    ("10.10  Grid-forming (VSG) & Black Start.",
     "The PCS supplied (Kehua C-series) is capable of grid-forming (VSG) and black-start operation "
     "via firmware. Activation, FAT or SAT verification, firmware enablement, licences and "
     "commissioning support for grid-forming and black-start capability are included in the Contract "
     "Price, unless a future DSO or grid-code requirement first published after the Effective Date "
     "materially exceeds the requirements known as at signing. Hardware is unaffected."),
    ("19.13  Licensing.",
     "The Contractor shall ensure BESS installation works are performed in a manner suitable for "
     "compliance with applicable licensing and permitting rules published by the Republic of Cyprus "
     "for energy storage installations, to the extent such compliance is within the Contractor's "
     "scope of supply. Installer sign-off by the Contractor's ETEK engineer is included; licensed "
     "electrical design and as-built drawing packages for the wider site are excluded (Client's "
     "engineer)."),
    # ── v6.2: remaining client (Timotheos / technical) redline items now incorporated ──
    ("4.6  Excluded items \u2014 scope certainty.",
     "Items are excluded from the Contractor's scope only if expressly identified in Schedule A with "
     "a fixed technical boundary, price impact and responsible party. Any genuine ambiguity in scope "
     "shall be construed reasonably in favour of turnkey delivery."),
    ("5A.5  Client-breach causation and mitigation.",
     "No extension of time, cost reimbursement, warranty limitation or liability exclusion arising "
     "from a Client act or omission shall apply unless the Contractor demonstrates that the Client "
     "breach directly and materially caused the relevant delay, defect or loss and that the "
     "Contractor took reasonable mitigation measures. Conversely, for Client operational "
     "non-compliance (for example low-SOC, denial of site access or loss of connectivity), the "
     "Client bears the burden of showing that it complied with its operational obligations."),
    ("10.6A  Witnessed performance testing and independent determination.",
     "Annual performance testing shall be performed under an agreed protocol and may be witnessed by "
     "the Client or its Independent Engineer. Disputes over SOH, usable AC capacity or RTE shall be "
     "determined by an independent expert or an ISO/IEC 17025-accredited laboratory under an agreed "
     "protocol where the Parties do not agree; the OEM's view is evidence but not determinative."),
    ("12.6  Force Majeure \u2014 limitations.",
     "Force Majeure shall not excuse: payment of amounts already due; warranty obligations for "
     "defects existing before the event; APG-extension obligations; insurance-maintenance "
     "obligations; reasonable mitigation; or delivery of documents and data capable of being "
     "delivered electronically."),
    ("1A.7A  Client withdrawal right.",
     "In addition to Section 1A.2, the Client may withdraw without liability (save for the refund "
     "mechanics in Section 1A.2) if: (a) the confirmed Delivery Schedule shows a Target PAC Date "
     "beyond the PAC Long-Stop defined in Section 1A.7; (b) the Confirmed Price Certificate is "
     "disputed by the Client in good faith and not resolved; (c) the APG validity does not cover "
     "the revised programme; or (d) the OEM quotation materially changes the equipment "
     "configuration or warranties from those on which this Agreement is based."),
    ("16.3A  Lender step-in and direct agreement.",
     "The Client's project-finance lender or security agent shall be entitled to step-in rights, "
     "cure rights, notice rights and the right to enforce or receive proceeds from the APGs, the OEM "
     "performance guarantee, warranties, insurance and performance undertakings. The Contractor "
     "shall, acting reasonably, enter into a direct agreement with the lender/security agent on "
     "customary terms that do not materially increase the Contractor's obligations or exposure."),
    ("10.8A  OEM Direct Warranty Undertaking \u2014 required features (procurement covenant).",
     "The Contractor shall procure that the OEM Direct Warranty Undertaking is irrevocable, directly "
     "enforceable by the Client and its lender/security agent, survives termination of any "
     "OEM/Distributor arrangement, includes a waiver of defences arising from disputes with the "
     "Contractor or Distributor, and does not permit unilateral amendment of the OEM Warranty Manual "
     "for this Project. This is an obligation to procure the DWU in that form; it does not warrant "
     "terms the OEM declines to give, in which case the Parties shall address the shortfall under "
     "Section 1A.5."),
    ("10.9B  APG issuing-bank acceptability (procurement covenant).",
     "The Contractor shall procure that each APG is issued by a bank acceptable to the Client and its "
     "lender and, where the issuing bank is outside the EU, is confirmed by or payable through an "
     "acceptable EU bank; permits partial and multiple drawings; remains independent of EPC disputes; "
     "and is extended automatically if PAC or Site delivery is delayed beyond the then-current expiry. "
     "This is an obligation to use all reasonable endeavours to procure these features."),
    ("18A  Cyber security, data and software (framework).",
     "The Contractor shall implement and maintain cyber-security controls for the BESS control "
     "systems consistent with recognised industrial-control-system good practice (IEC 62443 "
     "principles), including secure remote access, multi-factor authentication, role-based access "
     "control, audit logs, firmware-integrity controls, vulnerability management and incident "
     "reporting to the Client within twenty-four (24) hours of becoming aware of a material incident. "
     "The Client owns the operational data and is entitled to full raw-data export in open formats. "
     "The specific control set and any formal IEC 62443 / NIS2 certification shall be confirmed with "
     "the ICS/EMS provider (DISPERON) and recorded in the Technical Agreement before PAC."),
    # ── v6.3 — Stelios Constantinou track-change items (July 2026) ───────────────
    ("7A  Value Added Tax (VAT) — reverse charge.",
     "The EPC services provided by the Contractor (Lighthief Cyprus Ltd) to the Client (Galascope Ltd) "
     "are subject to Cyprus Value Added Tax on a reverse-charge basis under Article 11B of the Value "
     "Added Tax Law (L. 95(I)/2000). The Client is responsible for self-accounting the VAT on receipt "
     "of the Contractor's invoice; the Contractor will not charge or collect VAT on EPC services. All "
     "Contract Prices and milestone amounts in this Agreement are exclusive of VAT."),
    ("7.1  Payment schedule \u2014 decoupled streams (replaces body \u00a77.1).",
     "The Contract Price (EUR 3,462,849.40, ex VAT) is structured in two independent payment "
     "streams with separate milestone triggers:\n\n"
     "COMPONENT A \u2014 Equipment Supply Price (EUR 2,823,169.43, CIF Limassol):\n"
     "A1. Equipment Advance (25%, EUR 705,792.36): due within seven (7) days of the payment "
     "trigger (\u00a71A.3); secured by APG No.\u00a01; Contractor simultaneously pays the OEM advance.\n"
     "A2. Pre-Shipment (50%, EUR 1,411,584.72): due on joint Contractor/Client written "
     "confirmation of FAT and issuance of APG No.\u00a02, before shipment.\n"
     "A3. Delivery (15%, EUR 423,475.41): due within seven (7) days of equipment delivery to "
     "Site, completion of joint inspection and quantity/serial-number verification.\n"
     "A4. PAC (10%, EUR 282,316.94): due within seven (7) days of PAC issuance.\n\n"
     "COMPONENT B \u2014 EPC Services Price (EUR 639,679.97):\n"
     "M1a. Mobilisation + Engineering (20%, EUR 127,935.99): due at the payment trigger \u2014 "
     "professional services commence on this date: protection engineering, DSO/EAC application "
     "preparation, procurement management, ETEK design and project programme.\n"
     "M1b. Civil/Platforms complete (20%, EUR 127,935.99): due on inspection and written "
     "approval of concrete platforms by the Client or its representative.\n"
     "M2. Delivery + Placement (15%, EUR 95,951.99): due on completion of joint equipment "
     "inspection and placement of containers and MV skid on platforms.\n"
     "M3. Cabling, Earthing and LPS complete (15%, EUR 95,951.99): due on completion of all "
     "MV/LV cabling, earthing grid and DEHN LPS/SPD, with test records submitted.\n"
     "M4. EMS/SCADA + Cold Commissioning (15%, EUR 95,951.99): due on EMS/SCADA installation "
     "and cold-commissioning report.\n"
     "M5. PAC (10%, EUR 63,967.99): due within seven (7) days of PAC issuance.\n"
     "Retention (5%, EUR 31,984.00): released upon issuance of FAC at end of DLP (\u00a79.2A).\n\n"
     "APG No.\u00a01 secures A1 (EUR 705,792.36); APG No.\u00a02 secures A2 (EUR 1,411,584.72). Component "
     "B payments are secured by the Contractor\u2019s performance obligations and the performance "
     "guarantee. The first combined payment due at trigger = A1 + M1a = EUR 833,728.35. "
     "Payments are due within seven (7) Business Days of the relevant trigger. "
     "Cross-reference key: elsewhere in this Agreement, \u201cSection 7.1(a)\u201d means the "
     "advance payment (A1 above); \u201cSection 7.1(b)\u201d means the pre-shipment payment (A2); "
     "and \u201cSection 7.1(c)\u201d means the PAC payment (A4 and M5 together, i.e. the "
     "Component A and Component B amounts both due on PAC)."),
    ("7B  Component B \u2014 EPC Works Milestone Payment Schedule (cross-reference).",
     "For the detailed achievement specifications and verification criteria applicable to each "
     "Component B milestone (M1a through Retention), see Section 7.1 above. Payments are due "
     "within seven (7) Business Days of the Client\u2019s written confirmation of milestone "
     "completion. Each milestone is linked to the Delivery Schedule confirmed at Connection "
     "Terms; exact dates are set at that stage."),
    ("9.1A  PAC acceptance matrix.",
     "PAC shall be conditional on successful completion and documented sign-off of the following "
     "acceptance checks, each to be witnessed by the Client or its appointed representative: "
     "(a) usable AC energy capacity \u226598.50% of rated (Section 9.1); "
     "(b) Round-Trip Efficiency \u226584% at the AC system boundary (Section 9.1); "
     "(c) BMS/SCADA communications and EMS command-execution verified; "
     "(d) all protection relays set, tested and verified by the Contractor's ETEK engineer; "
     "(e) DSO/SCADA telemetry and remote-trip verified operational; "
     "(f) fire-detection/suppression system tested and certified; "
     "(g) emergency-stop and safety systems tested; "
     "(h) HVAC/liquid-cooling operational; "
     "(i) grid-forming/VSG and black-start readiness verified (where DSO activation is required). "
     "The full FAT/SAT test protocol and pass/fail values shall be set out in Schedule A, agreed "
     "before order placement. Where measured capacity or RTE falls below the Contract guarantees "
     "but above the PAC hard floor, Section 9.1B applies."),
    ("9.2A-EXP  FAC \u2014 independent expert sign-off.",
     "In addition to the matters in Section 9.2A, the Final Acceptance Certificate (FAC) shall "
     "require written acceptance by the Client or its appointed Independent Engineer confirming that "
     "all defects, punch-list items, outstanding commissioning obligations and performance test "
     "documentation have been completed to their reasonable satisfaction. Such acceptance shall not "
     "be unreasonably withheld or delayed beyond ten (10) Business Days. Retention is released only "
     "after FAC issuance."),
    ("10.6-SOH  SOH guarantee \u2014 full annual curve reference.",
     "The SOH guarantees in Section 10.6 correspond to the 0.25P, 1-cycle-per-day degradation "
     "profile of the Linyang Power Atlantic 5 MWh BESS (BOL DC capacity 5,015 kWh), which matches "
     "this Project's actual ~4-hour system duration (Galascope 1: 5 MW/20.06 MWh; Galascope 2: "
     "2.5 MW/10.03 MWh) \u2014 confirmed against Linyang's Power Atlantic 5MWh Degradation Curve "
     "document (Schedule A, SA-05). The full year-by-year guaranteed SOH curve is set out in "
     "Schedule 5 of the LTSA (LCY-LTSA-GAL-2026), which is incorporated by reference into this "
     "Agreement. The key contractual milestone values are: Year 0 (COD) 98.5%, Year 5 \u226586.78%, "
     "Year 10 \u226579.58%, Year 15 \u226573.61% (0.25P, 1 cycle/day, per Linyang Power Atlantic "
     "degradation curve document). Year 10 and Year 15 guarantees apply only where the OEM "
     "extended warranty for the relevant period has been purchased."),
    ("10.6-EXT  Extended BESS warranty \u2014 optional, priced (BESS only; PCS/MVS excluded).",
     "The base warranty under Section 10 runs for five (5) years from PAC. The Client may, at its "
     "sole option and by written election to the Contractor at any time before expiry of the "
     "then-current warranty period, extend the OEM (Linyang) BESS Performance and Product Warranty "
     "for one or both of the following periods, at the following rates (official Linyang Cyprus "
     "pricing for the Project, EUR/MWh/Year, excluding VAT): "
     "(a) Years 6\u201310: EUR 913.92/MWh/Year (EUR 27,417.60/Year for the combined 30 MWh System "
     "Capacity, i.e. Galascope 1: 20 MWh + Galascope 2: 10 MWh); "
     "(b) Years 11\u201315: EUR 1,157.62/MWh/Year (EUR 34,728.60/Year for the combined 30 MWh System "
     "Capacity). "
     "This extension covers the BESS (cells, modules, racks and BMS) only. The Client has confirmed "
     "it does not require, and this Agreement does not offer, an extended warranty on the PCS or "
     "MV Skid (MVS) beyond the base warranty period; the Contractor's quoted PCS/MVS extended-"
     "warranty rate (EUR 747.76/MWh/Year for Years 6\u201310 and EUR 926.10/MWh/Year for Years 11\u201315) "
     "is disclosed for completeness but is not part of the Contract Price or any priced option under "
     "this Agreement unless separately elected in writing. Neither extension is included in the "
     "Contract Price unless and until the Client elects it in writing; if elected, the applicable fee "
     "is invoiced and payable in addition to the Base Service Fee under the LTSA (Schedule 2), which "
     "is the vehicle through which the extended warranty is administered and delivered, consistent "
     "with Section 11.1. For the avoidance of doubt, the SOH guarantees for Year 10 and Year 15 "
     "under Section 10.6-SOH are conditional on the corresponding extension in this Section having "
     "been elected and paid for."),
    ("11.1-PERF  Park performance responsibility while LTSA is in force (concede).",
     "For as long as the operation and maintenance of the BESS is being carried out by the "
     "Contractor under a confirmed and executed LTSA (LCY-LTSA-GAL-2026), the Contractor "
     "acknowledges responsibility for park performance (including availability and SOH performance) "
     "and the Client's primary remedies in respect of such performance are those set out in the "
     "LTSA and its schedules. This acknowledgement is without prejudice to the Contractor's "
     "obligations and the Client's rights under this Agreement (including Sections 9.1, 9.1B, 10.6 "
     "and 13), and the aggregate liability caps in Section 13.2 remain in full force."),
    ("4.6A  Variation items \u2014 available at separately priced extra.",
     "The following items are excluded from the Contract Price and are available as separately "
     "priced variations upon written Change Order agreed in advance: "
     "(a) Crusher-run / crane-access ground preparation. Where the crane contractor's bearing-capacity "
     "requirements cannot be met by the existing site ground conditions, a compacted aggregate "
     "(crusher-run) base for crane outrigger pads and access routes will be required. This item is "
     "charged at cost plus fifteen percent (15%) on the basis of a site survey by the crane "
     "contractor and a written quotation agreed before works commence. "
     "(b) PCC switchgear bay extension / MV modification. Where connection of the BESS to the "
     "Client\u2019s existing PV transformer requires the installation of a new MV switchgear bay "
     "and/or reconfiguration of existing busbar in-feeds (including merging or re-routing of "
     "existing cable feeders to free a bay for the BESS incomer), this work will be separately "
     "priced per quotation. Indicative range: EUR 20,000\u201335,000 subject to site survey of "
     "existing switchgear by the Contractor\u2019s ETEK engineer and confirmation of panel "
     "compatibility, busbar rating and DSO approval requirements."),
    ("7C  Import VAT and customs \u2014 importer of record.",
     "Import duty applicable to the BESS equipment (Component A) is included within the Contract "
     "Price. Import VAT on the BESS equipment, applied at the port of discharge (CIF Limassol), is "
     "excluded from the Contract Price and is not part of the agreed payment schedule. The Parties "
     "shall agree in writing, before the equipment is dispatched, which Party acts as importer of "
     "record for Cypriot Customs purposes. The Party designated as importer of record shall be "
     "responsible for paying and reclaiming import VAT directly with the Cyprus Tax Department. "
     "Neither Party shall be required to pre-fund import VAT on the other\u2019s behalf."),
    ("10.9-APG  Consistency with agreed APG wording.",
     "The Advance Payment Guarantees issued by the OEM's bank shall be consistent in all material "
     "respects with the agreed APG form (as set out in the bank-change instruction issued by the "
     "Contractor to the OEM's bank). Any material discrepancy between the issued instruments and "
     "the agreed APG wording shall be notified to the Client within five (5) Business Days of the "
     "Contractor becoming aware, and the Contractor shall use all reasonable endeavours to procure "
     "correction before the relevant payment milestone falls due."),
]


def _replace_in_paragraph(p, old, new):
    if old in "".join(r.text for r in p.runs):
        full = "".join(r.text for r in p.runs).replace(old, new)
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = full
        else:
            p.add_run(full)


def _integrate_epc_annex(doc):
    """Relocate every ANNEX clause into its numbered position in the body.
    Colour-coding is preserved; only the clause's location changes."""
    A = dict(enumerate(ANNEX))

    def ins(idx, anchor_prefix, headline_override=None, occurrence=1):
        headline, text = A[idx]
        if headline_override:
            headline = headline_override
        anchor = _find_para(doc, anchor_prefix, occurrence)
        if anchor is None:
            raise RuntimeError(f"anchor not found for ANNEX[{idx}] ({headline!r}): {anchor_prefix!r}")
        col = _amd_color(headline, NAVY)
        return _insert_clause_after(anchor, headline, text, col)

    def rep(idx, target_prefix, headline_override=None, occurrence=1):
        headline, text = A[idx]
        if headline_override:
            headline = headline_override
        target = _find_para(doc, target_prefix, occurrence)
        if target is None:
            raise RuntimeError(f"replace-target not found for ANNEX[{idx}] ({headline!r}): {target_prefix!r}")
        col = _amd_color(headline, NAVY)
        return _replace_clause(target, headline, text, col)

    # --- Section 1A: conditions precedent / order trigger -------------------
    p = ins(0, "1A.3")                                             # -> 1A.4
    p = ins(1, "1A.4  Nature of EPC Execution.")                   # -> 1A.5
    p = ins(33, "1A.5  Companion Documents")                       # -> 1A.5A
    p = ins(2, "1A.5A  LTSA")                                      # -> 1A.5B
    p = ins(3, "1A.5B  Pre-Shipment Payment")                      # -> 1A.6
    p = ins(8, "1A.6  Bilateral Walk-Away.")                       # -> 1A.7
    p = ins(42, "1A.7  Client withdrawal on the delivery timeline.")  # -> 1A.7A
    p = ins(35, "1A.7A  Client withdrawal right.")                 # -> Schedule A note

    # --- Section 4: Scope -----------------------------------------------
    # Anchor on the base clause's own last list item, (e), not the bare
    # "4.4 EMS Integration" heading — so 4.4A/4.4B land after the original
    # (a)-(e) EMS list rather than pushing it below both new clauses.
    p = ins(30, "(e) The Contractor is not responsible for EMS optimisation",
            headline_override="4.4A  EMS Provider (DISPERON).")   # -> 4.4A
    p = ins(31, "4.4A  EMS Provider")                              # -> 4.4B
    p = ins(38, "4.5")                                             # -> 4.6
    p = ins(55, "4.6  Excluded items")                             # -> 4.6A

    # --- Section 5A: Client obligations ----------------------------------
    # Anchor on each base clause's own last list item, not the bare heading,
    # so the new clause lands after the original list rather than before it.
    p = ins(9, "(e) Obligation to provide electrical connection point")  # -> 5A.1A
    p = ins(39, "(d) No Contractor Liability for delays or losses arising")  # -> 5A.5

    # --- Section 6.1: price basis ------------------------------------------
    # Labelled 6.1A/6.1B/6.1C (not parenthetical letters) because the base §6.1
    # clause already has its own internal (a)-(f) list (embedded via line breaks
    # within the same paragraph as the "6.1 Contract Price" heading, and cross-
    # referenced elsewhere as Section 6.1(c)-(d)); reusing (f)/(g)/(h) at the
    # 6.1-prefix level would collide with that existing internal (f).
    p = ins(6, "6.1 Contract Price")                               # -> 6.1A
    p = ins(4, "6.1A  Price basis")                                # -> 6.1B
    p = ins(5, "6.1B  Price Basis Certificates.")                  # -> 6.1C

    # --- Section 7: Payment -------------------------------------------------
    rep(48, "7.1 Payment Milestones",
        headline_override="7.1  Payment schedule \u2014 decoupled streams.")
    p = ins(24, "7.1  Payment schedule",
            headline_override="7.1B  Independent Engineer sign-off \u2014 PAC payment.")  # -> 7.1B (see below: 7.1A
                                                                    # is inserted right after this same anchor
                                                                    # later, in the cleanup pass, and lands
                                                                    # ahead of this clause positionally)
    p = ins(49, "7.1B  Independent Engineer sign-off",
            headline_override="7B  Component B \u2014 EPC Works Milestone Payment Schedule (cross-reference).")  # -> 7B
    p = ins(56, "7B  Component B")                                 # -> 7C
    # 7A (VAT reverse-charge) intentionally not duplicated inline — its content
    # is already stated in body §7.2 (fixed to the Article 11B reverse-charge
    # treatment agreed with the Client).
    p = ins(21, "(d) The Client shall not withhold, set off, or deduct")  # -> 7.4A

    # --- Section 8: Delivery -------------------------------------------------
    p = ins(7, "8.3 Typical installation period",
            headline_override="8.3A  Delivery & Shipping Schedule and Target PAC.")  # -> 8.3A
    p = ins(17, "8.3A  Delivery & Shipping Schedule",
            headline_override="8.3B  Pre-Shipment Inspection / Factory Acceptance Test (FAT).")  # -> 8.3B
    # Pre-existing base-document defect (predates this v6.3 pass): item "(vii)
    # Upstream Force Majeure" had been tacked onto the "8.4.6 Exclusions" heading
    # line itself, and the old 8.4.7A paragraph sat between that heading and its
    # own (i)-(vi) exclusions list — so the list read heading -> (vii) -> 8.4.7A
    # -> (i)-(vi), with item (vii) appearing before, not after, items (i)-(vi).
    # Fix: restore a clean "8.4.6 Exclusions" heading, move (vii) to the end of
    # the (i)-(vi) list, and place the full 8.4.7A clause after that complete list.
    _p_646 = _find_para(doc, "8.4.6 Exclusions; (vii) Upstream Force Majeure")
    if _p_646 is not None:
        _replace_in_paragraph(_p_646, _p_646.text, "8.4.6 Exclusions")
    _p_ldslist = _find_para(doc, "LDs not applicable for:")
    if _p_ldslist is not None:
        _replace_in_paragraph(
            _p_ldslist,
            "LDs not applicable for: (i) Force Majeure; (ii) Client-caused delays; "
            "(iii) grid authority delays; (iv) Client change orders; (v) customs delays "
            "with timely shipment proved; (vi) Client\u2019s late confirmation of technical "
            "agreements or drawings.",
            "LDs not applicable for: (i) Force Majeure; (ii) Client-caused delays; "
            "(iii) grid authority delays; (iv) Client change orders; (v) customs delays "
            "with timely shipment proved; (vi) Client\u2019s late confirmation of technical "
            "agreements or drawings; (vii) Upstream Force Majeure declared by the OEM and "
            "notified to the Client per Section 8.4.7A.")
    _old_8471a = _find_para(doc, "8.4.7A Upstream Force Majeure")
    if _old_8471a is not None:
        _delete_para(_old_8471a)
    p = ins(10, "LDs not applicable for:",
            headline_override="8.4.7A Upstream Force Majeure (OEM Flow-Through).")  # -> 8.4.7
    # anchor on 8.4.7's own body text (not the bare header), so 8.4.8 doesn't
    # land in between the 8.4.7 header and its own content.
    p = ins(20, "Target PAC Date extends day-for-day. Contractor notifies")  # -> 8.4.8

    # --- Section 9: Acceptance -------------------------------------------------
    p = ins(50, "9.1 Provisional Acceptance (PAC)")                # -> 9.1A
    p = ins(27, "9.1A  PAC acceptance matrix.",
            headline_override="9.1A-EXT  Extended rejection window.")  # -> 9.1A-EXT
    p = ins(19, "9.1A-EXT  Extended rejection window.")            # -> 9.1B
    p = ins(18, "A Final Acceptance Certificate (FAC) shall be issued at the end")  # -> 9.2A
    p = ins(51, "9.2A  Final Acceptance Certificate")              # -> 9.2A-EXP

    # --- Section 10: Warranty / performance --------------------------------
    # Anchor on the base guarantee list's own last item, (d) Remedy — not the
    # bare "10.6 Performance Guarantees" heading — so the three new clauses
    # (10.6-SOH, 10.6-EXT, 10.6A) land after the original (a)-(d) guarantees
    # list rather than pushing it below all three amendments.
    p = ins(52, "(d) Remedy: Contractor procures replacement modules")  # -> 10.6-SOH
    p = ins(53, "10.6-SOH  SOH guarantee")                         # -> 10.6-EXT
    p = ins(40, "10.6-EXT  Extended BESS warranty")                # -> 10.6A
    p = ins(32, "The Contractor shall alert the Client within twenty-four (24) hours of SOC dropping below ten percent (10%) for as long as the LTSA is in effect")  # -> 10.5A
    p = ins(44, "10.8 OEM Direct Warranty Undertaking")            # -> 10.8A
    p = ins(12, "10.9 Advance Payment Guarantee",
            headline_override="10.9(a)  APG issuance & validity \u2014 confirmed mechanism.")  # -> 10.9(a)
    p = ins(11, "10.9(a)  APG issuance",
            headline_override="10.9(b)  Advance Payment Guarantee No. 1 \u2014 advance tranche.")  # -> 10.9(b)
    p = ins(13, "10.9(b)  Advance Payment Guarantee No. 1",
            headline_override="10.9(c)  Advance Payment Guarantee No. 2 \u2014 pre-shipment tranche.")  # -> 10.9(c)
    p = ins(14, "10.9(c)  Advance Payment Guarantee No. 2",
            headline_override="10.9(d)  APG validity \u2014 to PAC.")  # -> 10.9(d)
    p = ins(15, "10.9(d)  APG validity")                           # -> 10.9(e)
    p = ins(16, "10.9(e)  APG trigger.")                           # -> 10.9A
    p = ins(45, "10.9A  Security layering")                        # -> 10.9B
    p = ins(57, "10.9B  APG issuing-bank acceptability")           # -> 10.9-APG
    p = ins(36, "10.9-APG  Consistency")                            # -> 10.10

    # --- Section 11: integrated document suite ------------------------------
    p = ins(54, "11.1 The operational availability guarantee",
            headline_override="11.1-PERF  Park performance responsibility while the LTSA is in force.")  # -> 11.1-PERF
    p = ins(23, "11.4 This Agreement may be executed on its own",
            headline_override="11.5  Integrated document suite and interpretation.")  # -> 11.5

    # --- Section 12: Force Majeure -------------------------------------------
    p = ins(41, "12.5 Extended Force Majeure")                    # -> 12.6

    # --- Section 13: Liability ------------------------------------------------
    rep(28, "13.2 The Contractor\u2019s total aggregate liability",
        headline_override="13.2  Limitation of Liability.")
    p = ins(25, "(d) Manufacturing defects in OEM equipment")  # -> 13.3A
    p = ins(29, "13.4 The tiered structure")                       # -> 13.5

    # --- Section 14: Insurance -------------------------------------------------
    p = ins(22, "14.4 Proof of Insurance")                         # -> 14.5

    # --- Section 16: General --------------------------------------------------
    p = ins(43, "(b) Upon lender enforcement, Contractor recognises the lender as Client")  # -> 16.3A
    p = ins(26, "16.3A  Lender step-in")                           # -> 16.3B

    # --- Sections 18 / 19: Data protection / Regulatory --------------------
    p = ins(46, "18. DATA PROTECTION")                            # -> 18A
    p = ins(34, "(b) Contractor: All technical and regulatory responsibilities")  # -> 19.12
    p = ins(37, "19.12  Technical deliverables")                   # -> 19.13

    # Delete pre-v6 body paragraphs now fully superseded by the relocated
    # clauses above (they were never removed when 7.1/10.9/13.2 were first
    # amended, so they were silently duplicating/contradicting the new text).
    for prefix in [
        "(a) Payment schedule: The Contract Price is paid in two decoupled streams",
        "(b) 55% payment prior to shipment of BESS equipment",
        "(c) Combined PAC payment (EUR 346,284.94)",
        "(a) The Contractor shall procure that the OEM provides an unconditional, irrevocable Advance Payment Guarantee",
        "(b) Two Advance Payment Guarantees are provided, each issued by the OEM's bank",
        "(c) The APG shall be delivered prior to the advance payment becoming due",
        "(d) Each APG shall remain valid until the earlier of issuance of PAC",
        "(e) No APG is required against Component B",
        "(f) The Contractor shall also procure that the OEM provides a performance bond",
        "(a) Claims relating to the Contractor's own EPC services, installation and non-OEM works",
        "(b) All other contractual breaches (excluding warranty, fraud, wilful misconduct, and manufacturing defects)",
        "(c) Fraud, wilful misconduct, and manufacturing defects in OEM equipment: uncapped.",
    ]:
        stale = _find_para(doc, prefix)
        if stale is not None:
            _delete_para(stale)

    # The old 7.1(d) (interim performance-guarantee arrangement pending the
    # OEM's 9-month PG extension) is still live content, not superseded — move
    # it to sit right after §7.1 and drop the orphaned "(d)" lettering (its
    # siblings (a)-(c) have just been deleted above as pure duplicates). Labelled
    # Labelled 7.1A (not 7.1B) because it lands, positionally, immediately after
    # §7.1 and before the "Independent Engineer sign-off" clause above (which is
    # itself labelled 7.1B for that reason, even though 7.1B is defined first in
    # the script) — letters follow document order, not insertion order.
    old_d = _find_para(doc, "(d) 5% Retention released at the end of the Defects Liability Period")
    if old_d is not None:
        full_text = old_d.text[len("(d) "):]
        anchor = _find_para(doc, "7.1  Payment schedule")
        _insert_clause_after(anchor, "7.1A  Interim performance guarantee.", full_text, NAVY)
        _delete_para(old_d)

    # Now that every clause lives at its own numbered position, cross-references
    # to "Annex V6 §X" throughout the body should just read "§X" (most specific
    # patterns first, generic fallback last).
    _clean_annex_refs(doc, [
        ("Annex V6 \u00a7\u00a7", "\u00a7\u00a7"),
        ("Annex V6 \u00a7", "\u00a7"),
        ("Annex V6 Section", "Section"),
        (" (Annex V6)", ""),
        ("(Annex V6) ", ""),
        ("(Annex V6)", ""),
        ("Annex V6, ", ""),
        ("Annex V6", "this Agreement"),
    ])


def build_epc_v6():
    doc = Document(str(EPC_SRC))

    # Version stamp: bump cover page from v5.1 / May 2026 to v6.0 / June 2026
    for p in doc.paragraphs:
        _replace_in_paragraph(p, "Version: 5.1", "Version: 6.3")
        _replace_in_paragraph(p, "Date: May 2026", "Date: July 2026")
        # Add v6.0 entry to VERSION HISTORY block
        _replace_in_paragraph(
            p,
            "v5.1 — May 2026: Indicative Price + Confirmed Price mechanism (raw material / FX adjustment), "
            "LD Value split (Comp A delivery / Comp B commissioning), upstream FM flow-through, "
            "12-month Connection Terms longstop, 30-day advance payment trigger",
            "v5.1 — May 2026: Indicative Price + Confirmed Price mechanism (raw material / FX adjustment), "
            "LD Value split (Comp A delivery / Comp B commissioning), upstream FM flow-through, "
            "12-month Connection Terms longstop, 30-day advance payment trigger\n"
            "v6.0 — June 2026: Amendment Schedule (integrated into the body at v6.3) — dual APG to PAC, DWU as CP, "
            "manufacturing-defect carve-out, delivery at Connection Terms, price re-anchoring, "
            "VAT clarity, title/APG alignment, EMS entity update\n"
            "v6.1 — July 2026: Client redline (Dino/Anastasios) — accepted: grid-forming "
            "included in Contract Price, EMS-affiliate warranty carve-out, 4.4B interface "
            "responsibility, LTSA as simultaneous-execution CP, 19.12 deliverables list, "
            "Schedule A spec-freeze, insurance before first payment; negotiated: integrated-suite "
            "most-favourable-to-Client interpretation (11.5), Independent Engineer PAC-payment gate "
            "(7.1B), additional liability carve-outs latent + safety/grid-code (13.3A), FAC "
            "reinstated (9.2A), PAC shortfall remedy capped 10% Component A (9.1B), binding "
            "Delivery Schedule with OEM FM preserved (8.4.8), set-off with ascertained/agreed "
            "qualifier (7.4A), Lighthief International performance undertaking by confirmation "
            "letter (16.3B)\n"
            "v6.2 — July 2026: Timotheos / technical review items incorporated — EPC: excluded-items "
            "scope certainty (4.6), Client-breach causation & burden (5A.5), witnessed testing & "
            "independent expert (10.6A), Force Majeure limitations (12.6), Client withdrawal right "
            "(1A.7A), lender step-in & direct agreement (16.3A), DWU required-features & APG "
            "issuing-bank procurement covenants (10.8A/10.9B), cyber-security framework (18A); LTSA: "
            "restoration/resolution SLA (8.2A), monitoring/alerting (5.2A/8.1A), major-incident "
            "procedure (8.6A), 24/7 safety response (8.7A), warranty-preservation file (6.1A/6.3A), "
            "data ownership (5.8A), chronic underperformance (9.8), transition assistance (13.6), "
            "expert determination (20.5) and related technical clauses\n"
            "v6.3 — July 2026: Stelios Constantinou (Galascope) track-change items — VAT reverse-charge "
            "(7A), Component B milestone payment schedule M1-M5 (7B), PAC acceptance matrix (9.1A), "
            "FAC independent-expert sign-off (9.2A-EXP), PAC capacity criterion raised to 98.5% SOH, "
            "\u00a711 park-performance concession (11.1-PERF), SOH full annual curve "
            "reference (10.6-SOH), APG-consistency clause (10.9-APG), late-payment interest ECB+3%, "
            "joint client inspection at pre-shipment, 5% PG extended to 6 months post-PAC, "
            "fencing/security note, 90-120 day delivery-to-PAC installation timeframe, operation "
            "clauses LTSA+EMS cross-references, companion-document acceptance bundle at the "
            "Connection-Terms trigger \u2014 APG, LTSA, Technical Agreement sign-off, DWU and EMS "
            "Subscription Addendum all gate the advance payment (1A.3 / 1A.5); PAC Long-Stop defined "
            "at fifteen (15) months from receipt of Connection Terms, replacing the subjective "
            "withdrawal test in 1A.7 with an objective threshold shared with 1A.7A(a)\n"
            "v6.3 pre-signing corrections (30 Jul 2026, internal legal review): corrected the OEM "
            "Performance Guarantee back-to-back period claim in the retention/PG clause to match "
            "the OEM Performance Guarantee\u2019s actual current 3-month validity (extension "
            "requested separately from the OEM per Anastasios's 12-month ask); re-anchored the "
            "Contractor\u2019s Professional Indemnity insurance in-force date to the Order Date "
            "under Section 1A.4 (Section 14.4/14.5(e)), consistent with the CAR policy\u2019s own "
            "attachment point (this narrows only the informal \u201cfrom the Effective Date\u201d "
            "language the Contractor had added beyond Dino's actual §14.4 ask, which remains "
            "unchanged: insurance evidence before first payment)\n"
            "v6.3 optional extended BESS warranty pricing (30 Jul 2026): new Section 10.6-EXT discloses "
            "the official Linyang Cyprus pricing for an optional, Client-elected extension of the BESS "
            "Performance and Product Warranty for Years 6\u201310 (EUR 913.92/MWh/Year) and Years 11\u201315 "
            "(EUR 1,157.62/MWh/Year), delivered under the LTSA (Schedule 2) if elected. Scope is BESS "
            "only \u2014 the PCS/MVS extended-warranty option (EUR 747.76 and EUR 926.10/MWh/Year "
            "respectively) is disclosed for completeness but is not offered, per the Client's instruction; "
            "LTSA Schedule 2's PCS + MVS Product Warranty rows are marked \u201cnot offered\u201d "
            "accordingly. Neither option is included in the Contract Price or selected by default.\n"
            "v6.3 final sweep (30 Jul 2026): fixed EPC §7.2 to match the Article 11B domestic "
            "reverse-charge VAT treatment (Contractor does not add "
            "VAT; Client self-accounts); fixed EPC §10.5 to remove the \u201cTier A or above\u201d "
            "gating on the Contractor-monitoring-failure carve-out, so warranty protection under (i)"
            "\u2013(iv) applies for as long as any confirmed LTSA is in force, matching Stelios's "
            "insertion and Annex V6 §10.5A; corrected the stale LTSA §11.4-mirror \u201cCOMPANION "
            "DOCUMENTS\u201d section (was still tagged \u201cv6.0\u201d and \u201call three documents "
            "simultaneously\u201d) to reference v6.3 and the LTSA's own CP status per §1A.5A; added "
            "LTSA §5.5 Client (or nominated representatives) training on the Monitoring Platform, a "
            "Timotheos ask that had not yet been drafted or logged in the response matrix; clarified "
            "in the response matrix that the client's SOH-linked automatic LTSA-cancellation ask "
            "(§10.4) was a deliberate counter, not an oversight. Ran a full extraction of every "
            "tracked-change insertion in Stelios's, Timotheos's (EPC + LTSA) and Anastasios's source "
            "redlines (156 items) and confirmed each is now reflected in the contract text or is a "
            "disclosed, tracked counter/open item.\n"
            "v6.3 SOH curve correction (30 Jul 2026): direct read of Linyang's \u201cPower Atlantic "
            "5MWh Degradation Curve\u201d source PDF confirmed both Galascope 1 (5 MW/20.06 MWh) and "
            "Galascope 2 (2.5 MW/10.03 MWh) are ~4-hour systems, i.e. the 0.25P operating condition "
            "(not 0.5P, a prior working assumption). Corrected EPC §10.6(a)/§10.6-SOH and LTSA §10.2/"
            "Schedule 5 from a blended 0.5P/0.25P curve (Y5 86.26%, Y10 79.58%, Y15 72.45%) to the "
            "verified 0.25P, 1-cycle/day curve throughout (Y1 94.62%, Y2 91.77%, Y3 89.91%, Y4 88.00%, "
            "Y5 86.78%, Y6 84.97%, Y7 83.83%, Y8 82.25%, Y9 81.06%, Y10 79.58% unchanged, Y15 73.61%). "
            "This matches Stelios's original request exactly and is an OEM-data correction, not a new "
            "client concession \u2014 it raises, not lowers, the Client-facing guarantee. The OEM Direct "
            "Warranty Undertaking (04) is our own draft, so §2.1/§2.2 SOH figures and §2.4 cycle life "
            "(7,000\u21928,000) were corrected in the same pass, closing what would otherwise have been a "
            "~1.8\u20133.2 point back-to-back gap at Year 5/15 between the DWU and the EPC/LTSA. Linyang is "
            "asked (DWU-and-PG-Change-Instructions memo) to confirm agreement with the updated DWU figures "
            "before signing, or flag if they prefer a more conservative floor; not a blocker to signing "
            "the EPC.")

    # EMS provider entity change: Lighthief EU BESS Ltd -> Disperon Sp. z o.o.
    for p in doc.paragraphs:
        _replace_in_paragraph(
            p, "Lighthief EU BESS Ltd",
            "Disperon Sp. z o.o. (trading as DISPERON, under Lighthief International Ltd)")
        # §11 reframe — remove the "SEPARATION / not in this Agreement" wording the client objects to,
        # while keeping availability's remedy/cap anchored to the LTSA (protects the liability ring-fence).
        _replace_in_paragraph(p, "11. LTSA SEPARATION",
                              "11. RELATIONSHIP WITH THE LTSA AND COMPANION DOCUMENTS")
        _replace_in_paragraph(
            p,
            "11.1 97% availability guarantee, availability LDs, and extended warranty (years 6\u201315) are not in this Agreement.",
            "11.1 The operational availability guarantee (97%), availability liquidated damages, and the "
            "extended warranty (Years 6\u201315) are provided under the LTSA (LCY-LTSA-GAL-2026), which is "
            "executed before the advance payment falls due (Annex V6 \u00a71A.5A). The "
            "Client shall at all times have the benefit of these protections; the "
            "remedies and liability limits applicable to them are those set out in the LTSA.")
        _replace_in_paragraph(
            p, "11.2 Apply only under a separately executed LTSA.",
            "11.2 This Agreement and the LTSA are intended to operate together as a single commercial "
            "arrangement for the Project.")
        # v6.3 STELIOS inline replacements ──────────────────────────────────────────────────
        # PAC capacity criterion: 95% → 98.50% SOH (consistent with Year-0 COD = 98.5%)
        _replace_in_paragraph(p,
            "The BESS operates at not less than 95% of rated energy capacity;",
            "The BESS operates at not less than ninety-eight point five percent (98.50%) of original "
            "rated capacity (State of Health at COD); where measured SOH at PAC is between 95% (hard "
            "floor) and 98.50%, PAC is issued and Section 9.1B (shortfall remedy) applies;")
        # Late payment interest: ECB+8% → ECB+3% (agreed per Stelios)
        _replace_in_paragraph(p, "ECB + 8%", "ECB + 3%")
        _replace_in_paragraph(p, "ECB + 8 %", "ECB + 3%")
        _replace_in_paragraph(p, "ECB reference rate + 8%", "ECB reference rate + 3%")
        _replace_in_paragraph(p, "ECB + eight", "ECB + three (3)")
        # Pre-shipment: add joint client/representative inspection (curly apostrophe \u2019 in source)
        _replace_in_paragraph(p,
            "55% payment prior to shipment of BESS equipment, upon Contractor\u2019s written confirmation "
            "of successful Factory Acceptance Testing (FAT)",
            "55% payment prior to shipment of BESS equipment, upon Contractor\u2019s and Client\u2019s or "
            "Client\u2019s representative\u2019s joint written confirmation of successful Factory Acceptance "
            "Testing (FAT)")
        # DLP retention: add 5% PG for 6 months (counter to Stelios's 9 months)
        _replace_in_paragraph(p,
            "5% Retention released at the end of the Defects Liability Period (three (3) months "
            "after PAC), subject to Section 7.4.",
            "5% Retention released at the end of the Defects Liability Period (three (3) months "
            "after PAC), subject to Section 7.4. In addition, the Contractor shall maintain a "
            "five percent (5%) performance guarantee valid for nine (9) months after PAC in the "
            "form of a corporate performance guarantee issued by Lighthief Cyprus Ltd directly to "
            "the Client, confirming the Contractor\u2019s obligations to remedy defects and meet "
            "performance guarantees during that period (the Defects Liability Period plus six (6) "
            "additional months). The Lighthief Cyprus Ltd corporate performance guarantee is "
            "backed back-to-back by the OEM (Linyang) Performance Guarantee (Schedule C) for the "
            "initial three (3)-month Defects Liability Period, which is the OEM Performance "
            "Guarantee\u2019s validity period as currently drafted. For the additional six (6) "
            "months of the Lighthief Cyprus Ltd guarantee, Lighthief Cyprus Ltd bears the "
            "obligation directly pending the extension requested from the OEM to match the full "
            "nine (9)-month period; any such extension, once agreed, shall be reflected in an "
            "updated OEM Performance Guarantee without need to amend this Agreement. The "
            "Lighthief Cyprus Ltd guarantee does not limit the Client\u2019s rights under the "
            "OEM guarantee or under this Agreement.")
        # §7.1 body: replace blended payment clauses with pointer to Annex V6 §7.1
        _replace_in_paragraph(p,
            "(a) 30% advance payment within seven (7) days of the payment trigger under Section 1A.3;",
            "(a) Payment schedule: The Contract Price is paid in two decoupled streams "
            "(Component A \u2014 Equipment and Component B \u2014 EPC Services) per Annex V6 \u00a77.1, "
            "which replaces this body clause. First combined payment "
            "(EUR 833,728.35 = A1 EUR 705,792.36 + M1a EUR 127,935.99) "
            "is due within seven (7) days of the payment trigger under Section 1A.3.")
        _replace_in_paragraph(p,
            "(b) 55% payment prior to shipment of BESS equipment, upon Contractor\u2019s written "
            "confirmation of successful Factory Acceptance Testing (FAT) and photographic evidence "
            "of readiness;",
            "(b) Component A pre-shipment (50%, EUR 1,411,584.72): due on Contractor\u2019s and "
            "Client\u2019s or Client\u2019s representative\u2019s joint written confirmation of FAT "
            "and issuance of APG No.\u00a02. Component B works milestones payable per Annex V6 \u00a77.1.")
        _replace_in_paragraph(p,
            "(c) 10% payment upon issuance of PAC;",
            "(c) Combined PAC payment (EUR 346,284.94) = Component A PAC (10%, EUR 282,316.94) "
            "+ Component B PAC milestone (10%, EUR 63,967.99), due within seven (7) days of PAC.")
        # Schedule A payment summary lines (target ORIGINAL source amounts — run before amount updates)
        _replace_in_paragraph(p,
            "Advance (30%):     EUR 1,033,290.00  \u2014 Within 30 days of payment trigger (\u00a71A.3)",
            "COMPONENT A (EUR 2,823,169.43) \u2014 A1 Advance 25%: EUR 705,792 | A2 Pre-ship 50%: "
            "EUR 1,411,585 | A3 Delivery 15%: EUR 423,475 | A4 PAC 10%: EUR 282,317")
        _replace_in_paragraph(p,
            "Pre-Ship (55%):    EUR 1,894,365.00  \u2014 On FAT + Linyang written confirmation",
            "COMPONENT B (EUR 639,680) \u2014 M1a Mob+Eng 20%: EUR 127,936 | M1b Civil 20%: EUR 127,936 "
            "| M2\u201315%: EUR 95,952 | M3\u201315%: EUR 95,952 | M4\u201315%: EUR 95,952 "
            "| M5 PAC 10%: EUR 63,968 | Retention 5%: EUR 31,984")
        _replace_in_paragraph(p,
            "PAC (10%):         EUR 344,430.00  \u2014 System commissioned, grid-connected",
            "First combined payment at trigger (A1 + M1a): EUR 833,728.35")
        _replace_in_paragraph(p,
            "Retention (5%):    EUR 172,215.00  \u2014 Released after 3-month DLP",
            "TOTAL CONTRACT PRICE (ex VAT): EUR 3,462,849.40 \u2014 see Annex V6 \u00a77.1 for full schedule")
        # Target PAC: add 90-120 day timeframe note
        _replace_in_paragraph(p,
            "Contractor shall achieve PAC by Target PAC Date.",
            "Contractor shall achieve PAC by the Target PAC Date. The target timeframe from "
            "equipment delivery to Site to PAC is ninety (90) calendar days (worst case one hundred "
            "and twenty (120) calendar days). Exact dates are defined in the Delivery Schedule "
            "issued at Connection Terms.")
        # §1A.3 trigger clause: cross-reference the full companion-document CP bundle
        # (APG, LTSA, Technical Agreement sign-off, DWU, EMS) so the trigger clause itself
        # signposts that acceptance of all companion documents gates the advance payment.
        _replace_in_paragraph(
            p,
            "1A.3 The advance payment under Section 7.1(a) shall become due within thirty (30) days "
            "of the later of: (i) the Effective Date; and (ii) receipt of Connection Terms.",
            "1A.3 The advance payment under Section 7.1(a) shall become due within thirty (30) days "
            "of the later of: (i) the Effective Date; and (ii) receipt of Connection Terms; and in "
            "any event only once the Client has accepted or signed off the companion-document bundle "
            "(the OEM Direct Warranty Undertaking, the Advance Payment Guarantee (APG No. 1), the OEM "
            "5% performance guarantee, the Confirmed Price Certificate, the LTSA, the Technical "
            "Agreement(s) for Galascope 1 and Galascope 2, and the EMS Subscription Addendum) under "
            "Annex V6 Section 1A.5.")
        # §5A.3 (a) BESS operation: add LTSA+EMS reference
        _replace_in_paragraph(p,
            "(a) Operate BESS per OEM guidelines;",
            "(a) Operate BESS per OEM guidelines — for as long as the LTSA and EMS Subscription "
            "Addendum are in force, BESS operation and monitoring are managed under those agreements;")
        # §5A.3 (b) monitoring: add LTSA+EMS reference
        _replace_in_paragraph(p,
            "(b) Maintain monitoring systems;",
            "(b) Maintain monitoring systems — monitoring is provided under the LTSA and EMS "
            "Subscription Addendum for as long as those agreements are in force;")
        # §5A.3 (d) secure storage: add fencing note
        _replace_in_paragraph(p,
            "(d) Provide secure storage for equipment;",
            "(d) Provide secure storage for equipment (existing site fencing and alarm system "
            "confirmed as satisfactory for this purpose);")
        # §5A.3 (e) electrical connection: add "Obligation to" prefix
        _replace_in_paragraph(p,
            "(e) Provide electrical connection point and grid connection approval.",
            "(e) Obligation to provide electrical connection point and grid connection approval.")
        # Payments linked to milestones
        _replace_in_paragraph(p,
            "(e) Make all payments when due;",
            "(e) Make all payments when due, linked to the agreed milestone completion certificates "
            "in Annex V6 Section 7B;")
        # §4.5 excluded items: add milestone cross-reference
        _replace_in_paragraph(p,
            "4.5 Any works not expressly listed are excluded unless agreed in writing.",
            "4.5 Any works not expressly listed are excluded unless agreed in writing. The EPC "
            "works (Component B) scope and the corresponding milestone payment schedule are set "
            "out in Annex V6 Section 7B and the Delivery Schedule confirmed at Connection Terms.")
        _replace_in_paragraph(p, "Indicative Component B (EPC Services):     EUR 389,287.57",
                                   "Indicative Component B (EPC Services):     EUR 407,836.97")
        _replace_in_paragraph(p, "INDICATIVE CONTRACT PRICE (A+B, ex VAT):   EUR 2,238,000.00",
                                   "INDICATIVE CONTRACT PRICE (A+B, ex VAT):   EUR 2,256,549.40")
        _replace_in_paragraph(p, "TOTAL INDICATIVE CONTRACT PRICE (ex VAT):  EUR 3,444,300.00",
                                   "TOTAL INDICATIVE CONTRACT PRICE (ex VAT):  EUR 3,462,849.40")
        _replace_in_paragraph(p, "TOTAL:             EUR 3,444,300.00",
                                   "TOTAL:             EUR 3,462,849.40")
        _replace_in_paragraph(p, "Advance (30%):     EUR 1,033,290.00  \u2014 Within 30 days of payment trigger (\u00a71A.3)",
                                   "Advance (30%):     EUR 1,038,854.82  \u2014 Within 30 days of payment trigger (\u00a71A.3)")
        _replace_in_paragraph(p, "Pre-Ship (55%):    EUR 1,894,365.00  \u2014 On FAT + Linyang written confirmation",
                                   "Pre-Ship (55%):    EUR 1,904,567.17  \u2014 On FAT + Linyang written confirmation")
        _replace_in_paragraph(p, "PAC (10%):         EUR 344,430.00  \u2014 System commissioned, grid-connected",
                                   "PAC (10%):         EUR 346,284.94  \u2014 System commissioned, grid-connected")
        _replace_in_paragraph(p, "Retention (5%):    EUR 172,215.00  \u2014 Released after 3-month DLP",
                                   "Retention (5%):    EUR 173,142.47  \u2014 Released after 3-month DLP")
        # FINAL REVIEW: populate the two price fields (from Schedule A; APG = 25%/50% of Component A confirms)
        _replace_in_paragraph(
            p, "EUR [\u25cf] (\u201cEquipment Supply Price\u201d)",
            "EUR 2,823,169.43 (\u201cEquipment Supply Price\u201d)")
        _replace_in_paragraph(
            p, "EUR [\u25cf] (\u201cEPC Services Price\u201d)",
                 "EUR 639,679.97 (\u201cEPC Services Price\u201d)")
        # FINAL REVIEW: align EPC §10.6 SOH/cycle figures to the OEM Atlantic 5MWh datasheet + LTSA
        # Year-10 SOH kept at the group-presentation milestone (79.58%); no change.
        _replace_in_paragraph(
            p, "(c) Cycle Life: 7,000 cycles at 0.5C, 90% DoD, to 70% EOL.",
            "(c) Cycle Life: 8,000 cycles at 0.5C, 90% DoD, to 70% EOL.")
        # CONFIRMED via direct PDF read of Linyang's "Power Atlantic 5MWh Degradation Curve"
        # document (30 Jul 2026): Galascope 1 (5MW/20.06MWh) and Galascope 2 (2.5MW/10.03MWh) are
        # both ~4-hour systems = the 0.25P operating condition, not 0.5P. The verified 0.25P,
        # 1-cycle/day curve gives Year5=86.78%, Year10=79.58% (unchanged), Year15=73.61% — matching
        # Stelios's ask exactly and confirming it is OEM-backed, not a client concession beyond the
        # manufacturer's own data.
        _replace_in_paragraph(p, "Year 5 \u226585%", "Year 5 \u226586.78%")
        _replace_in_paragraph(p, "Year 15 \u226570%", "Year 15 \u226573.61%")
        # EMS fee clarity: the €400/MWh EMS subscription is included within the LTSA all-in fee
        _replace_in_paragraph(
            p, "3. EMS Subscription Addendum — DISPERON (EUR 400/MWh/yr from PAC)",
            "3. EMS Subscription Addendum — DISPERON (EUR 400/MWh/yr from PAC; while the LTSA is in "
            "force this is included within the LTSA all-in Service Fee of EUR 1,740/MWh/yr and not "
            "invoiced separately; the standalone rate applies only if the Client exits the LTSA)")
        # §11.4 — remove the stale "all three documents executed simultaneously" framing.
        # The EPC is signed 31 July 2026 on its own; the LTSA/DWU/APGs follow as CPs to the
        # advance payment per Annex V6 §1A.5 / §1A.5(e), not as same-day signatures.
        _replace_in_paragraph(
            p,
            "11.4 This Agreement is intended to be executed alongside the LTSA (LCY-LTSA-GAL-2026) "
            "and the OEM Direct Warranty Undertaking (LCY-OEM-DWU-001 v1.0). The Parties intend to "
            "execute all three documents simultaneously.",
            "11.4 This Agreement may be executed on its own. The LTSA (LCY-LTSA-GAL-2026) and the "
            "OEM Direct Warranty Undertaking (LCY-OEM-DWU-001 v1.0) are companion documents which "
            "must be executed and delivered, in each case, before the advance payment under Section "
            "7.1(a) falls due (Annex V6 §§1A.5 and 1A.5A); they are not required to be signed on "
            "the same date as this Agreement.")
        # Schedule A Part 6 heading — drop the "(executed simultaneously)" framing to match 11.4/1A.5(e)
        _replace_in_paragraph(
            p, "PART 6 — COMPANION DOCUMENTS (executed simultaneously)",
            "PART 6 — COMPANION DOCUMENTS (executed as conditions precedent to the advance payment)")
        # Schedule A Part 6 item 1 — drop stale "LTSA v4.0" version tag (current LTSA is v6.3)
        _replace_in_paragraph(
            p, "1. LTSA v4.0 (Ref: LCY-LTSA-GAL-2026)",
            "1. LTSA (Ref: LCY-LTSA-GAL-2026) — executed before the advance payment falls due")
        # Schedule C DWU annex note — drop stale "simultaneously with this Agreement" framing
        _replace_in_paragraph(
            p,
            "[Annexed: OEM-Direct-Warranty-Undertaking-Linyang.docx — to be executed by Linyang, "
            "Lighthief Cyprus Ltd, and Client simultaneously with this Agreement.]",
            "[Annexed: OEM-Direct-Warranty-Undertaking-Linyang.docx — to be executed by Linyang, "
            "Lighthief Cyprus Ltd, and Client before the advance payment under Section 7.1(a) falls "
            "due (Annex V6 §1A.5(a)).]")
        # Schedule A Part 3 — remove fixed calendar-date estimates that contradict the "not fixed at "
        # signing" principle (Annex V6 clause "8 / Schedule A Part 3"); tie everything to the Delivery
        # Schedule confirmed at Connection Terms instead.
        _replace_in_paragraph(
            p, "Production Start:       Q2 2026",
            "Production Start:       per the Delivery Schedule confirmed at Connection Terms (Annex V6)")
        _replace_in_paragraph(
            p, "Factory Acceptance:     Q3 2026",
            "Factory Acceptance:     per the Delivery Schedule confirmed at Connection Terms (Annex V6)")
        _replace_in_paragraph(
            p, "CIF Limassol:           August / September 2026",
            "CIF Limassol:           per the Delivery Schedule confirmed at Connection Terms (Annex V6)")

    # Re-anchor Confirmed-Price baseline from January 2026 to the May 2026 / Effective-Date basis.
    # The negotiated May price already absorbs the Jan->May index movement; the two-way adjustment
    # must run only from the Effective-Date reference values, not the stale January baseline.
    for p in doc.paragraphs:
        # Price basis worded to MATCH the signed LOI Clause 4.4 (no inconsistency).
        _replace_in_paragraph(
            p, "Linyang Quotation LY202601271 (January 2026)",
            "the January 2026 quotation basis (locked per the signed LOI Clause 4.4)")
        _replace_in_paragraph(
            p, "155,000 CNY/tonne (Mysteel China battery-grade spot, January 2026 monthly average)",
            "as recorded in the Indicative Price Basis Certificate at EPC signing (per LOI Clause 4.4)")
        _replace_in_paragraph(
            p, "8.18 CNY per EUR (January 2026 average)",
            "as recorded in the Indicative Price Basis Certificate at EPC signing (per LOI Clause 4.4)")
        if p.text.strip().startswith("Reference Date"):
            _replace_in_paragraph(
                p, "January 2026",
                "January 2026 quotation basis (price held per LOI Clause 4.4; adjustment measured "
                "from the indices in the Indicative Price Basis Certificate at EPC signing)")
        # Delivery timeline not fixed at signing -> set at Connection Terms (Annex V6)
        _replace_in_paragraph(
            p, "31 January 2027",
            "the date in the Delivery Schedule confirmed at Connection Terms (Annex V6)")
        # Strip indicative calendar-date estimates (all timing keyed to Connection Terms)
        _replace_in_paragraph(p, " (January 2032)", "")
        _replace_in_paragraph(p, " (approx. April 2027)", "")
        # Consolidate body to match Annex V6 (remove contradictions a reviewer would catch)
        _replace_in_paragraph(
            p,
            "(b) The APG covers: 20% advance on Component A + 50% pre-shipment on Component A = 70% of the Component A Equipment Supply Price.",
            "(b) Two Advance Payment Guarantees are provided, each issued by the OEM's bank naming the "
            "Client (Galascope Ltd) as beneficiary: APG No. 1 securing the equipment advance "
            "(EUR 705,792) and APG No. 2 securing the equipment pre-shipment tranche (EUR 1,411,585); "
            "combined cover EUR 2,117,377. APG No. 2 is issued at the Factory Acceptance Test, before "
            "the pre-shipment payment falls due.")
        _replace_in_paragraph(
            p, "(d) The APG shall remain valid until equipment delivery to Site plus thirty (30) days.",
            "(d) Each APG shall remain valid until the earlier of issuance of PAC or twelve (12) months "
            "after delivery of the equipment to Site; upon PAC the APGs are released.")
        _replace_in_paragraph(
            p,
            "(a) Warranty claims (defects in materials, workmanship, or OEM equipment): ten percent (10%) of the Contract Price;",
            "(a) Claims relating to the Contractor's own EPC services, installation and non-OEM works "
            "(excluding OEM equipment): ten percent (10%) of the Contract Price;")
        _replace_in_paragraph(p, "APG (70% \u00d7 Indicative Component A):", "APG (combined batch, per Annex V6):")
        _replace_in_paragraph(
            p, "EUR 1,294,098.70",
            "No. 1 advance EUR 705,792 + No. 2 pre-shipment EUR 1,411,585 (combined batch)")
        _replace_in_paragraph(p, "EUR 682,119.90", "see combined batch APG figure above (Annex V6)")
        _replace_in_paragraph(p, "Performance Bond (5% \u00d7 Component A):", "Performance Guarantee (5% \u00d7 Component A):")
        # §9.2 body — replace "no FAC / auto-release" with "see Annex V6 §9.2A for FAC"
        _replace_in_paragraph(
            p,
            "There is no Final Acceptance Certificate. Retention is released automatically at the "
            "end of the Defects Liability Period subject to Section 7.4(c).",
            "A Final Acceptance Certificate (FAC) shall be issued at the end of the Defects "
            "Liability Period in accordance with Annex V6 Section 9.2A. Retention is released "
            "within thirty (30) days of FAC issuance per Section 7.4.")
        # §8.4.5(a) — Dino: deduction should be cost-to-complete (already correct; add clarification)
        _replace_in_paragraph(
            p,
            "(a) The Contractor shall retain from payments received an amount equal to the value "
            "of works completed and materials delivered to Site;",
            "(a) The Contractor shall retain from payments received an amount equal to the value "
            "of works completed and materials delivered to Site, being the reasonable and documented "
            "cost incurred by the Contractor up to the date of termination;")
        # §14.4 Insurance — evidence provided BEFORE the first payment (Dino accept)
        _replace_in_paragraph(
            p,
            "The Contractor shall provide insurance certificates or cover notes within fourteen (14) days of the first payment under Section 7.1(a).",
            "The Contractor shall provide insurance certificates or cover notes (public liability, "
            "professional indemnity, and CAR/erection all-risks as required by Section 14.1) to the "
            "Client before the advance payment under Section 7.1(a) falls due. Public Liability and "
            "Employers\u2019 Liability insurance, being standing corporate policies of the Contractor, "
            "shall be in place from the Effective Date. The Contractor\u2019s Professional Indemnity "
            "insurance shall be bound and in force no later than the date the order is placed under "
            "Section 1A.4 (the same trigger as the advance payment under Section 7.1(a)), consistent "
            "with the CAR policy\u2019s own attachment point under Section 14.5(a); the Contractor "
            "shall undertake no design, engineering or site works giving rise to professional-"
            "negligence exposure before that date.")
        # §7.2 VAT — EPC services are domestic reverse-charge under Cyprus Article 11B; the Contractor
        # does not add VAT to invoices, so §7.2 must not say otherwise. (The former standalone "§7A"
        # VAT clause was folded directly into this §7.2 text rather than kept as its own clause, so
        # §7.2 must not cross-reference "§7A" as if it were a separate section elsewhere in the body.)
        _replace_in_paragraph(
            p, "7.2 Payments exclusive of VAT.",
            "7.2 The Contract Price and all milestone amounts are stated exclusive of VAT. The EPC "
            "services under this Agreement are subject to Cyprus VAT on a domestic reverse-charge basis "
            "under Article 11B of the Value Added Tax Law (L. 95(I)/2000); the Contractor shall not add "
            "or collect VAT on its invoices, and the Client is responsible for self-accounting the VAT "
            "directly with the Cyprus Tax Department. This is without prejudice to import VAT on the "
            "BESS equipment, which is addressed separately in Section 7C.")
        _replace_in_paragraph(
            p, "7.2 The Contract Price and all milestone amounts are stated exclusive of VAT. VAT at the "
            "applicable Cyprus rate (currently nineteen percent (19%)) shall be added to each invoice "
            "and paid by the Client, and is recoverable by the Client in accordance with applicable law.",
            "7.2 The Contract Price and all milestone amounts are stated exclusive of VAT. The EPC "
            "services under this Agreement are subject to Cyprus VAT on a domestic reverse-charge basis "
            "under Article 11B of the Value Added Tax Law (L. 95(I)/2000); the Contractor shall not add "
            "or collect VAT on its invoices, and the Client is responsible for self-accounting the VAT "
            "directly with the Cyprus Tax Department. This is without prejudice to import VAT on the "
            "BESS equipment, which is addressed separately in Section 7C.")
        # §10.5 low-SOC/low-voltage warranty-void exceptions — remove the "Tier A or above" gating on
        # (iv) and the following alert commitment. Stelios's client-proposed insertion said responsibility
        # sits with the Contractor "for as long as the O&M of the BESS is being carried out under a
        # confirmed LTSA contract with the Contractor" — i.e. for ANY LTSA Tier, not just Tier A+. The
        # Galascope default is Tier C (Schedule 2), so Tier-A+-gating this protection would leave the
        # Client's own selected Tier unprotected — contradicting both Stelios's ask and Timotheos's
        # parallel (untiered) "Contractor monitoring failures ... shall not void" ask, and the untiered
        # Annex V6 §10.5A carve-out itself.
        _replace_in_paragraph(
            p, "(ii) an EMS malfunction not attributable to the Client;",
            "(ii) an EMS malfunction not attributable to the Client, including a malfunction of DISPERON "
            "(the EMS provider) or of any other Lighthief-group entity (see Section 10.5A);")
        _replace_in_paragraph(
            p,
            "(iv) failure by the Contractor to provide a monitoring alert under an LTSA with 24/7 "
            "monitoring (Tier A or above), where such alert would have enabled the Client to prevent "
            "the condition.",
            "(iv) any failure by the Contractor to provide monitoring, alerting or a timely warning that "
            "would have enabled the Client to prevent the condition.\n"
            "For as long as the operation and maintenance of the BESS is carried out under a confirmed "
            "LTSA with the Contractor, responsibility for a condition falling within (i)-(iv) above lies "
            "with the Contractor, not the Client, and the applicable remedies are provided under the "
            "LTSA (and, where the condition involves the EMS, the EMS Subscription Addendum), consistent "
            "with Section 10.5A and Section 11.5 (single integrated Project-Document suite). This "
            "protection applies regardless of which LTSA Tier is selected under LTSA Schedule 2.")
        _replace_in_paragraph(
            p,
            "Where the Client has entered into an LTSA with 24/7 monitoring (Tier A+), the Contractor "
            "shall alert the Client within 24 hours of SOC dropping below 10%.",
            "The Contractor shall alert the Client within twenty-four (24) hours of SOC dropping below "
            "ten percent (10%) for as long as the LTSA is in effect, regardless of Tier; the specific "
            "monitoring and alerting capability provided is as set out in the LTSA Tier selected under "
            "Schedule 2, but this 24-hour alert obligation itself is not conditional on Tier selection.")
        # §7.5(b) title — "whichever comes first" so title never lags the APG (Anastasios point)
        _replace_in_paragraph(
            p,
            "(b) Title (ownership) passes to the Client upon issuance of PAC and receipt of the PAC payment under Section 7.1(c). The Retention does not defer title.",
            "(b) Title (ownership) passes to the Client on the earlier of: (i) issuance of PAC and "
            "receipt of the PAC payment under Section 7.1(c); or (ii) twelve (12) months after delivery "
            "of the equipment to Site (the latest expiry of the Advance Payment Guarantees), so that "
            "title passes no later than the end of APG cover. Where title passes under (ii) before "
            "PAC, the Contractor retains a security interest in the Equipment for any unpaid balance "
            "until the PAC payment is made. The Retention does not defer title.")

    # Inline-integrate every amendment clause into its numbered body position
    # (no separate annex / side schedule — single document, colour-coded).
    _integrate_epc_annex(doc)

    # Top banner marking v6.3
    first = doc.paragraphs[0]
    banner = first.insert_paragraph_before()
    _run(banner, "EPC AGREEMENT \u2014 VERSION 6.3 (July 2026)", bold=True, size=13, color=NAVY)
    b2 = first.insert_paragraph_before()
    _run(b2, "All agreed amendments through v6.3 are integrated directly into the numbered "
             "clauses below at their correct place in the Agreement \u2014 there is no separate "
             "annex or side schedule. Clauses are colour-coded by source (see the key after "
             "Section 1): red = included as the Client requested; purple = included with our "
             "changes; orange = included with our counter-proposal; teal = raised by Dino "
             "Constantinou; blue = raised by Anastasios (lawyer); navy = base v6 drafting.",
         italic=True, size=9.5, color=GREY)
    legend = first.insert_paragraph_before()
    _run(legend, "Colour key:  ", bold=True, size=8.5, color=GREY)
    _run(legend, "\u25A0 included as requested   ", size=8.5, color=AMD_RED)
    _run(legend, "\u25A0 included \u2014 our changes   ", size=8.5, color=AMD_PURPLE)
    _run(legend, "\u25A0 included \u2014 our counter   ", size=8.5, color=AMD_ORANGE)
    _run(legend, "\u25A0 from Dino   ", size=8.5, color=AMD_TEAL)
    _run(legend, "\u25A0 from Anastasios   ", size=8.5, color=AMD_BLUE)
    _run(legend, "\u25A0 base v6 drafting", size=8.5, color=NAVY)
    first.insert_paragraph_before()

    PKG.mkdir(parents=True, exist_ok=True)
    doc.save(str(EPC_OUT))
    print(f"EPC v6 -> {EPC_OUT}")


# ── clean client cover note ──────────────────────────────────────────────────
def build_cover_note():
    doc = Document()
    s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(10.5)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(t, "Galascope BESS \u2014 EPC Package", bold=True, size=15, color=NAVY)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(st, "Galascope 1 (5 MW / 20 MWh) & Galascope 2 (2.5 MW / 10 MWh) \u2014 Famagusta, Cyprus  \u00b7  Ref. LCY-EPC-GAL-B1-2026 (v6.3, July 2026)",
         italic=True, size=10, color=GOLD)
    doc.add_paragraph()

    para(doc, "To: Galascope Ltd \u2014 Ntinos Konstantinos (Director), and advisers / board.")
    para(doc, "From: Lighthief Cyprus Ltd \u2014 Alexander Papacosta (Director).")

    h(doc, "How this package works")
    for tx in [
        "Only the EPC Agreement is signed now. Signing it records the agreed terms and triggers no payment.",
        "The order begins at Connection Terms (grid terms from EAC/DSO). At that point Lighthief signs the "
        "supply contract with the manufacturer and the manufacturer's bank issues the Advance Payment "
        "Guarantee against both agreements.",
        "The companion documents below travel with this pack for review now; they are executed at the order "
        "trigger, not today.",
        "No advance is payable until the signed Direct Warranty Undertaking, the bank APG (with Galascope as "
        "beneficiary), the 5% performance guarantee and the Confirmed Price Certificate are presented. If any "
        "is not provided in the agreed form, either party may withdraw with no liability.",
    ]:
        b = doc.add_paragraph(style="List Bullet"); _run(b, tx)

    h(doc, "Documents in this package")
    rows = [
        ("01", "EPC Agreement v6.3 (with Annex V6 Amendment Schedule)", "For signature now"),
        ("02", "Ownership & Guarantee Flow (one-page explainer)", "Explainer"),
        ("03", "LTSA (Long-Term Service Agreement) v6.3", "For review"),
        ("04", "OEM Direct Warranty Undertaking \u2014 G1 & G2", "Draft for review; signed by OEM at order"),
        ("05a", "Advance Payment Guarantee No. 1 \u2014 equipment advance (EUR 705,792 = 25% of equipment supply, Galascope beneficiary)", "Form for review; bank-issued before advance"),
        ("05b", "Advance Payment Guarantee No. 2 \u2014 equipment pre-shipment (EUR 1,411,585 = 50% of equipment supply)", "Form for review; bank-issued at FAT"),
        ("05c", "APG No. 1 Bank Review Package (Stelios / Alpha Bank routing note + redline + clean form)", "For Stelios to route to Alpha Bank before sending to Linyang"),
        ("06", "OEM 5% Performance Guarantee", "Form for review; issued at order"),
        ("07", "EMS Subscription Addendum (DISPERON)", "For review"),
        ("08", "Redline Response Matrix v6.3 (client comments vs. our response, colour-coded)", "For review"),
        ("Sch. A", "Technical Specifications \u2014 Technical Agreements G1/G2, OEM/Kehua/Schneider datasheets, SOH degradation curve", "Reference"),
        ("Sch. B", "Certificates & Compliance \u2014 grid-code, CE, IEC and transport-safety certificates", "Reference"),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for c, txt in zip(hdr, ["#", "Document", "Status"]):
        _run(c.paragraphs[0], txt, bold=True, size=10)
    for num, name, status in rows:
        cells = tbl.add_row().cells
        _run(cells[0].paragraphs[0], num); _run(cells[1].paragraphs[0], name); _run(cells[2].paragraphs[0], status)

    h(doc, "Key protections built in for the Client")
    for tx in [
        "The manufacturer's equipment supply is secured by two first-demand bank guarantees totalling "
        "EUR 2,117,377: APG No. 1 (EUR 705,792, 25% of the equipment supply value) issued before the "
        "advance, and APG No. 2 (EUR 1,411,585, 50% of the equipment supply value) issued at the Factory "
        "Acceptance Test before the pre-shipment payment. These secure the equipment refund from the "
        "manufacturer. They are deliberately not full cover for each milestone: your milestone payments "
        "(30% / 55% of the total contract price) also include EPC services, and that portion \u2014 together "
        "with any balance above the APG figures \u2014 is protected by the other layers below (FAT before "
        "payment, the 5% performance guarantee, the 5% retention, your possession and title of the "
        "equipment, and the 5-year DWU), not by the APGs.",
        "Both APGs name Galascope directly as beneficiary and are kept valid to PAC (or 12 months after "
        "delivery), so faults found at commissioning remain covered.",
        "The manufacturer's bank issues the APG(s) once it holds both this EPC and the supply contract, "
        "naming Galascope as beneficiary from the outset \u2014 so you can claim directly under the APG even "
        "though you pay Lighthief (not the manufacturer). It is original issuance in your name, not a "
        "transfer, which confirms the guarantee is valid for you and supports the payment terms.",
        "Title passes to you on the earlier of PAC or 12 months after delivery \u2014 so ownership never "
        "lags the end of the APG cover (with Lighthief retaining a security interest for any unpaid "
        "balance if title passes before PAC). The APGs expire at the same moment as the latest "
        "title-transfer branch (PAC or 12 months); once the APGs release, protection continues under "
        "the 5-year DWU warranty and the performance guarantee (valid to DLP end).",
        "A Factory Acceptance Test (FAT) lets you (or your inspector) witness factory testing and have "
        "defects fixed before the pre-shipment payment and before shipment, and the second guarantee is "
        "issued at that point \u2014 the primary protection against a large proportion of equipment being "
        "faulty.",
        "Your right to reject and require remedy runs to commissioning / PAC, not just 30 days after delivery.",
        "Price moves both ways \u2014 the Indicative Price is the current agreed price; the confirmed price "
        "tracks the lithium index and EUR/CNY measured only from the Effective Date (signing) forward: "
        "downward moves are passed to you; upward moves are capped at 5%; either side may walk away above "
        "the cap; orders more than six months out are re-quoted on the same two-way basis.",
        "Delivery timeline is set at the order trigger, not fixed today. Production lead time is "
        "indicative only and can change from any manufacturing estimate. When Connection Terms are "
        "received, the confirmed manufacturing lead time, the shipping timeline and the Target PAC date "
        "are set in the Delivery Schedule, and the Confirmed Price is fixed at the same time (it can move "
        "only within the two-way adjustment and the 5% upward cap). All delay liquidated damages are "
        "measured from that confirmed schedule. If the confirmed timeline is too long for you, you may "
        "withdraw with no liability \u2014 this protects you against unpredictable manufacturer lead times.",
        "Continuous cover: APGs to PAC (advance + pre-shipment) \u2192 CAR/erection insurance at full "
        "replacement value from port of discharge (the sea voyage is covered by the manufacturer's CIF "
        "marine) \u2192 5% performance guarantee \u2192 5% retention \u2192 5-year OEM warranty & Direct Warranty "
        "Undertaking.",
        "The OEM's standard warranty caps aggregate liability at 10% of the payment received for "
        "defective products. The EPC and DWU draft include a requested manufacturing-defect carve-out "
        "under which the OEM bears full repair/replacement cost for confirmed manufacturing defects "
        "(not subject to the 10% cap). The primary protection against a large proportion of equipment "
        "being faulty is FAT: defects are identified and remedied before the pre-shipment payment is "
        "made and before shipment, preventing payment for defective goods. The APGs, extended "
        "rejection window to PAC, and 5-year DWU provide continuous layered cover.",
    ]:
        b = doc.add_paragraph(style="List Bullet"); _run(b, tx)

    h(doc, "Note \u2014 Kehua PCS and OEM scope")
    para(doc,
         "The PCS is Kehua C-series (BCS1250K-C-HUD), supplied within Linyang's commercial scope as part of "
         "Linyang's integrated BESS offering. The Kehua PCS is an OEM-supplied product and is covered under "
         "the Linyang product warranty and Direct Warranty Undertaking.")

    h(doc, "Note \u2014 grid-forming (VSG) & black start")
    para(doc,
         "The PCS is hardware-capable of grid-forming (VSG) and black start; these are firmware-enabled. "
         "The current Technical Agreement lists them as excluded and they will be added by amendment after "
         "review. The hardware is identical, so we expect no change to the equipment price; whether "
         "activation is at the factory or by later update depends on the DSO grid-code requirement, "
         "which we are confirming. This is flagged for transparency and does not affect signing the EPC now.")

    h(doc, "Note \u2014 governing law")
    para(doc,
         "This EPC, the LTSA, EMS Addendum, and the Advance Payment Guarantee instruments are governed by "
         "Cyprus law. The OEM Direct Warranty Undertaking follows the manufacturer's standard warranty terms "
         "under PRC law, with disputes resolved at the Shanghai International Arbitration Center (SHIAC) in "
         "English \u2014 consistent with international BESS procurement practice. The Cyprus-law EPC and APGs "
         "provide the primary enforceable remedies. Lenders requiring a legal opinion on SHIAC enforceability "
         "are directed to the OEM DWU section 6.3.")

    doc.add_paragraph()
    f = doc.add_paragraph(); f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(f, "Lighthief Cyprus Ltd \u00b7 HE 477423 \u00b7 28 October Ave 249, Lophitis Business Center 1, Office 201, "
            "3035 Limassol, Cyprus \u00b7 office@lighthief.com \u00b7 +357 99 164 158 \u00b7 solarfarms.cy",
         size=8, color=GREY)

    out = PKG / "00-Cover-Note-and-Index.docx"
    doc.save(str(out))
    print(f"Cover note -> {out}")


def _rewrite_paragraph(p, full_new):
    """Replace the entire text of a paragraph (keeps first run's formatting)."""
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = full_new
    else:
        p.add_run(full_new)


def _fix_docx_text(src: Path, dest: Path, replacements=None, rewrites=None):
    """Load src docx, apply substring replacements and/or full-paragraph rewrites, save to dest.

    replacements: list of (old_substring, new_substring) applied within runs.
    rewrites: list of (marker_substring, full_new_paragraph_text); if marker is in a
    paragraph's text, the whole paragraph is rewritten (used to strip group/portfolio leaks).
    """
    replacements = replacements or []
    rewrites = rewrites or []
    doc = Document(str(src))

    def handle(p):
        for marker, full_new in rewrites:
            if marker in p.text:
                _rewrite_paragraph(p, full_new)
                return
        for old, new in replacements:
            _replace_in_paragraph(p, old, new)

    for p in doc.paragraphs:
        handle(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    handle(p)
    doc.save(str(dest))


def _append_annex(dest: Path, title: str, intro: str, clauses):
    """Append a titled amendment annex (list of (headline, body)) to an existing docx."""
    doc = Document(str(dest))
    doc.add_page_break()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(title)
    tr.bold = True
    tr.font.name = "Calibri"
    tr.font.size = Pt(13)
    tr.font.color.rgb = NAVY
    ip = doc.add_paragraph()
    ir = ip.add_run(intro)
    ir.italic = True
    ir.font.name = "Calibri"
    ir.font.size = Pt(9.5)
    ir.font.color.rgb = GREY
    _annex_colour_legend(doc)
    for headline, body in clauses:
        col = _amd_color(headline, AMD_PURPLE)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        hr = p.add_run(headline + "  ")
        hr.bold = True
        hr.font.name = "Calibri"
        hr.font.size = Pt(10)
        hr.font.color.rgb = col
        br = p.add_run(body)
        br.font.name = "Calibri"
        br.font.size = Pt(10)
        br.font.color.rgb = col
    doc.save(str(dest))


# LTSA v6.2 — Timotheos / technical review clauses now incorporated (accepts + our counters).
LTSA_ANNEX = [
    ("2.4  Integrated Project-Document suite.",
     "This Agreement is part of a single integrated Project-Document suite comprising the EPC "
     "Agreement, the Technical Agreements, the OEM Direct Warranty Undertaking, the OEM Warranty "
     "Terms, the Advance Payment Guarantees, the OEM Performance Guarantee and the EMS Subscription "
     "Agreement, and shall be interpreted consistently with them so that there is no gap in "
     "responsibility for operation, maintenance, warranty preservation or performance."),
    ("Definitions \u2014 Restoration and Resolution.",
     "\u201cRestoration Time\u201d means the time from detection of a fault to restoration of safe "
     "operational capability (including safe partial operation where full operation is not "
     "immediately possible). \u201cResolution Time\u201d means the time from detection to permanent "
     "repair/replacement, recommissioning and incident closure. \u201cRepeated Service Failure\u201d means "
     "three or more Critical breaches of the same or similar cause within a rolling twelve (12) months."),
    ("4.1A  Selected Tier.",
     "Unless the Client expressly selects otherwise in Schedule 2, the Project shall be treated as "
     "Tier C from PAC for both Galascope 1 and Galascope 2 (corrective maintenance, local spare "
     "parts, 97% availability guarantee, priority response, warranty-preservation services and "
     "annual performance testing)."),
    ("4.3A  Corrective Maintenance \u2014 scope.",
     "Corrective Maintenance includes all diagnosis, labour, travel within Cyprus, removal, "
     "installation, recommissioning, configuration restoration, testing and incident documentation "
     "required to return the affected BESS, PCS, BMS, MV switchgear, SCADA gateway and related "
     "controls to service, except to the extent a cost is expressly excluded under Schedule 2."),
    ("4.4A  Exclusions \u2014 causation and burden.",
     "Exclusions apply only to the extent the excluded event directly caused the relevant loss, "
     "downtime or degradation. The Service Provider continues to provide monitoring, diagnosis, "
     "emergency response, safe isolation, mitigation, documentation and warranty-claim support "
     "notwithstanding any exclusion."),
    ("5.2A  Minimum monitored parameters.",
     "Monitoring shall include, at minimum, cell-level voltage and temperature, string/container "
     "SOC and SOH, current, power limits, alarm/fault codes, PCS status, transformer/MV status, "
     "HVAC/liquid-cooling status, fire-system and gas-detection status, emergency stop, auxiliary "
     "supply, communication status, EMS command execution and DSO/SCADA telemetry."),
    ("5.6A  Reporting content.",
     "Reports shall include raw-data extracts, incident log, downtime classification with "
     "excluded-hour justification, availability per Park and in aggregate, SOH trend, cell "
     "imbalance, thermal deviation, firmware/software version, open punch items, spare-parts usage, "
     "warranty notices, cyber/security events and recommendations. The Client may reasonably "
     "challenge a classification."),
    ("5.8A  Client data ownership.",
     "The Client owns all operational, technical, performance, market, dispatch, SCADA, BMS, PCS, "
     "SOH and availability data generated by or relating to the BESS. The Service Provider shall "
     "not use identifiable Client data for benchmarking, training, product development or "
     "third-party purposes without the Client's prior written consent."),
    ("6.1A  Warranty preservation.",
     "The Service Provider shall perform all maintenance required to preserve the OEM warranty, DWU, "
     "SOH guarantees, corrosion warranty, fire-system warranty, PCS/MV warranty and insurance cover, "
     "and shall be liable for loss of warranty, claim rejection, downtime or degradation caused by "
     "its failure to perform or document required maintenance."),
    ("6.3A  Warranty-Preservation File.",
     "The Service Provider shall maintain a warranty-preservation file per Park, including "
     "maintenance checklists, photographs, torque records, thermal images, grounding/insulation "
     "results, firmware records, alarm/SOC logs, environmental records, fire-system tests, "
     "spare-part records and OEM communications, in a format acceptable to the OEM (Linyang)."),
    ("7.4A  No additional charge for warranty-preserving services.",
     "No additional charge shall apply for services reasonably required to comply with the OEM "
     "manuals, preserve warranties, perform corrective maintenance within the selected Tier, "
     "prepare warranty claims, comply with agreed regulatory testing, deliver standard data "
     "exports, or remedy a Service Provider breach."),
    ("8.1A  Redundant monitoring.",
     "The Service Provider shall maintain redundant monitoring, alerting and communication channels, "
     "with automatic alerts to the Client for low SOC, low cell voltage, high temperature, "
     "HVAC/liquid-cooling fault, fire/gas alarm, PCS trip, communication loss and any condition that "
     "could void warranty. Failure of the monitoring platform shall itself be treated as a Major "
     "Alert (or Critical Alert where safety-related)."),
    ("8.2A  Maximum Restoration and Resolution Times.",
     "Critical Alerts: safe remote isolation immediately where required; restoration plan within "
     "four (4) hours; safe partial or full restoration within twenty-four (24) hours where "
     "technically possible; permanent resolution within five (5) Business Days unless OEM parts are "
     "required. Major Alerts: restoration plan within one (1) Business Day; restoration within five "
     "(5) Business Days; permanent resolution within fifteen (15) Business Days. Minor Alerts: "
     "resolution at the next maintenance visit or within thirty (30) days, whichever is earlier. Any "
     "deviation requires a written remediation plan accepted by the Client, such acceptance not to "
     "be unreasonably withheld or delayed. These times are service targets and are subject to OEM "
     "part lead-times, Force Majeure and Client/site-access delays (which stop the clock); any "
     "service credit under Section 8.5 for a breach of these targets is not additional to the "
     "Availability Liquidated Damages for the same downtime."),
    ("8.6A  Major Incident Procedure.",
     "For any fire, thermal event, repeated trip, forced outage exceeding twenty-four (24) hours, "
     "cyber incident, grid-code breach, warranty-threatening event, or unavailable capacity "
     "exceeding 10% of a Park, the Service Provider shall notify the Client immediately, issue an "
     "initial incident report within twenty-four (24) hours and a root-cause analysis thereafter, "
     "and keep the Client updated until closure."),
    ("8.7A  Safety-critical 24/7 response.",
     "Critical Alerts, safety incidents, fire-system alarms, warranty-voiding SOC/voltage alerts, "
     "cyber incidents and complete Park outages require 24/7/365 operational response. On-site "
     "attendance for safety-critical events shall not be limited to Business Days where local "
     "personnel or emergency-contractor attendance is reasonably available."),
    ("9.8  Chronic underperformance.",
     "If Availability for any Park is below 95% in any contract year, or below 97% for two "
     "consecutive years, or there are three or more Critical Alert SLA breaches in a rolling twelve "
     "(12) months, the Client may require a remediation plan, an independent technical audit at the "
     "Service Provider's cost and additional spare stock."),
    ("10.3A  Witnessed SOH/RTE testing.",
     "SOH, usable AC capacity and RTE testing shall be performed under a pre-agreed protocol "
     "consistent with the EPC, Technical Agreements and OEM Warranty Terms. The Client and lender's "
     "Independent Engineer may witness testing. If an additional test requested by the Client "
     "confirms underperformance or a warranty issue, the Service Provider shall bear the test cost."),
    ("10.4A  SOH remedy \u2014 scope.",
     "Remedies for SOH underperformance shall restore the Client to the guaranteed position and "
     "shall include parts, labour, installation, recommissioning, testing, documentation and data "
     "updates. Shipping, customs and craneage attributable to an OEM warranty return shall be "
     "OEM-funded; the Service Provider shall manage and evidence the OEM claim. This does not apply "
     "to the extent a Client-caused exclusion directly caused the underperformance."),
    ("10.6A  OEM reserve evidence.",
     "The Service Provider shall provide annual evidence that the OEM warranty reserve, extended "
     "warranty, spare-part access and OEM support arrangements remain valid for the Project, and "
     "shall notify the Client within five (5) Business Days of any OEM dispute, insolvency risk, "
     "supply restriction, product recall, safety notice or firmware issue affecting the Project."),
    ("11.3A  Local spare-parts warehouse.",
     "A local spare-parts warehouse shall be operational no later than PAC, holding an agreed minimum "
     "stock of critical sub-components (control boards, sensors, fans, fuses, contactors) with part "
     "numbers, compatibility references, storage conditions and replenishment deadlines. Failure to "
     "maintain the agreed stock is a Service Provider default and related downtime counts as "
     "Unavailable. Stocking of full PCS units and transformers is subject to further agreement; in "
     "the interim these are supported by a four (4) week OEM supply commitment."),
    ("11.6  Obsolescence and equivalent replacement.",
     "The Service Provider shall use commercially reasonable efforts to support the availability of "
     "compatible spare parts, firmware and replacement components for at least fifteen (15) years. "
     "If an original component is obsolete, it shall provide an equivalent or better compatible "
     "replacement without loss of warranty, certification or grid-code compliance."),
    ("12.5A / 12.7A  Client obligations and suspension \u2014 causation.",
     "The Client's operational obligations shall not be treated as breached where the relevant "
     "condition was caused by EMS/BMS/PCS malfunction, Service Provider monitoring failure, Force "
     "Majeure, or the Service Provider's failure to provide timely alerts or operating instructions. "
     "Suspension of the Availability, SOH or warranty-related guarantees applies only to the extent "
     "a Client breach directly caused the non-performance, and only after written notice, a "
     "reasonable cure opportunity and Service Provider mitigation. Non-payment of genuinely disputed "
     "amounts shall not suspend safety monitoring, warranty-preservation alerts or data access."),
    ("13.6  Transition assistance.",
     "On expiry or termination the Service Provider shall provide up to ninety (90) days of "
     "transition assistance at agreed rates, including data export, passwords and access "
     "credentials, configuration files, register maps, firmware records, maintenance records, "
     "open-incident files, spare-parts list and warranty-claim files."),
    ("14.3A / 14.4A  Additional liability carve-outs.",
     "The liability limitations shall not apply to: cyber-security breach or unauthorised remote "
     "access caused by Service Provider systems; loss, corruption or withholding of Client data; "
     "regulatory or grid-code non-compliance within Service Provider scope; or safety incidents "
     "caused or worsened by Service Provider breach. Third-party-system exclusions shall not apply "
     "where the Service Provider selected, configured, integrated, maintained, monitored or "
     "recommended the relevant system, or where the fault arises at an interface for which the "
     "Service Provider or its affiliates are responsible."),
    ("14.5A  OEM claim management.",
     "For claims relating to OEM equipment defects the Service Provider shall actively manage, "
     "prosecute and evidence the claim with the OEM, and remains liable for loss caused by its "
     "failure to preserve warranty rights, perform maintenance, maintain records/spares, mitigate "
     "downtime, implement temporary solutions or comply with its service obligations."),
    ("16.4  Force Majeure \u2014 limitations.",
     "Force Majeure shall not excuse obligations to maintain monitoring where technically possible, "
     "provide data, issue alerts, preserve warranty rights, mitigate downtime, maintain insurance, "
     "protect cyber-security, deliver records, or perform obligations capable of being performed "
     "remotely."),
    ("18A  Cyber security (framework).",
     "The Service Provider shall implement and maintain cyber-security controls for the industrial "
     "control systems consistent with recognised good practice (IEC 62443 principles), including "
     "multi-factor authentication, role-based access control, secure remote access, encryption in "
     "transit and at rest, audit logs, least-privilege accounts, vulnerability management, "
     "firmware-integrity controls and annual access review. The specific control set and any formal "
     "IEC 62443 / NIS2 certification shall be confirmed with the ICS/EMS provider (DISPERON)."),
    ("19.5A  SCADA/DSO communication failures.",
     "SCADA/DSO communication failures attributable to the Service Provider, EMS/SCADA "
     "configuration, protocol mapping, firmware changes, communication gateway, maintenance activity "
     "or cyber-security controls shall count as Unavailable Hours where they prevent dispatch, "
     "monitoring, regulatory reporting or DSO-approved operation."),
    ("19.8A  Editable documentation.",
     "Documentation shall be provided in editable electronic format where reasonably available, "
     "including native settings/configuration files, relay settings, SCADA point lists, "
     "Modbus/IEC 61850/IEC 60870-5-104 maps, firmware versions, access-control handover records, "
     "maintenance photographs and independent test certificates."),
    ("20.5  Expert determination.",
     "Technical disputes concerning Availability, SOH, RTE, capacity, root cause, warranty "
     "attribution, excluded hours or service failure may be referred by either Party to an "
     "independent technical expert agreed by the Parties or appointed by the President of ETEK. The "
     "expert's determination shall be binding absent manifest error."),
    ("21.3B  Lender step-in.",
     "Any project-finance lender or security agent may step in, receive notices, cure Client "
     "defaults, assume payment obligations, enforce this Agreement, direct payment of proceeds and "
     "require continued performance for a reasonable cure period following an enforcement event. The "
     "Service Provider shall, acting reasonably, enter into a direct agreement with the "
     "lender/security agent."),
    ("Assignment \u2014 Client financing.",
     "The Client may assign or charge its rights under this Agreement and any proceeds to any "
     "project-finance lender, security agent, refinancing lender, purchaser of the Project or "
     "successor owner without the Service Provider's consent."),
]


def copy_attachments():
    # Simple file copies
    mapping = [
        (V6 / "Galascope-ownership-guarantee-flow-jun2026.html", "02-Ownership-and-Guarantee-Flow.html"),
        (V6 / "Performance-Guarantee-Galascope-CLIENT-REVIEW-jun2026.docx", "06-OEM-5pct-Performance-Guarantee.docx"),
        (CONTRACTS / "Cyprus 5MW_20MWh Technical Agreement V2.1 (draft) 260429.docx", "08-Technical-Agreement-Galascope-1-5MW-20MWh.docx"),
        (CONTRACTS / "Cyprus 2.5MW_10MWh Technical Agreement V2.1 (draft) 260429.docx", "09-Technical-Agreement-Galascope-2-2.5MW-10MWh.docx"),
        (REPO / "docs" / "hardware" / "Linyang Warranty Terms v2.pdf", "10-Linyang-Product-Warranty-Terms-v2.pdf"),
    ]
    for src, dest in mapping:
        if src.exists():
            shutil.copy2(str(src), str(PKG / dest))
            print(f"  copied -> {dest}")
        else:
            print(f"  !! MISSING: {src}")

    # OEM Direct Warranty Undertaking (DWU) — this is our own draft for Linyang to sign, so we
    # correct it in step with the confirmed 0.25P, 1-cycle/day curve (Y5 86.78% / Y10 79.58% /
    # Y15 73.61%) now used in the EPC and LTSA, rather than leaving it at the old rounded floor
    # (85% / 79.58% / 70%) and only flagging the gap to Linyang.
    dwu_src = V6 / "OEM-DWU-Galascope-CLIENT-REVIEW-jun2026.docx"
    dwu_dest = PKG / "04-OEM-Direct-Warranty-Undertaking-G1-G2.docx"
    if dwu_src.exists():
        _fix_docx_text(dwu_src, dwu_dest,
            replacements=[
                ("End of Year 5 \u2265 85%; End of Year 10 \u2265 79.58%; End of Year 15 \u2265 70% "
                 "(1 cycle/day).",
                 "End of Year 5 \u2265 86.78%; End of Year 10 \u2265 79.58%; End of Year 15 \u2265 73.61% "
                 "(0.25P, 1 cycle/day \u2014 per Linyang's Power Atlantic 5MWh Degradation Curve "
                 "reference document, confirmed against this Project's 4-hour system duration)."),
                ("the guaranteed SOH point is Year 5 (\u2265 85%).",
                 "the guaranteed SOH point is Year 5 (\u2265 86.78%)."),
                ("Cycle life: 7,000 equivalent full cycles to 70% end-of-life at 0.5C, 90% DoD, "
                 "25\u00b0C as confirmed by the manufacturer.",
                 "Cycle life: 8,000 equivalent full cycles to 70% end-of-life at 0.5C, 90% DoD, "
                 "25\u00b0C as confirmed by the manufacturer."),
            ])
        print(f"  fixed + copied -> 04-OEM-Direct-Warranty-Undertaking-G1-G2.docx (SOH curve 0.25P)")
    else:
        print(f"  !! MISSING: {dwu_src}")

    # LTSA — copy with group-order/portfolio pricing leak stripped + Dino v6.1 ACCEPT fixes
    ltsa_src = CONTRACTS / "LTSA-Galascope-Esperia-may2026.docx"
    ltsa_dest = PKG / "03-LTSA-Galascope-G1-G2.docx"
    if ltsa_src.exists():
        _fix_docx_text(ltsa_src, ltsa_dest,
            replacements=[
                # ACCEPT: EPC ref v4.0 → v6.3 (Anastasios checklist) — kept in step with the EPC's own
                # version bumps (was frozen at "v6.0" through the v6.1/6.2/6.3 rounds; caught on final sweep).
                ("LCY-EPC-GAL-B1-2026 v4.0", "LCY-EPC-GAL-B1-2026 (v6.3, July 2026)"),
                ("EPC Agreement (LCY-EPC-GAL-B1-2026 v4.0)", "EPC Agreement (LCY-EPC-GAL-B1-2026 v6.3)"),
                # FINAL REVIEW: the standalone "COMPANION DOCUMENTS" mini-list (2 items, "all three
                # documents simultaneously") is stale and now duplicates/contradicts the fuller,
                # accurate suite list in Annex LT-2 §2.4 and the EPC's own §11.4/§1A.5(e) framing
                # (LTSA is a CP to the advance, not a same-day signature requirement).
                ("This LTSA is intended to be executed alongside and read in conjunction with:",
                 "This LTSA is a condition precedent to the EPC advance payment (EPC "
                 "§1A.5A) and is intended to be executed alongside and read in conjunction with:"),
                ("The Parties intend to execute all three documents simultaneously.",
                 "This LTSA is not required to be signed on the same date as the EPC Agreement; it "
                 "must, however, be executed and in force before the EPC advance payment falls due. "
                 "For the complete companion-document suite (including the Advance Payment "
                 "Guarantees, the OEM Performance Guarantee, the EMS Subscription Addendum and the "
                 "Technical Agreements), see Section 2.4 below."),
                # Leak-hygiene: avoid "portfolio" wording implying the wider group order
                ("the BESS portfolio covered under this Agreement",
                 "the BESS systems covered under this Agreement"),
                # Align PI limit to EUR 2,000,000 (consistent with EPC §14.1(b))
                ("General Commercial Liability and Professional Indemnity Insurance of minimum EUR 1,000,000 each.",
                 "General Commercial Liability Insurance of minimum EUR 1,000,000 and Professional "
                 "Indemnity Insurance of minimum EUR 2,000,000."),
                # Fix stale Critical on-site breach reference to match §8.2(a) 8-hour commitment
                ("on-site attendance by the next-Business-Day commitment for a Critical Alert",
                 "on-site attendance within the eight (8) hour commitment for a Critical Alert"),
                # FINAL REVIEW consistency: Scheduled Downtime definition 10 days -> agreed 2 days (48h)
                ("which shall not exceed ten (10) days per calendar year.",
                 "which shall not exceed two (2) days (48 hours) per Park per calendar year."),
                # SOH LD per-kWh: resolve the blank by pointing to the market-rate mechanism already stated
                ("x System Capacity x EUR [\u25cf] per kWh",
                 "x System Capacity x the prevailing market per-kWh rate for equivalent battery modules"),
                # Schedule 2 — System Capacity (30 MWh combined)
                ("| System Capacity | [\u25cf] MWh |",
                 "| System Capacity | 30 MWh |"),
                # Optional warranty extensions — BESS priced per official Linyang Cyprus rate; PCS+MVS
                # extended warranty is not offered (Client has confirmed it does not want this option).
                ("| BESS Performance & Product Warranty | 6-10 | \u20ac913.92 | \u2610 |",
                 "| BESS Performance & Product Warranty | 6-10 | \u20ac913.92/MWh/Yr (\u20ac27,417.60/Yr for 30 MWh) | \u2610 |"),
                ("| BESS Performance & Product Warranty | 11-15 | \u20ac1,157.62 | \u2610 |",
                 "| BESS Performance & Product Warranty | 11-15 | \u20ac1,157.62/MWh/Yr (\u20ac34,728.60/Yr for 30 MWh) | \u2610 |"),
                ("| PCS + MVS Product Warranty | 6-10 | \u20ac747.76 | \u2610 |",
                 "| PCS + MVS Product Warranty | 6-10 | Not offered \u2014 Client has elected BESS-only extension | n/a |"),
                ("| PCS + MVS Product Warranty | 11-15 | \u20ac926.10 | \u2610 |",
                 "| PCS + MVS Product Warranty | 11-15 | Not offered \u2014 Client has elected BESS-only extension | n/a |"),
                ("Note: Extended warranty is available for a maximum of fifteen (15) years from COD (1 cycle "
                 "per day condition). Years 16-20 warranty extension is not available from the OEM. "
                 "Alternative: upfront spares package at \u20ac1,000\u20131,250/MWh.",
                 "Note: Extended warranty is available for a maximum of fifteen (15) years from COD (1 cycle "
                 "per day condition). Years 16-20 warranty extension is not available from the OEM. Pricing "
                 "above is the official Linyang Cyprus quotation for the Project (30 MWh combined System "
                 "Capacity: Galascope 1, 20 MWh + Galascope 2, 10 MWh). Scope is BESS (cells, modules, racks, "
                 "BMS) only \u2014 the PCS/MVS extended-warranty option is not offered under this LTSA at the "
                 "Client's instruction. This extension is optional and is not selected by default; the "
                 "Client may elect either period in writing at any time before expiry of the then-current "
                 "warranty period, per EPC Section 10.6-EXT. Alternative: upfront spares package at "
                 "\u20ac1,000\u20131,250/MWh."),
                # Schedule 1 — EMS party is DISPERON (not a generic third party)
                ("| EMS | [Third Party \u2013 specify] |",
                 "| EMS | DISPERON (Disperon Sp. z o.o.) |"),
                # FINAL SWEEP: Timotheos \u00a75.5 asked to include Client (or Client representatives)
                # training on the Monitoring Platform \u2014 never captured in the response matrix or
                # drafted. Low-cost, easy accept.
                ("The Service Provider shall provide the Client with secure access to the Monitoring "
                 "Platform, enabling the Client to view real-time and historical Performance Data.",
                 "The Service Provider shall provide the Client with secure access to the Monitoring "
                 "Platform, enabling the Client to view real-time and historical Performance Data. The "
                 "Service Provider shall, at no additional charge, provide the Client (or up to three (3) "
                 "Client representatives nominated in writing) with training on the use of the Monitoring "
                 "Platform and portal, comprising one (1) session at or shortly after PAC and one (1) "
                 "refresher session per calendar year on reasonable notice, or promptly following any "
                 "material change to the Monitoring Platform."),
            ],
            rewrites=[
                # ACCEPT: Strip group/portfolio pricing leak (existing)
                ("portfolio range",
                 "Calculation basis: Annual subscription = 20% of total installed EMS/SCADA cost (EMS hardware "
                 "+ SCADA Local + SCADA Global), divided by system capacity in MWh. For this Project the "
                 "EMS/SCADA subscription rate is EUR 400/MWh/Year (EUR 12,000 per year for the 30 MWh combined "
                 "across Galascope 1 and Galascope 2), as set out in the EMS Subscription Addendum. For the "
                 "avoidance of doubt, this EUR 400/MWh/Year EMS/SCADA component is included within — and is not "
                 "additional to — the all-in LTSA Service Fee of EUR 1,740/MWh/Year agreed for this Project; the "
                 "EMS subscription is not invoiced separately on top of the LTSA Service Fee while this Agreement "
                 "is in force. The standalone EUR 400/MWh/Year rate applies only if the Client exits or "
                 "terminates the LTSA and the EMS subscription continues on a standalone basis."),
            # group-availability paragraph handled by the NEGOTIATE rewrite block below
                # ACCEPT: EMS-affiliate carve-out in LTSA §9.3(e) (mirrors EPC 10.5A) — Anastasios #4
                ("(e) Downtime caused by EMS or third-party system failures;",
                 "(e) Downtime caused by EMS or third-party system failures — except that this exclusion shall "
                 "not apply where the EMS failure is attributable to DISPERON (Disperon Sp. z o.o.) or "
                 "any other entity within the Lighthief group of companies, or is caused by the Service "
                 "Provider's configuration, monitoring, integration or cyber-security controls;"),
                # CONFIRMED via direct PDF read of Linyang's "Power Atlantic 5MWh Degradation Curve"
                # document (30 Jul 2026): Galascope 1 (5MW/20.06MWh) and Galascope 2 (2.5MW/10.03MWh)
                # are both ~4-hour systems = the 0.25P operating condition (not 0.5P — an earlier
                # working assumption based on the datasheet's dual rated-power listing was wrong;
                # "P" is normalized to energy content as a notional 1-hour rate, so 1,250 kW per
                # container = 0.25P, matching the installed BCS1250K/4-hour design exactly). The
                # verified 0.25P, 1-cycle/day curve (cross-checked against the source PDF page image,
                # not just the linyang.md transcription) is used throughout §10.2 and Schedule 5:
                # Yr1 94.62 / Yr2 91.77 / Yr3 89.91 / Yr4 88.00 / Yr5 86.78 / Yr6 84.97 / Yr7 83.83 /
                # Yr8 82.25 / Yr9 81.06 / Yr10 79.58 / Yr15 73.61. This matches Stelios's ask exactly
                # and is OEM-backed, not a client concession beyond the manufacturer's own data.
                # §10.2 body milestones:
                ("(a) End of Year 1: 98% of original rated capacity;",
                 "(a) End of Year 1: 94.62% of original rated capacity;"),
                ("(b) End of Year 2: 96% of original rated capacity;",
                 "(b) End of Year 2: 91.77% of original rated capacity;"),
                ("(c) End of Year 3: 94% of original rated capacity;",
                 "(c) End of Year 3: 89.91% of original rated capacity;"),
                ("(d) End of Year 4: 92% of original rated capacity;",
                 "(d) End of Year 4: 88.00% of original rated capacity;"),
                ("(e) End of Year 5: 85% of original rated capacity;",
                 "(e) End of Year 5: 86.78% of original rated capacity;"),
                ("(g) End of Year 15: 70% of original rated capacity.",
                 "(g) End of Year 15: 73.61% of original rated capacity."),
                # Schedule 5 guaranteed 1cpd table — align Y5 and Y15 to the confirmed 0.25P curve
                ("| 5 | 85% |", "| 5 | 86.78% |"),
                ("| 15 | 70% |", "| 15 | 73.61% |"),
                # Year 10 = 79.58% is unchanged (identical under 0.25P, confirmed correct as-is).
                ("(f) End of Year 10: 79.58% of original rated capacity;",
                 "(f) End of Year 10: 79.58% of original rated capacity;"),
                # Schedule 5 "1 cycle per day" guaranteed table — full precision, confirmed 0.25P curve:
                ("| 1 | 95% |", "| 1 | 94.62% |"),
                ("| 2 | 92% |", "| 2 | 91.77% |"),
                ("| 3 | 89% |", "| 3 | 89.91% |"),
                ("| 4 | 87% |", "| 4 | 88.00% |"),
                ("| 6 | 83% |", "| 6 | 84.97% |"),
                ("| 7 | 81% |", "| 7 | 83.83% |"),
                ("| 8 | 79% |", "| 8 | 82.25% |"),
                ("| 9 | 80% |", "| 9 | 81.06% |"),
                ("| 10 | 79.58% |", "| 10 | 79.58% |"),
                # Datasheet spec alignment (Power Atlantic 5.015 MWh):
                ("| Cycle Life (to 70% EOL) | 7,000 cycles (25\u00b0C, 90% DOD, 0.5C) |",
                 "| Cycle Life (to 70% EOL) | 8,000 cycles (25\u00b0C, 90% DOD, 0.5C) |"),
                ("Based on cycle life of 7,000 cycles to 70% EOL at reference conditions:",
                 "Based on cycle life of 8,000 cycles to 70% EOL at reference conditions:"),
                ("| Energy Density | 183 Wh/kg |",
                 "| Energy Density | 175 Wh/kg |"),
                # ACCEPT: Schedule 4 — clarify LD formula (remove ambiguity between % table and EUR/day rate)
                ("Availability LD Rate: \u20ac30/day/MWh of unavailable capacity (per OEM confirmed rate)",
                 "Availability LDs shall be calculated as follows: for each percentage point below 97% "
                 "achieved, the Service Provider shall pay EUR 30 per MWh per day of affected unavailable "
                 "capacity (minimum EUR 500 per day per Park). Where this daily rate would produce a lower "
                 "aggregate than the fee-reduction percentages in the table above, the fee-reduction table "
                 "governs; where the daily rate produces a higher aggregate, the daily rate governs. In all "
                 "cases the cap in Section 9.5 applies."),
                # ACCEPT: Data retention — 15 years + client accessible; warranty claim logs (Dino request)
                ("The Service Provider shall retain Performance Data for a minimum of five (5) years "
                 "from the date of collection.",
                 "The Service Provider shall retain Performance Data, alarm/event logs, BMS/PCS/EMS logs, "
                 "maintenance records, SOH certificates and availability calculations for a minimum of "
                 "fifteen (15) years from the Commissioning Date, and for five (5) years following "
                 "termination. The Client shall have continuous access to and the right to export all "
                 "such data at any time. Data shall be retained in formats acceptable to the OEM (Linyang) "
                 "for warranty claims and to any applicable regulatory authority."),
                # NEGOTIATE: Availability LD cap 20% → 50% of annual Service Fee
                ("The maximum Availability Liquidated Damages payable in any year shall not exceed "
                 "twenty percent (20%) of the annual Service Fee.",
                 "The maximum Availability Liquidated Damages payable in any year shall not exceed "
                 "fifty percent (50%) of the annual Service Fee per Park. This cap shall not apply "
                 "to wilful default or gross negligence by the Service Provider."),
                # NEGOTIATE: Remove 'sole remedy' (§9.6) — client may claim for wilful default/gross negligence/safety
                # Note: the source uses a standard apostrophe in "Client's"
                ("The Availability Liquidated Damages set out in this Section shall be the Client's "
                 "sole and exclusive remedy for failure to achieve the Availability Guarantee, and "
                 "the Client shall not be entitled to claim additional damages for unavailability "
                 "beyond the Liquidated Damages.",
                 "The Availability Liquidated Damages set out in this Section are the agreed primary "
                 "remedy for ordinary failure to achieve the Availability Guarantee. They shall not "
                 "exclude the Client's right to claim for wilful default, gross negligence, or "
                 "repeated safety incidents by the Service Provider, for which the general liability "
                 "provisions of Section 14 apply."),
                # NEGOTIATE: Liability cap 12 months → 24 months
                ("The Service Provider's total aggregate liability under this Agreement shall not "
                 "exceed the total Service Fees paid by the Client in the twelve (12) months "
                 "preceding the claim.",
                 "The Service Provider's total aggregate liability under this Agreement shall not "
                 "exceed the total Service Fees paid by the Client in the twenty-four (24) months "
                 "preceding the claim. This cap shall not apply to wilful misconduct, gross "
                 "negligence, fraud, breach of confidentiality, or the carve-outs in Section 14.3."),
                # NEGOTIATE: Remove CPI+2% escalation — fees locked
                ("The Service Fee may be adjusted annually on each anniversary of the Effective "
                 "Date, by an amount not exceeding the Cyprus Consumer Price Index (CPI) increase "
                 "for the preceding twelve (12) months, plus two percent (2%).",
                 "The Service Fee is fixed for the Initial Term (five (5) years). No adjustment "
                 "shall apply during the Initial Term. After the Initial Term, the Service Fee "
                 "may be adjusted only by written agreement of both Parties."),
                # NEGOTIATE: Scheduled downtime 10 days → 2 days (accepted)
                ("Total Scheduled Downtime shall not exceed ten (10) days per calendar year, "
                 "distributed across preventive maintenance visits, OEM service windows, and "
                 "system health checks.",
                 "Total Scheduled Downtime shall not exceed two (2) days (48 hours) per Park per "
                 "calendar year, excluding Force Majeure events. Maintenance visits shall be "
                 "pre-scheduled at least thirty (30) days in advance, confirmed by the Client, "
                 "and performed in low-dispatch windows (preferably between 04:00 and 08:00 local "
                 "time). Scheduled Downtime exceeding 48 consecutive hours per Park requires the "
                 "Client's prior written approval except for safety emergencies."),
                # NEGOTIATE: Availability stays group level — rename header and add explanatory note.
                # Target the original source text for the paragraph (before any other rewrite touches it).
                ("9.2A Group-Level Rationale",
                 "9.2A Group-Level Availability Rationale and LD Protection"),
                # This paragraph is the one that in the source reads "The 97% Availability Guarantee
                # applies at the group portfolio level..." — the "reflects the portfolio pricing basis"
                # rewrite also targets it, so we combine both into one final form here.
                ("The 97% Availability Guarantee applies at the group portfolio level, not per "
                 "individual park. This means that an isolated park outage does not constitute a "
                 "breach of the guarantee provided the weighted average availability across all "
                 "active parks exceeds 97%. This structure reflects the portfolio pricing basis "
                 "under which the Tier C service fee was calculated and enables the Service "
                 "Provider to allocate resources efficiently across the Client's park portfolio.",
                 "The 97% Availability Guarantee applies at the Project level across Galascope 1 "
                 "and Galascope 2 in aggregate, consistent with the Parties\u2019 agreement. An isolated "
                 "park outage does not constitute a breach of the guarantee provided the weighted "
                 "average availability across both Project parks exceeds 97%. This structure was "
                 "agreed by the Parties and reflects the aggregated (project-level) basis under which "
                 "the Tier C service fee was agreed. Note: the group-level calculation does not "
                 "eliminate the financial consequence of individual park underperformance \u2014 the "
                 "Availability LD cap has been increased to fifty percent (50%) of the annual "
                 "Service Fee (Section 9.5) to provide meaningful compensation even where one "
                 "park drives the shortfall. The Parties acknowledge that requiring per-park 97% "
                 "guarantees on individual smaller parks would require a materially higher service "
                 "fee and was not the basis on which the Tier C fee was agreed."),
                # NEGOTIATE: Spare parts — critical sub-components; counter on transformer (TBD)
                ("(d) Critical spare parts include, at minimum:",
                 "(d) Critical spare parts stocked in the local warehouse include, at minimum: "
                 "(Note: full transformer unit stocking is subject to further discussion given the "
                 "distinct transformer ratings across Galascope 1 and Galascope 2; the Parties shall "
                 "agree a transformer-spares arrangement within thirty (30) days of the LTSA "
                 "Effective Date. In the interim, the Service Provider commits to a four (4) week "
                 "OEM supply lead-time for transformer units.) Critical sub-components include:"),
                # NEGOTIATE (accept client): termination for convenience — Client 90 days / SP 12 months
                ("Either Party may terminate this Agreement for convenience by providing one "
                 "hundred eighty (180) days' written notice.",
                 "The Client may terminate this Agreement for convenience by providing ninety (90) "
                 "days' written notice. The Service Provider may terminate for convenience only "
                 "after the Initial Term, by providing twelve (12) months' written notice, provided "
                 "that such termination shall not prejudice warranty, data-handover or transition "
                 "obligations."),
                # NEGOTIATE (accept client): systemic SOH underperformance trigger 5% -> 3%
                ("(b) SYSTEMIC UNDERPERFORMANCE (affecting 5% or more of total capacity):",
                 "(b) SYSTEMIC UNDERPERFORMANCE (affecting 3% or more of total capacity):"),
                ("(Affecting \u2265 5% of total system capacity)",
                 "(Affecting \u2265 3% of total system capacity)"),
                # NEGOTIATE (item 7): €1,740/MWh/yr all-in rate INCLUDING the EMS/SCADA subscription
                ("SELECTED SERVICE TIER: \u2610 Tier A  \u2610 Tier B  \u2610 Tier C",
                 "SELECTED SERVICE TIER: \u2610 Tier A  \u2610 Tier B  \u2612 Tier C (default)\n"
                 "AGREED ALL-IN RATE: EUR 1,740 / MWh / Year, which is inclusive of the DISPERON "
                 "EMS/SCADA software subscription (i.e. the EMS subscription is NOT charged "
                 "separately on top of this rate). This all-in Tier C rate is the rate agreed for "
                 "this Project."),
                # ACCEPT: Scheduled downtime 240h/10 days reference in Schedule 4
                ("(maximum 240 hours / 10 days per year)",
                 "(maximum 48 hours / 2 days per year per Park)"),
                # AGREED SLA RESPONSE TIMES (settled Jul 2026)
                # Critical: immediate remote, 30-min ack, 24h on-site
                ("(a) Critical Alerts: Initial remote response within four (4) hours; on-site "
                 "attendance (if required) by the end of the next Business Day. For any safety "
                 "risk the Service Provider shall apply immediate remote isolation or safe-shutdown.",
                 "(a) Critical Alerts: Initial remote response immediately upon alert generation "
                 "(automated 24/7 platform); duty technician acknowledgement within thirty (30) "
                 "minutes; on-site attendance (if required) within eight (8) hours. "
                 "For any safety risk the Service Provider shall apply immediate remote isolation "
                 "or safe-shutdown command without waiting for acknowledgement."),
                # Major: 4h remote, 6h on-site
                ("(b) Major Alerts: Initial remote response within one (1) Business Day; "
                 "on-site attendance (if required) within five (5) Business Days.",
                 "(b) Major Alerts: Initial remote response within four (4) hours; on-site "
                 "attendance (if required) within six (6) hours of the alert."),
                # Minor: 12h remote, 5 Business Days on-site
                ("(c) Minor Alerts: Initial remote response within forty-eight (48) hours; "
                 "on-site attendance (if required) at the next scheduled maintenance visit "
                 "or within five (5) business days, whichever is sooner.",
                 "(c) Minor Alerts: Initial remote response within twelve (12) hours; on-site "
                 "attendance (if required) at the next scheduled maintenance visit or within "
                 "five (5) Business Days, whichever is sooner."),
                # Auto-escalation: 30 minutes (was 1 hour)
                ("If the primary duty technician does not acknowledge a Critical Alert within "
                 "one (1) hour, the Monitoring Platform shall automatically escalate the alert "
                 "to the backup technician ",
                 "If the primary duty technician does not acknowledge a Critical Alert within "
                 "thirty (30) minutes, the Monitoring Platform shall automatically escalate "
                 "the alert to the backup technician "),
                # Management escalation trigger: 2 hours (was 4 hours)
                ("If on-site attendance has not been dispatched within four (4) hours of a "
                 "Critical Alert, the operations manager shall escalate",
                 "If on-site attendance has not been dispatched within two (2) hours of a "
                 "Critical Alert, the operations manager shall escalate"),
                # ACCEPT: Schedule 2 note — now populated with the agreed fee
                ("SYSTEM CAPACITY: [●] MWh",
                 "SYSTEM CAPACITY: 30 MWh (Galascope 1: 20 MWh + Galascope 2: 10 MWh)\n"
                 "AGREED ANNUAL SERVICE FEE: EUR 52,200 / Year (30 MWh x EUR 1,740 / MWh all-in), "
                 "excluding VAT; EUR 62,118 / Year including 19% VAT. Fixed for the 15-year Initial Term "
                 "(no CPI escalation). Site, equipment serial numbers and dates in Schedule 1 are "
                 "completed at the order trigger (Connection Terms)."),
                # FINAL REVIEW: populate Schedule 2 fee cells (all-in EUR 1,740/MWh; EMS EUR 400/MWh included)
                ("| Preventive Maintenance \u2013 BESS | \u20ac1,157.62 | \u20ac[\u25cf] |",
                 "| Preventive Maintenance \u2013 BESS | (list ref.) | see all-in |"),
                ("| Preventive & Corrective Maintenance \u2013 PCS + MVS | \u20ac1,311.97 | \u20ac[\u25cf] |",
                 "| Preventive & Corrective Maintenance \u2013 PCS + MVS | (list ref.) | see all-in |"),
                ("| 97% Availability Guarantee (Years 1-15) | \u20ac2,201.73 | \u20ac[\u25cf] |",
                 "| 97% Availability Guarantee (Years 1-15) | (list ref.) | see all-in |"),
                ("| SOH/Performance Warranty | Price on Application | \u20ac[\u25cf] |",
                 "| SOH/Performance Warranty | Price on Application | not selected |"),
                ("BASE SERVICE PRICING (Per MWh per Year)",
                 "BASE SERVICE PRICING (Per MWh per Year) \u2014 list/reference build-up; the agreed Tier C "
                 "charge is the all-in EUR 1,740/MWh/Year in the Annual Fee Summary below"),
                ("| EMS/SCADA Software Subscription | \u20ac[\u25cf] | \u20ac[\u25cf] |",
                 "| EMS/SCADA Software Subscription | \u20ac400 | \u20ac12,000 |"),
                ("| Installed EMS/SCADA Cost | \u20ac[\u25cf] |",
                 "| Installed EMS/SCADA Cost | \u20ac60,000 |"),
                ("| Annual Subscription (20%) | \u20ac[\u25cf] |",
                 "| Annual Subscription (20%) | \u20ac12,000 |"),
                ("| EUR/MWh/Year | \u20ac[\u25cf] |",
                 "| EUR/MWh/Year | \u20ac400 |"),
                ("| Base Service Fee (Selected Tier) | \u20ac[\u25cf] |",
                 "| Base Service Fee (Tier C, all-in, incl. EMS) | \u20ac52,200 |"),
                ("| Warranty Extension (if selected) | \u20ac[\u25cf] |",
                 "| Warranty Extension (if selected) | \u2014 (none selected) |"),
                ("| EMS/SCADA Annual Subscription | \u20ac[\u25cf] |",
                 "| EMS/SCADA Annual Subscription | \u20ac12,000 (included in Base above) |"),
                ("| TOTAL ANNUAL FEE | \u20ac[\u25cf] |",
                 "| TOTAL ANNUAL FEE | \u20ac52,200 |"),
                ("| VAT (19%) | \u20ac[\u25cf] |",
                 "| VAT (19%) | \u20ac9,918 |"),
                ("| TOTAL INCLUDING VAT | \u20ac[\u25cf] |",
                 "| TOTAL INCLUDING VAT | \u20ac62,118 |"),
            ]
        )
        _append_annex(
            ltsa_dest,
            "ANNEX LT-2 \u2014 AGREED TECHNICAL & SERVICE AMENDMENTS (v6.3, July 2026)",
            "This Annex incorporates the technical and service clauses agreed on the client "
            "(Timotheos / technical review) redline of July 2026. It forms part of this LTSA. Where "
            "this Annex is inconsistent with the body, this Annex prevails. Defined terms have the "
            "meaning given in the body.",
            LTSA_ANNEX,
        )
        print("  fixed + copied -> 03-LTSA-Galascope-G1-G2.docx (+ Annex LT-2)")
    else:
        print(f"  !! MISSING: {ltsa_src}")

    # EMS addendum — copy with inline LTSA reference fix (LCY-LTSA-GAL-B1-2026 → LCY-LTSA-GAL-2026)
    ems_src = CONTRACTS / "EMS-Subscription-Galascope-may2026.docx"
    ems_dest = PKG / "07-EMS-Subscription-Addendum.docx"
    if ems_src.exists():
        _fix_docx_text(ems_src, ems_dest, replacements=[
            ("LCY-LTSA-GAL-B1-2026", "LCY-LTSA-GAL-2026"),
        ], rewrites=[
            ("invoiced annually in advance from the PAC date",
             "The subscription fee of EUR 12,000.00 per year (ex VAT) is invoiced annually in advance "
             "from the PAC date (set per the EPC Delivery Schedule confirmed at Connection Terms). "
             "Payment due within 30 days of invoice. Fee is exclusive of VAT (19%). For the avoidance "
             "of doubt, while the LTSA (LCY-LTSA-GAL-2026) is in force this EMS/SCADA subscription is "
             "included within the all-in LTSA Service Fee of EUR 1,740/MWh/Year and is not invoiced "
             "separately in addition to it. The standalone pricing in this Addendum applies only if "
             "the Client exits or terminates the LTSA and the EMS subscription continues on a "
             "standalone basis; for as long as the LTSA remains in force the EMS subscription is "
             "included in the LTSA Service Fee and is not charged separately."),
            ("persists for the lifetime",
             "The Linyang OEM product warranty and Direct Warranty Undertaking survive termination of the "
             "Lighthief\u2013Linyang Distribution Agreement, so OEM product cover does not depend on that "
             "arrangement. The BMS northbound communication interface (Modbus TCP / RS485) is a standard "
             "OEM interface, and its compatibility with the DISPERON EMS is confirmed in the Technical "
             "Agreement. Continuity of DISPERON EMS data access is provided under this Subscription "
             "Agreement and is not contingent on the Distribution Agreement."),
        ])
        print("  fixed + copied -> 07-EMS-Subscription-Addendum.docx")
    else:
        print(f"  !! MISSING: {ems_src}")


if __name__ == "__main__":
    build_epc_v6()
    build_cover_note()
    copy_attachments()
    print("\nDONE. Package folder:", PKG)
