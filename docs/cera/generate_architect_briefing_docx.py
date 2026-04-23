#!/usr/bin/env python3
"""Generate architect briefing DOCX for BESS under Κ.Δ.Π. 15/2026 & 17/2026. Run: python3 generate_architect_briefing_docx.py"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x1A, 0x36, 0x5D)
OUT = Path(__file__).resolve().parent / "briefing-architect-bess-kdp-2026.docx"


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_fill="1E3A5D") -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_fill)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6)

    def h1(text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = BLUE

    def h2(text: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)

    def h3(text: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
        p.paragraph_format.space_before = Pt(10)

    def para(text: str) -> None:
        doc.add_paragraph(text)

    def bl(text: str) -> None:
        p = doc.add_paragraph(text, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)

    h1("Briefing for the Authorised Architect / ETEK Designer")
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sp.add_run(
        "BESS in a licensed RES (solar) park — planning and building exemption route\n"
        "Κ.Δ.Π. 15/2026 and Κ.Δ.Π. 17/2026 (Official Gazette 5992, 16 January 2026)"
    )
    s.font.size = Pt(11)
    s.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()
    pco = doc.add_paragraph()
    pco.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = pco.add_run("Lighthief Cyprus Ltd")
    r0.bold = True
    r0.font.size = Pt(10)
    for line in [
        "Ref: LCY-ARCH-BRIEF-KDP-2026-001",
        "Version: 1.0   |   Date: 24 April 2026",
    ]:
        pl = doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.add_run(line).font.size = Pt(10)

    doc.add_page_break()

    h2("1. Purpose of this briefing")
    para(
        "Lighthief supports clients who add a containerised Battery Energy Storage System (BESS) to an existing "
        "CERA-licensed solar (RES) station. When the conditions of the 2026 special development decrees are met, a "
        "separate building permit for the storage may be treated as deemed issued, and the town-planning route is "
        "handled through a package centred on a signed Υπεύθυνη δήλωση μελετητή (responsible study declaration) and "
        "the drawings listed in the decrees."
    )
    para(
        "This document explains what we need from you as the ETEK-registered architect (or, where the design office "
        "splits scope, the lead planning and layout designer) so the submitted plans are consistent with "
        "Κ.Δ.Π. 15/2026 and 17/2026 and with the way we annotate our BESS layout drawings (distances, height, access)."
    )

    h2("2. CERA, EOA, and where forms are filed")
    para("The following points avoid mixing electricity licensing with the planning / building file:")
    bl(
        "CERA (ΡΑΕΚ) is the electricity regulator (generation licences, market). It is not the address for these "
        "particular planning and building exemption forms."
    )
    bl(
        "The decrees use αρμόδια αρχή (competent planning authority). For Κ.Δ.Π. 15/2026, the responsible declaration "
        "is submitted to an e-mail address published on the website of the local EOA (Επαρχιακός Οργανισμός "
        "Αυτοδιοίκησης — district administration), unless replaced by a centralised portal."
    )
    bl(
        "For Κ.Δ.Π. 17/2026, Παράρτημα II material is submitted through the ΙΠΠΟΔΑΜΟΣ information system, as stated in "
        "the decree. Always use the current instructions and addresses for the client’s district."
    )

    h2("3. What we ask you to do")
    bl(
        "Review the annotated layout and Χωροταξικό we provide: boundary setbacks, 3 m height envelope, access for fire, "
        "paving limit, and positions of battery containers and MV skids on the licensed plot."
    )
    bl(
        "Confirm that, on the information given, the proposal is consistent with the spatial and visual conditions in "
        "Κ.Δ.Π. 15/2026 and 17/2026, subject to correct survey, cadastre, and any Development Plan or Ministerial Order "
        "that applies to the specific parcel."
    )
    bl(
        "Sign the official state appendix as the Μελετητής where your ETEK specialty covers this scope, and provide "
        "ETEK registration and professional indemnity particulars as required on the form."
    )
    bl(
        "Coordinate with civil/structural and electrical consultants: your deliverable is planning and layout plus the "
        "declaration; foundations, earthing, and grid/DSO documentation are other disciplines unless your firm also holds "
        "those categories."
    )

    h2("4. Legal basis (at a glance)")
    add_table(
        doc,
        ["Instrument", "Role"],
        [
            (
                "Κ.Δ.Π. 15/2026",
                "Roads and buildings special development decree: BESS (and optional data centre) in an existing licensed "
                "RES site; Appendix: Υπεύθυνη δήλωση μελετητή; submission to competent authority (EOA e-mail in "
                "transitional period per decree).",
            ),
            (
                "Κ.Δ.Π. 17/2026",
                "Town planning special development decree; repeals Κ.Δ.Π. 215/2025. Appendix I: conditions by sector; "
                "Appendix II: required drawings; submission through ΙΠΠΟΔΑΜΟΣ as per Article 5.",
            ),
            (
                "Ν. 180(Ι)/2025 / Ν. 181(Ι)/2025",
                "Enabling deemed permit routes as referenced in the decrees (as in force).",
            ),
        ],
    )

    h2("5. Conditions you will see on our marked drawings")
    para(
        "The table is a working checklist. Final numbers must follow the current official Greek text of the decrees and "
        "any written clarification from the competent authority."
    )
    add_table(
        doc,
        ["Topic", "Typical rule (verify in decree and on site)"],
        [
            (
                "Setback from plot boundary",
                "At least 6,00 m from boundaries, including where the boundary is formed by the road network "
                "(see Κ.Δ.Π. 17 §4 and parallel conditions in Κ.Δ.Π. 15).",
            ),
            (
                "Height",
                "Maximum 3,00 m from the lower of natural or finished ground, including support structures and "
                "roof-mounted services (e.g. HVAC). Multiple units are not placed one above the other above that line.",
            ),
            (
                "Paving / hardstanding",
                "Cementing limited to the minimum necessary for the equipment installation.",
            ),
            (
                "Topography",
                "No material alteration of natural ground levels except what the installation reasonably requires.",
            ),
            (
                "Finishes / appearance",
                "Neutral, earth-tone claddings; no bright or highly reflective treatment where the decree so requires.",
            ),
            (
                "Fire (layout-related)",
                "Perimeter access for fire vehicles: commonly dimensioned at 6 m in the fire section of the exemption "
                "decrees; final detail per decree text and the Fire Service where they opine on the file.",
            ),
        ],
    )

    h2("6. Κ.Δ.Π. 17/2026 — Παράρτημα ΙΙ (supporting your submission)")
    para(
        "The official gazette lists the exact documents. Typically the pack includes, among other items: cadastral/ "
        "property particulars; a copy of existing planning and/or building permissions for the RES plant; a "
        "Χωροταξικό Διάγραμμα Δόμησης showing the licensed works, the proposed storage, set-out distances to plot "
        "boundaries, and spot levels on the development parcel, adjoining plots, and roads; owner or co-owner consent; "
        "and ETEK professional insurance proof for the signatory, as listed."
    )

    h2("7. Lighthief deliverables vs architect deliverables")
    h3("Lighthief / EPC (typical input)")
    bl("Marked plan: container footprints, key dimensions, height check versus 3 m, access and circulation, trench routes where shown.")
    bl("Reference dimensions from the OEM (e.g. 20 ft / 40 ft container external sizes) for drawing consistency.")
    h3("Your office (typical output)")
    bl("Final drawing sheet(s) in the form and language expected by the authority: Greek labels, scale, north, legend, alignment with the cadastral print where required.")
    bl("Signed Υπεύθυνη δήλωση on the official form attached to the published decree, consistent with your professional review.")

    h2("8. Outside this briefing (other consultants)")
    para(
        "Structural design of the reinforced concrete pad (a decree condition), external lightning and earthing, "
        "electrical and inverter/MV work, fire suppression system detail, and EAC/DSO grid application materials are not "
        "covered by this document unless your firm is retained for those specific disciplines."
    )

    h2("9. Next steps")
    bl("Issue us a marked-up list if any dimension or boundary on our plan conflicts with the exemption or the development plan for the site.")
    bl("Complete the state form with your ETEK and insurance block as published in Κ.Δ.Π. 15/2026 or 17/2026.")
    doc.add_paragraph()
    pc = doc.add_paragraph()
    pc.add_run("Contact: ").bold = True
    pc.add_run("office@lighthief.com")
    pc.add_run("  |  ")
    pc.add_run("https://solarfarms.cy")

    h2("10. Internal reference documents (Lighthief repo)")
    para(
        "docs/cera/ (official PDFs and README), docs/CyprusDSO.md (Section 18 — permit exemption terms), and "
        "docs/internal/civil-engineer-dso-bess-positioning-guide.html (spacing and TSO/DSO layout cross-checks)."
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
