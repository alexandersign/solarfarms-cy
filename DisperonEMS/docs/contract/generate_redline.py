"""
Generate DISPERON_SHA_Redline_Lighthief.docx
Voltus draft (2026-04-20) with Lighthief counter-proposals marked as:
  - RED STRIKETHROUGH  = Voltus text Lighthief rejects / deletes
  - BLUE UNDERLINE     = Lighthief counter / inserted text
  - BLACK normal       = accepted / unchanged text
  - [LH NOTE] boxes   = explanatory margin notes
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── colours ──────────────────────────────────────────────────────────────────
BLUE   = RGBColor(0x1F, 0x49, 0x7D)
RED    = RGBColor(0xC0, 0x00, 0x00)
LH_BLU = RGBColor(0x00, 0x50, 0xC0)   # insertion colour
BLACK  = RGBColor(0x00, 0x00, 0x00)
GREY   = RGBColor(0x50, 0x50, 0x50)
AMBER  = RGBColor(0xC0, 0x60, 0x00)
GREEN  = RGBColor(0x1A, 0x6B, 0x30)


# ── xml helpers ───────────────────────────────────────────────────────────────
def _set_strike(run, on=True):
    rPr = run._r.get_or_add_rPr()
    s = OxmlElement('w:strike')
    if on:
        rPr.append(s)


def _set_hl(run, colour_hex):
    """Set character highlight (shading) via rPr."""
    rPr = run._r.get_or_add_rPr()
    hl = OxmlElement('w:highlight')
    hl.set(qn('w:val'), colour_hex)
    rPr.append(hl)


def _set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)


# ── run factories ─────────────────────────────────────────────────────────────
def keep(para, text, bold=False, size=10):
    """Normal black — accepted text."""
    if not text:
        return
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = BLACK
    return run


def strike_red(para, text, size=10):
    """Red strikethrough — Voltus text being deleted."""
    if not text:
        return
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RED
    _set_strike(run)
    return run


def insert_blue(para, text, bold=False, size=10):
    """Blue underline — Lighthief counter / new text."""
    if not text:
        return
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = LH_BLU
    run.font.underline = True
    return run


# ── block builders ────────────────────────────────────────────────────────────
def setup(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    for s in doc.sections:
        s.top_margin    = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin   = Cm(2.8)
        s.right_margin  = Cm(2.8)


def lh_note(doc, text):
    """Yellow-highlighted [LIGHTHIEF NOTE] block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    lbl = p.add_run("[LIGHTHIEF NOTE]  ")
    lbl.font.bold  = True
    lbl.font.size  = Pt(9)
    lbl.font.color.rgb = AMBER
    _set_hl(lbl, 'yellow')
    body = p.add_run(text)
    body.font.size  = Pt(9)
    body.font.italic = True
    body.font.color.rgb = RGBColor(0x60, 0x40, 0x00)
    _set_hl(body, 'yellow')
    return p


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.bold = True
        run.font.size = Pt(13)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.bold = True
        run.font.size = Pt(11)
    return p


def plain(doc, *segments):
    """
    Each segment is a tuple: (style, text) where style is
    'keep', 'del', or 'ins'.  Or pass a plain string → keep.
    """
    p = doc.add_paragraph()
    for seg in segments:
        if isinstance(seg, str):
            keep(p, seg)
        elif seg[0] == 'keep':
            keep(p, seg[1], bold=seg[2] if len(seg) > 2 else False)
        elif seg[0] == 'del':
            strike_red(p, seg[1])
        elif seg[0] == 'ins':
            insert_blue(p, seg[1], bold=seg[2] if len(seg) > 2 else False)
    return p


def bul(doc, *segments):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    for seg in segments:
        if isinstance(seg, str):
            keep(p, seg)
        elif seg[0] == 'keep':
            keep(p, seg[1])
        elif seg[0] == 'del':
            strike_red(p, seg[1])
        elif seg[0] == 'ins':
            insert_blue(p, seg[1])
    return p


def tbl_redline(doc, data, col_widths=None):
    """
    data: list of row-lists; each cell is a list of segments or plain string.
    Row 0 is always the header (blue bg).
    A row starting with 'del_row' is fully struck-through.
    A row starting with 'ins_row' is fully underlined blue.
    """
    rows = [r for r in data if not isinstance(r[0], str) or r[0] not in ('del_row', 'ins_row', 'meta')]
    header_only_rows = len(data)
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.style = 'Table Grid'
    for i, row_data in enumerate(data):
        is_del = False
        is_ins = False
        if isinstance(row_data[0], str) and row_data[0] == 'del_row':
            is_del = True
            row_data = row_data[1:]
        if isinstance(row_data[0], str) and row_data[0] == 'ins_row':
            is_ins = True
            row_data = row_data[1:]
        for j, cell_val in enumerate(row_data):
            if j >= len(t.rows[i].cells):
                break
            c = t.rows[i].cells[j]
            if i == 0:
                _set_cell_bg(c, '1F497D')
            elif is_del:
                _set_cell_bg(c, 'FFE0E0')
            elif is_ins:
                _set_cell_bg(c, 'E0F0FF')
            p = c.paragraphs[0]; p.clear()
            if isinstance(cell_val, str):
                run = p.add_run(cell_val)
                run.font.size = Pt(9)
                run.font.bold = (i == 0)
                if i == 0:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                elif is_del:
                    run.font.color.rgb = RED
                    _set_strike(run)
                elif is_ins:
                    run.font.color.rgb = LH_BLU
                    run.font.underline = True
                else:
                    run.font.color.rgb = BLACK
            elif isinstance(cell_val, list):
                for seg in cell_val:
                    if isinstance(seg, str):
                        r = p.add_run(seg)
                        r.font.size = Pt(9)
                        if i == 0:
                            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            r.font.bold = True
                    elif seg[0] == 'del':
                        r = p.add_run(seg[1])
                        r.font.size = Pt(9); r.font.color.rgb = RED
                        _set_strike(r)
                    elif seg[0] == 'ins':
                        r = p.add_run(seg[1])
                        r.font.size = Pt(9); r.font.color.rgb = LH_BLU
                        r.font.underline = True
                    elif seg[0] == 'keep':
                        r = p.add_run(seg[1])
                        r.font.size = Pt(9); r.font.color.rgb = BLACK
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                if ci < len(row.cells):
                    row.cells[ci].width = Cm(w)
    return t


# ═══════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()
    setup(doc)

    # ── COVER ─────────────────────────────────────────────────────────────────
    cov = doc.add_paragraph()
    cov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cov.add_run(
        "SHAREHOLDERS AGREEMENT AND INTELLECTUAL PROPERTY LICENCE\n"
        "LIGHTHIEF EUBESS LTD — HE 474192\n"
    )
    r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s1 = sub.add_run("LIGHTHIEF COUNTER-REDLINE  ·  ")
    s1.font.size = Pt(11); s1.font.bold = True; s1.font.color.rgb = LH_BLU
    s2 = sub.add_run("Based on Voltus draft dated 20 April 2026")
    s2.font.size = Pt(11); s2.font.italic = True; s2.font.color.rgb = GREY

    leg = doc.add_paragraph()
    leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt, col in [
        ("■ RED STRIKETHROUGH", RED),
        (" = rejected / deleted text    ", BLACK),
        ("■ BLUE UNDERLINE", LH_BLU),
        (" = counter-proposed / inserted text    ", BLACK),
        ("■ BLACK", BLACK),
        (" = accepted unchanged", BLACK),
    ]:
        r2 = leg.add_run(txt)
        r2.font.size = Pt(9); r2.font.color.rgb = col
        if "STRIKETHROUGH" in txt:
            _set_strike(r2)
        if "UNDERLINE" in txt:
            r2.font.underline = True

    doc.add_paragraph()

    # ── PREAMBLE ──────────────────────────────────────────────────────────────
    h1(doc, "PREAMBLE")
    plain(doc, "This Shareholders Agreement and Intellectual Property Licence (the \"Agreement\") is entered "
          "into on the date last signed below by and among the following parties:")
    doc.add_paragraph()
    for party in [
        ('keep', 'PARTY 1 — LIGHTHIEF INTERNATIONAL ', True),
        ('keep', 'A company duly incorporated and existing under applicable law, represented by its '
         'Director, Mr. Arkadiusz Sybaris ("Lighthief International" or "Majority Shareholder").'),
    ]:
        plain(doc, party)
    for party in [
        ('keep', 'PARTY 2 — LIGHTHIEF EUBESS LTD ', True),
        ('keep', 'A private limited company registered in Cyprus under registration number HE 474192, '
         'having its registered office at 28 Oktovriou & Aemiliou Chourmouziou, Lophitis Business '
         'Center I, Floor 2, Office 1, 3035 Limassol, Cyprus, represented by its Director, '
         'Mr. Arkadiusz Sybaris (the "Company").'),
    ]:
        plain(doc, party)
    plain(doc, ('keep', 'PARTY 3 — MR. MARCIN SZUMIŁO  ', True),
          ('keep', 'An individual, holder of a Polish identity document ("Mr. Szumiło").'))
    plain(doc, ('keep', 'PARTY 4 — MR. ANDRZEJ LECHOWICZ  ', True),
          ('keep', 'An individual, holder of a Polish identity document ("Mr. Lechowicz").'))
    plain(doc, ('keep', 'PARTY 5 — MR. KAMIL TALAR  ', True),
          ('keep', 'An individual, holder of a Polish identity document ("Mr. Talar").'))
    plain(doc, ('keep', 'PARTY 6 — VOLTUS ENERGY SP. Z O.O.  ', True),
          ('keep', 'A limited liability company incorporated under the laws of the Republic of Poland, '
           'NIP: 1990133260, represented by Mr. Marcin Szumiło and Mr. Andrzej Lechowicz ("Voltus").'))

    # ── ARTICLE 1 ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "ARTICLE 1 — DEFINITIONS")

    # EMS Software definition — source code removed by Voltus, restore it
    h2(doc, "\"EMS Software\" or \"EMS\"")
    plain(doc,
          ('keep', 'means the energy management system and SCADA software platform developed by Voltus '
           'Energy Sp. z o.o., together with all associated '),
          ('del', 'documentation, algorithms, interfaces, configurations,'),
          ('ins', 'source code, object code, documentation, algorithms, interfaces, configurations, '
           'updates, enhancements, and derivative works,'),
          ('keep', ' currently branded as Energy Copilot and described at https://voltusenergy.pl/produkt/.'))
    lh_note(doc, ("Source code, object code, updates, enhancements, and derivative works restored to "
                  "definition. These are essential for the escrow and perpetual licence to have any value."))

    # Software Income / Subscription Income — rename and clarify
    h2(doc, "\"Software Margin\" / \"Software Income\"")
    plain(doc,
          ('del', '"Software Income"'),
          ('ins', '"Software Margin"'),
          ('keep', ' means the Software Licence Component of a customer contract, calculated as the '
           'applicable Per-MWh Licence Rate from Exhibit C multiplied by the installed MWh capacity '
           'of the relevant project. All customer invoices shall itemise the Software Licence Component '
           'and the Commissioning Component as separate line items. The Commissioning Component is '
           'excluded from Software Margin for all purposes of this Agreement.'))
    lh_note(doc, "Lighthief retains the defined term 'Software Margin' (not 'Software Income') for "
            "consistency with Exhibit C and Art 6.5.")

    h2(doc, "\"Subscription Income\" [NEW VOLTUS DEFINITION]")
    plain(doc,
          ('del', '"Subscription Income" means the Software Annual/month Subscription Licence Component '
           'of a customer contract... The Commissioning Component is excluded from Software Margin for '
           'all purposes of this Agreement.'),
          ('ins', '[DEFINITION ACCEPTED IN PRINCIPLE subject to clarification in Art 6.5(c) that '
           'subscription fees share the same Development Bonus rate as the Software Margin, not a '
           'higher rate, and that the base calculation method mirrors Exhibit C.]'))
    lh_note(doc, "A subscription revenue stream is acceptable but must be governed by the same "
            "Exhibit C formula. The definition as drafted is circular and references 'Software Margin' "
            "in a contradictory way.")

    h2(doc, "\"Development Bonus\"")
    plain(doc,
          ('keep', 'means the payment due from the Company to the Voltus Parties in respect of each '
           'Licence sold after the Milestone Date, calculated as '),
          ('del', '30% (thirty percent)'),
          ('ins', '25% (twenty-five percent)'),
          ('keep', ' of the Software Margin received from the relevant customer.'))
    lh_note(doc, "Rate restored to 25% as agreed in SHA v3. The 30%/50% sliding scale in Art 6.5 "
            "is separately addressed.")

    h2(doc, "\"BESS Manufacturer\" and \"Partners\"")
    plain(doc,
          ('del', '"BESS Manufacturer" means any manufacturer of battery energy storage systems (BESS), '
           'whether existing at the date of this Agreement or identified at any time thereafter. '
           'For the avoidance of doubt, cooperation with BESS Manufacturers shall not be subject to '
           'any territorial, exclusivity, or partner-list limitations under this Agreement. '
           '"Partners" means the business partners of Voltus Energy Sp. z o.o., including but not '
           'limited to integrators, distributors, and resellers, as expressly listed in Exhibit B...'),
          ('ins', '"BESS Manufacturer Partners" means the list of hardware manufacturers and business '
           'partners set out in Exhibit B, as agreed and signed by the Parties on the date of this '
           'Agreement and as may be updated from time to time by mutual written agreement of all Parties. '
           'Voltus direct supply rights apply only to the manufacturers expressly listed in Exhibit B.'))
    lh_note(doc, "CRITICAL REJECTION: The split of 'BESS Manufacturer' (unrestricted) and 'Partners' "
            "(Exhibit B only) nullifies all exclusivity provisions. The single 'BESS Manufacturer "
            "Partners' definition tied to Exhibit B is restored. Voltus cannot supply any BESS "
            "manufacturer outside Exhibit B without DISPERON's consent.")

    # ── ARTICLE 2 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 2 — SHARE STRUCTURE AND TRANSFER")
    h2(doc, "2.2  Transfer of Shares")
    lh_note(doc, "CRITICAL: Voltus changed the share structure from 76/8/8/8 to 75/6/6/13. "
            "Exhibit A was NOT updated and still shows 8%/8%/8%, creating an irreconcilable "
            "internal inconsistency. Lighthief rejects this restructure and restores the agreed "
            "76/8/8/8 allocation. If Talar's allocation is to increase, this requires a separate "
            "board resolution and updated Exhibit A — it cannot be done by mid-draft amendment.")
    tbl_redline(doc, [
        ["Shareholder", "Shareholding", "Percentage (%)", "Purchase Consideration"],
        ['del_row',
         [('del', 'Lighthief International')],
         [('del', '75 shares')],
         [('del', '75%')],
         [('del', 'N/A (Retained)')]],
        ['ins_row',
         [('ins', 'Lighthief International')],
         [('ins', '76 shares')],
         [('ins', '76%')],
         [('ins', 'N/A (Retained)')]],
        ['del_row',
         [('del', 'Mr. Marcin Szumiło')],
         [('del', '6 shares')],
         [('del', '6%')],
         [('del', 'EUR 200')]], # Voltus version
        ['ins_row',
         [('ins', 'Mr. Marcin Szumiło')],
         [('ins', '8 shares')],
         [('ins', '8%')],
         [('ins', 'EUR 200')]],
        ['del_row',
         [('del', 'Mr. Andrzej Lechowicz')],
         [('del', '6 shares')],
         [('del', '6%')],
         [('del', 'EUR 200')]],
        ['ins_row',
         [('ins', 'Mr. Andrzej Lechowicz')],
         [('ins', '8 shares')],
         [('ins', '8%')],
         [('ins', 'EUR 200')]],
        ['del_row',
         [('del', 'Mr. Kamil Talar')],
         [('del', '13 shares')],
         [('del', '13%')],
         [('del', 'EUR 200')]],
        ['ins_row',
         [('ins', 'Mr. Kamil Talar')],
         [('ins', '8 shares')],
         [('ins', '8%')],
         [('ins', 'EUR 200')]],
    ], col_widths=[5, 3, 3, 4])
    doc.add_paragraph()

    h2(doc, "2.3  Purchase Consideration")
    plain(doc,
          ('keep', 'The Parties confirm and acknowledge that the purchase consideration of '),
          ('del', 'EUR 25 (twenty five euros) per 1% shareholding package'),
          ('ins', 'EUR 200 (two hundred euros) per 8% shareholding package'),
          ('keep', ' has been received in full by the Company from each of Mr. Marcin Szumiło, '
           'Mr. Andrzej Lechowicz, and Mr. Kamil Talar prior to or simultaneously with the execution '
           'of this Agreement. The payment obligations under Article 2.2 are hereby acknowledged as '
           'fully discharged.'))

    # ── ARTICLE 4 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 4 — ROLES AND RESPONSIBILITIES OF SHAREHOLDERS")

    h2(doc, "4.1  Lighthief International")
    plain(doc, ('keep', 'Lighthief International, as Majority Shareholder, shall:'))
    bul(doc, ('del', 'Hold 75% of the issued share capital'),
        ('ins', 'Hold 76% of the issued share capital'))
    bul(doc, 'Provide strategic direction and oversight of the Company\'s operations;')
    bul(doc, 'Develop and manage sales channels for the EMS under the DISPERON brand;')
    bul(doc, 'Retain in full all Commissioning Fees generated from projects;')
    bul(doc, 'Be responsible for investor relations, brand development, and international expansion;')
    bul(doc, 'Provide corporate administration and registered office infrastructure.')

    h2(doc, "4.2  Mr. Kamil Talar")
    plain(doc, ('keep', 'Mr. Kamil Talar shall:'))
    bul(doc, ('del', 'Hold 13% of the issued share capital'),
        ('ins', 'Hold 8% of the issued share capital'))
    bul(doc, 'Actively develop and manage sales channels for the DISPERON EMS in European markets;')
    bul(doc, 'Deploy his existing business relationships, including his relationship with Linyang '
        'Energy Co., Ltd., to generate BESS project pipeline;')
    bul(doc, 'Report pipeline and sales activity to the Director on a monthly basis.')
    plain(doc,
          ('ins', 'Performance Obligation: Mr. Talar shall, during each calendar year, (a) maintain '
           'a documented active pipeline of not less than 100 MWh of BESS projects incorporating the '
           'EMS Software; and (b) introduce a minimum of 3 qualified customer leads per year. Failure '
           'to meet these targets in any two consecutive annual periods shall entitle the Board to '
           'reclassify his shares as non-voting ordinary shares and suspend dividend rights, subject '
           'to 30 days\' written notice and a 60-day cure period.'))
    lh_note(doc, "Performance KPI restored per SHA v3. This was not in the Voltus draft but is "
            "essential — Talar holds equity and must have measurable obligations.")

    h2(doc, "4.3  Mr. Marcin Szumiło and Mr. Andrzej Lechowicz")
    plain(doc, ('keep', 'Both Mr. Szumiło and Mr. Lechowicz, individually and as Management Board '
                'members of Voltus, shall:'))
    bul(doc, 'Contribute the EMS Software to the Company\'s commercial operations through the Licence;')
    bul(doc, 'Maintain primary responsibility for software development, technical architecture, and '
        'software quality;')
    bul(doc,
        ('del', 'Provide tier-3 technical support'),
        ('ins', 'Provide tier-2 technical support'),
        ('keep', ' to the Company\'s commissioning and operations teams;'))
    bul(doc, 'Support the sales process through technical demonstrations and client meetings;')
    bul(doc, 'Ensure the EMS Software remains compliant with NIS2, GDPR, and grid code standards.')
    plain(doc,
          ('del', '"All activities described above, including technical support, and commercial '
           'support, shall be performed by Voltus Energy Sp. z o.o. on a commercial basis and shall '
           'be subject to separate remuneration agreed between the Parties. Nothing in this Agreement '
           'shall be construed as requiring Voltus Energy Sp. z o.o. or its representatives to provide '
           'services free of charge."'))
    lh_note(doc, "CRITICAL REJECTION: This clause attempts to convert all Voltus obligations in "
            "Art 4.3 into separately billed services. The Licence Fee and Development Bonus already "
            "compensate Voltus for these contributions. Separate billing for every technical support "
            "interaction would make DISPERON commercially unviable. Specific support terms (if any "
            "are to be charged) must be set out in Exhibit D — not as a blanket disclaimer here.")

    # ── ARTICLE 5 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 5 — INTELLECTUAL PROPERTY")
    h2(doc, "5.1 – 5.4  [ACCEPTED UNCHANGED]")
    plain(doc, 'Articles 5.1 through 5.4 accepted as drafted by Voltus.')
    h2(doc, "5.5  IP Warranty and Indemnity  [ACCEPTED UNCHANGED]")
    plain(doc, 'Article 5.5 accepted as drafted — Voltus indemnity clause maintained.')

    # ── ARTICLE 6 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 6 — COMMERCIAL MODEL AND LICENCE FEES")
    h2(doc, "6.1(a)  Licence Fee — Payment Timing")
    plain(doc,
          ('keep', 'The Licence Fee shall be paid by the Company to Voltus '),
          ('del', 'on a proportional basis, corresponding to the advance payments received from the Client.'),
          ('ins', 'on a proportional basis corresponding to milestone payments received from the '
           'Client, provided that the full Licence Fee shall in any event be paid no later than '
           '90 (ninety) days from the commissioning date regardless of Client payment status.'))
    lh_note(doc, "Proportional payment flow accepted as reasonable. However a 90-day longstop from "
            "commissioning is added to protect against indefinite deferral if clients delay payment.")

    h2(doc, "6.2  Milestone Definition — Subclauses (d)–(h) REJECTED")
    plain(doc,
          ('keep', 'The Milestone shall be deemed achieved when aggregate installed capacity reaches '
           '500 MWh. Each project shall be counted upon commissioning evidenced by signed acceptance '
           'certificate. Projects cancelled prior to commissioning shall not count. The Milestone has '
           'no fixed time limit.'))
    plain(doc,
          ('del', '(d) Any reservation, exclusivity, or territorial rights granted to the Company in '
           'connection with the Milestone shall be valid for an initial period of 2 (two) years from '
           'the start date and shall be conditional upon the Company\'s continuous and active commercial '
           'engagement; (e) If the Milestone is not achieved within the initial 2-year period, the '
           'reservation may be extended for successive periods of 2 (two) years, subject to the Company '
           'demonstrating continued active commercial efforts; (f) After the first year, the Company '
           'may request an extension subject to Voltus\' approval; (g) Voltus Energy Sp. z o.o. shall '
           'have the right, at its sole discretion, to refuse or revoke any extension, limit the scope '
           'of reservation, or terminate the Milestone-related rights; (h) Any reservation or '
           'exclusivity granted under this Clause shall not constitute an unconditional or indefinite right.'))
    lh_note(doc, "CRITICAL REJECTION: Subclauses (d)–(h) give Voltus sole discretionary power to "
            "revoke exclusivity and milestone-related rights at any time on the grounds of 'no material "
            "commercial progress' or 'commercially not viable' — both undefined. This converts the "
            "Milestone into a rolling conditional right fully controlled by Voltus. These subclauses "
            "are deleted in their entirety. The Milestone is unconditional and time-unlimited as per "
            "SHA v3. Exclusivity is governed exclusively by Art 8.")

    h2(doc, "6.3  Perpetual Licence upon Achievement of Milestone")
    plain(doc,
          ('keep', 'Upon achievement of the Milestone, Voltus shall automatically grant to the Company '
           'a perpetual, irrevocable, non-exclusive, worldwide, royalty-free licence to use, deploy, '
           'adapt, '),
          ('del', 'the EMS Software in its then-current version'),
          ('ins', 'and sub-licence the EMS Software in its then-current version and all subsequent '
           'versions released by Voltus Energy Sp. z o.o.,'),
          ('keep', ' without any ongoing licence fee obligation.'))
    bul(doc, 'Cover an unlimited number of customer installations worldwide;')
    bul(doc,
        ('del', 'Apply to all versions of the EMS Software existing at the Milestone Date by Voltus '
         'Energy Sp. z o.o.;'),
        ('ins', 'Apply to all versions of the EMS Software existing at the Milestone Date and all '
         'subsequent versions released by Voltus Energy Sp. z o.o.;'))
    bul(doc, 'Require no further payment to Voltus in respect of the use of the EMS Software itself;')
    bul(doc, 'Survive any change of control, dissolution, or restructuring of Voltus, unless directly '
        'caused by a material, uncured breach of this Agreement by the Company.')
    plain(doc,
          ('del', 'Any adaptations, updates, and further developments of the EMS Software shall be '
           'made available to the Company only for so long as the applicable commission-based agreement '
           'remains in force and all payments due under such agreement are duly made in accordance with '
           'its terms.'))
    lh_note(doc, "Two critical restorations: (1) 'and sub-licence' — DISPERON must be able to "
            "sub-licence to customers under the perpetual licence or it has no commercial value. "
            "(2) 'all subsequent versions' — without future versions, a perpetual licence to static "
            "2025-era software is commercially dead within 2-3 years as grid codes evolve. The "
            "bolted-on paragraph conditioning future updates on continued payment is REJECTED — it "
            "converts the perpetual licence into a subscription that can be revoked for any payment "
            "breach.")

    h2(doc, "6.4  Source Code Escrow  [ACCEPTED UNCHANGED]")
    plain(doc, 'Article 6.4 (a)–(e) accepted as drafted by Voltus. Escrow structure maintained.')

    h2(doc, "6.5  Development Bonus (Post-Milestone)")
    plain(doc,
          ('keep', '(a) For each customer licence sold by the Company after the Milestone Date, the '
           'Company shall pay to the Voltus Parties a Development Bonus calculated as:'))
    bul(doc,
        ('del', '30% (thirty percent) for customer licences in countries not subject to exclusivity;'),
        ('ins', '25% (twenty-five percent) for customer licences in countries not subject to exclusivity;'))
    bul(doc,
        ('del', '50% (fifty percent) for customer licences in countries where the Company has been '
         'granted exclusivity or reservation rights.'),
        ('ins', '30% (thirty percent) for customer licences in countries where the Company has been '
         'granted exclusivity.'))
    lh_note(doc, "CRITICAL COUNTER: Lighthief rejects the 50% rate for exclusive territories. "
            "The stated rationale — that exclusivity creates 'risk and opportunity cost' for Voltus — "
            "is inverted. It is DISPERON that deploys capital, builds market presence, and bears "
            "commercial risk to earn exclusivity. Counter-proposal: 25% (non-exclusive) / 30% "
            "(exclusive), reflecting Voltus's legitimate interest in stronger markets without "
            "penalising DISPERON for its own commercial success.")
    plain(doc,
          ('keep', '(b) The Development Bonus rate shall be subject to review every 3 (three) years, '
           'adjustable by written agreement of all Parties within the range of '),
          ('del', '30% (thirty percent) to 50% (fifty percent)'),
          ('ins', '20% (twenty percent) to 30% (thirty percent)'),
          ('keep', ', provided Voltus produces a quarterly development expenditure summary.'))
    plain(doc,
          ('keep', '(c) The Development Bonus is calculated exclusively on the Software Margin component '
           'as defined in Article 1. Yearly software subscription fees and all other licence-related '
           'fees are expressly '),
          ('del', 'included into the Development Bonus calculation.'),
          ('ins', 'subject to the same Development Bonus rate set out in Art 6.5(a), and shall be '
           'calculated on the basis of the applicable Subscription Income definition in Article 1. '
           'For the avoidance of doubt, Commissioning Fees are excluded from the Development Bonus '
           'calculation under all circumstances.'))
    lh_note(doc, "Art 6.5(c): Lighthief accepts that subscription fees are within scope of the "
            "Development Bonus but counters the careless drafting. The original draft said subscription "
            "fees are 'included' without any rate or formula — this is accepted at the same 25%/30% "
            "rate with the Exhibit C formula applied.")
    plain(doc,
          ('keep', '(d) The Development Bonus shall be designated for ongoing research, development, '
           'and enhancement of the EMS Software. The Voltus Parties shall provide a quarterly summary '
           'of development activities funded by such payments.'))
    plain(doc,
          ('keep', '(e) The Development Bonus shall be payable within '),
          ('del', '14 (fourteen) days'),
          ('ins', '30 (thirty) days'),
          ('keep', ' of the Company receiving cleared payment from the relevant customer'),
          ('del', ' including any partial payments and prepayments.'),
          ('ins', '. Prepayments and partial payments from clients shall not trigger the Development '
           'Bonus until the relevant milestone or commissioning event is formally completed and '
           'confirmed in writing.'))
    lh_note(doc, "Art 6.5(e): 14 days is commercially unreasonable — DISPERON needs time to "
            "process payments and verify completion. 30 days restored. The prepayment / partial "
            "payment trigger is REJECTED — paying Voltus a bonus on a client advance before "
            "commissioning is complete would incentivise Voltus to take the money and deprioritise "
            "delivery.")

    # ── ARTICLE 7 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 7 — OPERATIONAL STRUCTURE")
    h2(doc, "7.2  Technology and Product Development — Voltus Obligations")
    plain(doc, ('keep', 'Voltus Energy Sp. z o.o. shall be primarily responsible for the continuous '
                'development, maintenance, security, and improvement of the EMS Software. '
                'Specifically, Voltus shall:'))
    bul(doc, 'Respond to severity-classified support requests within the timeframes set out in '
        'Exhibit D;')
    bul(doc,
        ('ins', 'Maintain cloud-hosted EMS components at a minimum availability of 99.5% measured '
         'per calendar month, excluding planned maintenance windows notified at least 5 days in advance;'))
    lh_note(doc, "Uptime obligation (99.5%) restored to Art 7.2 body text. Voltus removed this from "
            "the article; it must appear here as the primary contractual obligation, not only in Exhibit D.")
    bul(doc,
        ('keep', 'Implement changes required by new or amended EU or national grid code regulations '
         'within 90 (ninety) days of the official publication date '),
        ('del', 'only whenever applying such requirements are technically possible; and Company shall '
         'notify Voltus within 14 (fourteen) days of becoming aware of any such regulatory change.'),
        ('ins', 'of such requirements. Voltus shall notify the Company within 14 (fourteen) days of '
         'becoming aware of any regulatory change affecting the EMS Software.'))
    lh_note(doc, "Two issues corrected: (1) The escape clause 'only whenever technically possible' "
            "is deleted — this is a self-assessed exemption that could justify any delay. If a "
            "regulatory requirement is technically infeasible, the parties must agree a solution, not "
            "Voltus unilaterally decide. (2) Notification obligation restored to Voltus (not flipped "
            "to DISPERON) — Voltus as the software developer is better positioned to monitor EU grid "
            "code changes.")
    bul(doc, 'Provide the Company with a quarterly product roadmap and development activity report;')
    bul(doc, 'Maintain cybersecurity compliance with NIS2 and applicable EU standards.')
    plain(doc,
          ('del', 'The Company (Disperon) shall proactively identify, document, and submit to Voltus '
           'any new functional, technical, or regulatory requirements arising from market conditions '
           'or customer needs.'))
    lh_note(doc, "This sentence is REJECTED. It attempts to shift Voltus's regulatory compliance "
            "monitoring obligation onto DISPERON. DISPERON will reasonably cooperate and share market "
            "intelligence, but cannot be contractually obligated to do Voltus's product development "
            "work.")

    h2(doc, "7.3  Commissioning and Implementation")
    plain(doc,
          ('keep', 'On-site and remote commissioning, system integration, and project handover services '
           'shall be performed by the Company or a designated entity within the Lighthief International '
           'group. Voltus shall provide technical support, documentation, and training as reasonably '
           'required'),
          ('del', ' and shall be compensated based on agreed hourly labour rates or as otherwise '
           'mutually agreed in writing between the Parties'),
          ('keep', '. All Commissioning revenues belong to the Company.'))
    lh_note(doc, "Voltus's separate compensation clause for technical support during commissioning is "
            "REJECTED. Tier-2 technical support during commissioning is already within the scope of "
            "Voltus's obligations under Art 4.3 and compensated through the Licence Fee / Development "
            "Bonus. If specific advisory engagements are desired at day rates, these can be separately "
            "agreed in writing on a case-by-case basis — they should not be a blanket right.")

    h2(doc, "7.4  Customer Contracts")
    plain(doc,
          ('del', 'After passing and during maintaining exclusivity period described in Article 8, '),
          ('ins', 'All '),
          ('keep', 'customer contracts shall be entered into by the Company. The Company is the legal '
           'counterparty to all end customers. Voltus Energy Sp. z o.o. acts as software supplier to '
           'the Company, not to end customers directly, unless otherwise agreed in writing.'))
    lh_note(doc, "The conditional 'after passing exclusivity' qualifier is REJECTED. DISPERON is the "
            "contracting entity from day one of operations — not only after earning exclusivity. "
            "Restricting this right to the post-exclusivity period would prevent DISPERON from "
            "contracting with customers at all in the early phase.")

    # ── ARTICLE 8 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 8 — EXCLUSIVITY AND MARKET MODEL")
    h2(doc, "8.3  Exclusivity Volume Thresholds")
    plain(doc, 'Exclusivity shall be granted upon the Company achieving the following minimum '
          'installed capacity within 24 months of the Company\'s first commercial project in that country:')
    tbl_redline(doc, [
        ["Market Category", "Countries", "Minimum Volume", "Annual Maintenance", "Lapse After"],
        ['del_row',
         'Large markets', 'Germany, Italy, Spain, France',
         '500 MWh / 24 months', '200 MWh / year', '1 year failure'],
        ['ins_row',
         'Large markets', 'Germany, Italy, Spain, France',
         '500 MWh / 24 months', [('ins', '100 MWh / year')], [('ins', '2 consecutive years')]],
        ['del_row',
         'Mid-sized markets', 'Poland*, Netherlands, Romania, Greece',
         '250 MWh / 24 months', '100 MWh / year', '1 year failure'],
        ['ins_row',
         'Mid-sized markets', [('ins', 'Netherlands, Romania, Greece')],
         '250 MWh / 24 months', [('ins', '50 MWh / year')], [('ins', '2 consecutive years')]],
        ['del_row',
         'Smaller markets', 'Cyprus, other EU states',
         '100 MWh / 24 months', '40 MWh / year', '1 year failure'],
        ['ins_row',
         'Smaller markets', 'Cyprus, other EU states',
         '100 MWh / 24 months', [('ins', '20 MWh / year')], [('ins', '2 consecutive years')]],
    ], col_widths=[3.2, 3.5, 3.5, 3, 2.8])
    doc.add_paragraph()
    lh_note(doc, "Two counters: (1) Annual maintenance thresholds halved back to SHA v3 levels — "
            "Voltus's doubled thresholds (200/100/40 MWh/year) are unachievable for a new market "
            "entrant and designed to ensure rapid lapse. (2) Lapse trigger restored to TWO consecutive "
            "annual failures (not one). A single bad year (e.g. due to permitting delays, grid code "
            "changes, or macroeconomic factors) should not strip exclusivity. *Poland separately "
            "addressed in Art 8.5.")

    h2(doc, "8.5  Initial Country Limit")
    plain(doc,
          ('keep', 'At the time of execution of this Agreement, the Company may seek exclusivity in '
           'a maximum of 5 (five) countries simultaneously. Additional countries may be added as '
           'exclusivity is achieved and maintained. '),
          ('del', 'Poland and Baltics shall be excluded from list of exclusivity.'),
          ('ins', '[Poland and the Baltic states (Lithuania, Latvia, Estonia) shall be subject to a '
           '12-month exclusivity moratorium from the date of this Agreement, during which Voltus may '
           'continue existing commercial activities. At the end of the moratorium, DISPERON may apply '
           'for exclusivity in those countries under the standard terms of Article 8.3.]'))
    lh_note(doc, "Voltus's blanket permanent exclusion of Poland and the Baltics is rejected. Voltus "
            "may legitimately protect its home market during a transition period but cannot permanently "
            "carve out major EU markets. Counter: 12-month moratorium, then standard exclusivity "
            "process applies. This protects Voltus's existing Polish pipeline while giving DISPERON a "
            "path to exclusivity.")

    h2(doc, "8.6  Voltus Direct Rights — BESS Manufacturer Partners")
    plain(doc,
          ('del', 'Notwithstanding any exclusivity granted to the Company, Voltus Energy Sp. z o.o. '
           'retains the unrestricted right to supply the EMS Software directly to any BESS manufacturers. '
           'For the avoidance of doubt: This right shall not be subject to any exclusivity, threshold, '
           'or geographic restriction in respect of BESS manufacturers; Business partners (including '
           'but not limited to integrators, distributors, or resellers) with whom Voltus may cooperate '
           'directly shall be those listed in Exhibit B (the "Listed Partners"); Listed Partners '
           'included in Exhibit B as of the date of execution shall not be subject to any territorial '
           'or exclusivity restrictions; New business partners may be added to Exhibit B from time to '
           'time by mutual written agreement of the Parties; In countries where exclusivity has not '
           'yet been granted, Voltus shall be entitled to cooperate with all Listed Partners without '
           'restriction.'),
          ('ins', 'Notwithstanding any exclusivity granted to the Company, Voltus Energy Sp. z o.o. '
           'retains the right to supply the EMS Software directly to BESS Manufacturer Partners '
           'expressly listed in Exhibit B of this Agreement. This right: (i) applies only to the '
           'manufacturers and partners listed in Exhibit B; (ii) does not extend to any manufacturer '
           'or partner not listed in Exhibit B without written consent of all Parties; (iii) is not '
           'subject to geographic restriction in respect of listed Exhibit B entities only.'))
    plain(doc,
          ('keep', 'BESS Manufacturer Partner Referral Obligation: Where a project originating from '
           'a BESS Manufacturer Partner (including Linyang Energy Co., Ltd.) has been the subject of '
           'prior commercial engagement by the Company — evidenced by a written proposal, meeting '
           'record, or registered customer entry — Voltus shall '),
          ('del', '(a) notify the Company within 5 (five) business days of becoming aware of a '
           'specific commercial opportunity relating to such project; and (b) use reasonable efforts, '
           'acting in good faith, to cooperate with the Company in evaluating whether such opportunity '
           'may be pursued jointly under the DISPERON brand.'),
          ('ins', 'notify the Company within 5 (five) business days and shall refer such project to '
           'the Company under the DISPERON brand. The Parties shall agree commercial terms within 15 '
           'business days of such referral. In the absence of agreement within that period, the '
           'standard Development Bonus shall apply.'))
    plain(doc,
          ('del', 'The Parties shall use good faith efforts to agree on the applicable commercial '
           'structure within 15 business days. In the absence of agreement within this period, Voltus '
           'shall be entitled, at its sole discretion, to proceed independently with such opportunity, '
           'subject to the applicable Development Bonus provisions set out in Article 6.5.'))
    lh_note(doc, "CRITICAL REJECTION: Voltus's Art 8.6 gives them unrestricted rights to supply ANY "
            "BESS manufacturer — this is not a carve-out, it is a wholesale elimination of exclusivity. "
            "The Exhibit B-bounded right is restored. On the referral obligation: the weakened 'use "
            "reasonable efforts to evaluate' language (which then allows Voltus to proceed independently) "
            "is rejected and replaced with a firm referral obligation with a Development Bonus default "
            "as the backstop.")

    h2(doc, "8.7  Transparency and Customer Registration  [ACCEPTED WITH MINOR EDIT]")
    plain(doc, 'Article 8.7 accepted as drafted by Voltus. The expanded cooperation and good-faith '
          'framework is reasonable and more detailed than SHA v3.')

    h2(doc, "8.8  Non-Competition Undertaking of Voltus")
    plain(doc,
          ('keep', 'Voltus Energy Sp. z o.o. undertakes, for the duration of this Agreement, that: '
           '(a) it shall maintain full transparency with the Company regarding its direct sales '
           'activities within the Territory; (b) if a customer approaches Voltus directly for EMS '
           'services, and that same customer has previously been approached by or is in active '
           'discussion with the Company, Voltus shall decline to engage independently and shall refer '
           'such customer to the Company under the DISPERON brand. '),
          ('del', 'In case when even after referring customer to company such customer is insisting '
           'for Voltus to deliver services to him. Voltus has a right to do so. To avoid price '
           'competition Voltus will offer the same scope of delivery with price 10% higher that '
           'price offered by Company.'),
          ('ins', 'The referral obligation is unconditional and does not lapse if the customer '
           'subsequently approaches Voltus directly. Voltus shall maintain the referral regardless '
           'of customer preference.'))
    plain(doc,
          ('del', '"The above undertaking shall not apply to: (i) customers independently introduced '
           'by Voltus Energy Sp. z o.o.; and (ii) customers of BESS manufacturers, which shall remain '
           'unrestricted.'),
          ('ins', 'The above undertaking shall not apply to customers originally introduced by Voltus '
           'Energy Sp. z o.o. independently of the Company. It shall apply to all BESS Manufacturer '
           'Partner customers where DISPERON has prior commercial engagement, pursuant to Art 8.6.'))
    lh_note(doc, "CRITICAL REJECTION: The 'customer insisting' escape clause (Voltus can serve a "
            "referred customer directly at just +10% higher price) destroys the non-compete entirely. "
            "Any client who prefers Voltus for any reason can trigger this exception. The +10% premium "
            "is trivial on large BESS contracts. This is deleted. The non-compete on BESS manufacturer "
            "customers is also restored to be consistent with the restored Art 8.6.")

    # ── ARTICLE 9 — [ACCEPTED] ────────────────────────────────────────────────
    h1(doc, "ARTICLE 9 — TRANSFER OF SHARES  [ACCEPTED UNCHANGED]")
    plain(doc, 'Articles 9.1–9.4 (lock-up, ROFR, drag-along, tag-along) accepted as drafted by Voltus.')

    # ── ARTICLE 10 ────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 10 — GOVERNANCE AND BOARD OF DIRECTORS")

    h2(doc, "10.2  Board Composition")
    plain(doc,
          ('keep', '(a) Lighthief International shall have the right to appoint up to 3 (three) '
           'Directors to the Board. (b) The Shareholders other than Lighthief International shall '
           'have the right to appoint '),
          ('del', 'maximum 2 (two) Directors'),
          ('ins', '1 (one) Director'),
          ('keep', ' from among themselves to the Board.'))
    lh_note(doc, "Voltus increased minority board seats from 1 to 2. Lighthief counters at 1. "
            "With 24% combined minority shareholding, 1 director is proportionate. Two minority "
            "directors would create a 3:2 Lighthief:minority split, approaching an effective veto "
            "on operational decisions that require only board resolution.")

    h2(doc, "10.4  Decisions Requiring Majority Shareholder Consent")
    plain(doc,
          ('del', 'Only: entry into related-party transactions >EUR 50,000; and appointment/removal '
           'of auditors.'),
          ('ins', 'The following matters shall require the written consent of Shareholders holding a '
           'majority of Shares: (a) amendment of the Company\'s Articles of Association; (b) issuance '
           'of new shares or share capital increases; (c) material change to the Company\'s principal '
           'business activity; (d) entry into any related-party transaction exceeding EUR 50,000 in '
           'value; (e) appointment or removal of auditors.'))
    lh_note(doc, "Voltus stripped Art 10.4 to only 2 items, moving share issuance, Articles amendment, "
            "and business change to the 85% threshold. This makes those decisions harder for Lighthief "
            "to pass (85% vs. 51%). Restored to original 5-item list at simple majority — these are "
            "standard majority-vote governance matters.")

    h2(doc, "10.5  Reserved Matters Requiring 85% Shareholder Consent  [ACCEPTED]")
    plain(doc, 'Article 10.5 (85% reserved matters including commercial model, dissolution, and '
          'perpetual licence) accepted as drafted by Voltus.')

    # ── ARTICLES 11 / 12 ─────────────────────────────────────────────────────
    h1(doc, "ARTICLES 11 & 12 — CONFIDENTIALITY / GOVERNING LAW  [ACCEPTED UNCHANGED]")
    plain(doc, 'Articles 11 and 12 accepted as drafted by Voltus.')

    # ── ARTICLE 13 ────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 13 — TERM AND TERMINATION  [NEW ARTICLE BY VOLTUS — SUBSTANTIALLY REJECTED]")
    lh_note(doc, "CRITICAL: Voltus inserted an entirely new Article 13. Lighthief's position on each "
            "sub-article is set out below.")

    h2(doc, "13.1  Term of the Agreement")
    plain(doc,
          ('del', 'This Agreement shall be concluded for an initial period of 2 (two) years from the '
           'date of execution ("Initial Term"). Upon expiry of the Initial Term, the Agreement may be '
           'extended for successive periods of 2 (two) years, subject to mutual agreement of the '
           'Parties and satisfactory commercial performance.'),
          ('ins', 'This Agreement shall be concluded for an initial period of 5 (five) years from the '
           'date of execution ("Initial Term"). Upon expiry of the Initial Term, the Agreement shall '
           'automatically renew for successive periods of 2 (two) years unless either Party gives '
           'not less than 6 (six) months\' written notice of non-renewal prior to expiry of the '
           'then-current term.'))
    lh_note(doc, "A fixed term is accepted in principle. 2 years is too short — insufficient time to "
            "reach the 500 MWh Milestone and build a market position. Counter: 5-year initial term "
            "with automatic renewal. Renewal should be automatic (not subject to 'mutual agreement'), "
            "with a notice period for non-renewal.")

    h2(doc, "13.2  Volume-Based Termination  [REJECTED]")
    plain(doc,
          ('del', 'Voltus Energy Sp. z o.o. shall have the right to terminate this Agreement upon '
           'written notice if, within 2 (two) years: (a) the total volume of completed and '
           'commissioned projects is less than 200 MWh; and/or (b) the total volume of signed '
           'contracts is less than 500 MWh. The Company shall be granted a cure period of 6 (six) '
           'months from receipt of such notice.'))
    lh_note(doc, "REJECTED IN FULL: This provision is calibrated to fail. The 'and/or' means that "
            "even if DISPERON has EUR 500 MWh in contracts but has commissioned 199 MWh (due to "
            "permitting or grid delays outside DISPERON's control), Voltus can terminate. These "
            "thresholds treat commissioning delays — which are primarily driven by grid operators and "
            "permitting authorities — as DISPERON's breach. Delete entirely.")

    h2(doc, "13.3  Revenue-Based Termination  [REJECTED]")
    plain(doc,
          ('del', 'Voltus Energy Sp. z o.o. shall have the right to terminate this Agreement upon '
           'written notice if the annual value of Licence Fees, Subscription Fees, and Development '
           'Bonuses payable to Voltus does not meet: (a) during the first year — at least EUR '
           '1,000,000; (b) during each subsequent year — at least EUR 2,000,000 per year. A cure '
           'period of 6 (six) months shall apply.'))
    lh_note(doc, "REJECTED IN FULL: EUR 2M/year to Voltus requires ~EUR 6.7M in annual software "
            "revenue at a 30% rate — roughly 1,000+ MWh/year deployed. This is impossible to "
            "achieve in year 1–2. These are not targets; they are pre-set termination grounds. "
            "If minimum guaranteed revenue is desired, this must be separately negotiated as a "
            "Minimum Annual Guarantee (MAG) with appropriate ramp-up and market conditions "
            "adjustments — not as a unilateral termination right.")

    h2(doc, "13.4  Performance and Conduct-Based Non-Renewal  [REJECTED]")
    plain(doc,
          ('del', 'Voltus Energy Sp. z o.o. shall have the right not to extend this Agreement upon '
           'written notice if: (a) the Company fails to maintain a reasonable standard of commercial '
           'cooperation; (b) the Company engages in actions that are materially non-competitive or '
           'detrimental to the interests of Voltus; (c) repeated or substantiated complaints are '
           'received from customers; (d) the Company fails to act in good faith.'))
    lh_note(doc, "REJECTED IN FULL: This is a catch-all subjective non-renewal right based on "
            "undefined standards ('reasonable cooperation', 'good faith'). These are standard "
            "grounds for dispute resolution — not termination triggers. Any conduct-based concern "
            "must go through the Art 12 negotiation and arbitration process. If Voltus wants "
            "objective conduct standards, these should be defined with specific measurable criteria "
            "and a minimum 90-day cure period.")

    # ── ARTICLE 14 ────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 14 — GENERAL PROVISIONS  [ACCEPTED UNCHANGED]")
    plain(doc, 'Articles 14.1–14.6 (entire agreement, amendment, severability, waiver, notices, '
          'electronic signatures) accepted as drafted by Voltus.')

    # ── EXHIBIT A ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "EXHIBIT A — SHARE CERTIFICATES")
    plain(doc,
          ('del', 'Note: Exhibit A lists Szumiło 8%, Lechowicz 8%, Talar 8% but Article 2.2 body '
           'text (Voltus version) shows Szumiło 6%, Lechowicz 6%, Talar 13%. These are irreconcilably '
           'inconsistent.'),
          ('ins', 'Exhibit A shall be updated to reflect the agreed share structure: Lighthief '
           'International 76%, Mr. Szumiło 8%, Mr. Lechowicz 8%, Mr. Talar 8%, consistent with '
           'the restored Art 2.2.'))
    lh_note(doc, "CRITICAL: The Voltus draft contains an internal inconsistency — body text and "
            "Exhibit A show different shareholdings. This document cannot be signed until both are "
            "aligned. Lighthief's position: restore 76/8/8/8 throughout.")

    # ── EXHIBIT B ─────────────────────────────────────────────────────────────
    h1(doc, "EXHIBIT B — BESS MANUFACTURER PARTNERS")
    plain(doc,
          ('del', 'Exhibit header changed by Voltus to "CUSTOMER PARTNERS."'),
          ('ins', 'Exhibit B heading restored to "BESS MANUFACTURER PARTNERS" — consistent with '
           'the restored Art 1 definition.'))
    plain(doc, 'Linyang Energy Co., Ltd. — accepted as listed. Remaining rows to be completed '
          'by mutual agreement before signing.')
    lh_note(doc, "Voltus renamed Exhibit B to 'Customer Partners' as part of the strategy to split "
            "'BESS Manufacturers' (unrestricted) from 'Partners' (Exhibit B). With the Art 1 "
            "definition restored to a single 'BESS Manufacturer Partners' list, the original heading "
            "is restored.")

    # ── EXHIBITS C & D ────────────────────────────────────────────────────────
    h1(doc, "EXHIBITS C & D — LICENCE FEE SCHEDULE / SOFTWARE SUPPORT ANNEX  [ACCEPTED UNCHANGED]")
    plain(doc, 'Exhibit C (Licence Fee Schedule) and Exhibit D (Software Support Annex) accepted '
          'as drafted by Voltus, subject to the per-MWh rates being agreed and inserted within '
          '14 days of execution.')

    # ── CLOSING NOTE ─────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "LIGHTHIEF COUNTER-PROPOSAL SUMMARY")

    summary_data = [
        ["Article", "Voltus Change", "Lighthief Position", "Accept / Counter / Reject"],
        ["Art 1 — EMS definition",
         "Source code, updates removed",
         "Restore source code, object code, updates, derivative works",
         "COUNTER"],
        ["Art 1 — Dev Bonus def.",
         "30%",
         "Restore 25%",
         "COUNTER"],
        ["Art 1 — BESS Mfr def.",
         "Split: any mfr unrestricted",
         "Restore single Exhibit B list",
         "REJECT"],
        ["Art 2.2 — Shares",
         "75/6/6/13 + Exhibit A wrong",
         "Restore 76/8/8/8; fix Exhibit A",
         "REJECT"],
        ["Art 2.3 — Consideration",
         "EUR 25/1%",
         "Restore EUR 200/8%",
         "COUNTER"],
        ["Art 4.2 — Talar KPI",
         "Removed",
         "Restore 100 MWh pipeline / 3 leads/year",
         "RESTORE"],
        ["Art 4.3 — Voltus billing",
         "All services separately billed",
         "Delete clause; services included",
         "REJECT"],
        ["Art 4.3 — Tier support",
         "Tier-3",
         "Restore tier-2",
         "COUNTER"],
        ["Art 6.1(a) — Payment",
         "Proportional to client payments",
         "Accept + 90-day longstop",
         "ACCEPT+"],
        ["Art 6.2(d)–(h) — Milestone",
         "Voltus can revoke exclusivity",
         "Delete subclauses (d)–(h)",
         "REJECT"],
        ["Art 6.3 — Perp. licence",
         "No sub-licence; no future versions",
         "Restore sub-licence + future versions",
         "REJECT"],
        ["Art 6.3 — Updates clause",
         "Updates conditioned on payments",
         "Delete paragraph",
         "REJECT"],
        ["Art 6.5(a) — Bonus rates",
         "30% / 50% sliding scale",
         "25% (non-excl) / 30% (excl)",
         "COUNTER"],
        ["Art 6.5(b) — Review range",
         "30%–50%",
         "20%–30%",
         "COUNTER"],
        ["Art 6.5(c) — Subscriptions",
         "Included, no formula",
         "Included at same rate, Exhibit C formula",
         "ACCEPT+"],
        ["Art 6.5(e) — Payment days",
         "14 days incl. prepayments",
         "30 days; no prepayment trigger",
         "COUNTER"],
        ["Art 7.2 — Uptime",
         "Removed from Art 7.2",
         "Restore 99.5% in Art 7.2",
         "RESTORE"],
        ["Art 7.2 — Regulatory",
         "Escape clause; Company monitors",
         "No escape; Voltus monitors",
         "COUNTER"],
        ["Art 7.2 — Reqs burden",
         "DISPERON must submit requirements",
         "Delete",
         "REJECT"],
        ["Art 7.3 — Support billing",
         "Voltus charges for support",
         "Delete; support included",
         "REJECT"],
        ["Art 7.4 — Contracts",
         "Only during exclusivity period",
         "Always — from day 1",
         "REJECT"],
        ["Art 8.3 — Maint. thresholds",
         "Doubled; 1-year lapse",
         "SHA v3 levels; 2-year lapse",
         "COUNTER"],
        ["Art 8.5 — Poland / Baltics",
         "Permanently excluded",
         "12-month moratorium only",
         "COUNTER"],
        ["Art 8.6 — BESS Mfr rights",
         "Unrestricted to any mfr",
         "Restore Exhibit B limit",
         "REJECT"],
        ["Art 8.6 — Referral",
         "'Reasonable efforts'; Voltus can proceed",
         "Mandatory referral; Dev Bonus default",
         "COUNTER"],
        ["Art 8.8 — Customer insisting",
         "Voltus serves at +10%",
         "Delete escape clause",
         "REJECT"],
        ["Art 10.2 — Board seats",
         "Minority gets 2 seats",
         "Restore 1 minority seat",
         "COUNTER"],
        ["Art 10.4 — Majority votes",
         "Stripped to 2 items",
         "Restore 5-item list",
         "COUNTER"],
        ["Art 13.1 — Term",
         "2 years",
         "5 years + auto-renewal",
         "COUNTER"],
        ["Art 13.2 — Volume termination",
         "Voltus can terminate <200/500 MWh",
         "Delete entirely",
         "REJECT"],
        ["Art 13.3 — Revenue termination",
         "Voltus can terminate <EUR 1M/2M",
         "Delete entirely",
         "REJECT"],
        ["Art 13.4 — Conduct non-renewal",
         "Subjective non-renewal right",
         "Delete; disputes via Art 12",
         "REJECT"],
        ["Exhibit A",
         "Not updated — inconsistent shares",
         "Must align with restored Art 2.2",
         "FIX REQUIRED"],
        ["Exhibit B",
         "Renamed 'Customer Partners'",
         "Restore 'BESS Manufacturer Partners'",
         "COUNTER"],
        ["Exhibits C & D",
         "Accepted",
         "Accepted",
         "ACCEPT"],
    ]

    colours = {"ACCEPT": "E8F5E9", "ACCEPT+": "C8E6C9", "COUNTER": "FFF9C4",
               "REJECT": "FFCDD2", "RESTORE": "E3F2FD", "FIX REQUIRED": "FFE0B2"}

    st = doc.add_table(rows=len(summary_data), cols=4)
    st.style = 'Table Grid'
    for i, row in enumerate(summary_data):
        verdict = row[3] if i > 0 else ""
        bg = colours.get(verdict, "FFFFFF") if i > 0 else "1F497D"
        for j, val in enumerate(row):
            c = st.rows[i].cells[j]
            _set_cell_bg(c, bg)
            p = c.paragraphs[0]; p.clear()
            run = p.add_run(val)
            run.font.size = Pt(8)
            run.font.bold = (i == 0)
            if i == 0:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif verdict == "REJECT":
                run.font.color.rgb = RED
            elif verdict in ("COUNTER", "RESTORE"):
                run.font.color.rgb = LH_BLU
            elif verdict == "FIX REQUIRED":
                run.font.color.rgb = AMBER
            else:
                run.font.color.rgb = GREEN
    for ci, w in enumerate([2.5, 4.5, 5, 2.5]):
        for row in st.rows:
            row.cells[ci].width = Cm(w)

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = closing.add_run(
        "Lighthief Counter-Redline — Prepared April 2026\n"
        "This document is prepared for negotiation purposes. Not for circulation without Lighthief authorisation."
    )
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = GREY

    out = "/Volumes/T7 Grey/solinvest/DisperonEMS/docs/contract/DISPERON_SHA_Redline_Lighthief.docx"
    doc.save(out)
    print(f"Saved: {out}")

build()
