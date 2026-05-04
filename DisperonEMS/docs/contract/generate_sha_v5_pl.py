"""
DISPERON_SHA_v5_PL.docx
Umowa Udziałowców i Licencja Własności Intelektualnej — Wersja 5
Polskie tłumaczenie referencyjne (wersja angielska jest wiążąca)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE  = RGBColor(0x1F, 0x49, 0x7D)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x50, 0x50, 0x50)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _cell_bg(cell, hex_col):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_col)
    tcPr.append(shd)


def setup(doc):
    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(10)
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(3.0)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = BLUE; r.font.bold = True; r.font.size = Pt(13)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = BLUE; r.font.bold = True; r.font.size = Pt(11)


def para(doc, text, bold=False, italic=False, size=10, indent=0, centre=False):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    if centre: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.bold = bold
    run.font.italic = italic; run.font.color.rgb = BLACK
    return p


def bul(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(text); r.font.size = Pt(10); r.font.color.rgb = BLACK


def tbl(doc, data, col_widths=None):
    cols = len(data[0])
    t = doc.add_table(rows=len(data), cols=cols)
    t.style = 'Table Grid'
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            if i == 0: _cell_bg(c, '1F497D')
            p = c.paragraphs[0]; p.clear()
            r = p.add_run(val)
            r.font.size = Pt(9); r.font.bold = (i == 0)
            r.font.color.rgb = WHITE if i == 0 else BLACK
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                if ci < len(row.cells): row.cells[ci].width = Cm(w)
    return t


def sig_block(doc, party_name, role, extra=None, ref=None):
    p = doc.add_paragraph()
    r1 = p.add_run(party_name + "\n"); r1.font.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(role + "\n"); r2.font.size = Pt(10)
    if extra:
        re = p.add_run(extra + "\n"); re.font.size = Pt(10)
    if ref:
        rr = p.add_run(f"Nr ref.: {ref}\n"); rr.font.size = Pt(9); rr.font.italic = True
    p.add_run("\nPodpis: _________________________________\n\n")
    p.add_run("Imię i nazwisko (drukowane): _____________\n\n")
    p.add_run("Data: ____________________________________\n\n")
    p.add_run("Miejsce: _________________________________")
    doc.add_paragraph()


def build():
    doc = Document()
    setup(doc)

    # ── STRONA TYTUŁOWA ───────────────────────────────────────────────────────
    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UMOWA UDZIAŁOWCÓW\nI LICENCJA WŁASNOŚCI INTELEKTUALNEJ")
    r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = BLUE

    doc.add_paragraph()
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "LIGHTHIEF EUBESS LTD\nNr rejestracyjny HE 474192\n"
        "Lophitis Business Center I, piętro 2, biuro 1\n"
        "28 Oktovriou & Aemiliou Chourmouziou, 3035 Limassol, Cypr"
    )
    r2.font.size = Pt(12)

    doc.add_paragraph(); doc.add_paragraph()
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("WERSJA 5  ·  Data: ______________, 2025")
    r3.font.size = Pt(11); r3.font.italic = True

    doc.add_paragraph()
    p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(
        "TŁUMACZENIE REFERENCYJNE — Niniejszy dokument stanowi polskie tłumaczenie\n"
        "Umowy Udziałowców i Licencji Własności Intelektualnej w wersji 5.\n"
        "W przypadku rozbieżności wiążąca jest wersja angielska (DISPERON_SHA_v5_EN.docx)."
    )
    r4.font.size = Pt(9); r4.font.italic = True; r4.font.color.rgb = GREY

    doc.add_page_break()

    # ── PREAMBUŁA ─────────────────────────────────────────────────────────────
    h1(doc, "PREAMBUŁA")
    para(doc,
         'Niniejsza Umowa Udziałowców i Licencja Własności Intelektualnej (dalej: "Umowa") '
         'zostaje zawarta w dniu ostatniego podpisu wskazanego poniżej, pomiędzy następującymi stronami:')
    doc.add_paragraph()

    parties = [
        ("STRONA 1 — LIGHTHIEF INTERNATIONAL",
         'Spółka należycie zarejestrowana i działająca zgodnie z prawem właściwym, reprezentowana przez '
         'Dyrektora, Pana Arkadiusza Sybarisa ("Lighthief International" lub "Udziałowiec Większościowy").'),
        ("STRONA 2 — LIGHTHIEF EUBESS LTD",
         'Spółka z ograniczoną odpowiedzialnością zarejestrowana na Cyprze pod numerem HE 474192, '
         'z siedzibą w Lophitis Business Center I, piętro 2, biuro 1, 28 Oktovriou & Aemiliou '
         'Chourmouziou, 3035 Limassol, Cypr, reprezentowana przez Dyrektora Pana Arkadiusza Sybarisa '
         '(dalej: "Spółka").'),
        ("STRONA 3 — PAN MARCIN SZUMIŁO",
         'Osoba fizyczna, posiadacz polskiego dokumentu tożsamości, zamieszkały pod adresem '
         'wskazanym Spółce ("Pan Szumiło").'),
        ("STRONA 4 — PAN ANDRZEJ LECHOWICZ",
         'Osoba fizyczna, posiadacz polskiego dokumentu tożsamości, zamieszkały pod adresem '
         'wskazanym Spółce ("Pan Lechowicz").'),
        ("STRONA 5 — PAN KAMIL TALAR",
         'Osoba fizyczna, posiadacz polskiego dokumentu tożsamości, zamieszkały pod adresem '
         'wskazanym Spółce ("Pan Talar").'),
        ("STRONA 6 — PAN KOSTAS ALEXANDER PAPACOSTA",
         'Osoba fizyczna, posiadacz cypryjskiego dokumentu tożsamości, rezydent podatkowy Republiki '
         'Cypru, na stałe zamieszkały na Cyprze, pełniący funkcję Dyrektora Spółki '
         '("Pan Papacosta" lub "Dyrektor Cypryjski").'),
        ("STRONA 7 — VOLTUS ENERGY SP. Z O.O.",
         'Spółka z ograniczoną odpowiedzialnością z siedzibą w Gdańskim Parku Naukowo-Technologicznym, '
         'Budynek B, Lokal 2.10.3, ul. Trzy Lipy 3, 80-172 Gdańsk, Polska, NIP: 1990133260, '
         'reprezentowana przez członków Zarządu: Pana Marcina Szumiłę i Pana Andrzeja Lechowicza '
         '("Voltus").'),
    ]
    for name, desc in parties:
        p = doc.add_paragraph()
        p.add_run(name + "\n").font.bold = True
        p.runs[0].font.size = Pt(10)
        r2 = p.add_run(desc); r2.font.size = Pt(10); r2.font.color.rgb = BLACK
        doc.add_paragraph()

    para(doc,
         'Lighthief International, Pan Szumiło, Pan Lechowicz, Pan Talar i Pan Papacosta są zwani '
         'indywidualnie "Udziałowcem", a łącznie "Udziałowcami". Pan Szumiło i Pan Lechowicz '
         'wraz z Voltus Energy Sp. z o.o. są zwani łącznie "Stronami Voltusa". Strony 1–7 są '
         'zwane łącznie "Stronami".')

    # ── ARTYKUŁ 1 — DEFINICJE ─────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "ARTYKUŁ 1 — DEFINICJE")
    para(doc, "Na potrzeby niniejszej Umowy poniższe pojęcia mają następujące znaczenie:")
    doc.add_paragraph()

    defs = [
        ('"Spółka"', 'oznacza Lighthief EUBESS Ltd, HE 474192, opisaną w Preambule.'),
        ('"Udziały"', 'oznacza udziały zwykłe w kapitale zakładowym Spółki.'),
        ('"Oprogramowanie EMS" lub "EMS"',
         'oznacza platformę systemu zarządzania energią i oprogramowanie SCADA opracowane przez '
         'Voltus Energy Sp. z o.o., wraz z powiązanym kodem źródłowym, kodem wynikowym, '
         'dokumentacją, algorytmami, interfejsami, konfiguracjami, aktualizacjami, ulepszeniami '
         'i pracami pochodnymi, aktualnie oferowaną pod marką Energy Copilot.'),
        ('"DISPERON"',
         'oznacza komercyjną nazwę marki i znak towarowy, pod którym Spółka sprzedaje usługi EMS.'),
        ('"Licencja"', 'oznacza licencję na oprogramowanie udzieloną przez Voltus Spółce zgodnie z Art. 5.'),
        ('"Cena Katalogowa"',
         'oznacza cenę referencyjną dla danego projektu, obliczoną jako: (Stawka MW × zainstalowana '
         'moc w MW) + (Stawka MWh × zainstalowana pojemność w MWh), zgodnie ze stawkami w Załączniku C.'),
        ('"Opłata Licencyjna za Projekt"',
         'oznacza opłatę należną za każdy projekt przed datą Kamienia Milowego, '
         'obliczoną zgodnie z Cennikiem w Załączniku C.'),
        ('"Kamień Milowy"',
         'oznacza łączne wdrożenie 500 MWh pojemności magazynowania energii w ramach Oprogramowania EMS '
         'na podstawie umów zawartych przez Spółkę lub za jej pośrednictwem.'),
        ('"Data Kamienia Milowego"',
         'oznacza datę, w której Kamień Milowy zostaje pisemnie potwierdzony przez Strony.'),
        ('"Licencja Wieczysta"',
         'oznacza nieodwołalną, niewyłączną, bezpłatną, ogólnoświatową licencję udzieloną Spółce '
         'po osiągnięciu Kamienia Milowego.'),
        ('"Uruchomienie"',
         'oznacza instalację techniczną, konfigurację, integrację, weryfikację zgodności z kodeksem '
         'sieciowym i przekazanie Oprogramowania EMS na projekcie BESS klienta.'),
        ('"Opłata za Uruchomienie"',
         'oznacza opłatę pobieraną od klienta za usługi Uruchomienia, zatrzymywaną w całości przez Spółkę.'),
        ('"Marża Oprogramowania"',
         'oznacza Cenę Katalogową dla danego projektu. Wszystkie faktury klientów wyszczególniają '
         'Składnik Licencji Oprogramowania (równy Cenie Katalogowej) oraz Składnik Uruchomienia '
         'jako osobne pozycje. Składnik Uruchomienia jest wyłączony z Marży Oprogramowania.'),
        ('"Premia Rozwojowa"',
         'oznacza płatność należną od Spółki na rzecz Stron Voltusa z tytułu każdej licencji '
         'projektowej, obliczoną jako procent Ceny Katalogowej zgodnie z Art. 6.6.'),
        ('"Opłata Subskrypcyjna EMS" (od klienta)',
         'oznacza roczną opłatę pobieraną przez Spółkę od klienta za dostęp do platformy EMS '
         'i SCADA DISPERON, ustaloną na poziomie 400 EUR za MWh zainstalowanej pojemności rocznie. '
         'Opłata jest fakturowana osobno od jakichkolwiek opłat LTSA lub usług O&M, corocznie od '
         'daty uruchomienia (PAC) projektu.'),
        ('"Opłata Subskrypcyjna Voltusa"',
         'oznacza roczną opłatę płatną przez Spółkę na rzecz Voltusa w odniesieniu do każdego '
         'uruchomionego projektu, w wysokości 20% Opłaty Subskrypcyjnej EMS otrzymanej przez '
         'Spółkę od odpowiedniego klienta za dany rok, płatną w ciągu 14 dni od otrzymania '
         'każdej rocznej Opłaty Subskrypcyjnej EMS od klienta.'),
        ('"Własność Intelektualna" lub "IP"',
         'oznacza wszystkie patenty, znaki towarowe, prawa autorskie, prawa do baz danych, '
         'tajemnice handlowe, know-how, kod źródłowy i inne prawa własności, zarejestrowane '
         'lub niezarejestrowane.'),
        ('"Terytorium"',
         'oznacza Unię Europejską i jej państwa członkowskie oraz inne kraje uzgodnione na piśmie.'),
        ('"Wyłączność"', 'oznacza prawo Spółki do bycia jedynym autoryzowanym dystrybutorem '
         'Oprogramowania EMS w danym kraju, zgodnie z Art. 8.'),
        ('"Ograniczona Wyłączność"',
         'oznacza stan przejściowy, w którym Spółka zachowuje prawo do dokończenia istniejących '
         'projektów, a Voltus może samodzielnie pozyskiwać nowych klientów, zgodnie z Art. 8.4.'),
        ('"Partnerzy Producenci BESS"',
         'oznacza listę producentów sprzętu wskazanych w Załączniku B.'),
        ('"Agent Powierniczy"', 'oznacza niezależną stronę trzecią wyznaczoną do przechowywania '
         'Materiałów Powierniczych zgodnie z Art. 6.5.'),
        ('"Materiały Powiernicze"',
         'oznacza kod źródłowy, skrypty kompilacji, dokumentację techniczną i instrukcje wdrożenia '
         'aktualnej wersji Oprogramowania EMS.'),
        ('"Aneks Wsparcia Oprogramowania"',
         'oznacza harmonogram poziomu usług i konserwacji zawarty w Załączniku D.'),
    ]

    dt = doc.add_table(rows=len(defs), cols=2)
    dt.style = 'Table Grid'
    for i, (term, defn) in enumerate(defs):
        c0, c1 = dt.rows[i].cells[0], dt.rows[i].cells[1]
        if i % 2 == 0:
            _cell_bg(c0, 'EEF3FB'); _cell_bg(c1, 'EEF3FB')
        p0 = c0.paragraphs[0]; p0.clear()
        r0 = p0.add_run(term); r0.font.bold = True; r0.font.size = Pt(9)
        p1 = c1.paragraphs[0]; p1.clear()
        r1 = p1.add_run(defn); r1.font.size = Pt(9); r1.font.color.rgb = BLACK
    for row in dt.rows:
        row.cells[0].width = Cm(4.5); row.cells[1].width = Cm(10.5)
    doc.add_paragraph()

    # ── ARTYKUŁ 2 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 2 — STRUKTURA UDZIAŁÓW I PRZENIESIENIE")
    h2(doc, "2.1  Obecna własność")
    para(doc, "Bezpośrednio przed zawarciem niniejszej Umowy 100% wyemitowanego kapitału zakładowego "
         "Spółki było w posiadaniu Lighthief International.")
    h2(doc, "2.2  Przeniesienie udziałów")
    para(doc, "Z chwilą zawarcia niniejszej Umowy dokonuje się następujących przeniesień udziałów:")
    tbl(doc, [
        ["Udziałowiec", "Udziały", "%", "Wynagrodzenie"],
        ["Lighthief International",             "65", "65%", "N/D (Zatrzymane)"],
        ["Pan Kostas Alexander Papacosta",       "10", "10%", "250 EUR"],
        ["Pan Kamil Talar",                      "13", "13%", "325 EUR"],
        ["Pan Marcin Szumiło",                   "6",  "6%",  "150 EUR"],
        ["Pan Andrzej Lechowicz",                "6",  "6%",  "150 EUR"],
        ["ŁĄCZNIE",                              "100","100%", "—"],
    ], col_widths=[6, 2, 2, 4])
    doc.add_paragraph()

    h2(doc, "2.3  Wynagrodzenie — Potwierdzenie płatności")
    para(doc,
         "Wynagrodzenie wynosi 25 EUR za 1% udziałów. Strony potwierdzają otrzymanie przez Spółkę "
         "następujących kwot: (a) 250 EUR od Pana Papacosty (10%); (b) 325 EUR od Pana Talara "
         "(13%); (c) 150 EUR od każdego z Panów Szumiły i Lechowicza (po 6%). Wydano pokwitowania. "
         "Zobowiązania płatnicze z Art. 2.2 uznaje się za w pełni wykonane.")
    h2(doc, "2.4  Rejestracja przeniesienia")
    para(doc,
         "Pan Arkadiusz Sybaris, działając jako Dyrektor obu podmiotów, dokona stosownych wpisów "
         "w Rejestrze Udziałowców prowadzonym przez cypryjski Urząd Rejestrowy Spółek.")
    h2(doc, "2.5  Powołanie Dyrektora przez Udziałowców")
    para(doc,
         "Udziałowcy inni niż Lighthief International i Pan Papacosta — tj. Panowie Szumiło, "
         "Lechowicz i Talar — w ciągu 30 dni od zawarcia Umowy wspólnie wyznaczą jedną osobę "
         "spośród siebie do pełnienia funkcji Dyrektora zgodnie z Art. 10.2.")
    h2(doc, "2.6  Aktualni Dyrektorzy")
    para(doc,
         "Na dzień zawarcia Umowy Dyrektorami Spółki są:\n"
         "(a) Pan Arkadiusz Sybaris — Dyrektor i Sekretarz, rezydent podatkowy Cypru;\n"
         "(b) Pan Kostas Alexander Papacosta — Dyrektor i Dyrektor Zarządzający na Cyprze, "
         "rezydent podatkowy Cypru, na stałe zamieszkały na Cyprze.")

    # ── ARTYKUŁ 3 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 3 — MARKA DISPERON")
    h2(doc, "3.1  Własność marki")
    para(doc,
         "Spółka jest jedynym prawnym i rzeczywistym właścicielem marki DISPERON, znaku towarowego "
         "i wszelkich powiązanych praw własności intelektualnej, w tym nazwy handlowej, logo, "
         "domeny disperon.com oraz znaków towarowych zarejestrowanych lub do zarejestrowania "
         "w EUIPO w klasach Nicejskich 09 i 42. Wszyscy Udziałowcy potwierdzają tę własność.")
    h2(doc, "3.2  Cel marki")
    para(doc,
         "DISPERON jest wyłączną marką komercyjną, pod którą Spółka prowadzi sprzedaż i dostawę "
         "usług EMS i SCADA dla systemów BESS na terenie UE i całego świata.")
    h2(doc, "3.3  Ochrona marki")
    para(doc,
         "Żaden Udziałowiec nie może używać marki DISPERON ani podobnych oznaczeń do celów innych "
         "niż działalność Spółki bez uprzedniej pisemnej zgody Dyrektora. Udziałowiec naruszający "
         "tę zasadę odpowiada wobec Spółki za wynikłe szkody.")

    # ── ARTYKUŁ 4 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 4 — ROLE I OBOWIĄZKI UDZIAŁOWCÓW")
    h2(doc, "4.1  Lighthief International")
    para(doc, "Lighthief International, jako Udziałowiec Większościowy:")
    for item in [
        "Posiada 65% wyemitowanego kapitału zakładowego Spółki i wykonuje odpowiednie prawa głosu;",
        "Zapewnia kierunek strategiczny i nadzór nad działalnością Spółki;",
        "Rozwija kanały sprzedaży EMS pod marką DISPERON przez własne biura i sieci w Europie;",
        "Zatrzymuje w całości Opłaty za Uruchomienie wygenerowane przez realizowane projekty;",
        "Odpowiada za relacje inwestorskie, rozwój marki i ekspansję międzynarodową DISPERON;",
        "Zapewnia administrację korporacyjną i infrastrukturę siedziby rejestrowej.",
    ]:
        bul(doc, item)

    h2(doc, "4.2  Pan Kamil Talar")
    para(doc, "Pan Kamil Talar:")
    for item in [
        "Posiada 13% wyemitowanego kapitału zakładowego Spółki na własny rachunek;",
        "Aktywnie rozwija kanały sprzedaży EMS DISPERON na rynkach europejskich;",
        "Wykorzystuje istniejące relacje biznesowe, w tym z Linyang Energy Co., Ltd., "
        "do generowania pipeline projektów BESS;",
        "Składa raporty o pipeline i aktywności sprzedażowej Dyrektorowi co miesiąc;",
        "Ponosi współodpowiedzialność z Lighthief International za realizację celów sprzedażowych.",
    ]:
        bul(doc, item)
    para(doc,
         "Zobowiązanie do aktywności: Pan Talar dołoży wszelkich starań handlowych, aby aktywnie "
         "rozwijać i utrzymywać pipeline projektów BESS z Oprogramowaniem EMS oraz będzie "
         "miesięcznie raportował Dyrektorowi swoje działania handlowe i możliwości rynkowe.",
         indent=0.5)

    h2(doc, "4.3  Pan Marcin Szumiło i Pan Andrzej Lechowicz")
    para(doc,
         "Panowie Szumiło i Lechowicz posiadają po 6% wyemitowanego kapitału zakładowego Spółki. "
         "Obaj, indywidualnie i jako członkowie Zarządu Voltus Energy Sp. z o.o.:")
    for item in [
        "Wnoszą Oprogramowanie EMS do działalności Spółki poprzez Licencję;",
        "Ponoszą główną odpowiedzialność za rozwój oprogramowania, architekturę techniczną, "
        "roadmapę produktu i jakość Oprogramowania EMS;",
        "Zapewniają wsparcie techniczne drugiego poziomu dla zespołów uruchomieniowych Spółki;",
        "Wspierają proces sprzedaży poprzez demonstracje techniczne i udział w spotkaniach klientów;",
        "Dbają o zgodność Oprogramowania EMS z przepisami UE, w tym NIS2, RODO i kodeksami sieciowymi.",
    ]:
        bul(doc, item)

    h2(doc, "4.4  Pan Kostas Alexander Papacosta")
    para(doc,
         "Pan Kostas Alexander Papacosta posiada 10% wyemitowanego kapitału zakładowego Spółki. "
         "Pan Papacosta, jako Dyrektor Zarządzający na Cyprze:")
    for item in [
        "Pełni funkcję głównego dyrektora wykonawczego Lighthief EUBESS Ltd na Cyprze, posiadając "
        "wszystkie niezbędne rejestracje i zezwolenia wymagane do działalności Spółki na Cyprze;",
        "Kieruje całą realizacją projektów EPC pod marką DISPERON na Cyprze;",
        "Zarządza relacjami z cypryjskim Operatorem Systemu Dystrybucyjnego (EAC), Operatorem "
        "Systemu Przesyłowego (TSOC) i Urzędem Regulacji Energetyki (CERA);",
        "Pełni funkcję głównego menedżera relacji z klientami cypryjskiego portfela Spółki;",
        "Nadzoruje uruchomienie, integrację SCADA IEC 60870-5-104 i przekazanie wszystkich "
        "instalacji BESS na Cyprze;",
        "Składa miesięczne raporty o wynikach operacyjnych, pipeline i wynikach finansowych.",
    ]:
        bul(doc, item)

    # ── ARTYKUŁ 5 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 5 — WŁASNOŚĆ INTELEKTUALNA")
    h2(doc, "5.1  Własność IP Oprogramowania EMS")
    para(doc,
         "Strony potwierdzają, że wszelkie prawa własności intelektualnej do Oprogramowania EMS "
         "należą wyłącznie do Voltus Energy Sp. z o.o. Niniejsza Umowa nie stanowi przeniesienia, "
         "cesji ani nowacji tych praw na Spółkę ani na żadnego Udziałowca, z wyjątkiem przypadków "
         "wyraźnie przewidzianych w Art. 5 i Art. 6.")
    h2(doc, "5.2  Licencja udzielona Spółce")
    para(doc,
         "Voltus Energy Sp. z o.o. udziela Spółce niewyłącznej, ogólnoświatowej licencji na "
         "używanie, wdrażanie, konfigurowanie, adaptowanie i sublicencjonowanie Oprogramowania EMS "
         "klientom pod marką DISPERON, na warunkach określonych w Art. 6.")
    h2(doc, "5.3  Voltus zachowuje pełne IP")
    para(doc,
         "Voltus Energy Sp. z o.o. zachowuje pełną i wyłączną własność całego IP Oprogramowania EMS "
         "przez cały czas. Licencja Wieczysta przyznaje Spółce prawo do korzystania z Oprogramowania, "
         "lecz nie przenosi własności ani żadnych praw właścicielskich.")
    h2(doc, "5.4  IP Spółki i marka DISPERON")
    para(doc,
         "Wszelkie IP stworzone przez Spółkę lub na jej zlecenie w związku z marką DISPERON — "
         "w tym materiały marketingowe, treści strony internetowej, metodologie wdrożeniowe, "
         "dokumentacja klientów i znak towarowy DISPERON — należy wyłącznie do Spółki.")
    h2(doc, "5.5  Gwarancja IP i odszkodowanie")
    para(doc, "Voltus Energy Sp. z o.o. oświadcza i gwarantuje Spółce, że:")
    for item in [
        "Jest jedynym i nieobciążonym właścicielem wszelkich praw własności intelektualnej "
        "do Oprogramowania EMS;",
        "Oprogramowanie EMS nie narusza żadnych praw własności intelektualnej osób trzecich;",
        "Żadne składniki oprogramowania open source nie są zawarte w Oprogramowaniu EMS na "
        "warunkach licencji niezgodnych z komercyjną dystrybucją lub sublicencjonowaniem;",
        "Oprogramowanie EMS spełnia wszystkie obowiązujące przepisy UE dotyczące kontroli eksportu.",
    ]:
        bul(doc, item)
    para(doc,
         "Voltus zobowiązuje się bronić Spółkę i jej dyrektorów przed wszelkimi roszczeniami "
         "osób trzecich wynikającymi z naruszenia gwarancji określonych w Art. 5.5 oraz pokryć "
         "wszelkie straty, szkody, koszty i honoraria prawne z tym związane.")

    # ── ARTYKUŁ 6 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 6 — MODEL HANDLOWY I OPŁATY LICENCYJNE")

    h2(doc, "6.1  Cena Katalogowa")
    para(doc,
         "Dla każdego projektu Cenę Katalogową oblicza się jako: (Stawka MW × zainstalowana moc "
         "w MW) + (Stawka MWh × zainstalowana pojemność w MWh), zgodnie z Załącznikiem C. "
         "Voltus potwierdza Cenę Katalogową dla każdego przesłanego briefu projektowego w ciągu "
         "14 dni kalendarzowych od jego otrzymania. Brak odpowiedzi w ciągu 14 dni oznacza "
         "dorozumiane zatwierdzenie według stawek z Załącznika C.")

    h2(doc, "6.2  Model przychodów przed Kamieniem Milowym")
    para(doc,
         "(a) Opłata Licencyjna za Projekt: Przed datą Kamienia Milowego Spółka płaci Voltusowi "
         "Opłatę Licencyjną za każdy projekt równą Cenie Katalogowej, proporcjonalnie do płatności "
         "etapowych otrzymanych od klienta, przy czym całość opłaty musi być uiszczona nie później "
         "niż 14 dni od daty uruchomienia, niezależnie od statusu płatności klienta.")
    para(doc,
         "(b) Opłata za Uruchomienie: Spółka ustala i pobiera Opłatę za Uruchomienie bezpośrednio "
         "od klienta. Jest ona w całości zatrzymywana przez Spółkę. Voltus nie ma prawa do żadnej "
         "części Opłaty za Uruchomienie.")

    h2(doc, "6.3  Definicja Kamienia Milowego")
    para(doc,
         "Kamień Milowy uważa się za osiągnięty, gdy łączna zainstalowana pojemność projektów BESS "
         "wdrażających Oprogramowanie EMS w ramach umów zawartych przez Spółkę lub za jej "
         "pośrednictwem osiągnie 500 MWh od daty zawarcia Umowy. Kamień Milowy nie ma "
         "terminu granicznego.")

    h2(doc, "6.4  Licencja Wieczysta po osiągnięciu Kamienia Milowego")
    para(doc,
         "Po osiągnięciu Kamienia Milowego Voltus automatycznie udziela Spółce wieczystej, "
         "nieodwołalnej, niewyłącznej, ogólnoświatowej, bezpłatnej licencji na używanie, "
         "wdrażanie, adaptowanie i sublicencjonowanie Oprogramowania EMS we wszystkich jego "
         "wersjach — bieżącej i wszystkich przyszłych. Licencja Wieczysta przeżywa zmianę "
         "kontroli, rozwiązanie lub restrukturyzację Voltusa, chyba że wynika z istotnego "
         "i nienaprawionego naruszenia Umowy przez Spółkę.")

    h2(doc, "6.5  Depozyt kodu źródłowego (Escrow)")
    para(doc,
         "(a) Voltus, w ciągu 60 dni od zawarcia Umowy, złoży Materiały Powiernicze u uzgodnionego "
         "niezależnego Agenta Powierniczego (proponowany dostawca: NCC Group Escrow lub odpowiednik).")
    para(doc,
         "(b) Voltus aktualizuje Materiały Powiernicze w ciągu 30 dni od każdego istotnego "
         "wydania lub aktualizacji Oprogramowania EMS. Koszty depozytu ponoszą w równych "
         "częściach Voltus i Spółka.")
    para(doc,
         "(c) Warunki wyzwolenia: Agent Powierniczy przekazuje Materiały Powiernicze Spółce "
         "w przypadku: (i) upadłości lub likwidacji Voltusa bez przejęcia obowiązków "
         "utrzymaniowych przez następcę w ciągu 90 dni; (ii) zaprzestania świadczenia "
         "usług konserwacyjnych przez ponad 90 dni bez naprawienia sytuacji; (iii) istotnego "
         "i nienaprawionego naruszenia Art. 7.2 utrzymującego się ponad 30 dni po pisemnym "
         "powiadomieniu ze strony Spółki.")

    h2(doc, "6.6  Premia Rozwojowa (po Kamieniu Milowym)")
    para(doc,
         "(a) Po dacie Kamienia Milowego Spółka płaci Stronom Voltusa Premię Rozwojową obliczoną jako:")
    for item in [
        "25% Ceny Katalogowej dla projektów w krajach, w których Spółka nie posiada Wyłączności;",
        "30% Ceny Katalogowej dla projektów w krajach, w których Spółka posiada Wyłączność.",
    ]:
        bul(doc, item)
    para(doc,
         "(b) Stawki Premii Rozwojowej podlegają przeglądowi co 3 lata w zakresie 20%–35%, "
         "pod warunkiem że Voltus przedstawi kwartalne zestawienia wydatków na R&D.")
    para(doc,
         "(c) Premia Rozwojowa jest obliczana wyłącznie od Ceny Katalogowej. Opłaty za "
         "Uruchomienie są wyraźnie wyłączone.")
    para(doc,
         "(e) Premia Rozwojowa jest płatna w ciągu 14 dni od otrzymania przez Spółkę "
         "zapłaty od klienta. Przedpłaty i płatności częściowe nie uruchamiają Premii "
         "przed datą uruchomienia danego projektu.")

    h2(doc, "6.7  Opłata Subskrypcyjna EMS (od klienta) i Opłata Subskrypcyjna Voltusa")
    para(doc,
         "(a) Opłata Subskrypcyjna EMS pobierana od klientów: Spółka pobiera od każdego klienta "
         "roczną Opłatę Subskrypcyjną EMS w wysokości 400 EUR za MWh zainstalowanej pojemności "
         "rocznie. Opłata jest fakturowana corocznie od daty uruchomienia (PAC) projektu "
         "i osobno od wszelkich opłat LTSA, O&M lub usług fizycznych. Klient płaci w ciągu "
         "14 dni od każdej rocznej faktury.")
    para(doc,
         "(b) Opłata Subskrypcyjna Voltusa: W odniesieniu do każdego uruchomionego projektu "
         "Spółka płaci Voltusowi roczną Opłatę Subskrypcyjną Voltusa równą 20% Opłaty "
         "Subskrypcyjnej EMS otrzymanej od klienta w danym roku. Płatna w ciągu 14 dni "
         "od otrzymania każdej rocznej Opłaty Subskrypcyjnej EMS od klienta.")
    para(doc,
         "(c) Przykład: Dla projektu 30 MWh — Opłata EMS od klienta: 12 000 EUR/rok. "
         "Opłata Voltusa: 20% × 12 000 EUR = 2 400 EUR/rok. DISPERON zatrzymuje: 9 600 EUR/rok.")
    para(doc,
         "(d) Opłata Subskrypcyjna EMS obejmuje: (i) wszystkie aktualizacje EMS i SCADA; "
         "(ii) wsparcie produktu (diagnostyka, hotfixy); (iii) aktualizacje regulacyjne "
         "w ciągu 90 dni od publikacji; (iv) zdalny dostęp do dashboardu SCADA i utrzymanie "
         "integracji DSO. Nie obejmuje konfiguracji na miejscu, uruchamiania ani fizycznego "
         "wsparcia klienta końcowego.")
    para(doc,
         "(e) Brak otrzymania Opłaty EMS od klienta przez ponad 90 dni nie wpływa na licencję "
         "Spółki na używanie Oprogramowania EMS. Nieprzykazanie 20% Opłaty EMS Voltusowi "
         "w terminie 14 dni uprawnia Voltusa do zawieszenia nowych aktualizacji po 90 dniach "
         "pisemnego wezwania bez naprawy. Licencja na używanie bieżącej wersji pozostaje "
         "w mocy w każdych okolicznościach.")

    h2(doc, "6.8  Dokumentacja i audyt")
    para(doc,
         "Spółka prowadzi dokładną dokumentację wszystkich wdrożeń projektów, zapłaconych "
         "opłat licencyjnych, otrzymanych Opłat za Uruchomienie, płatności Premii Rozwojowej "
         "i Opłat Subskrypcyjnych. Voltus ma prawo zlecić niezależny audyt raz w roku "
         "kalendarzowym, po co najmniej 20 dniach roboczych pisemnego powiadomienia. "
         "Koszty audytu ponosi Voltus, chyba że wykaże niedopłatę przekraczającą 5%.")

    # ── ARTYKUŁ 7 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 7 — STRUKTURA OPERACYJNA")
    h2(doc, "7.1  Sprzedaż i rozwój rynku")
    para(doc,
         "Spółka, działając przez Lighthief International, Pana Talara i Pana Papacostę, "
         "odpowiada głównie za komercyjny rozwój marki DISPERON i generowanie pipeline klientów "
         "na Terytorium. Pan Papacosta kieruje rozwojem rynku, zarządzaniem klientami "
         "i realizacją EPC na Cyprze.")
    h2(doc, "7.2  Technologia i rozwój produktu")
    para(doc,
         "Voltus Energy Sp. z o.o. odpowiada głównie za ciągły rozwój, konserwację, "
         "bezpieczeństwo i doskonalenie Oprogramowania EMS, utrzymując je w stanie gotowym "
         "do wdrożenia komercyjnego i zgodnym z przepisami. W szczególności Voltus:")
    for item in [
        "Odpowiada na klasyfikowane zgłoszenia wsparcia w terminach określonych w Załączniku D;",
        "Utrzymuje dostępność chmurowych komponentów EMS na poziomie minimum 99,5% miesięcznie;",
        "Implementuje zmiany wynikające z nowych lub zmienionych przepisów UE i krajowych "
        "kodeksów sieciowych w ciągu 90 dni od oficjalnej publikacji;",
        "Dostarcza Spółce kwartalne raporty o roadmapie produktu i działaniach rozwojowych;",
        "Utrzymuje zgodność z NIS2 i obowiązującymi standardami UE w zakresie cyberbezpieczeństwa.",
    ]:
        bul(doc, item)
    h2(doc, "7.3  Uruchomienie i wdrożenie")
    para(doc,
         "Usługi instalacji, integracji systemów i przekazania projektów świadczy Spółka lub "
         "podmiot z grupy Lighthief International. Voltus zapewnia zdalne wsparcie techniczne, "
         "dokumentację i szkolenia w ramach Licencji i Opłaty Subskrypcyjnej. Wszystkie "
         "przychody z Uruchomienia należą do Spółki.")
    h2(doc, "7.4  Umowy z klientami")
    para(doc,
         "Wszystkie umowy z klientami zawiera Spółka od dnia zawarcia niniejszej Umowy. "
         "Spółka jest stroną prawną wobec wszystkich klientów końcowych. Voltus działa "
         "jako dostawca oprogramowania dla Spółki, nie dla klientów końcowych.")

    # ── ARTYKUŁ 8 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 8 — WYŁĄCZNOŚĆ I MODEL RYNKOWY")
    h2(doc, "8.1  Zasada zdobytej wyłączności")
    para(doc,
         "Wyłączność nie jest przyznawana automatycznie. Można ją zdobyć w konkretnych krajach "
         "na podstawie wykazanej aktywności komercyjnej i osiągniętych progów wolumenowych.")
    h2(doc, "8.2  Warunki wstępne wyłączności")
    para(doc, "Przed uznaniem wyłączności w danym kraju Spółka dostarcza Voltusowi:")
    for item in [
        "Udokumentowany pipeline projektów w danym kraju (nazwy klientów, rozmiary, terminy);",
        "Pisemny plan sprzedaży na minimum 12 miesięcy;",
        "Dowody aktywnego zaangażowania komercyjnego (spotkania, oferty, LOI).",
    ]:
        bul(doc, item)
    h2(doc, "8.3  Progi wolumenowe wyłączności")
    tbl(doc, [
        ["Kategoria rynku", "Kraje", "Min. wolumen", "Utrzymanie roczne"],
        ["Duże rynki", "Niemcy, Włochy, Hiszpania, Francja", "500 MWh / 24 mies.", "100 MWh / rok"],
        ["Średnie rynki", "Holandia, Rumunia, Grecja, Czechy, Austria",
         "250 MWh / 24 mies.", "50 MWh / rok"],
        ["Małe rynki", "Cypr, inne kraje UE", "100 MWh / 24 mies.", "20 MWh / rok"],
    ], col_widths=[3.2, 4.5, 3.8, 3.5])
    doc.add_paragraph()

    h2(doc, "8.4  Utrzymanie i wygaśnięcie wyłączności")
    para(doc,
         "(a) Ograniczona Wyłączność: Nieosiągnięcie progu utrzymania rocznego w jednym roku "
         "powoduje wprowadzenie statusu Ograniczonej Wyłączności na kolejne 12 miesięcy. "
         "W tym czasie Spółka zachowuje prawo do realizacji istniejących projektów, a Voltus "
         "może pozyskiwać nowych klientów spoza udokumentowanego pipeline Spółki.")
    para(doc,
         "(b) Pełne wygaśnięcie: Nieosiągnięcie progu utrzymania przez dwa kolejne lata "
         "skutkuje automatycznym wygaśnięciem wyłączności. Spółka może ponownie ubiegać "
         "się o wyłączność po 12 miesiącach od wygaśnięcia.")
    h2(doc, "8.5  Limit krajów i moratorium")
    para(doc,
         "Spółka może ubiegać się o wyłączność w maksymalnie 5 krajach jednocześnie. "
         "Polska, Litwa, Łotwa i Estonia ('Kraje Moratorium') podlegają 24-miesięcznemu "
         "moratorium od daty zawarcia Umowy. Po upływie moratorium Spółka może ubiegać "
         "się o wyłączność na standardowych warunkach z Art. 8.3.")
    h2(doc, "8.6  Bezpośrednie prawa Voltusa — Partnerzy Producenci BESS")
    para(doc,
         "Niezależnie od przyznanej wyłączności Voltus zachowuje prawo do bezpośredniej "
         "dostawy Oprogramowania EMS do Partnerów Producentów BESS wymienionych w Załączniku B. "
         "Prawo to dotyczy wyłącznie podmiotów wymienionych w Załączniku B i nie podlega "
         "ograniczeniom geograficznym w odniesieniu do wymienionych producentów. Dodanie "
         "nowego podmiotu wymaga pisemnej zgody wszystkich Stron.")
    para(doc,
         "Obowiązek przekierowania: Jeśli projekt Partnera Producenta BESS był przedmiotem "
         "wcześniejszego zaangażowania komercyjnego Spółki, Voltus powiadamia Spółkę w ciągu "
         "5 dni roboczych i kieruje projekt do Spółki. Brak porozumienia w ciągu 15 dni "
         "skutkuje zastosowaniem standardowej stawki Premii Rozwojowej.")
    h2(doc, "8.7  Przejrzystość i rejestracja klientów")
    para(doc,
         "Każda Strona może rejestrować szanse sprzedażowe, przy których podjęła aktywne "
         "zaangażowanie. W przypadku konfliktu Strony najpierw dążą do współpracy w dobrej "
         "wierze. W przypadku sporu Dyrektor rozstrzyga w ciągu 15 dni roboczych.")
    h2(doc, "8.8  Zobowiązanie Voltusa do niekonkurowania")
    para(doc,
         "Voltus zobowiązuje się przez cały czas obowiązywania Umowy do pełnej przejrzystości "
         "wobec Spółki w zakresie bezpośredniej sprzedaży na Terytorium oraz do kierowania do "
         "Spółki klientów, z którymi Spółka prowadziła aktywne rozmowy. Zobowiązanie to nie "
         "dotyczy klientów samodzielnie pozyskanych przez Voltus.")
    h2(doc, "8.9  Procedura niezadowolonego klienta")
    para(doc,
         "(a) W przypadku formalnej skargi klienta Spółka ma 60 dni na usunięcie uchybień. "
         "(b) Skargę uznaje się za 'potwierdzoną' wyłącznie jeżeli: (i) złożona jest pisemnie; "
         "(ii) Spółka nie usunęła uchybień w terminie; (iii) Zarząd niezależnie zweryfikował "
         "brak działań naprawczych. (c) Ponad 3 potwierdzone skargi dotyczące jednego klienta "
         "w ciągu 24 miesięcy uprawniają Voltusa, po pisemnym powiadomieniu Spółki, do "
         "bezpośredniej obsługi tego klienta. Taka bezpośrednia dostawa podlega standardowej "
         "Premii Rozwojowej płatnej na rzecz DISPERON. (d) Skargi wynikające z przyczyn "
         "niezależnych od Spółki (siła wyższa, opóźnienia regulacyjne, usterki oprogramowania "
         "Voltusa) nie są wliczane do liczby potwierdzonych skarg.")

    # ── ARTYKUŁ 9 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 9 — PRZENIESIENIE UDZIAŁÓW")
    h2(doc, "9.1  Okres zakazu przenoszenia")
    para(doc,
         "Przez 24 miesiące od daty Umowy żaden Udziałowiec nie może zbywać, przenosić, "
         "zastawiać ani rozporządzać Udziałami bez uprzedniej pisemnej zgody Lighthief International.")
    h2(doc, "9.2  Prawo pierwokupu")
    para(doc,
         "Po upływie okresu zakazu przenoszenia, Udziałowiec zamierzający zbyć Udziały "
         "powiadamia pisemnie wszystkich pozostałych Udziałowców o liczbie Udziałów, "
         "proponowanej cenie i tożsamości nabywcy. Każdy nie-sprzedający Udziałowiec "
         "ma 30 dni na wykonanie prawa pierwokupu proporcjonalnie na tych samych warunkach.")
    h2(doc, "9.3  Prawo przyciągnięcia (Drag-along)")
    para(doc,
         "Jeżeli Lighthief International proponuje sprzedaż 100% swoich Udziałów nabywcy "
         "wymagającemu 100% kapitału zakładowego, Lighthief International może wymagać od "
         "pozostałych Udziałowców sprzedaży ich Udziałów na tych samych warunkach, "
         "za 30-dniowym pisemnym wypowiedzeniem.")
    h2(doc, "9.4  Prawo przyłączenia (Tag-along)")
    para(doc,
         "Jeżeli Lighthief International proponuje przeniesienie ponad 50% swoich Udziałów, "
         "każdy pozostały Udziałowiec ma prawo w ciągu 20 dni dołączyć do transakcji "
         "na tych samych warunkach.")

    # ── ARTYKUŁ 10 ────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 10 — ZARZĄDZANIE I ZARZĄD")
    h2(doc, "10.1  Aktualni Dyrektorzy")
    para(doc,
         "Na dzień Umowy Dyrektorami Spółki są:\n"
         "(a) Pan Arkadiusz Sybaris — Dyrektor i Sekretarz, rezydent podatkowy Cypru;\n"
         "(b) Pan Kostas Alexander Papacosta — Dyrektor i Dyrektor Zarządzający na Cyprze.")
    h2(doc, "10.2  Skład Zarządu")
    para(doc,
         "(a) Lighthief International ma prawo powołać do 3 Dyrektorów, w tym Dyrektora "
         "i Sekretarza. Pan Papacosta jest jednym z Dyrektorów powoływanych przez Lighthief "
         "International na czas trwania Umowy.\n"
         "(b) Udziałowcy inni niż Lighthief International i Pan Papacosta (tj. Panowie "
         "Szumiło, Lechowicz i Talar łącznie) mają prawo powołać 1 Dyrektora.")
    h2(doc, "10.3  Powoływanie i odwoływanie")
    para(doc,
         "Dyrektorzy powołani przez Lighthief International mogą być odwołani w każdym czasie "
         "przez pisemną uchwałę Lighthief International, przy czym odwołanie Pana Papacosty "
         "wymaga dodatkowo uchwały Zarządu.")
    h2(doc, "10.4  Decyzje wymagające zgody Udziałowców")
    para(doc, "Poniższe sprawy wymagają pisemnej zgody Udziałowców posiadających "
         "większość Udziałów:")
    for item in [
        "Zmiana umowy Spółki;",
        "Emisja nowych udziałów lub podwyższenie kapitału zakładowego;",
        "Istotna zmiana głównego przedmiotu działalności Spółki;",
        "Zawarcie transakcji z podmiotami powiązanymi przekraczającej 50 000 EUR;",
        "Powołanie lub odwołanie biegłego rewidenta.",
    ]:
        bul(doc, item)
    h2(doc, "10.5  Sprawy zastrzeżone wymagające 85% zgody Udziałowców")
    para(doc, "Poniższe sprawy wymagają pisemnej zgody Udziałowców posiadających "
         "co najmniej 85% wszystkich Udziałów:")
    for item in [
        "Istotna zmiana modelu handlowego z Art. 6, w tym stawek Premii Rozwojowej, "
        "Opłaty Subskrypcyjnej lub Cennika z Załącznika C;",
        "Dobrowolne rozwiązanie lub likwidacja Spółki;",
        "Jakakolwiek transakcja lub uchwała, która naruszyłaby, obciążyła lub "
        "zakończyła Licencję Wieczystą lub depozyt kodu źródłowego z Art. 6.5.",
    ]:
        bul(doc, item)

    # ── ARTYKUŁ 11 ────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 11 — POUFNOŚĆ")
    h2(doc, "11.1  Informacje poufne")
    para(doc,
         'Każda Strona zachowuje w tajemnicy wszelkie niepubliczne informacje otrzymane '
         'w związku z Umową, w tym specyfikacje techniczne, kod źródłowy, dane klientów, '
         'warunki finansowe i strategie biznesowe ("Informacje Poufne").')
    h2(doc, "11.2  Czas trwania")
    para(doc,
         "Obowiązek zachowania poufności obowiązuje przez czas trwania Umowy i przez "
         "5 lat po jej rozwiązaniu lub wygaśnięciu.")
    h2(doc, "11.3  Dozwolone ujawnienie")
    para(doc,
         "Informacje Poufne mogą być ujawniane wyłącznie pracownikom, wykonawcom lub "
         "doradcom mającym uzasadnioną potrzebę dostępu, związanym równoważnymi "
         "zobowiązaniami do poufności, lub gdy wymagają tego przepisy prawa.")

    # ── ARTYKUŁ 12 ────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 12 — PRAWO WŁAŚCIWE I ROZWIĄZYWANIE SPORÓW")
    h2(doc, "12.1  Prawo właściwe")
    para(doc,
         "Niniejsza Umowa podlega prawu Republiki Cypru i jest interpretowana zgodnie z nim.")
    h2(doc, "12.2  Negocjacje")
    para(doc,
         'W przypadku jakiegokolwiek sporu wynikającego z Umowy ("Spór") Strony podejmują '
         "próbę rozwiązania go w drodze negocjacji w dobrej wierze w terminie 30 dni "
         "od pisemnego powiadomienia o Sporze.")
    h2(doc, "12.3  Arbitraż")
    para(doc,
         "Jeżeli Spór nie zostanie rozwiązany w drodze negocjacji, zostanie skierowany do "
         "ostatecznego i wiążącego arbitrażu w Limassol (Cypr) na podstawie cypryjskiego "
         "Prawa Arbitrażowego (Kap. 4), prowadzonego w języku angielskim.")

    # ── ARTYKUŁ 13 ────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 13 — CZAS TRWANIA UMOWY I MECHANIZM WYJŚCIA")

    h2(doc, "13.1  Okres obowiązywania i odnowienie")
    para(doc,
         "Niniejsza Umowa zostaje zawarta na wstępny okres 5 lat od daty zawarcia "
         '("Wstępny Okres"). Po upływie Wstępnego Okresu Umowa automatycznie odnawia się '
         "na kolejne okresy 2-letnie, chyba że którakolwiek ze Stron złoży pisemne "
         "wypowiedzenie z co najmniej 12-miesięcznym wyprzedzeniem przed końcem "
         "bieżącego okresu.")

    h2(doc, "13.2  Element A — Wyjście czasowe z Prowizją Następczą")
    para(doc,
         "(a) Po upływie Wstępnego Okresu każda Strona może rozwiązać Umowę, składając "
         "12-miesięczne pisemne wypowiedzenie wszystkim pozostałym Stronom.")
    para(doc,
         "(b) Jeżeli Voltus korzysta z prawa wyjścia na podstawie niniejszego artykułu, "
         "Voltus płaci Spółce Prowizję Następczą przez 12 miesięcy od daty rozwiązania "
         "w wysokości 10% Premii Rozwojowej od projektów obsługiwanych bezpośrednio "
         "przez Voltusa na byłym terytorium DISPERON w tym 12-miesięcznym okresie. "
         "Prowizja ta odzwierciedla relacje komercyjne zbudowane przez DISPERON.")
    para(doc,
         "(c) Jeżeli Spółka korzysta z prawa wyjścia, Prowizja Następcza nie jest należna.")

    h2(doc, "13.3  Element B — Wyjście na podstawie wyników")
    para(doc,
         "Każda Strona może rozwiązać Umowę bez odszkodowania i bez okresu wypowiedzenia "
         "w przypadku wystąpienia Zdarzenia Wyzwalającego po stronie drugiej Strony.")
    para(doc, "Zdarzenia Wyzwalające dające Voltusowi prawo do wyjścia:", bold=True)
    for item in [
        "Roczny wolumen MWh uruchomiony przez Spółkę spada poniżej 50% progu utrzymania "
        "z Art. 8.3 przez 2 kolejne lata;",
        "Spółka nie podpisuje żadnych nowych umów z klientami przez 9 kolejnych miesięcy "
        "w kraju z Wyłącznością;",
        "Spółka zalega z płatnościami z Art. 6.6 lub 6.7 przez ponad 90 dni po pisemnym "
        "powiadomieniu.",
    ]:
        bul(doc, item)
    para(doc, "Zdarzenia Wyzwalające dające Spółce prawo do wyjścia:", bold=True)
    for item in [
        "Voltus nie dostarcza zamówionych aktualizacji oprogramowania przez ponad 6 miesięcy "
        "bez uzasadnienia technicznego;",
        "Voltus traci certyfikacje produktowe wymagane dla zgodności z przepisami UE i nie "
        "odzyskuje ich w ciągu 90 dni;",
        "Voltus podwyższa stawki Cennika z Załącznika C o ponad 50% w jednej rocznej korekcie "
        "bez pisemnej zgody wszystkich Stron.",
    ]:
        bul(doc, item)
    para(doc,
         "Procedura: Strona powołująca się na Zdarzenie Wyzwalające dostarcza pisemne "
         "powiadomienie z dokumentacją. Druga Strona ma 30 dni na naprawienie sytuacji. "
         "W razie braku naprawy wyjście staje się skuteczne po 30 dniach od upływu "
         "okresu naprawczego.")

    h2(doc, "13.4  Element C — 12-miesięczna stopniowa tranzycja rynkowa")
    para(doc,
         "Po wyjściu na podstawie Art. 13.2 lub 13.3 obowiązuje następujący okres przejściowy:")
    tbl(doc, [
        ["Faza", "Miesiące", "Prawa Spółki", "Prawa Voltusa"],
        ["Pełna ochrona", "0–4",
         "Pełna Wyłączność. Spółka może podpisywać umowy z klientami w aktywnej negocjacji.",
         "Voltus przygotowuje własne kanały sprzedaży. Brak sprzedaży w krajach wyłącznych."],
        ["Ograniczona", "5–8",
         "Spółka realizuje istniejące umowy. Brak nowych podpisań. ROFR na klientów z ostatnich 24 mies.",
         "Voltus może pozyskiwać nowych klientów spoza pipeline Spółki."],
        ["Wspólny rynek", "9–12",
         "ROFR Spółki ograniczone do klientów aktywnie obsługiwanych w ostatnich 12 miesiącach.",
         "Voltus może konkurować we wszystkich nowych szansach."],
        ["Pełna wolność", "13+",
         "Brak ROFR, brak ochrony wyłączności.",
         "Voltus ma pełną wolność rynkową we wszystkich krajach."],
    ], col_widths=[2.5, 2, 5, 5])
    doc.add_paragraph()

    # ── ARTYKUŁ 14 ────────────────────────────────────────────────────────────
    h1(doc, "ARTYKUŁ 14 — POSTANOWIENIA OGÓLNE")
    h2(doc, "14.1  Całość Umowy")
    para(doc,
         "Niniejsza Umowa stanowi całość porozumienia Stron w zakresie jej przedmiotu "
         "i zastępuje wszelkie wcześniejsze ustalenia, negocjacje i oświadczenia.")
    h2(doc, "14.2  Zmiany")
    para(doc, "Zmiany Umowy wymagają pisemnego instrumentu podpisanego przez wszystkie Strony.")
    h2(doc, "14.3  Rozdzielność postanowień")
    para(doc,
         "Jeżeli jakiekolwiek postanowienie Umowy zostanie uznane za nieważne lub niewykonalne, "
         "pozostałe postanowienia zachowują pełną moc.")
    h2(doc, "14.4  Zrzeczenie się")
    para(doc,
         "Niewykonanie przez Stronę prawa lub środka zaradczego nie stanowi zrzeczenia się tego prawa.")
    h2(doc, "14.5  Powiadomienia")
    para(doc,
         "Wszelkie powiadomienia wymagają formy pisemnej i są dostarczane e-mailem z potwierdzeniem "
         "odczytu lub listem poleconym na adresy wskazane w Preambule.")
    h2(doc, "14.6  Egzemplarze i podpisy elektroniczne")
    para(doc,
         "Umowę można zawrzeć w kilku egzemplarzach. Podpisy elektroniczne są ważne i wiążące. "
         "Wersja angielska niniejszej Umowy ma pierwszeństwo przed tłumaczeniem polskim.")

    # ── PODPISY ───────────────────────────────────────────────────────────────
    doc.add_page_break()
    p_s = doc.add_paragraph(); p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = p_s.add_run("PODPISY")
    rs.font.size = Pt(16); rs.font.bold = True; rs.font.color.rgb = BLUE

    doc.add_paragraph()
    para(doc,
         "NA DOWÓD CZEGO Strony zawarły niniejszą Umowę w dniu wskazanym na początku. "
         "Każdy podpisujący oświadcza i gwarantuje, że posiada pełne upoważnienie do "
         "zawarcia Umowy w imieniu reprezentowanej strony.", italic=True)
    doc.add_paragraph()

    sig_block(doc, "STRONA 1 — LIGHTHIEF INTERNATIONAL",
              "Reprezentowana przez: Pana Arkadiusza Sybarisa",
              "Stanowisko: Dyrektor", ref="Strona 1")
    sig_block(doc, "STRONA 2 — LIGHTHIEF EUBESS LTD (HE 474192)",
              "Reprezentowana przez: Pana Arkadiusza Sybarisa",
              "Stanowisko: Dyrektor i Sekretarz", ref="Strona 2")
    sig_block(doc, "STRONA 3 — PAN MARCIN SZUMIŁO",
              "Udziałowiec indywidualny", ref="Strona 3")
    sig_block(doc, "STRONA 4 — PAN ANDRZEJ LECHOWICZ",
              "Udziałowiec indywidualny", ref="Strona 4")
    sig_block(doc, "STRONA 5 — PAN KAMIL TALAR",
              "Udziałowiec indywidualny", ref="Strona 5")
    sig_block(doc, "STRONA 6 — PAN KOSTAS ALEXANDER PAPACOSTA",
              "Udziałowiec indywidualny i Dyrektor Zarządzający na Cyprze", ref="Strona 6")
    sig_block(doc, "STRONA 7 — VOLTUS ENERGY SP. Z O.O.",
              "Reprezentowana przez: Pana Marcina Szumiłę, Członka Zarządu",
              extra="Reprezentowana przez: Pana Andrzeja Lechowicza, Członka Zarządu",
              ref="Strona 7")

    # ── ZAŁĄCZNIK A ───────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "ZAŁĄCZNIK A — CERTYFIKATY UDZIAŁÓW")
    para(doc,
         "Poniższe certyfikaty udziałów Lighthief EUBESS Ltd (HE 474192) zostaną wydane "
         "w ciągu 30 dni od rejestracji w cypryjskim Urzędzie Rejestrowym Spółek:")
    for item in [
        "Aktualny certyfikat udziałów wydany dla Lighthief International (100% przed przeniesieniem);",
        "Nowy certyfikat udziałów wydany dla Pana Kostasa Alexandra Papacosty (10%);",
        "Nowy certyfikat udziałów wydany dla Pana Kamila Talara (13%);",
        "Nowy certyfikat udziałów wydany dla Pana Marcina Szumiły (6%);",
        "Nowy certyfikat udziałów wydany dla Pana Andrzeja Lechowicza (6%);",
        "Zaktualizowany certyfikat udziałów wydany dla Lighthief International (65% po przeniesieniu).",
    ]:
        bul(doc, item)

    # ── ZAŁĄCZNIK B ───────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "ZAŁĄCZNIK B — PARTNERZY PRODUCENCI BESS")
    para(doc,
         "Poniżej znajduje się kompletna i wyłączna lista producentów sprzętu BESS, "
         "wobec których Voltus Energy Sp. z o.o. zachowuje bezpośrednie prawa do dostawy "
         "zgodnie z Art. 8.6. Dodanie nowego podmiotu wymaga pisemnej zgody wszystkich Stron.")
    doc.add_paragraph()
    tbl(doc, [
        ["Nr", "Nazwa producenta", "Kraj rejestracji", "Data dodania"],
        ["1.", "Linyang Energy Co., Ltd.", "Chińska Republika Ludowa", "Data zawarcia Umowy"],
        ["2.", "________________________", "________________________", "________________"],
        ["3.", "________________________", "________________________", "________________"],
        ["4.", "________________________", "________________________", "________________"],
    ], col_widths=[1.2, 6, 4.5, 3.3])
    doc.add_paragraph()
    para(doc, "Potwierdzono przez Voltus Energy Sp. z o.o.:")
    para(doc, "Pan Marcin Szumiło  ________________________________     Data: __________________")
    para(doc, "Pan Andrzej Lechowicz  ________________________________   Data: __________________")
    doc.add_paragraph()
    para(doc, "Potwierdzono przez Lighthief EUBESS Ltd:")
    para(doc, "Pan Arkadiusz Sybaris  ________________________________    Data: __________________")

    # ── ZAŁĄCZNIK C ───────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "ZAŁĄCZNIK C — CENNIK I HARMONOGRAM OPŁAT LICENCYJNYCH")
    para(doc, "Niniejszy Załącznik C stanowi część Umowy i reguluje obliczanie Ceny Katalogowej "
         "na podstawie Art. 6.1, 6.2, 6.6 i 6.7.")
    h2(doc, "C.1  Stawki Ceny Katalogowej")
    tbl(doc, [
        ["Składnik", "Stawka", "Jednostka", "Data obowiązywania"],
        ["Stawka MW",  "EUR [___]", "za MW zainstalowane",  "Data zawarcia Umowy"],
        ["Stawka MWh", "EUR [___]", "za MWh zainstalowane", "Data zawarcia Umowy"],
    ], col_widths=[5, 3, 4, 4])
    doc.add_paragraph()
    para(doc,
         "Stawki w nawiasach zostaną uzgodnione przez Strony w ciągu 14 dni od zawarcia Umowy "
         "i wpisane aneksem podpisanym przez wszystkie Strony. Do czasu uzgodnienia żaden "
         "brief projektowy nie jest uznawany za złożony dla celów Art. 6.2.")
    h2(doc, "C.2  Roczna waloryzacja")
    para(doc,
         "Stawki podlegają corocznej waloryzacji w dniu 1 stycznia o wskaźnik HICP UE "
         "(Eurostat) za poprzedni rok kalendarzowy. HICP ≤ 0% oznacza brak zmiany stawek. "
         "Pierwsza waloryzacja następuje 1 stycznia roku następującego po zawarciu Umowy.")
    h2(doc, "C.3  Obliczanie Premii Rozwojowej i Opłaty Subskrypcyjnej")
    para(doc,
         "Premia Rozwojowa = Stawka (25% lub 30%) × Cena Katalogowa\n"
         "Gdzie: Cena Katalogowa = (Stawka MW × MW) + (Stawka MWh × MWh)\n"
         "Opłata Subskrypcyjna = 15% × Premia Rozwojowa (rocznie, od pierwszej rocznicy uruchomienia)")
    doc.add_paragraph()
    para(doc, "Uzgodniono i podpisano przez wszystkie Strony:")
    doc.add_paragraph()
    for name in [
        "Lighthief International — Pan Arkadiusz Sybaris",
        "Voltus Energy Sp. z o.o. — Pan Marcin Szumiło",
        "Voltus Energy Sp. z o.o. — Pan Andrzej Lechowicz",
    ]:
        para(doc, f"{name}  ________________________________     Data: __________________")

    # ── ZAŁĄCZNIK D ───────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "ZAŁĄCZNIK D — ANEKS WSPARCIA OPROGRAMOWANIA")
    para(doc,
         "Niniejszy Aneks stanowi część Umowy i reguluje obowiązki Voltus Energy Sp. z o.o. "
         "w zakresie wsparcia i konserwacji zgodnie z Art. 7.2.")
    h2(doc, "D.1  Klasyfikacja wagi i cele SLA")
    tbl(doc, [
        ["Waga", "Opis", "Przykłady", "Czas reakcji", "Czas rozwiązania"],
        ["P1 — Krytyczny", "System niedostępny lub krytyczna funkcja bezpieczeństwa uszkodzona",
         "EMS offline; błąd komendy DSO", "2 godziny", "24 godziny"],
        ["P2 — Poważny", "Podstawowa funkcja mocno ograniczona; brak obejścia",
         "Błąd SOC; awaria pętli sterowania", "8 godzin rob.", "5 dni rob."],
        ["P3 — Drobny", "Funkcja drugorzędna ograniczona; istnieje obejście",
         "Anomalia raportowania; defekt UI", "2 dni rob.", "30 dni rob."],
        ["P4 — Kosmetyczny", "Drobny problem; brak wpływu operacyjnego",
         "Formatowanie wyświetlacza", "5 dni rob.", "Kolejne wydanie"],
    ], col_widths=[2.5, 3.5, 3, 2.5, 3.5])
    doc.add_paragraph()
    h2(doc, "D.2  Cel dostępności")
    para(doc,
         "Chmurowe komponenty EMS utrzymują minimalną dostępność 99,5% mierzoną miesięcznie, "
         "z wyłączeniem planowanych okien konserwacyjnych. Planowane przerwy wymagają "
         "co najmniej 5 dni pisemnego powiadomienia i nie mogą przekraczać 4 godzin.")
    h2(doc, "D.3  Obowiązek aktualizacji regulacyjnych")
    para(doc,
         "Voltus implementuje zmiany wynikające z nowych lub zmienionych przepisów UE "
         "i krajowych kodeksów sieciowych w ciągu 90 dni od oficjalnej publikacji "
         "i powiadamia Spółkę w ciągu 14 dni od powzięcia wiadomości o takiej zmianie.")
    h2(doc, "D.4  Kredyty serwisowe")
    para(doc,
         "(a) Dostępność: Za każdy miesiąc poniżej 99,5% Spółka może potrącić z kolejnej "
         "Opłaty Subskrypcyjnej 5% tej opłaty za każdy pełny punkt procentowy poniżej 99,5%.")
    para(doc,
         "(b) Awarie P1: Za każdy incydent P1, przy którym przekroczono SLA, Voltus "
         "kredytuje Spółce 500 EUR za każdy dzień opóźnienia ponad określony SLA.")
    h2(doc, "D.5  Roczny przegląd")
    para(doc,
         "Cele SLA i wartości kredytów podlegają rocznemu przeglądowi na pisemny wniosek "
         "którejkolwiek ze Stron. Wszelkie zmiany wymagają pisemnej zgody wszystkich Stron.")

    # ── ZAPIS ─────────────────────────────────────────────────────────────────
    out = "/Volumes/T7 Grey/solinvest/DisperonEMS/docs/contract/DISPERON_SHA_v5_PL.docx"
    doc.save(out)
    print(f"Saved: {out}")


build()
