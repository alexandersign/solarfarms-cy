"""Generate DISPERON_SHA_v3_FINAL.docx — clean execution-ready contract."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE  = RGBColor(0x1F, 0x49, 0x7D)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x40, 0x40, 0x40)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def setup(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.2)
        section.right_margin  = Cm(3.2)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.bold = True
        run.font.size = Pt(13)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.bold = True
        run.font.size = Pt(11)


def para(doc, text, bold=False, italic=False, size=10, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK
    return p


def bul(doc, text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(level)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK


def tbl(doc, data, header_bg='1F497D', col_widths=None):
    cols = len(data[0])
    t = doc.add_table(rows=len(data), cols=cols)
    t.style = 'Table Grid'
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            c = t.rows[i].cells[j]
            if i == 0:
                set_cell_bg(c, header_bg)
            p = c.paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.bold = (i == 0)
            run.font.color.rgb = WHITE if i == 0 else BLACK
    if col_widths:
        for ci, w in enumerate(col_widths):
            set_col_width(t, ci, w)
    return t


def sig(doc, party_name, role, extra=None):
    p = doc.add_paragraph()
    p.add_run(f"{party_name}\n").font.bold = True
    run_role = p.add_run(f"{role}\n")
    run_role.font.size = Pt(10)
    if extra:
        run_extra = p.add_run(f"{extra}\n")
        run_extra.font.size = Pt(10)
    p.add_run("Signature: ___________________________     Date: __________________")
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()
    setup(doc)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("SHAREHOLDERS AGREEMENT\nAND INTELLECTUAL PROPERTY LICENCE")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = BLUE

    doc.add_paragraph()

    co = doc.add_paragraph()
    co.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = co.add_run(
        "LIGHTHIEF EUBESS LTD\n"
        "Registration No. HE 474192\n"
        "Lophitis Business Center I, Floor 2, Office 1\n"
        "28 Oktovriou & Aemiliou Chourmouziou, 3035 Limassol, Cyprus"
    )
    rc.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    dated = doc.add_paragraph()
    dated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = dated.add_run("Dated: ______________, 2025")
    rd.font.size = Pt(11)
    rd.font.italic = True

    doc.add_page_break()

    # ── PREAMBLE ──────────────────────────────────────────────────────────────
    h1(doc, "PREAMBLE")
    para(doc,
         'This Shareholders Agreement and Intellectual Property Licence (the "Agreement") is entered '
         'into on the date last signed below by and among the following parties:')
    doc.add_paragraph()

    parties = [
        ("PARTY 1 — LIGHTHIEF INTERNATIONAL",
         'A company duly incorporated and existing under applicable law, represented by its Director, '
         'Mr. Arkadiusz Sybaris ("Lighthief International" or "Majority Shareholder").'),
        ("PARTY 2 — LIGHTHIEF EUBESS LTD",
         'A private limited company registered in Cyprus under registration number HE 474192, having its '
         'registered office at 28 Oktovriou & Aemiliou Chourmouziou, Lophitis Business Center I, Floor 2, '
         'Office 1, 3035 Limassol, Cyprus, represented by its Director, Mr. Arkadiusz Sybaris (the "Company").'),
        ("PARTY 3 — MR. MARCIN SZUMIŁO",
         'An individual, holder of a Polish identity document, residing at an address notified to the '
         'Company ("Mr. Szumiło").'),
        ("PARTY 4 — MR. ANDRZEJ LECHOWICZ",
         'An individual, holder of a Polish identity document, residing at an address notified to the '
         'Company ("Mr. Lechowicz").'),
        ("PARTY 5 — MR. KAMIL TALAR",
         'An individual, holder of a Polish identity document, residing at an address notified to the '
         'Company ("Mr. Talar").'),
        ("PARTY 6 — MR. KOSTAS ALEXANDER PAPACOSTA",
         'An individual, holder of a Cypriot identity document, tax resident of the Republic of Cyprus, '
         'permanently residing in Cyprus, serving as Director of Lighthief EUBESS Ltd '
         '("Mr. Papacosta" or "Cyprus Director").'),
        ("PARTY 7 — VOLTUS ENERGY SP. Z O.O.",
         'A limited liability company incorporated and existing under the laws of the Republic of Poland, '
         'with its registered office at Gdański Park Naukowo-Technologiczny, Budynek B, Lokal 2.10.3, '
         'ul. Trzy Lipy 3, 80-172 Gdańsk, Poland, NIP: 1990133260, represented by its Management Board '
         'members, Mr. Marcin Szumiło and Mr. Andrzej Lechowicz ("Voltus").'),
    ]
    for name, desc in parties:
        p = doc.add_paragraph()
        r1 = p.add_run(name + "\n")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)
        r2.font.color.rgb = BLACK
        doc.add_paragraph()

    para(doc,
         'Lighthief International, Mr. Szumiło, Mr. Lechowicz, Mr. Talar, and Mr. Papacosta are referred '
         'to individually as a "Shareholder" and collectively as the "Shareholders". Mr. Szumiło and '
         'Mr. Lechowicz together with Voltus Energy Sp. z o.o. are referred to collectively as the '
         '"Voltus Parties". Parties 1 through 7 are collectively referred to as the "Parties".')

    # ── ARTICLE 1 — DEFINITIONS ───────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "ARTICLE 1 — DEFINITIONS")
    para(doc, "For the purposes of this Agreement, the following terms shall have the meanings assigned below:")
    doc.add_paragraph()

    definitions = [
        ('"Company"',
         'means Lighthief EUBESS Ltd, HE 474192, as described in the Preamble.'),
        ('"Shares"',
         'means the ordinary shares in the capital of the Company.'),
        ('"EMS Software" or "EMS"',
         'means the energy management system and SCADA software platform developed by Voltus Energy '
         'Sp. z o.o., together with all associated source code, object code, documentation, algorithms, '
         'interfaces, configurations, updates, enhancements, and derivative works, currently branded as '
         'Energy Copilot and described at https://voltusenergy.pl/produkt/.'),
        ('"DISPERON"',
         'means the commercial brand name and trademark under which the Company markets and sells EMS '
         'services and solutions.'),
        ('"Licence"',
         'means the software licence granted by Voltus Energy Sp. z o.o. to the Company pursuant to '
         'Article 5 of this Agreement.'),
        ('"Per-Project Licence Fee"',
         'means the fee payable per individual installation project prior to the Milestone Date, '
         'calculated in accordance with the Fee Schedule set out in Exhibit C.'),
        ('"Milestone"',
         'means the aggregate deployment of 500 MWh of energy storage capacity under the EMS Software '
         'pursuant to contracts concluded by or through the Company.'),
        ('"Milestone Date"',
         'means the date on which the Milestone is certified in writing by the Parties.'),
        ('"Perpetual Licence"',
         'means the irrevocable, non-exclusive, royalty-free, worldwide licence granted to the Company '
         'upon achievement of the Milestone as set out in Article 6.'),
        ('"Commissioning"',
         'means the on-site or remote technical installation, configuration, integration, grid code '
         'compliance verification, and handover of the EMS Software on a customer\'s BESS project.'),
        ('"Commissioning Fee"',
         'means the fee charged to the customer for Commissioning services, retained in full by the Company.'),
        ('"Software Margin"',
         'means the Software Licence Component of a customer contract, calculated as the applicable '
         'Per-MWh Licence Rate from Exhibit C multiplied by the installed MWh capacity of the relevant '
         'project. All customer invoices shall itemise the Software Licence Component and the '
         'Commissioning Component as separate line items. The Commissioning Component is excluded from '
         'Software Margin for all purposes of this Agreement.'),
        ('"Development Bonus"',
         'means the payment due from the Company to the Voltus Parties in respect of each Licence sold '
         'after the Milestone Date, calculated as 25% (twenty-five percent) of the Software Margin '
         'received from the relevant customer.'),
        ('"Intellectual Property" or "IP"',
         'means all patents, trademarks, copyrights, database rights, trade secrets, know-how, source '
         'code, and other proprietary rights, whether registered or unregistered.'),
        ('"Territory"',
         'means the European Union and its member states, and any other country or region agreed in '
         'writing by the Parties.'),
        ('"Exclusivity"',
         'means the right of the Company to be the sole authorised distributor of the EMS Software in '
         'a designated country, subject to Article 8.'),
        ('"BESS Manufacturer Partners"',
         'means the list of hardware manufacturers set out in Exhibit B, as agreed and signed by the '
         'Parties on the date of this Agreement.'),
        ('"Escrow Agent"',
         'means the independent third party appointed to hold the Escrow Materials pursuant to '
         'Article 6.4.'),
        ('"Escrow Materials"',
         'means the source code, build scripts, technical documentation, and deployment instructions '
         'for the then-current version of the EMS Software, as updated by Voltus pursuant to Article 6.4.'),
        ('"Software Support Annex"',
         'means the service level and maintenance schedule set out in Exhibit D, forming part of '
         'this Agreement.'),
    ]

    dt = doc.add_table(rows=len(definitions), cols=2)
    dt.style = 'Table Grid'
    for i, (term, defn) in enumerate(definitions):
        c0, c1 = dt.rows[i].cells[0], dt.rows[i].cells[1]
        if i % 2 == 0:
            set_cell_bg(c0, 'EEF3FB')
            set_cell_bg(c1, 'EEF3FB')
        p0 = c0.paragraphs[0]; p0.clear()
        r0 = p0.add_run(term); r0.font.bold = True; r0.font.size = Pt(9)
        p1 = c1.paragraphs[0]; p1.clear()
        r1 = p1.add_run(defn); r1.font.size = Pt(9); r1.font.color.rgb = BLACK
    set_col_width(dt, 0, 4.5)
    set_col_width(dt, 1, 10.5)
    doc.add_paragraph()

    # ── ARTICLE 2 — SHARE STRUCTURE ───────────────────────────────────────────
    h1(doc, "ARTICLE 2 — SHARE STRUCTURE AND TRANSFER")

    h2(doc, "2.1  Current Ownership")
    para(doc,
         "Immediately prior to execution of this Agreement, 100% of the issued share capital of "
         "the Company was held by Lighthief International.")

    h2(doc, "2.2  Transfer of Shares")
    para(doc,
         "Upon execution of this Agreement, the following share transfers are effected and agreed "
         "by all Parties:")
    tbl(doc, [
        ["Shareholder", "Shareholding", "Percentage (%)", "Purchase Consideration"],
        ["Lighthief International",         "65 shares", "65%", "N/A (Retained)"],
        ["Mr. Kostas Alexander Papacosta",  "10 shares", "10%", "EUR 250"],
        ["Mr. Marcin Szumiło",              "6 shares",  "6%",  "EUR 150"],
        ["Mr. Andrzej Lechowicz",           "6 shares",  "6%",  "EUR 150"],
        ["Mr. Kamil Talar",                 "13 shares", "13%", "EUR 325"],
    ], col_widths=[5, 3, 3, 4])
    doc.add_paragraph()

    h2(doc, "2.3  Purchase Consideration — Confirmed Payment")
    para(doc,
         "The Parties confirm and acknowledge that the purchase consideration of EUR 25 (twenty-five "
         "euros) per 1% shareholding has been received in full by the Company from each paying "
         "Shareholder prior to or simultaneously with the execution of this Agreement, as follows: "
         "(a) EUR 250 (two hundred and fifty euros) from Mr. Kostas Alexander Papacosta in respect "
         "of his 10% shareholding; (b) EUR 150 (one hundred and fifty euros) from each of "
         "Mr. Marcin Szumiło and Mr. Andrzej Lechowicz in respect of their respective 6% "
         "shareholdings; (c) EUR 325 (three hundred and twenty-five euros) from Mr. Kamil Talar in "
         "respect of his 13% shareholding. Written receipts have been issued by the Company to each "
         "paying Shareholder. The payment obligations under Article 2.2 are hereby acknowledged as "
         "fully discharged.")

    h2(doc, "2.4  Registration of Transfer")
    para(doc,
         "Mr. Arkadiusz Sybaris, acting in his capacity as Director of both Lighthief International "
         "and the Company, shall cause the necessary entries to be made in the Register of Members "
         "maintained by the Cyprus Registrar of Companies. Updated share certificates shall be issued "
         "to each Shareholder as soon as practicable following registration.")

    h2(doc, "2.5  Appointment of Shareholder Director")
    para(doc,
         "The Shareholders other than Lighthief International and Mr. Papacosta — being Mr. Marcin "
         "Szumiło, Mr. Andrzej Lechowicz, and Mr. Kamil Talar — shall, within 30 (thirty) days of "
         "execution of this Agreement, jointly designate one individual from among themselves to serve "
         "as a Director of the Company in accordance with Article 10.2. Written notice of such "
         "designation shall be delivered to Mr. Arkadiusz Sybaris, who shall thereupon cause the newly "
         "appointed Director to be registered with the Cyprus Registrar of Companies.")

    h2(doc, "2.6  Current Directors")
    para(doc,
         "As of the date of this Agreement, the duly appointed Directors of the Company are:\n"
         "(a) Mr. Arkadiusz Sybaris — Director and Secretary, tax resident of the Republic of Cyprus;\n"
         "(b) Mr. Kostas Alexander Papakosta — Director, tax resident of the Republic of Cyprus.\n"
         "Both existing Directors shall continue in their roles following execution of this Agreement "
         "unless otherwise resolved in accordance with Article 10.")

    h2(doc, "2.7  Share Certificates")
    para(doc,
         "Share certificates reflecting the revised shareholding shall be issued within 30 (thirty) "
         "days of registration with the Cyprus Registrar of Companies and shall thereupon be appended "
         "as Exhibit A and circulated to all Shareholders. Each Shareholder acknowledges that "
         "Exhibit A shall be completed and distributed as soon as practicable following registration.")

    # ── ARTICLE 3 — DISPERON BRAND ────────────────────────────────────────────
    h1(doc, "ARTICLE 3 — THE DISPERON BRAND")

    h2(doc, "3.1  Brand Ownership")
    para(doc,
         "The Company is the sole legal and beneficial owner of the DISPERON brand, trademark, and "
         "all associated intellectual property, including but not limited to the trade name, logo, "
         "domain name disperon.com, and any related trademarks registered or to be registered with "
         "the European Union Intellectual Property Office (EUIPO) under Nice Classes 09 and 42. "
         "All Shareholders acknowledge and confirm this ownership.")

    h2(doc, "3.2  Brand Purpose")
    para(doc,
         "DISPERON is the exclusive commercial brand under which the Company shall market, sell, and "
         "deliver EMS and SCADA services and solutions for Battery Energy Storage Systems (BESS) "
         "across the European Union and globally. All customer-facing materials, contracts, invoices, "
         "and marketing communications shall use the DISPERON brand.")

    h2(doc, "3.3  Brand Protection")
    para(doc,
         "No Shareholder shall use the DISPERON brand, logo, or any confusingly similar designation "
         "for any purpose other than the activities of the Company without the prior written consent "
         "of the Director. Any Shareholder who undertakes activities constituting brand dilution or "
         "misuse shall be liable to the Company for resulting damages.")

    # ── ARTICLE 4 — ROLES ─────────────────────────────────────────────────────
    h1(doc, "ARTICLE 4 — ROLES AND RESPONSIBILITIES OF SHAREHOLDERS")

    h2(doc, "4.1  Lighthief International")
    para(doc, "Lighthief International, as Majority Shareholder, shall:")
    for item in [
        "Hold 65% of the issued share capital of the Company and exercise corresponding voting rights;",
        "Provide strategic direction and oversight of the Company's operations;",
        "Develop and manage sales channels for the EMS under the DISPERON brand through its own "
        "offices and networks across Europe and beyond;",
        "Retain in full all Commissioning Fees generated from projects delivered by or through its "
        "infrastructure;",
        "Be responsible for investor relations, brand development, and international expansion of DISPERON;",
        "Provide corporate administration and registered office infrastructure through the Lighthief "
        "International group.",
    ]:
        bul(doc, item)

    h2(doc, "4.2  Mr. Kamil Talar")
    para(doc, "Mr. Kamil Talar shall:")
    for item in [
        "Hold 13% of the issued share capital of the Company in his own name and for his own account;",
        "Actively develop and manage sales channels for the DISPERON EMS in European markets, in "
        "coordination with Lighthief International;",
        "Deploy his existing business relationships, offices, and market presence — including his "
        "relationship with Linyang Energy Co., Ltd. and its customers — to generate BESS project pipeline;",
        "Report pipeline and sales activity to the Director on a monthly basis;",
        "Be jointly responsible with Lighthief International for achieving market sales targets.",
    ]:
        bul(doc, item)
    para(doc,
         "Performance Obligation: Mr. Talar shall, during each calendar year following execution of "
         "this Agreement: (a) maintain a documented active pipeline of not less than 100 MWh of BESS "
         "projects incorporating the EMS Software; and (b) introduce a minimum of 3 (three) qualified "
         "customer leads per year to the Company. Failure to meet these targets in any two consecutive "
         "annual periods shall entitle the Board to reclassify his shares as non-voting ordinary shares "
         "and to suspend dividend rights, subject to 30 (thirty) days' written notice and a 60 "
         "(sixty) day cure period.", indent=0.5)

    h2(doc, "4.3  Mr. Marcin Szumiło and Mr. Andrzej Lechowicz")
    para(doc,
         "Mr. Marcin Szumiło and Mr. Andrzej Lechowicz each individually hold 6% of the issued share "
         "capital of the Company in their own names and for their own respective accounts. Both "
         "Mr. Marcin Szumiło and Mr. Andrzej Lechowicz, individually and in their capacity as "
         "Management Board members of Voltus Energy Sp. z o.o., shall:")
    for item in [
        "Contribute the EMS Software developed by Voltus Energy Sp. z o.o. to the Company's "
        "commercial operations through the Licence as described in Article 5;",
        "Maintain primary responsibility for software development, technical architecture, product "
        "roadmap, and software quality of the EMS Software;",
        "Provide tier-2 technical support to the Company's commissioning and operations teams;",
        "Support the sales process through technical demonstrations, product presentations, and "
        "participation in client meetings and market visits as reasonably requested by the Director;",
        "Ensure the EMS Software remains compliant with applicable EU regulatory requirements, "
        "including NIS2, GDPR, and relevant grid code standards.",
    ]:
        bul(doc, item)

    h2(doc, "4.4  Mr. Kostas Alexander Papacosta")
    para(doc,
         "Mr. Kostas Alexander Papacosta holds 10% of the issued share capital of the Company in his "
         "own name and for his own account. Mr. Papacosta, in his capacity as Cyprus Managing Director "
         "and Director of the Company, shall:")
    for item in [
        "Serve as the primary executive director of Lighthief EUBESS Ltd in the Republic of Cyprus, "
        "holding all necessary registrations and authorisations required for the Company's Cyprus "
        "operations;",
        "Lead all engineering, procurement, and construction (EPC) delivery for BESS projects "
        "commissioned under the DISPERON brand in Cyprus and such other markets as the Board may "
        "designate;",
        "Manage all relationships with the Cyprus Distribution System Operator (EAC), Transmission "
        "System Operator (TSOC), and Energy Regulatory Authority (CERA), including DSO applications, "
        "SCADA integration approvals, and grid connection certifications;",
        "Act as the primary client relationship manager for the Company's Cyprus customer portfolio, "
        "including executing EPC and LTSA contracts on behalf of the Company;",
        "Oversee commissioning, IEC 60870-5-104 SCADA integration, and project handover for all "
        "BESS installations in Cyprus, coordinating with Voltus for EMS configuration and with "
        "Lighthief International for equipment supply;",
        "Report operational performance, project pipeline, and financial results to the Board on a "
        "monthly basis.",
    ]:
        bul(doc, item)
    para(doc,
         "Performance Obligation: Mr. Papacosta shall, during each calendar year following execution "
         "of this Agreement: (a) maintain the Company's Cyprus operating licences, DSO registrations, "
         "and SCADA approvals in good standing; and (b) achieve commissioning of not less than 10 MWh "
         "of BESS capacity under the DISPERON brand per calendar year from the second year of "
         "operations onwards. Persistent failure to maintain regulatory registrations shall constitute "
         "a material breach of his obligations as Director, subject to Board review.", indent=0.5)

    # ── ARTICLE 5 — IP ────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 5 — INTELLECTUAL PROPERTY")

    h2(doc, "5.1  Ownership of EMS Software IP")
    para(doc,
         "The Parties acknowledge and confirm that all Intellectual Property rights in and to the "
         "EMS Software are owned exclusively by Voltus Energy Sp. z o.o. Nothing in this Agreement "
         "shall be construed as a transfer, assignment, or novation of such IP rights to the Company "
         "or to any individual Shareholder, except as expressly provided in this Article 5 and "
         "Article 6.")

    h2(doc, "5.2  Licence Granted to the Company")
    para(doc,
         "Voltus Energy Sp. z o.o. hereby grants to the Company a non-exclusive, worldwide licence "
         "to use, deploy, configure, adapt, and sub-licence the EMS Software to customers under the "
         "DISPERON brand, subject to the commercial terms set out in Article 6 and the conditions of "
         "this Agreement.")

    h2(doc, "5.3  Voltus Retains Full IP")
    para(doc,
         "Voltus Energy Sp. z o.o. retains full and exclusive ownership of all IP in the EMS Software "
         "at all times, including after the grant of the Perpetual Licence pursuant to Article 6.3. "
         "The Perpetual Licence grants the Company a right to use the EMS Software but does not "
         "transfer ownership, title, or any other proprietary rights.")

    h2(doc, "5.4  Company IP and DISPERON Brand")
    para(doc,
         "All Intellectual Property created by or on behalf of the Company in connection with the "
         "DISPERON brand — including marketing materials, website content, implementation "
         "methodologies, customer documentation, and the DISPERON trademark — shall be owned "
         "exclusively by the Company. The Voltus Parties and individual Shareholders shall have no "
         "claim to such IP.")

    h2(doc, "5.5  IP Warranty and Indemnity")
    para(doc, "Voltus Energy Sp. z o.o. represents and warrants to the Company that:")
    for item in [
        "It is the sole and unencumbered owner of all Intellectual Property rights in and to the "
        "EMS Software;",
        "The EMS Software does not infringe any third-party intellectual property right, including "
        "patents, copyrights, trade secrets, or database rights;",
        "No open-source software components are included in the EMS Software under licences "
        "incompatible with commercial distribution or sub-licensing to end customers;",
        "The EMS Software complies with all applicable EU export control regulations.",
    ]:
        bul(doc, item)
    para(doc,
         "Voltus shall indemnify, defend, and hold harmless the Company and its directors, officers, "
         "and employees against any and all third-party claims, losses, damages, costs, and reasonable "
         "legal fees arising from any breach of the warranties set out in this Article 5.5.")

    # ── ARTICLE 6 — COMMERCIAL MODEL ─────────────────────────────────────────
    h1(doc, "ARTICLE 6 — COMMERCIAL MODEL AND LICENCE FEES")

    h2(doc, "6.1  Project-by-Project Revenue Model (Pre-Milestone)")
    para(doc,
         "(a) Per-Project Licence Fee: For each project, the Licence Fee shall be calculated in "
         "accordance with the Fee Schedule attached as Exhibit C. The applicable fee is the Per-MWh "
         "Licence Rate set out in Exhibit C multiplied by the project's installed MWh capacity. "
         "Voltus shall confirm the applicable Licence Fee in writing within 14 (fourteen) calendar "
         "days of receiving a written project brief from the Company. If Voltus fails to respond "
         "within 14 days, the fee shall be deemed confirmed at the applicable rate stated in Exhibit C. "
         "The Licence Fee shall be paid by the Company to Voltus within 30 (thirty) days of the "
         "commissioning date.")
    para(doc,
         "(b) Commissioning Fee: The Company sets and collects the Commissioning Fee directly from "
         "the customer. The Commissioning Fee is retained in full by the Company and does not form "
         "part of any Licence Fee calculation. Voltus Energy Sp. z o.o. has no entitlement to any "
         "portion of the Commissioning Fee, whether before or after the Milestone Date.")

    h2(doc, "6.2  Milestone Definition")
    para(doc,
         "The Milestone shall be deemed achieved when the aggregate installed capacity of BESS "
         "projects deploying the EMS Software under active contracts initiated by or through the "
         "Company reaches 500 MWh (five hundred megawatt-hours), calculated cumulatively from the "
         "date of execution of this Agreement. The following conditions apply:")
    for item in [
        "Each project shall be counted upon successful commissioning and formal handover to the "
        "customer, evidenced by a signed commissioning acceptance certificate;",
        "Projects contracted but subsequently cancelled prior to commissioning shall not count "
        "toward the Milestone;",
        "The Milestone has no time limit and shall be measured indefinitely from the date of this "
        "Agreement unless the Agreement is terminated;",
        "The Director shall maintain a written cumulative log of commissioned projects and MWh "
        "capacity, accessible to all Parties upon 5 (five) business days' written request.",
    ]:
        bul(doc, item)

    h2(doc, "6.3  Perpetual Licence upon Achievement of Milestone")
    para(doc,
         "Upon achievement of the Milestone, Voltus Energy Sp. z o.o. shall automatically grant to "
         "the Company a perpetual, irrevocable, non-exclusive, worldwide, royalty-free licence to "
         "use, deploy, adapt, and sub-licence the EMS Software in its then-current version and all "
         "subsequent versions and updates, without any ongoing licence fee obligation. This Perpetual "
         "Licence shall:")
    for item in [
        "Cover an unlimited number of customer installations worldwide;",
        "Apply to all versions of the EMS Software existing at the Milestone Date and all subsequent "
        "versions released by Voltus Energy Sp. z o.o.;",
        "Require no further payment from the Company to Voltus in respect of the use of the EMS "
        "Software itself;",
        "Survive any change of control, dissolution, or restructuring of Voltus Energy Sp. z o.o., "
        "unless directly caused by a material, uncured breach of this Agreement by the Company.",
    ]:
        bul(doc, item)

    h2(doc, "6.4  Source Code Escrow")
    para(doc,
         "(a) Deposit Obligation: Voltus shall, within 60 (sixty) days of execution of this "
         "Agreement, deposit the Escrow Materials with a mutually agreed independent Escrow Agent. "
         "The initial proposed Escrow Agent shall be NCC Group Escrow or such equivalent provider "
         "as the Parties may agree in writing within 30 days of execution.")
    para(doc,
         "(b) Maintenance: Voltus shall update the Escrow Materials within 30 (thirty) days of each "
         "material release or update to the EMS Software. Annual verification of the escrow contents "
         "shall be conducted by the Escrow Agent, at a cost shared equally between Voltus and the "
         "Company.")
    para(doc,
         "(c) Release Triggers: The Escrow Agent shall release the Escrow Materials to the Company "
         "upon the occurrence of any of the following events: (i) Voltus Energy Sp. z o.o. enters "
         "insolvency, administration, or liquidation proceedings and no successor entity assumes the "
         "maintenance obligations under this Agreement within 90 days; (ii) Voltus ceases to provide "
         "contracted maintenance or support for a continuous period exceeding 90 (ninety) days without "
         "cure; (iii) Voltus commits a material uncured breach of Article 7.2 persisting for more "
         "than 30 (thirty) days after written notice from the Company.")
    para(doc,
         "(d) Use of Released Materials: Following release of the Escrow Materials, the Company may "
         "use them solely to continue operating and maintaining the EMS Software for existing and "
         "future customers under the DISPERON brand. The Company shall not sub-licence the source "
         "code itself to any third party.")
    para(doc,
         "(e) Costs: The costs of the escrow arrangement shall be borne equally between Voltus and "
         "the Company.")

    h2(doc, "6.5  Development Bonus (Post-Milestone)")
    para(doc,
         "(a) Following the Milestone Date and for the duration of this Agreement, for each customer "
         "licence sold by the Company, the Company shall pay to the Voltus Parties — or to such legal "
         "entity or natural person as the Voltus Parties shall designate in writing — a Development "
         "Bonus of 25% (twenty-five percent) of the Software Margin received from that customer.")
    para(doc,
         "(b) The Development Bonus rate of 25% shall be subject to review every 3 (three) years, "
         "adjustable by written agreement of all Parties, within the range of 20% (twenty percent) "
         "to 30% (thirty percent), provided that Voltus produces a quarterly development expenditure "
         "summary demonstrating active research and development investment.")
    para(doc,
         "(c) The Development Bonus is calculated exclusively on the Software Margin component as "
         "defined in Article 1. Commissioning Fees and any other service fees are expressly excluded "
         "from the Development Bonus calculation and shall not form part of the basis of computation "
         "under any circumstances.")
    para(doc,
         "(d) The Development Bonus shall be designated by the Voltus Parties exclusively for the "
         "purposes of ongoing research, development, and enhancement of the EMS Software. The Voltus "
         "Parties shall provide the Company with a quarterly summary of development activities funded "
         "by such payments.")
    para(doc,
         "(e) The Development Bonus shall be payable within 45 (forty-five) days of the Company "
         "receiving cleared payment from the relevant customer.")

    h2(doc, "6.6  Records and Audit")
    para(doc,
         "The Company shall maintain accurate records of all project deployments, Licence Fees paid, "
         "Commissioning Fees received, and Development Bonus payments. Such records shall be available "
         "for inspection by the Voltus Parties and any Shareholder upon 10 (ten) business days' "
         "written notice.")

    # ── ARTICLE 7 — OPERATIONAL STRUCTURE ────────────────────────────────────
    h1(doc, "ARTICLE 7 — OPERATIONAL STRUCTURE")

    h2(doc, "7.1  Sales and Market Development")
    para(doc,
         "The Company, acting through Lighthief International, Mr. Kamil Talar, and Mr. Kostas "
         "Alexander Papacosta, shall be primarily responsible for the commercial development of the "
         "DISPERON brand and generation of customer pipeline across the Territory, including direct "
         "sales, partner channels, tender responses, EPC delivery, and industry representation. "
         "Mr. Papacosta shall lead all market development, client management, and EPC execution "
         "activities in the Republic of Cyprus.")

    h2(doc, "7.2  Technology and Product Development")
    para(doc,
         "Voltus Energy Sp. z o.o. shall be primarily responsible for the continuous development, "
         "maintenance, security, and improvement of the EMS Software, maintaining it in a commercially "
         "deployable and regulatory-compliant state at all times. Specifically, Voltus shall:")
    for item in [
        "Respond to severity-classified support requests within the timeframes set out in the "
        "Software Support Annex (Exhibit D);",
        "Maintain cloud-hosted EMS components at a minimum availability of 99.5% (ninety-nine point "
        "five percent) measured per calendar month, excluding planned maintenance windows notified "
        "at least 5 (five) calendar days in advance;",
        "Implement changes required by new or amended EU or national grid code regulations within "
        "90 (ninety) days of the official publication date of such requirements, and notify the "
        "Company within 14 (fourteen) days of becoming aware of any such regulatory change;",
        "Provide the Company with a quarterly product roadmap and development activity report;",
        "Maintain cybersecurity compliance with NIS2 and applicable EU standards at all times.",
    ]:
        bul(doc, item)
    para(doc,
         "Service credits for SLA failures are set out in Exhibit D. Persistent failure by Voltus "
         "to meet the obligations in this Article 7.2 shall constitute a material breach of this "
         "Agreement, entitling the Company to activate the escrow release mechanism set out in "
         "Article 6.4(c).")

    h2(doc, "7.3  Commissioning and Implementation")
    para(doc,
         "On-site and remote commissioning, system integration, and project handover services shall "
         "be performed by the Company or a designated entity within the Lighthief International "
         "group. Voltus shall provide technical support, documentation, and training as reasonably "
         "required. All Commissioning revenues belong to the Company.")

    h2(doc, "7.4  Customer Contracts")
    para(doc,
         "All customer contracts shall be entered into by the Company. The Company is the legal "
         "counterparty to all end customers. Voltus Energy Sp. z o.o. acts as software supplier "
         "to the Company, not to end customers directly, unless otherwise agreed in writing.")

    # ── ARTICLE 8 — EXCLUSIVITY ───────────────────────────────────────────────
    h1(doc, "ARTICLE 8 — EXCLUSIVITY AND MARKET MODEL")

    h2(doc, "8.1  Principle of Earned Exclusivity")
    para(doc,
         "Exclusivity is not granted automatically. It may be earned by the Company in specific "
         "countries on the basis of demonstrated commercial activity and achieved volume thresholds, "
         "in accordance with this Article.")

    h2(doc, "8.2  Pre-Conditions for Exclusivity")
    para(doc,
         "Before any exclusivity is recognised in a specific country, the Company shall provide "
         "Voltus Energy Sp. z o.o. with:")
    for item in [
        "A documented project pipeline for that country, including client names, project sizes, "
        "and estimated timelines;",
        "A written sales plan covering a minimum of 12 (twelve) months;",
        "Evidence of active commercial engagement, including customer meetings, proposals, or "
        "signed letters of intent.",
    ]:
        bul(doc, item)

    h2(doc, "8.3  Exclusivity Volume Thresholds")
    para(doc,
         "Exclusivity shall be granted upon the Company achieving the following minimum installed "
         "capacity within 24 (twenty-four) months of the Company's first commercial project in "
         "that country:")
    tbl(doc, [
        ["Market Category", "Countries", "Minimum Volume", "Annual Maintenance"],
        ["Large markets",
         "Germany, Italy, Spain, France",
         "500 MWh within 24 months",
         "100 MWh / year"],
        ["Mid-sized markets",
         "Poland, Netherlands, Romania, Greece",
         "250 MWh within 24 months",
         "50 MWh / year"],
        ["Smaller markets",
         "Cyprus, other EU states",
         "100 MWh within 24 months",
         "20 MWh / year"],
    ], col_widths=[3.5, 4, 4, 3.5])
    doc.add_paragraph()

    h2(doc, "8.4  Maintenance of Exclusivity")
    para(doc,
         "Following the initial grant of exclusivity, the Company must achieve the minimum annual "
         "installed capacity set out in the Annual Maintenance column of Article 8.3 in each "
         "subsequent 12-month period to retain exclusivity in that country. Failure to meet this "
         "maintenance threshold for two consecutive annual periods shall result in automatic lapsing "
         "of exclusivity for that country.")

    h2(doc, "8.5  Initial Country Limit")
    para(doc,
         "At the time of execution of this Agreement, the Company may seek exclusivity in a maximum "
         "of 5 (five) countries simultaneously. Additional countries may be added as exclusivity is "
         "achieved and maintained.")

    h2(doc, "8.6  Voltus Direct Rights — BESS Manufacturer Partners")
    para(doc,
         "Notwithstanding any exclusivity granted to the Company, Voltus Energy Sp. z o.o. retains "
         "the right to supply the EMS Software directly to BESS Manufacturers listed in Exhibit B "
         "of this Agreement. This right:")
    for item in [
        "Applies only to the manufacturers expressly listed in Exhibit B, which shall be completed, "
        "signed, and attached by the Parties on the date of execution of this Agreement;",
        "Does not extend to any other hardware manufacturer or distributor not listed in Exhibit B "
        "unless the Parties agree in writing to amend Exhibit B;",
        "Is not subject to any exclusivity threshold or geographic restriction in respect of the "
        "listed BESS Manufacturer Partners.",
    ]:
        bul(doc, item)
    para(doc,
         "BESS Manufacturer Partner Referral Obligation: Notwithstanding the foregoing, where a "
         "project originating from a BESS Manufacturer Partner has been the subject of prior "
         "commercial engagement by the Company — evidenced by a written proposal, meeting record, "
         "or registered customer entry in the Company's pipeline log — Voltus shall, within 5 "
         "(five) business days of receiving an enquiry from such manufacturer in respect of that "
         "project, notify the Company and refer such project to the Company under the DISPERON "
         "brand. The Parties shall agree on applicable commercial terms within 15 (fifteen) "
         "business days of such referral. In the absence of agreement within that period, the "
         "standard Development Bonus rate set out in Article 6.5 shall apply.")
    para(doc,
         "Voltus shall provide the Company with a monthly written summary of any EMS Software "
         "supply activities directed to BESS Manufacturer Partners, including project name, country, "
         "MWh capacity, and customer identity to the extent permitted by applicable confidentiality "
         "obligations.")

    h2(doc, "8.7  Transparency and Customer Registration")
    para(doc,
         "Prior to any formal exclusivity being granted in a given country:")
    for item in [
        "Each Party may register customers in any country within the Territory;",
        "A customer shall be assigned to the Party actively leading the sales process, confirmed "
        "in writing;",
        "In the event of dispute over customer assignment, the Director shall resolve the matter "
        "within 15 (fifteen) business days.",
    ]:
        bul(doc, item)

    h2(doc, "8.8  Non-Competition Undertaking of Voltus")
    para(doc,
         "Voltus Energy Sp. z o.o. undertakes, for the duration of this Agreement, that:")
    for item in [
        "It shall maintain full transparency with the Company regarding its direct sales activities "
        "within the Territory;",
        "If a customer approaches Voltus Energy Sp. z o.o. directly for EMS services, and that same "
        "customer has previously been approached by or is in active discussion with the Company, "
        "Voltus Energy Sp. z o.o. shall decline to engage independently and shall refer such "
        "customer to the Company under the DISPERON brand;",
        "The above undertaking does not apply to customers originally introduced by Voltus Energy "
        "Sp. z o.o. independently of the Company, or to customers of BESS Manufacturer Partners "
        "listed in Exhibit B where the BESS Manufacturer Partner Referral Obligation in Article 8.6 "
        "does not apply.",
    ]:
        bul(doc, item)

    # ── ARTICLE 9 — TRANSFER OF SHARES ───────────────────────────────────────
    h1(doc, "ARTICLE 9 — TRANSFER OF SHARES")

    h2(doc, "9.1  Lock-Up Period")
    para(doc,
         "For 24 (twenty-four) months from the date of this Agreement (the \"Lock-Up Period\"), no "
         "Shareholder shall sell, transfer, assign, pledge, or otherwise dispose of any Shares "
         "without the prior written consent of Lighthief International.")

    h2(doc, "9.2  Right of First Refusal")
    para(doc,
         "Following the Lock-Up Period, if any Shareholder proposes to transfer Shares to a third "
         "party: written notice must be given to all other Shareholders, specifying the number of "
         "Shares, proposed price, and buyer identity; each non-selling Shareholder has 30 (thirty) "
         "days to exercise a right of first refusal on a pro-rata basis at the same price and terms; "
         "remaining unexercised Shares may be transferred to the third party within 90 (ninety) days "
         "at no less than the offered price.")

    h2(doc, "9.3  Drag-Along Right")
    para(doc,
         "If Lighthief International proposes to sell 100% of its Shares to a bona fide purchaser "
         "requiring 100% of the issued capital, Lighthief International may require all other "
         "Shareholders to sell their Shares on the same terms, on 30 (thirty) days' written notice.")

    h2(doc, "9.4  Tag-Along Right")
    para(doc,
         "If Lighthief International proposes to transfer more than 50% of its Shares to a third "
         "party, each other Shareholder has the right, within 20 (twenty) days, to include their "
         "Shares in such sale on the same terms.")

    # ── ARTICLE 10 — GOVERNANCE ───────────────────────────────────────────────
    h1(doc, "ARTICLE 10 — GOVERNANCE AND BOARD OF DIRECTORS")

    h2(doc, "10.1  Current Directors")
    para(doc,
         "As of the date of this Agreement, the duly appointed Directors of the Company are:\n"
         "(a) Mr. Arkadiusz Sybaris — Director and Secretary, tax resident of the Republic of Cyprus, "
         "representing Lighthief International;\n"
         "(b) Mr. Kostas Alexander Papacosta — Director and Cyprus Managing Director, tax resident "
         "of the Republic of Cyprus, permanently residing in Cyprus.")

    h2(doc, "10.2  Board Composition")
    para(doc,
         "The Board of Directors of the Company shall be composed as follows:\n"
         "(a) Lighthief International, as Majority Shareholder, shall have the right to appoint up "
         "to 3 (three) Directors to the Board, including the Director and Secretary role. "
         "Mr. Kostas Alexander Papacosta, as Cyprus Managing Director, shall be one of the Directors "
         "appointed by Lighthief International for the duration of this Agreement;\n"
         "(b) The Shareholders other than Lighthief International and Mr. Papacosta — being "
         "Mr. Marcin Szumiło, Mr. Andrzej Lechowicz, and Mr. Kamil Talar jointly — shall have the "
         "right to appoint 1 (one) Director from among themselves to the Board.\n"
         "The Shareholder Director appointed under (b) above shall be the person designated "
         "pursuant to Article 2.5.")

    h2(doc, "10.3  Appointment and Removal")
    para(doc,
         "Directors appointed by Lighthief International may be appointed and removed at any time "
         "by written resolution of Lighthief International, provided that the removal of "
         "Mr. Papacosta as Director shall additionally require a resolution of the Board. "
         "The Shareholder Director may be replaced by written agreement of Mr. Marcin Szumiło, "
         "Mr. Andrzej Lechowicz, and Mr. Kamil Talar. Any appointment or removal shall take effect "
         "upon registration with the Cyprus Registrar of Companies.")

    h2(doc, "10.4  Decisions Requiring Majority Shareholder Consent")
    para(doc,
         "The following matters shall require the written consent of Shareholders holding a majority "
         "of Shares:")
    for item in [
        "Amendment of the Company's Articles of Association;",
        "Issuance of new shares or share capital increases;",
        "Material change to the Company's principal business activity;",
        "Entry into any related-party transaction exceeding EUR 50,000 in value;",
        "Appointment or removal of auditors.",
    ]:
        bul(doc, item)

    h2(doc, "10.5  Reserved Matters Requiring 85% Shareholder Consent")
    para(doc,
         "The following matters shall require the written consent of Shareholders holding not less "
         "than 85% (eighty-five percent) of all issued Shares:")
    for item in [
        "Material amendment to the commercial model set out in Article 6, including the Development "
        "Bonus rate, the Software Margin definition, or the fee schedule in Exhibit C;",
        "Voluntary dissolution or winding up of the Company;",
        "Any transaction or resolution that would impair, encumber, or terminate the Perpetual "
        "Licence or the source code escrow arrangement under Article 6.4.",
    ]:
        bul(doc, item)

    h2(doc, "10.6  Shareholder Meetings")
    para(doc,
         "The Company shall hold at least one annual meeting of Shareholders. Additional meetings "
         "may be called by the Director or by any Shareholder holding at least 20% of issued capital, "
         "on not less than 14 (fourteen) days' notice. Meetings may be held in person or by video "
         "conference.")

    # ── ARTICLE 11 — CONFIDENTIALITY ─────────────────────────────────────────
    h1(doc, "ARTICLE 11 — CONFIDENTIALITY")

    h2(doc, "11.1  Confidential Information")
    para(doc,
         "Each Party shall keep confidential all non-public information received in connection with "
         "this Agreement, including technical specifications, source code, customer data, financial "
         'terms, and business strategies ("Confidential Information").')

    h2(doc, "11.2  Duration")
    para(doc,
         "The confidentiality obligation applies during the term of this Agreement and for 5 (five) "
         "years following its termination or expiry.")

    h2(doc, "11.3  Permitted Disclosure")
    para(doc,
         "Confidential Information may be disclosed only to employees, contractors, or advisors with "
         "a need to know who are bound by equivalent obligations, or where required by applicable law.")

    # ── ARTICLE 12 — GOVERNING LAW ────────────────────────────────────────────
    h1(doc, "ARTICLE 12 — GOVERNING LAW AND DISPUTE RESOLUTION")

    h2(doc, "12.1  Governing Law")
    para(doc,
         "This Agreement shall be governed by and construed in accordance with the laws of the "
         "Republic of Cyprus.")

    h2(doc, "12.2  Negotiation")
    para(doc,
         "In the event of any dispute arising out of or in connection with this Agreement (a "
         '"Dispute"), the Parties shall attempt resolution through good-faith negotiation within '
         "30 (thirty) days of written notice of the Dispute.")

    h2(doc, "12.3  Arbitration")
    para(doc,
         "If a Dispute is not resolved by negotiation within the period set out in Article 12.2, "
         "it shall be referred to final and binding arbitration in Limassol, Cyprus, under the "
         "Cyprus Arbitration Law (Cap. 4), conducted in the English language.")

    # ── ARTICLE 13 — GENERAL PROVISIONS ──────────────────────────────────────
    h1(doc, "ARTICLE 13 — GENERAL PROVISIONS")

    h2(doc, "13.1  Entire Agreement")
    para(doc,
         "This Agreement constitutes the entire agreement of the Parties with respect to its subject "
         "matter and supersedes all prior understandings, negotiations, and representations.")

    h2(doc, "13.2  Amendment")
    para(doc,
         "This Agreement may only be amended by a written instrument signed by all Parties.")

    h2(doc, "13.3  Severability")
    para(doc,
         "If any provision of this Agreement is found to be invalid, illegal, or unenforceable, the "
         "remaining provisions shall continue in full force and effect.")

    h2(doc, "13.4  Waiver")
    para(doc,
         "No failure or delay by any Party in exercising any right or remedy under this Agreement "
         "shall constitute a waiver thereof.")

    h2(doc, "13.5  Notices")
    para(doc,
         "All notices under this Agreement shall be in writing, delivered by email with read receipt "
         "or by registered post to the addresses stated in the Preamble or as otherwise notified in "
         "writing.")

    h2(doc, "13.6  Counterparts and Electronic Signatures")
    para(doc,
         "This Agreement may be executed in counterparts, each of which shall constitute an original. "
         "Electronic signatures shall be valid and binding on all Parties. The English version of this "
         "Agreement shall prevail over any translation.")

    # ── SIGNATURES ────────────────────────────────────────────────────────────
    doc.add_page_break()
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sig = p_sig.add_run("SIGNATURES")
    r_sig.font.size = Pt(14); r_sig.font.bold = True; r_sig.font.color.rgb = BLUE

    doc.add_paragraph()
    para(doc,
         "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written "
         "above.", italic=True)
    doc.add_paragraph()

    sig(doc, "LIGHTHIEF INTERNATIONAL",
        "Represented by: Mr. Arkadiusz Sybaris, Director")
    sig(doc, "LIGHTHIEF EUBESS LTD (HE 474192)",
        "Represented by: Mr. Arkadiusz Sybaris, Director and Secretary")
    sig(doc, "MR. KOSTAS ALEXANDER PAPACOSTA",
        "Individual Shareholder and Cyprus Managing Director")
    sig(doc, "MR. MARCIN SZUMIŁO",
        "Individual Shareholder")
    sig(doc, "MR. ANDRZEJ LECHOWICZ",
        "Individual Shareholder")
    sig(doc, "MR. KAMIL TALAR",
        "Individual Shareholder")
    sig(doc, "VOLTUS ENERGY SP. Z O.O.",
        "Represented by: Mr. Marcin Szumiło, Management Board Member",
        extra="Represented by: Mr. Andrzej Lechowicz, Management Board Member")

    # ── EXHIBIT A ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "EXHIBIT A — SHARE CERTIFICATES")
    para(doc,
         "The following share certificates of Lighthief EUBESS Ltd (HE 474192) shall be issued "
         "within 30 (thirty) days of registration with the Cyprus Registrar of Companies and "
         "appended hereto:")
    for item in [
        "Current share certificate issued to Lighthief International (100% prior to transfer);",
        "New share certificate issued to Mr. Kostas Alexander Papacosta (10%);",
        "New share certificate issued to Mr. Marcin Szumiło (6%);",
        "New share certificate issued to Mr. Andrzej Lechowicz (6%);",
        "New share certificate issued to Mr. Kamil Talar (13%);",
        "Updated share certificate issued to Lighthief International (65% post-transfer).",
    ]:
        bul(doc, item)

    # ── EXHIBIT B ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "EXHIBIT B — BESS MANUFACTURER PARTNERS")
    para(doc,
         "The following is the complete and exclusive list of BESS hardware manufacturers in respect "
         "of whom Voltus Energy Sp. z o.o. retains direct supply rights pursuant to Article 8.6. "
         "Only the manufacturers listed below are covered by Article 8.6. No other manufacturer or "
         "entity shall be added to this list without the prior written consent of all Parties and a "
         "signed amendment to this Agreement.")
    doc.add_paragraph()
    tbl(doc, [
        ["No.", "Manufacturer Name", "Registered Country", "Date Added"],
        ["1.", "Linyang Energy Co., Ltd.", "People's Republic of China", "Date of Agreement"],
        ["2.", "________________________", "________________________",   "________________"],
        ["3.", "________________________", "________________________",   "________________"],
        ["4.", "________________________", "________________________",   "________________"],
    ], col_widths=[1, 5.5, 4.5, 4])
    doc.add_paragraph()
    para(doc,
         "The above list has been agreed, confirmed, and signed by all Parties on the date of "
         "execution of this Agreement. Any manufacturer not listed above shall not benefit from the "
         "direct supply rights set out in Article 8.6.")
    doc.add_paragraph()
    para(doc, "Confirmed by Voltus Energy Sp. z o.o.:")
    para(doc, "Mr. Marcin Szumiło  ________________________________")
    para(doc, "Mr. Andrzej Lechowicz  ________________________________")
    para(doc, "Date: __________________")
    doc.add_paragraph()
    para(doc, "Confirmed by Lighthief EUBESS Ltd:")
    para(doc, "Mr. Arkadiusz Sybaris  ________________________________")
    para(doc, "Date: __________________")

    # ── EXHIBIT C — FEE SCHEDULE ──────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "EXHIBIT C — LICENCE FEE SCHEDULE")
    para(doc,
         "This Exhibit C forms part of the Agreement and governs all Per-Project Licence Fee "
         "calculations under Article 6.1 and the Software Margin calculation under Article 6.5.")

    h2(doc, "C.1  Per-MWh Licence Rate (Pre-Milestone)")
    tbl(doc, [
        ["Project Scale (MWh installed)", "Per-MWh Licence Rate", "Effective Date"],
        ["≤ 10 MWh",                      "EUR [___] / MWh",      "Date of Agreement"],
        ["10 MWh – 50 MWh",              "EUR [___] / MWh",      "Date of Agreement"],
        ["> 50 MWh",                      "EUR [___] / MWh (negotiated case-by-case)",
         "Date of Agreement"],
    ], col_widths=[5, 5, 5])
    doc.add_paragraph()
    para(doc,
         "The rates in brackets above shall be agreed by the Parties within 14 (fourteen) days of "
         "execution of this Agreement and inserted by written addendum signed by all Parties. "
         "Until such agreement is reached, no project brief shall be deemed submitted for the "
         "purposes of Article 6.1.")

    h2(doc, "C.2  Annual Escalation")
    para(doc,
         "The Per-MWh Licence Rate shall be adjusted annually on 1 January of each year by the "
         "change in the EU Harmonised Index of Consumer Prices (HICP) published by Eurostat for "
         "the preceding calendar year. The first adjustment shall occur on 1 January of the year "
         "following execution of this Agreement.")

    h2(doc, "C.3  14-Day Confirmation Procedure")
    para(doc,
         "Voltus shall confirm the applicable Licence Fee for each submitted project brief within "
         "14 (fourteen) calendar days of receipt. A project brief shall include: project name, "
         "location, system size in MWh, customer name, and estimated commissioning date. Failure "
         "by Voltus to respond within 14 days constitutes deemed approval at the applicable "
         "schedule rate.")

    h2(doc, "C.4  Software Margin Calculation")
    para(doc,
         "The Software Margin for the purposes of Article 6.5 shall be calculated as: "
         "Applicable Per-MWh Rate × Project MWh Capacity. This amount shall appear as a separate "
         "line item on the customer invoice, labelled 'EMS Software Licence'. The Commissioning "
         "Component shall appear as a separate line item. No other allocation methodology shall "
         "be used without the written agreement of all Parties.")

    doc.add_paragraph()
    para(doc, "Agreed and signed by all Parties:")
    doc.add_paragraph()
    for name in ["Lighthief International — Mr. Arkadiusz Sybaris",
                 "Voltus Energy Sp. z o.o. — Mr. Marcin Szumiło",
                 "Voltus Energy Sp. z o.o. — Mr. Andrzej Lechowicz"]:
        para(doc, f"{name}  ________________________________     Date: __________________")

    # ── EXHIBIT D — SOFTWARE SUPPORT ANNEX ───────────────────────────────────
    doc.add_page_break()
    h1(doc, "EXHIBIT D — SOFTWARE SUPPORT ANNEX")
    para(doc,
         "This Software Support Annex forms part of the Agreement and governs Voltus Energy "
         "Sp. z o.o.'s support and maintenance obligations under Article 7.2.")

    h2(doc, "D.1  Severity Classification and SLA Targets")
    tbl(doc, [
        ["Severity", "Description", "Examples", "Response SLA", "Resolution SLA"],
        ["P1 — Critical",
         "System completely unavailable or safety-critical function impaired",
         "EMS offline; DSO command failure",
         "2 hours",
         "24 hours"],
        ["P2 — Major",
         "Core function severely degraded; no workaround available",
         "SOC reporting error; control loop failure",
         "8 business hours",
         "5 business days"],
        ["P3 — Minor",
         "Non-critical function impaired; workaround available",
         "Reporting anomaly; UI defect",
         "2 business days",
         "30 business days"],
        ["P4 — Cosmetic",
         "Minor issue; no operational impact",
         "Display formatting; translation",
         "5 business days",
         "Next planned release"],
    ], col_widths=[2.5, 3.5, 3.5, 2.5, 3])
    doc.add_paragraph()

    h2(doc, "D.2  Availability Target")
    para(doc,
         "Cloud-hosted EMS components shall maintain a minimum availability of 99.5% (ninety-nine "
         "point five percent) measured per calendar month, excluding scheduled maintenance windows. "
         "Scheduled maintenance windows shall require a minimum of 5 (five) calendar days' advance "
         "written notice to the Company and shall not exceed 4 (four) hours per occurrence.")

    h2(doc, "D.3  Regulatory Update Obligation")
    para(doc,
         "Voltus shall implement changes required by new or amended EU or national grid code "
         "regulations within 90 (ninety) days of the official publication date of such requirements. "
         "Voltus shall notify the Company in writing within 14 (fourteen) days of becoming aware "
         "of any regulatory change that affects the EMS Software.")

    h2(doc, "D.4  Service Credits")
    para(doc,
         "(a) Availability: In the event monthly availability falls below 99.5%, the Company may "
         "offset from the next Licence Fee or Development Bonus payment an amount equal to 5% "
         "(five percent) of such payment for each full percentage point below 99.5%.")
    para(doc,
         "(b) P1 SLA Failures: For each P1 incident where the response or resolution SLA is "
         "exceeded, Voltus shall credit the Company EUR 500 (five hundred euros) per day of "
         "delay beyond the applicable SLA.")
    para(doc,
         "(c) Cap and Remedy: Credits are the Company's sole financial remedy for SLA failures "
         "under this Annex and shall not limit any other right or remedy available to the Company "
         "for material breach of the Agreement.")

    h2(doc, "D.5  Annual Review")
    para(doc,
         "The SLA targets and credit values set out in this Annex shall be reviewed annually at "
         "the written request of either Party. Any amendment shall require the written agreement "
         "of all Parties.")

    # ─── SAVE ─────────────────────────────────────────────────────────────────
    out = "/Volumes/T7 Grey/solinvest/DisperonEMS/docs/contract/DISPERON_SHA_v3_FINAL_rev2.docx"
    doc.save(out)
    print(f"Saved: {out}")

build()
