"""Generate DISPERON_SHA_v3_PROPOSED.docx — clean final contract with all fixes."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE   = RGBColor(0x1F, 0x49, 0x7D)
BLACK  = RGBColor(0x00, 0x00, 0x00)
GREY   = RGBColor(0x40, 0x40, 0x40)
GREEN  = RGBColor(0x1A, 0x6B, 0x30)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)

def set_doc_defaults(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = BLACK
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

def heading(doc, text, level=1, colour=BLUE, centre=False):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centre else WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = colour
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(13)
        elif level == 2:
            run.font.size = Pt(11)

def body(doc, text, bold=False, italic=False, colour=BLACK, size=10, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = colour
    return p

def bullet(doc, text, indent=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    return p

def note(doc, text):
    """New/amended clause annotation."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"▶  NEW / AMENDED (v3):  {text}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = GREEN
    return p

def simple_table(doc, data, header_bg='1F497D'):
    cols = len(data[0])
    tbl = doc.add_table(rows=len(data), cols=cols)
    tbl.style = 'Table Grid'
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            c = tbl.rows[i].cells[j]
            if i == 0:
                set_cell_bg(c, header_bg)
            p = c.paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.bold = (i == 0)
            run.font.color.rgb = WHITE if i == 0 else BLACK
    return tbl

def sig_block(doc, name, role):
    p = doc.add_paragraph()
    p.add_run(f"{name}\n").font.bold = True
    p.add_run(f"{role}\n")
    p.add_run("Signature: ___________________________\n")
    p.add_run("Date: ________________________________\n")
    doc.add_paragraph()

# ── BUILD ────────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    set_doc_defaults(doc)

    # ── COVER ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SHAREHOLDERS AGREEMENT\nAND INTELLECTUAL PROPERTY LICENCE")
    run.font.size = Pt(18); run.font.bold = True; run.font.color.rgb = BLUE

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("LIGHTHIEF EUBESS LTD\nRegistration No. HE 474192\n"
                   "Lophitis Business Center I, Floor 2, Office 1\n"
                   "28 Oktovriou & Aemiliou Chourmouziou, 3035 Limassol, Cyprus")
    r.font.size = Pt(11)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("VERSION 3  ·  Dated: ______________, 2025")
    r3.font.size = Pt(10); r3.font.italic = True; r3.font.color.rgb = GREY

    doc.add_paragraph()
    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note_p.add_run(
        "This is the proposed final version (v3). Clauses marked ▶ NEW / AMENDED (v3) "
        "address issues identified in the DISPERON SHA v2 Issues Register."
    )
    nr.font.size = Pt(9); nr.font.italic = True; nr.font.color.rgb = GREEN

    doc.add_page_break()

    # ── PREAMBLE ───────────────────────────────────────────────────────────
    heading(doc, "PREAMBLE")
    body(doc,
         "This Shareholders Agreement and Intellectual Property Licence (the \"Agreement\") is entered "
         "into on the date last signed below by and among the following parties:")
    body(doc,
         "PARTY 1 — LIGHTHIEF INTERNATIONAL\n"
         "A company duly incorporated and existing under applicable law, represented by its Director, "
         "Mr. Arkadiusz Sybaris (\"Lighthief International\" or \"Majority Shareholder\").", bold=False)
    body(doc,
         "PARTY 2 — LIGHTHIEF EUBESS LTD\n"
         "A private limited company registered in Cyprus under registration number HE 474192, having its "
         "registered office at 28 Oktovriou & Aemiliou Chourmouziou, Lophitis Business Center I, Floor 2, "
         "Office 1, 3035 Limassol, Cyprus, represented by its Director, Mr. Arkadiusz Sybaris (the \"Company\").")
    body(doc, "PARTY 3 — MR. MARCIN SZUMIŁO\nAn individual, holder of a Polish identity document, "
         "residing at an address notified to the Company (\"Mr. Szumiło\").")
    body(doc, "PARTY 4 — MR. ANDRZEJ LECHOWICZ\nAn individual, holder of a Polish identity document, "
         "residing at an address notified to the Company (\"Mr. Lechowicz\").")
    body(doc, "PARTY 5 — MR. KAMIL TALAR\nAn individual, holder of a Polish identity document, "
         "residing at an address notified to the Company (\"Mr. Talar\").")
    body(doc, "PARTY 6 — VOLTUS ENERGY SP. Z O.O.\nA limited liability company incorporated and existing "
         "under the laws of the Republic of Poland, with its registered office at Gdański Park "
         "Naukowo-Technologiczny, Budynek B, Lokal 2.10.3, ul. Trzy Lipy 3, 80-172 Gdańsk, Poland, "
         "NIP: 1990133260, represented by its Management Board members, Mr. Marcin Szumiło and "
         "Mr. Andrzej Lechowicz (\"Voltus\").")
    body(doc, "Lighthief International, Mr. Szumiło, Mr. Lechowicz, and Mr. Talar are referred to "
         "individually as a \"Shareholder\" and collectively as the \"Shareholders\". Mr. Szumiło and "
         "Mr. Lechowicz together with Voltus Energy Sp. z o.o. are referred to collectively as the "
         "\"Voltus Parties\". Parties 1 through 6 are collectively referred to as the \"Parties\".")

    # ── ARTICLE 1 — DEFINITIONS ────────────────────────────────────────────
    heading(doc, "ARTICLE 1 — DEFINITIONS")
    defs = [
        ("\"Company\"", "means Lighthief EUBESS Ltd, HE 474192, as described in the Preamble."),
        ("\"Shares\"", "means the ordinary shares in the capital of the Company."),
        ("\"EMS Software\" or \"EMS\"",
         "means the energy management system and SCADA software platform developed by Voltus Energy "
         "Sp. z o.o., together with all associated source code, object code, documentation, algorithms, "
         "interfaces, configurations, updates, enhancements, and derivative works, currently branded as "
         "Energy Copilot and described at https://voltusenergy.pl/produkt/."),
        ("\"DISPERON\"",
         "means the commercial brand name and trademark under which the Company markets and sells EMS "
         "services and solutions."),
        ("\"Licence\"",
         "means the software licence granted by Voltus Energy Sp. z o.o. to the Company pursuant to "
         "Article 5 of this Agreement."),
        ("\"Licence Fee\" or \"Per-Project Licence Fee\"",
         "means the fee payable per individual installation project prior to the Milestone Date, "
         "calculated in accordance with the Fee Schedule set out in Exhibit C."),
        ("\"Milestone\"",
         "means the aggregate deployment of 500 MWh of energy storage capacity under the EMS Software "
         "pursuant to contracts concluded by or through the Company."),
        ("\"Milestone Date\"",
         "means the date on which the Milestone is certified in writing by the Parties."),
        ("\"Perpetual Licence\"",
         "means the irrevocable, non-exclusive, royalty-free, worldwide licence granted to the Company "
         "upon achievement of the Milestone as set out in Article 6."),
        ("\"Commissioning\"",
         "means the on-site or remote technical installation, configuration, integration, grid code "
         "compliance verification, and handover of the EMS Software on a customer's BESS project."),
        ("\"Commissioning Fee\"",
         "means the fee charged to the customer for Commissioning services, retained in full by the Company."),
        ("\"Software Margin\"",
         "means the Software Licence Component of a customer contract, calculated as the applicable "
         "per-MWh Licence Rate from Exhibit C multiplied by the MWh capacity of the relevant project. "
         "All customer invoices shall itemise the Software Licence Component and the Commissioning "
         "Component as separate line items. The Commissioning Component is excluded from Software Margin "
         "for all purposes of this Agreement."),
        ("\"Development Bonus\"",
         "means the payment due from the Company to the Voltus Parties in respect of each Licence sold "
         "after the Milestone Date, calculated as 25% (twenty-five percent) of the Software Margin."),
        ("\"Intellectual Property\" or \"IP\"",
         "means all patents, trademarks, copyrights, database rights, trade secrets, know-how, source "
         "code, and other proprietary rights, whether registered or unregistered."),
        ("\"Territory\"",
         "means the European Union and its member states, and any other country or region agreed in "
         "writing by the Parties."),
        ("\"Exclusivity\"",
         "means the right of the Company to be the sole authorised distributor of the EMS Software in "
         "a designated country, subject to Article 8."),
        ("\"BESS Manufacturer Partners\"",
         "means the list of hardware manufacturers set out in Exhibit B, as agreed and signed by the "
         "Parties on the date of this Agreement."),
        ("\"Escrow Agent\"",
         "means the independent third party appointed to hold the Escrow Materials pursuant to Article 6.4."),
        ("\"Escrow Materials\"",
         "means the source code, build scripts, technical documentation, and deployment instructions "
         "for the then-current version of the EMS Software, as updated by Voltus pursuant to Article 6.4."),
        ("\"Software Support Annex\"",
         "means the service level and maintenance schedule set out in Exhibit D, forming part of this Agreement."),
    ]
    tbl = doc.add_table(rows=len(defs), cols=2)
    tbl.style = 'Table Grid'
    for i, (term, defn) in enumerate(defs):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        if i % 2 == 0:
            set_cell_bg(c0, 'EEF3FB')
            set_cell_bg(c1, 'EEF3FB')
        p0 = c0.paragraphs[0]; p0.clear()
        r0 = p0.add_run(term); r0.font.bold = True; r0.font.size = Pt(9)
        p1 = c1.paragraphs[0]; p1.clear()
        r1 = p1.add_run(defn); r1.font.size = Pt(9)
    doc.add_paragraph()

    # ── ARTICLE 2 — SHARE STRUCTURE ────────────────────────────────────────
    heading(doc, "ARTICLE 2 — SHARE STRUCTURE AND TRANSFER")
    heading(doc, "2.1  Current Ownership", level=2)
    body(doc, "Immediately prior to execution of this Agreement, 100% of the issued share capital of "
         "the Company was held by Lighthief International.")
    heading(doc, "2.2  Transfer of Shares", level=2)
    body(doc, "Upon execution of this Agreement, the following share transfers are effected and agreed "
         "by all Parties:")
    simple_table(doc, [
        ["Shareholder", "Shareholding", "Percentage (%)", "Purchase Consideration"],
        ["Lighthief International", "76 shares", "76%", "N/A (Retained)"],
        ["Mr. Marcin Szumiło",      "8 shares",  "8%",  "EUR 200"],
        ["Mr. Andrzej Lechowicz",   "8 shares",  "8%",  "EUR 200"],
        ["Mr. Kamil Talar",         "8 shares",  "8%",  "EUR 200"],
    ])
    doc.add_paragraph()
    heading(doc, "2.3  Purchase Consideration — Confirmed Payment", level=2)
    body(doc, "The Parties confirm and acknowledge that the purchase consideration of EUR 200 (two "
         "hundred euros) per 8% shareholding package has been received in full by the Company from "
         "each of Mr. Marcin Szumiło, Mr. Andrzej Lechowicz, and Mr. Kamil Talar prior to or "
         "simultaneously with the execution of this Agreement. Written receipts have been issued. "
         "The payment obligations under Article 2.2 are hereby acknowledged as fully discharged.")
    heading(doc, "2.4  Registration of Transfer", level=2)
    body(doc, "Mr. Arkadiusz Sybaris, acting in his capacity as Director of both Lighthief "
         "International and the Company, shall cause the necessary entries to be made in the Register "
         "of Members maintained by the Cyprus Registrar of Companies.")
    heading(doc, "2.5  Appointment of Shareholder Director", level=2)
    body(doc, "The incoming Shareholders — Mr. Marcin Szumiło, Mr. Andrzej Lechowicz, and Mr. Kamil "
         "Talar — shall, within 30 (thirty) days of execution of this Agreement, jointly designate "
         "one individual from among themselves to serve as a Director of the Company in accordance "
         "with Article 10.2.")
    heading(doc, "2.6  Current Directors", level=2)
    body(doc, "(a) Mr. Arkadiusz Sybaris — Director and Secretary, tax resident of Cyprus;\n"
         "(b) Mr. Kostas Alexander Papakosta — Director, tax resident of Cyprus.")
    heading(doc, "2.7  Share Certificates", level=2)
    body(doc, "Share certificates reflecting the revised shareholding shall be issued within 30 days "
         "of registration with the Cyprus Registrar of Companies and shall thereupon be appended as "
         "Exhibit A and circulated to all Shareholders. The acknowledgement of receipt in this clause "
         "refers to the certificates as and when issued.")
    note(doc, "Art 2.7 amended — certificates to be issued post-registration rather than falsely "
         "stated as attached at signing (Issue 9).")

    # ── ARTICLE 3 — DISPERON BRAND ────────────────────────────────────────
    heading(doc, "ARTICLE 3 — THE DISPERON BRAND")
    heading(doc, "3.1  Brand Ownership", level=2)
    body(doc, "The Company is the sole legal and beneficial owner of the DISPERON brand, trademark, "
         "and all associated intellectual property, including the trade name, logo, domain name "
         "disperon.com, and any related trademarks registered or to be registered with the EUIPO "
         "under Nice Classes 09 and 42. All Shareholders acknowledge and confirm this ownership.")
    heading(doc, "3.2  Brand Purpose", level=2)
    body(doc, "DISPERON is the exclusive commercial brand under which the Company shall market, sell, "
         "and deliver EMS and SCADA services and solutions for BESS across the European Union and "
         "globally. All customer-facing materials, contracts, invoices, and marketing communications "
         "shall use the DISPERON brand.")
    heading(doc, "3.3  Brand Protection", level=2)
    body(doc, "No Shareholder shall use the DISPERON brand, logo, or any confusingly similar "
         "designation for any purpose other than the activities of the Company without the prior "
         "written consent of the Director. Any Shareholder who undertakes activities constituting "
         "brand dilution or misuse shall be liable to the Company for resulting damages.")

    # ── ARTICLE 4 — ROLES ──────────────────────────────────────────────────
    heading(doc, "ARTICLE 4 — ROLES AND RESPONSIBILITIES OF SHAREHOLDERS")
    heading(doc, "4.1  Lighthief International", level=2)
    for item in [
        "Hold 76% of the issued share capital and exercise corresponding voting rights;",
        "Provide strategic direction and oversight of the Company's operations;",
        "Develop and manage sales channels for the EMS under the DISPERON brand;",
        "Retain in full all Commissioning Fees generated from projects delivered by or through its infrastructure;",
        "Be responsible for investor relations, brand development, and international expansion of DISPERON;",
        "Provide corporate administration and registered office infrastructure.",
    ]:
        bullet(doc, item)
    heading(doc, "4.2  Mr. Kamil Talar", level=2)
    for item in [
        "Hold 8% of the issued share capital in his own name and for his own account;",
        "Actively develop and manage sales channels for the DISPERON EMS in European markets, in "
        "coordination with Lighthief International;",
        "Deploy his existing business relationships, offices, and market presence — including his "
        "relationship with Linyang Energy Co., Ltd. and its customers — to generate BESS project pipeline;",
        "Report pipeline and sales activity to the Director on a monthly basis;",
        "Be jointly responsible with Lighthief International for achieving market sales targets.",
    ]:
        bullet(doc, item)
    body(doc,
         "Performance Obligation: Mr. Talar shall, during each calendar year following execution of "
         "this Agreement, (a) maintain a documented active pipeline of not less than 100 MWh of BESS "
         "projects incorporating the EMS Software; and (b) introduce a minimum of 3 qualified customer "
         "leads per year to the Company. Failure to meet these targets in any two consecutive annual "
         "periods shall entitle the Board to reclassify his shares as non-voting ordinary shares and "
         "to suspend dividend rights, subject to 30 days written notice and a 60-day cure period.",
         italic=True)
    note(doc, "Art 4.2 amended — Performance KPI and consequence mechanism added (Issue 8).")

    heading(doc, "4.3  Mr. Marcin Szumiło and Mr. Andrzej Lechowicz", level=2)
    body(doc, "Mr. Marcin Szumiło and Mr. Andrzej Lechowicz each individually hold 8% of the issued "
         "share capital of the Company in their own names. Both shall:")
    for item in [
        "Contribute the EMS Software to the Company's commercial operations through the Licence;",
        "Maintain primary responsibility for software development, technical architecture, product "
        "roadmap, and software quality of the EMS Software;",
        "Provide tier-2 technical support to the Company's commissioning and operations teams;",
        "Support the sales process through technical demonstrations and participation in client meetings;",
        "Ensure the EMS Software remains compliant with applicable EU regulatory requirements, "
        "including NIS2, GDPR, and relevant grid code standards.",
    ]:
        bullet(doc, item)

    # ── ARTICLE 5 — IP ─────────────────────────────────────────────────────
    heading(doc, "ARTICLE 5 — INTELLECTUAL PROPERTY")
    heading(doc, "5.1  Ownership of EMS Software IP", level=2)
    body(doc, "The Parties acknowledge and confirm that all Intellectual Property rights in and to the "
         "EMS Software are owned exclusively by Voltus Energy Sp. z o.o. Nothing in this Agreement "
         "shall be construed as a transfer, assignment, or novation of such IP rights to the Company "
         "or to any individual Shareholder, except as expressly provided in Articles 5 and 6.")
    heading(doc, "5.2  Licence Granted to the Company", level=2)
    body(doc, "Voltus Energy Sp. z o.o. hereby grants to the Company a non-exclusive, worldwide "
         "licence to use, deploy, configure, adapt, and sub-licence the EMS Software to customers "
         "under the DISPERON brand, subject to the commercial terms set out in Article 6.")
    heading(doc, "5.3  Voltus Retains Full IP", level=2)
    body(doc, "Voltus Energy Sp. z o.o. retains full and exclusive ownership of all IP in the EMS "
         "Software at all times, including after the grant of the Perpetual Licence pursuant to "
         "Article 6.3. The Perpetual Licence grants the Company a right to use the EMS Software "
         "but does not transfer ownership, title, or any other proprietary rights.")
    heading(doc, "5.4  Company IP and DISPERON Brand", level=2)
    body(doc, "All Intellectual Property created by or on behalf of the Company in connection with "
         "the DISPERON brand — including marketing materials, website content, implementation "
         "methodologies, customer documentation, and the DISPERON trademark — shall be owned "
         "exclusively by the Company. The Voltus Parties and individual Shareholders shall have "
         "no claim to such IP.")
    heading(doc, "5.5  IP Warranty and Indemnity", level=2)
    body(doc, "Voltus Energy Sp. z o.o. warrants that:")
    for item in [
        "It is the sole and unencumbered owner of all Intellectual Property in the EMS Software;",
        "The EMS Software does not infringe any third-party intellectual property right, including "
        "patents, copyrights, trade secrets, or database rights;",
        "No open-source software components are included in the EMS Software under licences "
        "incompatible with commercial distribution or sub-licensing;",
        "The EMS Software complies with applicable EU export control regulations.",
    ]:
        bullet(doc, item)
    body(doc, "Voltus shall indemnify, defend, and hold harmless the Company and its directors, "
         "officers, and employees against any third-party claims, losses, damages, costs, and legal "
         "fees arising from any breach of the warranties set out in this Article 5.5.")
    note(doc, "Art 5.5 is NEW — IP warranty and third-party indemnity added (Issue 7).")

    # ── ARTICLE 6 — COMMERCIAL MODEL ──────────────────────────────────────
    heading(doc, "ARTICLE 6 — COMMERCIAL MODEL AND LICENCE FEES")
    heading(doc, "6.1  Project-by-Project Revenue Model (Pre-Milestone)", level=2)
    body(doc,
         "(a) Per-Project Licence Fee: For each project, the Licence Fee shall be calculated in "
         "accordance with the Fee Schedule attached as Exhibit C. The applicable rate is the "
         "per-MWh Licence Rate multiplied by the project's installed MWh capacity. Voltus shall "
         "confirm the applicable fee in writing within 14 (fourteen) calendar days of receiving "
         "a written project brief from the Company. If Voltus fails to respond within 14 days, "
         "the fee shall be deemed confirmed at the rate stated in Exhibit C. The Licence Fee "
         "shall be paid by the Company to Voltus within 30 (thirty) days of the commissioning date.")
    body(doc,
         "(b) Commissioning Fee: The Company sets and collects the Commissioning Fee directly from "
         "the customer. The Commissioning Fee is retained in full by the Company and does not form "
         "part of any Licence Fee calculation. Voltus Energy Sp. z o.o. has no entitlement to any "
         "portion of the Commissioning Fee, whether before or after the Milestone Date.")
    note(doc, "Art 6.1(a) amended — fee now governed by Exhibit C schedule with 14-day "
         "response obligation and deemed-approval mechanism (Issue 1).")

    heading(doc, "6.2  Milestone Definition", level=2)
    body(doc, "The Milestone shall be deemed achieved when the aggregate installed capacity of BESS "
         "projects deploying the EMS Software under active contracts initiated by or through the "
         "Company reaches 500 MWh, calculated cumulatively from the date of execution of this "
         "Agreement. Projects contracted but cancelled prior to commissioning shall not count. "
         "The Director shall maintain a written cumulative log of commissioned projects accessible "
         "to all Parties upon 5 (five) business days' request.")

    heading(doc, "6.3  Perpetual Licence upon Achievement of Milestone", level=2)
    body(doc, "Upon achievement of the Milestone, Voltus Energy Sp. z o.o. shall automatically "
         "grant to the Company a perpetual, irrevocable, non-exclusive, worldwide, royalty-free "
         "licence to use, deploy, adapt, and sub-licence the EMS Software in its then-current "
         "version and all subsequent versions and updates, without any ongoing licence fee. "
         "This Perpetual Licence shall cover an unlimited number of customer installations "
         "worldwide and shall survive any change of control, dissolution, or restructuring of "
         "Voltus Energy Sp. z o.o., unless directly caused by a material uncured breach of "
         "this Agreement by the Company.")

    heading(doc, "6.4  Source Code Escrow", level=2)
    body(doc,
         "(a) Voltus shall, within 60 (sixty) days of execution of this Agreement, deposit the "
         "Escrow Materials with a mutually agreed independent Escrow Agent (the initial proposed "
         "provider being NCC Group Escrow or such equivalent as the Parties may agree).")
    body(doc,
         "(b) Voltus shall update the Escrow Materials within 30 (thirty) days of each material "
         "release or update to the EMS Software. Annual verification of escrow contents shall be "
         "conducted by the Escrow Agent at cost shared equally between Voltus and the Company.")
    body(doc,
         "(c) The Escrow Agent shall release the Escrow Materials to the Company upon the "
         "occurrence of any of the following Release Triggers: (i) Voltus Energy Sp. z o.o. "
         "enters insolvency, liquidation, or administration proceedings; (ii) Voltus ceases to "
         "provide contracted maintenance or support for a continuous period exceeding 90 days "
         "without cure; (iii) Voltus commits a material uncured breach of Article 7.2 persisting "
         "for more than 30 days after written notice from the Company.")
    body(doc,
         "(d) Following release of Escrow Materials, the Company may use them solely to continue "
         "operating and maintaining the EMS Software for existing and future customers under the "
         "DISPERON brand. The Company shall not sub-licence the source code itself to any third party.")
    note(doc, "Art 6.4 is NEW — Source code escrow mechanism added (Issue 2).")

    heading(doc, "6.5  Development Bonus (Post-Milestone)", level=2)
    body(doc,
         "(a) Following the Milestone Date, for each customer licence sold by the Company, the "
         "Company shall pay to the Voltus Parties a Development Bonus of 25% (twenty-five percent) "
         "of the Software Margin received from that customer.")
    body(doc,
         "(b) The Development Bonus rate of 25% is fixed as of the date of this Agreement. It "
         "shall be subject to review every 3 (three) years, adjustable by written agreement of "
         "all Parties within the range of 20% to 30%, provided Voltus produces a quarterly "
         "development expenditure summary demonstrating active R&D investment.")
    body(doc,
         "(c) The Development Bonus is calculated exclusively on the Software Margin as defined "
         "in Article 1. Commissioning Fees and all other service fees are expressly excluded.")
    body(doc,
         "(d) The Development Bonus shall be payable within 45 (forty-five) days of the Company "
         "receiving cleared payment from the relevant customer.")
    note(doc, "Art 6.5 amended — Development Bonus rate locked at 25% (Issue 5). "
         "Software Margin now clearly defined by formula (Issue 4).")

    heading(doc, "6.6  Records and Audit", level=2)
    body(doc, "The Company shall maintain accurate records of all project deployments, Licence "
         "Fees paid, Commissioning Fees received, and Development Bonus payments. Such records "
         "shall be available for inspection by the Voltus Parties and any Shareholder upon "
         "10 (ten) business days' written notice.")

    # ── ARTICLE 7 — OPERATIONAL STRUCTURE ─────────────────────────────────
    heading(doc, "ARTICLE 7 — OPERATIONAL STRUCTURE")
    heading(doc, "7.1  Sales and Market Development", level=2)
    body(doc, "The Company, acting through Lighthief International and Mr. Kamil Talar, shall be "
         "primarily responsible for the commercial development of the DISPERON brand and generation "
         "of customer pipeline across the Territory.")
    heading(doc, "7.2  Technology and Product Development — Voltus Obligations", level=2)
    body(doc, "Voltus Energy Sp. z o.o. shall be primarily responsible for the continuous development, "
         "maintenance, security, and improvement of the EMS Software, maintaining it in a commercially "
         "deployable and regulatory-compliant state at all times. Specifically, Voltus shall:")
    for item in [
        "Respond to severity-classified support requests within the timeframes set out in the "
        "Software Support Annex (Exhibit D);",
        "Maintain cloud-hosted components at a minimum availability of 99.5% measured monthly, "
        "excluding scheduled maintenance windows notified at least 5 days in advance;",
        "Implement required changes arising from EU and national grid code regulatory updates "
        "within 90 (ninety) days of official publication of such requirements;",
        "Provide the Company with a quarterly product roadmap and development activity report;",
        "Maintain cybersecurity compliance with NIS2 and applicable EU standards at all times.",
    ]:
        bullet(doc, item)
    body(doc, "Service credits for SLA failures are set out in Exhibit D. Persistent failure by "
         "Voltus to meet the obligations in this Article 7.2 constitutes a material breach "
         "entitling the Company to activate the Escrow release mechanism in Article 6.4.")
    note(doc, "Art 7.2 amended — SLA obligations and regulatory update commitments added (Issue 6). "
         "Detailed SLA schedule in Exhibit D.")

    heading(doc, "7.3  Commissioning and Implementation", level=2)
    body(doc, "On-site and remote commissioning, system integration, and project handover services "
         "shall be performed by the Company or a designated entity within the Lighthief International "
         "group. Voltus shall provide technical support, documentation, and training as reasonably "
         "required. All Commissioning revenues belong to the Company.")
    heading(doc, "7.4  Customer Contracts", level=2)
    body(doc, "All customer contracts shall be entered into by the Company. The Company is the legal "
         "counterparty to all end customers. Voltus acts as software supplier to the Company, not to "
         "end customers directly, unless otherwise agreed in writing.")

    # ── ARTICLE 8 — EXCLUSIVITY ────────────────────────────────────────────
    heading(doc, "ARTICLE 8 — EXCLUSIVITY AND MARKET MODEL")
    heading(doc, "8.1  Principle of Earned Exclusivity", level=2)
    body(doc, "Exclusivity is not granted automatically. It may be earned by the Company in specific "
         "countries on the basis of demonstrated commercial activity and achieved volume thresholds.")
    heading(doc, "8.2  Pre-Conditions for Exclusivity", level=2)
    body(doc, "Before any exclusivity is recognised in a specific country, the Company shall provide "
         "Voltus with a documented project pipeline, a 12-month sales plan, and evidence of active "
         "commercial engagement.")
    heading(doc, "8.3  Exclusivity Volume Thresholds", level=2)
    simple_table(doc, [
        ["Market Category", "Minimum Volume for Exclusivity", "Maintenance (annual)"],
        ["Large markets (Germany, Italy, Spain, France)",
         "500 MWh within 24 months", "100 MWh/year"],
        ["Mid-sized markets (Poland, Netherlands, Romania, Greece)",
         "250 MWh within 24 months", "50 MWh/year"],
        ["Smaller markets (Cyprus, other EU states)",
         "100 MWh within 24 months", "20 MWh/year"],
    ])
    doc.add_paragraph()
    note(doc, "Art 8.3 amended — Mid-sized market threshold reduced to 250 MWh; "
         "Smaller markets reduced to 100 MWh to reflect realistic market sizes (Issue 10).")

    heading(doc, "8.4  Maintenance of Exclusivity", level=2)
    body(doc, "Following the initial grant of exclusivity, the Company must achieve the minimum "
         "annual installed capacity set out in the Maintenance column of Article 8.3 in each "
         "subsequent 12-month period to retain exclusivity. Failure for two consecutive annual "
         "periods shall result in automatic lapsing of exclusivity for that country.")
    heading(doc, "8.5  Initial Country Limit", level=2)
    body(doc, "At the time of execution of this Agreement, the Company may seek exclusivity in a "
         "maximum of 5 (five) countries simultaneously.")
    heading(doc, "8.6  Voltus Direct Rights — BESS Manufacturer Partners", level=2)
    body(doc,
         "Notwithstanding any exclusivity granted to the Company, Voltus Energy Sp. z o.o. retains "
         "the right to supply the EMS Software directly to BESS Manufacturers listed in Exhibit B. "
         "This right:")
    for item in [
        "Applies only to the manufacturers expressly listed in Exhibit B;",
        "Does not extend to any manufacturer not listed in Exhibit B unless agreed by all Parties in writing;",
        "Is not subject to any exclusivity threshold or geographic restriction in respect of "
        "the listed BESS Manufacturer Partners.",
    ]:
        bullet(doc, item)
    body(doc,
         "BESS Manufacturer Partner Referral Obligation: Notwithstanding the above, where a project "
         "originating from a BESS Manufacturer Partner (including Linyang Energy Co., Ltd.) has been "
         "the subject of prior commercial engagement by the Company — evidenced by a written proposal, "
         "meeting record, or registered customer entry — Voltus shall promptly notify the Company "
         "within 5 (five) business days of receiving an inquiry from such manufacturer in respect of "
         "that project, and shall refer such project to the Company under the DISPERON brand. The "
         "Parties shall agree revenue-sharing terms within 15 business days of such referral. In the "
         "absence of agreement, the standard Development Bonus shall apply.")
    body(doc,
         "Voltus shall provide the Company with a monthly summary of any EMS Software supply "
         "activities to BESS Manufacturer Partners, including project name, customer identity "
         "(where permitted), country, and MWh capacity.")
    note(doc, "Art 8.6 amended — Referral obligation and transparency reporting added for BESS "
         "Manufacturer Partner (Linyang) projects with prior DISPERON engagement (Issue 3).")

    heading(doc, "8.7  Transparency and Customer Registration", level=2)
    body(doc, "Prior to any formal exclusivity being granted in a given country: (a) each Party may "
         "register customers; (b) a customer shall be assigned to the Party actively leading the "
         "sales process, confirmed in writing; (c) disputes over customer assignment shall be "
         "resolved by the Director within 15 (fifteen) business days.")
    heading(doc, "8.8  Non-Competition Undertaking of Voltus", level=2)
    body(doc, "Voltus Energy Sp. z o.o. undertakes, for the duration of this Agreement, that: "
         "(a) it shall maintain full transparency with the Company regarding its direct sales "
         "activities within the Territory; (b) if a customer approaches Voltus directly for EMS "
         "services, and that same customer has previously been approached by or is in active "
         "discussion with the Company, Voltus shall decline to engage independently and shall "
         "refer such customer to the Company under the DISPERON brand. This undertaking does not "
         "apply to customers originally introduced by Voltus independently of the Company, or to "
         "customers of BESS Manufacturer Partners listed in Exhibit B where the BESS Manufacturer "
         "Partner Referral Obligation in Article 8.6 does not apply.")

    # ── ARTICLE 9 — TRANSFER OF SHARES ────────────────────────────────────
    heading(doc, "ARTICLE 9 — TRANSFER OF SHARES")
    heading(doc, "9.1  Lock-Up Period", level=2)
    body(doc, "For 24 (twenty-four) months from the date of this Agreement, no Shareholder shall "
         "sell, transfer, assign, pledge, or otherwise dispose of any Shares without the prior "
         "written consent of Lighthief International.")
    heading(doc, "9.2  Right of First Refusal", level=2)
    body(doc, "Following the Lock-Up Period, if any Shareholder proposes to transfer Shares to a "
         "third party: written notice must be given to all other Shareholders specifying the number "
         "of Shares, proposed price, and buyer identity; each non-selling Shareholder has 30 (thirty) "
         "days to exercise a right of first refusal on a pro-rata basis at the same price and terms.")
    heading(doc, "9.3  Drag-Along Right", level=2)
    body(doc, "If Lighthief International proposes to sell 100% of its Shares to a bona fide "
         "purchaser requiring 100% of issued capital, Lighthief International may require all other "
         "Shareholders to sell their Shares on the same terms, on 30 (thirty) days' written notice.")
    heading(doc, "9.4  Tag-Along Right", level=2)
    body(doc, "If Lighthief International proposes to transfer more than 50% of its Shares to a "
         "third party, each other Shareholder has the right, within 20 (twenty) days, to include "
         "their Shares in such sale on the same terms.")

    # ── ARTICLE 10 — GOVERNANCE ────────────────────────────────────────────
    heading(doc, "ARTICLE 10 — GOVERNANCE AND BOARD OF DIRECTORS")
    heading(doc, "10.1  Current Directors", level=2)
    body(doc, "(a) Mr. Arkadiusz Sybaris — Director and Secretary;\n"
         "(b) Mr. Kostas Alexander Papakosta — Director.")
    heading(doc, "10.2  Board Composition", level=2)
    body(doc, "(a) Lighthief International, as Majority Shareholder, shall have the right to appoint "
         "up to 3 (three) Directors to the Board;\n"
         "(b) The Shareholders other than Lighthief International shall have the right to appoint "
         "1 (one) Director from among themselves to the Board.")
    heading(doc, "10.3  Appointment and Removal", level=2)
    body(doc, "Directors appointed by Lighthief International may be appointed and removed at any "
         "time by written resolution of Lighthief International. The Shareholder Director may be "
         "replaced by written agreement of Mr. Szumiło, Mr. Lechowicz, and Mr. Talar.")
    heading(doc, "10.4  Decisions Requiring Shareholder Consent", level=2)
    body(doc, "The following matters shall require the written consent of Shareholders holding a "
         "majority of Shares:")
    for item in [
        "Amendment of the Company's Articles of Association;",
        "Issuance of new shares or share capital increases;",
        "Material change to the Company's principal business activity;",
        "Entry into any related-party transaction exceeding EUR 50,000 in value;",
        "Appointment or removal of auditors.",
    ]:
        bullet(doc, item)
    heading(doc, "10.5  Reserved Matters Requiring 85% Shareholder Consent", level=2)
    body(doc, "The following matters shall require the written consent of Shareholders holding not "
         "less than 85% (eighty-five percent) of all issued Shares:")
    for item in [
        "Material amendment to the commercial model set out in Article 6, including the "
        "Development Bonus rate, Software Margin definition, or Exhibit C fee schedule;",
        "Voluntary dissolution or winding up of the Company;",
        "Any transaction that would impair, encumber, or terminate the Perpetual Licence "
        "or the Escrow arrangement under Article 6.4.",
    ]:
        bullet(doc, item)
    note(doc, "Art 10.5 is NEW — Reserved matters at 85% threshold protect Voltus parties' "
         "core economic interest without day-to-day blocking power (Issue 11).")

    heading(doc, "10.6  Shareholder Meetings", level=2)
    body(doc, "The Company shall hold at least one annual meeting of Shareholders. Additional "
         "meetings may be called by the Director or by any Shareholder holding at least 20% of "
         "issued capital, on not less than 14 (fourteen) days' notice. Meetings may be held in "
         "person or by video conference.")

    # ── ARTICLE 11 — CONFIDENTIALITY ──────────────────────────────────────
    heading(doc, "ARTICLE 11 — CONFIDENTIALITY")
    heading(doc, "11.1  Confidential Information", level=2)
    body(doc, "Each Party shall keep confidential all non-public information received in connection "
         "with this Agreement, including technical specifications, source code, customer data, "
         "financial terms, and business strategies.")
    heading(doc, "11.2  Duration", level=2)
    body(doc, "The confidentiality obligation applies during the term of this Agreement and for "
         "5 (five) years following its termination or expiry.")
    heading(doc, "11.3  Permitted Disclosure", level=2)
    body(doc, "Confidential Information may be disclosed only to employees, contractors, or advisors "
         "with a need to know who are bound by equivalent obligations, or where required by law.")

    # ── ARTICLE 12 — GOVERNING LAW ─────────────────────────────────────────
    heading(doc, "ARTICLE 12 — GOVERNING LAW AND DISPUTE RESOLUTION")
    heading(doc, "12.1  Governing Law", level=2)
    body(doc, "This Agreement shall be governed by and construed in accordance with the laws of "
         "the Republic of Cyprus.")
    heading(doc, "12.2  Negotiation", level=2)
    body(doc, "In the event of any Dispute, the Parties shall attempt resolution through good-faith "
         "negotiation within 30 (thirty) days of written notice.")
    heading(doc, "12.3  Arbitration", level=2)
    body(doc, "If unresolved by negotiation, any Dispute shall be referred to final and binding "
         "arbitration in Limassol, Cyprus, under the Cyprus Arbitration Law (Cap. 4), conducted "
         "in English.")

    # ── ARTICLE 13 — GENERAL ──────────────────────────────────────────────
    heading(doc, "ARTICLE 13 — GENERAL PROVISIONS")
    for num, title, text in [
        ("13.1", "Entire Agreement",
         "This Agreement constitutes the entire agreement of the Parties with respect to its "
         "subject matter and supersedes all prior understandings, negotiations, and representations."),
        ("13.2", "Amendment",
         "This Agreement may only be amended by a written instrument signed by all Parties."),
        ("13.3", "Severability",
         "If any provision is found invalid, illegal, or unenforceable, the remaining provisions "
         "shall continue in full force."),
        ("13.4", "Waiver",
         "No failure to exercise any right or remedy shall constitute a waiver thereof."),
        ("13.5", "Notices",
         "All notices shall be in writing, delivered by email with read receipt or registered "
         "post to the addresses stated in the Preamble."),
        ("13.6", "Counterparts and Electronic Signatures",
         "This Agreement may be executed in counterparts. Electronic signatures shall be valid "
         "and binding. The English version shall prevail over any translation."),
    ]:
        heading(doc, f"{num}  {title}", level=2)
        body(doc, text)

    # ── SIGNATURES ─────────────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "SIGNATURES", centre=True)
    body(doc, "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first "
         "written above.", italic=True)
    doc.add_paragraph()
    for name, role in [
        ("LIGHTHIEF INTERNATIONAL", "Represented by: Mr. Arkadiusz Sybaris, Director"),
        ("LIGHTHIEF EUBESS LTD (HE 474192)", "Represented by: Mr. Arkadiusz Sybaris, Director and Secretary"),
        ("MR. MARCIN SZUMIŁO", "Individual Shareholder"),
        ("MR. ANDRZEJ LECHOWICZ", "Individual Shareholder"),
        ("MR. KAMIL TALAR", "Individual Shareholder"),
        ("VOLTUS ENERGY SP. Z O.O.", "Represented by: Mr. Marcin Szumiło & Mr. Andrzej Lechowicz, Management Board"),
    ]:
        sig_block(doc, name, role)

    # ── EXHIBITS ───────────────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "EXHIBIT A — SHARE CERTIFICATES")
    body(doc, "Share certificates reflecting the post-transfer shareholding shall be issued within "
         "30 days of registration with the Cyprus Registrar of Companies and appended hereto. "
         "Until issued, the share transfer table in Article 2.2 shall constitute the authoritative "
         "record.", italic=True)

    doc.add_page_break()
    heading(doc, "EXHIBIT B — BESS MANUFACTURER PARTNERS")
    body(doc, "The following is the complete and exclusive list of BESS hardware manufacturers in "
         "respect of whom Voltus Energy Sp. z o.o. retains direct supply rights pursuant to "
         "Article 8.6.")
    simple_table(doc, [
        ["No.", "Manufacturer Name", "Registered Country", "Date Added"],
        ["1.", "Linyang Energy Co., Ltd.", "People's Republic of China", "Date of Agreement"],
        ["2.", "________________________", "________________________", "________________"],
        ["3.", "________________________", "________________________", "________________"],
        ["4.", "________________________", "________________________", "________________"],
    ])
    doc.add_paragraph()
    body(doc, "Any manufacturer not listed above shall not benefit from the direct supply rights "
         "in Article 8.6. Any addition requires written consent of all Parties and a signed amendment.")

    doc.add_page_break()
    heading(doc, "EXHIBIT C — LICENCE FEE SCHEDULE")
    note(doc, "Exhibit C is NEW — resolves Issue 1 (undefined fee) and Issue 4 (Software Margin ambiguity).")
    body(doc, "This Exhibit C forms part of the Agreement and governs all Licence Fee calculations "
         "under Article 6.")
    heading(doc, "C.1  Per-MWh Licence Rate (Pre-Milestone)", level=2)
    simple_table(doc, [
        ["Project Scale", "Licence Rate (EUR/MWh installed)", "Effective Date"],
        ["≤ 10 MWh",      "EUR [___] / MWh",                 "Date of Agreement"],
        ["10 MWh – 50 MWh", "EUR [___] / MWh",              "Date of Agreement"],
        ["> 50 MWh",      "EUR [___] / MWh (negotiated)",    "Date of Agreement"],
    ])
    doc.add_paragraph()
    body(doc, "The rates in brackets above shall be agreed by the Parties within 14 days of "
         "execution and inserted by written addendum. The Development Bonus (post-Milestone) shall "
         "be calculated as 25% of the Software Licence Component = Applicable Rate × Project MWh.")
    heading(doc, "C.2  Annual Escalation", level=2)
    body(doc, "The per-MWh Licence Rate shall be adjusted annually on 1 January of each year by "
         "the change in the EU Harmonised Index of Consumer Prices (HICP) published by Eurostat "
         "for the preceding calendar year. The first adjustment shall occur on 1 January of the "
         "year following execution.")
    heading(doc, "C.3  14-Day Confirmation Obligation", level=2)
    body(doc, "Voltus shall confirm the applicable Licence Fee for each submitted project brief "
         "within 14 (fourteen) calendar days of receipt. Failure to respond constitutes deemed "
         "approval at the applicable schedule rate. Project briefs shall include: project name, "
         "location, system size (MWh), customer name, and estimated commissioning date.")

    doc.add_page_break()
    heading(doc, "EXHIBIT D — SOFTWARE SUPPORT ANNEX")
    note(doc, "Exhibit D is NEW — resolves Issue 6 (no SLA).")
    body(doc, "This Software Support Annex forms part of the Agreement and governs Voltus's "
         "support and maintenance obligations under Article 7.2.")
    heading(doc, "D.1  Severity Classification", level=2)
    simple_table(doc, [
        ["Severity", "Description", "Examples", "Response SLA", "Resolution SLA"],
        ["P1 — Critical",
         "System completely unavailable or safety-critical function impaired",
         "EMS offline, DSO command failure",
         "2 hours",
         "24 hours"],
        ["P2 — Major",
         "Core function severely degraded; no workaround",
         "SOC reporting error, control loop failure",
         "8 business hours",
         "5 business days"],
        ["P3 — Minor",
         "Non-critical function impaired; workaround exists",
         "Reporting anomaly, UI defect",
         "2 business days",
         "30 business days"],
        ["P4 — Cosmetic",
         "Minor issue; no operational impact",
         "Display formatting, translation",
         "5 business days",
         "Next planned release"],
    ])
    doc.add_paragraph()
    heading(doc, "D.2  Availability Target", level=2)
    body(doc, "Cloud-hosted EMS components shall maintain minimum availability of 99.5% measured "
         "per calendar month, excluding planned maintenance windows. Planned maintenance windows "
         "require a minimum of 5 (five) calendar days' advance notice to the Company.")
    heading(doc, "D.3  Regulatory Update Obligation", level=2)
    body(doc, "Voltus shall implement changes required by new or amended EU or national grid code "
         "regulations within 90 (ninety) days of the official publication date of such requirements. "
         "Voltus shall notify the Company within 14 days of becoming aware of any regulatory change "
         "affecting the EMS Software.")
    heading(doc, "D.4  Service Credits", level=2)
    body(doc, "In the event Voltus fails to meet an SLA in this Annex: (a) for availability below "
         "99.5%: the Company may offset from the next Licence Fee or Development Bonus payment an "
         "amount equal to 5% of the monthly Licence Fee or Development Bonus payment for each "
         "percentage point (or fraction thereof) below 99.5%; (b) for P1 response/resolution "
         "failures: EUR 500 credit per failure per day of delay beyond the applicable SLA. Credits "
         "are the Company's sole financial remedy for SLA failures and do not limit other remedies "
         "for material breach.")
    heading(doc, "D.5  Annual Review", level=2)
    body(doc, "SLA targets and credits shall be reviewed annually at the request of either Party. "
         "Any amendment requires written agreement of all Parties.")

    out_path = "/Volumes/T7 Grey/solinvest/DisperonEMS/docs/contract/DISPERON_SHA_v3_PROPOSED.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")

build()
