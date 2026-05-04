"""
DISPERON_SHA_v5_EN.docx
Shareholders Agreement & IP Licence — Version 5 (Lighthief proposed final)
English execution copy — ready for signing
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE  = RGBColor(0x1F, 0x49, 0x7D)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x50, 0x50, 0x50)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _cell_bg(cell, hex_col):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_col)
    tcPr.append(shd)


def setup(doc):
    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(10)
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(3.0)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = BLUE; r.font.bold = True; r.font.size = Pt(13)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = BLUE; r.font.bold = True; r.font.size = Pt(11)


def para(doc, text, bold=False, italic=False, size=10, indent=0, centre=False):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    if centre: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.bold = bold
    run.font.italic = italic; run.font.color.rgb = BLACK
    return p


def bul(doc, text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(level)
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.color.rgb = BLACK


def tbl(doc, data, col_widths=None, hdr='1F497D'):
    cols = len(data[0])
    t = doc.add_table(rows=len(data), cols=cols)
    t.style = 'Table Grid'
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            if i == 0: _cell_bg(c, hdr)
            p = c.paragraphs[0]; p.clear()
            r = p.add_run(val)
            r.font.size = Pt(9)
            r.font.bold = (i == 0)
            r.font.color.rgb = WHITE if i == 0 else BLACK
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                if ci < len(row.cells):
                    row.cells[ci].width = Cm(w)
    return t


def sig_block(doc, party_name, role, extra=None, ref=None):
    p = doc.add_paragraph()
    r1 = p.add_run(party_name + "\n"); r1.font.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(role + "\n"); r2.font.size = Pt(10)
    if extra:
        re = p.add_run(extra + "\n"); re.font.size = Pt(10)
    if ref:
        rr = p.add_run(f"Reference: {ref}\n"); rr.font.size = Pt(9); rr.font.italic = True
    p.add_run("\nSignature: _________________________________\n\n")
    p.add_run("Full name (print): _______________________\n\n")
    p.add_run("Date: ____________________________________\n\n")
    p.add_run("Place: ___________________________________")
    doc.add_paragraph()


def divider(doc):
    doc.add_paragraph()


# ── BUILD ─────────────────────────────────────────────────────────────────────
def build():
    doc = Document()
    setup(doc)

    # ── COVER ─────────────────────────────────────────────────────────────────
    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SHAREHOLDERS AGREEMENT\nAND INTELLECTUAL PROPERTY LICENCE")
    r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = BLUE

    doc.add_paragraph()
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "LIGHTHIEF EUBESS LTD\nRegistration No. HE 474192\n"
        "Lophitis Business Center I, Floor 2, Office 1\n"
        "28 Oktovriou & Aemiliou Chourmouziou, 3035 Limassol, Cyprus"
    )
    r2.font.size = Pt(12)

    doc.add_paragraph(); doc.add_paragraph()
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("VERSION 5  ·  Dated: ______________, 2025")
    r3.font.size = Pt(11); r3.font.italic = True

    doc.add_paragraph()
    p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(
        "PROPOSED FINAL — Lighthief EUBESS Ltd / Voltus Energy Sp. z o.o.\n"
        "This version incorporates Lighthief's negotiated positions on all open items."
    )
    r4.font.size = Pt(9); r4.font.italic = True; r4.font.color.rgb = GREY

    doc.add_page_break()

    # ── PREAMBLE ──────────────────────────────────────────────────────────────
    h1(doc, "PREAMBLE")
    para(doc,
         'This Shareholders Agreement and Intellectual Property Licence (the "Agreement") is entered '
         'into on the date last signed below by and among the following parties:')
    doc.add_paragraph()

    parties = [
        ("PARTY 1 — LIGHTHIEF INTERNATIONAL",
         'A company duly incorporated and existing under applicable law, represented by its Director, '
         'Mr. Arkadiusz Sybaris ("Lighthief International" or "Majority Shareholder").'),
        ("PARTY 2 — LIGHTHIEF EUBESS LTD",
         'A private limited company registered in Cyprus under registration number HE 474192, having its '
         'registered office at 28 Oktovriou & Aemiliou Chourmouziou, Lophitis Business Center I, Floor 2, '
         'Office 1, 3035 Limassol, Cyprus, represented by its Director, Mr. Arkadiusz Sybaris (the "Company").'),
        ("PARTY 3 — MR. MARCIN SZUMIŁO",
         'An individual, holder of a Polish identity document, residing at an address notified to the '
         'Company ("Mr. Szumiło").'),
        ("PARTY 4 — MR. ANDRZEJ LECHOWICZ",
         'An individual, holder of a Polish identity document, residing at an address notified to the '
         'Company ("Mr. Lechowicz").'),
        ("PARTY 5 — MR. KAMIL TALAR",
         'An individual, holder of a Polish identity document, residing at an address notified to the '
         'Company ("Mr. Talar").'),
        ("PARTY 6 — MR. KOSTAS ALEXANDER PAPACOSTA",
         'An individual, holder of a Cypriot identity document, tax resident of the Republic of Cyprus, '
         'permanently residing in Cyprus, serving as Director of Lighthief EUBESS Ltd '
         '("Mr. Papacosta" or "Cyprus Director").'),
        ("PARTY 7 — VOLTUS ENERGY SP. Z O.O.",
         'A limited liability company incorporated and existing under the laws of the Republic of Poland, '
         'with its registered office at Gdański Park Naukowo-Technologiczny, Budynek B, Lokal 2.10.3, '
         'ul. Trzy Lipy 3, 80-172 Gdańsk, Poland, NIP: 1990133260, represented by its Management Board '
         'members, Mr. Marcin Szumiło and Mr. Andrzej Lechowicz ("Voltus").'),
    ]
    for name, desc in parties:
        p = doc.add_paragraph()
        p.add_run(name + "\n").font.bold = True
        p.runs[0].font.size = Pt(10)
        r2 = p.add_run(desc); r2.font.size = Pt(10); r2.font.color.rgb = BLACK
        doc.add_paragraph()

    para(doc,
         'Lighthief International, Mr. Szumiło, Mr. Lechowicz, Mr. Talar, and Mr. Papacosta are referred '
         'to individually as a "Shareholder" and collectively as the "Shareholders". Mr. Szumiło and '
         'Mr. Lechowicz together with Voltus Energy Sp. z o.o. are referred to collectively as the '
         '"Voltus Parties". Parties 1 through 7 are collectively referred to as the "Parties".')

    # ── ARTICLE 1 — DEFINITIONS ───────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "ARTICLE 1 — DEFINITIONS")
    para(doc, "For the purposes of this Agreement, the following terms shall have the meanings assigned below:")
    doc.add_paragraph()

    defs = [
        ('"Company"', 'means Lighthief EUBESS Ltd, HE 474192, as described in the Preamble.'),
        ('"Shares"', 'means the ordinary shares in the capital of the Company.'),
        ('"EMS Software" or "EMS"',
         'means the energy management system and SCADA software platform developed by Voltus Energy '
         'Sp. z o.o., together with all associated source code, object code, documentation, algorithms, '
         'interfaces, configurations, updates, enhancements, and derivative works, currently branded as '
         'Energy Copilot and described at https://voltusenergy.pl/produkt/.'),
        ('"DISPERON"',
         'means the commercial brand name and trademark under which the Company markets and sells EMS '
         'services and solutions.'),
        ('"Licence"',
         'means the software licence granted by Voltus Energy Sp. z o.o. to the Company pursuant to '
         'Article 5 of this Agreement.'),
        ('"Catalogue Price"',
         'means the reference price for a given project, calculated as: (Catalogue MW Rate × project '
         'installed MW capacity) + (Catalogue MWh Rate × project installed MWh capacity), using the '
         'rates set out in Exhibit C. The Catalogue Price is the base from which the Development Bonus '
         'and Subscription Fee are calculated.'),
        ('"Per-Project Licence Fee"',
         'means the fee payable per individual installation project prior to the Milestone Date, '
         'calculated in accordance with the Fee Schedule set out in Exhibit C.'),
        ('"Milestone"',
         'means the aggregate deployment of 500 MWh of energy storage capacity under the EMS Software '
         'pursuant to contracts concluded by or through the Company.'),
        ('"Milestone Date"',
         'means the date on which the Milestone is certified in writing by the Parties.'),
        ('"Perpetual Licence"',
         'means the irrevocable, non-exclusive, royalty-free, worldwide licence granted to the Company '
         'upon achievement of the Milestone as set out in Article 6.'),
        ('"Commissioning"',
         'means the on-site or remote technical installation, configuration, integration, grid code '
         'compliance verification, and handover of the EMS Software on a customer\'s BESS project.'),
        ('"Commissioning Fee"',
         'means the fee charged to the customer for Commissioning services, retained in full by the Company.'),
        ('"Software Margin"',
         'means the Catalogue Price for a given project as calculated per the definition above and '
         'Exhibit C. All customer invoices shall separately itemise the Software Licence Component '
         '(equal to the Catalogue Price) and the Commissioning Component. The Commissioning Component '
         'is excluded from Software Margin for all purposes of this Agreement.'),
        ('"Development Bonus"',
         'means the payment due from the Company to the Voltus Parties in respect of each project '
         'licence, calculated as a percentage of the Catalogue Price as set out in Article 6.5.'),
        ('"EMS Subscription Fee" (client-facing)',
         'means the annual fee charged by the Company to the customer for access to the DISPERON EMS '
         'and SCADA software platform, currently set at EUR 400 (four hundred euros) per MWh of '
         'installed project capacity per year. This fee is billed separately from any LTSA or O&M '
         'service fees and is invoiced annually from the commissioning date (PAC) of each project.'),
        ('"Voltus Subscription Fee"',
         'means the annual fee payable by the Company to Voltus in respect of each commissioned '
         'project, calculated as 20% (twenty percent) of the EMS Subscription Fee received by '
         'the Company from the relevant customer in respect of that project, payable within '
         '14 (fourteen) days of receipt of each annual EMS Subscription Fee payment from the customer.'),
        ('"Intellectual Property" or "IP"',
         'means all patents, trademarks, copyrights, database rights, trade secrets, know-how, source '
         'code, and other proprietary rights, whether registered or unregistered.'),
        ('"Territory"',
         'means the European Union and its member states, and any other country or region agreed in '
         'writing by the Parties.'),
        ('"Exclusivity"',
         'means the right of the Company to be the sole authorised distributor of the EMS Software in '
         'a designated country, subject to Article 8.'),
        ('"Restricted Exclusivity"',
         'means a transitional state in which the Company retains the right to complete existing '
         'contracted projects in a country but Voltus may independently pursue new commercial '
         'opportunities, as set out in Article 8.4.'),
        ('"BESS Manufacturer Partners"',
         'means the list of hardware manufacturers set out in Exhibit B, as agreed and signed by the '
         'Parties on the date of this Agreement and as may be updated from time to time by mutual '
         'written agreement of all Parties.'),
        ('"Escrow Agent"',
         'means the independent third party appointed to hold the Escrow Materials pursuant to Article 6.4.'),
        ('"Escrow Materials"',
         'means the source code, build scripts, technical documentation, and deployment instructions '
         'for the then-current version of the EMS Software, as updated by Voltus pursuant to Article 6.4.'),
        ('"Software Support Annex"',
         'means the service level and maintenance schedule set out in Exhibit D, forming part of this Agreement.'),
    ]

    dt = doc.add_table(rows=len(defs), cols=2)
    dt.style = 'Table Grid'
    for i, (term, defn) in enumerate(defs):
        c0, c1 = dt.rows[i].cells[0], dt.rows[i].cells[1]
        if i % 2 == 0:
            _cell_bg(c0, 'EEF3FB'); _cell_bg(c1, 'EEF3FB')
        p0 = c0.paragraphs[0]; p0.clear()
        r0 = p0.add_run(term); r0.font.bold = True; r0.font.size = Pt(9)
        p1 = c1.paragraphs[0]; p1.clear()
        r1 = p1.add_run(defn); r1.font.size = Pt(9); r1.font.color.rgb = BLACK
    for row in dt.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(10.5)
    divider(doc)

    # ── ARTICLE 2 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 2 — SHARE STRUCTURE AND TRANSFER")
    h2(doc, "2.1  Current Ownership")
    para(doc, "Immediately prior to execution of this Agreement, 100% of the issued share capital of "
         "the Company was held by Lighthief International.")
    h2(doc, "2.2  Transfer of Shares")
    para(doc, "Upon execution of this Agreement, the following share transfers are effected and agreed "
         "by all Parties:")
    tbl(doc, [
        ["Shareholder", "Shares", "%", "Purchase Consideration"],
        ["Lighthief International",        "65", "65%", "N/A (Retained)"],
        ["Mr. Kostas Alexander Papacosta", "10", "10%", "EUR 250"],
        ["Mr. Kamil Talar",                "13", "13%", "EUR 325"],
        ["Mr. Marcin Szumiło",             "6",  "6%",  "EUR 150"],
        ["Mr. Andrzej Lechowicz",          "6",  "6%",  "EUR 150"],
        ["TOTAL",                          "100","100%", "—"],
    ], col_widths=[6, 2, 2, 4])
    divider(doc)

    h2(doc, "2.3  Purchase Consideration — Confirmed Payment")
    para(doc,
         "The purchase consideration is EUR 25 (twenty-five euros) per 1% shareholding. The Parties "
         "confirm and acknowledge that the following amounts have been received in full by the Company "
         "from each paying Shareholder prior to or simultaneously with the execution of this Agreement: "
         "(a) EUR 250 from Mr. Kostas Alexander Papacosta (10%); (b) EUR 325 from Mr. Kamil Talar "
         "(13%); (c) EUR 150 from each of Mr. Marcin Szumiło and Mr. Andrzej Lechowicz (6% each). "
         "Written receipts have been issued. The payment obligations under Article 2.2 are hereby "
         "acknowledged as fully discharged.")
    h2(doc, "2.4  Registration of Transfer")
    para(doc,
         "Mr. Arkadiusz Sybaris, acting in his capacity as Director of both Lighthief International "
         "and the Company, shall cause the necessary entries to be made in the Register of Members "
         "maintained by the Cyprus Registrar of Companies. Updated share certificates shall be issued "
         "to each Shareholder as soon as practicable following registration.")
    h2(doc, "2.5  Appointment of Shareholder Director")
    para(doc,
         "The Shareholders other than Lighthief International and Mr. Papacosta — being Mr. Marcin "
         "Szumiło, Mr. Andrzej Lechowicz, and Mr. Kamil Talar — shall, within 30 (thirty) days of "
         "execution of this Agreement, jointly designate one individual from among themselves to serve "
         "as a Director of the Company in accordance with Article 10.2.")
    h2(doc, "2.6  Current Directors")
    para(doc,
         "As of the date of this Agreement, the duly appointed Directors of the Company are:\n"
         "(a) Mr. Arkadiusz Sybaris — Director and Secretary, tax resident of the Republic of Cyprus;\n"
         "(b) Mr. Kostas Alexander Papacosta — Director and Cyprus Managing Director, tax resident of "
         "the Republic of Cyprus, permanently residing in Cyprus.")
    h2(doc, "2.7  Share Certificates")
    para(doc,
         "Share certificates reflecting the revised shareholding shall be issued within 30 (thirty) "
         "days of registration with the Cyprus Registrar of Companies and appended as Exhibit A.")

    # ── ARTICLE 3 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 3 — THE DISPERON BRAND")
    h2(doc, "3.1  Brand Ownership")
    para(doc,
         "The Company is the sole legal and beneficial owner of the DISPERON brand, trademark, and "
         "all associated intellectual property, including the trade name, logo, domain name "
         "disperon.com, and any related trademarks registered or to be registered with the European "
         "Union Intellectual Property Office (EUIPO) under Nice Classes 09 and 42. All Shareholders "
         "acknowledge and confirm this ownership.")
    h2(doc, "3.2  Brand Purpose")
    para(doc,
         "DISPERON is the exclusive commercial brand under which the Company shall market, sell, and "
         "deliver EMS and SCADA services and solutions for Battery Energy Storage Systems (BESS) "
         "across the European Union and globally. All customer-facing materials, contracts, invoices, "
         "and marketing communications shall use the DISPERON brand.")
    h2(doc, "3.3  Brand Protection")
    para(doc,
         "No Shareholder shall use the DISPERON brand, logo, or any confusingly similar designation "
         "for any purpose other than the activities of the Company without the prior written consent "
         "of the Director. Any Shareholder who undertakes activities constituting brand dilution or "
         "misuse shall be liable to the Company for resulting damages.")

    # ── ARTICLE 4 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 4 — ROLES AND RESPONSIBILITIES OF SHAREHOLDERS")
    h2(doc, "4.1  Lighthief International")
    para(doc, "Lighthief International, as Majority Shareholder, shall:")
    for item in [
        "Hold 65% of the issued share capital of the Company and exercise corresponding voting rights;",
        "Provide strategic direction and oversight of the Company's operations;",
        "Develop and manage sales channels for the EMS under the DISPERON brand through its own "
        "offices and networks across Europe and beyond;",
        "Retain in full all Commissioning Fees generated from projects delivered by or through its "
        "infrastructure;",
        "Be responsible for investor relations, brand development, and international expansion of DISPERON;",
        "Provide corporate administration and registered office infrastructure through the Lighthief "
        "International group.",
    ]:
        bul(doc, item)

    h2(doc, "4.2  Mr. Kamil Talar")
    para(doc, "Mr. Kamil Talar shall:")
    for item in [
        "Hold 13% of the issued share capital of the Company in his own name and for his own account;",
        "Actively develop and manage sales channels for the DISPERON EMS in European markets, in "
        "coordination with Lighthief International;",
        "Deploy his existing business relationships, offices, and market presence — including his "
        "relationship with Linyang Energy Co., Ltd. and its customers — to generate BESS project pipeline;",
        "Report pipeline and sales activity to the Director on a monthly basis;",
        "Be jointly responsible with Lighthief International for achieving market sales targets.",
    ]:
        bul(doc, item)
    para(doc,
         "Activity Commitment: Mr. Talar shall use his best commercial efforts to actively develop "
         "and maintain a BESS project pipeline incorporating the EMS Software and to report his "
         "commercial activities and market opportunities to the Director on a monthly basis.",
         indent=0.5)

    h2(doc, "4.3  Mr. Marcin Szumiło and Mr. Andrzej Lechowicz")
    para(doc,
         "Mr. Marcin Szumiło and Mr. Andrzej Lechowicz each individually hold 6% of the issued share "
         "capital of the Company in their own names and for their own respective accounts. Both, "
         "individually and in their capacity as Management Board members of Voltus Energy Sp. z o.o., shall:")
    for item in [
        "Contribute the EMS Software to the Company's commercial operations through the Licence;",
        "Maintain primary responsibility for software development, technical architecture, product "
        "roadmap, and software quality of the EMS Software;",
        "Provide tier-2 technical support to the Company's commissioning and operations teams;",
        "Support the sales process through technical demonstrations and participation in client meetings;",
        "Ensure the EMS Software remains compliant with applicable EU regulatory requirements, "
        "including NIS2, GDPR, and relevant grid code standards.",
    ]:
        bul(doc, item)

    h2(doc, "4.4  Mr. Kostas Alexander Papacosta")
    para(doc,
         "Mr. Kostas Alexander Papacosta holds 10% of the issued share capital of the Company in his "
         "own name and for his own account. Mr. Papacosta, in his capacity as Cyprus Managing Director "
         "and Director of the Company, shall:")
    for item in [
        "Serve as the primary executive director of Lighthief EUBESS Ltd in the Republic of Cyprus, "
        "holding all necessary registrations and authorisations required for the Company's Cyprus operations;",
        "Lead all engineering, procurement, and construction (EPC) delivery for BESS projects "
        "commissioned under the DISPERON brand in Cyprus and such other markets as the Board may designate;",
        "Manage all relationships with the Cyprus Distribution System Operator (EAC), Transmission "
        "System Operator (TSOC), and Energy Regulatory Authority (CERA), including DSO applications, "
        "SCADA integration approvals, and grid connection certifications;",
        "Act as the primary client relationship manager for the Company's Cyprus customer portfolio, "
        "including executing EPC and LTSA contracts on behalf of the Company;",
        "Oversee commissioning, IEC 60870-5-104 SCADA integration, and project handover for all "
        "BESS installations in Cyprus, coordinating with Voltus for EMS configuration;",
        "Report operational performance, project pipeline, and financial results to the Board monthly.",
    ]:
        bul(doc, item)
    para(doc,
         "Performance Obligation: Mr. Papacosta shall maintain the Company's Cyprus operating licences, "
         "DSO registrations, and SCADA approvals in good standing at all times.", indent=0.5)

    # ── ARTICLE 5 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 5 — INTELLECTUAL PROPERTY")
    h2(doc, "5.1  Ownership of EMS Software IP")
    para(doc,
         "The Parties acknowledge and confirm that all Intellectual Property rights in and to the "
         "EMS Software are owned exclusively by Voltus Energy Sp. z o.o. Nothing in this Agreement "
         "shall be construed as a transfer, assignment, or novation of such IP rights to the Company "
         "or to any individual Shareholder, except as expressly provided in this Article 5 and Article 6.")
    h2(doc, "5.2  Licence Granted to the Company")
    para(doc,
         "Voltus Energy Sp. z o.o. hereby grants to the Company a non-exclusive, worldwide licence "
         "to use, deploy, configure, adapt, and sub-licence the EMS Software to customers under the "
         "DISPERON brand, subject to the commercial terms set out in Article 6.")
    h2(doc, "5.3  Voltus Retains Full IP")
    para(doc,
         "Voltus Energy Sp. z o.o. retains full and exclusive ownership of all IP in the EMS Software "
         "at all times, including after the grant of the Perpetual Licence pursuant to Article 6.3. "
         "The Perpetual Licence grants the Company a right to use the EMS Software but does not "
         "transfer ownership, title, or any other proprietary rights.")
    h2(doc, "5.4  Company IP and DISPERON Brand")
    para(doc,
         "All Intellectual Property created by or on behalf of the Company in connection with the "
         "DISPERON brand — including marketing materials, website content, implementation "
         "methodologies, customer documentation, and the DISPERON trademark — shall be owned "
         "exclusively by the Company. The Voltus Parties and individual Shareholders shall have no "
         "claim to such IP.")
    h2(doc, "5.5  IP Warranty and Indemnity")
    para(doc, "Voltus Energy Sp. z o.o. represents and warrants to the Company that:")
    for item in [
        "It is the sole and unencumbered owner of all Intellectual Property rights in and to the EMS Software;",
        "The EMS Software does not infringe any third-party intellectual property right, including "
        "patents, copyrights, trade secrets, or database rights;",
        "No open-source software components are included in the EMS Software under licences "
        "incompatible with commercial distribution or sub-licensing to end customers;",
        "The EMS Software complies with all applicable EU export control regulations.",
    ]:
        bul(doc, item)
    para(doc,
         "Voltus shall indemnify, defend, and hold harmless the Company and its directors, officers, "
         "and employees against any and all third-party claims, losses, damages, costs, and reasonable "
         "legal fees arising from any breach of the warranties set out in this Article 5.5.")

    # ── ARTICLE 6 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 6 — COMMERCIAL MODEL AND LICENCE FEES")

    h2(doc, "6.1  Catalogue Price")
    para(doc,
         "For each project, the Catalogue Price shall be calculated as: "
         "(Catalogue MW Rate × project installed MW) + (Catalogue MWh Rate × project installed MWh), "
         "using the rates set out in Exhibit C. The Catalogue Price forms the basis for all financial "
         "calculations under this Article 6. Voltus shall confirm the Catalogue Price for each "
         "submitted project brief within 14 (fourteen) calendar days of receipt. Failure to respond "
         "within 14 days constitutes deemed approval at the Exhibit C schedule rates.")

    h2(doc, "6.2  Pre-Milestone Revenue Model")
    para(doc,
         "(a) Per-Project Licence Fee: Prior to the Milestone Date, for each project the Company "
         "shall pay Voltus a Per-Project Licence Fee equal to the Catalogue Price for that project, "
         "payable on a proportional basis corresponding to milestone payments received from the client, "
         "provided that the full Per-Project Licence Fee shall be paid no later than 14 (fourteen) days "
         "from the commissioning date regardless of client payment status.")
    para(doc,
         "(b) Commissioning Fee: The Company sets and collects the Commissioning Fee directly from "
         "the customer. The Commissioning Fee is retained in full by the Company. Voltus has no "
         "entitlement to any portion of the Commissioning Fee, whether before or after the Milestone Date.")

    h2(doc, "6.3  Milestone Definition")
    para(doc,
         "The Milestone shall be deemed achieved when the aggregate installed capacity of BESS projects "
         "deploying the EMS Software under active contracts initiated by or through the Company reaches "
         "500 MWh (five hundred megawatt-hours), calculated cumulatively from the date of execution "
         "of this Agreement. The Milestone has no fixed time limit.")
    for item in [
        "Each project shall be counted upon successful commissioning evidenced by a signed commissioning "
        "acceptance certificate;",
        "Projects contracted but subsequently cancelled prior to commissioning shall not count;",
        "The Director shall maintain a written cumulative log of commissioned projects and MWh capacity, "
        "accessible to all Parties upon 5 (five) business days' written request.",
    ]:
        bul(doc, item)

    h2(doc, "6.4  Perpetual Licence upon Achievement of Milestone")
    para(doc,
         "Upon achievement of the Milestone, Voltus Energy Sp. z o.o. shall automatically grant to "
         "the Company a perpetual, irrevocable, non-exclusive, worldwide, royalty-free licence to "
         "use, deploy, adapt, and sub-licence the EMS Software in its then-current version and all "
         "subsequent versions and updates released by Voltus, without any ongoing licence fee "
         "obligation. This Perpetual Licence shall:")
    for item in [
        "Cover an unlimited number of customer installations worldwide;",
        "Apply to all versions of the EMS Software existing at the Milestone Date and all subsequent "
        "versions released by Voltus Energy Sp. z o.o.;",
        "Require no further payment from the Company to Voltus in respect of the use of the EMS "
        "Software itself;",
        "Survive any change of control, dissolution, or restructuring of Voltus Energy Sp. z o.o., "
        "unless directly caused by a material, uncured breach of this Agreement by the Company.",
    ]:
        bul(doc, item)

    h2(doc, "6.5  Source Code Escrow")
    para(doc,
         "(a) Voltus shall, within 60 (sixty) days of execution of this Agreement, deposit the Escrow "
         "Materials with a mutually agreed independent Escrow Agent (initial proposed provider: NCC "
         "Group Escrow or equivalent agreed in writing within 30 days).")
    para(doc,
         "(b) Voltus shall update the Escrow Materials within 30 (thirty) days of each material release "
         "or update. Annual verification shall be conducted by the Escrow Agent, at cost shared equally.")
    para(doc,
         "(c) Release Triggers: The Escrow Agent shall release the Escrow Materials to the Company "
         "upon: (i) Voltus insolvency, administration, or liquidation with no successor assuming "
         "maintenance obligations within 90 days; (ii) Voltus ceasing contracted maintenance or "
         "support for more than 90 (ninety) consecutive days without cure; (iii) material uncured "
         "breach of Article 7.2 persisting for more than 30 (thirty) days after written notice.")
    para(doc,
         "(d) Following release, the Company may use the Escrow Materials solely to continue operating "
         "and maintaining the EMS Software for existing and future customers under the DISPERON brand. "
         "The Company shall not sub-licence the source code itself to any third party.")
    para(doc,
         "(e) Costs of the escrow arrangement shall be borne equally between Voltus and the Company.")

    h2(doc, "6.6  Development Bonus (Post-Milestone)")
    para(doc,
         "(a) Following the Milestone Date and for the duration of this Agreement, for each customer "
         "project the Company shall pay to the Voltus Parties a Development Bonus calculated as:")
    for item in [
        "25% (twenty-five percent) of the Catalogue Price for projects in countries where the Company "
        "does not hold Exclusivity;",
        "30% (thirty percent) of the Catalogue Price for projects in countries where the Company "
        "holds Exclusivity.",
    ]:
        bul(doc, item)
    para(doc,
         "(b) The Development Bonus rates shall be subject to review every 3 (three) years by written "
         "agreement of all Parties, within the range of 20% to 35%, provided Voltus produces quarterly "
         "development expenditure summaries demonstrating active R&D investment.")
    para(doc,
         "(c) The Development Bonus is calculated exclusively on the Catalogue Price as defined in "
         "Article 1. Commissioning Fees and any other service fees are expressly excluded.")
    para(doc,
         "(d) The Development Bonus shall be designated for ongoing research, development, and "
         "enhancement of the EMS Software. The Voltus Parties shall provide a quarterly summary of "
         "development activities funded by such payments.")
    para(doc,
         "(e) The Development Bonus shall be payable within 14 (fourteen) days of the Company receiving "
         "cleared payment from the relevant customer. Prepayments and partial payments shall not "
         "trigger the Development Bonus until the commissioning date of the relevant project.")

    h2(doc, "6.7  EMS Subscription Fee (Client-Facing) and Voltus Subscription Fee")
    para(doc,
         "(a) EMS Subscription Fee charged to customers: The Company shall charge each customer "
         "an annual EMS Subscription Fee of EUR 400 (four hundred euros) per MWh of installed "
         "project capacity per year (the 'EMS Subscription Fee'). The EMS Subscription Fee shall:")
    for item in [
        "Be invoiced to the customer annually, starting from the commissioning date (PAC) of the "
        "relevant project and on each subsequent anniversary;",
        "Be billed separately from any LTSA, O&M, or physical maintenance fees;",
        "Be payable by the customer within 14 (fourteen) days of each annual invoice.",
    ]:
        bul(doc, item)
    para(doc,
         "(b) Voltus Subscription Fee: In respect of each commissioned project, the Company shall "
         "pay to Voltus an annual Voltus Subscription Fee equal to 20% (twenty percent) of the "
         "EMS Subscription Fee received from the relevant customer in that year. The Voltus "
         "Subscription Fee shall be payable within 14 (fourteen) days of the Company's receipt "
         "of each annual EMS Subscription Fee payment from the customer.")
    para(doc,
         "(c) Example: For a 30 MWh project — EMS Subscription Fee charged to customer: "
         "EUR 12,000/year. Voltus Subscription Fee: 20% × EUR 12,000 = EUR 2,400/year. "
         "DISPERON retains: EUR 9,600/year.")
    para(doc,
         "(d) The EMS Subscription Fee covers: (i) all EMS and SCADA software updates (minor and "
         "major releases, security patches); (ii) product support (diagnostic and hotfix services "
         "for software defects); (iii) regulatory compliance updates (EU and national grid code "
         "changes implemented within 90 days of official publication); (iv) remote SCADA dashboard "
         "access and DSO integration maintenance. It does not cover on-site configuration, "
         "commissioning, or end-customer physical support.")
    para(doc,
         "(e) Non-receipt of the EMS Subscription Fee from the customer for more than 90 "
         "(ninety) consecutive days shall not affect the Company's licence to use the EMS Software. "
         "If the Company fails to pass 20% of received EMS Subscription Fees to Voltus within the "
         "14-day payment window, Voltus may suspend delivery of new software updates after 90 days "
         "written notice without cure. The licence to use the then-current version shall remain "
         "in full force under all circumstances.")

    h2(doc, "6.8  Records and Audit")
    para(doc,
         "The Company shall maintain accurate records of all project deployments, Catalogue Prices, "
         "Licence Fees paid, Commissioning Fees received, Development Bonus payments, and Subscription "
         "Fee payments. Voltus shall have the right to commission an independent audit of these records "
         "once per calendar year, on not less than 20 (twenty) business days' written notice. Audit "
         "costs shall be borne by Voltus unless the audit reveals an underpayment exceeding 5% (five "
         "percent) of amounts due, in which case the Company shall bear the audit costs and promptly "
         "remedy the underpayment.")

    # ── ARTICLE 7 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 7 — OPERATIONAL STRUCTURE")
    h2(doc, "7.1  Sales and Market Development")
    para(doc,
         "The Company, acting through Lighthief International, Mr. Kamil Talar, and Mr. Kostas "
         "Alexander Papacosta, shall be primarily responsible for the commercial development of the "
         "DISPERON brand and generation of customer pipeline across the Territory, including direct "
         "sales, partner channels, tender responses, EPC delivery, and industry representation. "
         "Mr. Papacosta shall lead all market development, client management, and EPC execution in "
         "the Republic of Cyprus.")
    h2(doc, "7.2  Technology and Product Development")
    para(doc,
         "Voltus Energy Sp. z o.o. shall be primarily responsible for the continuous development, "
         "maintenance, security, and improvement of the EMS Software, maintaining it in a commercially "
         "deployable and regulatory-compliant state at all times. Specifically, Voltus shall:")
    for item in [
        "Respond to severity-classified support requests within the timeframes set out in Exhibit D;",
        "Maintain cloud-hosted EMS components at a minimum availability of 99.5% measured per calendar "
        "month, excluding planned maintenance windows notified at least 5 days in advance;",
        "Implement changes required by new or amended EU or national grid code regulations within 90 "
        "(ninety) days of official publication, and notify the Company within 14 days of becoming "
        "aware of any such regulatory change;",
        "Provide the Company with a quarterly product roadmap and development activity report;",
        "Maintain cybersecurity compliance with NIS2 and applicable EU standards.",
    ]:
        bul(doc, item)
    para(doc,
         "Persistent failure by Voltus to meet the obligations in this Article 7.2 constitutes a "
         "material breach entitling the Company to activate the escrow release mechanism in Article 6.5.")
    h2(doc, "7.3  Commissioning and Implementation")
    para(doc,
         "On-site and remote commissioning, system integration, and project handover services shall "
         "be performed by the Company or a designated entity within the Lighthief International group. "
         "Voltus shall provide remote technical support, documentation, and training as included in "
         "the Licence and Subscription Fee. All Commissioning revenues belong to the Company.")
    h2(doc, "7.4  Customer Contracts")
    para(doc,
         "All customer contracts shall be entered into by the Company from the date of execution of "
         "this Agreement. The Company is the legal counterparty to all end customers. Voltus acts as "
         "software supplier to the Company, not to end customers directly, unless otherwise agreed "
         "in writing.")

    # ── ARTICLE 8 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 8 — EXCLUSIVITY AND MARKET MODEL")
    h2(doc, "8.1  Principle of Earned Exclusivity")
    para(doc,
         "Exclusivity is not granted automatically. It may be earned by the Company in specific "
         "countries on the basis of demonstrated commercial activity and achieved volume thresholds.")
    h2(doc, "8.2  Pre-Conditions for Exclusivity")
    para(doc, "Before any exclusivity is recognised in a specific country, the Company shall provide:")
    for item in [
        "A documented project pipeline for that country, including client names, project sizes, and "
        "estimated timelines;",
        "A written sales plan covering a minimum of 12 (twelve) months;",
        "Evidence of active commercial engagement, including customer meetings, proposals, or signed "
        "letters of intent.",
    ]:
        bul(doc, item)
    h2(doc, "8.3  Exclusivity Volume Thresholds")
    para(doc,
         "Exclusivity shall be granted upon the Company achieving the following minimum installed "
         "capacity within 24 (twenty-four) months of the Company's first commercial project in that country:")
    tbl(doc, [
        ["Market Category", "Countries", "Minimum Volume", "Annual Maintenance"],
        ["Large markets", "Germany, Italy, Spain, France", "500 MWh within 24 months", "100 MWh / year"],
        ["Mid-sized markets", "Netherlands, Romania, Greece, Czech Republic, Austria",
         "250 MWh within 24 months", "50 MWh / year"],
        ["Smaller markets", "Cyprus, other EU states", "100 MWh within 24 months", "20 MWh / year"],
    ], col_widths=[3.2, 4.5, 3.8, 3.5])
    divider(doc)

    h2(doc, "8.4  Maintenance and Lapse of Exclusivity")
    para(doc,
         "(a) Following the initial grant of exclusivity, the Company must achieve the minimum annual "
         "installed capacity set out in Article 8.3 in each subsequent 12-month period.")
    para(doc,
         "(b) Restricted Exclusivity: If the Company fails to meet the Annual Maintenance threshold "
         "in any single 12-month period, the country shall enter Restricted Exclusivity status for "
         "the following 12-month period. During Restricted Exclusivity: (i) the Company retains the "
         "right to complete all existing contracted projects; (ii) Voltus may independently pursue "
         "new commercial opportunities in that country with customers not in the Company's documented "
         "pipeline at the date of restriction.")
    para(doc,
         "(c) Full Lapse: Failure to meet the Annual Maintenance threshold for two consecutive annual "
         "periods shall result in automatic lapsing of exclusivity for that country. The Company may "
         "re-apply for exclusivity after 12 (twelve) months following lapse.")

    h2(doc, "8.5  Initial Country Limit and Moratorium")
    para(doc,
         "At the time of execution of this Agreement, the Company may seek exclusivity in a maximum "
         "of 5 (five) countries simultaneously. Poland, Lithuania, Latvia, and Estonia (the 'Moratorium "
         "Countries') shall be subject to a 24 (twenty-four) month moratorium from the date of "
         "execution of this Agreement, during which: (a) Voltus may continue and develop existing "
         "commercial activities in those countries; (b) the Company may conduct sales activities but "
         "may not claim exclusivity. Following expiry of the moratorium, the Company may apply for "
         "exclusivity in the Moratorium Countries under the standard terms of Article 8.3.")

    h2(doc, "8.6  Voltus Direct Rights — BESS Manufacturer Partners")
    para(doc,
         "Notwithstanding any exclusivity granted to the Company, Voltus Energy Sp. z o.o. retains "
         "the right to supply the EMS Software directly to BESS Manufacturer Partners listed in "
         "Exhibit B. This right: (i) applies only to the manufacturers listed in Exhibit B; (ii) does "
         "not extend to any entity not listed in Exhibit B without the written consent of all Parties; "
         "(iii) is not subject to geographic restriction for listed Exhibit B manufacturers.")
    para(doc,
         "BESS Manufacturer Partner Referral Obligation: Where a project from a BESS Manufacturer "
         "Partner has been the subject of prior commercial engagement by the Company — evidenced by "
         "a written proposal, meeting record, or registered customer entry — Voltus shall notify the "
         "Company within 5 (five) business days and refer such project to the Company under the "
         "DISPERON brand. In the absence of agreement on commercial terms within 15 business days, "
         "the standard Development Bonus rate in Article 6.6 shall apply. Voltus shall provide the "
         "Company with a monthly written summary of EMS supply activities to BESS Manufacturer Partners.")

    h2(doc, "8.7  Transparency and Customer Registration")
    para(doc,
         "(a) Each Party may register customer opportunities where it has initiated active commercial "
         "engagement, evidenced by a written proposal, meeting record, or documented pipeline entry. "
         "(b) A registered customer opportunity shall be respected by the other Party provided it is "
         "actively pursued and updated. (c) Where both Parties engage with the same customer or "
         "project, the Parties shall first seek to cooperate in good faith to jointly pursue the "
         "opportunity. (d) In the event of dispute, the Director shall resolve the matter within "
         "15 (fifteen) business days.")

    h2(doc, "8.8  Non-Competition Undertaking of Voltus")
    para(doc,
         "Voltus Energy Sp. z o.o. undertakes, for the duration of this Agreement, that: (a) it shall "
         "maintain full transparency with the Company regarding its direct sales activities within the "
         "Territory; (b) if a customer approaches Voltus directly for EMS services, and that same "
         "customer has previously been approached by or is in active discussion with the Company, "
         "Voltus shall decline to engage independently and shall refer such customer to the Company "
         "under the DISPERON brand. This undertaking does not apply to: (i) customers originally "
         "introduced by Voltus independently of the Company; (ii) customers of BESS Manufacturer "
         "Partners listed in Exhibit B where the BESS Manufacturer Partner Referral Obligation in "
         "Article 8.6 does not apply.")

    h2(doc, "8.9  Dissatisfied Customer Procedure")
    para(doc,
         "(a) If an end customer raises a formal complaint regarding DISPERON's performance, the "
         "Company shall have 60 (sixty) days to remedy the identified deficiency (the 'Customer "
         "Cure Period').")
    para(doc,
         "(b) A complaint shall only be considered 'confirmed' if: (i) it is submitted in writing "
         "by the customer; (ii) the Company fails to remedy the issue within the Customer Cure "
         "Period; and (iii) the failure is independently verified by the Board.")
    para(doc,
         "(c) If more than 3 (three) confirmed complaints are recorded against the Company in "
         "respect of a single customer within any 24-month period, Voltus may, after written "
         "notice to the Company, offer to serve that specific customer directly. Any such direct "
         "supply shall be subject to the standard Development Bonus rate payable to DISPERON.")
    para(doc,
         "(d) Complaints arising from factors outside the Company's control — including regulatory "
         "delays, force majeure, grid operator requirements, or Voltus's own software deficiencies "
         "— shall not count as confirmed complaints.")

    # ── ARTICLE 9 ─────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 9 — TRANSFER OF SHARES")
    h2(doc, "9.1  Lock-Up Period")
    para(doc,
         "For 24 (twenty-four) months from the date of this Agreement (the 'Lock-Up Period'), no "
         "Shareholder shall sell, transfer, assign, pledge, or otherwise dispose of any Shares "
         "without the prior written consent of Lighthief International.")
    h2(doc, "9.2  Right of First Refusal")
    para(doc,
         "Following the Lock-Up Period, if any Shareholder proposes to transfer Shares to a third "
         "party: written notice must be given to all other Shareholders specifying the number of "
         "Shares, proposed price, and buyer identity; each non-selling Shareholder has 30 (thirty) "
         "days to exercise a right of first refusal on a pro-rata basis at the same price and terms.")
    h2(doc, "9.3  Drag-Along Right")
    para(doc,
         "If Lighthief International proposes to sell 100% of its Shares to a bona fide purchaser "
         "requiring 100% of the issued capital, Lighthief International may require all other "
         "Shareholders to sell their Shares on the same terms, on 30 (thirty) days' written notice.")
    h2(doc, "9.4  Tag-Along Right")
    para(doc,
         "If Lighthief International proposes to transfer more than 50% of its Shares to a third "
         "party, each other Shareholder has the right, within 20 (twenty) days, to include their "
         "Shares in such sale on the same terms.")

    # ── ARTICLE 10 ────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 10 — GOVERNANCE AND BOARD OF DIRECTORS")
    h2(doc, "10.1  Current Directors")
    para(doc,
         "As of the date of this Agreement, the duly appointed Directors of the Company are:\n"
         "(a) Mr. Arkadiusz Sybaris — Director and Secretary, tax resident of the Republic of Cyprus;\n"
         "(b) Mr. Kostas Alexander Papacosta — Director and Cyprus Managing Director, tax resident "
         "of the Republic of Cyprus, permanently residing in Cyprus.")
    h2(doc, "10.2  Board Composition")
    para(doc,
         "(a) Lighthief International, as Majority Shareholder, shall have the right to appoint up "
         "to 3 (three) Directors to the Board, including the Director and Secretary role. "
         "Mr. Kostas Alexander Papacosta shall be one of the Directors appointed by Lighthief "
         "International for the duration of this Agreement.\n"
         "(b) The Shareholders other than Lighthief International and Mr. Papacosta — being "
         "Mr. Marcin Szumiło, Mr. Andrzej Lechowicz, and Mr. Kamil Talar jointly — shall have "
         "the right to appoint 1 (one) Director from among themselves to the Board.")
    h2(doc, "10.3  Appointment and Removal")
    para(doc,
         "Directors appointed by Lighthief International may be appointed and removed at any time "
         "by written resolution of Lighthief International, provided that the removal of "
         "Mr. Papacosta as Director shall additionally require a resolution of the Board. "
         "The Shareholder Director may be replaced by written agreement of Mr. Szumiło, "
         "Mr. Lechowicz, and Mr. Talar.")
    h2(doc, "10.4  Decisions Requiring Majority Shareholder Consent")
    para(doc, "The following matters shall require the written consent of Shareholders holding a "
         "majority of Shares:")
    for item in [
        "Amendment of the Company's Articles of Association;",
        "Issuance of new shares or share capital increases;",
        "Material change to the Company's principal business activity;",
        "Entry into any related-party transaction exceeding EUR 50,000 in value;",
        "Appointment or removal of auditors.",
    ]:
        bul(doc, item)
    h2(doc, "10.5  Reserved Matters Requiring 85% Shareholder Consent")
    para(doc, "The following matters shall require the written consent of Shareholders holding not "
         "less than 85% (eighty-five percent) of all issued Shares:")
    for item in [
        "Material amendment to the commercial model set out in Article 6, including the Development "
        "Bonus rate, the Subscription Fee rate, or the Catalogue Price schedule in Exhibit C;",
        "Voluntary dissolution or winding up of the Company;",
        "Any transaction or resolution that would impair, encumber, or terminate the Perpetual "
        "Licence or the source code escrow arrangement under Article 6.5.",
    ]:
        bul(doc, item)
    h2(doc, "10.6  Shareholder Meetings")
    para(doc,
         "The Company shall hold at least one annual meeting of Shareholders. Additional meetings "
         "may be called by the Director or by any Shareholder holding at least 20% of issued capital, "
         "on not less than 14 (fourteen) days' notice. Meetings may be held in person or by video conference.")

    # ── ARTICLE 11 ────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 11 — CONFIDENTIALITY")
    h2(doc, "11.1  Confidential Information")
    para(doc,
         "Each Party shall keep confidential all non-public information received in connection with "
         'this Agreement, including technical specifications, source code, customer data, financial '
         'terms, and business strategies ("Confidential Information").')
    h2(doc, "11.2  Duration")
    para(doc,
         "The confidentiality obligation applies during the term of this Agreement and for 5 (five) "
         "years following its termination or expiry.")
    h2(doc, "11.3  Permitted Disclosure")
    para(doc,
         "Confidential Information may be disclosed only to employees, contractors, or advisors with "
         "a need to know who are bound by equivalent obligations, or where required by applicable law.")

    # ── ARTICLE 12 ────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 12 — GOVERNING LAW AND DISPUTE RESOLUTION")
    h2(doc, "12.1  Governing Law")
    para(doc,
         "This Agreement shall be governed by and construed in accordance with the laws of the "
         "Republic of Cyprus.")
    h2(doc, "12.2  Negotiation")
    para(doc,
         'In the event of any dispute arising out of or in connection with this Agreement (a "Dispute"), '
         "the Parties shall attempt resolution through good-faith negotiation within 30 (thirty) days "
         "of written notice of the Dispute.")
    h2(doc, "12.3  Arbitration")
    para(doc,
         "If a Dispute is not resolved by negotiation within the period set out in Article 12.2, "
         "it shall be referred to final and binding arbitration in Limassol, Cyprus, under the "
         "Cyprus Arbitration Law (Cap. 4), conducted in the English language.")

    # ── ARTICLE 13 — TERM AND EXIT ────────────────────────────────────────────
    h1(doc, "ARTICLE 13 — TERM AND EXIT MECHANISM")

    h2(doc, "13.1  Initial Term and Renewal")
    para(doc,
         "This Agreement shall be concluded for an initial period of 5 (five) years from the date of "
         "execution (the 'Initial Term'). Upon expiry of the Initial Term, the Agreement shall "
         "automatically renew for successive periods of 2 (two) years unless either Party gives "
         "not less than 12 (twelve) months' written notice of non-renewal prior to the expiry of "
         "the then-current term.")

    h2(doc, "13.2  Element A — Time-Based Exit with Trail Commission")
    para(doc,
         "(a) Following the expiry of the Initial Term, either Party may terminate this Agreement "
         "by providing 12 (twelve) months' written notice to all other Parties.")
    para(doc,
         "(b) If Voltus exercises the right to exit under this Article 13.2, Voltus shall pay to "
         "the Company a Trail Commission for a period of 12 (twelve) months following the effective "
         "date of termination. The Trail Commission shall be calculated as 10% (ten percent) of the "
         "Development Bonus that would have been payable to Voltus on any customer projects that "
         "Voltus directly supplies in the former DISPERON territory during the 12-month period. "
         "The Trail Commission reflects the commercial relationships built by DISPERON.")
    para(doc,
         "(c) If the Company exercises the right to exit under this Article 13.2, no Trail Commission "
         "shall be payable.")

    h2(doc, "13.3  Element B — Performance-Based Exit")
    para(doc,
         "Either Party may terminate this Agreement without compensation or notice period upon the "
         "occurrence of a Triggering Event attributable to the other Party. The following constitute "
         "Triggering Events:")
    para(doc, "Triggering Events giving Voltus the right to exit:", bold=True)
    for item in [
        "The Company's annual commissioned MWh volume falls below 50% of the applicable Annual "
        "Maintenance threshold in Article 8.3 for 2 (two) consecutive years;",
        "The Company fails to sign any new customer contracts for 9 (nine) consecutive months in "
        "any country where it holds Exclusivity;",
        "The Company is in material breach of its financial obligations under Articles 6.6 or 6.7 "
        "for more than 90 (ninety) days after written notice.",
    ]:
        bul(doc, item)
    para(doc, "Triggering Events giving the Company the right to exit:", bold=True)
    for item in [
        "Voltus fails to deliver contracted software updates for more than 6 (six) months without "
        "reasonable technical justification;",
        "Voltus loses product certifications required for EU regulatory compliance and fails to "
        "restore them within 90 (ninety) days;",
        "Voltus increases the Catalogue Price rates in Exhibit C by more than 50% in any single "
        "annual revision without the written agreement of all Parties.",
    ]:
        bul(doc, item)
    para(doc,
         "Procedure: The Party asserting a Triggering Event must provide written notice with "
         "supporting documentation. The other Party has 30 (thirty) days to cure. If uncured, "
         "exit becomes effective 30 (thirty) days after the cure period expires.")

    h2(doc, "13.4  Element C — 12-Month Graduated Market Transition")
    para(doc,
         "Following any exit under Articles 13.2 or 13.3, the following graduated transition "
         "period shall apply (the '12-Month Transition'):")
    tbl(doc, [
        ["Phase", "Months", "Company Rights", "Voltus Rights"],
        ["Full Protection", "0–4",
         "Full Exclusivity maintained. Company may sign new contracts with clients in active negotiation.",
         "Voltus prepares own sales infrastructure. No new sales activities in Company's exclusive countries."],
        ["Restricted", "5–8",
         "Company completes existing contracts. No new signings. ROFR on clients from past 24 months (30-day window).",
         "Voltus may pursue new leads with clients not in Company's documented pipeline."],
        ["Shared Market", "9–12",
         "Company ROFR narrowed to actively served clients (past 12 months only).",
         "Voltus may compete on all new opportunities."],
        ["Full Freedom", "13+",
         "No ROFR, no exclusivity protection.",
         "Voltus has full market freedom in all countries."],
    ], col_widths=[2.5, 2, 5, 5])
    divider(doc)
    para(doc,
         "During the 12-Month Transition, the Development Bonus and Subscription Fee obligations "
         "continue on all projects commissioned or contracted during the transition period.")

    # ── ARTICLE 14 ────────────────────────────────────────────────────────────
    h1(doc, "ARTICLE 14 — GENERAL PROVISIONS")
    h2(doc, "14.1  Entire Agreement")
    para(doc,
         "This Agreement constitutes the entire agreement of the Parties with respect to its subject "
         "matter and supersedes all prior understandings, negotiations, and representations.")
    h2(doc, "14.2  Amendment")
    para(doc, "This Agreement may only be amended by a written instrument signed by all Parties.")
    h2(doc, "14.3  Severability")
    para(doc,
         "If any provision of this Agreement is found to be invalid, illegal, or unenforceable, the "
         "remaining provisions shall continue in full force and effect.")
    h2(doc, "14.4  Waiver")
    para(doc,
         "No failure or delay by any Party in exercising any right or remedy under this Agreement "
         "shall constitute a waiver thereof.")
    h2(doc, "14.5  Notices")
    para(doc,
         "All notices under this Agreement shall be in writing, delivered by email with read receipt "
         "or by registered post to the addresses stated in the Preamble or as otherwise notified in writing.")
    h2(doc, "14.6  Counterparts and Electronic Signatures")
    para(doc,
         "This Agreement may be executed in counterparts, each of which shall constitute an original. "
         "Electronic signatures shall be valid and binding on all Parties. The English version of this "
         "Agreement shall prevail over any translation, including the Polish version.")

    # ── SIGNATURES ────────────────────────────────────────────────────────────
    doc.add_page_break()
    p_s = doc.add_paragraph(); p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = p_s.add_run("SIGNATURES")
    rs.font.size = Pt(16); rs.font.bold = True; rs.font.color.rgb = BLUE

    doc.add_paragraph()
    para(doc,
         "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written "
         "above. Each signatory below represents and warrants that they have full authority to execute "
         "this Agreement on behalf of the party they represent.", italic=True)
    doc.add_paragraph()

    sig_block(doc, "PARTY 1 — LIGHTHIEF INTERNATIONAL",
              "Represented by: Mr. Arkadiusz Sybaris",
              "Title: Director",
              ref="Party 1")
    sig_block(doc, "PARTY 2 — LIGHTHIEF EUBESS LTD (HE 474192)",
              "Represented by: Mr. Arkadiusz Sybaris",
              "Title: Director and Secretary",
              ref="Party 2")
    sig_block(doc, "PARTY 3 — MR. MARCIN SZUMIŁO",
              "Individual Shareholder",
              ref="Party 3")
    sig_block(doc, "PARTY 4 — MR. ANDRZEJ LECHOWICZ",
              "Individual Shareholder",
              ref="Party 4")
    sig_block(doc, "PARTY 5 — MR. KAMIL TALAR",
              "Individual Shareholder",
              ref="Party 5")
    sig_block(doc, "PARTY 6 — MR. KOSTAS ALEXANDER PAPACOSTA",
              "Individual Shareholder and Cyprus Managing Director",
              ref="Party 6")
    sig_block(doc, "PARTY 7 — VOLTUS ENERGY SP. Z O.O.",
              "Represented by: Mr. Marcin Szumiło, Management Board Member",
              extra="Represented by: Mr. Andrzej Lechowicz, Management Board Member",
              ref="Party 7")

    # ── EXHIBIT A ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "EXHIBIT A — SHARE CERTIFICATES")
    para(doc,
         "The following share certificates of Lighthief EUBESS Ltd (HE 474192) shall be issued "
         "within 30 (thirty) days of registration with the Cyprus Registrar of Companies:")
    for item in [
        "Current share certificate issued to Lighthief International (100% prior to transfer);",
        "New share certificate issued to Mr. Kostas Alexander Papacosta (10%);",
        "New share certificate issued to Mr. Kamil Talar (13%);",
        "New share certificate issued to Mr. Marcin Szumiło (6%);",
        "New share certificate issued to Mr. Andrzej Lechowicz (6%);",
        "Updated share certificate issued to Lighthief International (65% post-transfer).",
    ]:
        bul(doc, item)

    # ── EXHIBIT B ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "EXHIBIT B — BESS MANUFACTURER PARTNERS")
    para(doc,
         "The following is the complete and exclusive list of BESS hardware manufacturers in respect "
         "of whom Voltus Energy Sp. z o.o. retains direct supply rights pursuant to Article 8.6. "
         "Any addition to this list requires the prior written consent of all Parties and a signed "
         "amendment to this Agreement.")
    divider(doc)
    tbl(doc, [
        ["No.", "Manufacturer Name", "Registered Country", "Date Added"],
        ["1.", "Linyang Energy Co., Ltd.", "People's Republic of China", "Date of Agreement"],
        ["2.", "________________________", "________________________", "________________"],
        ["3.", "________________________", "________________________", "________________"],
        ["4.", "________________________", "________________________", "________________"],
    ], col_widths=[1.2, 6, 4.5, 3.3])
    divider(doc)
    para(doc, "Confirmed by Voltus Energy Sp. z o.o.:")
    para(doc, "Mr. Marcin Szumiło  ________________________________     Date: __________________")
    para(doc, "Mr. Andrzej Lechowicz  ________________________________   Date: __________________")
    divider(doc)
    para(doc, "Confirmed by Lighthief EUBESS Ltd:")
    para(doc, "Mr. Arkadiusz Sybaris  ________________________________    Date: __________________")

    # ── EXHIBIT C — FEE SCHEDULE ──────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "EXHIBIT C — CATALOGUE PRICE AND LICENCE FEE SCHEDULE")
    para(doc,
         "This Exhibit C forms part of the Agreement and governs all Catalogue Price calculations "
         "under Articles 6.1, 6.2, 6.6, and 6.7.")
    h2(doc, "C.1  Catalogue Price Rates")
    tbl(doc, [
        ["Component", "Rate", "Unit", "Effective Date"],
        ["Catalogue MW Rate",  "EUR [___]", "per MW installed",  "Date of Agreement"],
        ["Catalogue MWh Rate", "EUR [___]", "per MWh installed", "Date of Agreement"],
    ], col_widths=[5, 3, 4, 4])
    divider(doc)
    para(doc,
         "The rates in brackets above shall be agreed by all Parties within 14 (fourteen) days of "
         "execution of this Agreement and inserted by written addendum signed by all Parties. "
         "Until agreed, no project brief shall be deemed submitted for the purposes of Article 6.2.")
    h2(doc, "C.2  Annual Escalation")
    para(doc,
         "The Catalogue Rates shall be adjusted annually on 1 January of each year by the change "
         "in the EU Harmonised Index of Consumer Prices (HICP) published by Eurostat for the "
         "preceding calendar year. HICP ≤ 0% shall result in no change (no downward adjustment). "
         "The first adjustment shall occur on 1 January of the year following execution.")
    h2(doc, "C.3  Rate Review")
    para(doc,
         "Either Party may initiate a review of the Catalogue Rates by written notice given at least "
         "60 (sixty) days before the relevant anniversary date. The Parties have 30 (thirty) days to "
         "agree revised rates. In the absence of agreement, the HICP escalation mechanism applies. "
         "Any agreed rate revision requires a signed amendment to this Exhibit C.")
    h2(doc, "C.4  Development Bonus Calculation")
    para(doc,
         "Development Bonus = Applicable Rate (25% or 30%) × Catalogue Price\n"
         "Where: Catalogue Price = (Catalogue MW Rate × MW) + (Catalogue MWh Rate × MWh)\n"
         "Subscription Fee = 15% × Development Bonus (annual, per project, from first anniversary)")
    h2(doc, "C.5  Confirmation Procedure")
    para(doc,
         "Voltus shall confirm the applicable Catalogue Price for each submitted project brief within "
         "14 (fourteen) calendar days of receipt. A project brief shall include: project name, "
         "location, MW and MWh capacity, customer name, and estimated commissioning date. Failure "
         "to respond within 14 days constitutes deemed approval at the schedule rates.")
    divider(doc)
    para(doc, "Agreed and signed by all Parties:")
    divider(doc)
    for name in [
        "Lighthief International — Mr. Arkadiusz Sybaris",
        "Voltus Energy Sp. z o.o. — Mr. Marcin Szumiło",
        "Voltus Energy Sp. z o.o. — Mr. Andrzej Lechowicz",
    ]:
        para(doc, f"{name}  ________________________________     Date: __________________")

    # ── EXHIBIT D ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "EXHIBIT D — SOFTWARE SUPPORT ANNEX")
    para(doc,
         "This Software Support Annex forms part of the Agreement and governs Voltus Energy "
         "Sp. z o.o.'s support and maintenance obligations under Article 7.2.")
    h2(doc, "D.1  Severity Classification and SLA Targets")
    tbl(doc, [
        ["Severity", "Description", "Examples", "Response SLA", "Resolution SLA"],
        ["P1 — Critical", "System unavailable or safety-critical function impaired",
         "EMS offline; DSO command failure", "2 hours", "24 hours"],
        ["P2 — Major", "Core function degraded; no workaround",
         "SOC error; control loop failure", "8 business hours", "5 business days"],
        ["P3 — Minor", "Non-critical function impaired; workaround available",
         "Reporting anomaly; UI defect", "2 business days", "30 business days"],
        ["P4 — Cosmetic", "Minor issue; no operational impact",
         "Display formatting", "5 business days", "Next planned release"],
    ], col_widths=[2.5, 3.5, 3, 2.5, 3.5])
    divider(doc)
    h2(doc, "D.2  Availability Target")
    para(doc,
         "Cloud-hosted EMS components shall maintain a minimum availability of 99.5% measured per "
         "calendar month, excluding scheduled maintenance windows. Scheduled maintenance requires "
         "a minimum of 5 (five) calendar days' advance written notice and shall not exceed 4 hours "
         "per occurrence.")
    h2(doc, "D.3  Regulatory Update Obligation")
    para(doc,
         "Voltus shall implement changes required by new or amended EU or national grid code "
         "regulations within 90 (ninety) days of official publication. Voltus shall notify the "
         "Company within 14 (fourteen) days of becoming aware of any such regulatory change.")
    h2(doc, "D.4  Service Credits")
    para(doc,
         "(a) Availability: For each month where availability falls below 99.5%, the Company may "
         "offset from the next Subscription Fee payment an amount equal to 5% of the monthly "
         "Subscription Fee for each full percentage point below 99.5%.")
    para(doc,
         "(b) P1 SLA Failures: For each P1 incident where the response or resolution SLA is "
         "exceeded, Voltus shall credit the Company EUR 500 (five hundred euros) per day of delay.")
    para(doc,
         "(c) Credits are the Company's sole financial remedy for SLA failures and do not limit "
         "other remedies available for material breach.")
    h2(doc, "D.5  Annual Review")
    para(doc,
         "SLA targets and credit values shall be reviewed annually at the written request of either "
         "Party. Any amendment requires written agreement of all Parties.")

    # ── SAVE ──────────────────────────────────────────────────────────────────
    out = "/Volumes/T7 Grey/solinvest/DisperonEMS/docs/contract/DISPERON_SHA_v5_EN.docx"
    doc.save(out)
    print(f"Saved: {out}")


build()
