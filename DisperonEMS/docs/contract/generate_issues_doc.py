"""Generate DISPERON_SHA_v2_ISSUES.docx — annotated issues analysis."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

RED    = RGBColor(0xC0, 0x00, 0x00)
AMBER  = RGBColor(0xFF, 0x80, 0x00)
GREEN  = RGBColor(0x37, 0x86, 0x4E)
BLUE   = RGBColor(0x1F, 0x49, 0x7D)
BLACK  = RGBColor(0x00, 0x00, 0x00)
GREY   = RGBColor(0x40, 0x40, 0x40)
LTGREY = RGBColor(0xF2, 0xF2, 0xF2)

def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)

def add_heading(doc, text, level=1, colour=BLUE):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = colour
        run.font.bold = True
    return p

def add_para(doc, text, bold=False, colour=BLACK, size=10, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = colour
    return p

def add_issue_block(doc, number, title, priority_label, priority_colour,
                    article_ref, quote, problem, fix):
    """Render a single issue block as a bordered table."""
    table = doc.add_table(rows=0, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    row = table.add_row()
    cell = row.cells[0]
    set_cell_bg(cell, 'D6E4F0')
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(f"ISSUE {number}  ·  {title}    [{priority_label}]")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = priority_colour

    # Article reference row
    if article_ref:
        row = table.add_row()
        cell = row.cells[0]
        set_cell_bg(cell, 'EAF0FB')
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(f"Contract Reference: {article_ref}")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = GREY

    # Quote row
    if quote:
        row = table.add_row()
        cell = row.cells[0]
        set_cell_bg(cell, 'FFF9E6')
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(f'"{quote}"')
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x40, 0x00)

    # Problem
    row = table.add_row()
    cell = row.cells[0]
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run("Problem:  ")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RED
    run2 = p.add_run(problem)
    run2.font.size = Pt(10)
    run2.font.color.rgb = BLACK

    # Fix
    row = table.add_row()
    cell = row.cells[0]
    set_cell_bg(cell, 'E8F5E9')
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run("Proposed Fix:  ")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREEN
    run2 = p.add_run(fix)
    run2.font.size = Pt(10)
    run2.font.color.rgb = BLACK

    doc.add_paragraph()  # spacer

def build_issues_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DISPERON SHA v2 — LEGAL & COMMERCIAL ISSUES REGISTER")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        "Lighthief EUBESS Ltd (HE 474192)  ·  Prepared: April 2026  ·  CONFIDENTIAL"
    )
    run2.font.size = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = GREY

    doc.add_paragraph()

    # Executive summary box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, 'EEF3FB')
    p = c.paragraphs[0]
    p.clear()
    run = p.add_run("EXECUTIVE SUMMARY\n")
    run.font.bold = True; run.font.size = Pt(11); run.font.color.rgb = BLUE
    run2 = p.add_run(
        "This document identifies 11 legal and commercial issues in the DISPERON SHA v2 "
        "between Lighthief EUBESS Ltd, the Voltus Parties, and individual shareholders. "
        "Issues are graded CRITICAL / HIGH / MEDIUM / LOW. "
        "A revised clean draft (SHA v3) incorporating all proposed fixes accompanies this register."
    )
    run2.font.size = Pt(10); run2.font.color.rgb = BLACK
    doc.add_paragraph()

    # Priority legend
    add_heading(doc, "PRIORITY LEGEND", level=2)
    legend = doc.add_table(rows=1, cols=4)
    legend.style = 'Table Grid'
    labels = [
        ("🔴  CRITICAL", "FF0000", "Contract cannot safely be signed"),
        ("🟡  HIGH",      "FF8000", "Significant commercial exposure"),
        ("🔵  MEDIUM",    "1F497D", "Manageable but should be addressed"),
        ("⚪  LOW",       "606060", "Administrative / clarification"),
    ]
    for i, (lbl, col, desc) in enumerate(labels):
        c = legend.rows[0].cells[i]
        p = c.paragraphs[0]
        p.clear()
        r1 = p.add_run(lbl + "\n")
        r1.font.bold = True; r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(int(col[0:2],16), int(col[2:4],16), int(col[4:6],16))
        r2 = p.add_run(desc)
        r2.font.size = Pt(8); r2.font.color.rgb = GREY

    doc.add_paragraph()

    # Summary table
    add_heading(doc, "ISSUES SUMMARY", level=2)
    headers = ["#", "Issue", "Article(s)", "Priority", "Status"]
    summary = doc.add_table(rows=1, cols=len(headers))
    summary.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = summary.rows[0].cells[i]
        set_cell_bg(c, '1F497D')
        p = c.paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    rows_data = [
        ("1",  "Per-project licence fee — no formula or schedule",  "Art 6.1(a)",   "🔴 CRITICAL", "OPEN"),
        ("2",  "No source code escrow — perpetual licence is hollow","Art 6.3",      "🔴 CRITICAL", "OPEN"),
        ("3",  "Linyang Exhibit B bypasses DISPERON commercial logic","Art 8.6, Ex B","🔴 CRITICAL", "OPEN"),
        ("4",  "Software Margin definition is ambiguous",            "Art 6.4(b)",   "🟡 HIGH",     "OPEN"),
        ("5",  "Development Bonus rate deferred post-signing",       "Art 6.4(b)",   "🟡 HIGH",     "OPEN"),
        ("6",  "No SLA or uptime obligations from Voltus",           "Art 7.2",      "🟡 HIGH",     "OPEN"),
        ("7",  "No IP warranty or indemnity from Voltus",            "Art 5.1",      "🔵 MEDIUM",   "OPEN"),
        ("8",  "Kamil Talar — no performance KPI",                   "Art 4.2",      "🔵 MEDIUM",   "OPEN"),
        ("9",  "Exhibit A (share certificates) not attached",        "Art 2.7",      "🔵 MEDIUM",   "OPEN"),
        ("10", "Exclusivity thresholds identical for large/mid markets","Art 8.3",   "⚪ LOW",      "OPEN"),
        ("11", "Voltus board representation vs economic interest",   "Art 10.2",     "⚪ LOW",      "NOTE"),
    ]
    priority_colours = {
        "🔴 CRITICAL": "FFE0E0",
        "🟡 HIGH":     "FFF3CD",
        "🔵 MEDIUM":   "E8F0FE",
        "⚪ LOW":      "F5F5F5",
    }
    for row_data in rows_data:
        row = summary.add_row()
        for i, val in enumerate(row_data):
            c = row.cells[i]
            if i == 3:
                bg = priority_colours.get(val, "FFFFFF")
                set_cell_bg(c, bg)
            p = c.paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.font.size = Pt(9)
            if i == 0: run.font.bold = True

    doc.add_paragraph()
    doc.add_page_break()

    # ── DETAILED ISSUES ─────────────────────────────────────────────────────────
    add_heading(doc, "DETAILED ISSUE ANALYSIS", level=1)
    doc.add_paragraph()

    add_issue_block(
        doc, "1",
        "Per-Project Licence Fee — No Formula or Schedule",
        "CRITICAL", RED,
        "Article 6.1(a)",
        "a Licence Fee shall be agreed between the Parties on a per-project basis.",
        "There is no formula, floor price, ceiling, or reference rate defined anywhere in the "
        "agreement. Every project requires a fresh bilateral negotiation with Voltus before DISPERON "
        "can commit to customer pricing. Voltus can withhold agreement or inflate fees on any project, "
        "creating a structural choke point. It is impossible to build a customer price book or respond "
        "to tenders without this certainty.",
        "Insert Exhibit C: Fee Schedule. Define a per-MWh licence fee (e.g. EUR X/MWh installed) "
        "indexed to an agreed annual escalator (e.g. HICP). Include a 14-day response obligation for "
        "Voltus to confirm project-specific fees and a deemed-approval mechanism if they fail to respond."
    )

    add_issue_block(
        doc, "2",
        "No Source Code Escrow — Perpetual Licence Is Hollow",
        "CRITICAL", RED,
        "Article 6.3",
        "Voltus Energy Sp. z o.o. shall automatically grant … a perpetual, irrevocable … licence … "
        "This Perpetual Licence shall … survive any change of control, dissolution, or restructuring "
        "of Voltus Energy Sp. z o.o.",
        "A licence to use software that no one can access, maintain, or update is worthless. If Voltus "
        "dissolves, Szumiło and Lechowicz leave, or the company becomes insolvent, DISPERON holds a "
        "contractual right it cannot practically exercise. There is no mechanism for the Company to "
        "obtain the source code, documentation, or build environment under any scenario.",
        "Add a Source Code Escrow clause. A neutral third party (e.g. NCC Group Escrow, Iron Mountain, "
        "or Escrow London) holds a continuously updated copy of the EMS Software source code, build "
        "scripts, and technical documentation. Release triggers: (a) Voltus insolvency or cessation of "
        "operations; (b) failure to provide contracted maintenance for >90 days; (c) material uncured "
        "breach. Costs split 50/50 Voltus/Company."
    )

    add_issue_block(
        doc, "3",
        "Linyang Exhibit B Carve-Out Bypasses DISPERON's Core Commercial Logic",
        "CRITICAL", RED,
        "Article 8.6 and Exhibit B",
        "Voltus Energy Sp. z o.o. retains the right to supply the EMS Software directly to BESS "
        "Manufacturers listed in Exhibit B … Is not subject to any exclusivity threshold or geographic "
        "restriction …",
        "Linyang Energy is listed as Exhibit B Row 1. The stated business model relies on Kamil Talar's "
        "Linyang relationships directing customers to DISPERON. But Voltus can supply EMS directly to any "
        "Linyang project — regardless of geography, exclusivity, or prior DISPERON contact — and DISPERON "
        "receives nothing. This makes Linyang a liability: every Linyang sale Talar develops risks being "
        "taken directly by Voltus. The carve-out also has no revenue-sharing obligation on Voltus for "
        "Linyang-originated revenue.",
        "Add to Article 8.6: 'Notwithstanding the above, where a project originating from a BESS "
        "Manufacturer Partner has been the subject of prior commercial engagement by the Company — "
        "evidenced by written proposal, meeting record, or customer registration — Voltus shall refer "
        "such project to the Company and the parties shall agree a revenue-sharing arrangement within "
        "15 business days. In the absence of agreement, the standard Development Bonus shall apply.' "
        "Separately, add a transparency obligation: Voltus to notify DISPERON of any Linyang project "
        "within 5 business days of engagement."
    )

    add_issue_block(
        doc, "4",
        "Software Margin — Definition Is Ambiguous",
        "HIGH", AMBER,
        "Article 1 (Definitions) and Article 6.4(b)",
        "Software Margin means the revenue received by the Company from the customer in respect of the "
        "Licence component of a customer contract, excluding Commissioning Fees and any other service fees.",
        "When a customer is invoiced a bundled price, there is no agreed mechanism to determine what "
        "portion is 'Licence component' vs. 'Commissioning.' DISPERON could allocate 90% to commissioning "
        "to minimise the bonus base; Voltus could argue the opposite. This will generate disputes on every "
        "post-milestone invoice.",
        "Define Software Margin with a formula or fixed ratio. Insert into the Definitions: 'Software "
        "Margin means X% of the total customer contract value for the relevant project, as set in the Fee "
        "Schedule (Exhibit C), which shall specify both the Commissioning component and the Software "
        "Licence component as separate line items.' Alternatively, use a cost-plus formula referenced to "
        "the agreed per-project licence fee."
    )

    add_issue_block(
        doc, "5",
        "Development Bonus Rate Deferred to Post-Signing Negotiation",
        "HIGH", AMBER,
        "Article 6.4(b)",
        "The Development Bonus percentage shall be agreed between the Parties within 30 (thirty) days of "
        "execution of this Agreement and shall fall within the range of 20% to 30%.",
        "Signing a shareholders agreement with a key economic term unresolved is poor practice. If the "
        "parties cannot agree within 30 days they default to 25%, but having signed the agreement already, "
        "there is no leverage to negotiate. The 30-day window creates a brief period of false urgency "
        "followed by a permanent default.",
        "Lock the Development Bonus at 25% in the signed agreement. Remove the deferred negotiation "
        "process entirely. If either party wants a different rate, negotiate it now and state it clearly. "
        "The range (20-30%) can be retained as a periodic review mechanism (e.g. every 3 years) tied to "
        "demonstrated R&D spend by Voltus."
    )

    add_issue_block(
        doc, "6",
        "No SLA, Uptime, or Maintenance Obligations from Voltus",
        "HIGH", AMBER,
        "Article 7.2",
        "Voltus Energy Sp. z o.o. shall be primarily responsible for the continuous development, "
        "maintenance, security, and improvement of the EMS Software, maintaining it in a commercially "
        "deployable and regulatory-compliant state at all times.",
        "DISPERON signs customer contracts promising software performance, uptime, and regulatory "
        "compliance. If Voltus fails to patch a critical bug, respond to a security incident, or update "
        "for new EU grid code requirements, DISPERON bears full customer liability with no contractual "
        "backstop. There are no defined response times, severity levels, or remedies.",
        "Add Software Support Annex (Exhibit D) defining: (a) severity levels (P1 system down, P2 "
        "major function impaired, P3 minor issue, P4 cosmetic); (b) response and resolution SLAs per "
        "severity; (c) planned maintenance windows with 5-day advance notice; (d) uptime target of "
        "≥99.5% for cloud-hosted components measured monthly; (e) regulatory update obligation — Voltus "
        "to implement required EU/national grid code changes within 90 days of publication; (f) credit "
        "or fee-offset mechanism if SLAs are missed."
    )

    add_issue_block(
        doc, "7",
        "No IP Warranty or Third-Party Infringement Indemnity from Voltus",
        "MEDIUM", BLUE,
        "Article 5.1",
        "The Parties acknowledge and confirm that all Intellectual Property rights in and to the EMS "
        "Software are owned exclusively by Voltus Energy Sp. z o.o.",
        "Voltus asserts ownership but gives no warranty that: (a) the software is free of third-party "
        "IP claims (patent, copyright, trade secret); (b) there is no contaminating open-source licence; "
        "(c) EU export control laws are complied with. If a patent holder or OSS licensor brings a claim, "
        "DISPERON is the commercially visible entity facing customers and potential injunctions.",
        "Add Article 5.5: 'Voltus warrants that (a) it is the sole and unencumbered owner of all IP in "
        "the EMS Software; (b) the EMS Software does not infringe any third-party IP right; (c) no "
        "open-source components are included under licences incompatible with commercial distribution. "
        "Voltus shall indemnify, defend, and hold harmless the Company against any third-party IP claim "
        "relating to the EMS Software, including costs, damages, and legal fees.'"
    )

    add_issue_block(
        doc, "8",
        "Kamil Talar — No Performance KPI or Consequence for Inactivity",
        "MEDIUM", BLUE,
        "Article 4.2",
        "Mr. Kamil Talar shall … actively develop and manage sales channels … Report pipeline and sales "
        "activity to the Director on a monthly basis.",
        "Talar holds 8% with no defined minimum output. The obligation to 'actively develop' and 'report "
        "monthly' has no measurable standard and no consequence clause. A shareholder can remain passive, "
        "hold equity, and face no contractual remedy. Given that Talar's key value is the Linyang "
        "relationship, this gap is commercially significant.",
        "Add to Article 4.2: 'Mr. Talar shall, during each calendar year following execution of this "
        "Agreement, maintain a documented active pipeline of not less than [X] MWh of BESS projects "
        "incorporating the EMS Software and shall introduce a minimum of [Y] qualified customer leads per "
        "year to the Company. Failure to meet these targets in any two consecutive annual periods shall "
        "entitle the Board to reclassify his shares as non-voting ordinary shares and to suspend dividend "
        "rights, subject to 30 days written notice and a cure period.'"
    )

    add_issue_block(
        doc, "9",
        "Exhibit A (Share Certificates) Not Attached at Execution",
        "MEDIUM", BLUE,
        "Article 2.7",
        "Current and updated share certificates shall be appended as Exhibit A. Each Shareholder "
        "acknowledges receipt of a copy of Exhibit A upon signing.",
        "The agreement states certificates are appended, but Exhibit A is blank. Shareholders are "
        "acknowledging receipt of a document that does not exist. This creates an evidentiary gap and "
        "could complicate Registrar filings.",
        "Either attach executed certificates at signing (preferred), or replace Article 2.7 language "
        "with: 'Share certificates reflecting the revised shareholding shall be issued within 30 days "
        "of registration with the Cyprus Registrar of Companies and shall thereupon be appended as "
        "Exhibit A and circulated to all Shareholders.'"
    )

    add_issue_block(
        doc, "10",
        "Exclusivity Thresholds Identical for Large and Mid-Sized Markets",
        "LOW", GREY,
        "Article 8.3 (Table)",
        "Large markets (Germany, Italy, Spain, France): 500 MWh within 24 months. "
        "Mid-sized markets (Poland, Netherlands, Romania, Greece): 500 MWh within 24 months.",
        "Both market tiers require the same 500 MWh threshold. Achieving 500 MWh in Romania or Greece "
        "is structurally harder than in Germany or France due to market size and regulatory maturity. "
        "Either this is a drafting error or a deliberate decision that should be confirmed by all parties. "
        "Also notable: the 500 MWh threshold equals the Perpetual Licence milestone, conflating two "
        "logically separate events.",
        "Confirm intent. If mid-sized markets are intentionally held to the same standard as large "
        "markets, document rationale. Recommended differentiation: Large markets 500 MWh/24 months; "
        "Mid-sized markets 250 MWh/24 months; Smaller markets 100 MWh/24 months (replacing current 200). "
        "Decouple the exclusivity threshold from the Perpetual Licence milestone in definitions."
    )

    add_issue_block(
        doc, "11",
        "Voltus Founders Hold 16% Economic Interest with Minimal Governance Protection",
        "LOW", GREY,
        "Article 10.2",
        "The Shareholders other than Lighthief International … shall have the right to appoint 1 (one) "
        "Director from among themselves to the Board.",
        "Szumiło and Lechowicz together hold 16% but share one board seat with Talar (8%). Lighthief "
        "can appoint up to 3 of 4 directors and controls all decisions requiring a Shareholder majority "
        "(76% > 50%). The Voltus founders' governance rights are effectively vestigial. This is noted "
        "as a structural observation — both parties should enter this agreement with eyes open.",
        "No change required if intentional. For Voltus party comfort, consider adding a reserved matter "
        "requiring 85% shareholder consent for: (a) material change to the commercial model in Art 6; "
        "(b) dissolution of the Company; (c) entry into any transaction that would impair the Perpetual "
        "Licence. This protects the Voltus parties' core interest (software revenue) without giving them "
        "day-to-day blocking power."
    )

    # Footer note
    doc.add_paragraph()
    add_heading(doc, "NEXT STEPS", level=2)
    step_rows = [
        ("Action", "Owner", "Deadline"),
        ("Agree Exhibit C fee schedule (Issue 1 + 4)", "Sybaris + Voltus", "Before signing"),
        ("Appoint source code escrow provider (Issue 2)", "Sybaris", "Before signing"),
        ("Amend Art 8.6 re Linyang referral obligation (Issue 3)", "Legal counsel", "Before signing"),
        ("Lock Development Bonus at 25% (Issue 5)", "All parties", "Before signing"),
        ("Draft Software Support Annex / Exhibit D (Issue 6)", "Voltus", "Within 30 days of signing"),
    ]
    steps = doc.add_table(rows=len(step_rows), cols=3)
    steps.style = 'Table Grid'
    for i, (a, o, d) in enumerate(step_rows):
        row = steps.rows[i]
        for j, val in enumerate([a, o, d]):
            c = row.cells[j]
            if i == 0:
                set_cell_bg(c, '1F497D')
            p = c.paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.font.bold = (i == 0)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF) if i == 0 else BLACK

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "This issues register was prepared for internal review by Lighthief International / Lighthief EUBESS Ltd. "
        "It does not constitute legal advice. Parties should obtain independent legal counsel before execution."
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = GREY

    out_path = "/Volumes/T7 Grey/solinvest/DisperonEMS/docs/contract/DISPERON_SHA_v2_ISSUES.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")

build_issues_doc()
