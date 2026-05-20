#!/usr/bin/env python3
"""
DISPERON EMS/SCADA Order Agreement Generator — Lighthief EU BESS Ltd
====================================================================
Adapted from Voltus Energy order agreements (legal/voltus-contracts/) for
client-facing execution under Lighthief EU BESS Ltd (trading as DISPERON).

Usage:
    python generate_disperon_order_agreement.py

Output:
    DISPERON-Order-Agreement-Part-I-Functional-Scope.docx
    DISPERON-Order-Agreement-Part-II-General-Terms.docx
"""

from __future__ import annotations

import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x40, 0x40, 0x40)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

CONTRACTOR = (
    "Lighthief EU BESS Ltd (trading as DISPERON)\n"
    "Registration No. HE 474192\n"
    "28 Oktovriou & Aemiliou Chourmouziou\n"
    "Lophitis Business Center I, Floor 2, Office 1\n"
    "3035 Limassol, Republic of Cyprus\n"
    "Email: office@lighthief.com | Tel: +357 99 164 158"
)
REPRESENTATIVE = "Alexander Papacosta, Cyprus Managing Director"
DOC_DATE = date.today().strftime("%d %B %Y")
OUT_DIR = os.path.dirname(os.path.abspath(__file__))


def cell_bg(cell, hex_col: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_col)
    tc_pr.append(shd)


def setup(doc: Document) -> None:
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(10)
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(3.0)


def title(doc: Document, text: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(11)
        r2.font.color.rgb = GOLD
        r2.italic = True


def h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = GOLD
        r.font.bold = True
        r.font.size = Pt(13)


def h2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.bold = True
        r.font.size = Pt(11)


def para(doc: Document, text: str, bold: bool = False, indent: float = 0) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.bold = bold
    r.font.color.rgb = BLACK


def bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.4)
    r = p.add_run(text)
    r.font.size = Pt(10)


def checkbox_line(doc: Document, text: str) -> None:
    para(doc, f"☐  {text}")


def meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        c0, c1 = t.rows[i].cells[0], t.rows[i].cells[1]
        if i == 0:
            cell_bg(c0, "1A365D")
            cell_bg(c1, "1A365D")
        p0 = c0.paragraphs[0]
        p0.clear()
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(9)
        r0.font.color.rgb = WHITE if i == 0 else BLACK
        p1 = c1.paragraphs[0]
        p1.clear()
        r1 = p1.add_run(value)
        r1.font.size = Pt(9)
        r1.font.color.rgb = WHITE if i == 0 else BLACK
    doc.add_paragraph()


def sig_block(doc: Document, party: str, role: str) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"{party}\n")
    r.bold = True
    p.add_run(f"{role}\n\n")
    p.add_run("Signature: _________________________________\n\n")
    p.add_run("Full name (print): _______________________\n\n")
    p.add_run("Date: ____________________________________")


def build_part_i() -> Document:
    doc = Document()
    setup(doc)

    title(
        doc,
        "DISPERON EMS/SCADA SYSTEM\nORDER AGREEMENT",
        "Part I — Functional, Technical and Commercial Scope",
    )
    doc.add_paragraph()

    meta_table(
        doc,
        [
            ("Field", "Value"),
            ("Document Reference", "LEB-EMS-ORDER-I"),
            ("Version", "1.0"),
            ("Date", DOC_DATE),
            ("Contractor", "Lighthief EU BESS Ltd (DISPERON)"),
        ],
    )

    h1(doc, "Parties")
    para(doc, "Date of Agreement: .....................................................")
    para(doc, "Customer: ....................................................................")
    para(doc, "Contractor:")
    for line in CONTRACTOR.split("\n"):
        para(doc, line, indent=0.5)
    para(doc, f"Represented by: {REPRESENTATIVE}")
    para(doc, 'Hereinafter referred to as the "Contractor".')

    h1(doc, "1. Subject of the Agreement")
    para(
        doc,
        "The subject of this Agreement is the delivery, configuration and commissioning of the "
        "DISPERON EMS/SCADA system in accordance with the functional scope defined in this Agreement "
        "and the selected integration options. DISPERON is the commercial brand under which "
        "Lighthief EU BESS Ltd supplies energy management and SCADA services for battery energy "
        "storage and hybrid renewable installations.",
    )

    h2(doc, "Scope of devices included in the system")
    for item in [
        "BESS energy storage systems",
        "Battery inverters / PCS",
        "PV inverters",
        "Photovoltaic installation (SPC — Solar Power Controller)",
        "Wind turbines (WPC — Wind Power Controller)",
        "EV chargers (EVMS — Electric Vehicle Management System)",
        "Generator set / cogeneration",
        "Energy meters (number of meters): ………..",
        "MV/LV transformer station",
    ]:
        checkbox_line(doc, item)

    h2(doc, "System operator modules (Cyprus)")
    for item in [
        "EAC DSO module — Electricity Authority of Cyprus distribution network",
        "TSOC module — Transmission System Operator (installations >8 MW)",
        "System services module (FCR / aFRR / mFRR — where applicable)",
        "Energy trader / market participant module (TSOC MMS interfaces)",
    ]:
        checkbox_line(doc, item)

    h2(doc, "Data sources and integrations")
    for item in [
        "Weather forecasts",
        "PV irradiance forecasts",
        "Electricity market prices (TSOC Day-Ahead Market / balancing)",
        "Energy meter data",
        "Operator systems data (IEC 60870-5-104)",
        "Integration with external SCADA systems",
        "Integration with trading / portfolio systems",
    ]:
        checkbox_line(doc, item)

    h2(doc, "DISPERON EMS/SCADA functionalities")
    h2(doc, "Monitoring & Diagnostics")
    for item in [
        "Installation parameter monitoring",
        "Monitoring of individual devices (PV, BESS, EV, meters)",
        "Detailed diagnostics and alarms",
    ]:
        checkbox_line(doc, item)

    h2(doc, "Data Collection & Analysis")
    for item in [
        "Production and consumption history",
        "History of alarms, errors, and events",
        "Energy cost analysis",
        "Efficiency analysis (kWh/kWp, performance ratio)",
    ]:
        checkbox_line(doc, item)

    h2(doc, "Device Control")
    for item in [
        "Remote PV control",
        "Remote BESS control",
        "Remote EV charger control",
        "Device parameter configuration",
        "Power balancing between inverters",
    ]:
        checkbox_line(doc, item)

    h2(doc, "Active Power — Operating Strategies")
    checkbox_line(doc, "Zero import / Zero export")
    checkbox_line(doc, "Set import / Set export")
    checkbox_line(doc, "Peak Shaving")
    checkbox_line(
        doc,
        "External Setpoint via ☐ External SCADA / ☐ Trader / "
        "☐ Dynamic control according to energy price / ☐ Dynamic control according to weather",
    )

    h2(doc, "Reactive Power and Power Factor")
    for item in [
        "Reactive power compensation (Q)",
        "Q/U regulation (Cyprus EN 50549-2)",
        "cos φ / tan φ control",
    ]:
        checkbox_line(doc, item)

    h2(doc, "Schedules")
    for item in [
        "Weekly schedules",
        "Advanced schedules based on prices and weather",
        "EV charging schedules",
    ]:
        checkbox_line(doc, item)

    h2(doc, "Safety and Alarms")
    for item in [
        "E-mail notifications",
        "SMS notifications",
        "Emergency mode (Backup Mode / Off-Grid)",
        "Overload and overvoltage protections",
    ]:
        checkbox_line(doc, item)

    h2(doc, "Reports")
    for item in [
        "PV production report",
        "BESS storage report",
        "Energy balance report",
        "Cost and revenue report",
        "Investor dashboard (KPI, ROI)",
    ]:
        checkbox_line(doc, item)

    h2(doc, "Grid compliance (Cyprus)")
    for item in [
        "IEC 60870-5-104 SCADA gateway for EAC/TSOC",
        "Active and reactive power setpoint control",
        "Discrete curtailment levels (100% / 60% / 30% / 0%)",
        "Loss-of-mains and anti-islanding coordination",
        "NIS2-aligned access control and audit logging",
    ]:
        checkbox_line(doc, item)

    h2(doc, "Installation parameters")
    para(doc, "PV — installation capacity: ............ kWp")
    para(doc, "PV — inverter capacity: ............ kW")
    para(doc, "BESS — power: ............ kW")
    para(doc, "BESS — capacity: ............ kWh")
    para(doc, "Type of BESS inverters: ........................................")
    para(doc, "Type of batteries / BMS: .............................................")
    para(doc, "Number of energy meters: ............")
    para(doc, "MV/LV station — type: ..........................................")

    h1(doc, "2. Implementation Stages")
    para(doc, "Stage 1 — Installation", bold=True)
    bullet(doc, "Preparation and configuration of controllers at the Contractor's facility")
    bullet(doc, "Shipment of devices to the site")
    bullet(doc, "System installation and device integration")
    bullet(doc, "Communication configuration (Modbus TCP, IEC 104, VPN as applicable)")

    para(doc, "Stage 2 — Testing", bold=True)
    bullet(doc, "Communication tests and control tests")
    bullet(doc, "On-site system configuration")
    bullet(doc, "Alarm and notification tests")
    bullet(doc, "Integration with EAC/TSOC SCADA (where in scope)")

    para(doc, "Stage 3 — Commissioning", bold=True)
    bullet(doc, "System startup and acceptance testing")
    bullet(doc, "Personnel training")
    bullet(doc, "Startup monitoring period")

    h1(doc, "3. Scope of Delivery and Execution")
    para(
        doc,
        "3.1. The Contractor shall deliver DISPERON EMS/SPC/WPC/EVMS system controllers required "
        "for installation management in line with the functional scope.",
    )
    para(doc, "Controller configuration and commissioning:", bold=True)
    for item in [
        "HPC — Hybrid Power Controller",
        "Limit Operator Module (PPC — Power Plant Controller)",
        "EMS — Energy Management System",
        "SPC PV / EVMS / WPC",
        "SCADA — Cloud and/or On-Premises",
    ]:
        checkbox_line(doc, item)

    para(
        doc,
        "The Contractor shall perform communication configuration, device integration, software setup, "
        "licensing and system configuration in accordance with this Agreement. As-built documentation "
        "shall be delivered upon completion. Programming, configuration and commissioning works are "
        "performed remotely unless otherwise agreed; on-site works are executed as defined in Section 5.",
    )

    h1(doc, "4. Customer's Technical Infrastructure Readiness")
    para(doc, "The Customer shall ensure:")
    for item in [
        "Access to LAN/WAN infrastructure and required VLANs",
        "Static IP addressing for all devices included in the system",
        "Remote access (VPN), if applicable",
        "Completed electrical and I&C works (wiring, connections, installation)",
        "Energization of the facility and proper device startup",
        "Full operational readiness and correct response to control commands",
        "Appointment of a person responsible for coordination",
        "Local and/or remote access for the Contractor",
        "Agreement with EAC/TSOC on signals, functionalities and control logic (where applicable)",
    ]:
        bullet(doc, item)

    h1(doc, "5. On-site Works")
    para(doc, "On-site works include technical visits, commissioning and training:")
    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    cell_bg(hdr[0], "1A365D")
    cell_bg(hdr[1], "1A365D")
    for ci, txt in enumerate(["Description", "Number of Days"]):
        p = hdr[ci].paragraphs[0]
        p.clear()
        r = p.add_run(txt)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9)
    for ri, desc in enumerate(
        ["First Visit", "Integration & Commissioning", "Training"], start=1
    ):
        t.rows[ri].cells[0].text = desc
        t.rows[ri].cells[1].text = ".."
    doc.add_paragraph()
    for item in [
        "Visits are scheduled at least 14 days in advance and require confirmed infrastructure readiness",
        "Changes to days require mutual agreement and may affect pricing",
        "Travel and accommodation costs are billed separately unless included in the fixed price",
    ]:
        bullet(doc, item)

    h1(doc, "6. Exclusions from the Scope of Delivery")
    para(doc, "The Agreement does not include:")
    for item in [
        "Additional installation works beyond EMS/SCADA scope",
        "Additional configuration or commissioning of the Customer's devices not listed herein",
        "Integrations with third-party systems not listed in this Agreement",
        "On-site works not included in Section 5",
        "Configuration of communication devices, inverters, or protection relays (unless agreed)",
        "Works related solely to the DSO/TSOC operator's own SCADA infrastructure",
        "Physical BESS, PCS, or MV equipment supply (covered under separate EPC agreement)",
    ]:
        bullet(doc, item)

    para(
        doc,
        "Additional notes: Final cost may change depending on actual device configuration. "
        "The Agreement becomes effective upon signature and advance payment. The Customer provides "
        "all technical data and access. Changes require written form. Additional requirements are "
        "billed separately. Failure to provide conditions enabling work does not release the Customer "
        "from payment obligations for work already performed.",
    )

    h1(doc, "7. Additional Work Rates")
    bullet(doc, "Hourly rate (per person): EUR 80 net")
    bullet(doc, "Daily rate (up to 8h per person): EUR 1,000 net")
    bullet(doc, "Travel and accommodation costs billed separately unless agreed in writing")

    h1(doc, "8. Compensation")
    para(doc, "Total price for the execution of the Agreement: ............................ EUR net")
    para(doc, "Price includes:", bold=True)
    for item in [
        "DISPERON EMS/SCADA system delivery",
        "Configuration, integration, tests, documentation",
        "Perpetual licence for local EMS installation at the site",
        "SCADA Local (on-site gateway) — standard rate EUR 15,000 per park (where selected)",
        "SCADA Global (portfolio/group) — standard rate EUR 60,000 per client group, first park only (where selected)",
    ]:
        bullet(doc, item)

    para(doc, "Payment terms:", bold=True)
    bullet(doc, "50% advance — within 7 days of signing")
    bullet(doc, "20% before delivery / remote configuration start")
    bullet(doc, "30% after acceptance (signed acceptance protocol)")

    para(doc, "Annual software subscription (separate from upfront price):", bold=True)
    bullet(doc, "EUR 400 per MWh of installed BESS capacity per year (ex VAT)")
    bullet(doc, "SCADA Local maintenance: EUR 3,000 per park per year (where applicable)")
    bullet(doc, "SCADA Global maintenance: EUR 12,000 per client group per year (where applicable)")
    para(
        doc,
        "Subscription period included with upfront package (if any): "
        "☐ 1 year   ☐ 3 years   ☐ 5 years   ☐ 10 years   ☐ Billed separately from PAC",
    )
    para(
        doc,
        "Detailed subscription and support terms are defined in Part II — General Terms of "
        "Implementation, Service and Warranty, and in the EMS Software Subscription Agreement "
        "where executed separately.",
    )

    h1(doc, "9. Warranty and Support")
    para(
        doc,
        "The Contractor provides warranty for delivered software and hardware (if applicable) for a "
        "period of …… years. Warranty and support conditions are defined in Part II.",
    )

    h1(doc, "10. Liability of the Parties")
    para(
        doc,
        "The Contractor is responsible for the correct execution of the DISPERON system in accordance "
        "with this Agreement. The Customer is responsible for data accuracy and correct operation of "
        "the technical infrastructure.",
    )

    h1(doc, "11. Confidentiality")
    para(
        doc,
        "Both Parties agree to maintain confidentiality of all technical and commercial information "
        "related to the execution of this Agreement.",
    )

    h1(doc, "12. Final Provisions")
    para(doc, "Changes to this Agreement require written form.")
    para(
        doc,
        "For matters not regulated herein, the laws of the Republic of Cyprus apply. "
        "Disputes shall be subject to the exclusive jurisdiction of the courts of Limassol, Cyprus.",
    )
    para(doc, "This Agreement is executed in two identical copies, one for each Party.")

    h1(doc, "13. Components of the Agreement")
    para(doc, "This Agreement consists of:")
    bullet(doc, "Part I — Functional, Technical and Commercial Scope (this document)")
    bullet(doc, "Part II — General Terms of Implementation, Service and Warranty")
    bullet(doc, "Schedule A — Site-specific pricing and scope (if attached)")
    para(
        doc,
        "Both parts form an integral whole. In case of discrepancies between Part I and Part II, "
        "the provisions of Part II shall apply unless Schedule A expressly states otherwise.",
    )

    h1(doc, "Signatures")
    sig_block(doc, "Customer", "Authorised signatory")
    sig_block(doc, "Contractor — Lighthief EU BESS Ltd (DISPERON)", REPRESENTATIVE)

    return doc


def build_part_ii() -> Document:
    doc = Document()
    setup(doc)

    title(
        doc,
        "DISPERON EMS/SCADA\nGENERAL TERMS",
        "Part II — Implementation, Service and Warranty",
    )
    doc.add_paragraph()

    meta_table(
        doc,
        [
            ("Field", "Value"),
            ("Document Reference", "LEB-EMS-ORDER-II"),
            ("Version", "1.0"),
            ("Date", DOC_DATE),
            ("Applies with", "LEB-EMS-ORDER-I (Part I)"),
        ],
    )

    sections = [
        (
            "1. Implementation Deadlines",
            [
                "1.1 Implementation begins after the Agreement is signed and the advance payment is received.",
                "1.2 Standard project start is 6–8 weeks from signing, unless otherwise agreed.",
                "1.3 Final completion date is confirmed after verifying full technical and electrical readiness on the Customer's side.",
                "1.4 Full readiness includes in particular:",
            ],
        ),
    ]

    readiness = [
        "Delivery of complete technical documentation (PV, BESS, EV, switchgear, metering).",
        "Completion of all electrical works, wiring, connections, and initial device configuration.",
        "Completion and acceptance of the grid connection (EAC / TSOC as applicable).",
        "Energization of the facility and commissioning of devices per technical documentation.",
        "Full operability of devices and correct response to DISPERON control commands.",
        "Agreement with EAC/TSOC on signals, functionalities and control logic.",
        "Operational telecommunications infrastructure (LAN/WAN, VLANs, VPN, static IP addressing).",
        "Configuration of IP addressing for all EMS/SCADA-connected devices.",
        "Stable power supply and correct network parameters.",
        "Appointment of a coordinator on the Customer's side.",
        "Provision of local or remote access for the Contractor.",
        "Completion of technical acceptance procedures required by applicable law.",
    ]

    for title_text, lines in sections:
        h1(doc, title_text)
        for line in lines:
            para(doc, line)
        for item in readiness:
            bullet(doc, item, level=1)
        para(
            doc,
            "1.5 Any delay resulting from failure to meet the above conditions extends the implementation "
            "schedule without liability for the Contractor.",
        )
        para(
            doc,
            "1.6 The Contractor undertakes to deliver and implement the DISPERON EMS/SCADA system within "
            "up to 28 days from the moment the Customer meets all readiness conditions.",
        )
        para(
            doc,
            "1.7 Commissioning and acceptance tests shall be carried out within 14 days from confirming "
            "installation readiness.",
        )

    h1(doc, "2. Scope and Method of Implementation")
    for line in [
        "2.1 The Contractor implements the DISPERON EMS/SCADA system per Part I.",
        "2.2 Programming, configuration and commissioning are performed remotely unless otherwise agreed.",
        "2.3 On-site works follow the schedule and scope in Part I.",
        "2.4 Scope is based on Customer documentation and may be clarified after site verification.",
        "2.5 The Customer shall provide infrastructure, documentation, remote access, and a coordinator.",
        "2.6 Discrepancies between documentation and the actual installation may require scope, deadline, and cost adjustment.",
    ]:
        para(doc, line)

    h1(doc, "3. On-site Works")
    for line in [
        "3.1 On-site works include: first technical visit, integration and commissioning, personnel training.",
        "3.2 Number of on-site days is defined in Part I or Schedule A.",
        "3.3 Visits require 14 days' notice and confirmed readiness.",
        "3.4 Changes to on-site days require mutual agreement and may affect cost.",
        "3.5 Travel and delegation expenses are billed separately unless included in the fixed price.",
    ]:
        para(doc, line)

    h1(doc, "4. Exclusions from the Agreement")
    for line in [
        "4.1 The Agreement does not include additional installation works, re-commissioning of Customer devices, "
        "integrations not listed in Part I, on-site works beyond Part I, configuration of third-party protection "
        "relays unless agreed, or works on the DSO/TSOC operator's own SCADA system.",
        "4.2 Cost may change depending on actual device configuration.",
        "4.3 The Customer must provide all data and access required to begin works.",
        "4.4 Additional requirements are priced individually.",
        "4.5 Customer-caused inability to perform does not release payment for work already done.",
        "4.6 The system is installed per communication standards and manufacturer requirements.",
    ]:
        para(doc, line)

    h1(doc, "5. Subscription and Technical Support")
    h2(doc, "5.1 Subscription scope")
    para(doc, "The annual subscription (where not prepaid in Part I) includes:")
    for item in [
        "System updates and functional development",
        "Cybersecurity maintenance (access policies, NIS2-aligned controls)",
        "Database administration and backup",
        "Maintenance of integration protocols (Modbus, MQTT, TCP/IP, IEC 60870-5-104)",
        "Access to advanced EMS modes (arbitrage, peak shaving, import/export limits)",
        "External data (market prices, forecasts, operator signals where licensed)",
        "Service and operational support per Section 5.3",
        "Multi-site management (SCADA Global tier)",
    ]:
        bullet(doc, item)

    h2(doc, "5.2 Subscription period and billing")
    para(doc, "5.2.1 Available periods: 1–15 years (aligned with LTSA where applicable).")
    para(doc, "5.2.2 Billing: monthly, annual upfront, or one-time for the selected period.")
    para(doc, "5.2.3 Standard annual rate: EUR 400 per MWh of installed BESS capacity (ex VAT).")
    para(doc, "5.2.4 After subscription expiry:")
    for item in [
        "Service, updates and cloud integration functions expire",
        "The system may operate as local EMS only, without cloud SCADA, historical cloud data, or external data feeds",
    ]:
        bullet(doc, item, level=1)

    h2(doc, "5.3 Technical support")
    for line in [
        "5.3.1 Basic subscription includes up to 1 hour per month remote support per park.",
        "5.3.2 Support includes remote connections, log analysis, diagnostics, and operational consultations.",
        "5.3.3 Unused support time does not carry over.",
        "5.3.4 Excess work is billed at EUR 80/hour net per Part I.",
        "5.3.5 Support excludes Customer infrastructure faults, third-party device failures, and unauthorised interference.",
    ]:
        para(doc, line)

    h1(doc, "6. Acceptance of Works")
    para(doc, "6.1 Acceptance is confirmed by a protocol signed by both Parties.")
    para(doc, "6.2 Comments may be submitted within 7 days of receiving test results.")

    h1(doc, "7. Warranty")
    para(doc, "7.1 Warranty covers:")
    for item in [
        "Removal of critical software errors attributable to the Contractor",
        "Corrections resulting from the Contractor's fault",
        "Software stability for the warranted period",
        "Repair or replacement of defective hardware supplied by the Contractor",
    ]:
        bullet(doc, item)
    para(doc, "7.2 Target response time: up to 36 working hours for critical issues.")
    para(doc, "7.3 Warranty may be extended for 1 / 3 / 5 / 10 years where agreed in Part I.")
    para(doc, "7.4 Extended warranty requires an active subscription and current software version.")
    para(doc, "7.5 Warranty excludes third-party interference, improper operation, external device failures, and Customer-side communication faults.")

    h1(doc, "8. Liability of the Parties")
    para(doc, "8.1 The Contractor is responsible for implementing the system per the Agreement.")
    para(doc, "8.2 The Customer is responsible for data accuracy and technical infrastructure operation.")
    para(
        doc,
        "8.3 Except for fraud or wilful misconduct, each Party's aggregate liability under this Agreement "
        "is limited to the amounts paid or payable under Part I for the affected site, unless mandatory "
        "law provides otherwise.",
    )

    h1(doc, "9. Confidentiality")
    para(doc, "The Parties undertake to maintain confidentiality of technical and commercial information.")

    h1(doc, "10. Final Provisions")
    para(doc, "10.1 Amendments require written form.")
    para(doc, "10.2 Governing law: Republic of Cyprus. Courts: Limassol, Cyprus.")
    para(doc, "10.3 Executed in two identical copies.")

    h1(doc, "11. Relationship of Documents")
    para(
        doc,
        "Part II is an integral part of the DISPERON Order Agreement and applies together with Part I. "
        "Where an EMS Software Subscription Agreement is also executed, subscription billing terms in "
        "that agreement prevail for annual recurring fees if expressly stated therein.",
    )

    h1(doc, "Signatures")
    sig_block(doc, "Customer", "Authorised signatory")
    sig_block(doc, "Contractor — Lighthief EU BESS Ltd (DISPERON)", REPRESENTATIVE)

    return doc


def main() -> None:
    os.makedirs(OUT_DIR, exist_ok=True)
    part_i_path = os.path.join(
        OUT_DIR, "DISPERON-Order-Agreement-Part-I-Functional-Scope.docx"
    )
    part_ii_path = os.path.join(
        OUT_DIR, "DISPERON-Order-Agreement-Part-II-General-Terms.docx"
    )
    build_part_i().save(part_i_path)
    build_part_ii().save(part_ii_path)
    print(f"Generated:\n  {part_i_path}\n  {part_ii_path}")


if __name__ == "__main__":
    main()
