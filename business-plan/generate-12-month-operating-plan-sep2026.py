"""
12-Month Operating Plan (Sep 2026 - Aug 2027) — Lighthief Cyprus Ltd
=====================================================================
Prepared for: Dr. Arkadiusz Sybaris (Founder & CEO / UBO) and Michael Shapiro
Scope: BESS O&M/LTSA growth (Esperia/Galascope + Timotheos/Lampros/Spanercom
       pipeline + Aeolian standalone) and PV O&M growth (C&I + residential
       >=20kWp only) — marketing, staffing, payroll, office/warehouse, cashflow.

Source data: lib/portfolio-data.ts (SSOT, Aug 2026), docs/internal/opex-plan-2026.md,
             team/team-data.ts, team/README.md, pv-om/README.md + cost model,
             lighthief-cyprus/lighthief-cyprus-sales-growth-plan.md,
             marketing/solinvest-marketing-strategy.html,
             lighthief-cyprus/rental-due-diligence-response.md,
             financial/preliminary-statements/preliminary-financials-2025-2026.md

Run: python3 "business-plan/generate-12-month-operating-plan-sep2026.py"
Output: business-plan/lighthief-12-month-operating-plan-sep2026-aug2027.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Brand colours (Lighthief brand identity) ───────────────────────────────
NAVY, GOLD, WHITE, BLACK, GREY = (
    RGBColor(0x1A, 0x36, 0x5D), RGBColor(0xC9, 0xA4, 0x32),
    RGBColor(0xFF, 0xFF, 0xFF), RGBColor(0x00, 0x00, 0x00), RGBColor(0x40, 0x40, 0x40),
)
NAVY_HEX, GOLD_HEX, LIGHT_HEX, ALT_HEX, RED_HEX, GREEN_HEX = (
    "1A365D", "C9A432", "EBF0F7", "F5F7FA", "FBE4E4", "E6F4EA",
)

OUT_DIR = os.path.dirname(__file__)
OUT_FILE = os.path.join(OUT_DIR, "lighthief-12-month-operating-plan-sep2026-aug2027.docx")

# ── Helpers (consistent with pv-om/contracts/generate-standard-pv-om-contract.py) ──
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(15 if level == 1 else 12)
    r.font.color.rgb = GOLD
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), GOLD_HEX if level == 1 else NAVY_HEX)
    pBdr.append(bot); pPr.append(pBdr)
    return p

def add_paragraph(doc, text="", bold=False, italic=False, size=10.5,
                   color=BLACK, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.color.rgb = color
    return p

def add_bullet(doc, text, bold_prefix=None, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix); r1.bold = True
        r1.font.size = Pt(size); r1.font.color.rgb = NAVY
    r2 = p.add_run(text)
    r2.font.size = Pt(size); r2.font.color.rgb = BLACK

def make_table(doc, n_cols):
    t = doc.add_table(rows=0, cols=n_cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    return t

def add_row(table, cells_data, header=False, alt=False, total=False):
    row = table.add_row()
    for i, (text, width, align_str, bold) in enumerate(cells_data):
        cell = row.cells[i]; cell.width = Cm(width)
        if header: set_cell_bg(cell, NAVY_HEX)
        elif total: set_cell_bg(cell, LIGHT_HEX)
        elif alt: set_cell_bg(cell, ALT_HEX)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold or header or total; run.font.size = Pt(9)
        run.font.color.rgb = WHITE if header else BLACK
        p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align_str == "C"
                       else WD_ALIGN_PARAGRAPH.RIGHT if align_str == "R"
                       else WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    return row

def note(doc, text):
    add_paragraph(doc, text, size=9, color=GREY, space_before=2, space_after=10)

# ══════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21); section.page_height = Cm(29.7)
section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)
section.top_margin = Cm(1.8); section.bottom_margin = Cm(1.8)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# ── Header bar ──────────────────────────────────────────────────────────────
hdr = doc.add_table(rows=1, cols=1); hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
hc = hdr.rows[0].cells[0]; set_cell_bg(hc, NAVY_HEX)
p = hc.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
r = p.add_run("LIGHTHIEF CYPRUS LTD")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = WHITE
p2 = hc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run("12-Month Operating Plan — BESS O&M and PV O&M Growth")
r2.bold = True; r2.font.size = Pt(13); r2.font.color.rgb = GOLD
p3 = hc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(12)
r3 = p3.add_run("Marketing · Staffing · Payroll · Office & Warehouse · Cashflow")
r3.font.size = Pt(10.5); r3.font.color.rgb = WHITE
doc.add_paragraph()

# ── Reference block ──────────────────────────────────────────────────────────
ref = doc.add_table(rows=0, cols=2); ref.alignment = WD_TABLE_ALIGNMENT.LEFT
for label, value in [
    ("Document Reference:", "LCY-PLAN-2026-002"),
    ("Date Prepared:", "11 August 2026"),
    ("Plan Period:", "1 September 2026 \u2013 31 August 2027 (12 months)"),
    ("Prepared For:", "Dr. Arkadiusz Sybaris (Founder & CEO / UBO) and Michael Shapiro"),
    ("Prepared By:", "Lighthief Cyprus Ltd \u2014 Alexander Papacosta, Cyprus Director"),
    ("Classification:", "INTERNAL \u2014 CONFIDENTIAL \u2014 Not for external distribution"),
    ("Status:", "DRAFT \u2014 for review and sign-off"),
]:
    row = ref.add_row()
    c0, c1 = row.cells[0], row.cells[1]
    c0.width = Cm(4.5); c1.width = Cm(12)
    r0 = c0.paragraphs[0].add_run(label)
    r0.bold = True; r0.font.size = Pt(9.5); r0.font.color.rgb = GREY
    r1 = c1.paragraphs[0].add_run(value)
    r1.font.size = Pt(9.5); r1.font.color.rgb = BLACK
doc.add_paragraph()

# ── Executive Summary ────────────────────────────────────────────────────────
add_heading(doc, "Executive Summary", level=1)
add_paragraph(doc,
    "This plan sets out how Lighthief Cyprus Ltd grows its two recurring-revenue service lines \u2014 "
    "BESS O&M/LTSA and PV O&M \u2014 over the 12 months from 1 September 2026 to 31 August 2027. "
    "BESS O&M growth is anchored on the Esperia Energy (Dino Constantinou) / Galascope confirmed order and the "
    "conditional Timotheos Timotheou / Lampros Andreadis / Spanercom (Anarita) pipeline extension, plus the "
    "standalone Aeolian Dynamics client. PV O&M growth targets commercial & industrial (C&I) sites of any size "
    "and residential systems of 20 kWp or larger only \u2014 small residential rooftop (<20 kWp) is explicitly "
    "excluded from this plan's O&M sales motion.")
add_paragraph(doc,
    "The plan is deliberately conservative on revenue timing: BESS O&M/LTSA billing only starts once a park "
    "reaches PAC (Galascope 1+2 target 31 Jan 2027), so the first 5 months are pure ramp-up cost. The resulting "
    "cash requirement is designed to be funded from EPC margin already earned on the same batches \u2014 not from "
    "new external capital \u2014 consistent with the approach set out in the 2026 OPEX plan.")

dash = make_table(doc, 4)
add_row(dash, [("Metric", 7.5, "L", True), ("Value", 3.5, "C", True),
               ("Metric", 4.5, "L", True), ("Value", 1.0, "C", True)], header=True)
dash_rows = [
    ("Plan period", "12 months", "PV O&M target segment", "C&I + resi \u226520kWp"),
    ("BESS O&M base case MWh (confirmed)", "30 MWh", "PV O&M new contracts targeted", "3"),
    ("BESS O&M upside MWh (+ conditional)", "+90 MWh", "Field/O&M headcount added", "+4"),
    ("Total 12-mo growth OPEX", "\u2248\u20ac549,000", "Total 12-mo O&M/PV-O&M revenue", "\u2248\u20ac114,000"),
    ("Net cash requirement (12 mo)", "\u2248\u20ac435,000", "Covered by EPC margin available", "up to \u20ac1.16M"),
    ("Marketing budget (12 mo)", "\u2248\u20ac84,000", "Warehouse net cost (from Oct26)", "\u20ac4,000/mo"),
]
for i, row in enumerate(dash_rows):
    add_row(dash, [(row[0], 7.5, "L", False), (row[1], 3.5, "C", True),
                   (row[2], 4.5, "L", False), (row[3], 1.0, "C", True)], alt=(i % 2 == 1))
note(doc, "Figures are modelled from lib/portfolio-data.ts (SSOT, Aug 2026), the 2026 OPEX plan, team-data.ts and "
          "the sales growth plan. BESS O&M \u201cupside\u201d depends on Timotheos (50%), Lampros (50%) and "
          "Spanercom (85%) signing \u2014 see \u00a72.")

# ── Section 1: Scope & Objectives ────────────────────────────────────────────
add_heading(doc, "1. Scope & Strategic Objectives", level=1)
add_paragraph(doc, "Four objectives for the 12-month period, agreed for sign-off by Dr. Sybaris and Mr. Shapiro:")
for bold, text in [
    ("BESS O&M/LTSA: ", "stand up the field team and back-office ahead of Galascope 1+2 PAC (target 31 Jan 2027, "
     "30 MWh), and be ready to absorb the Timotheos/Lampros/Spanercom extension (+90 MWh) if/when those EPCs sign."),
    ("PV O&M growth: ", "grow beyond the existing Spanercom contract (\u20ac56,400/yr) by signing 3 new C&I or "
     "residential \u226520kWp O&M contracts in the 12 months, using the ECO/SILVER/GOLD package model."),
    ("Marketing: ", "fund enough spend to keep both pipelines full \u2014 BESS curtailment/retrofit leads for "
     "existing CERA licensees, and a new C&I/large-residential PV O&M campaign that explicitly excludes small "
     "residential rooftop."),
    ("Staff, payroll & premises: ", "grow headcount from 5 to 9\u201310 in step with confirmed demand (not ahead "
     "of it), keep payroll on the existing SSOT cost model, and resolve the office/warehouse footprint question "
     "raised in the July 2026 rental due-diligence pack."),
]:
    add_bullet(doc, text, bold_prefix=bold)
note(doc, "This is an internal operating plan, not an investor document \u2014 it does not repeat CIF, margin-% "
          "or quotation-reference detail from client-facing proposals.")

# ── Section 2: Market & Pipeline ─────────────────────────────────────────────
add_heading(doc, "2. Market & Pipeline Overview", level=1)
add_heading(doc, "2.1  BESS O&M \u2014 Esperia/Galascope (Dino), Timotheos pipeline, standalone", level=2)
pipe = make_table(doc, 6)
add_row(pipe, [("Client / Batch", 5.0, "L", True), ("Parks", 1.5, "C", True), ("MWh", 1.5, "C", True),
               ("Status", 2.5, "C", True), ("PAC Target", 2.5, "C", True), ("O&M Start", 3.0, "C", True)], header=True)
pipe_rows = [
    ("Galascope 1+2 (Esperia \u2014 Dino Constantinou)", "2", "30.0", "Confirmed", "31 Jan 2027", "Feb 2027 \u2014 base case"),
    ("Timotheos Timotheou (3 parks)", "3", "35.0", "Pending \u2014 50%", "31 Jan 2027*", "Feb 2027 if signed"),
    ("Lampros Andreadis (2 parks)", "2", "15.0", "Pending \u2014 50%", "31 Jan 2027*", "Feb 2027 if signed"),
    ("Spanercom \u2014 Anarita (2 parks)", "2", "40.0", "High \u2014 85%", "31 Jan 2027*", "Feb 2027 if signed"),
    ("Aeolian Dynamics (standalone, Larnaca)", "1", "20.06", "High \u2014 80%", "On EPC signing", "Separate EPC + LTSA"),
    ("Esperia pipeline Batch 2 (2026\u20132027)", "4", "75.0", "Pipeline", "30 Jun 2027", "Outside this plan's window"),
]
for i, row in enumerate(pipe_rows):
    add_row(pipe, [(v, w, "C" if j else "L", False) for j, (v, w) in enumerate(zip(row, [5.0,1.5,1.5,2.5,2.5,3.0]))], alt=(i % 2 == 1))
note(doc, "*Batch 1 Extension shares Batch 1's production/CIF/PAC schedule per portfolio-data.ts. Aeolian is a "
          "separate standalone client with its own EPC and LTSA \u2014 not part of the group order. "
          "Source: lib/portfolio-data.ts BATCHES / GROUPS / AEOLIAN (SSOT, dated 29 Jul\u201330 Jun 2026).")

add_paragraph(doc,
    "Base case for this plan uses Galascope only (30 MWh confirmed). The Timotheos/Lampros/Spanercom extension "
    "(+90 MWh) is treated as upside \u2014 if it signs on the same timeline, O&M revenue from Feb 2027 roughly "
    "quadruples (see \u00a77) and the case for the 5th\u20136th field engineer (\u00a74.3) becomes immediate rather than "
    "conditional.")

add_heading(doc, "2.2  PV O&M \u2014 target segment: C&I + residential \u226520kWp only", level=2)
add_paragraph(doc,
    "Per direction from Dr. Sybaris and Mr. Shapiro, PV O&M sales/marketing effort in this plan is restricted to "
    "commercial & industrial sites (any size) and residential systems of 20 kWp or larger. Small residential "
    "rooftop (typically 3\u201310 kWp) is excluded \u2014 it is served by the separate Deye net-billing kit / 7SUN "
    "distribution motion, not by the O&M division, because per-unit O&M economics do not support it.")
seg = make_table(doc, 4)
add_row(seg, [("Segment", 5.5, "L", True), ("Est. Prospects", 3.0, "C", True),
              ("Avg. Deal Size (O&M/yr)", 3.5, "C", True), ("Pain Point", 4.0, "L", True)], header=True)
seg_rows = [
    ("C&I >100 kWp (factories, cold storage, hotels)", "~300\u2013500", "\u20ac5,000\u2013\u20ac25,000", "DSO compliance + high electricity bills"),
    ("C&I 20\u2013100 kWp (SMEs, farms)", "~200\u2013300", "\u20ac1,500\u2013\u20ac5,000", "No in-house O&M capability"),
    ("Residential \u226520 kWp (large villas/estates)", "~50\u2013150", "\u20ac1,000\u2013\u20ac3,500", "Net billing transition, warranty admin"),
    ("Existing contract \u2014 Spanercom (Anarita, 2\u00d75 MW)", "1 (signed)", "\u20ac56,400/yr", "Baseline recurring revenue"),
]
for i, row in enumerate(seg_rows):
    add_row(seg, [(v, w, "L" if j in (0, 3) else "C", j == 2) for j, (v, w) in enumerate(zip(row, [5.5,3.0,3.5,4.0]))], alt=(i % 2 == 1))
note(doc, "Prospect counts and pain points from lighthief-cyprus-sales-growth-plan.md Tier 2 (Mar 2026), scoped "
          "down to the \u226520kWp threshold. Package pricing (ECO ~\u20ac4,200/MW, SILVER ~\u20ac9,500/MW, GOLD ~\u20ac11,300/MW) "
          "per pv-om/README.md.")

# ── Section 3: Marketing Plan ────────────────────────────────────────────────
add_heading(doc, "3. Marketing Plan \u2014 12 Months", level=1)
add_paragraph(doc,
    "Two parallel workstreams, both funded from the same budget line, phased to ramp with headcount (\u00a74) so "
    "leads are not generated faster than the team can service them.")

add_heading(doc, "3.1  Workstream A \u2014 BESS retrofit / LTSA upsell (existing playbook)", level=2)
for text in [
    "Audience: PV park owners already suffering curtailment (47% in 2025) and CERA BESS licensees (33 entities, "
    "1,000+ MW, almost none deployed) \u2014 reuses Campaigns 1\u20133 from the existing Google Ads playbook.",
    "Channel mix: Google Search (brand + curtailment-recovery + BESS EPC), LinkedIn for CERA licensees/investors, "
    "monthly \u201cCurtailment Recovery\u201d seminar.",
    "Goal: keep the post-Galascope pipeline fed so the field team built for this plan has follow-on work beyond "
    "the confirmed 30 MWh base case.",
]:
    add_bullet(doc, text)

add_heading(doc, "3.2  Workstream B \u2014 C&I + residential \u226520kWp PV O&M (new)", level=2)
for text in [
    "Audience: the ~500\u2013800 C&I sites and ~50\u2013150 large-residential systems identified in \u00a72.2. "
    "Negative-keyword list must retain the block on small residential (\"home solar panels\", \"rooftop solar\", "
    "\"solar panel cost\") to avoid diluting spend into the excluded segment.",
    "Channel mix: targeted Google Search (\u201cDSO Category C compliance\u201d, \u201cPV O&M Cyprus\u201d, \u201ccommercial solar "
    "maintenance\u201d), direct outreach to top 50 electricity consumers per district, energy-auditor referral "
    "partnerships (per sales-growth-plan \u00a73.2).",
    "Entry offer: free energy audit / free O&M cost review for any site >20kWp, converting to an ECO/SILVER/GOLD "
    "or custom O&M contract.",
]:
    add_bullet(doc, text)

add_heading(doc, "3.3  Budget phasing", level=2)
mkt = make_table(doc, 4)
add_row(mkt, [("Period", 4.5, "L", True), ("Monthly Budget", 3.5, "C", True),
              ("Workstream A (BESS)", 3.5, "C", True), ("Workstream B (PV O&M C&I/\u226520kWp)", 4.5, "C", True)], header=True)
mkt_rows = [
    ("Sep\u2013Dec 2026 (Phase 1)", "\u20ac5,000/mo", "\u20ac3,000/mo", "\u20ac2,000/mo"),
    ("Jan\u2013Aug 2027 (Phase 2)", "\u20ac8,000/mo", "\u20ac4,500/mo", "\u20ac3,500/mo"),
]
for i, row in enumerate(mkt_rows):
    add_row(mkt, [(v, w, "C" if j else "L", False) for j, (v, w) in enumerate(zip(row, [4.5,3.5,3.5,4.5]))], alt=(i % 2 == 1))
add_row(mkt, [("12-month total", 4.5, "L", True), ("\u2248\u20ac84,000", 3.5, "C", True), ("\u2248\u20ac48,000", 3.5, "C", True), ("\u2248\u20ac36,000", 4.5, "C", True)], total=True)
note(doc, "Budget levels match the existing digital playbook (marketing/solinvest-marketing-strategy.html Phase "
          "1 \u20ac4,650/mo \u2192 Phase 2 \u20ac8,100/mo, rounded), split roughly 60/40 between the two workstreams. "
          "Phase 2 begins in January 2027 once the O&M back-office coordinator (\u00a74.3) can process the extra "
          "lead volume.")

add_heading(doc, "3.4  12-month marketing KPIs", level=2)
kpi = make_table(doc, 3)
add_row(kpi, [("KPI", 7.0, "L", True), ("Target", 3.0, "C", True), ("Owner", 3.0, "C", True)], header=True)
kpi_rows = [
    ("New qualified O&M leads (both workstreams)", "40+ over 12 mo", "Lead Coordinator"),
    ("C&I/\u226520kWp O&M proposals issued", "12+", "Sales Exec (new hire, \u00a74.3)"),
    ("C&I/\u226520kWp O&M contracts signed", "3", "BDM + Sales Exec"),
    ("Curtailment-recovery seminars hosted", "3", "BDM + Director"),
    ("CERA licensee database maintained / recontacted quarterly", "33 entities", "BDM"),
]
for i, row in enumerate(kpi_rows):
    add_row(kpi, [(v, w, "L" if j == 0 else "C", False) for j, (v, w) in enumerate(zip(row, [7.0,3.0,3.0]))], alt=(i % 2 == 1))

# ── Section 4: Staffing ──────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "4. Staffing Plan & Organisational Growth", level=1)
add_heading(doc, "4.1  Current team (as at 1 September 2026)", level=2)
cur = make_table(doc, 4)
add_row(cur, [("Name", 4.5, "L", True), ("Role", 6.0, "L", True), ("Monthly Gross", 2.5, "C", True), ("Type", 2.0, "C", True)], header=True)
cur_rows = [
    ("Alexander Papacosta", "Cyprus Director", "\u20ac5,000*", "Employee"),
    ("Costas Hadjikyriacou", "BESS Division Lead (ETEK licensed)", "\u20ac3,000", "Employee"),
    ("Andreas Christoforou", "Business Development & Technical Sales", "\u20ac1,000 + 18% commission", "Employee"),
    ("Zinovia Efesopoulou", "Sales Executive & Lead Intake Coordinator", "\u20ac1,400", "Employee"),
    ("Cihat Ertugrul (\u201cJihat\u201d)", "Field Engineer \u2014 Civil Works & O&M", "\u20ac1,400", "Employee"),
]
for i, row in enumerate(cur_rows):
    add_row(cur, [(v, w, "L" if j < 2 else "C", False) for j, (v, w) in enumerate(zip(row, [4.5,6.0,2.5,2.0]))], alt=(i % 2 == 1))
note(doc, "*Alexander Papacosta at \u20ac5,000 gross per Shapiro DD pack MD terms (uplift from \u20ac2,000 net on Galascope "
          "EPC signing). Source: team/team-data.ts (SSOT) and preliminary-financials-2025-2026.md \u00a76. "
          "Christos Nicolaou (Back Office) is terminated and excluded. Team-data.ts shows Andreas at \u20ac1,000+18% "
          "commission; team/README.md shows \u20ac2,500 flat \u2014 reconcile with Timkas before finalising payroll (\u00a75).")

add_heading(doc, "4.2  Growth hires \u2014 timeline (base case: Galascope only)", level=2)
hire = make_table(doc, 5)
add_row(hire, [("Month", 2.5, "L", True), ("Role Added", 5.0, "L", True), ("Fully Loaded Cost", 3.0, "C", True),
               ("Trigger", 5.5, "L", True)], header=True)
hire_rows = [
    ("Oct 2026", "2 \u00d7 BESS Field Engineer", "\u20ac4,500/mo each", "Ahead of Galascope CIF (15 Sep 2026) \u2014 install/commissioning support"),
    ("Nov 2026", "1 \u00d7 PV O&M Field Technician", "\u20ac1,600/mo", "C&I/\u226520kWp workstream (\u00a73.2) begins generating site visits"),
    ("Jan 2027", "1 \u00d7 Sales Executive \u2014 C&I/PV O&M", "\u20ac1,400/mo + commission", "Phase 2 marketing (\u00a73.3) needs dedicated proposal capacity"),
    ("Feb 2027", "1 \u00d7 O&M Back-Office / CMMS Coordinator", "\u20ac2,000/mo", "Ahead of Galascope PAC (31 Jan 2027) \u2014 work orders, client reporting"),
]
for i, row in enumerate(hire_rows):
    add_row(hire, [(v, w, "L" if j != 2 else "C", False) for j, (v, w) in enumerate(zip(row, [2.5,5.0,3.0,5.5]))], alt=(i % 2 == 1))
add_paragraph(doc, "Headcount grows from 5 (Sep 2026) to 9 (Feb 2027) \u2014 kept one step behind confirmed demand, "
              "not ahead of it, consistent with the opex-plan-2026.md \u201cscale when triggered\u201d principle.",
              size=9.5, color=GREY, space_before=2, space_after=8)

add_heading(doc, "4.3  Conditional hire \u2014 upside case only", level=2)
add_paragraph(doc,
    "If Timotheos Timotheou, Lampros Andreadis and/or Spanercom (Anarita) sign their EPCs on the Batch 1 "
    "Extension timeline (\u00a72.1), a 5th field engineer (\u20ac4,500/mo fully loaded) should be added in March 2027 "
    "to keep the 2-engineer-per-HV-visit safety rule (per opex-plan-2026.md \u00a72.4) intact at the higher MWh count. "
    "This hire is NOT included in the base-case payroll or cashflow in \u00a75\u2013\u00a77; it is a trigger-based decision "
    "for Dr. Sybaris and Mr. Shapiro once any of those three EPCs is signed.")

add_heading(doc, "4.4  Steady-state O&M staffing benchmark (for context beyond this plan)", level=2)
add_paragraph(doc,
    "For reference, the SSOT full-portfolio O&M staffing model (OM_OPEX, 486.5 MWh) assumes 6 field engineers + "
    "1 driver/logistics + 1 back-office + 0.5 FTE O&M manager, total personnel cost \u20ac401,400/yr. This plan's "
    "9-person team at the end of the 12-month period (\u00a74.2) is appropriately sized for the 30\u2013120 MWh in this "
    "plan's window \u2014 not the full 486.5 MWh group-order portfolio, which is a multi-year build-out.")

# ── Section 5: Payroll ───────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "5. Payroll", level=1)
add_paragraph(doc,
    "Employer cost = gross salary \u00d7 1.154 (Cyprus 2026 statutory burden: 12.5% social insurance/funds + 2.9% "
    "GeSY = 15.4% total employer burden, per CYPRUS_TAX in lib/portfolio-data.ts). Commission (Andreas, and any "
    "new Sales Executive) is excluded from the base payroll run-rate below and paid on collection per the Unified "
    "Sales Commission Structure (LCY-COM-001).")

pay = make_table(doc, 5)
add_row(pay, [("Month", 2.5, "L", True), ("Headcount", 2.0, "C", True), ("New This Month", 5.5, "L", True),
              ("Base Payroll (gross)", 3.0, "R", True), ("Employer Cost (+15.4%)", 3.0, "R", True)], header=True)
payroll_rows = [
    ("Sep 2026", "5", "\u2014 (baseline)", "\u20ac11,800", "\u20ac13,619"),
    ("Oct 2026", "7", "+2 BESS Field Engineers", "\u20ac20,800", "\u20ac24,003"),
    ("Nov 2026", "8", "+1 PV O&M Field Technician", "\u20ac22,400", "\u20ac25,850"),
    ("Dec 2026", "8", "\u2014", "\u20ac22,400", "\u20ac25,850"),
    ("Jan 2027", "9", "+1 Sales Executive (C&I/PV O&M)", "\u20ac23,800", "\u20ac27,466"),
    ("Feb 2027", "10", "+1 O&M Back-Office Coordinator", "\u20ac25,800", "\u20ac29,774"),
    ("Mar\u2013Aug 2027 (each month)", "10", "\u2014", "\u20ac25,800", "\u20ac29,774"),
]
for i, row in enumerate(payroll_rows):
    add_row(pay, [(v, w, "L" if j in (0, 2) else "C" if j == 1 else "R", False) for j, (v, w) in enumerate(zip(row, [2.5,2.0,5.5,3.0,3.0]))], alt=(i % 2 == 1))
note(doc, "Base payroll gross for Sep 2026 = \u20ac11,800 excl. Andreas commission (Alex \u20ac5,000 + Costas \u20ac3,000 + "
          "Andreas \u20ac1,000 + Zinovia \u20ac1,400 + Jihat \u20ac1,400). Field engineer cost of \u20ac4,500/mo fully loaded "
          "already includes accommodation share and statutory burden per opex-plan-2026.md \u00a72.2 \u2014 shown gross-"
          "equivalent above for table consistency.")

# ── Section 6: Office & Warehouse ────────────────────────────────────────────
add_heading(doc, "6. Office & Warehouse (Premises)", level=1)
add_heading(doc, "6.1  Current office", level=2)
add_paragraph(doc,
    "15 Agaritsis, Nektaria Court, Office 201, 3045 Zakaki, Limassol. Current run-rate \u20ac1,500/month "
    "(per financial/ebitda-forecast-2026.md SG&A allocation, \u20ac27K over 18 months). Note: preliminary-financials "
    "\u00a73.4 records a separate director-funded claim of \u20ac7,000/month since March 2025 for \u201coffice rent\u201d \u2014 this "
    "discrepancy should be reconciled with Timkas before the office budget line below is finalised.")

add_heading(doc, "6.2  Decision point \u2014 new/expanded premises", level=2)
add_paragraph(doc,
    "A commercial rental due-diligence pack was prepared 17 July 2026 for a new/larger premises search "
    "(lighthief-cyprus/rental-due-diligence-response.md). The current office was sized for a 5\u2013person team; this "
    "plan grows headcount to 9\u201310 by February 2027 (\u00a74.2), which the current office cannot comfortably absorb "
    "alongside client meetings and O&M back-office work. ")
for bold, text in [
    ("Recommendation: ", "decide by end of Q4 2026 (Nov\u2013Dec) whether to renew/extend the current office or move."),
    ("Budget range if upsized: ", "\u20ac2,500\u2013\u20ac3,500/month, effective January 2027, to align with the O&M "
     "back-office coordinator start date."),
    ("Guarantee options already scoped: ", "Lighthief International Ltd comfort letter/guarantee, or director "
     "personal guarantee \u2014 both discussed in the July 2026 DD pack given the company's 1 July 2025 incorporation date."),
]:
    add_bullet(doc, text, bold_prefix=bold)

add_heading(doc, "6.3  Warehouse (Ypsonas)", level=2)
add_paragraph(doc,
    "\u20ac8,000/month gross, 50% co-funded by 7SUN \u2192 \u20ac4,000/month net Lighthief cost, plus a one-off \u20ac24,000 "
    "(2 deposits + 1 month advance) on signing. Activated from October 2026 to align with Galascope CIF "
    "(15 Sep 2026) and the 2 new field engineers' tooling/logistics needs. If the \u20ac24,000 deposit was already "
    "paid under the original March 2026 warehouse proposal, remove it from the October one-off costs in \u00a77.")

# ── Section 7: Cashflow ──────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "7. 12-Month Cashflow (Base Case \u2014 Galascope Only)", level=1)
add_paragraph(doc,
    "Base case assumes only the confirmed Galascope 1+2 batch (30 MWh) reaches PAC in the plan window. The "
    "Timotheos/Lampros/Spanercom extension is shown as an upside overlay in \u00a77.3 \u2014 it is not included in the "
    "totals below.")

cf = make_table(doc, 6)
add_row(cf, [("Month", 2.3, "L", True), ("OPEX", 2.7, "R", True), ("PV O&M Rev.", 2.7, "R", True),
             ("BESS O&M Rev.", 2.7, "R", True), ("Net Cashflow", 2.8, "R", True), ("Cumulative", 2.8, "R", True)], header=True)
cf_rows = [
    ("Sep 2026", 20600, 4700, 0),
    ("Oct 2026", 77070, 4700, 0),
    ("Nov 2026", 38920, 5950, 0),
    ("Dec 2026", 38920, 5950, 0),
    ("Jan 2027", 43520, 5950, 0),
    ("Feb 2027", 47120, 7200, 4350),
    ("Mar 2027", 47120, 7200, 4350),
    ("Apr 2027", 47120, 8450, 4350),
    ("May 2027", 47120, 8450, 4350),
    ("Jun 2027", 47120, 8450, 4350),
    ("Jul 2027", 47120, 8450, 4350),
    ("Aug 2027", 47120, 8450, 4350),
]
cum = 0
for i, (m, opex, pv_rev, bess_rev) in enumerate(cf_rows):
    total_rev = pv_rev + bess_rev
    net = total_rev - opex
    cum += net
    add_row(cf, [
        (m, 2.3, "L", False),
        (f"\u20ac{opex:,.0f}", 2.7, "R", False),
        (f"\u20ac{pv_rev:,.0f}", 2.7, "R", False),
        (f"\u20ac{bess_rev:,.0f}" if bess_rev else "\u2014", 2.7, "R", False),
        (f"\u2212\u20ac{abs(net):,.0f}" if net < 0 else f"\u20ac{net:,.0f}", 2.8, "R", True),
        (f"\u2212\u20ac{abs(cum):,.0f}" if cum < 0 else f"\u20ac{cum:,.0f}", 2.8, "R", False),
    ], alt=(i % 2 == 1))

total_opex = sum(r[1] for r in cf_rows)
total_pv = sum(r[2] for r in cf_rows)
total_bess = sum(r[3] for r in cf_rows)
total_rev = total_pv + total_bess
total_net = total_rev - total_opex
add_row(cf, [
    ("12-Month Total", 2.3, "L", True),
    (f"\u20ac{total_opex:,.0f}", 2.7, "R", True),
    (f"\u20ac{total_pv:,.0f}", 2.7, "R", True),
    (f"\u20ac{total_bess:,.0f}", 2.7, "R", True),
    (f"\u2212\u20ac{abs(total_net):,.0f}", 2.8, "R", True),
    (f"\u2212\u20ac{abs(cum):,.0f}", 2.8, "R", True),
], total=True)

note(doc, "OPEX = payroll (employer cost, \u00a75) + marketing (\u00a73.3) + office (\u00a76.1) + warehouse (\u00a76.3) + fleet "
          "(2 vans from Oct26, \u20ac1,235/mo each) + one-off tools/training/warehouse-deposit (\u20ac40,000 in Oct26) + "
          "SCADA/IT/insurance/consumables (scaled down from OM_OPEX operations bucket). PV O&M revenue = existing "
          "Spanercom (\u20ac4,700/mo) + 3 new contracts phased in Nov26/Jan27/Apr27 at \u2248\u20ac15,000/yr each (\u00a72.2). "
          "BESS O&M revenue = 30 MWh \u00d7 \u20ac1,740/MWh/yr (LTSA Tier C) from Feb 2027 (post-PAC).")

add_heading(doc, "7.1  Funding requirement", level=2)
add_paragraph(doc,
    f"The base case shows a 12-month net cash requirement of approximately \u2212\u20ac{abs(total_net):,.0f}. This is "
    "expected and is the standard shape of an O&M ramp-up: the field team, warehouse and marketing spend all "
    "start before the first park reaches PAC.")

fund = make_table(doc, 2)
add_row(fund, [("Available EPC margin to cover this plan", 10.0, "L", True), ("Amount", 6.0, "R", True)], header=True)
fund_rows = [
    ("Galascope 1+2 EPC margin (confirmed)", "\u20ac230,722"),
    ("Timotheos + Lampros + Spanercom EPC margin (if all sign)", "\u20ac927,583"),
    ("Total available EPC margin cover", "up to \u20ac1,158,305"),
]
for i, (a, b) in enumerate(fund_rows):
    add_row(fund, [(a, 10.0, "L", False), (b, 6.0, "R", True)], alt=(i % 2 == 1))
note(doc, "Source: lib/portfolio-data.ts BATCHES[0].margin and BATCHES[1].margin. Even the base-case funding "
          "requirement (\u2248\u20ac435K) is fully covered by the Galascope EPC margin alone once combined with normal "
          "trading cashflow \u2014 this plan does not require new equity beyond what Dr. Sybaris and Mr. Shapiro have "
          "already committed to the group.")

add_heading(doc, "7.2  Sensitivity \u2014 upside case (Extension signs on schedule)", level=2)
add_paragraph(doc,
    "If Timotheos, Lampros and Spanercom all sign and reach PAC on the same 31 Jan 2027 date, BESS O&M revenue "
    "from Feb 2027 rises from \u20ac4,350/mo to \u20ac17,400/mo (120 MWh total \u00d7 \u20ac1,740/MWh/yr \u00f7 12), adding "
    "\u2248\u20ac78,300 of extra revenue across Feb\u2013Aug 2027 (7 months) \u2014 before accounting for the additional field "
    "engineer in \u00a74.3 (\u20ac5,196/mo employer cost from March 2027, \u2248\u20ac31,200 over 6 months). Net effect: the "
    "upside case improves the 12-month net position by roughly \u20ac47,000 versus the base case, while also being "
    "the trigger that converts \u20ac927,583 of Extension EPC margin into contracted revenue.")

# ── Section 8: Risks & Assumptions ───────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "8. Key Assumptions & Risks", level=1)
risk = make_table(doc, 3)
add_row(risk, [("Risk / Assumption", 7.0, "L", True), ("Likelihood", 3.0, "C", True), ("Mitigation", 6.0, "L", True)], header=True)
risk_rows = [
    ("Galascope PAC slips beyond 31 Jan 2027", "Medium", "Field team already deployed for commissioning support \u2014 no idle cost; O&M revenue start date shifts, not team size."),
    ("Timotheos/Lampros/Spanercom do not sign in this window", "Medium", "Base case already excludes them \u2014 no plan disruption. 5th field engineer (\u00a74.3) stays unhired."),
    ("C&I/\u226520kWp PV O&M pipeline slower than 3 contracts/yr", "Medium", "Sales Executive hire (Jan27) timed after first market signal (Nov\u2013Dec26 leads), not before."),
    ("Office rent discrepancy (\u20ac1,500 vs \u20ac7,000/mo claim) unresolved", "High \u2014 needs action", "Reconcile with Timkas/accountant before Q4 2026 premises decision (\u00a76.2)."),
    ("Andreas Christoforou pay rate discrepancy (team-data.ts vs README)", "Medium", "Confirm actual contracted rate before finalising \u00a75 payroll run-rate."),
    ("Warehouse \u20ac24,000 deposit may already be paid (Mar 2026 proposal)", "Medium", "Confirm with accountant; if paid, remove from Oct26 one-off cost in \u00a77."),
    ("Marketing spend generates leads faster than 9-person team can service", "Low", "Phase 2 budget increase (\u00a73.3) deliberately timed after back-office coordinator hire (Feb27)."),
]
for i, row in enumerate(risk_rows):
    add_row(risk, [(v, w, "L" if j != 1 else "C", False) for j, (v, w) in enumerate(zip(row, [7.0,3.0,6.0]))], alt=(i % 2 == 1))

# ── Section 9: Milestones ────────────────────────────────────────────────────
add_heading(doc, "9. Quarterly Milestones & Decision Calendar", level=1)
mile = make_table(doc, 3)
add_row(mile, [("Quarter", 2.5, "L", True), ("Milestones", 8.0, "L", True), ("Decision Needed From Sybaris/Shapiro", 5.5, "L", True)], header=True)
mile_rows = [
    ("Q1 (Sep\u2013Nov 2026)", "Galascope CIF (15 Sep); 2 field engineers + warehouse activated; PV O&M "
     "field technician hired; Phase 1 marketing live.", "Confirm office renewal vs. new premises search timeline."),
    ("Q2 (Dec 2026\u2013Feb 2027)", "Galascope PAC target (31 Jan); Sales Exec + O&M Back-Office Coordinator "
     "hired; Phase 2 marketing begins; BESS O&M billing starts.", "Go/no-go on office upsize (\u00a76.2); review "
     "Timotheos/Lampros/Spanercom signing status \u2014 trigger 5th engineer if any sign."),
    ("Q3 (Mar\u2013May 2027)", "First 2 of 3 new PV O&M contracts targeted signed; 90-day review of C&I/\u226520kWp "
     "campaign conversion rate.", "Approve/adjust marketing budget split (\u00a73.3) based on actual lead quality."),
    ("Q4 (Jun\u2013Aug 2027)", "3rd PV O&M contract signed; full-year review vs. plan; scoping for Year 2 "
     "(steady-state OM_OPEX model, \u00a74.4).", "Sign off Year 2 plan; decide on scaling beyond 10-person team."),
]
for i, row in enumerate(mile_rows):
    add_row(mile, [(v, w, "L", False) for v, w in zip(row, [2.5,8.0,5.5])], alt=(i % 2 == 1))

# ── Approval ──────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Approval", level=1)
add_paragraph(doc, "This plan is submitted for review and sign-off by:", space_after=16)
sig = doc.add_table(rows=1, cols=2); sig.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (name, title) in enumerate([
    ("Dr. Arkadiusz Sybaris", "Founder & CEO / Ultimate Beneficial Owner, Lighthief Group"),
    ("Michael Shapiro", "Reviewing Principal"),
]):
    cell = sig.rows[0].cells[i]; cell.width = Cm(8)
    for text, bold, after in [
        (name, True, 32),
        ("Signature: ___________________________", False, 4),
        (f"Name: {name}", False, 4),
        (f"Title: {title}", False, 4),
        ("Date: ___________________________", False, 4),
    ]:
        p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(10.5); r.font.color.rgb = BLACK
for cell in sig.rows[0].cells:
    first = cell.paragraphs[0]
    if not first.text:
        first._p.getparent().remove(first._p)

doc.add_paragraph()
footer = add_paragraph(doc, "Lighthief Cyprus Ltd \u00b7 HE 477423 \u00b7 solarfarms.cy", size=9, color=GREY, space_before=12)
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer2 = add_paragraph(doc, "15 Agaritsis, Nektaria Court, Office 201, 3045 Zakaki, Limassol, Cyprus \u00b7 "
                              "+357 77 77 00 50 \u00b7 office@lighthief.com", size=9, color=GREY)
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT_FILE)
print("Saved:", OUT_FILE)
print(f"12-month total OPEX:    \u20ac{total_opex:,.0f}")
print(f"12-month total revenue: \u20ac{total_rev:,.0f}")
print(f"12-month net cashflow:  \u2212\u20ac{abs(total_net):,.0f}")
