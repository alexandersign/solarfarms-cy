#!/usr/bin/env python3
"""Generate OEM Direct Warranty Undertaking (Linyang Step-In Letter) DOCX"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BLUE = RGBColor(0x1A, 0x36, 0x5D)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)

def para(text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)

# HEADER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('OEM PRODUCT WARRANTY CONFIRMATION\nAND DIRECT UNDERTAKING')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BLUE

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ['Document Reference: LCY-OEM-DWU-001',
             'Version: 1.0',
             'Date: [\u25cf] March 2026']:
    r = meta.add_run(line + '\n')
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

doc.add_paragraph()
doc.add_paragraph()

# FROM / TO
para('FROM:', bold=True)
para('Jiangsu Linyang Energy Storage Technology Co., Ltd')
para('No. 1 Building, International R&D Headquarters Park, Xincheng Science and Technology Park, 68 Aoti Street, Jianye District, Nanjing City, Jiangsu Province, China')
para('("Linyang" or "OEM")')
doc.add_paragraph()

para('TO:', bold=True)
para('[End-Customer Legal Name]')
para('[End-Customer Address]')
para('("End-Customer")')
doc.add_paragraph()

para('COPY TO:', bold=True)
para('Lighthief Cyprus Ltd')
para('28 October Avenue 249, Lophitis Business Center 1, Office 201, 3035 Limassol, Cyprus')
para('("Distributor")')
doc.add_paragraph()

para('RE: Direct Warranty Undertaking for BESS Equipment purchased through Lighthief Cyprus Ltd (Authorised Exclusive Distributor)', bold=True)
doc.add_paragraph()

# RECITALS
para('RECITALS', bold=True, size=12)
doc.add_paragraph()
para('A. Linyang is the manufacturer and supplier of battery energy storage systems, power conversion systems, and related equipment ("Products").')
para('B. Lighthief Cyprus Ltd is Linyang\u2019s exclusive authorised distributor for BESS products in the Republic of Cyprus under a Distribution Agreement dated [22 February 2026].')
para('C. The End-Customer has entered into (or intends to enter into) an EPC Agreement with Lighthief Cyprus Ltd for the supply and installation of BESS Products at [Site Name / Park Name], [District], Cyprus ("Project").')
para('D. Linyang wishes to confirm that its standard product warranty applies to Products purchased by the End-Customer through the Distributor, and that such warranty is enforceable directly by the End-Customer in the circumstances described below.')
doc.add_paragraph()

# UNDERTAKING
para('UNDERTAKING', bold=True, size=12)
doc.add_paragraph()

para('1. WARRANTY CONFIRMATION', bold=True)
para('1.1 Linyang confirms that all Products supplied for the Project through the Distributor carry Linyang\u2019s standard product warranty of five (5) years from the Warranty Start Date (the earlier of commissioning (PAC) or six (6) months after shipment from Linyang\u2019s factory).')
para('1.2 The warranty covers defects in design, materials, and workmanship in accordance with Linyang\u2019s Warranty Terms (Version 2) and the Sales Contract between Linyang and the Distributor.')

para('2. SCOPE OF WARRANTY', bold=True)
para('2.1 The following Products are covered:')
bullet('(a) Battery containers, modules, and cells;')
bullet('(b) Power Conversion System (PCS);')
bullet('(c) Battery Management System (BMS);')
bullet('(d) Thermal management and cooling systems;')
bullet('(e) Fire suppression systems;')
bullet('(f) OEM-supplied cables, connectors, and ancillaries;')
bullet('(g) Installation and commissioning workmanship (where performed by Linyang or its authorised personnel).')

para('3. PERFORMANCE GUARANTEES', bold=True)
para('3.1 Linyang guarantees the following performance levels for Products supplied to the Project:')
bullet('(a) State of Health (SOH) at 1 cycle per day: Year 5 \u226585%, Year 10 \u226579.58%, Year 15 \u226570%;')
bullet('(b) Round-Trip Efficiency (RTE): \u226586.32% at rated power (0.5C, 25\u00b12\u00b0C);')
bullet('(c) Cycle Life: 7,000 equivalent full cycles at 0.5C, 90% DoD, to 70% EOL.')
para('3.2 Remedy: If Products fail to meet the above thresholds, Linyang shall supply replacement modules (FOB or CIF Limassol as applicable) at Linyang\u2019s cost to restore performance above the guaranteed threshold.')

para('4. DIRECT ENFORCEMENT RIGHT', bold=True)
para('4.1 If the Distributor (Lighthief Cyprus Ltd) is unable to fulfil its warranty obligations to the End-Customer for any reason, including but not limited to:', bold=False)
bullet('(a) insolvency, liquidation, or dissolution of the Distributor;')
bullet('(b) cessation of the Distributor\u2019s business operations;')
bullet('(c) termination of the Distribution Agreement between Linyang and the Distributor;')
bullet('(d) failure or refusal by the Distributor to process a valid warranty claim within thirty (30) days of notification;')
para('then the End-Customer may enforce the warranty directly against Linyang by written notice to Linyang at the address above, providing:')
bullet('(i) evidence of purchase through the Distributor (EPC Agreement, delivery records, or PAC certificate);')
bullet('(ii) description of the defect and supporting evidence (photographs, BMS logs, test reports);')
bullet('(iii) evidence that the Distributor is unable to fulfil the claim (e.g., confirmation of insolvency, or evidence of 30-day non-response).')

para('4.2 Upon receipt of a valid direct claim, Linyang shall respond within fourteen (14) days and, if the claim is accepted, shall perform the warranty remedy in accordance with its standard warranty terms.')

para('5. WARRANTY CONDITIONS AND EXCLUSIONS', bold=True)
para('5.1 The warranty under this Undertaking is subject to the same conditions and exclusions as Linyang\u2019s standard warranty, including:')
bullet('(a) Operation of Products in accordance with Linyang\u2019s user manuals and guidelines;')
bullet('(b) Maintenance by qualified, authorised personnel;')
bullet('(c) Serial numbers remaining intact;')
bullet('(d) Products not modified, relocated, or repaired without Linyang\u2019s written consent.')
para('5.2 The warranty-voiding battery conditions apply as per Linyang\u2019s Warranty Terms (Version 2), including cell voltage and SOC thresholds.')
para('5.3 Coastal installation restrictions (C5 enclosure) apply as per the Sales Contract.')

para('6. INSURANCE', bold=True)
para('6.1 Linyang confirms that it maintains product liability insurance of EUR 5,000,000 per occurrence (currently with AXA Tianping), covering Products supplied for installation in the Republic of Cyprus.')
para('6.2 Linyang shall provide evidence of insurance coverage to the End-Customer upon written request.')

para('7. NO ADDITIONAL OBLIGATIONS', bold=True)
para('7.1 This Undertaking does not create any obligation on Linyang to provide services, maintenance, monitoring, or support beyond the standard product warranty.')
para('7.2 The 97% availability guarantee and associated liquidated damages are governed exclusively by the LTSA between the End-Customer and the Distributor, and do not form part of this Undertaking.')
para('7.3 Extended warranty (years 6\u201315) is a separate arrangement between the End-Customer and Linyang, subject to separate agreement and payment.')

para('8. DURATION', bold=True)
para('8.1 This Undertaking is effective from the date of delivery of Products to the Project Site and remains in force for the duration of the warranty period (5 years from the Warranty Start Date).')
para('8.2 This Undertaking survives termination of the Distribution Agreement and/or the Sales Contract between Linyang and the Distributor.')

para('9. GOVERNING LAW', bold=True)
para('9.1 This Undertaking shall be governed by the laws of Singapore.')
para('9.2 Any dispute arising from this Undertaking shall be resolved by arbitration under the SIAC Rules, seated in Singapore, conducted in English by three (3) arbitrators.')
para('9.3 Notwithstanding the above, the End-Customer may apply to the courts of Cyprus for interim relief, asset preservation, or enforcement of awards.')

para('10. CONTACT FOR WARRANTY CLAIMS', bold=True)
para('10.1 Primary contact:')
bullet('Conor Yang \u2014 conoryang@linyang.com.cn')
bullet('Kamil Tyburski \u2014 kamil@linyang.com')
bullet('Tomasz Wieckowski \u2014 tomasz.wieckowski@linyang.com')
para('10.2 Notice address: Jiangsu Linyang Energy Storage Technology Co., Ltd, No. 1 Building, International R&D Headquarters Park, Xincheng Science and Technology Park, 68 Aoti Street, Jianye District, Nanjing City, Jiangsu Province, China.')

doc.add_paragraph()
doc.add_paragraph()

# SIGNATURES
para('SIGNED for and on behalf of Jiangsu Linyang Energy Storage Technology Co., Ltd:', bold=True)
doc.add_paragraph()
para('Name: __________________________')
para('Title: __________________________')
para('Signature: _____________________')
para('Date: __________________________')
para('Company Seal:')
doc.add_paragraph()
doc.add_paragraph()

para('ACKNOWLEDGED by Lighthief Cyprus Ltd (Distributor):', bold=True)
doc.add_paragraph()
para('Name: Alexander Papacosta')
para('Title: Director')
para('Signature: _____________________')
para('Date: __________________________')
doc.add_paragraph()
doc.add_paragraph()

para('ACKNOWLEDGED by [End-Customer]:', bold=True)
doc.add_paragraph()
para('Name: __________________________')
para('Title: __________________________')
para('Signature: _____________________')
para('Date: __________________________')

doc.save('OEM-Direct-Warranty-Undertaking-Linyang.docx')
print(f'Saved: OEM-Direct-Warranty-Undertaking-Linyang.docx ({len(doc.paragraphs)} paragraphs)')
