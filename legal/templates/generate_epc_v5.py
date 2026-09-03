#!/usr/bin/env python3
"""
Generate EPC Agreement v5.0 DOCX
Changes from v4.0:
- Two-component price split (Component A: Equipment Supply / Component B: EPC Services)
- APG covers only Component A pre-delivery payments (advance 20% + pre-shipment 50% = 70% of Comp A)
- FAC removed; retention released 3 months after PAC
- Liability cap: tiered 10%/50%/uncapped matching upstream Linyang
- Termination non-payment: pro-rata refund of excess over works done (QS valuation)
- Warranty Start Date: Contractor-fault exception added
- RTE 84% hard floor clarified (no cumulative deduction)
- Warranty void: grid/EMS/FM carve-outs added
- FM termination: excess advance refund added
- EMS affiliate: Lighthief Cyprus parent guarantee added
- Client indemnity: proportionality qualifier added
- Title for lenders: early-title negotiation clause added
- Schedule C (OEM Direct Warranty Undertaking) annexed
- Version history updated
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BLUE = RGBColor(0x1A, 0x36, 0x5D)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BLUE

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(16)

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(10)

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10)

def bullet(text, level=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)

# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════
doc.add_paragraph()
h1('ENGINEERING, PROCUREMENT & CONSTRUCTION AGREEMENT')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Battery Energy Storage System (BESS)')
r.font.size = Pt(12); r.font.color.rgb = BLUE
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    'Document Reference: LCY-EPC-001',
    'Version: 5.0',
    'Date: 6 May 2026',
    '',
    'VERSION HISTORY',
    'v1.0 \u2014 Oct 2025: Initial draft',
    'v2.0 \u2014 Jan 2026: Linyang RFI confirmed terms',
    'v3.0 \u2014 22 Feb 2026: Installation scope, delay LD framework, bankability',
    'v4.0 \u2014 17 Mar 2026: Single counterparty warranty, title at PAC, unified warranty, lender assignment',
    'v5.0 \u2014 6 May 2026: FINAL BANKABLE VERSION \u2014 two-component price split,',
    '       APG on equipment only, 3-month retention release (no FAC),',
    '       tiered liability cap (10%/50%/uncapped), pro-rata refund on termination,',
    '       RTE hard floor clarified, warranty void carve-outs, all Anastasis v4.0 comments addressed',
]:
    r = meta.add_run(line + '\n')
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    if 'v5.0' in line:
        r.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. PARTIES
# ═══════════════════════════════════════════════════════════
h2('1. PARTIES')
para('This Engineering, Procurement and Construction Agreement (\u201cAgreement\u201d) is made as of [\u25cf] 2026 (\u201cEffective Date\u201d)')
para('BETWEEN:')
para('Lighthief Cyprus Ltd, a company incorporated under the laws of the Republic of Cyprus, with registered office at Agiou Andreou 241, AG TRIAS COURT, Flat/Office 31, 3036 Limassol, Cyprus, and operational address at 15 Agaritsis, Nektaria Court, Office 201, 3045 Zakaki, Limassol, Cyprus, Company No. HE 477423, TIN 60187188Q (\u201cContractor\u201d);')
para('and')
para('[\u25cf End Customer Legal Name], a company incorporated under the laws of [\u25cf], with registered office at [\u25cf] (\u201cClient\u201d).')
para('Contractor and Client are referred to individually as a \u201cParty\u201d and collectively as the \u201cParties.\u201d')

# 1A. CONDITIONS PRECEDENT
h2('1A. CONDITIONS PRECEDENT')
para('1A.1 This Agreement is conditional upon the Client obtaining grid connection terms from EAC/DSO for the Site (\u201cConnection Terms\u201d) within [\u25cf] months of the Effective Date (\u201cLong-Stop Date\u201d).')
para('1A.2 If Connection Terms are not obtained by the Long-Stop Date, either Party may terminate this Agreement by written notice without liability. The Contractor shall refund the advance payment within thirty (30) days, less reasonable and documented costs incurred.')
para('1A.3 The advance payment under Section 7.1(a) shall become due within seven (7) days of the later of: (i) the Effective Date; and (ii) receipt of Connection Terms.')

# 2. BACKGROUND
h2('2. BACKGROUND')
para('2.1 The Client owns and operates one or more renewable energy facilities and wishes to procure a battery energy storage system.')
para('2.2 The Contractor has the capability to design, integrate, install, commission, and deliver battery energy storage projects on a turnkey EPC basis.')
para('2.3 The Parties wish to set out the terms under which the Contractor shall deliver the BESS to the Client.')

# 3. DEFINITIONS
h2('3. DEFINITIONS AND INTERPRETATION')
h3('3.1 Defined Terms')
defs = [
    ('"Agreement"', 'means this EPC Agreement, including all schedules, appendices, and amendments.'),
    ('"BESS"', 'means the battery energy storage system supplied under this Agreement, comprising the Linyang Power Atlantic LFP liquid-cooled battery container system, including battery packs, BMS, thermal management, fire suppression, and related components as described in Schedule A.'),
    ('"Business Day"', 'means a day other than a Saturday, Sunday, or public holiday in the Republic of Cyprus.'),
    ('"Component A \u2014 Equipment Supply Price"', 'means the portion of the Contract Price attributable to supply of BESS equipment CIF Limassol, as set out in Section 6.2.'),
    ('"Component B \u2014 EPC Services Price"', 'means the portion of the Contract Price attributable to EPC services, as set out in Section 6.3.'),
    ('"Connection Terms"', 'means the grid connection terms issued by EAC/DSO for the Site.'),
    ('"Contract Price"', 'has the meaning given in Section 6.1.'),
    ('"Defects Liability Period"', 'means the period of three (3) months commencing from the date of PAC, during which the Contractor shall remedy all defects notified by the Client.'),
    ('"Delay LDs"', 'means the liquidated damages payable by the Contractor for delay per Section 8.4.'),
    ('"Effective Date"', 'has the meaning given in Section 1.'),
    ('"EMS"', 'means Energy Management System. EMS integration is provided by Lighthief EU BESS Ltd under a separate EMS Integration Agreement, guaranteed by the Contractor per Section 4.4.'),
    ('"Force Majeure"', 'has the meaning given in Section 12.1.'),
    ('"LTSA"', 'means a Long-Term Service Agreement entered into separately between the Parties.'),
    ('"OEM"', 'means Original Equipment Manufacturer (Linyang Energy).'),
    ('"OEM Direct Warranty Undertaking"', 'means the direct warranty confirmation letter issued by the OEM to the Client (Schedule C).'),
    ('"PAC"', 'means Provisional Acceptance Certificate per Section 9.1.'),
    ('"Products"', 'means BESS, PCS, battery containers, control equipment, and related components.'),
    ('"Retention"', 'means five percent (5%) of the Contract Price withheld by the Client and released at the end of the Defects Liability Period per Section 7.4.'),
    ('"Site"', 'means the installation location specified in Schedule A.'),
    ('"Target PAC Date"', 'means the date by which PAC is to be achieved per Section 8.4.'),
    ('"Warranty Period"', 'means five (5) years from the Warranty Start Date.'),
    ('"Warranty Start Date"', 'means: (a) the date of PAC; or (b) if PAC is delayed beyond six (6) months from equipment delivery to Site, the date six (6) months after delivery \u2014 but only where such delay is not caused by the Contractor. Where PAC is delayed by causes attributable to the Contractor, the Warranty Start Date shall be the actual date of PAC.'),
]
for term, defn in defs:
    p = doc.add_paragraph()
    r1 = p.add_run(term + ' ')
    r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(defn)
    r2.font.size = Pt(10)

h3('3.2 Interpretation')
for t in ['(a) Headings are for convenience only.', '(b) "Including" means without limitation.', '(c) References to laws include amendments.', '(d) Official language is English.']:
    bullet(t)

# 4. SCOPE
h2('4. SCOPE OF WORK')
para('4.1 Turnkey EPC Solution: The Contractor shall provide a complete turnkey EPC solution for the design, procurement, delivery, installation, commissioning, and handover of the BESS. The Contractor is the Client\u2019s sole counterparty for all works under this Agreement.', bold=True)

h3('4.2 Contractor Scope')
para('The Contractor shall be responsible for, and warrants the performance of:')
for item in [
    '(a) Project management and coordination;',
    '(b) Detailed engineering and system integration design;',
    '(c) Procurement of BESS equipment from the OEM (Component A);',
    '(d) International logistics, import duties, customs clearance, and port-to-site transport (Component B);',
    '(e) Civil works, foundations, fencing, and site preparation (Component B);',
    '(f) AC/DC cabling, grounding, and auxiliary electrical systems (Component B);',
    '(g) Equipment installation, assembly, and commissioning by OEM-trained personnel under the Contractor\u2019s coordination and responsibility (Component B);',
    '(h) SCADA integration and grid compliance coordination (Component B);',
    '(i) Authority liaison, EAC submissions, and regulatory compliance (Component B);',
    '(j) Project documentation, training, and handover (Component B).',
]:
    bullet(item)

h3('4.3 Single-Point Responsibility')
para('(a) The Contractor is responsible to the Client for the entire BESS system including OEM equipment and installation workmanship. The Client\u2019s sole recourse for any defect, delay, or performance issue is against the Contractor.')
para('(b) Subcontracting of installation or commissioning to OEM personnel does not relieve the Contractor of its obligations.')
para('(c) The Contractor warrants the performance of OEM equipment as if it were the Contractor\u2019s own work.')

h3('4.4 EMS Integration')
para('(a) EMS integration and SCADA commissioning are provided by Lighthief EU BESS Ltd, an affiliate of the Contractor, under a separate EMS Integration Agreement.')
para('(b) The Contractor hereby guarantees the performance of Lighthief EU BESS Ltd\u2019s obligations under the EMS Integration Agreement. If Lighthief EU BESS Ltd fails to perform, the Contractor shall fulfil those obligations directly or procure equivalent performance at the Contractor\u2019s cost.', bold=True)
para('(c) The Contractor warrants that the BESS is interface-compatible with the EMS at commissioning and that BMS-to-EMS communication protocols function correctly.')
para('(d) Advanced EMS features, analytics, optimisation algorithms, market participation logic, or software licences shall be contracted separately.')
para('(e) The Contractor is not responsible for EMS optimisation outcomes or revenue generation beyond interface compatibility.')

para('4.5 Any works not expressly listed are excluded unless agreed in writing.')

# 5A. CLIENT OBLIGATIONS
h2('5A. CLIENT OBLIGATIONS')
h3('5A.1 Pre-Commencement')
para('The Client shall:')
for item in ['(a) Provide complete and accurate site information;', '(b) Obtain planning permissions, building permits, and landowner consents;', '(c) Ensure clear site access for delivery and installation;', '(d) Provide secure storage for equipment;', '(e) Provide electrical connection point and grid connection approval.']:
    bullet(item)
h3('5A.2 During Works')
para('The Client shall:')
for item in ['(a) Provide site access during agreed working hours;', '(b) Designate a single point of contact;', '(c) Respond to queries within five (5) Business Days;', '(d) Provide timely approvals;', '(e) Not materially interfere with works.']:
    bullet(item)
h3('5A.3 Post-Commissioning')
para('The Client shall:')
for item in ['(a) Operate BESS per OEM guidelines;', '(b) Maintain monitoring systems;', '(c) Maintain site security and environmental conditions;', '(d) Maintain insurance;', '(e) Make all payments when due;', '(f) Permit Contractor access for warranty/maintenance.']:
    bullet(item)
h3('5A.4 Consequences of Client Breach')
para('(a) Delay: If Contractor\u2019s works are delayed more than five (5) Business Days by Client\u2019s material failure, Target PAC Date extends day-for-day and Client reimburses reasonable documented costs.')
para('(b) Payment Default: If the Client fails to make any payment when due:')
bullet('(i) Interest at EU Directive 2011/7/EU rate (ECB + 8% pa, simple interest);', 1)
bullet('(ii) Contractor may suspend works on seven (7) days\u2019 notice;', 1)
bullet('(iii) If overdue more than thirty (30) days, Contractor may terminate. Upon termination, the Contractor shall retain from payments received an amount equal to the value of works completed and materials delivered to Site, as agreed between the Parties within fourteen (14) days (or as determined by an independent quantity surveyor within thirty (30) days). Any excess payments received over such amount shall be refunded to the Client within thirty (30) days of determination.', 1)
para('(c) Warranty: Client breach of Sections 10.4 or 10.5 may limit warranty to the extent the breach caused or contributed to the defect.')
para('(d) No Contractor Liability for delays or losses arising from Client\u2019s material failure to fulfil obligations.')

# 6. CONTRACT PRICE — TWO COMPONENTS
h2('6. CONTRACT PRICE')
para('6.1 The total contract price (\u201cContract Price\u201d) is the sum of Component A and Component B as set out in Sections 6.2 and 6.3, exclusive of VAT. All milestone percentages in Section 7.1 are calculated on the Contract Price exclusive of VAT. VAT shall be added to each invoice in accordance with Cyprus law.')

h3('6.2 Component A \u2014 Equipment Supply Price')
para('EUR [\u25cf] (\u201cEquipment Supply Price\u201d), being the CIF Limassol value of BESS equipment supplied under this Agreement, comprising:')
for item in ['(a) Battery containers and modules;', '(b) Power Conversion System (PCS);', '(c) Medium Voltage Skid and switchgear;', '(d) Main step-up transformer;', '(e) Battery Management System (BMS) and thermal management;', '(f) OEM-supplied cables, connectors, and ancillaries.']:
    bullet(item)
para('The Advance Payment Guarantee (APG) under Section 10.9 is secured against the Component A Equipment Supply Price only.')

h3('6.3 Component B \u2014 EPC Services Price')
para('EUR [\u25cf] (\u201cEPC Services Price\u201d), being the cost of EPC services under this Agreement, comprising:')
for item in ['(a) International logistics, import duties, customs clearance, and port-to-site transport;', '(b) Civil works, foundations, fencing, and site preparation;', '(c) AC/DC cabling, grounding, and auxiliary electrical systems;', '(d) Installation coordination and commissioning support;', '(e) SCADA integration and grid compliance;', '(f) Authority liaison, EAC submissions, and project management;', '(g) Documentation and handover.']:
    bullet(item)
para('No APG is required against the Component B EPC Services Price, which is earned progressively against services performed.')

para('6.4 The Contract Price is based on current technical, regulatory, customs, and tax conditions. Material changes in scope, regulations, or law occurring after the Effective Date entitle the Parties to agree a variation by written change order.')

# 7. PAYMENT
h2('7. PAYMENT TERMS')
h3('7.1 Payment Milestones')
para('Payments shall be made by bank transfer as follows:')
para('(a) 30% advance payment within seven (7) days of the payment trigger under Section 1A.3;')
para('(b) 55% payment prior to shipment of BESS equipment, upon Contractor\u2019s written confirmation of successful Factory Acceptance Testing (FAT) and photographic evidence of readiness;')
para('(c) 10% payment upon issuance of PAC;')
para('(d) 5% Retention released at the end of the Defects Liability Period (three (3) months after PAC), subject to Section 7.4.')
para('7.2 Payments exclusive of VAT.')
para('7.3 Late payments: Contractor may suspend on seven (7) days\u2019 notice. Interest per Section 5A.4(b)(i).')

h3('7.4 Retention')
para('(a) The Client shall retain five percent (5%) of the Contract Price (\u201cRetention\u201d) as security for defect rectification during the Defects Liability Period.')
para('(b) Retention shall be released within thirty (30) days of the end of the Defects Liability Period (three (3) months after PAC).')
para('(c) Where punch-list items remain outstanding at the end of the Defects Liability Period, the Client may deduct from the Retention only the reasonable and documented cost of completing or procuring completion of those outstanding items. All other amounts shall be released.')
para('(d) The Client shall not withhold, set off, or deduct from the Retention any amounts other than those directly attributable to punch-list items the Contractor has failed to rectify after written notice.')

h3('7.5 TITLE AND RISK TRANSFER')
para('(a) Risk passes to Client upon delivery to Site and unloading.')
para('(b) Title (ownership) passes to the Client upon issuance of PAC and receipt of the PAC payment under Section 7.1(c). The Retention does not defer title.', bold=True)
para('(c) Defects Liability: The Contractor undertakes to rectify all defects notified during the Defects Liability Period at its own cost. This obligation is the consideration for the Retention.')
para('(d) Lender Security: Upon title transfer, the Client may grant security interests, charges, or pledges over the Equipment to lenders without the Contractor\u2019s consent.')
para('(e) Early Title for Lenders: If the Client\u2019s project finance lender requires title to pass at an earlier stage (such as at equipment delivery or FAT), the Parties shall negotiate an alternative security or title transfer arrangement in good faith, provided it does not materially increase the Contractor\u2019s financial exposure.')

# 8. DELIVERY AND DELAY LDs
h2('8. DELIVERY AND INSTALLATION')
para('8.1 Equipment supply CIF Limassol, Cyprus (Incoterms\u00ae 2020).')
para('8.2 Installation commences upon site readiness and delivery.')
para('8.3 Typical installation period: 4\u20136 weeks.')

h3('8.4 Delay Liquidated Damages')
h3('8.4.1 Definitions')
para('(a) "Target PAC Date" \u2014 agreed PAC date per Schedule A.')
para('(b) "Critical Path Equipment" \u2014 PCS, MV switchgear/skid, main step-up transformer. Delay = Full Park Delay.')
para('(c) "Non-Critical Path Equipment" \u2014 individual battery containers where PCS/MV can commission remaining units. Delay = Partial Delay.')
para('(d) "Full Park Delay" \u2014 no portion of Park can commission.')
para('(e) "Partial Delay" \u2014 reduced capacity commissioning possible.')
para('(f) "Affected Capacity" \u2014 delayed MWh as % of total MWh.')
para('(g) "LD Value" \u2014 Full Park: total Contract Price. Partial: Contract Price \u00d7 Affected Capacity %.')

h3('8.4.2 Obligation')
para('Contractor shall achieve PAC by Target PAC Date.')

h3('8.4.3 Calculation')
para('(a) Delay LDs on LD Value:')
bullet('(i) Days 1\u201330: 0.1% per calendar day;', 1)
bullet('(ii) Days 31\u201360: 0.15% per calendar day;', 1)
bullet('(iii) Day 61+: 0.2% per calendar day.', 1)
para('(b) Full Park: LDs on entire Contract Price. Partial: LDs on proportional value only. Escalation: recalculates if Partial becomes Full. See Schedule B for worked examples.')

h3('8.4.4 Cap')
para('Maximum aggregate Delay LDs per Park shall not exceed ten percent (10%) of the Contract Price. This cap and the termination right in Section 8.4.5 are the primary remedies for delay and together provide adequate protection.')

h3('8.4.5 Termination for Prolonged Delay')
para('If PAC is not achieved within one hundred and eighty (180) days of the Target PAC Date, the Client may terminate this Agreement upon written notice. Upon such termination:')
bullet('(a) The Contractor shall retain from payments received an amount equal to the value of works completed and materials delivered to Site;')
bullet('(b) Valuation shall be agreed between the Parties within fourteen (14) days. If not agreed, an independent quantity surveyor shall determine the value within thirty (30) additional days, applying a pro-rata assessment of completed works against the project programme;')
bullet('(c) All amounts received in excess of the determined value shall be refunded to the Client within thirty (30) days of the valuation determination;')
bullet('(d) Deductions shall be limited to actual and documented costs of works completed and materials delivered, evidenced by delivery records and progress reports.')

h3('8.4.6 Exclusions')
para('LDs not applicable for: (i) Force Majeure; (ii) Client-caused delays; (iii) grid authority delays; (iv) Client change orders; (v) customs delays with timely shipment proved; (vi) Client\u2019s late confirmation of technical agreements or drawings.')

h3('8.4.7 Extension of Time')
para('Target PAC Date extends day-for-day. Contractor notifies within seven (7) Business Days.')

# 9. ACCEPTANCE
h2('9. ACCEPTANCE')
h3('9.1 Provisional Acceptance (PAC)')
para('PAC shall be issued when the following criteria are met:')
bullet('(a) The BESS operates at not less than 95% of rated energy capacity;')
bullet('(b) Round-Trip Efficiency is not less than 84% under site conditions. For the avoidance of doubt, 84% is the hard floor for PAC; this threshold already incorporates a field tolerance adjustment from the OEM reference RTE of 86.32% and no further deduction applies;', )
bullet('(c) BMS/SCADA communications verified and operational;')
bullet('(d) All safety and protection systems functional and tested;')
bullet('(e) All mandatory commissioning tests per Schedule A passed.')

h3('9.2 Defects Liability Period')
para('Following PAC, the Defects Liability Period of three (3) months commences. During this period:')
bullet('(a) The Contractor shall remedy all defects notified by the Client promptly and at no cost;')
bullet('(b) Minor defects that do not materially affect system operation shall be recorded on a punch list and remedied within the Defects Liability Period;')
bullet('(c) At the end of the Defects Liability Period, the Retention is released per Section 7.4.')
para('There is no Final Acceptance Certificate. Retention is released automatically at the end of the Defects Liability Period subject to Section 7.4(c).')

# 10. WARRANTY
h2('10. WARRANTY')
h3('10.1 Unified Warranty')
para('The Contractor warrants the entire BESS system \u2014 including all OEM equipment, installation workmanship, and EPC works \u2014 against defects in design, materials, and workmanship for the Warranty Period. The Contractor is the Client\u2019s sole warranty counterparty.', bold=True)

h3('10.2 Coastal Installation')
bullet('(a) C5-rated enclosures as standard for Cyprus.')
bullet('(b) >500m from sea: full 5-year warranty.')
bullet('(c) <500m: PCS/transformer/switchgear warranty reduced to 2 years.')
bullet('(d) Client confirms coastal distance at signing (Schedule A).')

h3('10.3 Exclusions')
para('Warranty excludes: (a) consumables; (b) misuse or negligence by the Client; (c) unauthorised modifications; (d) Force Majeure damage; (e) relocation without approval.')

h3('10.4 Conditions')
para('Subject to: (a) operation per OEM guidelines; (b) qualified maintenance; (c) serial numbers intact; (d) access for inspection; (e) full payment.')

h3('10.5 WARRANTY-VOIDING BATTERY CONDITIONS')
para('The following conditions void the battery warranty:')
bullet('(a) Cell voltage below 2.8V for 120 consecutive hours;')
bullet('(b) Cluster SOC at 0% for 120 consecutive hours;')
bullet('(c) Discharge cell voltage \u22642.5V (any occurrence);')
bullet('(d) SOC below 5% for more than 7 consecutive days.')
para('The above conditions shall NOT void the warranty where the low SOC or low voltage condition was caused by:', bold=True)
bullet('(i) a grid fault, outage, or curtailment instruction from DSO;', 1)
bullet('(ii) an EMS malfunction not attributable to the Client;', 1)
bullet('(iii) any Force Majeure event; or', 1)
bullet('(iv) failure by the Contractor to provide a monitoring alert under an LTSA with 24/7 monitoring (Tier A or above), where such alert would have enabled the Client to prevent the condition.', 1)
para('Where the Client has entered into an LTSA with 24/7 monitoring (Tier A+), the Contractor shall alert the Client within 24 hours of SOC dropping below 10%.')

h3('10.6 Performance Guarantees')
para('The Contractor guarantees (backed by OEM warranty):')
para('(a) State of Health (SOH) at 1 cycle per day: Year 5 \u226585%, Year 10 \u226579.58%, Year 15 \u226570% (Year 15 subject to extended warranty under LTSA).', bold=True)
para('(b) Round-Trip Efficiency: \u226586.32% at 0.5C, 25\u00b12\u00b0C. PAC field floor per Section 9.1.', bold=True)
para('(c) Cycle Life: 7,000 cycles at 0.5C, 90% DoD, to 70% EOL.', bold=True)
para('(d) Remedy: Contractor procures replacement modules from OEM at Contractor\u2019s cost.')

h3('10.7 Availability Guarantee')
para('97% availability and associated LDs apply only under a separately executed LTSA (Tier C+).')

h3('10.8 OEM Direct Warranty Undertaking')
para('The OEM has provided a Direct Warranty Undertaking (Schedule C) confirming its standard 5-year product warranty is enforceable directly by the Client if the Contractor cannot fulfil obligations for any reason, including insolvency.')

h3('10.9 Advance Payment Guarantee (APG) and Performance Bond')
para('(a) The Contractor shall procure that the OEM provides an unconditional, irrevocable Advance Payment Guarantee (APG) equal to the sum of the advance payment and pre-shipment payment made under Sections 7.1(a) and 7.1(b) in respect of Component A (Equipment Supply Price) only.', bold=True)
para('(b) The APG covers: 20% advance on Component A + 50% pre-shipment on Component A = 70% of the Component A Equipment Supply Price.')
para('(c) The APG shall be delivered prior to the advance payment becoming due (condition precedent per Section 9B.4 of the upstream supply agreement).')
para('(d) The APG shall remain valid until equipment delivery to Site plus thirty (30) days.')
para('(e) No APG is required against Component B (EPC Services Price).')
para('(f) The Contractor shall also procure that the OEM provides a performance bond equal to five percent (5%) of the Component A Equipment Supply Price, valid until the end of the Defects Liability Period (three (3) months after PAC), covering defects in materials and workmanship.')

h2('11. LTSA SEPARATION')
para('11.1 97% availability guarantee, availability LDs, and extended warranty (years 6\u201315) are not in this Agreement.')
para('11.2 Apply only under a separately executed LTSA.')
para('11.3 SOH, RTE, and cycle life guarantees (Section 10.6) are in this Agreement, independent of the LTSA.')
para('11.4 This Agreement is intended to be executed alongside the LTSA (LCY-LTSA-001 v4.0) and the OEM Direct Warranty Undertaking (LCY-OEM-DWU-001 v1.0). The Parties intend to execute all three documents simultaneously.')

h2('12. FORCE MAJEURE')
h3('12.1 Definition')
para('Force Majeure means events genuinely beyond a Party\u2019s reasonable control: war, terrorism, riots, natural disasters, pandemics, epidemics, governmental sanctions, embargoes, and acts of government.')
para('Force Majeure does NOT include: regulatory changes, grid connection delays, port congestion, transportation disruptions, financing difficulties, or ordinary commercial risks. Grid authority delays are addressed under Section 8.4.6.')
h3('12.2 Effect')
para('Obligations suspended without penalties for duration of event.')
h3('12.3 Notification')
para('Prompt notice with event details, expected duration, and mitigation steps.')
h3('12.4 Mitigation')
para('Commercially reasonable efforts to resume performance.')
h3('12.5 Extended Force Majeure')
para('If a Force Majeure event persists for more than six (6) months, either Party may terminate upon thirty (30) days\u2019 written notice. In such event:')
bullet('(a) The Client shall pay the value of works completed and materials delivered as at the date of termination;')
bullet('(b) If payments received by the Contractor exceed the value of works completed and materials delivered, the Contractor shall refund the excess to the Client within thirty (30) days of termination;')
bullet('(c) Valuation shall follow the same process as Section 8.4.5(b).')

h2('13. LIMITATION OF LIABILITY')
para('13.1 Neither Party shall be liable for indirect or consequential damages, including loss of profits, loss of revenue, or loss of business.')
para('13.2 The Contractor\u2019s total aggregate liability under this Agreement shall be subject to the following tiers, each calculated as a percentage of the Contract Price:', bold=True)
bullet('(a) Warranty claims (defects in materials, workmanship, or OEM equipment): ten percent (10%) of the Contract Price;')
bullet('(b) All other contractual breaches (excluding warranty, fraud, wilful misconduct, and manufacturing defects): fifty percent (50%) of the Contract Price;')
bullet('(c) Fraud, wilful misconduct, and manufacturing defects in OEM equipment: uncapped.')
para('13.3 Carve-Outs (not subject to any cap):')
bullet('(a) Fraud or wilful misconduct;')
bullet('(b) Death or personal injury caused by negligence;')
bullet('(c) Client\u2019s payment obligations;')
bullet('(d) Manufacturing defects in OEM equipment \u2014 Contractor pursues OEM recovery; OEM indemnification flows through to Client per §13A.1(e).')
para('13.4 The tiered structure in Section 13.2 mirrors the upstream liability framework agreed with the OEM. For manufacturing defects specifically, the OEM bears full replacement cost including shipping under its supply agreement with the Contractor, with no cap applied.')

h2('13A. INDEMNIFICATION')
h3('13A.1 Contractor Indemnity to Client')
para('The Contractor indemnifies the Client against:')
bullet('(a) Defects in the BESS (including OEM equipment) during Warranty Period;')
bullet('(b) Contractor\u2019s negligence during installation and commissioning;')
bullet('(c) Breach of Contractor\u2019s warranties;')
bullet('(d) Personal injury or property damage caused by Contractor or OEM personnel on-site;')
bullet('(e) Manufacturing defects in OEM equipment (recoverable from OEM under supply agreement).')

h3('13A.2 Client Indemnity to Contractor')
para('The Client indemnifies the Contractor against claims arising from:')
bullet('(a) Client\u2019s material breach of this Agreement;')
bullet('(b) Client\u2019s negligence or wilful acts;')
bullet('(c) Operation contrary to OEM guidelines;')
bullet('(d) Unauthorised modifications;')
bullet('(e) Failure to maintain insurance.')
para('The Client\u2019s indemnity obligations under this Section apply only to the extent that the relevant loss was caused or contributed to by the Client\u2019s act or omission. Where fault is shared, the Client\u2019s indemnity shall be reduced proportionally to reflect the Client\u2019s degree of responsibility.', bold=True)

h2('14. INSURANCE')
h3('14.1 Contractor Insurance')
bullet('(a) Public Liability: EUR 5,000,000 per occurrence;')
bullet('(b) Professional Indemnity: EUR 2,000,000 per occurrence;')
bullet('(c) CAR Insurance: full Contract Price, LEG3 defects clause;')
bullet('(d) Employers\u2019 Liability: statutory minimum.')
h3('14.2 OEM Insurance')
para('OEM maintains Product Liability EUR 5,000,000/occurrence (AXA Tianping) and PI EUR 2,000,000/occurrence for Cyprus.')
h3('14.3 Client Insurance')
bullet('(a) Site/property insurance from delivery;')
bullet('(b) Operational liability from commissioning;')
bullet('(c) Business interruption (recommended).')
h3('14.4 Proof of Insurance')
para('The Contractor shall provide insurance certificates or cover notes within fourteen (14) days of the first payment under Section 7.1(a).', bold=True)

h2('15. GOVERNING LAW AND JURISDICTION')
para('15.1 Laws of the Republic of Cyprus.')
para('15.2 Courts of Cyprus have exclusive jurisdiction.')
para('15.3 ADR: 30-day negotiation before proceedings.')
para('15.4 English is the official language.')

h2('16. MISCELLANEOUS')
para('16.1 Entire Agreement.')
para('16.2 Amendments: in writing, signed by authorised representatives.')
h3('16.3 Assignment')
para('(a) No assignment without consent, except:')
bullet('(i) to an Affiliate (corporate reorganisation); or', 1)
bullet('(ii) by the Client to a project finance lender for security purposes \u2014 no Contractor consent required.', 1)
para('(b) Upon lender enforcement, Contractor recognises the lender as Client if payment obligations are assumed.')
para('16.4 Counterparts and electronic signature.')
para('16.5 Severability.')
para('16.6 Notices: in writing to specified addresses.')

h2('17. COMPLIANCE AND ANTI-BRIBERY')
para('[Sections 17.1\u201317.6 as per v4.0 \u2014 anti-corruption, AML, environmental, H&S, Cyprus energy regulations, EU directives.]')
h2('18. DATA PROTECTION')
para('[Sections 18.1\u201318.4 as per v4.0 \u2014 GDPR, data processing, security, breach notification.]')
h2('19. REGULATORY COMPLIANCE')
para('[Sections 19.1\u201319.11 as per v4.0 \u2014 Cyprus laws, building permit, site requirements, fire safety, environmental, certifications, grid connection, SCADA, electrical studies, DSO commissioning.]')
para('19.11 Responsibilities:')
para('(a) Client: Planning permits, landowner consents, environmental permits.')
para('(b) Contractor: All technical and regulatory responsibilities including EAC, DSO, ETEK, engineering, grid connection, and OEM coordination.')

# SCHEDULES
doc.add_page_break()
h2('SCHEDULE A \u2014 TECHNICAL SPECIFICATIONS AND PROJECT DETAILS')
para('[Linyang Power Atlantic 5MWh specifications per v4.0. Project-specific details to be completed:')
bullet('Park name(s):')
bullet('Total BESS capacity (MWh):')
bullet('Number of containers:')
bullet('MV skid configuration:')
bullet('PCS model and quantity:')
bullet('Coastal distance (km):')
bullet('Target PAC Date:')
bullet('Component A \u2014 Equipment Supply Price (EUR):')
bullet('Component B \u2014 EPC Services Price (EUR):')
bullet('Contract Price (A + B) (EUR):')

doc.add_page_break()
h2('SCHEDULE B \u2014 DELAY LD WORKED EXAMPLES')
para('Example 1 \u2014 Full Park Delay:', bold=True)
para('Park Contract Price: EUR 2,000,000. PCS delayed 45 days.')
para('Days 1\u201330: 30 \u00d7 0.1% \u00d7 EUR 2,000,000 = EUR 60,000')
para('Days 31\u201345: 15 \u00d7 0.15% \u00d7 EUR 2,000,000 = EUR 45,000')
para('Total: EUR 105,000 (5.25% of Contract Price)')
para('')
para('Example 2 \u2014 Partial Delay:', bold=True)
para('4-container Park (20 MWh, EUR 2,000,000). 1 container (25%) delayed 20 days.')
para('LD Value: EUR 2,000,000 \u00d7 25% = EUR 500,000')
para('LDs: 20 \u00d7 0.1% \u00d7 EUR 500,000 = EUR 10,000')
para('Remaining 3 containers proceed to PAC without LD liability.')
para('')
para('Example 3 \u2014 Escalation:', bold=True)
para('If delayed container prevents PCS minimum load, escalates to Full Park from that date.')

doc.add_page_break()
h2('SCHEDULE C \u2014 OEM DIRECT WARRANTY UNDERTAKING')
para('[Annexed: OEM-Direct-Warranty-Undertaking-Linyang.docx \u2014 to be executed by Linyang, Lighthief Cyprus Ltd, and Client simultaneously with this Agreement.]')
para('The OEM Direct Warranty Undertaking confirms:')
bullet('(a) Linyang\u2019s 5-year product warranty applies to all equipment supplied for this Project through the Distributor;')
bullet('(b) The Client may enforce the warranty directly against Linyang if the Contractor cannot perform;')
bullet('(c) The undertaking survives termination of any agreement between Linyang and Lighthief;')
bullet('(d) Governing law: Singapore (SIAC arbitration, Cyprus carve-out for interim relief).')

# SIGNATURES
doc.add_page_break()
h2('SIGNATURES')
para('')
para('For and on behalf of Lighthief Cyprus Ltd (\u201cContractor\u201d):')
para('')
para('Name: __________________________')
para('Title: __________________________')
para('Signature: _____________________')
para('Date: __________________________')
para('')
para('')
para('For and on behalf of [\u25cf End Customer] (\u201cClient\u201d):')
para('')
para('Name: __________________________')
para('Title: __________________________')
para('Signature: _____________________')
para('Date: __________________________')

doc.save('client_sales_v5.0.docx')
print(f'Saved: client_sales_v5.0.docx ({len(doc.paragraphs)} paragraphs)')
