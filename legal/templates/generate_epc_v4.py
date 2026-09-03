#!/usr/bin/env python3
"""Generate EPC Agreement v4.0 DOCX — Bankability Revision"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BLUE = RGBColor(0x1A, 0x36, 0x5D)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BLUE

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(16)

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(10)

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(10)

def bullet(text, level=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)

# TITLE
doc.add_paragraph()
h1('ENGINEERING, PROCUREMENT & CONSTRUCTION AGREEMENT')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Battery Energy Storage System (BESS)')
r.font.size = Pt(12); r.font.color.rgb = BLUE
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    'Document Reference: LCY-EPC-001',
    'Version: 4.0',
    'Date: 17 March 2026',
    '',
    'VERSION HISTORY',
    'v1.0 — Oct 2025: Initial draft',
    'v2.0 — Jan 2026: Updated to Linyang RFI confirmed terms',
    'v2.5 — Feb 2026: Updated warranty, LDs, regulatory compliance',
    'v3.0 — 22 Feb 2026: Installation scope split, delay LD framework, bankability',
    'v4.0 — 17 Mar 2026: BANKABILITY REVISION — single counterparty warranty,',
    '       title at PAC, unified 5-year warranty, SOH/RTE/cycle guarantees,',
    '       liability cap 15%, lender assignment, OEM step-in, insurance at',
    '       first payment, connection terms CP, narrowed FM, RTE field tolerance',
]:
    r = meta.add_run(line + '\n')
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    if 'v4.0' in line:
        r.bold = True

doc.add_page_break()

# 1. PARTIES
h2('1. PARTIES')
para('This Engineering, Procurement and Construction Agreement ("Agreement") is made as of [\u25cf] 2026 ("Effective Date")')
para('BETWEEN:')
para('Lighthief Cyprus Ltd, a company incorporated under the laws of the Republic of Cyprus, with registered office at Agiou Andreou 241, AG TRIAS COURT, Flat/Office 31, 3036 Limassol, Cyprus, and operational address at 15 Agaritsis, Nektaria Court, Office 201, 3045 Zakaki, Limassol, Cyprus, Company No. HE 477423, TIN 60187188Q ("Contractor");')
para('and')
para('[End Customer Legal Name], a company incorporated under the laws of [\u25cf], with registered office at [\u25cf] ("Client").')
para('Contractor and Client are referred to individually as a "Party" and collectively as the "Parties."')

# 1A. CONDITIONS PRECEDENT
h2('1A. CONDITIONS PRECEDENT')
para('1A.1 This Agreement is conditional upon the Client obtaining grid connection terms from EAC/DSO for the Site ("Connection Terms") within [\u25cf] months of the Effective Date ("Long-Stop Date").')
para('1A.2 If Connection Terms are not obtained by the Long-Stop Date, either Party may terminate this Agreement by written notice without liability. The Contractor shall refund the advance payment within thirty (30) days, less reasonable and documented costs incurred.')
para('1A.3 The advance payment under Section 7.1(a) shall become due within seven (7) days of the later of: (i) the Effective Date; and (ii) receipt of Connection Terms.')

# 2. BACKGROUND
h2('2. BACKGROUND')
para('2.1 The Client owns and operates one or more renewable energy facilities and wishes to procure a battery energy storage system.')
para('2.2 The Contractor has the capability to design, integrate, install, commission, and deliver battery energy storage projects on a turnkey EPC basis.')
para('2.3 The Parties wish to set out the terms under which the Contractor shall deliver the BESS to the Client.')

# 3. DEFINITIONS
h2('3. DEFINITIONS AND INTERPRETATION')
h3('3.1 Defined Terms')
defs = [
    ('"Agreement"', 'means this EPC Agreement, including all schedules, appendices, and amendments.'),
    ('"BESS"', 'means the battery energy storage system supplied under this Agreement, comprising the Linyang Power Atlantic LFP liquid-cooled battery container system, including battery packs, BMS, thermal management, fire suppression, and related components as described in Schedule A.'),
    ('"Business Day"', 'means a day other than a Saturday, Sunday, or public holiday in the Republic of Cyprus.'),
    ('"Connection Terms"', 'means the grid connection terms issued by EAC/DSO for the Site.'),
    ('"Contract Price"', 'has the meaning given in Section 6.1.'),
    ('"Delay LDs"', 'means the liquidated damages payable by the Contractor for delay per Section 8.4.'),
    ('"Effective Date"', 'has the meaning given in Section 1.'),
    ('"EMS"', 'means Energy Management System. EMS integration is provided by Lighthief EU BESS Ltd under a separate EMS Integration Agreement.'),
    ('"FAC"', 'means Final Acceptance Certificate per Section 9.2.'),
    ('"Force Majeure"', 'has the meaning given in Section 12.1.'),
    ('"LTSA"', 'means a Long-Term Service Agreement entered into separately between the Parties.'),
    ('"OEM"', 'means Original Equipment Manufacturer (Linyang Energy).'),
    ('"OEM Direct Warranty Undertaking"', 'means the direct warranty confirmation letter issued by the OEM to the Client (Schedule C).'),
    ('"PAC"', 'means Provisional Acceptance Certificate per Section 9.1.'),
    ('"Products"', 'means BESS, PCS, battery containers, control equipment, and related components.'),
    ('"Retention"', 'has the meaning given in Section 7.4.'),
    ('"Site"', 'means the installation location specified in Schedule A.'),
    ('"Target PAC Date"', 'means the date by which PAC is to be achieved per Section 8.4.'),
    ('"Warranty Period"', 'means five (5) years from the Warranty Start Date.'),
    ('"Warranty Start Date"', 'means the date of PAC; provided that if PAC is delayed beyond six (6) months from equipment delivery to Site due to causes not attributable to the Client, the Warranty Start Date shall be the date six (6) months after delivery.'),
]
for term, defn in defs:
    p = doc.add_paragraph()
    r1 = p.add_run(term + ' ')
    r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(defn)
    r2.font.size = Pt(10)

h3('3.2 Interpretation')
for t in ['(a) Headings are for convenience only.', '(b) "Including" means without limitation.', '(c) References to laws include amendments.', '(d) Official language is English.']:
    bullet(t)

# 4. SCOPE — SINGLE COUNTERPARTY
h2('4. SCOPE OF WORK')
para('4.1 Turnkey EPC Solution: The Contractor shall provide a complete turnkey EPC solution for the design, procurement, delivery, installation, commissioning, and handover of the BESS. The Contractor is the Client\u2019s sole counterparty for all works under this Agreement.', bold=True)

h3('4.2 Contractor Scope')
para('The Contractor shall be responsible for, and warrants the performance of:')
for item in [
    '(a) Project management and coordination;',
    '(b) Detailed engineering and system integration design;',
    '(c) Procurement of BESS equipment from the OEM;',
    '(d) International logistics, import duties, customs clearance, and port-to-site transport;',
    '(e) Civil works, foundations, fencing, and site preparation;',
    '(f) AC/DC cabling, grounding, and auxiliary electrical systems;',
    '(g) Equipment installation, assembly, and commissioning (performed by OEM-trained personnel under the Contractor\u2019s coordination and responsibility);',
    '(h) SCADA integration and grid compliance coordination;',
    '(i) Authority liaison, EAC submissions, and regulatory compliance;',
    '(j) Project documentation, training, and handover.',
]:
    bullet(item)

h3('4.3 Single-Point Responsibility')
para('(a) The Contractor is responsible to the Client for the entire BESS system, including OEM equipment and installation workmanship. The Client\u2019s sole recourse for any defect, delay, or performance issue is against the Contractor.')
para('(b) Subcontracting of installation or commissioning to OEM personnel does not relieve the Contractor of its obligations.')
para('(c) The Contractor warrants the performance of OEM equipment as if it were the Contractor\u2019s own work.')

h3('4.4 EMS Integration')
para('(a) EMS integration and SCADA commissioning are provided by Lighthief EU BESS Ltd, an affiliate of the Contractor, under a separate EMS Integration Agreement.')
para('(b) The Contractor warrants BMS-to-EMS interface compatibility at commissioning.')
para('(c) Advanced EMS features, analytics, and optimisation are contracted separately.')
para('(d) The Contractor is not responsible for EMS optimisation outcomes or revenue generation.')
para('4.5 Any works not expressly listed are excluded unless agreed in writing.')

# 5A. CLIENT OBLIGATIONS
h2('5A. CLIENT OBLIGATIONS')
h3('5A.1 Pre-Commencement')
para('The Client shall:')
for item in ['(a) Provide complete and accurate site information;', '(b) Obtain planning permissions, building permits, and landowner consents;', '(c) Ensure clear site access for delivery and installation;', '(d) Provide secure storage for equipment;', '(e) Provide electrical connection point and grid connection approval.']:
    bullet(item)
h3('5A.2 During Works')
para('The Client shall:')
for item in ['(a) Provide site access during agreed working hours;', '(b) Designate a single point of contact;', '(c) Respond to queries within five (5) Business Days;', '(d) Provide timely approvals;', '(e) Not materially interfere with works.']:
    bullet(item)
h3('5A.3 Post-Commissioning')
para('The Client shall:')
for item in ['(a) Operate BESS per OEM guidelines;', '(b) Maintain monitoring systems;', '(c) Maintain site security and environmental conditions;', '(d) Maintain insurance;', '(e) Make all payments when due;', '(f) Permit Contractor access for warranty/maintenance.']:
    bullet(item)
h3('5A.4 Consequences of Client Breach')
para('(a) Delay: If Contractor\u2019s works are delayed more than five (5) Business Days by Client\u2019s material failure, Target PAC Date extends day-for-day and Client reimburses reasonable documented costs.')
para('(b) Payment Default:')
bullet('(i) Interest at EU Directive 2011/7/EU rate (ECB + 8% pa, simple);', 1)
bullet('(ii) Contractor may suspend works on seven (7) days\u2019 notice;', 1)
bullet('(iii) If overdue more than thirty (30) days, Contractor may terminate and retain payments received.', 1)
para('(c) Warranty: Client breach of Sections 10.4 or 10.5 may limit warranty to the extent the breach caused or contributed to the defect.')
para('(d) No Contractor Liability for delays or losses arising from Client\u2019s material failure to fulfil obligations.')

# 6. CONTRACT PRICE
h2('6. CONTRACT PRICE')
para('6.1 The total contract price is EUR [\u25cf] plus VAT ("Contract Price").')
para('6.2 Based on current technical, regulatory, customs, and tax conditions.')
para('6.3 Material changes in scope, regulations, or law entitle the Parties to agree a variation by written change order.')

# 7. PAYMENT — TITLE AT PAC
h2('7. PAYMENT TERMS')
h3('7.1 Payment Milestones')
para('(a) 30% advance within seven (7) days of the payment trigger under Section 1A.3;')
para('(b) 55% prior to shipment, upon written FAT confirmation and photographic evidence;')
para('(c) 10% upon issuance of PAC;')
para('(d) 5% retention released upon issuance of FAC.')
para('7.2 Payments exclusive of VAT.')
para('7.3 Late payments: Contractor may suspend on seven (7) days\u2019 notice. Interest per Section 5A.4(b)(i).')
h3('7.4 Retention')
para('(a) 5% retained as security for punch-list and defect rectification.')
para('(b) Released within thirty (30) days of FAC.')
para('(c) No set-off except costs to remedy defects Contractor failed to rectify after notice.')
para('(d) If FAC not issued within twelve (12) months of PAC due to minor items only, Retention released less amounts for outstanding works.')
h3('7.5 TITLE AND RISK TRANSFER')
para('(a) Risk passes to Client upon delivery to Site and unloading.')
para('(b) Title passes to Client upon issuance of PAC and receipt of PAC payment under Section 7.1(c). The 5% Retention does not defer title.', bold=True)
para('(c) Defects Liability Undertaking: The Contractor undertakes to rectify all defects during the 24-month Defects Liability Period at its own cost.')
para('(d) Lender Security: Upon title transfer, the Client may grant security interests over the Equipment to lenders without Contractor consent.')

# 8. DELIVERY & DELAY LDs
h2('8. DELIVERY AND INSTALLATION')
para('8.1 CIF Limassol, Cyprus (Incoterms\u00ae 2020).')
para('8.2 Installation commences upon site readiness and delivery.')
para('8.3 Typical installation period: 4\u20136 weeks.')
h3('8.4 Delay Liquidated Damages')
h3('8.4.1 Definitions')
para('(a) "Target PAC Date" \u2014 agreed PAC date per Schedule A.')
para('(b) "Critical Path Equipment" \u2014 PCS, MV switchgear/skid, main step-up transformer. Delay = Full Park Delay.')
para('(c) "Non-Critical Path Equipment" \u2014 individual battery containers where PCS/MV can commission remaining units. Delay = Partial Delay.')
para('(d) "Full Park Delay" \u2014 no portion of Park can commission.')
para('(e) "Partial Delay" \u2014 reduced capacity commissioning possible.')
para('(f) "Affected Capacity" \u2014 delayed MWh as % of total MWh.')
para('(g) "LD Value" \u2014 Full Park: total Contract Price. Partial: Contract Price \u00d7 Affected Capacity %.')
h3('8.4.2 Obligation')
para('Contractor shall achieve PAC by Target PAC Date.')
h3('8.4.3 Calculation')
para('(a) Delay LDs on LD Value:')
bullet('(i) Days 1\u201330: 0.1% per day;', 1)
bullet('(ii) Days 31\u201360: 0.15% per day;', 1)
bullet('(iii) Day 61+: 0.2% per day.', 1)
para('(b) Full Park: LDs on entire Contract Price. Partial: LDs on proportional value only. Escalation: recalculates if Partial becomes Full.')
h3('8.4.4 Cap')
para('10% of Contract Price per Park (aggregate of all Full + Partial LDs).')
h3('8.4.5 Remedies')
para('Delay LDs are the Client\u2019s primary remedy. No consequential damages for delay. If PAC not achieved within 180 days of Target PAC Date, Client may terminate and Contractor refunds payments less value of completed works.')
h3('8.4.6 Exclusions')
para('LDs not applicable for: (i) Force Majeure; (ii) Client-caused delays; (iii) grid authority delays; (iv) Client change orders; (v) customs delays with timely shipment proved; (vi) Client\u2019s late confirmation of technical agreements or drawings.')
h3('8.4.7 Extension of Time')
para('Target PAC Date extends day-for-day. Contractor notifies within seven (7) Business Days.')
h3('8.4.8 Worked Examples')
para('See Schedule B for Full Park Delay, Partial Delay, and Escalation worked examples.')

# 9. ACCEPTANCE — PAC CRITERIA
h2('9. ACCEPTANCE')
h3('9.1 Provisional Acceptance (PAC)')
para('PAC issued when:')
bullet('(a) BESS operates at \u226595% of rated capacity;')
bullet('(b) RTE \u226584% under site conditions (86.32% OEM reference less 2.32pt field tolerance);')
bullet('(c) BMS/SCADA communications verified;')
bullet('(d) All safety/protection systems functional;')
bullet('(e) Mandatory commissioning tests per Schedule A passed.')
para('Minor RTE deviations from cable losses, ambient temperature >25\u00b0C, or grid voltage fluctuations shall not prevent PAC if within 3 percentage points of 86.32%.')
h3('9.2 Final Acceptance (FAC)')
para('FAC issued 24 months after PAC upon resolution of punch-list items.')
h3('9.3 Minor Defects')
para('Minor cosmetic defects not materially affecting operation do not prevent PAC if recorded on punch list.')

# 10. WARRANTY — UNIFIED 5 YEARS
h2('10. WARRANTY')
h3('10.1 Unified Warranty')
para('The Contractor warrants the entire BESS \u2014 including all OEM equipment, installation workmanship, and EPC works \u2014 against defects in design, materials, and workmanship for the Warranty Period. The Contractor is the Client\u2019s sole warranty counterparty.', bold=True)
h3('10.2 Coastal Installation')
bullet('(a) C5-rated enclosures as standard for Cyprus.')
bullet('(b) >500m from sea: full 5-year warranty.')
bullet('(c) <500m: PCS/transformer/switchgear warranty reduced to 2 years.')
bullet('(d) Client confirms coastal distance at signing (Schedule A).')
h3('10.3 Exclusions')
para('Warranty excludes: (a) consumables; (b) misuse/negligence; (c) unauthorised modifications; (d) Force Majeure damage; (e) relocation without approval.')
h3('10.4 Conditions')
para('Subject to: (a) operation per OEM guidelines; (b) qualified maintenance; (c) serial numbers intact; (d) access for inspection; (e) full payment.')
h3('10.5 WARRANTY-VOIDING BATTERY CONDITIONS')
para('The following void the battery warranty:')
bullet('(a) Cell voltage <2.8V for 120 consecutive hours;')
bullet('(b) Cluster SOC at 0% for 120 consecutive hours;')
bullet('(c) Discharge cell voltage \u22642.5V (any occurrence);')
bullet('(d) SOC <5% for more than 7 consecutive days.')
para('Where an LTSA with 24/7 monitoring (Tier A+) is in place, the Contractor shall alert the Client within 24 hours of SOC dropping below 10%. If the Contractor fails to alert and a voiding condition occurs, the voiding shall not apply to the extent caused by the Contractor\u2019s failure.')
h3('10.6 Performance Guarantees')
para('The Contractor guarantees (backed by OEM warranty):')
para('(a) SOH at 1 cycle/day: Year 5 \u226585%, Year 10 \u226579.58%, Year 15 \u226570% (Year 15 subject to extended warranty under LTSA).', bold=True)
para('(b) RTE: \u226586.32% at 0.5C, 25\u00b12\u00b0C. Field tolerance per Section 9.1.', bold=True)
para('(c) Cycle Life: 7,000 cycles at 0.5C, 90% DoD, to 70% EOL.', bold=True)
para('(d) Remedy: Contractor procures replacement modules from OEM at Contractor\u2019s cost.')
h3('10.7 Availability Guarantee')
para('97% availability and associated LDs apply only under a separately executed LTSA (Tier C+).')
h3('10.8 OEM Direct Warranty Undertaking')
para('The OEM has provided a Direct Warranty Undertaking (Schedule C) confirming its standard 5-year warranty is enforceable directly by the Client if the Contractor cannot fulfil obligations for any reason, including insolvency.')
h3('10.9 Performance Bond')
para('(a) 5% of Contract Price, corporate guarantee backed by bank.')
para('(b) Valid until FAC (24 months post-PAC).')
para('(c) Covers defects in materials, workmanship, and non-performance.')

# 11. LTSA SEPARATION
h2('11. LTSA SEPARATION')
para('11.1 97% availability guarantee, availability LDs, and extended warranty (years 6\u201315) are not in this Agreement.')
para('11.2 Apply only under a separately executed LTSA.')
para('11.3 SOH, RTE, and cycle life guarantees (Section 10.6) are in this Agreement, independent of the LTSA.')

# 12. FORCE MAJEURE — NARROWED
h2('12. FORCE MAJEURE')
h3('12.1 Definition')
para('Force Majeure means events genuinely beyond a Party\u2019s reasonable control: war, terrorism, riots, natural disasters, pandemics, epidemics, governmental sanctions, embargoes, and acts of government.')
para('Force Majeure does NOT include: regulatory changes, grid connection delays, port congestion, transportation disruptions, financing difficulties, or ordinary commercial risks. These are addressed under Section 8.4.6 (Extension of Time).')
h3('12.2 Effect')
para('Obligations suspended without penalties for duration of event.')
h3('12.3 Notification')
para('Prompt notice with event details, expected duration, and mitigation steps.')
h3('12.4 Mitigation')
para('Commercially reasonable efforts to resume performance.')
h3('12.5 Extended Force Majeure')
para('If FM persists >6 months, either Party may terminate on 30 days\u2019 notice. Client pays for completed works and delivered materials.')

# 13. LIABILITY — 15% CAP
h2('13. LIMITATION OF LIABILITY')
para('13.1 Neither Party liable for indirect/consequential damages.')
para('13.2 Contractor\u2019s total aggregate liability shall not exceed fifteen percent (15%) of the Contract Price.', bold=True)
para('13.3 Carve-Outs (not subject to cap):')
bullet('(a) Fraud or wilful misconduct;')
bullet('(b) Death or personal injury;')
bullet('(c) Client\u2019s payment obligations;')
bullet('(d) Manufacturing defects in OEM equipment \u2014 Contractor pursues OEM recovery; OEM indemnification flows through to Client.')

# 13A. INDEMNIFICATION
h2('13A. INDEMNIFICATION')
h3('13A.1 Contractor Indemnity to Client')
para('The Contractor indemnifies the Client against:')
bullet('(a) Defects in the BESS (including OEM equipment) during Warranty Period;')
bullet('(b) Contractor\u2019s negligence during installation/commissioning;')
bullet('(c) Breach of Contractor\u2019s warranties;')
bullet('(d) Personal injury/property damage by Contractor/OEM personnel on-site;')
bullet('(e) Manufacturing defects (recoverable from OEM).')
h3('13A.2 Client Indemnity to Contractor')
para('The Client indemnifies the Contractor against:')
bullet('(a) Client\u2019s material breach;')
bullet('(b) Client\u2019s negligence;')
bullet('(c) Operation contrary to OEM guidelines;')
bullet('(d) Unauthorised modifications;')
bullet('(e) Failure to maintain insurance.')

# 14. INSURANCE
h2('14. INSURANCE')
h3('14.1 Contractor Insurance')
bullet('(a) Public Liability: EUR 5,000,000 per occurrence;')
bullet('(b) Professional Indemnity: EUR 2,000,000 per occurrence;')
bullet('(c) CAR Insurance: full Contract Price, with LEG3 defects clause;')
bullet('(d) Employers\u2019 Liability: statutory minimum.')
h3('14.2 OEM Insurance')
para('OEM maintains Product Liability EUR 5,000,000/occurrence (AXA Tianping) and PI EUR 2,000,000/occurrence for Cyprus.')
h3('14.3 Client Insurance')
bullet('(a) Site/property insurance from delivery;')
bullet('(b) Operational liability from commissioning;')
bullet('(c) Business interruption (recommended).')
h3('14.4 Proof of Insurance')
para('The Contractor shall provide insurance certificates or cover notes within fourteen (14) days of the first payment under Section 7.1(a).', bold=True)

# 15. GOVERNING LAW
h2('15. GOVERNING LAW AND JURISDICTION')
para('15.1 Laws of the Republic of Cyprus.')
para('15.2 Courts of Cyprus have exclusive jurisdiction.')
para('15.3 ADR: 30-day negotiation before proceedings.')
para('15.4 English is the official language.')

# 16. MISCELLANEOUS — LENDER ASSIGNMENT
h2('16. MISCELLANEOUS')
para('16.1 Entire Agreement.')
para('16.2 Amendments: in writing, signed by authorised representatives.')
h3('16.3 Assignment')
para('(a) No assignment without consent, except:')
bullet('(i) to an Affiliate (corporate reorganisation); or', 1)
bullet('(ii) by the Client to a project finance lender for security purposes \u2014 no Contractor consent required.', 1)
para('(b) Upon lender enforcement, Contractor recognises the lender as Client if payment obligations are assumed.')
para('16.4 Counterparts and electronic signature.')
para('16.5 Severability.')
para('16.6 Notices: in writing to specified addresses.')

# 17-19 COMPLIANCE
h2('17. COMPLIANCE AND ANTI-BRIBERY')
para('[Sections 17.1\u201317.6 as per v3.0 \u2014 anti-corruption, AML, environmental, H&S, Cyprus energy regulations, EU directives.]')
h2('18. DATA PROTECTION')
para('[Sections 18.1\u201318.4 as per v3.0 \u2014 GDPR, data processing, security, breach notification.]')
h2('19. REGULATORY COMPLIANCE')
para('[Sections 19.1\u201319.11 as per v3.0 \u2014 Cyprus laws, building permit, site requirements, fire safety, environmental, certifications, grid connection, SCADA, electrical studies, DSO commissioning.]')
para('19.11 Responsibilities (updated for single-counterparty):')
para('(a) Client: Planning permits, landowner consents, environmental permits.')
para('(b) Contractor: All technical and regulatory responsibilities including EAC, DSO, ETEK, engineering, grid connection, and OEM coordination.')

# SCHEDULES
doc.add_page_break()
h2('SCHEDULE A \u2014 TECHNICAL SPECIFICATIONS')
para('[Linyang Power Atlantic 5MWh specifications per v3.0. Project-specific: park names, MW, MWh, containers, coastal distance, Target PAC Date, Contract Price.]')

doc.add_page_break()
h2('SCHEDULE B \u2014 DELAY LD WORKED EXAMPLES')
para('Example 1 \u2014 Full Park Delay:', bold=True)
para('Park Contract Price: EUR 2,000,000. PCS delayed 45 days.')
para('Days 1\u201330: 30 \u00d7 0.1% \u00d7 EUR 2,000,000 = EUR 60,000')
para('Days 31\u201345: 15 \u00d7 0.15% \u00d7 EUR 2,000,000 = EUR 45,000')
para('Total: EUR 105,000 (5.25% of Contract Price)')
para('')
para('Example 2 \u2014 Partial Delay:', bold=True)
para('4-container Park (20 MWh, EUR 2,000,000). 1 container (25%) delayed 20 days.')
para('LD Value: EUR 2,000,000 \u00d7 25% = EUR 500,000')
para('LDs: 20 \u00d7 0.1% \u00d7 EUR 500,000 = EUR 10,000')
para('Remaining 3 containers proceed to PAC without LD liability.')
para('')
para('Example 3 \u2014 Escalation:', bold=True)
para('If delayed container prevents PCS minimum load, escalates to Full Park from that date.')

doc.add_page_break()
h2('SCHEDULE C \u2014 OEM DIRECT WARRANTY UNDERTAKING')
para('[Annexed separately: OEM-Direct-Warranty-Undertaking-Linyang.docx]')

# SIGNATURES
doc.add_page_break()
h2('SIGNATURES')
para('')
para('For and on behalf of Lighthief Cyprus Ltd ("Contractor"):')
para('')
para('Name: __________________________')
para('Title: __________________________')
para('Signature: _____________________')
para('Date: __________________________')
para('')
para('')
para('For and on behalf of [End Customer] ("Client"):')
para('')
para('Name: __________________________')
para('Title: __________________________')
para('Signature: _____________________')
para('Date: __________________________')

doc.save('client_sales_v4.0.docx')
print(f'Saved: client_sales_v4.0.docx ({len(doc.paragraphs)} paragraphs)')
