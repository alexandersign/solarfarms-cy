"""
Generate PV Residential Contract - Legally corrected DOCX with Lighthief branding.
Fixes all issues identified in legal review of 21.04.2026 template.
"""

import io
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colours ──────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1A, 0x36, 0x5D)
GOLD  = RGBColor(0xC9, 0xA4, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x40, 0x40, 0x40)
LIGHT_GREY_BG = "F0F4F8"
NAVY_HEX = "1A365D"
GOLD_HEX  = "C9A432"

OUTPUT = os.path.join(os.path.dirname(__file__),
                      "PV Residential - Contract - TEMPLATE - REVISED.docx")
LOGO_SVG = os.path.join(os.path.dirname(__file__),
                        "..", "..", "public", "favicon.svg")


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=True, bottom=True, left=True, right=True,
                     color="CCCCCC", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, active in [("top", top), ("bottom", bottom),
                         ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if active:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def para_spacing(para, before=0, after=60, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(13)
        run.font.color.rgb = GOLD
        # underline rule via bottom border on paragraph
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), NAVY_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
    para_spacing(p, before=160, after=60)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    para_spacing(p, before=100, after=40)
    return p


def add_body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    para_spacing(p, before=0, after=50)
    return p


def add_bullet(doc, text, indent=0.4):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    p.paragraph_format.left_indent = Cm(indent)
    para_spacing(p, before=0, after=30)
    return p


def add_notice_box(doc, text, bg=LIGHT_GREY_BG, border_hex=NAVY_HEX):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_borders(cell, color=border_hex, size="12")
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = BLACK
    para_spacing(p, before=60, after=60)
    doc.add_paragraph()
    return tbl


def add_two_col_table(doc, rows_data, header_row=None, col_widths=None):
    """rows_data: list of (label, value) tuples"""
    ncols = 2
    tbl = doc.add_table(rows=0, cols=ncols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    if header_row:
        row = tbl.add_row()
        for i, txt in enumerate(header_row):
            c = row.cells[i]
            set_cell_bg(c, NAVY_HEX)
            p = c.paragraphs[0]
            run = p.add_run(txt)
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = WHITE
            para_spacing(p, before=40, after=40)
    for label, value in rows_data:
        row = tbl.add_row()
        for i, txt in enumerate([label, value]):
            c = row.cells[i]
            p = c.paragraphs[0]
            if i == 0:
                run = p.add_run(txt)
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = NAVY
                set_cell_bg(c, "F5F7FA")
            else:
                run = p.add_run(txt)
                run.font.size = Pt(9.5)
                run.font.color.rgb = BLACK
            para_spacing(p, before=40, after=40)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return tbl


def add_payment_table(doc, rows):
    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = "Table Grid"
    # Header
    hrow = tbl.add_row()
    for i, txt in enumerate(["Instalment", "Trigger / Milestone", "Amount"]):
        c = hrow.cells[i]
        set_cell_bg(c, NAVY_HEX)
        p = c.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, before=40, after=40)
    for inst, trigger, pct in rows:
        row = tbl.add_row()
        for i, (txt, bold) in enumerate([(inst, True), (trigger, False), (pct, True)]):
            c = row.cells[i]
            p = c.paragraphs[0]
            run = p.add_run(txt)
            run.bold = bold
            run.font.size = Pt(9.5)
            run.font.color.rgb = NAVY if bold else BLACK
            if i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_spacing(p, before=40, after=40)
    for row in tbl.rows:
        row.cells[0].width = Inches(1.4)
        row.cells[1].width = Inches(4.0)
        row.cells[2].width = Inches(1.1)
    doc.add_paragraph()
    return tbl


def add_warranty_table(doc):
    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = "Table Grid"
    hrow = tbl.add_row()
    for i, txt in enumerate(["Item", "Minimum Coverage", "Basis"]):
        c = hrow.cells[i]
        set_cell_bg(c, NAVY_HEX)
        p = c.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        para_spacing(p, before=40, after=40)
    data = [
        ("Workmanship", "2 years from commissioning date", "Contractual (minimum legal standard)"),
        ("PV Modules — product defect", "As per OEM warranty — [input]", "OEM manufacturer warranty"),
        ("PV Modules — power output", "As per OEM performance warranty — [input]", "OEM manufacturer warranty"),
        ("Inverter", "As per OEM warranty — [input]", "OEM manufacturer warranty"),
        ("Battery Modules", "As per OEM warranty — [input]", "OEM manufacturer warranty"),
        ("Goods (all equipment)", "2 years from delivery (minimum)", "Cyprus Law 154(I)/2021"),
    ]
    for item, cov, basis in data:
        row = tbl.add_row()
        for i, txt in enumerate([item, cov, basis]):
            c = row.cells[i]
            p = c.paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9.5)
            run.font.color.rgb = NAVY if i == 0 else BLACK
            if i == 0:
                run.bold = True
            para_spacing(p, before=40, after=40)
    for row in tbl.rows:
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(2.8)
        row.cells[2].width = Inches(2.1)
    doc.add_paragraph()


def add_spec_table(doc):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    hrow = tbl.add_row()
    for i, txt in enumerate(["Parameter", "Specification"]):
        c = hrow.cells[i]
        set_cell_bg(c, NAVY_HEX)
        p = c.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True; run.font.size = Pt(9.5); run.font.color.rgb = WHITE
        para_spacing(p, before=40, after=40)
    specs = [
        ("Installation Address", "[Address]"),
        ("Solar Panels", "[No. of Panels] × [Panel Power W] [Model]"),
        ("PV Capacity", "[Total kWp]"),
        ("Inverter(s)", "[No. of Inverters] × [Inverter kW] [Model]"),
        ("Battery Unit(s)", "[No. of BESS] × [BESS Capacity kWh] [Model]  (or: Not applicable)"),
        ("Battery Capacity", "[Total kWh]  (or: Not applicable)"),
        ("Generator Interface (ATS)", "☐ Included   ☐ Not included"),
        ("Remote Monitoring", "☐ Included   ☐ Not included"),
        ("Estimated Annual Yield (indicative)", "[kWh/year] — see §1.4 disclaimer"),
    ]
    for label, val in specs:
        row = tbl.add_row()
        c0 = row.cells[0]; c1 = row.cells[1]
        set_cell_bg(c0, "F5F7FA")
        p0 = c0.paragraphs[0]; p1 = c1.paragraphs[0]
        r0 = p0.add_run(label); r0.bold = True; r0.font.size = Pt(9.5); r0.font.color.rgb = NAVY
        r1 = p1.add_run(val);   r1.font.size = Pt(9.5);   r1.font.color.rgb = BLACK
        para_spacing(p0, before=40, after=40); para_spacing(p1, before=40, after=40)
    for row in tbl.rows:
        row.cells[0].width = Inches(2.2); row.cells[1].width = Inches(4.3)
    doc.add_paragraph()


def add_horizontal_rule(doc, color_hex=NAVY_HEX):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    para_spacing(p, before=0, after=80)
    return p


def logo_png_bytes():
    """Render the Lighthief lightning bolt SVG to a PNG via svglib/reportlab."""
    from svglib.svglib import svg2rlg
    from reportlab.graphics import renderPM
    # Modify SVG to use transparent bg and gold bolt
    with open(LOGO_SVG, "r") as f:
        svg_text = f.read()
    # Remove the black background rects and recolour bolt to gold
    svg_text = svg_text.replace(
        '<rect width="512" height="512" fill="#0B0B0B"/>',
        '<rect width="512" height="512" fill="#1A365D"/>'
    )
    svg_text = svg_text.replace(
        '<rect width="256" height="256" fill="#0B0B0B"/>',
        ''
    )
    svg_text = svg_text.replace('fill="white"', 'fill="#C9A432"')
    tmp = os.path.join(os.path.dirname(__file__), "_tmp_logo.svg")
    with open(tmp, "w") as f:
        f.write(svg_text)
    drawing = svg2rlg(tmp)
    buf = io.BytesIO()
    renderPM.drawToFile(drawing, buf, fmt="PNG")
    buf.seek(0)
    os.remove(tmp)
    return buf


def add_header(doc):
    """Navy header bar: Lighthief logo + company name + document title."""
    # Header table: logo cell | text cell
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    logo_cell = tbl.cell(0, 0)
    text_cell = tbl.cell(0, 1)
    set_cell_bg(logo_cell, NAVY_HEX)
    set_cell_bg(text_cell, NAVY_HEX)
    logo_cell.width = Inches(1.1)
    text_cell.width = Inches(5.4)

    # Logo image
    try:
        img_buf = logo_png_bytes()
        p_logo = logo_cell.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(img_buf, width=Inches(0.75))
        para_spacing(p_logo, before=60, after=60)
    except Exception:
        p_logo = logo_cell.paragraphs[0]
        run = p_logo.add_run("☀")
        run.font.size = Pt(32)
        run.font.color.rgb = GOLD
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Text block
    p1 = text_cell.paragraphs[0]
    r1 = p1.add_run("LIGHTHIEF CYPRUS LTD")
    r1.bold = True
    r1.font.size = Pt(16)
    r1.font.color.rgb = GOLD
    para_spacing(p1, before=40, after=0)

    p2 = text_cell.add_paragraph()
    r2 = p2.add_run("Solar Photovoltaic System Commissioning Contract")
    r2.font.size = Pt(10)
    r2.font.color.rgb = WHITE
    para_spacing(p2, before=0, after=0)

    p3 = text_cell.add_paragraph()
    r3 = p3.add_run("Contractor: Procurement & Installation")
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = RGBColor(0xB0, 0xBE, 0xD4)
    para_spacing(p3, before=0, after=60)

    doc.add_paragraph()  # spacer
    return tbl


# ── Main document generator ────────────────────────────────────────────────────

def build():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Default style ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = BLACK

    # ════════════════════════════════════════════════════════════════════════
    # HEADER
    # ════════════════════════════════════════════════════════════════════════
    add_header(doc)
    add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════════════
    # PARTIES TABLE
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "PARTIES")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c_contractor = tbl.cell(0, 0)
    c_client     = tbl.cell(0, 1)
    set_cell_bg(c_contractor, "F0F4F8")
    set_cell_bg(c_client,     "FAFBFC")

    def party_block(cell, title, rows):
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY
        para_spacing(p, before=40, after=20)
        for label, val in rows:
            p2 = cell.add_paragraph()
            r_lbl = p2.add_run(label + ": ")
            r_lbl.bold = True; r_lbl.font.size = Pt(9); r_lbl.font.color.rgb = GREY
            r_val = p2.add_run(val)
            r_val.font.size = Pt(9); r_val.font.color.rgb = BLACK
            para_spacing(p2, before=0, after=20)

    party_block(c_contractor, "CONTRACTOR", [
        ("Company",       "LIGHTHIEF CYPRUS LTD"),
        ("Reg. No.",      "HE 477423"),
        ("TIN",           "60187188Q"),
        ("Address",       "15 Agaritsis, Nektaria Court, Office 201,\nOffice 201, 3035 Limassol, Cyprus"),
        ("Email",         "office@lighthief.com"),
        ("Tel",           "+357 77 77 00 50"),
        ("Representative","Alexander Papacosta, Director"),
    ])
    party_block(c_client, "CLIENT", [
        ("Full Name / Company", "[Fill]"),
        ("ID / Reg. No.",       "[Fill]"),
        ("Address",             "[Fill]"),
        ("Email",               "[Fill]"),
        ("Tel",                 "[Fill]"),
        ("Representative",      "[Fill]"),
    ])
    for row in tbl.rows:
        row.cells[0].width = Inches(3.25)
        row.cells[1].width = Inches(3.25)
    doc.add_paragraph()

    # Contractor and Client are referred to individually as a "Party"
    p = doc.add_paragraph()
    r = p.add_run(
        "Contractor and Client are referred to individually as a \"Party\" and collectively as "
        "the \"Parties.\" This Contract is governed by the laws of the Republic of Cyprus."
    )
    r.font.size = Pt(9.5); r.font.color.rgb = GREY
    r.italic = True
    para_spacing(p, before=0, after=80)
    add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 1. SCOPE OF WORK
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1.  SCOPE OF WORK")

    add_subheading(doc, "1.1  System Specifications")
    add_body(doc, (
        "The Contractor shall supply, deliver, install, and commission a photovoltaic and battery "
        "energy storage system (where batteries are applicable) at the premises specified below. "
        "The scope includes all labour, materials, cabling, protection devices, and programming "
        "required for a fully operational system."
    ))
    add_spec_table(doc)

    add_subheading(doc, "1.2  Services Included")
    add_body(doc, "The Contractor shall provide:")
    for item in [
        "Site assessment and system design;",
        "Submission of application to EAC (Electricity Authority of Cyprus);",
        "Handling of all necessary permits and approvals;",
        "Supply of all equipment and materials;",
        "Professional installation in accordance with Cyprus standards;",
        "System commissioning and testing;",
        "Assistance to the Client with any available government subsidy applications "
        "(see §3.3 for scope and limitations);",
        "Training on system operation and monitoring.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "1.3  Installation Standards")
    add_body(doc, "All work shall comply with:")
    for item in [
        "Cyprus Electricity Authority (EAC) regulations;",
        "Cyprus Energy Regulatory Authority (CERA) requirements;",
        "EU safety and quality standards;",
        "Local building and planning regulations;",
        "All applicable Cyprus laws and regulations.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "1.4  Yield Disclaimer")
    add_notice_box(doc, (
        "IMPORTANT: Any estimated annual yield figures provided in this Contract or in any "
        "accompanying proposal, quotation, or communication are indicative only. Actual energy "
        "generation depends on solar irradiation, shading, weather conditions, panel orientation, "
        "system degradation, and grid availability. The Contractor makes no warranty as to actual "
        "energy output unless a written performance guarantee is separately executed."
    ), bg="FFF8E1", border_hex=GOLD_HEX)

    # ════════════════════════════════════════════════════════════════════════
    # 2. TIMELINE AND MILESTONES
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2.  TIMELINE AND MILESTONES")

    add_subheading(doc, "2.1  Project Timeline")
    rows_tl = [
        ("Contract Signing Date",    "[Date]"),
        ("EAC Application Submission","Within 1–2 weeks after all client documents received"),
        ("Expected EAC Approval",    "5–8 weeks from submission (subject to EAC processing times)"),
        ("Equipment Delivery",       "Within agreed programme; Client notified ≥ 3 days in advance"),
        ("Installation",             "Within 4 months of contract signing; duration 2–3 working days"),
        ("EAC Inspection",           "Scheduled by EAC following installation"),
        ("System Commissioning",     "Upon successful EAC inspection and grid connection approval"),
    ]
    add_two_col_table(doc, rows_tl, header_row=["Milestone", "Target Date / Trigger"],
                      col_widths=[2.4, 4.1])

    add_subheading(doc, "2.2  EAC Application Rejection")
    add_body(doc, (
        "If EAC rejects the grid connection application or requires modifications to the installed "
        "system that render the project technically or economically unviable, the following shall apply:"
    ))
    for item in [
        "The Contractor shall notify the Client in writing within five (5) Business Days of receiving "
        "EAC's decision, setting out the reasons and any remedial options;",
        "The Parties shall negotiate in good faith for up to thirty (30) days to agree on a revised "
        "technical solution;",
        "If no solution is agreed within thirty (30) days, either Party may terminate this Contract "
        "by written notice. Upon termination: (i) the Client shall pay for all work completed and "
        "materials specifically procured; (ii) the Contractor shall refund all amounts received in "
        "excess of such costs within fourteen (14) days;",
        "The Contractor shall not be liable for EAC processing delays, EAC fee changes, or EAC "
        "rejection decisions beyond its reasonable control.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "2.3  Delays")
    add_body(doc, "The Contractor shall not be liable for delays caused by:")
    for item in [
        "EAC processing times;",
        "Weather conditions or force majeure events (see §13.6);",
        "Client-requested changes;",
        "Delays in obtaining permits beyond the Contractor's reasonable control;",
        "Client's failure to provide required documents or access.",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # 3. PAYMENT TERMS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3.  PAYMENT TERMS")

    add_subheading(doc, "3.1  Total Contract Price (VAT Inclusive)")
    add_two_col_table(doc, [
        ("Total Contract Price (VAT incl. 19%)", "€ [Price]"),
        ("Net Amount (excl. VAT)",               "€ [Price]"),
        ("VAT (19%)",                            "€ [Price]"),
    ], col_widths=[3.5, 3.0])

    add_subheading(doc, "3.2  Payment Schedule")
    add_body(doc, (
        "Payments shall be made by bank transfer to the Contractor's account. "
        "No works shall commence or materials ordered until the Deposit (Instalment 1) is received. "
        "All amounts are VAT-inclusive."
    ))
    add_payment_table(doc, [
        ("#1 — Deposit (40%)",       "Due upon contract signing",                    "40%"),
        ("#2 — Pre-Installation (55%)",
         "Due 3 days before installation date.\n"
         "The Contractor shall provide equipment model, serial numbers, and delivery "
         "confirmation before this payment falls due.",                               "55%"),
        ("#3 — Final (5%)",          "Due upon EAC grid connection approval (PAC)",  "5%"),
    ])

    add_subheading(doc, "3.2.1  Late Payment Interest")
    add_body(doc, (
        "If the Client fails to make any payment when due, interest shall accrue on the overdue "
        "amount from the due date until the date of actual payment at the rate prescribed by "
        "EU Directive 2011/7/EU on combating late payment in commercial transactions — being "
        "the European Central Bank main refinancing rate plus eight percentage points (8%) per "
        "annum, calculated on a simple interest basis. The Contractor may also suspend works "
        "upon seven (7) days' written notice of non-payment."
    ))

    add_subheading(doc, "3.2.2  Separate EAC/DSO Application Fee")
    add_notice_box(doc, (
        "Not included in the Contract Price: The EAC/DSO one-time connection application fee "
        "(covering DSO processing, inspection, system connection to the distribution network, "
        "and ripple control connection) is payable directly by the Client to EAC upon issuance "
        "of connection terms. The current fee is approximately €365 + VAT 19% = €434.35. "
        "This fee is set by EAC/CERA and may change without notice."
    ), bg="FFF8E1", border_hex=GOLD_HEX)

    add_subheading(doc, "3.3  Government Subsidies")
    add_body(doc, (
        "The Contractor shall provide reasonable assistance to the Client in applying for "
        "government subsidy schemes available at the time of application, including preparation "
        "of required technical documentation. The Contractor's assistance is limited to document "
        "preparation and submission support; the Contractor makes no warranty that any subsidy "
        "will be approved, that subsidy schemes will remain available, or that the Client will "
        "qualify. Subsidies are paid directly to the Client by the relevant authority and are "
        "not deducted from the Contract Price."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # 4. WARRANTIES AND GUARANTEES
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4.  WARRANTIES AND GUARANTEES")

    add_subheading(doc, "4.1  Warranty Summary")
    add_warranty_table(doc)

    add_subheading(doc, "4.2  Statutory Legal Guarantee (Cyprus Law 154(I)/2021)")
    add_body(doc, (
        "In accordance with Cyprus Law 154(I)/2021 implementing EU Directive 2019/771 on the "
        "sale of goods, the Client is entitled to a minimum two-year legal guarantee from the "
        "date of delivery of each item of equipment, covering any lack of conformity. This "
        "statutory guarantee is in addition to any commercial warranty provided by the OEM and "
        "cannot be excluded or reduced by contract. If any equipment is found to be non-conforming "
        "within two years of delivery, the Client may require repair, replacement, a price "
        "reduction, or rescission of the contract, subject to the conditions set out in the Law."
    ))

    add_subheading(doc, "4.3  Workmanship Warranty")
    add_body(doc, (
        "The Contractor warrants all installation workmanship for a period of two (2) years from "
        "the commissioning date. This covers defects directly attributable to the quality of "
        "installation performed by the Contractor's personnel. During the warranty period, the "
        "Contractor shall remedy any covered defect within a reasonable time following written "
        "notice from the Client, at no additional cost."
    ))

    add_subheading(doc, "4.4  Warranty Conditions and Exclusions")
    add_body(doc, "Warranties are void or limited if:")
    for item in [
        "The system is modified or repaired without the Contractor's prior written consent;",
        "Damage is caused by Client negligence or misuse;",
        "Damage is caused by acts of God, fire, flood, or other events beyond both Parties' control;",
        "The Client has failed to fulfil its payment obligations under §3.",
    ]:
        add_bullet(doc, item)

    add_body(doc, (
        "Warranties do not cover routine consumables, panel soiling, shading from new structures "
        "erected after installation, or performance variations within the OEM's published tolerance range."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # 5. CLIENT OBLIGATIONS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5.  CLIENT OBLIGATIONS")
    add_body(doc, "The Client shall:")
    for item in [
        "Provide accurate electricity consumption data (last 6 bills);",
        "Ensure clear, safe, and unobstructed access to the installation site;",
        "Obtain and provide to the Contractor all necessary consents from co-owners, neighbours, "
        "landlords, or property management authorities (including any condominium owners' assembly "
        "resolution required under Cyprus Condominium Law L.6/1993). By signing this Contract, the "
        "Client warrants that all required consents have been obtained or will be obtained before "
        "installation commences;",
        "Provide all necessary documentation required for the EAC application in a timely manner;",
        "Make payments in accordance with the agreed schedule;",
        "Not modify or repair the system without the Contractor's prior written approval;",
        "Notify the Contractor of any distribution board extension requirements before "
        "installation — such works, if required, are not included in the Contract Price and will "
        "be quoted separately. Shading mitigation and inverter positioning are the Client's "
        "responsibility, with the Contractor available to advise.",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # 6. CONTRACTOR OBLIGATIONS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6.  CONTRACTOR OBLIGATIONS")
    add_body(doc, "The Contractor shall:")
    for item in [
        "Perform all work with professional competence and in accordance with applicable standards;",
        "Use only certified, high-quality equipment conforming to EU CE marking requirements;",
        "Comply with all applicable safety regulations;",
        "Obtain all necessary permits and regulatory approvals within its scope;",
        "Provide the Client with all required documentation and certificates upon commissioning;",
        "Train the Client on system operation and monitoring;",
        "Respond to warranty claims in writing within ten (10) Business Days of notification;",
        "Provide the Client with equipment serial numbers and delivery confirmation before "
        "Instalment 2 falls due.",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # 7. MAINTENANCE AND SUPPORT
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7.  MAINTENANCE AND SUPPORT")

    add_subheading(doc, "7.1  Optional Maintenance Services")
    add_body(doc, "The Contractor offers optional maintenance packages, including:")
    for item in ["Annual inspection and cleaning;",
                 "Performance monitoring and reporting;",
                 "Priority response for service calls."]:
        add_bullet(doc, item)
    add_body(doc, "Maintenance packages are quoted separately and are not included in the Contract Price.")

    add_subheading(doc, "7.2  Client Self-Maintenance")
    add_body(doc, "The Client should:")
    for item in ["Keep panels reasonably clean and free from obstruction;",
                 "Monitor system performance via the installed monitoring application;",
                 "Report any faults, alarms, or unusual readings promptly;",
                 "Not install new structures that shade the panels."]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # 8. NET BILLING TERMS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8.  NET BILLING TERMS")

    add_subheading(doc, "8.1  EAC Net Billing Contract")
    add_body(doc, (
        "Upon grid connection approval, the Client will enter into a Net Billing contract "
        "directly with EAC under the current EAC/CERA Net Billing framework applicable to the "
        "Client's supply category (Residential or Commercial)."
    ))

    add_subheading(doc, "8.2  How Net Billing Works")
    add_body(doc, (
        "Under the EAC Net Billing framework, electricity produced by the photovoltaic system "
        "is first consumed on-site. At the end of the two-month billing period:"
    ))
    for item in [
        "Electricity imported from the grid is charged at the standard retail tariff applicable "
        "to the Client's supply category;",
        "Surplus electricity exported to the grid is credited at the official Net Billing export "
        "tariff set by EAC/CERA, which is different from and generally lower than the import tariff;",
        "Export credits are applied against import charges in the Client's EAC bill;",
        "Credits are not paid out in cash unless explicitly permitted by EAC regulations in force "
        "at the time.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "8.3  EAC Fees and Tariff Changes")
    add_body(doc, (
        "The Client shall pay all standard EAC charges applicable under Net Billing, including "
        "supply charges, network charges, ancillary service charges, and any Net Billing-specific "
        "administrative or settlement fees. EAC/CERA tariffs and fees are subject to regulatory "
        "change; the Contractor has no control over and accepts no liability for future tariff "
        "adjustments."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # 9. INSURANCE AND LIABILITY
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9.  INSURANCE AND LIABILITY")

    add_subheading(doc, "9.1  Contractor Insurance")
    add_body(doc, "The Contractor maintains the following insurance policies:")
    for item in [
        "Professional Liability (Errors & Omissions) Insurance;",
        "Public / General Liability Insurance;",
        "Workers' Compensation / Employer's Liability Insurance.",
    ]:
        add_bullet(doc, item)
    add_body(doc, "Proof of current insurance is available upon written request.")

    add_subheading(doc, "9.2  Limitation of Liability")
    add_body(doc, (
        "Except for fraud, wilful misconduct, or death or personal injury caused by the "
        "Contractor's negligence, the Contractor's total aggregate liability under this Contract "
        "shall not exceed the total Contract Price paid by the Client. Neither Party shall be "
        "liable for indirect, consequential, or loss-of-profit damages, except as required by "
        "mandatory consumer protection law."
    ))

    add_subheading(doc, "9.3  Indemnification")
    add_body(doc, (
        "Each Party shall indemnify the other against third-party claims arising directly from "
        "their own negligence or wilful misconduct in the performance of their respective "
        "obligations under this Contract."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # 10. RIGHT OF WITHDRAWAL AND CANCELLATION
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10.  RIGHT OF WITHDRAWAL AND CANCELLATION")

    add_subheading(doc, "10.1  14-Day Right of Withdrawal (Statutory — Off-Premises Contracts)")
    add_notice_box(doc, (
        "STATUTORY RIGHT: Under Cyprus Law 133(I)/2013 implementing EU Consumer Rights Directive "
        "2011/83/EU, if this Contract is concluded at the Client's premises or any location other "
        "than the Contractor's business premises (an off-premises contract), the Client has a "
        "statutory right to withdraw from this Contract without giving any reason within FOURTEEN "
        "(14) CALENDAR DAYS from the date of signing.\n\n"
        "To exercise this right, the Client must notify the Contractor before the 14-day period "
        "expires using the Withdrawal Form in Annex A, or by any other unambiguous written "
        "statement (email to office@lighthief.com is acceptable).\n\n"
        "Effect of withdrawal: The Contractor shall refund all payments received within fourteen "
        "(14) days of receiving the withdrawal notice, subject to §10.1.1 below."
    ), bg="E8F4FD", border_hex=NAVY_HEX)

    add_subheading(doc, "10.1.1  Early Start and Partial Withdrawal Refund")
    add_body(doc, (
        "If the Client expressly requests in writing that works commence before the 14-day "
        "withdrawal period expires, and the Client subsequently exercises the withdrawal right, "
        "the Client shall pay the Contractor a proportionate amount for services actually performed "
        "up to the point of withdrawal. No cancellation fee shall apply for withdrawals within "
        "the 14-day period beyond this proportionate payment."
    ))

    add_subheading(doc, "10.2  Cancellation After the 14-Day Period")
    add_body(doc, (
        "If the Client cancels this Contract after the 14-day withdrawal period has expired, "
        "the following cancellation charges apply to cover costs and losses incurred:"
    ))
    add_two_col_table(doc, [
        ("After 14-day period, before EAC application submitted",
         "10% of Contract Price (administration and design costs)"),
        ("After EAC application submitted, before installation",
         "30% of Contract Price (application, engineering, material procurement costs)"),
        ("After installation completed",
         "No refund — full Contract Price is due"),
    ], header_row=["Stage of Cancellation", "Cancellation Charge"],
       col_widths=[3.3, 3.2])
    add_body(doc, (
        "These cancellation charges represent a genuine pre-estimate of the Contractor's losses "
        "at each stage and are not a penalty."
    ))

    add_subheading(doc, "10.3  Contractor Termination Rights")
    add_body(doc, "The Contractor may terminate this Contract by written notice for:")
    for item in [
        "Client non-payment following thirty (30) days' written notice;",
        "Material breach of Client obligations that remains unremedied after fourteen (14) days' "
        "written notice;",
        "Discovery of unsafe site conditions that cannot be resolved within a reasonable timeframe.",
    ]:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # 11. DISPUTE RESOLUTION
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11.  DISPUTE RESOLUTION")

    add_subheading(doc, "11.1  Good-Faith Negotiation")
    add_body(doc, (
        "The Parties shall first attempt to resolve any dispute through good-faith negotiation "
        "within thirty (30) days of one Party giving written notice of the dispute to the other."
    ))

    add_subheading(doc, "11.2  Mediation and Consumer Protection")
    add_body(doc, (
        "If negotiation fails, the Parties shall, before commencing legal proceedings, attempt "
        "to resolve the dispute through mediation. The Client may also contact the following "
        "authorities for assistance at any time:"
    ))
    for item in [
        "Cyprus Consumer Protection Service: tel. 1429 (free consumer helpline);",
        "Cyprus Energy Regulatory Authority (CERA): +357 22 666 363;",
        "Online Dispute Resolution platform (EU ODR): https://ec.europa.eu/consumers/odr/.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "11.3  Legal Proceedings")
    add_body(doc, (
        "Any dispute not resolved by negotiation or mediation shall be subject to the exclusive "
        "jurisdiction of the courts of the Republic of Cyprus."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # 12. DATA PROTECTION
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "12.  DATA PROTECTION")
    add_body(doc, (
        "The Contractor shall comply with the General Data Protection Regulation (EU) 2016/679 "
        "(GDPR) and applicable Cyprus data protection laws regarding any personal data collected "
        "in connection with this Contract. Personal data will be processed solely for the purpose "
        "of performing this Contract and will not be disclosed to third parties without the "
        "Client's consent, except as required by law. The Client may request access to, "
        "correction, or deletion of their personal data at any time by contacting "
        "office@lighthief.com."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # 13. GENERAL PROVISIONS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "13.  GENERAL PROVISIONS")

    clauses = [
        ("13.1  Entire Agreement",
         "This Contract constitutes the entire agreement between the Parties regarding its "
         "subject matter and supersedes all prior negotiations, representations, or agreements."),
        ("13.2  Amendments",
         "Any amendment must be in writing and signed by authorised representatives of both Parties."),
        ("13.3  Governing Law",
         "This Contract is governed by the laws of the Republic of Cyprus."),
        ("13.4  Language",
         "This Contract is executed in English. A Greek translation is available upon request. "
         "In the event of any inconsistency, the Greek version shall prevail for Clients domiciled "
         "in the Republic of Cyprus."),
        ("13.5  Severability",
         "If any provision is found invalid or unenforceable, the remaining provisions continue "
         "in full force and effect."),
        ("13.6  Force Majeure",
         "Neither Party shall be liable for delays or failure to perform obligations due to "
         "circumstances beyond their reasonable control, including acts of God, war, terrorism, "
         "pandemic, governmental action, or grid authority decisions, provided the affected Party "
         "gives prompt written notice and uses reasonable efforts to mitigate the impact."),
        ("13.7  Assignment",
         "Neither Party may assign this Contract without the prior written consent of the other."),
        ("13.8  Notices",
         "All formal notices under this Contract shall be in writing and sent by email (with "
         "delivery confirmation) or by registered post to the addresses in the Parties section. "
         "Email notices are deemed received on the next Business Day."),
    ]
    for title, text in clauses:
        add_subheading(doc, title)
        add_body(doc, text)

    # ════════════════════════════════════════════════════════════════════════
    # 14. SPECIAL CONDITIONS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "14.  SPECIAL CONDITIONS")

    add_subheading(doc, "14.1  Roof Tile Guarantee")
    add_body(doc, (
        "The Contractor guarantees to replace any roof tiles broken or directly damaged during "
        "panel installation with tiles of the same style and colour. Where exactly matching tiles "
        "are no longer commercially available, the Contractor will replace with tiles of the "
        "closest available colour and grade. This guarantee covers only tiles directly damaged "
        "during installation; it does not extend to pre-existing roof membrane, waterproofing, "
        "or structural conditions."
    ))

    # ════════════════════════════════════════════════════════════════════════
    # PAYMENT BANK DETAILS — FRAUD WARNING
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "PAYMENT DETAILS")

    add_two_col_table(doc, [
        ("Account Name", "LIGHTHIEF CYPRUS LTD"),
        ("Account Number", "357044102353"),
        ("IBAN", "CY86002001950000357044102353"),
        ("SWIFT/BIC", "BCYPCY2N"),
        ("Bank", "Bank of Cyprus"),
    ], col_widths=[2.2, 4.3])

    add_notice_box(doc, (
        "⚠  FRAUD PREVENTION NOTICE: Lighthief Cyprus Ltd will NEVER change bank account details "
        "by email or SMS. Before making any payment, verify the account details above are correct "
        "by calling +357 77 77 00 50. Do not transfer funds to any account other than the one "
        "above without first confirming directly with a Lighthief representative by phone."
    ), bg="FFF3CD", border_hex="FF8C00")

    # ════════════════════════════════════════════════════════════════════════
    # SIGNATURES
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "SIGNATURES")
    add_body(doc, (
        "By signing below, both Parties confirm they have read, understood, and agreed to all "
        "terms of this Contract, including the Client's right of withdrawal in §10.1."
    ))
    doc.add_paragraph()

    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (party, name, cap) in enumerate([
        ("FOR THE CLIENT", "[Name]", "[Capacity]"),
        ("FOR THE CONTRACTOR — LIGHTHIEF CYPRUS LTD", "Alexander Papacosta", "Director"),
    ]):
        cell = sig_tbl.cell(0, i)
        p_title = cell.paragraphs[0]
        r = p_title.add_run(party)
        r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY
        para_spacing(p_title, before=0, after=80)
        for label, val in [("Name:", name), ("Capacity:", cap),
                            ("Signature:", "___________________________"),
                            ("Date:", "___________________________")]:
            p = cell.add_paragraph()
            r_lbl = p.add_run(label + "  ")
            r_lbl.bold = True; r_lbl.font.size = Pt(9.5); r_lbl.font.color.rgb = GREY
            r_val = p.add_run(val)
            r_val.font.size = Pt(9.5); r_val.font.color.rgb = BLACK
            para_spacing(p, before=0, after=50)
        cell.width = Inches(3.25)

    doc.add_paragraph()
    p_note = doc.add_paragraph()
    r_note = p_note.add_run("* Each Party shall retain one signed original of this Contract.")
    r_note.font.size = Pt(8.5); r_note.font.color.rgb = GREY; r_note.italic = True

    # ════════════════════════════════════════════════════════════════════════
    # CLIENT NOTICE BOX
    # ════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "IMPORTANT NOTICE TO CLIENT")
    add_notice_box(doc, (
        "This Contract includes important terms regarding your rights and obligations.\n\n"
        "You are advised to:\n"
        "  •  Read all terms carefully before signing.\n"
        "  •  Keep a copy for your records.\n"
        "  •  Note your 14-day right of withdrawal (§10.1) — you may cancel this Contract "
        "without penalty within 14 days of signing if it was concluded at your premises.\n"
        "  •  Understand the payment schedule and your right to receive equipment details "
        "before Instalment 2 falls due.\n"
        "  •  Be aware of your consumer rights under Cyprus law.\n\n"
        "For more information on your consumer rights, contact:\n"
        "  •  Cyprus Consumer Protection Service: 1429\n"
        "  •  Cyprus Energy Regulatory Authority (CERA): +357 22 666 363\n"
        "  •  EU Online Dispute Resolution: https://ec.europa.eu/consumers/odr/"
    ), bg="E8F4FD", border_hex=NAVY_HEX)

    # ════════════════════════════════════════════════════════════════════════
    # ANNEX A — WITHDRAWAL FORM
    # ════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "ANNEX A — STANDARD WITHDRAWAL FORM")
    add_body(doc, (
        "(Complete and return this form only if you wish to withdraw from the Contract)"
    ))
    add_notice_box(doc, (
        "TO: Lighthief Cyprus Ltd, 15 Agaritsis, Nektaria Court, Office 201, "
        "3045 Zakaki, Limassol, Cyprus\n"
        "Email: office@lighthief.com\n\n"
        "I / We (*) hereby give notice that I / We (*) withdraw from my / our (*) contract for "
        "the supply and installation of the following:\n\n"
        "System address: _______________________________________________\n\n"
        "Contract date: _______________________________________________\n\n"
        "Client name(s): _______________________________________________\n\n"
        "Client address: _______________________________________________\n\n"
        "Client signature(s) (only if this form is notified on paper):\n\n"
        "_______________________________________________\n\n"
        "Date: _______________________________________________\n\n"
        "(*) Delete as appropriate."
    ), bg="FAFBFC", border_hex=NAVY_HEX)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
