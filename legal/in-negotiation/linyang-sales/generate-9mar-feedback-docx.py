"""
Generate redline feedback on Linyang Sales Contract 9 March 2026.
Original text in black, Lighthief feedback/comments in red.
Focuses on: items still missing, items needing correction, and acceptance notes.
"""

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(SCRIPT_DIR, "Linyang sales contract - Lighthief - 09.03.docx")
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "Linyang-sales-contract-09Mar-Lighthief-FEEDBACK.docx")

RED = RGBColor(0xFF, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x80, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)


def add_run(para, text, color=BLACK, bold=False, italic=False, size=Pt(11)):
    run = para.add_run(text)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = size
    return run


def add_black_para(doc, text, bold=False):
    para = doc.add_paragraph()
    add_run(para, text, BLACK, bold=bold)
    return para


def add_red_para(doc, text, bold=False):
    para = doc.add_paragraph()
    add_run(para, text, RED, bold=bold)
    return para


def add_green_note(doc, text):
    para = doc.add_paragraph()
    add_run(para, "[ACCEPTED] ", GREEN, bold=True)
    add_run(para, text, GREEN, italic=True)
    return para


def add_red_comment(doc, text):
    para = doc.add_paragraph()
    add_run(para, "[LIGHTHIEF COMMENT] ", RED, bold=True)
    add_run(para, text, RED)
    return para


def add_red_insert(doc, heading, text):
    para = doc.add_paragraph()
    add_run(para, f"[LIGHTHIEF — {heading}] ", RED, bold=True)
    add_run(para, text, RED)
    return para


def copy_para(doc, text, bold=False):
    if not text.strip():
        doc.add_paragraph()
        return
    add_black_para(doc, text, bold=bold)


def generate():
    src = Document(INPUT_FILE)
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Calibri'

    # Title
    title = doc.add_paragraph()
    add_run(title, "SALES CONTRACT — LIGHTHIEF FEEDBACK ON 9 MARCH 2026 VERSION", BLACK, bold=True, size=Pt(14))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    add_run(sub, "Base document: Linyang sales contract - Lighthief - 09.03.docx\n", BLACK, italic=True, size=Pt(9))
    add_run(sub, "Feedback by: Lighthief Cyprus Ltd | Date: March 2026\n", BLACK, italic=True, size=Pt(9))
    add_run(sub, "BLACK = original contract text. ", BLACK, size=Pt(9))
    add_run(sub, "GREEN = accepted items. ", GREEN, bold=True, size=Pt(9))
    add_run(sub, "RED = items requiring completion, amendment, or insertion.", RED, bold=True, size=Pt(9))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    paras = src.paragraphs
    i = 0
    while i < len(paras):
        text = paras[i].text
        is_bold = any(r.bold for r in paras[i].runs if r.bold is not None)

        # ─── Art 1A.3: Accepted ───
        if i == 28 and '1A.3' in text:
            copy_para(doc, text)
            add_green_note(doc, "Art. 1A.3 accepted as-is. Anti-circumvention is covered by the Distribution Agreement (Sections 7-9).")
            i += 1
            continue

        # ─── Installation Scope: Accepted ───
        if i == 61 and 'Seller shall be responsible' in text:
            copy_para(doc, text)
            add_green_note(doc, "Installation scope items (a)-(f) now complete — accepted. Includes workmanship warranty 5 years from PAC.")
            i += 1
            continue

        # ─── No Cross-Liability: Accepted ───
        if i == 77 and 'No Cross-Liability' in text:
            copy_para(doc, text)
            add_green_note(doc, "No Cross-Liability clause restored — accepted.")
            i += 1
            continue

        # ─── Payment: STILL BLANK ───
        if i == 83 and '% advance payment' in text:
            copy_para(doc, text)
            add_red_comment(doc,
                "PAYMENT PERCENTAGES STILL BLANK. Please populate. As discussed on our call, "
                "Linyang proposed 30/60/10. Lighthief proposes 30/50/10/10 (shifting 10% to a post-delivery DAP milestone "
                "for Cyprus import VAT cash flow management). Payment percentages must be agreed and inserted before "
                "this contract can be executed. We are happy to discuss and finalise on a call.")
            i += 1
            continue

        # ─── End-customer auto-extension: Accepted ───
        if i == 89 and 'automatically extend' in text:
            copy_para(doc, text)
            add_green_note(doc, "End-customer payment dependency with automatic extension — accepted.")
            i += 1
            continue

        # ─── Interest rate: MISSING ───
        if i == 93 and 'Buyer shall not be deemed in default' in text:
            copy_para(doc, text)
            add_red_comment(doc,
                "LATE PAYMENT INTEREST RATE IS MISSING. The previous version's 2% per month clause was removed but "
                "no replacement rate was inserted. The contract currently has NO interest rate for late payments. "
                "Please insert: 'Late payments shall bear interest at the rate prescribed by EU Directive 2011/7/EU "
                "(the European Central Bank main refinancing rate plus eight percent (8%) per annum), calculated on "
                "a simple interest basis.' This rate is legally defensible, consistent across EU jurisdictions, and "
                "approximately 12% APR.")
            i += 1
            continue

        # ─── Performance Warranties 8D: Accepted with note ───
        if i == 159 and 'State of Health' in text:
            copy_para(doc, text)
            add_green_note(doc, "Section 8D Performance Warranties accepted — SOH 85/79.58/70%, RTE 86.32%, cycle life 7,000, PCS 98% all confirmed as contractual product warranties.")
            i += 1
            continue

        # ─── Voiding condition: Query ───
        if i == 167 and 'SOC < 5%' in text:
            copy_para(doc, text)
            add_red_comment(doc,
                "NEW VOIDING CONDITION NOTED. We accept this in principle — SOC < 5% for 7 consecutive days voids "
                "warranty for affected cells/modules. However, please clarify: (1) Does 'null and void' mean the "
                "warranty is permanently voided for those cells, or voided only for degradation occurring during "
                "the deep discharge period? (2) Is this condition per-container or per-system? (3) Can the warranty "
                "be reinstated if the Buyer demonstrates the cells were not damaged by the deep discharge (e.g., via "
                "SOH test showing no capacity loss)? We will mirror this condition in our downstream EPC contracts.")
            i += 1
            continue

        # ─── Liability Cap: Accepted ───
        if i == 185 and 'total aggregate liability' in text:
            copy_para(doc, text)
            add_green_note(doc, "Tiered liability cap accepted: 10% warranty / 50% general / 100% fraud+manufacturing defect. This is our Compromise B position and we confirm acceptance.")
            i += 1
            continue

        # ─── 9B.1-9B.3 Financial Guarantees: BLANK ───
        if i == 208 and not text.strip():
            # Paras 209-220 are all empty — this is where 9B.1-9B.3 should be
            copy_para(doc, text)
            add_red_comment(doc,
                "SECTIONS 9B.1 (APG), 9B.2 (PERFORMANCE BOND), AND 9B.3 (PRODUCT LIABILITY INSURANCE) ARE BLANK. "
                "These sections were proposed in our 6 March comments and are referenced by 9B.4 (Condition Precedent) "
                "and 9B.5 (Buyer Financial Standing) which ARE present. Please insert:\n\n"
                "9B.1 Advance Payment Guarantee (APG): The Seller shall provide an unconditional, irrevocable bank "
                "guarantee equal to 100% of the advance payment, issued by an internationally recognised bank. "
                "To be delivered prior to or simultaneously with the advance payment; valid until full delivery +30 days; "
                "callable on demand in the event of Seller default.\n\n"
                "9B.2 Performance Bond: 5% of total contract value, issued as a corporate guarantee backed by a bank. "
                "To be delivered within 14 days of receipt of the advance payment under Article 6(a); valid until "
                "final acceptance; covers defects in materials, workmanship, and non-performance.\n\n"
                "9B.3 Product Liability Insurance (AXA): The Seller shall maintain product liability insurance with "
                "minimum EUR 5,000,000 per occurrence (currently held with AXA Tianping). The Seller shall provide "
                "a copy of the insurance certificate within 14 days of contract signing and upon each renewal. "
                "NOTE: We have received the AXA policy terms (9 March 2026) but still require the declarations page "
                "showing named insured, coverage limits, territory, and product schedule.")
            # Skip empty paras 209-220
            while i < 221 and not paras[i].text.strip():
                i += 1
            continue

        # ─── 9B.4 Condition Precedent: Accepted ───
        if i == 221 and '9B.4' in text:
            copy_para(doc, text)
            add_green_note(doc, "9B.4 Condition Precedent accepted — APG must be delivered before advance payment is due.")
            i += 1
            continue

        # ─── 9B.5 Credit Check: Accepted ───
        if i == 222 and '9B.5' in text:
            copy_para(doc, text)
            add_green_note(doc, "9B.5 Buyer Financial Standing accepted.")
            i += 1
            continue

        # ─── Delay LD Rates: MISSING ───
        if i == 232 and text.strip() == '.':
            copy_para(doc, text)
            add_red_comment(doc,
                "DELAY LD DEFINITIONS AND RATES ARE MISSING. The exclusions section (below) is present, confirming "
                "that delay LDs are intended. However, the definitions, rate schedule, and cap are absent. "
                "These were agreed in the RFI (February 2026). Please insert before the Exclusions:\n\n"
                "9C.0 Definitions: 'Critical Path Equipment' (PCS, MV Skid, transformer — delay = Full Order Delay), "
                "'Non-Critical Path Equipment' (battery containers — delay = Partial Delay), "
                "'LD Value' (Full Order = total order value; Partial = delayed goods value).\n\n"
                "9C.1 Delay Liquidated Damages:\n"
                "(a) Days 1-30: 0.1% of LD Value per day;\n"
                "(b) Days 31-60: 0.15% of LD Value per day;\n"
                "(c) Days 61+: 0.2% of LD Value per day;\n"
                "(d) Cap: 10% of LD Value per order.\n\n"
                "9C.2 Application: PCS/MV delayed = Full Order Delay (LD on total value); "
                "individual container delayed = Partial Delay (LD on delayed goods only).\n\n"
                "These rates were confirmed by Linyang in the RFI response (February 2026): "
                "0.1% and 0.15% explicitly agreed; 0.2% for 61+ days is Lighthief's position; "
                "10% combined cap was RFI-confirmed.")
            i += 1
            continue

        # ─── Delay LD Exclusions: Accepted with note on (d) ───
        if i == 238 and 'Exclusions' in text:
            copy_para(doc, text)
            add_green_note(doc, "Delay LD Exclusions (a)-(c) accepted: Buyer late payment, Force Majeure, site readiness.")
            i += 1
            continue

        if i == 241 and 'delays in the Owner' in text:
            copy_para(doc, text)
            add_red_comment(doc,
                "EXCLUSION (d) — WORDING TOO BROAD. We accept the principle that delays caused by the Buyer's late "
                "confirmation of technical agreements or drawings should exclude delay LDs. However, the phrase "
                "'or any other causes attributable to the Buyer' is overly broad and could be interpreted to cover "
                "any Buyer-related circumstance. Please amend to: '(d) delays in the Buyer's confirmation of the "
                "technical agreement, system design drawings, or grid connection specifications, where such delays "
                "directly prevent the Seller from commencing production or shipment preparation.'")
            i += 1
            continue

        # ─── No Set-Off: Accepted ───
        if i == 242 and 'No set-Off' in text:
            copy_para(doc, text)
            add_green_note(doc, "No Set-Off clause accepted.")
            i += 1
            continue

        # ─── Commissioning Section 10: Accepted ───
        if i == 262 and '10.1 Commissioning' in text:
            copy_para(doc, text)
            add_green_note(doc, "Section 10 Commissioning and Provisional Acceptance accepted — PAC process, minor defects, site readiness all as requested.")
            i += 1
            continue

        # ─── FM 90-day negotiation: Query ───
        if i == 307 and 'ninety (90) days' in text:
            copy_para(doc, text)
            add_red_comment(doc,
                "NEW CLAUSE — 90-DAY NEGOTIATION. We note this is a new addition. Please clarify the relationship "
                "between this 90-day negotiation clause and the 60-day Buyer Cure Right in Section 15.4(a) below. "
                "As drafted, the Buyer can cancel at 60 days but the Parties are also required to negotiate at 90 days — "
                "this creates a conflict. We suggest: the 60-day cure right (Section 15.4) takes precedence, and the "
                "90-day negotiation applies only if the Buyer has NOT exercised the 60-day cancellation right.")
            i += 1
            continue

        # ─── FM production preparation: Query ───
        if i == 309 and 'procurement and production preparation' in text:
            copy_para(doc, text)
            add_red_comment(doc,
                "NEW CLAUSE — PRE-ADVANCE PRODUCTION. We understand the commercial reality that the Seller may "
                "commence procurement before receiving the advance payment. However, this increases the non-recoverable "
                "cost exposure in a Force Majeure refund scenario (Section 15.4(b)). Please confirm: (1) The Buyer "
                "is not liable for any production costs incurred before the advance payment is received; (2) The "
                "'documented costs for materials already procured' in Section 15.4(b) are limited to costs incurred "
                "AFTER receipt of the advance payment, not before.")
            i += 1
            continue

        # ─── FM Buyer Cure: Accepted ───
        if i == 311 and 'Buyer Cure Right' in text:
            copy_para(doc, text)
            add_green_note(doc, "Force Majeure Buyer Cure Right accepted — 60-day cure, 30-day refund, 14-day evidence.")
            i += 1
            continue

        # ─── Confidentiality: Accepted ───
        if i == 331 and 'confidentiality clause shall survive' in text:
            copy_para(doc, text)
            add_green_note(doc, "Confidentiality clause corrected — 'Republic of Poland' removed, survivorship clause added. Accepted.")
            i += 1
            continue

        # ─── Insurance Section 16: Accepted ───
        if i == 332 and 'Insurance Requirements' in text:
            copy_para(doc, text)
            add_green_note(doc, "Insurance Section 16 accepted — Seller EUR 5M PL + EUR 2M PI, Buyer EUR 1M each, proof of insurance. NOTE: We have received the AXA policy terms and will provide verification comments separately.")
            i += 1
            continue

        # ─── Product Compliance 20.5: Accepted ───
        if i == 378 and 'Product Standards Compliance' in text:
            copy_para(doc, text)
            add_green_note(doc, "Product Standards Compliance (IEC 62619, IEC 62477, CE, DSO grid codes) accepted.")
            i += 1
            continue

        # ─── Default: copy paragraph ───
        copy_para(doc, text, bold=is_bold)
        i += 1

    # Summary page
    doc.add_page_break()
    summary_title = doc.add_paragraph()
    add_run(summary_title, "SUMMARY OF FEEDBACK — 9 MARCH 2026 CONTRACT", RED, bold=True, size=Pt(14))
    summary_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_run(doc.add_paragraph(), "ITEMS ACCEPTED (no further changes needed):", GREEN, bold=True, size=Pt(12))
    for item in [
        "Art. 1A.3 — Affiliate coverage (anti-circumvention covered by Distribution Agreement)",
        "Art. 5A — Installation scope (a)-(f) complete, including workmanship warranty",
        "Art. 5C — No Cross-Liability restored",
        "Art. 6 — End-customer payment auto-extension",
        "Art. 8D — Performance Warranties (SOH 85/79.58/70%, RTE 86.32%, 7,000 cycles, PCS 98%)",
        "Art. 9 — Tiered liability cap (10%/50%/100%)",
        "Art. 9A — Indemnities",
        "Art. 9B.4 — Condition Precedent (APG before advance payment)",
        "Art. 9B.5 — Buyer Financial Standing (credit check response)",
        "Art. 9C Exclusions (a)-(c) — Delay LD exclusions",
        "Art. 8C.2 — No Set-Off for LTSA matters",
        "Section 10 — Commissioning and Provisional Acceptance",
        "Art. 15.4 — Force Majeure Buyer Cure Right (60 days, deposit refund)",
        "Art. 16 — Confidentiality (Poland reference fixed)",
        "Section 16 — Insurance Requirements (Seller + Buyer obligations)",
        "Art. 18 — Pre-Shipment Inspection & FAT",
        "Art. 19 — Data Protection / GDPR",
        "Art. 20.4-20.5 — Product Compliance Standards",
    ]:
        p = doc.add_paragraph()
        add_run(p, f"  ✓  {item}", GREEN, size=Pt(10))

    doc.add_paragraph()
    add_run(doc.add_paragraph(), "ITEMS REQUIRING COMPLETION:", RED, bold=True, size=Pt(12))
    for item in [
        "Art. 6(a) — PAYMENT PERCENTAGES BLANK — must be populated (propose 30/50/10/10)",
        "Art. 6 — LATE PAYMENT INTEREST RATE MISSING — insert EU Directive 2011/7/EU rate (ECB+8%)",
        "Art. 9B.1 — APG TERMS BLANK — insert 100% bank guarantee terms",
        "Art. 9B.2 — PERFORMANCE BOND BLANK — insert 5% corporate bank guarantee terms",
        "Art. 9B.3 — AXA INSURANCE CLAUSE BLANK — insert EUR 5M product liability + certificate obligation",
        "Art. 9C — DELAY LD DEFINITIONS AND RATES MISSING — insert 0.1%/0.15%/0.2% + 10% cap + Critical Path definitions",
    ]:
        p = doc.add_paragraph()
        add_run(p, f"  ✗  {item}", RED, size=Pt(10))

    doc.add_paragraph()
    add_run(doc.add_paragraph(), "ITEMS REQUIRING CLARIFICATION:", RED, bold=True, size=Pt(12))
    for item in [
        "Art. 8D — SOC < 5% voiding condition: permanent void vs. period-specific? Per-container or per-system?",
        "Art. 9C Exclusion (d) — 'any other causes attributable to the Buyer' too broad — propose narrower wording",
        "Art. 15 — 90-day negotiation vs. 60-day cure right — clarify precedence",
        "Art. 15 — Pre-advance production costs — confirm not recoverable from Buyer in FM refund",
    ]:
        p = doc.add_paragraph()
        add_run(p, f"  ?  {item}", RED, size=Pt(10), italic=True)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    add_run(footer, "Lighthief Cyprus Ltd | HE 477423 | office@lighthief.com | +357 77 77 00 50", BLACK, italic=True, size=Pt(8))
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_FILE)
    print(f"Feedback document saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    generate()
