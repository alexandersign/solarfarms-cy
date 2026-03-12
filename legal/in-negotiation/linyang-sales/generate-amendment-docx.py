"""
Generate Linyang Sales Contract Amendment DOCX
Reads the 6 March 2026 contract (rev260306) and produces an amendment document
with original text in black and all Lighthief proposed changes in red.
"""

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(SCRIPT_DIR, "rev260306-Linyang sales contract - Lighthief - V1 (1).docx")
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "rev260306-Linyang-sales-contract-Lighthief-AMENDMENT.docx")

RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)


def add_run(para, text, color=BLACK, bold=False, italic=False, size=Pt(11)):
    run = para.add_run(text)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = size
    return run


def add_black_para(doc, text, bold=False, size=Pt(11)):
    para = doc.add_paragraph()
    add_run(para, text, BLACK, bold=bold, size=size)
    return para


def add_red_para(doc, text, bold=False, size=Pt(11)):
    para = doc.add_paragraph()
    add_run(para, text, RED, bold=bold, size=size)
    return para


def add_red_heading(doc, text):
    para = doc.add_paragraph()
    add_run(para, text, RED, bold=True, size=Pt(12))
    return para


def add_amended_para(doc, original_text, new_text):
    """Original in black, then amendment marker + new text in red."""
    para = doc.add_paragraph()
    add_run(para, original_text, BLACK)
    add_run(para, "  [LIGHTHIEF AMENDMENT: ", RED, bold=True)
    add_run(para, new_text, RED)
    add_run(para, "]", RED, bold=True)
    return para


def copy_original_para(doc, text, bold=False):
    if not text.strip():
        doc.add_paragraph()
        return
    add_black_para(doc, text, bold=bold)


# ─── INSERTION BLOCKS ────────────────────────────────────────────────

def insert_anti_circumvention(doc):
    add_red_heading(doc, "[LIGHTHIEF INSERT — New Section 1A.4]")
    add_red_para(doc, "1A.4 Anti-Circumvention:", bold=True)
    add_red_para(doc, (
        "The Seller shall not, and shall procure that its Affiliates, subsidiaries, parent companies, "
        "and any related entities shall not, directly or indirectly:"
    ))
    add_red_para(doc, (
        "(a) enter into any sales, supply, distribution, service, or maintenance contract with any "
        "end-customer introduced by the Buyer within the Territory (Republic of Cyprus);"
    ))
    add_red_para(doc, (
        "(b) solicit, approach, or accept orders from any end-customer with whom the Buyer has an "
        "existing or pending commercial relationship within the Territory; or"
    ))
    add_red_para(doc, (
        "(c) supply Commodities or services to any third party for resale or installation within the "
        "Territory, where such supply would circumvent the Buyer's exclusive distribution rights "
        "under the Distribution Agreement."
    ))
    add_red_para(doc, (
        "This obligation shall survive termination of this Contract for a period of two (2) years "
        "and is without prejudice to the broader exclusivity and client protection provisions in "
        "the Distribution Agreement, which shall prevail in the event of any conflict."
    ))


def insert_installation_items_ef(doc):
    add_red_para(doc, "(e) Installation supervision and quality assurance;")
    add_red_para(doc, "(f) Installation workmanship warranty (5 years from PAC).")


def insert_cross_liability(doc):
    add_red_heading(doc, "[LIGHTHIEF INSERT — New Section 5C — No Cross-Liability]")
    add_red_para(doc, "No Cross-Liability:", bold=True)
    add_red_para(doc, "(a) The Seller is NOT liable for defects in the Buyer's EPC works;")
    add_red_para(doc, "(b) The Buyer is NOT liable for defects in the Seller's installation and commissioning.")


def insert_commissioning_pac(doc):
    add_red_heading(doc, "[LIGHTHIEF INSERT — New Section 10 — COMMISSIONING AND PROVISIONAL ACCEPTANCE]")
    add_red_para(doc, "10.1 Commissioning:", bold=True)
    add_red_para(doc, (
        "Commissioning shall be performed in accordance with the Seller's standard commissioning "
        "procedures and technical documentation, following completion of installation and site "
        "readiness confirmation by the Buyer."
    ))
    add_red_para(doc, "10.2 Provisional Acceptance:", bold=True)
    add_red_para(doc, (
        "Upon successful completion of commissioning and demonstration that the Products operate "
        "in accordance with agreed specifications, the Parties shall execute a Provisional "
        'Acceptance Certificate ("PAC").'
    ))
    add_red_para(doc, "10.3 Effect of Provisional Acceptance:", bold=True)
    add_red_para(doc, "Provisional Acceptance shall:")
    add_red_para(doc, "(a) confirm completion of installation and commissioning;")
    add_red_para(doc, "(b) trigger the Buyer's obligation to make the holdback payment under Article 6(d); and")
    add_red_para(doc, "(c) mark the commencement of the base warranty period under Article 8.")
    add_red_para(doc, "10.4 Minor Defects:", bold=True)
    add_red_para(doc, (
        "Minor defects that do not materially affect system operation shall not prevent Provisional "
        "Acceptance, provided such defects are recorded in a punch list and remedied by the Seller "
        "within a reasonable time."
    ))
    add_red_para(doc, "10.5 Site Readiness:", bold=True)
    add_red_para(doc, (
        "The Buyer shall ensure that the site is ready for installation in accordance with the "
        "Seller's technical requirements, including:"
    ))
    add_red_para(doc, "(a) Completed foundations and civil works;")
    add_red_para(doc, "(b) AC connection point ready for termination;")
    add_red_para(doc, "(c) Grounding system installed;")
    add_red_para(doc, "(d) Clear access for equipment and personnel.")
    add_red_para(doc, (
        "Delays arising from site unavailability or non-compliance shall automatically extend "
        "the installation and commissioning schedule without penalty to the Seller."
    ))


def insert_late_payment_framework(doc):
    add_red_heading(doc, "[LIGHTHIEF AMENDMENT — Replace existing 'no default' clause with Late Payment Framework]")
    add_red_para(doc, "Late Payment and End-Customer Pass-Through:", bold=True)
    add_red_para(doc, (
        "(a) The Buyer's payment obligations under this Contract are absolute and are not "
        "conditional upon the Buyer's receipt of payment from any end-customer."
    ))
    add_red_para(doc, (
        "(b) Where the Buyer fails to make any payment under Article 6(b)-(d) by the due date, "
        "interest shall accrue at the rate prescribed by EU Directive 2011/7/EU (ECB main "
        "refinancing rate + 8% per annum, simple interest), regardless of whether the delay "
        "is caused by the Buyer or by the Buyer's end-customer."
    ))
    add_red_para(doc, (
        "(c) Where late payment occurs, the Seller's delivery and commissioning obligations shall "
        "be extended day-for-day by the period of such late payment, without penalty to the Seller."
    ))
    add_red_para(doc, (
        "(d) If any payment under Article 6(b)-(d) remains unpaid for more than sixty (60) calendar "
        "days after the due date, the Seller may terminate this Contract in accordance with Article 6(g)."
    ))
    add_red_para(doc, (
        "(e) The Buyer shall include equivalent late payment and default provisions in its downstream "
        "contracts with end-customers, ensuring that EU Directive 2011/7/EU statutory interest rates "
        "and sixty (60) day default triggers are mirrored in the Buyer-Client agreement."
    ))


def insert_payment_late_default(doc):
    add_red_heading(doc, "[LIGHTHIEF INSERT — Late Payment and Default Provisions — Article 6(e)-(h)]")
    add_red_para(doc, "Late Payment:", bold=True)
    add_red_para(doc, (
        "(e) Where the Buyer fails to make any payment under Article 6(b)-(d) by the due date, "
        "the Buyer shall pay interest on the overdue amount at the rate prescribed by EU Directive "
        "2011/7/EU (the European Central Bank main refinancing rate plus eight percent (8%) per annum), "
        "calculated on a simple interest basis from the day following the due date until the date "
        "of actual payment."
    ))
    add_red_para(doc, (
        "(f) Where the Buyer's payment under Article 6(b)-(d) is delayed, the Seller's delivery "
        "and commissioning obligations shall be extended day-for-day by the period of such late "
        "payment, without penalty to the Seller."
    ))
    add_red_para(doc, "Default:", bold=True)
    add_red_para(doc, (
        "(g) If any payment under Article 6(b)-(d) remains unpaid for more than sixty (60) calendar "
        "days after the due date, the Seller may terminate this Contract by written notice. Upon such "
        "termination, the advance payment made under Article 6(a) shall be non-refundable and shall "
        "be retained by the Seller as agreed liquidated damages for the Buyer's default."
    ))
    add_red_para(doc, (
        "(h) For the avoidance of doubt, the advance payment under Article 6(a) is a pre-payment "
        "and no late payment interest or default provisions apply to it, as it is due and payable "
        "prior to the commencement of any Seller obligations."
    ))


def insert_performance_warranties(doc):
    add_red_heading(doc, "[LIGHTHIEF INSERT — New Section 8D — PRODUCT PERFORMANCE WARRANTIES]")
    add_red_para(doc, "8D.1 State of Health (SOH) Guarantees:", bold=True)
    add_red_para(doc, (
        "The Seller warrants that the battery cells and modules, when operated in accordance with "
        "the Seller's guidelines (1 cycle/day average, 90% DoD, 25±2°C), shall retain the following "
        "minimum State of Health:"
    ))
    add_red_para(doc, "(a) Year 5: ≥ 85% SOH;")
    add_red_para(doc, "(b) Year 10: ≥ 79.58% SOH;")
    add_red_para(doc, "(c) Year 15: ≥ 70% SOH.")
    add_red_para(doc, (
        "SOH remedy: The Seller shall supply replacement cells/modules (FOB or CIF Limassol as "
        "applicable) to restore capacity above the guaranteed threshold. Replacement parts are at "
        "Seller's cost; installation labour is at Buyer's cost unless an LTSA provides otherwise."
    ))
    add_red_para(doc, "8D.2 Round-Trip Efficiency (RTE):", bold=True)
    add_red_para(doc, (
        "The Seller warrants that the full system AC-AC Round-Trip Efficiency, including PCS "
        "conversion, auxiliary consumption, and internal DC/AC cabling losses within the Seller's "
        "supply scope, shall be not less than 86.32% at rated power under standard test conditions "
        "(0.5C charge / 0.5C discharge, 25±2°C)."
    ))
    add_red_para(doc, (
        "This warranty is conditional upon the Buyer's external MV cabling (from MV Skid to Point "
        "of Common Coupling) conforming to the cable specifications and maximum allowable loss "
        "parameters set out in the Seller's system design drawings. The Seller shall provide the "
        "Buyer with the applicable cable loss budget and specifications within fourteen (14) days "
        "of contract signing. Any RTE shortfall attributable solely to Buyer-supplied cables "
        "exceeding the specified loss ratings shall not constitute a breach of this warranty."
    ))
    add_red_para(doc, "8D.3 Cycle Life:", bold=True)
    add_red_para(doc, (
        "The Seller warrants a minimum cycle life of 7,000 equivalent full cycles at a standard "
        "charge/discharge rate of 0.5C/0.5C, at 90% Depth of Discharge (DoD) to 70% End of Life "
        "(EOL), at 25±2°C, under the operating conditions specified in the Seller's Warranty "
        "Terms (Version 2)."
    ))
    add_red_para(doc, (
        "For the avoidance of doubt, the maximum charge/discharge rate is 1C. Operation at rates "
        "exceeding 0.5C may affect achievable cycle life and shall be governed by the degradation "
        "curves provided in the Seller's Warranty Terms V2."
    ))
    add_red_para(doc, "8D.4 PCS Efficiency:", bold=True)
    add_red_para(doc, "The Seller warrants PCS conversion efficiency of not less than 98% at rated power.")
    add_red_para(doc, "8D.5 Remedies:", bold=True)
    add_red_para(doc, (
        "If the Products fail to meet the performance thresholds in §8D.1–8D.4, the Seller's "
        "obligation shall be limited to repair or replacement of defective components in accordance "
        "with Article 8 (Warranty). For the avoidance of doubt, availability LDs and revenue-based "
        "compensation remain governed exclusively by any applicable LTSA under Article 8C."
    ))
    add_red_para(doc, "8D.6 Reference Documents:", bold=True)
    add_red_para(doc, (
        "The following Seller documents form part of the performance warranty and are incorporated "
        "by reference:"
    ))
    add_red_para(doc, "(a) Warranty Terms — Version 2;")
    add_red_para(doc, "(b) Battery Container Specification (3.34 MWh);")
    add_red_para(doc, "(c) PCS Datasheets (BCS-C-HUD series);")
    add_red_para(doc, "(d) RTE Calculation Report (including cable loss budget assumptions).")


def insert_financial_guarantees(doc):
    add_red_heading(doc, "[LIGHTHIEF INSERT — New Section 9B — FINANCIAL GUARANTEES]")
    add_red_para(doc, "9B.1 Advance Payment Guarantee (APG):", bold=True)
    add_red_para(doc, (
        'The Seller shall provide an unconditional, irrevocable bank guarantee ("APG") equal to '
        "one hundred percent (100%) of the advance payment, issued by an internationally recognised "
        "bank acceptable to the Buyer. The APG shall:"
    ))
    add_red_para(doc, "(a) be delivered to the Buyer prior to or simultaneously with the advance payment;")
    add_red_para(doc, "(b) remain valid until full delivery of the Commodities plus thirty (30) days;")
    add_red_para(doc, "(c) be callable on demand in the event of Seller default or failure to deliver.")
    add_red_para(doc, "9B.2 Performance Bond:", bold=True)
    add_red_para(doc, (
        "The Seller shall provide a performance bond equal to five percent (5%) of the total "
        "contract value, issued as a corporate guarantee backed by a bank. The performance bond shall:"
    ))
    add_red_para(doc, "(a) be delivered within fourteen (14) days of receipt of the advance payment under Article 6(a);")
    add_red_para(doc, "(b) remain valid until final acceptance of the Commodities;")
    add_red_para(doc, "(c) cover defects in materials, workmanship, and non-performance during the warranty period.")
    add_red_para(doc, "9B.3 Product Liability Insurance (AXA):", bold=True)
    add_red_para(doc, (
        "The Seller shall maintain product liability insurance with a minimum coverage of "
        "EUR 5,000,000 per occurrence. The Seller has confirmed that such insurance is currently "
        "held with AXA. The Seller represents and warrants that:"
    ))
    add_red_para(doc, "(a) the policy covers products supplied under this Contract for installation in the Republic of Cyprus;")
    add_red_para(doc, "(b) coverage shall be maintained for the duration of this Contract and the warranty period;")
    add_red_para(doc, (
        "(c) the Seller shall provide a copy of the AXA insurance certificate to the Buyer within "
        "fourteen (14) days of contract signing and upon each policy renewal."
    ))
    add_red_para(doc, "9B.4 Condition Precedent:", bold=True)
    add_red_para(doc, (
        "The Buyer's obligation to make the advance payment under Article 6(a) shall not become "
        "due until the Seller delivers the APG under Section 9B.1."
    ))
    add_red_para(doc, "9B.5 Buyer Financial Standing:", bold=True)
    add_red_para(doc, (
        "Under the agreed payment structure, eighty to ninety percent (80-90%) of the contract value "
        "is prepaid before delivery. The Seller's credit exposure to the Buyer is limited to the ten "
        "percent (10%) holdback under Article 6(d). Notwithstanding the above, the Buyer shall, upon "
        "reasonable request:"
    ))
    add_red_para(doc, "(a) provide a bank reference letter or confirmation of good standing from its principal banking institution; and/or")
    add_red_para(doc, "(b) provide evidence of funded end-customer contracts supporting each Purchase Order placed under this Contract.")
    add_red_para(doc, (
        "For the avoidance of doubt, Lighthief Cyprus Ltd (Company No. HE 477423) is a wholly-owned "
        "subsidiary of Lighthief International Ltd. A letter of support from the parent entity may be "
        "provided upon request."
    ))


def insert_delay_lds(doc):
    add_red_heading(doc, "[LIGHTHIEF INSERT — New Section 9C — DELIVERY DELAY LIQUIDATED DAMAGES]")
    add_red_para(doc, "9C.0 Definitions for LD Purposes:", bold=True)
    for defn in [
        ('"Critical Path Equipment"', "Equipment whose delay prevents commissioning of the entire order or park. Includes: Power Conversion System (PCS), Medium-Voltage Switchgear/Skid (MV Skid), and main step-up transformer. Delay to any Critical Path Equipment constitutes a Full Order Delay."),
        ('"Non-Critical Path Equipment"', "Equipment whose delay prevents commissioning of only a portion of the total ordered capacity. Battery containers (BESS units) are Non-Critical Path Equipment when delivered individually as part of a multi-container order, provided the PCS and MV infrastructure are available to commission the remaining containers."),
        ('"Full Order Delay"', "Delay caused by late delivery, defect, or unavailability of any Critical Path Equipment, resulting in no portion of the order being capable of commissioning."),
        ('"Partial Delay"', "Delay caused by late delivery, defect, or unavailability of Non-Critical Path Equipment, where the remaining equipment is capable of commissioning at reduced capacity."),
        ('"Affected Capacity"', "The proportion of total ordered energy capacity (MWh) that cannot be commissioned due to the delayed equipment, expressed as a percentage of total capacity."),
        ('"LD Value"', "For a Full Order Delay: the total order value (all goods). For a Partial Delay: the value of the delayed goods only (= order value × Affected Capacity %)."),
    ]:
        para = doc.add_paragraph()
        add_run(para, defn[0], RED, bold=True)
        add_run(para, " — " + defn[1], RED)

    add_red_para(doc, "9C.1 Delay Liquidated Damages:", bold=True)
    add_red_para(doc, "(a) Days 1–30 of delay: 0.1% of LD Value per day;")
    add_red_para(doc, "(b) Days 31–60 of delay: 0.15% of LD Value per day;")
    add_red_para(doc, "(c) Days 61+ of delay: 0.2% of LD Value per day;")
    add_red_para(doc, "(d) Cap: aggregate delay LDs shall not exceed 10% of LD Value per order.")
    add_red_para(doc, "9C.2 Application:", bold=True)
    add_red_para(doc, "(a) PCS or MV Skid delayed → Full Order Delay → LD Value = total order value.")
    add_red_para(doc, "(b) Individual battery container delayed (others can commission) → Partial Delay → LD Value = value of delayed container(s) only.")
    add_red_para(doc, "(c) Delay LDs begin on the day following the agreed delivery date and accrue daily until actual delivery or termination.")
    add_red_para(doc, "9C.3 Exclusions:", bold=True)
    add_red_para(doc, "Delay LDs shall not apply where the delay is caused by:")
    add_red_para(doc, "(a) Buyer's failure to pay on time;")
    add_red_para(doc, "(b) Force Majeure under Article 15;")
    add_red_para(doc, "(c) Buyer's failure to provide site readiness as required.")
    add_red_para(doc, "9C.4 Total Combined LD Cap:", bold=True)
    add_red_para(doc, (
        "The aggregate of all liquidated damages under this Contract (delay + availability, if "
        "applicable) shall not exceed ten percent (10%) of the total contract value."
    ))


# ─── MAIN GENERATION ─────────────────────────────────────────────────

def generate():
    src = Document(INPUT_FILE)
    doc = Document()

    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Calibri'

    title = doc.add_paragraph()
    add_run(title, "SALES CONTRACT — LIGHTHIEF AMENDMENT MARKUP", BLACK, bold=True, size=Pt(14))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    add_run(sub, "Base document: rev260306 — Linyang sales contract - Lighthief - V1 (6 March 2026)", BLACK, italic=True, size=Pt(9))
    add_run(sub, "\n", BLACK, size=Pt(9))
    add_run(sub, "Prepared by: Lighthief Cyprus Ltd — Legal & Commercial", BLACK, italic=True, size=Pt(9))
    add_run(sub, "\n", BLACK, size=Pt(9))
    add_run(sub, "All text in BLACK is the original contract. ", BLACK, size=Pt(9))
    add_run(sub, "All text in RED represents Lighthief proposed amendments and insertions.", RED, bold=True, size=Pt(9))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    paras = src.paragraphs
    i = 0
    while i < len(paras):
        text = paras[i].text
        is_bold = any(r.bold for r in paras[i].runs if r.bold is not None)

        # ─── COMMENT 2: After 1A.3, insert 1A.4 anti-circumvention ───
        if i == 28:
            copy_original_para(doc, text, bold=is_bold)
            insert_anti_circumvention(doc)
            i += 1
            continue

        # ─── COMMENT 3: Fix blank installation scope item (a) ───
        if i == 62 and text.strip() == "(a) ;":
            add_amended_para(doc, "(a) ;",
                             "(a) Equipment installation and assembly by factory-trained technicians;")
            i += 1
            continue

        # ─── COMMENT 3: After item (d), insert (e) and (f) ───
        if i == 65 and "OEM testing protocols" in text:
            copy_original_para(doc, text)
            insert_installation_items_ef(doc)
            i += 1
            continue

        # ─── COMMENT 4: After Buyer scope (f), insert cross-liability ───
        if i == 75 and "Coordination with local authorities" in text:
            copy_original_para(doc, text)
            insert_cross_liability(doc)
            insert_commissioning_pac(doc)  # COMMENT 12: Section 10 after scope
            i += 1
            continue

        # ─── COMMENT 1: Payment milestone (a) — fill in 30% ───
        if i == 82 and "% advance payment" in text:
            add_amended_para(doc, text,
                "(a) Thirty percent (30%) advance payment by T/T within seven (7) days "
                "of contract effectiveness or purchase order signing;")
            i += 1
            continue

        # ─── COMMENT 1: After payment milestone (a), insert (b)-(d) ───
        # Paragraphs 83-85 are mostly blank in the original; para 86 is the suspension clause
        if i == 83:
            add_red_para(doc, (
                "(b) Fifty percent (50%) payment by T/T when the Products are ready for shipment "
                "(Ex-Works) from the manufacturing facility, upon Seller's written confirmation of "
                "successful Factory Acceptance Testing (FAT) and photographic evidence of readiness;"
            ))
            add_red_para(doc, (
                "(c) Ten percent (10%) payment by T/T upon arrival of the Products on site "
                "(CIF Limassol) and completion of port clearance, following Buyer confirmation "
                "that the Commodities conform to the applicable Quotation;"
            ))
            add_red_para(doc, (
                "(d) Ten percent (10%) performance holdback by T/T upon successful completion of "
                "Site Acceptance Testing (SAT) and Provisional Acceptance of the Products on site."
            ))
            # Skip blank paragraphs 83-85
            while i < 86:
                i += 1
            continue

        # ─── COMMENT 5: Replace "no default" clause ───
        if i == 90 and "Buyer shall not be deemed in default" in text:
            para = doc.add_paragraph()
            add_run(para, text, BLACK)
            add_run(para, "\n\n[LIGHTHIEF AMENDMENT — Replace the above clause with:]", RED, bold=True)
            insert_late_payment_framework(doc)
            i += 1
            continue

        # ─── COMMENT 6: Interest rate amendment ───
        if i == 94 and "2% per month" in text:
            add_amended_para(doc, text,
                "For contracts where the end destination is Cyprus, the interest rate on late "
                "payments shall be the rate prescribed by EU Directive 2011/7/EU on combating "
                "late payment in commercial transactions — being the European Central Bank (ECB) "
                "main refinancing rate plus eight percent (8%) per annum, calculated on a simple "
                "interest basis. This rate shall apply to all late payment obligations under this "
                "Article 6 and shall be the sole interest rate applicable to overdue amounts.")
            i += 1
            continue

        # ─── COMMENT 1: Insert late payment + default 6(e)-(h) after interest clause ───
        if i == 95 and text.strip() == ".":
            copy_original_para(doc, text)
            insert_payment_late_default(doc)
            i += 1
            continue

        # ─── COMMENT 8: After LTSA Separation, insert performance warranties 8D ───
        if i == 152 and "Availability guarantees" in text:
            copy_original_para(doc, text)
            insert_performance_warranties(doc)
            i += 1
            continue

        # ─── COMMENT 11: Liability cap — amend 10% to 100% ───
        if i == 162 and "ten percent of total charges" in text:
            add_amended_para(doc, text,
                "Subject to other clauses in this section, the total aggregate liability of the "
                "SELLER for all breaches of duty under this Contract shall not exceed one hundred "
                "percent (100%) of the total price paid by the BUYER to the SELLER for the Products "
                "giving rise to the claim. For the avoidance of doubt, this general liability cap is "
                "without prejudice to the warranty-specific liability cap in Section 8B, which shall "
                "apply to warranty claims independently.")
            i += 1
            continue

        # ─── COMMENTS 9+10: After 9A indemnities, insert 9B + 9C ───
        if i == 184 and "Third-party systems not supplied" in text:
            copy_original_para(doc, text)
            insert_financial_guarantees(doc)
            insert_delay_lds(doc)
            i += 1
            continue

        # ─── COMMENT 13: Force Majeure cure — complete the blank ───
        if i == 260 and "exceeding () days" in text:
            add_amended_para(doc, text,
                "(a) If a Force Majeure event affecting the Seller's delivery obligations "
                "continues for a period exceeding sixty (60) days, the Buyer may, upon written "
                "notice, cancel the affected order without liability.")
            add_red_para(doc, (
                "(b) Upon such cancellation, the Seller shall refund the Buyer's advance payment "
                "within thirty (30) days, less any documented costs for materials already procured "
                "and non-recoverable."
            ))
            add_red_para(doc, (
                "(c) The Seller shall provide reasonable evidence of claimed non-recoverable costs "
                "within fourteen (14) days of the cancellation notice."
            ))
            i += 1
            continue

        # ─── COMMENT 15: Confidentiality — Poland to Cyprus/Singapore ───
        if i == 279 and "Republic of Poland" in text:
            add_amended_para(doc, text,
                "This confidentiality clause shall survive termination of the Contract for a period "
                "of five (5) years, unless the confidential information enters the public domain "
                "through no fault of the receiving Party or is required to be disclosed under "
                "applicable law or regulation, including the laws of the Republic of Cyprus and Singapore.")
            i += 1
            continue

        # ─── Default: copy original paragraph ───
        copy_original_para(doc, text, bold=is_bold)
        i += 1

    doc.save(OUTPUT_FILE)
    print(f"Amendment document saved to: {OUTPUT_FILE}")
    print(f"Total paragraphs processed: {len(paras)}")


if __name__ == "__main__":
    generate()
