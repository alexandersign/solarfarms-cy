"""Shared helpers and park data for Esperia / Galascope LOI generators."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x40, 0x40, 0x40)
AMBER_TXT = RGBColor(0x78, 0x35, 0x0F)
AMBER_HDR = RGBColor(0x92, 0x40, 0x0E)
NAVY_HEX = "1A365D"
LIGHT_HEX = "EBF0F7"
AMBER_HEX = "FEF3C7"

REF_LITHIUM_CNY = (
    "155,000 CNY/tonne (Mysteel China battery-grade spot, January 2026 monthly average)"
)
REF_EUR_CNY = "8.18 CNY per EUR (January 2026 average)"
# EPC §6.1 bands: lithium >10% from ref; EUR/CNY >8% from ref (Jan 2026 baseline).
# May 2026 market (internal — disclosed in LOI §4.4): lithium ~183k CNY/t (~+18% vs Jan);
# EUR/CNY ~7.88 (ECB 22 May 2026, within 8% band). Linyang lock if LOI+EPC by 29 May 2026.
PRICE_LOCK_DEADLINE = "29 May 2026"

# Same authorised representative (Galascope Ltd + Esperia Energy Group)
CLIENT_REP_NAME = "Ntinos Konstantinos"
CLIENT_REP_TITLE_GALASCOPE = "Director"
CLIENT_REP_TITLE_ESPERIA = "Owner"

ESPERIA_CLIENT_ADDRESS = "Christaki Kranou 18, Limassol 4046, Cyprus"

# Galascope batch — negotiated & presented to Dino (May 2026). LOCKED; never apply to other parks.
GALASCOPE_LOCKED_EUR_PER_MWH = 111_900  # G1 €2,238,000 = 111,900 × 20 MWh nameplate
GALASCOPE_G1_LOCKED_TOTAL = 2_238_000
GALASCOPE_G2_LOCKED_TOTAL = 1_206_300  # 2.5/10 — separate negotiated line (not 111,900/MWh)

# Tier 1: 4×BESS + 1×MV @ ~20 MWh (≤5 MW PCS class — 4.75/20, 5/20). NOT Galascope €/MWh.
STANDARD_4BESS_1MV_20MWh_CLIENT_EUR_PER_MWH = 115_322  # LY Famagusta 2 5/20 + v5 adders @ 11.5%
STANDARD_4BESS_1MV_20MWh_CIF_REF = "Esperia Famagusta 2"  # LY202601271 — 4 BESS + 1 MV, 5 cnt
STANDARD_5_20_CLIENT_EUR_PER_MWH = STANDARD_4BESS_1MV_20MWh_CLIENT_EUR_PER_MWH
STANDARD_5_20_CIF_REF = STANDARD_4BESS_1MV_20MWh_CIF_REF

# Tier 2: 5×BESS + 2×MV @ 25 MWh (≥6 MW needs dual MV skids — 6.3/25, 7/25, 7.5/25 same €/MWh).
STANDARD_2MV_25MWh_CLIENT_EUR_PER_MWH = 122_929  # LY Tseri 2-B 7.5/25 + v5 adders @ 11.5%
STANDARD_2MV_25MWh_CIF_REF = "Esperia Tseri 2b"  # LY202601271 — 5 BESS + 2 MV, 7 cnt

# 6/20 class: 4×BESS + 1×MV @ 20 MWh (6.0–6.5 MW / T8 skid — same €/MWh as Tseri 2-C).
STANDARD_6_20_CLIENT_EUR_PER_MWH = 118_746  # LY Esperia Tseri 2c 6/20 + v5 adders @ 11.5%
STANDARD_6_20_CIF_REF = "Esperia Tseri 2c"

# Small park 2.5/10 class (Galascope G2 negotiated rate).
GALASCOPE_G2_EUR_PER_MWH = 120_630  # €1,206,300 ÷ 10 MWh

MW_THRESHOLD_DUAL_MV = 6.0


def standard_6_20_client_price(mwh: float) -> int:
    """6.0–6.5 MW / 20 MWh, 4 BESS + 1 MV (incl. Famagusta Ph1 T8 skid)."""
    return round(STANDARD_6_20_CLIENT_EUR_PER_MWH * mwh)


def galascope_g1_client_price(mwh: float) -> int:
    """Galascope G1 locked rate — also Tseri 3 (relationship pricing)."""
    return round(GALASCOPE_LOCKED_EUR_PER_MWH * mwh)


def galascope_g2_client_price(mwh: float) -> int:
    """Galascope G2 / 2.5/10 class — Tseri 2-A."""
    return round(GALASCOPE_G2_EUR_PER_MWH * mwh)


def standard_4bess_1mv_client_price(mwh: float) -> int:
    """≤5 MW / 4+1 MV / ~20 MWh class (excl. Galascope locked rate)."""
    return round(STANDARD_4BESS_1MV_20MWh_CLIENT_EUR_PER_MWH * mwh)


def standard_2mv_25mwh_client_price(mwh: float) -> int:
    """≥6 MW with 2 MV skids — same €/MWh basis as 7.5/25 (excl. extra MWh modules)."""
    return round(STANDARD_2MV_25MWh_CLIENT_EUR_PER_MWH * mwh)


def client_rep_line(title: str) -> str:
    return f"represented by: {CLIENT_REP_NAME} ({title})\n\n"


GALASCOPE = [
    {
        "name": "Galascope 1",
        "district": "Famagusta",
        "mw": 5.0,
        "mwh": 20,
        "effective_mwh": 20.06,
        "containers": 5,
        "client_price": 2_238_000,
    },
    {
        "name": "Galascope 2",
        "district": "Famagusta",
        "mw": 2.5,
        "mwh": 10,
        "effective_mwh": 10.03,
        "containers": 3,
        "client_price": 1_206_300,
    },
]

# Footnote text for parks not exactly quoted at final MW/MWh (LY202601271).
PRICE_BASIS_LY_EXACT = "LY202601271 — exact configuration quoted."
PRICE_BASIS_2MV_25_TIER = (
    "Nearest LY configuration: Esperia Tseri 2-B (7.5 MW / 25 MWh, 5 BESS + 2 MV)."
)
PRICE_BASIS_2MV_30 = (
    "Nearest LY configuration: Esperia Tseri 2-B (7.5 MW / 25 MWh, 5 BESS + 2 MV)."
)
PRICE_BASIS_6_20_CLASS = (
    "Nearest LY configuration: Esperia Tseri 2-C (6.0 MW / 20 MWh, 4 BESS + 1 MV)."
)
PRICE_BASIS_G2_10 = (
    "Nearest LY configuration: Galascope 2 / Dianary (2.5 MW / 10 MWh, 2 BESS + 1 MV)."
)
PRICE_BASIS_G1_20 = (
    "Nearest LY configuration: Galascope 1 / Famagusta 2 (5.0 MW / 20 MWh, 4 BESS + 1 MV)."
)
# Pipeline LOI — three batches by indicative delivery (Dino Schedule 1, May 2026).
PIPELINE = {
    "phase1": [
        {
            "name": "Esperia Energy (Tseri)",
            "district": "Nicosia",
            "mw": 7.2,
            "mwh": 25,
            "containers": 7,
            "client_price": standard_2mv_25mwh_client_price(25),
            "cif_ref": STANDARD_2MV_25MWh_CIF_REF,
            "price_estimate": True,
            "price_basis": PRICE_BASIS_2MV_25_TIER,
            "delivery": "2026",
            "licence": "Licence ~6 months",
        },
        {
            "name": "Esperia Energy (Tseri 2-A)",
            "district": "Nicosia",
            "mw": 2.75,
            "mwh": 10,
            "containers": 3,
            "client_price": galascope_g2_client_price(10),
            "cif_ref": "Galascope 2",
            "price_estimate": True,
            "price_basis": PRICE_BASIS_G2_10,
            "delivery": "Q3 2026",
            "licence": "Fully licensed",
        },
        {
            "name": "Esperia Energy (Tseri 3)",
            "district": "Nicosia",
            "mw": 4.75,
            "mwh": 20,
            "containers": 5,
            "client_price": galascope_g1_client_price(20),
            "cif_ref": "Galascope 1",
            "price_estimate": True,
            "price_basis": PRICE_BASIS_G1_20,
            "delivery": "Q3 2026",
            "licence": "Fully licensed",
        },
        {
            "name": "Esperia Energy (Famagusta)",
            "district": "Famagusta",
            "mw": 6.5,
            "mwh": 20,
            "containers": 5,
            "client_price": standard_6_20_client_price(20),
            "cif_ref": STANDARD_6_20_CIF_REF,
            "price_estimate": True,
            "price_basis": PRICE_BASIS_6_20_CLASS,
            "delivery": "Q4 2026",
            "licence": "Per EPC schedule",
        },
    ],
    "phase2": [
        {
            "name": "Esperia Green Energy (Limassol)",
            "district": "Limassol",
            "mw": 8.0,
            "mwh": 60,
            "containers": 12,
            "client_price": 5_644_044,
            "cif_ref": "Esperia Limassol",
            "price_basis": PRICE_BASIS_LY_EXACT,
            "delivery": "Q2 2027",
            "licence": "Per EPC schedule",
        },
        {
            "name": "Esperia Energy (Tseri 2-B)",
            "district": "Nicosia",
            "mw": 7.99,
            "mwh": 30,
            "containers": 7,
            "client_price": standard_2mv_25mwh_client_price(30),
            "cif_ref": STANDARD_2MV_25MWh_CIF_REF,
            "price_estimate": True,
            "price_basis": PRICE_BASIS_2MV_30,
            "delivery": "2027",
            "licence": "~9 months from licensing",
        },
        {
            "name": "Esperia Energy (Tseri 2-C)",
            "district": "Nicosia",
            "mw": 6.3,
            "mwh": 25,
            "containers": 7,
            "client_price": standard_2mv_25mwh_client_price(25),
            "cif_ref": STANDARD_2MV_25MWh_CIF_REF,
            "price_estimate": True,
            "price_basis": PRICE_BASIS_2MV_25_TIER,
            "delivery": "2027",
            "licence": "Licence ~6 months",
        },
    ],
    "phase3": [
        {
            "name": "Esperia Energy (Frenaros)",
            "district": "Famagusta",
            "mw": 25.0,
            "mwh": 100,
            "containers": 20,
            "client_price": 10_088_014,
            "cif_ref": "Esperia Frenaros",
            "price_basis": PRICE_BASIS_LY_EXACT,
            "delivery": "Q2 2028",
            "licence": "Per EPC schedule",
        },
        {
            "name": "Esperia Green Energy (Famagusta 2)",
            "district": "Famagusta",
            "mw": 7.0,
            "mwh": 25,
            "containers": 7,
            "client_price": standard_2mv_25mwh_client_price(25),
            "cif_ref": STANDARD_2MV_25MWh_CIF_REF,
            "price_estimate": True,
            "price_basis": PRICE_BASIS_2MV_25_TIER,
            "delivery": "Q2 2028",
            "licence": "Phase 2",
        },
    ],
}

PIPELINE_PHASE_ORDER = ("phase1", "phase2", "phase3")
PIPELINE_PHASE_LABELS = {
    "phase1": "Batch 1 — 2026",
    "phase2": "Batch 2 — 2027",
    "phase3": "Batch 3 — 2028",
}
PIPELINE_PHASE_EPC_TARGET = {
    "phase1": "2026 (per park)",
    "phase2": "2027 (per park)",
    "phase3": "Q2 2028",
}

PIPELINE_SCHEDULE_COLS = [3.8, 1.4, 1.0, 1.0, 1.8, 2.0, 2.0]
PIPELINE_SCHEDULE_RC = {2, 3, 6}


def fmt_eur(v):
    return f"\u20ac{v:,.0f}"


def pipeline_price_display(park: dict) -> str:
    s = fmt_eur(park["client_price"])
    if park.get("price_estimate"):
        return s + "*"
    return s


def iter_pipeline_parks():
    for phase_key in PIPELINE_PHASE_ORDER:
        for park in PIPELINE[phase_key]:
            yield park


def add_pipeline_price_footnotes(doc):
    """One consolidated Schedule 1 footnote for estimated parks."""
    estimated = [p for p in iter_pipeline_parks() if p.get("price_estimate")]
    if not estimated:
        return
    short_names = []
    for p in estimated:
        n = p["name"].replace("Esperia Green Energy (", "").replace("Esperia Energy (", "").rstrip(")")
        short_names.append(n)
    note(
        doc,
        "* Indicative Prices marked * use the nearest Linyang configuration (LY202601271) where "
        "MW/MWh is not an exact quoted line: "
        + ", ".join(short_names)
        + ". Limassol and Frenaros are exact LY lines.",
        space_after=4,
    )


def cm_to_twips(cm):
    return int(cm * 567)


def lock_table_widths(table, col_widths_cm):
    tbl = table._tbl
    tblPr = tbl.tblPr
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(cm_to_twips(w) for w in col_widths_cm)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    tblPr.append(lay)
    existing = tbl.find(qn("w:tblGrid"))
    if existing is not None:
        tbl.remove(existing)
    tblGrid = OxmlElement("w:tblGrid")
    for w_cm in col_widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(cm_to_twips(w_cm)))
        tblGrid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(col_widths_cm):
                break
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(cm_to_twips(col_widths_cm[i])))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_run(para, text, bold=False, italic=False, size_pt=10, color=BLACK, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    return run


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, bold=True, size_pt=12, color=GOLD)
    return p


def body(doc, text, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, text)
    return p


def note(doc, text, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, text, size_pt=9, color=GREY)
    return p


def bullet(doc, text, indent_cm=0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(indent_cm)
    add_run(p, text, size_pt=10)
    return p


def tbl_hdr(table, headers):
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, NAVY_HEX)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, hdr, bold=True, size_pt=8, color=WHITE)


def tbl_row(table, values, right_cols=None, total=False, shade=False):
    right_cols = right_cols or set()
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        if total:
            set_cell_bg(cell, NAVY_HEX)
        elif shade:
            set_cell_bg(cell, LIGHT_HEX)
        p = cell.paragraphs[0]
        p.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT
            if i in right_cols
            else WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        )
        color = WHITE if total else BLACK
        add_run(p, str(val), bold=total, size_pt=8.5, color=color)
    return row


def new_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    return doc


def add_header_bar(doc, ref_line):
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_tbl.autofit = False
    lc = hdr_tbl.rows[0].cells[0]
    set_cell_bg(lc, NAVY_HEX)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(6)
    lp.paragraph_format.space_after = Pt(6)
    add_run(lp, "Lighthief", bold=True, size_pt=22, color=WHITE)
    add_run(lp, " Cyprus Ltd", bold=False, size_pt=11, color=WHITE)
    rc = hdr_tbl.rows[0].cells[1]
    set_cell_bg(rc, NAVY_HEX)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(6)
    rp.paragraph_format.space_after = Pt(6)
    for line, is_gold in [
        ("LETTER OF INTENT", True),
        (ref_line, False),
        ("Date: May 2026", False),
        ("STRICTLY CONFIDENTIAL", True),
    ]:
        add_run(rp, line + "\n", bold=is_gold, size_pt=8, color=GOLD if is_gold else WHITE)
    lock_table_widths(hdr_tbl, [10.0, 4.0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_parties_table(doc, client_block, client_label="Client"):
    h1(doc, "PARTIES")
    parties_tbl = doc.add_table(rows=2, cols=2)
    parties_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    parties_tbl.autofit = False
    for row_i, (label, text) in enumerate(
        [
            (
                "Contractor",
                "Lighthief Cyprus Ltd\nRegistration No. HE 477423\n"
                "15 Agaritsis, Nektaria Court, Office 201,\n"
                "Office 201, 3035 Limassol, Cyprus\n"
                "office@lighthief.com | +357 77 77 00 50\n"
                "(hereinafter \"Lighthief\" or the \"Contractor\")",
            ),
            (client_label, client_block),
        ]
    ):
        lc_p = parties_tbl.rows[row_i].cells[0]
        rc_p = parties_tbl.rows[row_i].cells[1]
        set_cell_bg(lc_p, NAVY_HEX)
        lc_p.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lpp = lc_p.paragraphs[0]
        lpp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(lpp, label, bold=True, size_pt=10, color=WHITE)
        rpp = rc_p.paragraphs[0]
        add_run(rpp, text, size_pt=9, color=BLACK)
    lock_table_widths(parties_tbl, [2.5, 11.5])


def add_signatures(doc, client_sig_name="Esperia Energy Group"):
    h1(doc, "SIGNATURES")
    body(
        doc,
        "IN WITNESS WHEREOF, the authorised representatives of the Parties have executed "
        "this Letter of Intent as of the date first written above.",
        space_after=10,
    )
    sig_tbl = doc.add_table(rows=5, cols=2)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    sig_data = [
        ("For and on behalf of\nLighthief Cyprus Ltd", f"For and on behalf of\n{client_sig_name}"),
        ("Name:   Alexander Papacosta", "Name:   ___________________________"),
        ("Title:    Managing Director", "Title:    ___________________________"),
        ("Signature:   ___________________________", "Signature:   ___________________________"),
        ("Date:   ___________________________", "Date:   ___________________________"),
    ]
    for row_i, (lt, rt) in enumerate(sig_data):
        lc_s = sig_tbl.rows[row_i].cells[0]
        rc_s = sig_tbl.rows[row_i].cells[1]
        if row_i == 0:
            set_cell_bg(lc_s, NAVY_HEX)
            set_cell_bg(rc_s, NAVY_HEX)
            for cell, txt in ((lc_s, lt), (rc_s, rt)):
                sp = cell.paragraphs[0]
                add_run(sp, txt, bold=True, size_pt=9, color=WHITE)
        else:
            for cell, txt in ((lc_s, lt), (rc_s, rt)):
                sp = cell.paragraphs[0]
                add_run(sp, txt, size_pt=10, color=BLACK)
    lock_table_widths(sig_tbl, [7.0, 7.0])


def add_footer(doc, ref_line):
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        fp,
        "Lighthief Cyprus Ltd  |  HE 477423  |  15 Agaritsis, Nektaria Court, Office 201, "
        "Office 201, 3035 Limassol, Cyprus\n"
        "office@lighthief.com  |  +357 77 77 00 50  |  solarfarms.cy  |  "
        + ref_line,
        size_pt=7.5,
        color=GREY,
    )


def pipeline_phase_totals():
    totals = {}
    for k in PIPELINE_PHASE_ORDER:
        parks = PIPELINE[k]
        totals[k] = {
            "label": PIPELINE_PHASE_LABELS[k],
            "mw": sum(p["mw"] for p in parks),
            "mwh": sum(p["mwh"] for p in parks),
            "containers": sum(p["containers"] for p in parks),
            "val": sum(p["client_price"] for p in parks),
        }
    return totals


def galascope_totals():
    return {
        "mw": sum(p["mw"] for p in GALASCOPE),
        "mwh": sum(p["mwh"] for p in GALASCOPE),
        "containers": sum(p["containers"] for p in GALASCOPE),
        "val": sum(p["client_price"] for p in GALASCOPE),
    }


def add_clause_1_3(doc, *, reimbursement_cap: int | None = None):
    """Binding reliance costs — short form for fast client sign-off."""
    cap_clause = ""
    if reimbursement_cap:
        cap_clause = f" Total reimbursement under this Clause is capped at {fmt_eur(reimbursement_cap)}."
    body(
        doc,
        "1.3  (Binding) If the Client withdraws after Lighthief has relied on this LOI, the Client "
        "reimburses third-party costs paid or non-cancellable third-party commitments for Schedule 1 "
        "parks (evidenced by invoice or signed contract), unless withdrawal is for a Valid Reason."
        + cap_clause,
    )
    body(
        doc,
        "Valid Reason: (a) grid Connection Terms / licence not obtained within twelve (12) months "
        "of the EPC Effective Date (or such other longstop as the EPC specifies); (b) change in law "
        "or DSO/EAC rules makes installation unlawful or impossible (adviser certificate); "
        "(c) mutual written termination; (d) Lighthief material breach not remedied within thirty "
        "(30) days; (e) project financing for a park is not secured despite reasonable efforts "
        "evidenced in writing; (f) Linyang insolvency or sanctions preventing supply. New third-party "
        f"spend above {fmt_eur(50_000)} per park requires the Client's prior written approval.",
    )


def add_clause_6_good_faith(doc, *, per_phase_exit: bool = False):
    """§6 — whole LOI (Galascope) or per-batch exit (Esperia pipeline)."""
    h1(doc, "6.   GOOD FAITH")
    if per_phase_exit:
        body(
            doc,
            "6.1  The Parties will negotiate each batch EPC in good faith. If terms for a batch are "
            "not agreed within ninety (90) days of that batch's EPC target in Section 3, either Party "
            "may end this LOI for that batch only, without liability except Clause 1.3.",
        )
    else:
        body(
            doc,
            "6.1  The Parties will negotiate and sign the EPC package in good faith. If not executed "
            "within ninety (90) days of this LOI (unless extended in writing), either Party may "
            "withdraw without liability except Clause 1.3.",
        )


def add_clause_7_exclusivity(doc, *, carveout_galascope=False):
    """Exclusivity until EPC signing — no post-PAC tail (signing-friendly)."""
    h1(doc, "7.   EXCLUSIVITY")
    body(
        doc,
        "7.1  (Binding) Until this LOI expires or the relevant Schedule 1 EPC is signed, the Client "
        "will not appoint another EPC contractor for those parks or procure parallel BESS equipment "
        "for them from an OEM other than Linyang (Cyprus portfolio).",
    )
    body(
        doc,
        "7.2  Exclusivity does not apply on Valid Reason (Clause 1.3) or if Lighthief has not "
        "proposed an EPC for signature within ninety (90) days of this LOI (unless extended in writing).",
    )
    body(
        doc,
        "7.3  Exclusivity ends when the Schedule 1 EPC is signed, on Valid Reason, or if either "
        "Party materially breaches this LOI and does not remedy within thirty (30) days of notice.",
    )
    if carveout_galascope:
        body(
            doc,
            "7.4  Applies to the nine (9) pipeline parks only — not Galascope Ltd (LCY-LOI-GAL-B1-2026-R15).",
        )


SCHEDULE_TABLE_COLS = [6.0, 1.8, 1.4, 1.4, 3.4]
SCHEDULE_TABLE_RC = {2, 3, 4}


def add_price_mechanism_section(doc, *, schedule_a_note=""):
    """Indicative/Confirmed two-way pricing + May 2026 Linyang price-lock window."""
    h1(doc, "4.   COMMERCIAL SUMMARY")
    body(
        doc,
        "4.1  Schedule 1 shows Indicative Prices (excl. VAT) from Linyang quotation LY202601271. "
        "Payment terms, civil works scope, and the binding Contract Price are fixed in each executed "
        "EPC using the Indicative and Confirmed Price mechanism (EPC §6.1).",
    )
    if schedule_a_note:
        body(doc, schedule_a_note.strip())
    body(
        doc,
        "4.2  Schedule 1 indicatives are valid for ninety (90) days from this LOI for planning. "
        "The Confirmed Contract Price for each park or batch is set when Lighthief issues the "
        "Confirmed Price Certificate under Clause 4.3.",
    )
    body(
        doc,
        "4.3  (Two-way price commitment) Each EPC will state an Indicative Price (as in Schedule 1 "
        "or updated in writing before signing). Within fourteen (14) days after the Client receives "
        "grid Connection Terms for that park or batch, Lighthief will issue a Confirmed Price "
        "Certificate fixing the Contract Price. The Confirmed Price equals the Indicative Price "
        "unless the battery raw-material index (Mysteel lithium carbonate) or EUR/CNY moves beyond "
        "the thresholds in EPC §6.1 from the January 2026 reference values ("
        f"{REF_LITHIUM_CNY}; {REF_EUR_CNY}). Where a threshold is exceeded, the Confirmed Price "
        "adjusts to reflect verified movement on the equipment cost base: downward adjustments "
        "are passed through to the Client on milestones not yet invoiced; upward adjustments are "
        "capped at five percent (5%) of the Indicative Price. If the calculated upward adjustment "
        "would exceed five percent (5%), either Party may decline to proceed with that EPC without "
        "penalty. Lighthief will provide a Price Basis Certificate at EPC signing (and at equipment "
        "delivery where applicable) showing the reference indices and calculation so the Client can "
        "verify how the price follows the market.",
    )
    body(
        doc,
        f"4.4  (Price lock — signature by {PRICE_LOCK_DEADLINE}) The Parties acknowledge that, as of "
        "the date of this LOI, the Mysteel lithium carbonate index has moved materially above the "
        "January 2026 reference in Clause 4.3, while EUR/CNY remains within the EPC FX threshold. "
        "Schedule 1 therefore reflects pricing locked with Linyang on the January 2026 quotation "
        "basis, not today's spot index. Linyang has confirmed that where this LOI and the relevant "
        f"park or batch EPC are executed on or before {PRICE_LOCK_DEADLINE}, the Indicative Price "
        "(as in Schedule 1 or updated in writing before signing) shall be the Confirmed Price for "
        "equipment cost purposes, without upward raw-material or FX adjustment for index movement "
        "occurring before that EPC Effective Date. Clause 4.3 two-way adjustments apply after "
        "signing only to verified movement from the indices stated in the Price Basis Certificate "
        "at EPC signing.",
    )
