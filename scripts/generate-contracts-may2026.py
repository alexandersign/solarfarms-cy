"""
Contract Pack Generator — Galascope & Spanercom  (May 2026)
Generates: EPC v5.0, LTSA v4.0, EMS Addendum, OEM DWU x2 per client + Updated Esperia LOI
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE  = '/Volumes/T7 Grey/solinvest'
TEMPL = f'{BASE}/legal/templates'
GAL_OUT = f'{BASE}/docs/clients/group-order/Group2_Esperia_Energy/contracts'
SPA_OUT = f'{BASE}/docs/clients/Individual_Spanercom/contracts'
DATE  = 'May 2026'

os.makedirs(GAL_OUT, exist_ok=True)
os.makedirs(SPA_OUT, exist_ok=True)

# ── helpers ─────────────────────────────────────────────────────────────────

def rep_para(para, old, new):
    full = para.text
    if old not in full:
        return False
    if para.runs:
        combined = ''
        for r in para.runs:
            combined += r.text
        if combined == full:
            # simple: rewrite first run
            para.runs[0].text = full.replace(old, new)
            for r in para.runs[1:]:
                r.text = ''
            return True
    # fallback: direct XML text
    for node in para._p.iter():
        if node.text and old in node.text:
            node.text = node.text.replace(old, new)
            return True
    return False

def rep_doc(doc, old, new):
    for para in doc.paragraphs:
        rep_para(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    rep_para(para, old, new)

def fmt(n):
    return f'{n:,.2f}'

# ── 1 & 6: EPC v5.0 ─────────────────────────────────────────────────────────

def gen_epc(client_name, client_reg, parks, total_price, pac_date, ref, ltsa_ref, out_path):
    doc = Document(f'{TEMPL}/client_sales_v5.0.docx')

    rep_doc(doc, '[? End Customer Legal Name]', client_name)
    rep_doc(doc, '[? End Customer]', client_name)
    rep_doc(doc, '[? end Customer]', client_name)
    # The template uses [?] bullet points for some fields
    rep_doc(doc, 'a company incorporated under the laws of [?], with registered office at [?] ("Client").',
            f'{client_reg} ("Client").')
    rep_doc(doc, '[? 2026 ("Effective Date")]', f'[?] {DATE} ("Effective Date")')
    rep_doc(doc, 'Date: 6 May 2026', f'Date: {DATE}')
    rep_doc(doc, 'LCY-LTSA-001 v4.0', ltsa_ref)
    # reference id
    rep_doc(doc, 'Document Reference: LCY-EPC-001', f'Document Reference: {ref}')

    # Populate Schedule A
    sched_lines = [
        f'SCHEDULE A - TECHNICAL SPECIFICATIONS AND COMMERCIAL DETAILS',
        f'Document Reference: {ref} | Version: 1.0 | Date: {DATE}',
        '',
        '=' * 58,
        'PART 1 - PARK DETAILS',
        '=' * 58,
    ]
    for pk in parks:
        sched_lines += [
            '',
            f"Park Name:                    {pk['name']}",
            f"Site Location:                {pk['district']} District, Cyprus",
            f"BESS Capacity:                {pk['mwh']} MWh",
            f"BESS Power Rating:            {pk['mw']} MW",
            f"BESS Containers:              {pk['bess']} x Linyang Power Atlantic 5.015 MWh (20HC)",
            f"MV Skid:                      {pk['mv_desc']}",
            f"PCS:                          {pk['pcs']}",
            f"Coastal Distance:             {pk['coastal']} km (C5 enclosure - full 5yr warranty)",
            '',
            f"  Component A (Equipment, CIF):   EUR {fmt(pk['comp_a'])}",
            f"  Component B (EPC Services):     EUR {fmt(pk['comp_b'])}",
            f"  CONTRACT PRICE (A+B, ex VAT):   EUR {fmt(pk['price'])}",
            '',
            f"  APG amount (70% x Comp A):      EUR {fmt(pk['apg'])}",
            f"  Performance Bond (5% x Comp A): EUR {fmt(pk['pb'])}",
            '-' * 58,
        ]
    sched_lines += [
        '',
        '=' * 58,
        'PART 2 - COMBINED CONTRACT SUMMARY',
        '=' * 58,
        '',
        f"TOTAL CONTRACT PRICE (all parks, ex VAT):   EUR {fmt(total_price)}",
        '',
        'PAYMENT MILESTONES:',
        f"  Advance (30%):      EUR {fmt(total_price*0.30)}  - Within 7 days of signing",
        f"  Pre-Shipment (55%): EUR {fmt(total_price*0.55)}  - On FAT completion + Linyang written confirmation",
        f"  PAC (10%):          EUR {fmt(total_price*0.10)}  - System commissioned and grid-connected",
        f"  Retention (5%):     EUR {fmt(total_price*0.05)}  - Released at end of 3-month DLP",
        f"  TOTAL:              EUR {fmt(total_price)}",
        '',
        '=' * 58,
        'PART 3 - KEY DATES',
        '=' * 58,
        '',
        'Equipment Production Start: Q2 2026',
        'Factory Acceptance Test:    Q3 2026',
        'CIF Limassol Arrival:       August / September 2026',
        f'Target PAC Date:            {pac_date}',
        'Warranty Start:             On PAC',
        'Base Warranty Expiry:       5 years from PAC (January 2032)',
        'Defects Liability Period:   3 months post-PAC (April 2027)',
        'Retention Release:          On expiry of DLP - April 2027 (approx.)',
        '',
        '=' * 58,
        'PART 4 - BESS TECHNICAL SPECIFICATIONS',
        '=' * 58,
        '',
        'Equipment: Linyang Power Atlantic 5.015 MWh LFP Battery Container',
        '  Cell: EVE LF314 (314 Ah LFP), 12P416S configuration',
        '  Thermal management: liquid cooling (45 kW capacity)',
        '  Fire suppression: FK5112 perfluoroketone + water backup',
        '  BMS: BMU / BCMU / BAMS, Modbus TCP / IEC 60870-5-104',
        '  IP Rating: IP54; C5-rated enclosure (Cyprus standard)',
        '  Weight: ~43 tonnes; Dimensions: 6058 x 2438 x 2896 mm',
        '',
        'PCS: Kehua BCS1250K-C-HUD',
        '  PCS efficiency: >=98% at rated power',
        '  Grid code: EN 50549-2 (TUV cert D 115067 0077)',
        '  IEC 62909-1/-2, IEC 62477-1, IEC 62116/61727, CE-LVD, CE-EMC',
        '',
        'Incoterms: CIF Limassol, Cyprus (Incoterms 2020)',
        '',
        '=' * 58,
        'PART 5 - ACCOMPANYING DOCUMENTS (executed simultaneously)',
        '=' * 58,
        '',
        f'1. LTSA v4.0 - Long-Term Service Agreement (Ref: {ltsa_ref})',
        '2. OEM Direct Warranty Undertaking - Schedule C (one per park)',
        '   Ref: LCY-OEM-DWU-001 - Linyang direct step-in rights',
        '3. EMS Subscription Addendum - Disperon EMS (EUR 400/MWh/yr)',
        '4. APG from Linyang Energy - delivered before first payment',
        '5. Performance Bond from Linyang (5% of Component A)',
    ]

    # Find and replace Schedule A placeholder paragraph
    sched_new_text = '\n'.join(sched_lines)
    for para in doc.paragraphs:
        if '[Linyang Power Atlantic 5MWh specifications per v4.0.' in para.text:
            if para.runs:
                para.runs[0].text = sched_new_text
                for r in para.runs[1:]:
                    r.text = ''
            break

    # Fix signature block
    rep_doc(doc, 'For and on behalf of [? End Customer] ("Client"):',
            f'For and on behalf of {client_name} ("Client"):')
    rep_doc(doc, '[? End Customer]', client_name)

    doc.save(out_path)
    print(f'  OK: {os.path.basename(out_path)}')


# ── 2 & 7: LTSA v4.0 ────────────────────────────────────────────────────────

def gen_ltsa(client_name, client_reg, parks, total_mwh, pac_date, ref, epc_ref, out_path):
    doc = Document(f'{TEMPL}/ClientLTSA_v4.0.docx')

    rep_doc(doc, '[End Customer Legal Name]', client_name)
    rep_doc(doc, 'a company incorporated under the laws of [?], with registered office at [?]', client_reg)
    rep_doc(doc, '[?] 2026 ("Effective Date")', f'[?] {DATE} ("Effective Date")')
    rep_doc(doc, 'Document Reference: LCY-LTSA-001', f'Document Reference: {ref}')
    rep_doc(doc, 'Version: 4.0\nDate: 17 March 2026', f'Version: 4.0\nDate: {DATE}')
    rep_doc(doc, 'Date: 17 March 2026', f'Date: {DATE}')
    rep_doc(doc, 'LCY-EPC-001', epc_ref)

    # Schedule 1 - site info
    rep_doc(doc, 'Site Name: [?]', 'Site Name(s): ' + ' | '.join(p['name'] for p in parks))
    rep_doc(doc, 'Site Address: [?]', 'Site(s): ' + ', '.join(f"{p['name']}, {p['district']} District" for p in parks))
    rep_doc(doc, 'GPS Coordinates: Lat [?]  Long [?]', 'GPS Coordinates: To be confirmed per site')
    rep_doc(doc, 'Distance from Sea: [?] km (affects warranty - see EPC Agreement)',
            'Distance from Sea: >500m all parks (C5 enclosure - full 5-year warranty applies)')

    # Schedule 1 - equipment
    rep_doc(doc, '| Battery Container | Linyang Power Atlantic 5MWh | [?] | [?] |',
            f"| Battery Container | Linyang Power Atlantic 5.015 MWh | {sum(p['bess'] for p in parks)} | TBC at delivery |")
    rep_doc(doc, '| Power Conversion System | [?] MW PCS | [?] | [?] |',
            f"| Power Conversion System | Kehua BCS1250K-C-HUD 1.25 MW | {sum(p['bess'] for p in parks)} | TBC at delivery |")
    rep_doc(doc, '| Transformer | [?] MVA Transformer | [?] | [?] |',
            f"| Transformer | MV Step-Up Transformer | {sum(p['mv'] for p in parks)} | TBC at delivery |")
    rep_doc(doc, '| EMS | [Third Party - specify] | [?] | [?] |',
            '| EMS | Disperon EMS (Lighthief EU BESS Ltd) | 1 per park | N/A |')
    rep_doc(doc, '[Third Party - specify]', 'Disperon EMS (Lighthief EU BESS Ltd)')

    # Key dates
    rep_doc(doc, '| Equipment Delivered | [?] |', '| Equipment Delivered | September 2026 (target) |')
    rep_doc(doc, '| Commissioning Date (PAC) | [?] |', f'| Commissioning Date (PAC) | {pac_date} (target) |')
    rep_doc(doc, '| Warranty Start | [?] |', f'| Warranty Start | On PAC - {pac_date} |')
    rep_doc(doc, '| Base Warranty Expiry (5 years) | [?] |', '| Base Warranty Expiry (5 years) | January 2032 (target) |')
    rep_doc(doc, '| LTSA Start | [?] |', f'| LTSA Start | On PAC - {pac_date} |')
    rep_doc(doc, '| LTSA Initial Term End | [?] |', '| LTSA Initial Term End | January 2042 (15 years from PAC) |')

    # Schedule 2 - service fees
    tier_c = total_mwh * 1740
    ems_sub = total_mwh * 400
    total_annual = tier_c + ems_sub
    vat = total_annual * 0.19

    rep_doc(doc, 'SYSTEM CAPACITY: [?] MWh', f'SYSTEM CAPACITY: {total_mwh} MWh')
    rep_doc(doc, 'SELECTED SERVICE TIER: ? Tier A  ? Tier B  ? Tier C',
            'SELECTED SERVICE TIER: [X] Tier C  (Full Service with 97% Availability Guarantee)')
    rep_doc(doc, '| 97% Availability Guarantee (Years 1-15) | ?2,201.73 | ?[?] |',
            f'| 97% Availability Guarantee (Years 1-15) | EUR 1,740 | EUR {fmt(tier_c)} |')
    rep_doc(doc, '| EMS/SCADA Software Subscription | ?[?] | ?[?] |',
            f'| EMS/SCADA Software Subscription (Disperon) | EUR 400 | EUR {fmt(ems_sub)} |')
    rep_doc(doc, 'SYSTEM CAPACITY: [?] MWh\n| Item | Value |',
            f'SYSTEM CAPACITY: {total_mwh} MWh\n| Item | Value |')
    rep_doc(doc, '| System Capacity | [?] MWh |', f'| System Capacity | {total_mwh} MWh |')
    rep_doc(doc, '| Installed EMS/SCADA Cost | ?[?] |', '| Installed EMS/SCADA Cost | Per EMS Subscription Addendum |')
    rep_doc(doc, '| Annual Subscription (20%) | ?[?] |', f'| Annual Subscription (EUR 400/MWh/yr) | EUR {fmt(ems_sub)} |')
    rep_doc(doc, '| EUR/MWh/Year | ?[?] |', '| EUR/MWh/Year | EUR 400.00 |')
    rep_doc(doc, '| Base Service Fee (Selected Tier) | ?[?] |', f'| Base Service Fee (Tier C - EUR 1,740/MWh) | EUR {fmt(tier_c)} |')
    rep_doc(doc, '| Warranty Extension (if selected) | ?[?] |', '| Warranty Extension (if selected) | EUR 0.00 (not selected) |')
    rep_doc(doc, '| EMS/SCADA Annual Subscription | ?[?] |', f'| EMS/SCADA Annual Subscription (Disperon) | EUR {fmt(ems_sub)} |')
    rep_doc(doc, '| TOTAL ANNUAL FEE | ?[?] |', f'| TOTAL ANNUAL FEE | EUR {fmt(total_annual)} |')
    rep_doc(doc, '| VAT (19%) | ?[?] |', f'| VAT (19%) | EUR {fmt(vat)} |')
    rep_doc(doc, '| TOTAL INCLUDING VAT | ?[?] |', f'| TOTAL INCLUDING VAT | EUR {fmt(total_annual + vat)} |')

    # SOH LD rate placeholder
    rep_doc(doc, 'LD = (Guaranteed SOH - Actual SOH) x System Capacity x EUR [?] per kWh',
            f'LD = (Guaranteed SOH - Actual SOH) x {total_mwh} MWh x Market Rate (EUR/kWh at time of claim)')

    doc.save(out_path)
    print(f'  OK: {os.path.basename(out_path)}')


# ── 3 & 8: EMS Addendum ─────────────────────────────────────────────────────

def gen_ems(client_name, parks, total_mwh, pac_date, ref, epc_ref, out_path):
    doc = Document()
    s = doc.styles['Normal']
    s.font.size = Pt(10)

    def h1(txt):
        p = doc.add_heading(txt, level=1)
        p.paragraph_format.space_before = Pt(8)
        return p

    def h2(txt):
        p = doc.add_heading(txt, level=2)
        p.paragraph_format.space_before = Pt(6)
        return p

    def para(txt):
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(4)
        return p

    def bullet(txt):
        p = doc.add_paragraph(txt, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        return p

    h1('EMS SOFTWARE SUBSCRIPTION AGREEMENT')
    doc.add_paragraph('DISPERON Energy Management System - Annual Software Subscription')
    doc.add_paragraph('')

    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    for lbl, val in [
        ('Document Ref:', ref),
        ('Date:', DATE),
        ('Version:', '1.0'),
        ('EPC Agreement:', epc_ref),
        ('Client:', client_name),
        ('Service Provider:', 'Lighthief EU BESS Ltd (trading as DISPERON)'),
    ]:
        r = t.add_row()
        r.cells[0].text = lbl
        r.cells[1].text = val
    doc.add_paragraph('')

    h2('PARTIES')
    pt = doc.add_table(rows=0, cols=2)
    pt.style = 'Table Grid'
    r = pt.add_row()
    r.cells[0].text = (
        'Service Provider:\nLighthief EU BESS Ltd\n(trading as DISPERON)\nReg. No. HE 474192\n'
        '28 Oktovriou & Aemiliou Chourmouziou\nLophitis Business Center I, Floor 2\n3035 Limassol, Cyprus'
    )
    r.cells[1].text = f'Client:\n{client_name}\n[?] Cyprus'
    doc.add_paragraph('')

    h2('BACKGROUND')
    para(
        f'This EMS Software Subscription Agreement ("Addendum") governs the annual software '
        f'subscription for the DISPERON Energy Management System ("EMS") deployed at '
        f'{client_name}\'s BESS installations under EPC Agreement {epc_ref}. '
        f'The EMS hardware installation cost is included in the EPC Contract Price (Component B). '
        f'This Addendum covers only the ongoing annual software licence, platform access, '
        f'updates, and cloud services, billed separately from the EPC and LTSA.'
    )

    h2('1. WHAT THE SUBSCRIPTION COVERS')
    for item in [
        '24/7 cloud-hosted EMS platform access and real-time BESS monitoring dashboard',
        'All software updates, security patches, and regulatory compliance updates (EU grid codes, NIS2, GDPR)',
        'DSO/SCADA integration maintenance - IEC 60870-5-104 interface with Cyprus EAC',
        'BMS data integration - Modbus TCP / RS485 northbound interface monitoring',
        '10 read-only SCADA dashboard accounts per client group',
        'Remote EMS configuration support (1 hour/month per park; additional at EUR 80/hr)',
        'External data feeds: market price signals, grid frequency data, weather forecasts',
        'Battery degradation analytics and SOH trend reporting',
        'Quarterly EMS performance report per park',
        'Regulatory update notifications (grid code changes within 14 days of publication)',
    ]:
        bullet(item)

    h2('2. NOT COVERED BY THIS ADDENDUM')
    for item in [
        'On-site physical installation, configuration, or commissioning (covered under EPC)',
        'Physical maintenance of BESS, PCS, or MV switchgear (covered under LTSA)',
        'Custom software development beyond standard EMS configuration',
        'Advanced trading algorithms, revenue optimisation, or market participation logic (available separately)',
        'SCADA hardware replacement (covered under EPC warranty)',
    ]:
        bullet(item)

    h2('3. SUBSCRIPTION FEE')
    ft = doc.add_table(rows=0, cols=3)
    ft.style = 'Table Grid'
    r0 = ft.add_row()
    for ci, h in enumerate(['Park / Site', 'Capacity (MWh)', 'Annual Subscription (EUR 400/MWh)']):
        run = r0.cells[ci].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for pk in parks:
        row = ft.add_row()
        row.cells[0].text = pk['name']
        row.cells[1].text = f"{pk['mwh']} MWh"
        row.cells[2].text = f"EUR {pk['mwh'] * 400:,.2f}"
    total_row = ft.add_row()
    total_row.cells[0].paragraphs[0].add_run('TOTAL').bold = True
    total_row.cells[1].paragraphs[0].add_run(f'{total_mwh} MWh').bold = True
    total_row.cells[2].paragraphs[0].add_run(f'EUR {total_mwh * 400:,.2f} per year').bold = True
    doc.add_paragraph('')
    para(
        f'The subscription fee of EUR {total_mwh * 400:,.2f} per year (ex VAT) is invoiced annually '
        f'in advance from PAC date ({pac_date}, target). Payment due within 30 days of invoice. '
        f'Fee is exclusive of VAT (19%). Fixed Years 1-3, thereafter annual CPI adjustment (Cyprus HICP, max 3%).'
    )

    h2('4. TERM AND RENEWAL')
    para(
        f'Initial term: 15 years from PAC, aligned with LTSA {epc_ref.replace("EPC","LTSA")}. '
        f'Auto-renews for successive 5-year periods unless either Party gives 12 months written notice.'
    )

    h2('5. LINYANG WARRANTY AND BMS ACCESS')
    para(
        'The Linyang OEM Direct Warranty Undertaking (Schedule C of the EPC) confirms that '
        'Linyang\'s BMS northbound communication interface is compatible with the DISPERON EMS. '
        'This access persists for the lifetime of the equipment regardless of the status of the '
        'Lighthief-Linyang Distribution Agreement. The OEM Undertaking survives distribution agreement '
        'termination and ensures uninterrupted DISPERON EMS data access.'
    )

    h2('6. GOVERNING LAW')
    para('Laws of the Republic of Cyprus. Disputes resolved in Cyprus courts.')

    doc.add_paragraph('')
    h2('SIGNATURES')
    st = doc.add_table(rows=0, cols=2)
    st.style = 'Table Grid'
    for lbl, val in [
        ('For Lighthief EU BESS Ltd (DISPERON)', f'For {client_name}'),
        ('Name: Alexander Papacosta', 'Name: ___________________________'),
        ('Title: Cyprus Managing Director', 'Title: ___________________________'),
        ('Signature: ___________________________', 'Signature: ___________________________'),
        ('Date: ___________________________', 'Date: ___________________________'),
    ]:
        r = st.add_row()
        r.cells[0].text = lbl
        r.cells[1].text = val

    doc.save(out_path)
    print(f'  OK: {os.path.basename(out_path)}')


# ── 4 & 9: OEM DWU ──────────────────────────────────────────────────────────

def gen_dwu(client_name, client_addr, park_name, district, mwh, bess, mv, ref, epc_ref, out_path):
    doc = Document(f'{TEMPL}/OEM-Direct-Warranty-Undertaking-Linyang.docx')

    rep_doc(doc, '[End-Customer Legal Name]', client_name)
    rep_doc(doc, '[End-Customer Address]', client_addr)
    rep_doc(doc, '[Site Name / Park Name], [District], Cyprus',
            f'{park_name}, {district} District, Cyprus')
    rep_doc(doc, '[Site Name / Park Name]', park_name)
    rep_doc(doc, '[District]', district)
    rep_doc(doc, 'Document Reference: LCY-OEM-DWU-001', f'Document Reference: {ref}')
    rep_doc(doc, 'Version: 1.0\nDate: [?] March 2026', f'Version: 1.0\nDate: {DATE}\nProject Ref: {ref}')
    rep_doc(doc, 'Date: [?] March 2026', f'Date: {DATE}')
    rep_doc(doc, 'for the supply and installation of BESS Products at [Site Name / Park Name], [District], Cyprus',
            f'for the supply and installation of BESS Products at {park_name}, {district} District, Cyprus '
            f'({bess} x Linyang Power Atlantic 5.015 MWh battery containers + {mv} x MV Skid(s))')
    rep_doc(doc, '(EPC Agreement, delivery records, or PAC certificate)',
            f'(EPC Agreement Ref: {epc_ref}, delivery records, or PAC certificate)')
    rep_doc(doc, 'ACKNOWLEDGED by [End-Customer]:', f'ACKNOWLEDGED by {client_name}:')

    doc.save(out_path)
    print(f'  OK: {os.path.basename(out_path)}')


# ── 5: Updated Esperia LOI ───────────────────────────────────────────────────

def gen_updated_loi(out_path):
    src = f'{GAL_OUT}/LOI-Esperia-Energy-remaining-pipeline-apr2026.docx'
    doc = Document(src)

    # 1. Change 24-month DLP to 3-month
    rep_doc(doc, 'released upon expiry of the 24-month Defect Liability Period ("DLP").',
            'released upon expiry of the 3-month Defects Liability Period ("DLP") following PAC. '
            '[UPDATED per EPC v5.0 - FAC removed, DLP is 3 months from PAC]')
    rep_doc(doc, 'released upon expiry of the 24-month Defect Liability Period',
            'released upon expiry of the 3-month Defects Liability Period')
    rep_doc(doc, 'upon expiry of the 24-month Defect Liability Period',
            'upon expiry of the 3-month Defects Liability Period')

    # 2. Update date and ref
    rep_doc(doc, 'Date: 23 April 2026', f'Date: {DATE}  [Revision 1]')
    rep_doc(doc, 'Ref: LCY-LOI-ESP-PIPELINE-2026', 'Ref: LCY-LOI-ESP-PIPELINE-2026-R1')

    # 3. Add EMS subscription to LTSA section
    rep_doc(doc,
            'Term: a minimum of ten (10) years from the PAC date of the first commissioned park.',
            'Term: a minimum of fifteen (15) years from the PAC date of the first commissioned park.\n'
            '6.3 EMS Software Subscription: In addition to the LTSA service fee, the Client confirms '
            'its intention to subscribe to the DISPERON EMS platform (Lighthief EU BESS Ltd) at '
            'EUR 400/MWh/year, invoiced annually from PAC. Governed by separate EMS Subscription '
            'Addendum. DISPERON SHA signed May 2026.')

    # 4. Add DLP definition
    rep_doc(doc,
            '"FAT"  means Factory Acceptance Test conducted at Linyang\'s facility prior to shipment.',
            '"FAT"  means Factory Acceptance Test conducted at Linyang\'s facility prior to shipment.\n'
            '"DLP"  means Defects Liability Period. Under EPC v5.0: 3 months from PAC. '
            '5% Retention released at end of DLP (not 24 months as in prior drafts).')

    # 5. Update Esperia Famagusta price in pipeline table
    rep_doc(doc, 'EUR 2,316,815', 'EUR 2,500,334')

    # 6. Note at top
    rep_doc(doc,
            'Letter of Intent\nEsperia Energy Group - BESS Portfolio EPC Pipeline Commitment',
            'Letter of Intent - REVISED (Revision 1, May 2026)\n'
            'Esperia Energy Group - BESS Portfolio EPC Pipeline Commitment\n'
            'KEY CHANGES IN REVISION 1:\n'
            '- Retention DLP changed from 24 months to 3 months (EPC v5.0)\n'
            '- LTSA term extended to 15 years\n'
            '- EMS Subscription confirmed at EUR 400/MWh/yr (Disperon SHA signed)\n'
            '- Esperia Famagusta pricing updated to EUR 2,500,334')

    doc.save(out_path)
    print(f'  OK: {os.path.basename(out_path)}')


# ══ RUN ALL ══════════════════════════════════════════════════════════════════

gal_parks = [
    {'name': 'Galascope 1', 'district': 'Famagusta', 'mw': 5.0, 'mwh': 20, 'bess': 4, 'mv': 1,
     'mv_desc': '1 x T4 MV Skid (4 x BCS1250K = 5 MW, SL-5000 transformer)',
     'pcs': '4 x Kehua BCS1250K-C-HUD (1.25 MW each) = 5.0 MW', 'coastal': '>5',
     'comp_a': 1848712.43, 'comp_b': 410187.57, 'price': 2258900,
     'apg': 1848712.43 * 0.70, 'pb': 1848712.43 * 0.05},
    {'name': 'Galascope 2', 'district': 'Famagusta', 'mw': 2.5, 'mwh': 10, 'bess': 2, 'mv': 1,
     'mv_desc': '1 x T2 MV Skid (2 x BCS1250K = 2.5 MW, SL-2500 transformer)',
     'pcs': '2 x Kehua BCS1250K-C-HUD (1.25 MW each) = 2.5 MW', 'coastal': '>5',
     'comp_a': 974457.00, 'comp_b': 231843.00, 'price': 1206300,
     'apg': 974457.00 * 0.70, 'pb': 974457.00 * 0.05},
]

spa_parks = [
    {'name': 'Anarita 1', 'district': 'Paphos', 'mw': 5.0, 'mwh': 20, 'bess': 4, 'mv': 1,
     'mv_desc': '1 x T4 MV Skid (4 x BCS1250K = 5 MW, SL-5000 transformer)',
     'pcs': '4 x Kehua BCS1250K-C-HUD (1.25 MW each) = 5.0 MW', 'coastal': '>5',
     'comp_a': 1848712.43, 'comp_b': 531287.57, 'price': 2380000,
     'apg': 1848712.43 * 0.70, 'pb': 1848712.43 * 0.05},
    {'name': 'Anarita 2', 'district': 'Paphos', 'mw': 5.0, 'mwh': 20, 'bess': 4, 'mv': 1,
     'mv_desc': '1 x T4 MV Skid (4 x BCS1250K = 5 MW, SL-5000 transformer)',
     'pcs': '4 x Kehua BCS1250K-C-HUD (1.25 MW each) = 5.0 MW', 'coastal': '>5',
     'comp_a': 1848712.43, 'comp_b': 531287.57, 'price': 2380000,
     'apg': 1848712.43 * 0.70, 'pb': 1848712.43 * 0.05},
]

print('=== GALASCOPE ===')
gen_epc('Galascope Ltd',
        'a company incorporated under the laws of Cyprus, with registered office at [?], Cyprus',
        gal_parks, 3465200, '31 January 2027',
        'LCY-EPC-GAL-B1-2026', 'LCY-LTSA-GAL-2026',
        f'{GAL_OUT}/EPC-Galascope-Esperia-batch1-may2026.docx')

gen_ltsa('Galascope Ltd',
         'a company incorporated under the laws of Cyprus, with registered office at [?], Cyprus',
         gal_parks, 30, '31 January 2027',
         'LCY-LTSA-GAL-2026', 'LCY-EPC-GAL-B1-2026',
         f'{GAL_OUT}/LTSA-Galascope-Esperia-may2026.docx')

gen_ems('Galascope Ltd', gal_parks, 30, '31 January 2027',
        'LCY-EMS-GAL-2026', 'LCY-EPC-GAL-B1-2026',
        f'{GAL_OUT}/EMS-Subscription-Galascope-may2026.docx')

for pk in gal_parks:
    gen_dwu('Galascope Ltd', '[?] Famagusta, Cyprus',
            pk['name'], pk['district'], pk['mwh'], pk['bess'], pk['mv'],
            f"LCY-OEM-DWU-{pk['name'].replace(' ', '-').upper()}-2026",
            'LCY-EPC-GAL-B1-2026',
            f"{GAL_OUT}/OEM-DWU-{pk['name'].replace(' ','-')}-may2026.docx")

print('  Updating LOI...')
gen_updated_loi(f'{GAL_OUT}/LOI-Esperia-Energy-pipeline-may2026.docx')

print('\n=== SPANERCOM ===')
gen_epc('Spanercom Ltd',
        'a company incorporated under the laws of Cyprus, with registered office at [?], Paphos, Cyprus',
        spa_parks, 4760000, '31 January 2027',
        'LCY-EPC-SPA-2026', 'LCY-LTSA-SPA-2026',
        f'{SPA_OUT}/EPC-Spanercom-Anarita-may2026.docx')

gen_ltsa('Spanercom Ltd',
         'a company incorporated under the laws of Cyprus, with registered office at [?], Paphos, Cyprus',
         spa_parks, 40, '31 January 2027',
         'LCY-LTSA-SPA-2026', 'LCY-EPC-SPA-2026',
         f'{SPA_OUT}/LTSA-Spanercom-Anarita-may2026.docx')

gen_ems('Spanercom Ltd', spa_parks, 40, '31 January 2027',
        'LCY-EMS-SPA-2026', 'LCY-EPC-SPA-2026',
        f'{SPA_OUT}/EMS-Subscription-Spanercom-may2026.docx')

for pk in spa_parks:
    gen_dwu('Spanercom Ltd', '[?] Paphos, Cyprus',
            pk['name'], pk['district'], pk['mwh'], pk['bess'], pk['mv'],
            f"LCY-OEM-DWU-{pk['name'].replace(' ', '-').upper()}-2026",
            'LCY-EPC-SPA-2026',
            f"{SPA_OUT}/OEM-DWU-{pk['name'].replace(' ','-')}-may2026.docx")

print('\n=== OUTPUT FILES ===')
for folder, label in [(GAL_OUT, 'GALASCOPE'), (SPA_OUT, 'SPANERCOM')]:
    print(f'\n{label}:')
    for f in sorted(os.listdir(folder)):
        if 'may2026' in f.lower() and f.endswith('.docx'):
            size = os.path.getsize(os.path.join(folder, f))
            print(f'  {f}  ({size/1024:.0f} KB)')
print('\nDone.')
