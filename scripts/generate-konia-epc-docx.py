#!/usr/bin/env python3
"""
Generate brand-styled Word (.docx) deliverables for TOTALCON LTD — 1 MW Konia PV EPC.

Outputs into docs/clients/TOTALCON/:
  - Lighthief-PV-EPC-Proposal-TOTALCON-Konia-FixedTilt-Jul2026.docx (Option A, EUR 550,000)
  - Lighthief-PV-EPC-Proposal-TOTALCON-Konia-Trackers-Jul2026.docx (Option B, EUR 615,000)
  - Lighthief-PV-EPC-Contract-SAMPLE-TOTALCON-Konia-Jul2026.docx

Brand tokens per .cursor/rules/lighthief-brand-identity.mdc (NAVY/GOLD, Calibri).
Contacts per .cursor/rules/lighthief-contacts.mdc (SSOT).
Pricing per docs/internal/konia-epc-cost-model.md (Rev C): PV EUR550k/MW + EUR90k/MW
tracker adder, BESS EUR70k flat, 4 MWh container.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand tokens ────────────────────────────────────────────────
NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x40, 0x40, 0x40)
NAVY_HEX = "1A365D"

REPO = "/Volumes/T7 Grey/solinvest"
LOGO = os.path.join(REPO, "public/images/logo/lighthief-logo.png")
OUTDIR = os.path.join(REPO, "docs/clients/TOTALCON")
os.makedirs(OUTDIR, exist_ok=True)


# ── Helpers ─────────────────────────────────────────────────────
def new_doc():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)
    st.font.color.rgb = BLACK
    st.paragraph_format.space_after = Pt(6)
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    return doc


def logo_and_title(doc, title, subtitle, meta_lines):
    if os.path.exists(LOGO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run()
        r.add_picture(LOGO, width=Inches(2.1))
    h = doc.add_paragraph()
    r = h.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = GOLD
    if subtitle:
        s = doc.add_paragraph()
        rs = s.add_run(subtitle)
        rs.font.size = Pt(13)
        rs.font.color.rgb = NAVY
        rs.bold = True
    for line in meta_lines:
        m = doc.add_paragraph()
        rm = m.add_run(line)
        rm.font.size = Pt(9.5)
        rm.font.color.rgb = GREY
    _hr(doc)


def _hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C9A432')
    pbdr.append(bottom)
    pPr.append(pbdr)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = GOLD
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    return p


def para(doc, text, bold=False, italic=False, size=11, color=BLACK):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, color=BLACK, size=10, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = 'Calibri'


def make_table(doc, headers, rows, col_align=None, total_row=False, widths=None):
    """headers: list[str]; rows: list[list[str]]; col_align optional per-col alignment."""
    ncols = len(headers)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, htext in enumerate(headers):
        c = table.rows[0].cells[j]
        set_cell_bg(c, NAVY_HEX)
        align = col_align[j] if col_align else None
        _set_cell_text(c, htext, bold=True, color=WHITE, size=10, align=align)
    # body
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        is_total = total_row and (i == len(rows) - 1)
        for j, val in enumerate(row):
            if is_total:
                set_cell_bg(cells[j], NAVY_HEX)
                _set_cell_text(cells[j], val, bold=True, color=WHITE, size=10,
                               align=(col_align[j] if col_align else None))
            else:
                _set_cell_text(cells[j], val, bold=False, color=BLACK, size=10,
                               align=(col_align[j] if col_align else None))
    if widths:
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w
    return table


def callout(doc, text, warn=True):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFBEB' if warn else 'F0F4F8')
    pPr.append(shd)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    return p


def footer_block(doc):
    _hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Lighthief Cyprus Ltd  —  HE 477423  —  solarfarms.cy")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("28 October Ave 249, Lophitis Business Center 1, Office 201, 3035 Limassol, Cyprus")
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = GREY


R = WD_ALIGN_PARAGRAPH.RIGHT
C = WD_ALIGN_PARAGRAPH.CENTER


# ════════════════════════════════════════════════════════════════
# PROPOSAL
# ════════════════════════════════════════════════════════════════
def build_one_pager(v):
    doc = new_doc()
    logo_and_title(
        doc,
        "Solar PV — EPC Fixed-Price Offer",
        f"{v['opt']} · {v['cap']} · Konia, Ierokipia, Paphos",
        [
            "Prepared for: TOTALCON LTD (HE 458896)   |   Attention: Mr. Qiu",
            f"July 2026   |   Ref: {v['ref']}",
        ],
    )

    p = doc.add_paragraph()
    r = p.add_run(v['price'] + "  ")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAVY
    r2 = p.add_run("fixed lump sum, PV turnkey (ex VAT) · " + v['perwp'])
    r2.font.size = Pt(11); r2.font.color.rgb = GREY
    para(doc, v['cap'] + " · " + v['struct_line'], size=10, color=GREY)

    h2(doc, "Included in this fixed price")
    for t in v['included']:
        bullet(doc, t)

    h2(doc, "Key terms")
    make_table(doc, ["Item", "Detail"], v['terms'])

    is_tracker = "Tracker" in v['opt']
    proc_label = "Module & tracker procurement & shipping" if is_tracker else "Module procurement & shipping (China)"
    install_dur = "7–9 weeks (after modules land)" if is_tracker else "6–8 weeks (after modules land)"
    h2(doc, "Indicative timeline (from advance payment)")
    make_table(doc, ["Phase", "Duration"],
               [["Engineering, permits & civils", "6–10 weeks (parallel with procurement)"],
                [proc_label, "~18–20 weeks — critical path"],
                ["Install, cabling & switchgear", install_dur],
                ["Commissioning & EAC energisation", "3–4 weeks"],
                ["Total (indicative)", "~7–9 months"]])

    h2(doc, "Optional add-ons (ex VAT)")
    make_table(doc, ["Add-on", "Price", "Detail"],
               [["BESS EPC — install only", "EUR 70,000", "4 MWh container (1 MW / 4 MWh, per CERA). Client supplies all battery equipment; Lighthief installs foundations, cabling, protection & commissioning support."],
                ["PV O&M — annual service", v['om'], v['om_note']]],
               col_align=[None, R, None])

    h2(doc, "Not included (client responsibility)")
    para(doc, "EAC LV connection & transformer (charged by EAC on the connection offer); EAC testing/admin fees; building permit fees; off-site grid reinforcement; BESS hardware; VAT (19%).")

    h2(doc, "Reference")
    para(doc, "CERA exemption E5229/2026 — 1 MWp PV + 1 MW / 4 MWh storage. Site: Temachio 150, Konia, Ierokipia, Paphos. Layout: OGT Energy Group, September 2025. Work commences on receipt of the 30% advance payment; permits and grid connection are the Client's responsibility and do not delay the works.")

    para(doc, "")
    para(doc, "Lighthief Cyprus Ltd — EPC", bold=True, color=NAVY)
    para(doc, "Alexander Papacosta — Cyprus Director   |   Tel: +357 99 164 158   |   office@lighthief.com")
    para(doc, "28 October Ave 249, Lophitis Business Center 1, Office 201, 3035 Limassol")

    footer_block(doc)
    para(doc, "Confidential — prepared for TOTALCON LTD. Fixed price ex VAT, valid 60 days.", italic=True, size=8.5, color=GREY)

    out = os.path.join(OUTDIR, v['file'])
    doc.save(out)
    return out


_COMMON_INCLUDED = [
    "ETEK-certified engineering, SLDs, permit support, as-builts",
    "Huawei string inverters + DC/AC BOS + monitoring",
    "Site civils — grading, cut/fill, internal roads, drainage",
    "Foundations & piling",
    "Perimeter fencing, gates & security system (CCTV + perimeter intrusion)",
    "Mechanical & electrical install, cabling, switchgear",
    "Commissioning, SCADA, EAC witness-test & energisation",
    "Logistics, cranage, HSE & project management",
]

OPTION_A = {
    'opt': "Option A — Fixed-Tilt",
    'cap': "825 kWp",
    'ref': "LCY-EPC-TOTALCON-KONIA-2026-07-A",
    'price': "EUR 550,000",
    'perwp': "EUR 0.67/Wp",
    'struct_line': "1,100 x RECOM Lion 750 W · fixed-tilt · LV, no transformer substation",
    'om': "EUR 6,600 / yr",
    'om_note': "Preventive maintenance, panel cleaning, vegetation, SCADA monitoring, 6 corrective call-outs/yr. Fixed-tilt structure \u2014 no drive/actuator maintenance required. Separate LTSA from PAC.",
    'file': "Lighthief-PV-EPC-Proposal-TOTALCON-Konia-FixedTilt-Jul2026.docx",
    'included': ["PV modules — 1,100 x RECOM Lion 750 W (Tier-1)",
                 "Fixed-tilt galvanised steel mounting structures"] + _COMMON_INCLUDED,
    'terms': [["Capacity", "825 kWp (fixed-tilt)"],
              ["Grid", "LV — no on-site transformer substation"],
              ["Programme", "~7–9 months from advance payment"],
              ["Payment", "30% signature · 40% delivery · 20% mechanical completion · 10% PAC"],
              ["Retention / holdback", "None — 12-month defects liability + warranty"],
              ["Warranty", "Modules 12/25-yr · inverters 5-yr · structures 10-yr · workmanship 2-yr from PAC"],
              ["Validity", "60 days"]],
}

OPTION_B = {
    'opt': "Option B — Single-Axis Trackers",
    'cap': "809 kWp",
    'ref': "LCY-EPC-TOTALCON-KONIA-2026-07-B",
    'price': "EUR 615,000",
    'perwp': "EUR 0.76/Wp",
    'struct_line': "1,156 x AIKO 700 W · single-axis trackers · +15-25% yield · LV, no TS",
    'om': "EUR 8,900 / yr",
    'om_note': "Preventive maintenance, panel cleaning, vegetation, SCADA monitoring, 8 corrective call-outs/yr. Includes tracker drive/actuator inspection, gearbox lubrication & row alignment calibration (2x/yr). Separate LTSA from PAC.",
    'file': "Lighthief-PV-EPC-Proposal-TOTALCON-Konia-Trackers-Jul2026.docx",
    'included': ["PV modules — 1,156 x AIKO 700 W (Tier-1)",
                 "Single-axis tracker system (structures + drives)"] + _COMMON_INCLUDED,
    'terms': [["Capacity", "809 kWp (single-axis trackers)"],
              ["Grid", "LV — no on-site transformer substation"],
              ["Yield", "+15-25% vs fixed tilt"],
              ["Programme", "~7–9 months from advance payment"],
              ["Payment", "30% signature · 40% delivery · 20% mechanical completion · 10% PAC"],
              ["Retention / holdback", "None — 12-month defects liability + warranty"],
              ["Warranty", "Modules 12/25-yr · inverters 5-yr · structures 10-yr · trackers 5-yr · workmanship 2-yr"],
              ["Validity", "60 days"]],
}


# ════════════════════════════════════════════════════════════════
# CONTRACT (SAMPLE)
# ════════════════════════════════════════════════════════════════
def build_contract():
    doc = new_doc()
    logo_and_title(
        doc,
        "Engineering, Procurement & Construction Agreement",
        "Solar Photovoltaic Park — Konia, Ierokipia, Paphos",
        [
            "Document Ref: LCY-EPC-PV-TOTALCON-KONIA-2026-07   |   Version: Sample Draft 1.0   |   July 2026",
        ],
    )

    callout(doc, "SAMPLE / DRAFT FOR DISCUSSION ONLY. This document is a non-binding sample prepared to illustrate the intended contract structure and commercial terms. It does not constitute an offer or a legally binding agreement. Fields marked [*] are to be completed on execution. Final terms are subject to legal review, receipt of EAC connection terms, and formal execution by both Parties. Figures are ex VAT.")

    h2(doc, "1. Parties")
    para(doc, 'This Engineering, Procurement and Construction Agreement (the "Agreement") is made as of [*] 2026 (the "Effective Date")')
    para(doc, "BETWEEN:", bold=True)
    para(doc, "Lighthief Cyprus Ltd, a company incorporated under the laws of the Republic of Cyprus, registered office at Agiou Andreou 241, AG TRIAS COURT, Flat/Office 31, 3036 Limassol, Cyprus, operational address at 28 October Avenue 249, Lophitis Business Center 1, Office 201, 3035 Limassol, Cyprus, Company No. HE 477423 (the \u201cContractor\u201d);")
    para(doc, "and", bold=True)
    para(doc, 'TOTALCON LTD, a company incorporated under the laws of the Republic of Cyprus, Company No. HE 458896, with registered office at [*] (the "Client").')
    para(doc, 'The Contractor and the Client are referred to individually as a "Party" and collectively as the "Parties".')

    h2(doc, "1A. Commencement")
    para(doc, "1A.1 The Works commence upon the Contractor's receipt of the advance payment under Section 7.1(a). The Client triggers commencement by releasing the advance payment when it wishes the project to proceed; no works, procurement, or mobilisation are undertaken before the advance payment is received.")
    para(doc, "1A.2 This Agreement is not conditional upon any planning permission, building permit, or EAC/DSO grid connection approval. Obtaining and maintaining all permits, consents, and the grid connection is the Client's responsibility under Section 5, and the absence or delay of any such permit or connection does not affect the commencement of the Works, the programme, or the Contractor's entitlement to payment.")

    h2(doc, "2. Background")
    para(doc, "2.1 The Client holds CERA exemption E5229/2026 authorising a 1 MWp solar photovoltaic generating station and a 1 MW / 4 MWh energy storage installation at the Site (Temachio 150, Konia, Ierokipia, Paphos).")
    para(doc, "2.2 The Contractor has the capability to design, procure, install, and commission solar PV parks on a turnkey EPC basis.")
    para(doc, "2.3 The Parties wish to set out the terms under which the Contractor shall deliver the Works to the Client.")

    h2(doc, "3. Definitions and Interpretation")
    for term, dfn in [
        ("\u201cWorks\u201d", "means the design, procurement, construction, installation, testing, and commissioning of the solar PV park and, if selected, the BESS installation works, as described in Section 4 and Schedule A."),
        ("\u201cSite\u201d", "means the land at Temachio 150, Konia, Ierokipia, Paphos, as identified in Schedule A."),
        ("\u201cBusiness Day\u201d", "means a day other than a Saturday, Sunday, or public holiday in the Republic of Cyprus."),
        ("\u201cComponent A — Equipment Supply\u201d", "means the portion of the Contract Price attributable to supply of PV equipment, as set out in Section 6."),
        ("\u201cComponent B — EPC Services\u201d", "means the portion of the Contract Price attributable to EPC works and services, as set out in Section 6."),
        ("\u201cPAC\u201d", "means Provisional Acceptance Certificate, issued under Section 9."),
        ("\u201cDefects Liability Period\u201d", "means the period of twelve (12) months from the date of PAC."),
        ("\u201cPoint of Connection (POC)\u201d", "means the low-voltage connection point defined in the EAC Connection Terms."),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(term + " ")
        r.bold = True
        r.font.size = Pt(10.5)
        r2 = p.add_run(dfn)
        r2.font.size = Pt(10.5)

    h2(doc, "4. Scope of Works")
    para(doc, "The Works are delivered in three packages. The Client has selected: [* Option A Fixed-Tilt / Option B Trackers] and [* including / excluding] Package 3 (BESS install-only).")
    h3(doc, "4.1 Package 1 — PV Equipment Supply (Component A)")
    for t in ["PV modules — Tier-1, CE/IEC certified (RECOM Lion 750 W for Option A; AIKO 700 W for Option B);",
              "Mounting structures — fixed-tilt galvanised steel (Option A) or single-axis tracker system (Option B);",
              "Inverters — Huawei string inverters with EAC-compliant grid codes;",
              "DC BOS — combiner boxes, DC cabling, connectors, SPD, earthing materials;",
              "AC BOS — AC cabling, LV switchgear, protection devices, metering hardware;",
              "Monitoring — data logger, sensors, communications hardware."]:
        bullet(doc, t)
    h3(doc, "4.2 Package 2 — PV EPC Works (Component B)")
    for t in ["ETEK-certified engineering design, SLDs, cable schedules, permit support, as-builts;",
              "Site civils — grading, cut/fill and elevation, internal roads, drainage;",
              "Foundations and piling per geotechnical design;",
              "Perimeter security fencing, access gates, and signage;",
              "Security system — CCTV, perimeter intrusion detection, and remote monitoring integration;",
              "Mechanical installation — structure assembly, module mounting, inverter placement;",
              "Electrical installation — DC stringing, AC/LV cabling and termination, trenching, switchgear;",
              "Commissioning, SCADA configuration, and EAC witness-test support;",
              "Logistics, cranage, HSE, and project management."]:
        bullet(doc, t)
    h3(doc, "4.3 Package 3 — BESS EPC Works Only (optional)")
    for t in ["Client supplies all BESS equipment (4 MWh container/PCS, electrical equipment, OEM documentation);",
              "Contractor provides foundations, crane offload, LV cabling and terminations, earthing, and protection integration;",
              "Commissioning support and EAC coordination at the PV point of interconnection;",
              "Excludes BESS hardware, import duty, and OEM factory commissioning crew unless separately agreed."]:
        bullet(doc, t)
    h3(doc, "4.4 Single-Point Responsibility")
    para(doc, "(a) The Contractor is responsible to the Client for the entire Works, including procured equipment and installation workmanship. The Client's sole recourse for any defect, delay, or performance issue is against the Contractor.")
    para(doc, "(b) Subcontracting of any part of the Works does not relieve the Contractor of its obligations.")
    para(doc, "(c) Any works not expressly listed are excluded unless agreed in writing by change order.")

    h2(doc, "5. Client Obligations")
    para(doc, "The Client shall: (a) provide complete and accurate site information and the topographic survey; (b) obtain planning permissions, building permits, and any landowner consents; (c) obtain and maintain the EAC Connection Terms and pay all EAC connection and transformer charges directly; (d) ensure clear and safe site access for delivery and installation; (e) designate a single point of contact and respond to queries within five (5) Business Days; (f) make all payments when due; and (g) permit Contractor access for warranty and maintenance.")

    h2(doc, "6. Contract Price")
    para(doc, "6.1 The total contract price (the \u201cContract Price\u201d) is the sum of Component A and Component B for the selected option, plus Package 3 if selected, exclusive of VAT. VAT (19%) shall be added to each invoice in accordance with Cyprus law. All milestone percentages in Section 7 are calculated on the Contract Price exclusive of VAT.")
    para(doc, "6.2 PV EPC is offered as a fixed lump-sum price per option: EUR 550,000 for the fixed-tilt option (Option A) and EUR 615,000 for the single-axis tracker option (Option B). BESS EPC (Package 3, install-only) is a flat EUR 70,000 for the client's 4 MWh battery container (1 MW / 4 MWh, per CERA).")
    h3(doc, "6.3 Price by Option (ex VAT)")
    make_table(doc, ["Configuration", "PV EPC (Pkg 1+2)", "BESS (Pkg 3)", "Total"],
               [["Option A — Fixed-Tilt 825 kWp", "550,000", "70,000", "620,000"],
                ["Option B — Single-Axis Trackers 809 kWp", "615,000", "70,000", "685,000"],
                ["Selected Contract Price [*]", "[*]", "[*]", "[*]"]],
               col_align=[None, R, R, R], total_row=True)
    para(doc, "6.4 The Contract Price is based on current technical, regulatory, customs, and tax conditions and on the module price and EUR/USD exchange rate prevailing at the Effective Date. Material changes in scope, regulation, law, module price, or exchange rate occurring after the Effective Date entitle the Parties to agree a variation by written change order.")

    h2(doc, "7. Payment Terms")
    h3(doc, "7.1 Payment Milestones")
    para(doc, "Payments shall be made by bank transfer, exclusive of VAT, as follows:")
    make_table(doc, ["Milestone", "%", "Trigger"],
               [["(a) Advance", "30%", "On signature; the Works commence upon the Contractor's receipt of this payment"],
                ["(b) Equipment delivered to Site", "40%", "Package 1 goods on Site"],
                ["(c) Mechanical completion", "20%", "Structures, modules, DC stringing complete"],
                ["(d) Provisional Acceptance", "10%", "Energisation + EAC witness test passed (PAC) — final payment, no holdback"]],
               col_align=[None, R, None])
    para(doc, "7.2 No retention. No retention or holdback is applied. The full Contract Price is paid across the milestones in Section 7.1, with the final payment due at PAC. The Client's security for defect rectification is the Contractor's obligations during the Defects Liability Period (Section 9.2) and the workmanship warranty (Section 10).")
    para(doc, "7.3 Package 3 (BESS works) is invoiced separately on matching milestones once the client equipment delivery date is confirmed.")
    para(doc, "7.4 Late payment. Interest accrues at the EU Directive 2011/7/EU rate (ECB + 8% p.a., simple interest). The Contractor may suspend the Works on seven (7) days' notice and, if any amount is overdue by more than thirty (30) days, terminate under Section 14.")
    h3(doc, "7.5 Title and Risk")
    para(doc, "(a) Risk passes to the Client upon delivery to Site and unloading.")
    para(doc, "(b) Title to the equipment passes to the Client upon issuance of PAC and receipt of the PAC payment under Section 7.1(d).")
    para(doc, "(c) Upon title transfer, the Client may grant security over the equipment to project-finance lenders without the Contractor's consent. If a lender requires earlier title transfer, the Parties shall negotiate an alternative arrangement in good faith.")

    h2(doc, "8. Delivery and Programme")
    para(doc, "8.1 Equipment supply and installation commence upon the Contractor's receipt of the advance payment (Section 7.1(a)) and availability of the Site, and proceed on the basis of continuous, unobstructed site access.")
    para(doc, "8.2 The indicative programme from commencement (receipt of the advance payment) to Provisional Acceptance is approximately seven to nine (7–9) months (Schedule C), the critical path being module procurement and shipping. Programme dates are good-faith estimates only; time is not of the essence. The programme extends for any delay caused by the Client, by permitting or grid-authority timelines, by Force Majeure, or by any matter outside the Contractor's reasonable control.")

    h2(doc, "9. Acceptance")
    para(doc, "9.1 Provisional Acceptance (PAC) is issued when: (a) the PV park is energised and exporting per the EAC Connection Terms; (b) all inverters and, if applicable, trackers are operational; (c) SCADA/monitoring communications are verified; (d) all safety and protection systems are functional and tested; and (e) all mandatory commissioning tests per Schedule A are passed.")
    para(doc, "9.2 Defects Liability Period. Following PAC, a twelve (12) month Defects Liability Period commences, during which the Contractor shall remedy all notified defects promptly at no cost. Minor defects not materially affecting operation are recorded on a punch list and remedied within the period. There is no Final Acceptance Certificate and no retention is held; the Contractor remedies all notified defects during the period at its own cost.")

    h2(doc, "10. Warranty")
    para(doc, "10.1 Unified warranty. The Contractor warrants the entire Works — procured equipment, installation workmanship, and EPC works — against defects in design, materials, and workmanship, and is the Client's sole warranty counterparty.")
    para(doc, "10.2 Warranty periods:")
    for t in ["PV modules — 12-year product / 25-year linear performance guarantee (manufacturer);",
              "Inverters — 5-year manufacturer warranty;",
              "Mounting structures — 10-year structural warranty (manufacturer);",
              "Tracker system (Option B) — 5-year mechanical warranty (manufacturer);",
              "Security system — 2-year equipment warranty (manufacturer);",
              "Contractor workmanship — 2 years from PAC (Packages 2 and 3)."]:
        bullet(doc, t)
    para(doc, "10.3 Exclusions. Warranty excludes consumables, Client misuse or negligence, unauthorised modifications, Force Majeure damage, and grid-side faults beyond the POC.")
    para(doc, "10.4 Conditions. Warranty is subject to operation per manufacturer guidelines, qualified maintenance, intact serial numbers, access for inspection, and payment in full.")

    h2(doc, "11. Liability")
    para(doc, "11.1 The Contractor's aggregate liability under this Agreement is capped at one hundred percent (100%) of the Contract Price, save that the cap does not apply to death or personal injury caused by negligence, fraud, or any liability that cannot be limited under Cyprus law.")
    para(doc, "11.2 Neither Party is liable for indirect or consequential loss, loss of profit, or loss of revenue, except where arising from the Contractor's wilful misconduct.")

    h2(doc, "12. Insurance")
    para(doc, "The Contractor shall maintain, during the Works, Contractor's All-Risk (CAR) / erection all-risks insurance and public liability insurance of not less than EUR [*], and employer's liability insurance as required by Cyprus law. Certificates shall be provided to the Client on request.")

    h2(doc, "13. Force Majeure")
    para(doc, "Neither Party is liable for failure or delay in performance caused by events beyond its reasonable control (including acts of God, war, epidemic, government action, grid unavailability, or extreme weather). The affected Party shall notify the other within seven (7) Business Days. If Force Majeure continues for more than one hundred and eighty (180) days, either Party may terminate, with settlement on the same basis as Section 14.2.")

    h2(doc, "14. Termination")
    para(doc, "14.1 Either Party may terminate on written notice if the other commits a material breach not remedied within thirty (30) days of notice, or becomes insolvent.")
    para(doc, "14.2 On termination, the Contractor retains the value of works completed and materials delivered to Site (agreed between the Parties within fourteen (14) days, failing which determined by an independent quantity surveyor), and refunds any excess payments within thirty (30) days.")

    h2(doc, "15. Governing Law and Disputes")
    para(doc, "This Agreement is governed by the laws of the Republic of Cyprus. The Parties shall attempt to resolve disputes amicably; failing which, disputes shall be submitted to the exclusive jurisdiction of the competent courts of Limassol, Cyprus, or, if both Parties agree in writing, to arbitration under the [*] Rules.")

    h2(doc, "16. General")
    para(doc, "16.1 This Agreement, including its Schedules, constitutes the entire agreement between the Parties and supersedes all prior discussions and proposals.")
    para(doc, "16.2 No variation is effective unless in writing and signed by both Parties.")
    para(doc, "16.3 Notices shall be in writing to the addresses in Section 1 (Contractor contact: office@lighthief.com; +357 77 77 00 50).")
    para(doc, "16.4 If any provision is held invalid, the remainder continues in full force.")

    h2(doc, "Schedules")
    bullet(doc, "Schedule A — Technical specification, site details, selected option, module/inverter/structure schedule, commissioning tests, and indicative programme.")
    bullet(doc, "Schedule B — Payment milestone schedule.")
    bullet(doc, "Schedule C — Programme (Gantt) from Effective Date to PAC.")
    bullet(doc, "Schedule D — Optional PV O&M terms (annual service).")

    h2(doc, "Execution")
    para(doc, "Agreed by the Parties as of the Effective Date.")
    sig = doc.add_table(rows=1, cols=2)
    left = sig.rows[0].cells[0]
    right = sig.rows[0].cells[1]
    _set_cell_text(left, "For the Contractor\nLighthief Cyprus Ltd (HE 477423)\n\nName: Alexander Papacosta\nTitle: Cyprus Director\nSignature: ____________________\nDate: ____________________", size=10)
    _set_cell_text(right, "For the Client\nTOTALCON LTD (HE 458896)\n\nName: [*]\nTitle: [*]\nSignature: ____________________\nDate: ____________________", size=10)

    footer_block(doc)
    para(doc, "SAMPLE / DRAFT — non-binding, for discussion only. Prepared for TOTALCON LTD. Ex VAT.", italic=True, size=8.5, color=GREY)

    out = os.path.join(OUTDIR, "Lighthief-PV-EPC-Contract-SAMPLE-TOTALCON-Konia-Jul2026.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    a = build_one_pager(OPTION_A)
    b = build_one_pager(OPTION_B)
    c = build_contract()
    print("Generated:")
    for x in (a, b, c):
        print(" -", x)
