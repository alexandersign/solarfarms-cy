"""
Generate PV + BESS Flexible Connection Teaser — DOCX + XLSX
Two connection models:
  Model A: 4-hour Hybrid (PV day sales + BESS curtailment arbitrage)
  Model B: 8-hour Flexible (100% nighttime BESS dispatch, zero daytime export)

Sources (SSOT):
  lib/portfolio-data.ts       — RTE 86.32%, LTSA €1,740/MWh, CIT 15%
  lib/deals/agios-theodoros-rtb.ts  — PV EPC €720/kWp, BESS base €127k/MWh
  lib/market/cyprus-tsoc-dam-sample.ts — DAM evening €182.99, daytime €140.88
  scripts/generate-softades-arkatzia-scenarios.py — AE Solar Meteor specs
"""

import os, math
from pathlib import Path

# ── paths ──────────────────────────────────────────────────────────────────────
ROOT   = Path(__file__).parent.parent
OUTDIR = ROOT / "docs" / "clients"
OUTDIR.mkdir(parents=True, exist_ok=True)
LOGO_PATH = ROOT / "public" / "images" / "lighthief-logo.png"

DOCX_OUT = OUTDIR / "pv-bess-flexible-connection-teaser-apr2026.docx"
XLSX_OUT = OUTDIR / "pv-bess-flexible-connection-teaser-apr2026.xlsx"

# ── brand colours ──────────────────────────────────────────────────────────────
NAVY_HEX   = "1A365D"
NAVY2_HEX  = "244F87"
GOLD_HEX   = "C9A432"
WHITE_HEX  = "FFFFFF"
GREY_HEX   = "404040"
GREEN_HEX  = "1E7D4A"
LGREY_HEX  = "EEF3FB"  # light navy tint
MGREY_HEX  = "D0DDEF"  # border
INPUT_HEX  = "FFF9E6"  # yellow input
CALC_HEX   = "EEF6EE"  # green calc

# ── SSOT constants ─────────────────────────────────────────────────────────────
RTE                = 0.8632   # AC-AC round-trip efficiency (OEM confirmed)
LTSA_PER_MWH_YR    = 1_740    # Lighthief LTSA Tier C, 15yr (portfolio-data.ts)
PV_OM_PER_KWP_YR   = 15       # PV O&M €/kWp/yr
DAM_EVENING        = 182.99   # TSOC DAM evening peak €/MWh
DAM_DAYTIME        = 140.88   # TSOC DAM daytime avg 06:00-17:00 €/MWh
AGGREGATOR_FEE     = 0.10     # 10% aggregator / market access fee
PV_YIELD_KWH_KWP   = 2_150    # AE Solar Meteor bifacial, 25° south, albedo 0.70
PV_EPC_PER_KWP     = 720      # €/kWp turnkey (Agios Theodoros confirmed)
RTB_PER_MW         = 380_000  # €/MW RTB licence acquisition
EAC_PER_MW         = 80_000   # €/MW EAC grid connection terms (≤8 MW tier)
MV_CABLE_PER_KM    = 20_000   # €/km MV cable
MV_CABLE_KM        = 1        # illustrative 1 km
PERMITTING         = 35_000   # permitting + docs per park
LTV                = 0.65     # 65% senior debt
EQUITY_PCT         = 0.35     # 35% equity
LOAN_RATE          = 0.055    # 5.5% p.a.
LOAN_YEARS         = 12       # 12-year annuity
DA_YEARS           = 20       # depreciation years
CIT                = 0.15     # Cyprus CIT 2026
CURTAILMENT        = 0.65     # 65% curtailment (Agios Theodoros / Galascope model)
BESS_CAPTURE       = 0.95     # BESS captures 95% of curtailed energy

# annuity factor: r(1+r)^n / ((1+r)^n - 1)
_r, _n = LOAN_RATE, LOAN_YEARS
ANNUITY_FACTOR = _r * (1 + _r)**_n / ((1 + _r)**_n - 1)  # ≈ 0.1160

# ── scenario definitions ───────────────────────────────────────────────────────
MODEL_A_SCENARIOS = [
    {"label": "1 MW + 4 MWh",  "pv_mwp": 1.0,  "bess_mwh": 4,  "bess_eur_mwh": 150_000},
    {"label": "2.5 MW + 10 MWh","pv_mwp": 2.5, "bess_mwh": 10, "bess_eur_mwh": 125_000},
    {"label": "5 MW + 20 MWh", "pv_mwp": 5.0,  "bess_mwh": 20, "bess_eur_mwh": 120_000},
]
MODEL_B_SCENARIOS = [
    {"label": "1 MW + 8 MWh",   "pv_mwp": 1.0, "bess_mwh": 8,  "bess_eur_mwh": 151_000},
    {"label": "2.5 MW + 20 MWh","pv_mwp": 2.5, "bess_mwh": 20, "bess_eur_mwh": 141_000},
    {"label": "5 MW + 40 MWh",  "pv_mwp": 5.0, "bess_mwh": 40, "bess_eur_mwh": 128_000},
]


def calc(s, model):
    """Return dict of all financial metrics for one scenario."""
    pv_mwp    = s["pv_mwp"]
    bess_mwh  = s["bess_mwh"]
    pv_kwp    = pv_mwp * 1_000

    # CAPEX
    pv_epc    = pv_kwp * PV_EPC_PER_KWP
    bess_epc  = bess_mwh * s["bess_eur_mwh"]
    rtb       = pv_mwp * RTB_PER_MW
    eac       = pv_mwp * EAC_PER_MW
    cable     = MV_CABLE_KM * MV_CABLE_PER_KM
    capex     = pv_epc + bess_epc + rtb + eac + cable + PERMITTING

    # Production
    pv_prod_mwh = pv_kwp * PV_YIELD_KWH_KWP / 1_000  # MWh/yr

    if model == "A":
        direct_mwh   = pv_prod_mwh * (1 - CURTAILMENT)
        curtailed    = pv_prod_mwh * CURTAILMENT
        captured     = curtailed * BESS_CAPTURE
        discharged   = captured * RTE
        direct_rev   = direct_mwh * DAM_DAYTIME
        bess_rev     = discharged * DAM_EVENING
        gross_rev    = direct_rev + bess_rev
    else:  # B — 100% through BESS
        direct_mwh   = 0
        discharged   = pv_prod_mwh * RTE
        direct_rev   = 0
        bess_rev     = discharged * DAM_EVENING
        gross_rev    = bess_rev

    net_rev  = gross_rev * (1 - AGGREGATOR_FEE)

    # OPEX
    pv_om    = pv_kwp * PV_OM_PER_KWP_YR
    bess_om  = bess_mwh * LTSA_PER_MWH_YR
    opex     = pv_om + bess_om

    ebitda   = net_rev - opex

    # Finance
    equity   = capex * EQUITY_PCT
    debt     = capex * LTV
    svc      = debt * ANNUITY_FACTOR
    dscr     = ebitda / svc if svc else 0
    da       = capex / DA_YEARS
    pretax_c = ebitda - svc
    ebt      = pretax_c - da
    tax      = max(0, ebt * CIT)
    net_cash = pretax_c - tax

    upayback = capex / ebitda if ebitda else 99
    epayback = equity / net_cash if net_cash > 0 else 99

    return {
        "pv_kwp": pv_kwp, "bess_mwh": bess_mwh, "bess_eur_mwh": s["bess_eur_mwh"],
        "pv_epc": pv_epc, "bess_epc": bess_epc, "rtb": rtb, "eac": eac,
        "cable": cable, "permitting": PERMITTING, "capex": capex,
        "pv_prod_mwh": pv_prod_mwh, "direct_mwh": direct_mwh,
        "discharged": discharged,
        "direct_rev": direct_rev, "bess_rev": bess_rev,
        "gross_rev": gross_rev, "net_rev": net_rev,
        "pv_om": pv_om, "bess_om": bess_om, "opex": opex,
        "ebitda": ebitda, "equity": equity, "debt": debt,
        "svc": svc, "dscr": dscr, "da": da,
        "pretax_cash": pretax_c, "tax": tax, "net_cash": net_cash,
        "upayback": upayback, "epayback": epayback,
    }


def eur(n, dp=0):
    """Format as €X,XXX,XXX"""
    if dp == 0:
        return f"€{int(round(n)):,}"
    return f"€{n:,.{dp}f}"


def eurf(n):
    """€X.XXM or €XXXk"""
    if abs(n) >= 1_000_000:
        return f"€{n/1_000_000:.2f}M"
    if abs(n) >= 1_000:
        return f"€{n/1_000:.0f}k"
    return f"€{n:.0f}"


# ══════════════════════════════════════════════════════════════════════════════
# 1. GENERATE LOGO PNG
# ══════════════════════════════════════════════════════════════════════════════

def make_logo():
    from PIL import Image, ImageDraw, ImageFont
    W, H = 600, 120
    img  = Image.new("RGB", (W, H), "#1A365D")
    draw = ImageDraw.Draw(img)

    # Lightning bolt (simplified polygon, white)
    bolt = [
        (52, 20), (32, 62), (52, 62), (28, 100), (72, 50),
        (50, 50), (72, 20)
    ]
    draw.polygon(bolt, fill="#C9A432")

    # Try to load a bold font; fall back to default
    try:
        font_lg = ImageFont.truetype("arialbd.ttf", 52)
        font_sm = ImageFont.truetype("arial.ttf", 16)
    except Exception:
        try:
            font_lg = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 52)
            font_sm = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 16)
        except Exception:
            font_lg = ImageFont.load_default()
            font_sm = font_lg

    draw.text((92, 18), "LIGHTHIEF", font=font_lg, fill="#C9A432")
    draw.text((94, 82), "CYPRUS LTD  ·  Exclusive Linyang Energy Distributor", font=font_sm, fill="#AABBCC")

    img.save(LOGO_PATH)
    print(f"  Logo -> {LOGO_PATH}")


# ══════════════════════════════════════════════════════════════════════════════
# 2. GENERATE WORD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def make_docx(results_a, results_b):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY  = RGBColor(0x1A, 0x36, 0x5D)
    GOLD  = RGBColor(0xC9, 0xA4, 0x32)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    BLACK = RGBColor(0x00, 0x00, 0x00)
    GREY  = RGBColor(0x40, 0x40, 0x40)
    GREEN = RGBColor(0x1E, 0x7D, 0x4A)

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)

    # ── helpers ────────────────────────────────────────────────────────────────

    def set_cell_bg(cell, hex_c):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_c)
        tcPr.append(shd)

    def set_borders(cell, color="D0DDEF", sz="4"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for side in ["top", "bottom", "left", "right"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), sz)
            el.set(qn("w:color"), color)
            borders.append(el)
        tcPr.append(borders)

    def cell_text(cell, text, bold=False, color=None, size=10,
                  align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Twips(20)
        p.paragraph_format.space_after  = Twips(20)
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color or BLACK

    def add_para(text="", bold=False, color=None, size=11, italic=False,
                 align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic
            run.font.size = Pt(size)
            run.font.color.rgb = color or BLACK
        return p

    def section_banner(label, sub=""):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.rows[0].cells[0]
        set_cell_bg(cell, NAVY_HEX)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Twips(30)
        p.paragraph_format.space_after  = Twips(30)
        r1 = p.add_run(label.upper() + "  ")
        r1.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = GOLD
        if sub:
            r2 = p.add_run(sub)
            r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)
        # remove outer borders
        for side in ["top","bottom","left","right"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "none")
            tbl._tbl.tblPr.append(el)

    # ── HEADER TABLE (logo + title block) ──────────────────────────────────────
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.style = "Table Grid"
    lc, rc = hdr_tbl.rows[0].cells
    set_cell_bg(lc, NAVY_HEX)
    set_cell_bg(rc, NAVY_HEX)
    lc.width = Inches(4.5)
    rc.width = Inches(2.0)

    # Logo image in left cell
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Twips(40)
    lp.paragraph_format.space_after  = Twips(20)
    if LOGO_PATH.exists():
        run_img = lp.add_run()
        run_img.add_picture(str(LOGO_PATH), width=Inches(3.4))
    else:
        r = lp.add_run("LIGHTHIEF")
        r.bold = True; r.font.size = Pt(22); r.font.color.rgb = GOLD
        lp.add_run("\nCyprus Ltd · Exclusive Linyang Energy Distributor").font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

    # Ref + date in right cell
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Twips(60)
    rp.paragraph_format.space_after  = Twips(20)
    r1 = rp.add_run("Ref: LH-CY-PVBESS-APR2026\n")
    r1.font.size = Pt(8); r1.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)
    r2 = rp.add_run("April 2026 · Strictly Confidential")
    r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)

    doc.add_paragraph()

    # ── DOCUMENT TITLE ─────────────────────────────────────────────────────────
    p_badge = add_para("INVESTMENT TEASER  ·  CYPRUS RES MARKET",
                       bold=True, color=RGBColor(0x9C, 0x7D, 0x22), size=8)
    p_badge.paragraph_format.space_after = Pt(2)

    p_title = add_para()
    r_t = p_title.add_run("Cyprus PV + Battery Energy Storage")
    r_t.bold = True; r_t.font.size = Pt(20); r_t.font.color.rgb = NAVY
    p_title.paragraph_format.space_after = Pt(3)

    p_sub = add_para("4-Hour Hybrid Connection  ·  8-Hour Flexible Nighttime Dispatch  ·  "
                     "AE Solar Meteor N-type TOPCon Bifacial  ·  ≤8 MW Standard EAC Tier",
                     color=RGBColor(0x7A, 0x8F, 0xA6), size=9)
    p_sub.paragraph_format.space_after = Pt(2)

    p_conf = add_para("Confidential - prepared by Lighthief Cyprus Ltd for qualified investors. "
                      "Indicative figures based on confirmed EPC pricing and TSOC market data; not a binding offer.",
                      italic=True, color=RGBColor(0xAA, 0xBB, 0xCC), size=8)
    p_conf.paragraph_format.space_after = Pt(6)

    # ── KPI TILES (1 row × 4 cols) ─────────────────────────────────────────────
    kpi_tbl = doc.add_table(rows=2, cols=4)
    kpi_tbl.style = "Table Grid"
    kpis = [
        ("DAM Evening Peak", "€183/MWh", "TSOC sample Oct–Feb 2026"),
        ("BESS Round-Trip Efficiency", "86.32%", "AC-AC · Linyang LFP"),
        ("PV Specific Yield", "2,150 kWh/kWp", "Bifacial TOPCon · 25° · albedo 0.70"),
        ("Park Size Cap", "≤ 8 MW", "Standard EAC connection tier"),
    ]
    for i, (label, val, note) in enumerate(kpis):
        c = kpi_tbl.rows[0].cells[i]
        set_cell_bg(c, NAVY_HEX)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Twips(40)
        p.paragraph_format.space_after  = Twips(10)
        p.add_run(label.upper() + "\n").font.color.rgb = RGBColor(0x88, 0x99, 0xAA)
        p.runs[0].font.size = Pt(7.5)
        rv = p.add_run(val + "\n")
        rv.bold = True; rv.font.size = Pt(14); rv.font.color.rgb = GOLD
        rn = p.add_run(note)
        rn.font.size = Pt(7.5); rn.font.color.rgb = RGBColor(0x66, 0x77, 0x88)

        # bottom row: merge for spacing
        cb = kpi_tbl.rows[1].cells[i]
        set_cell_bg(cb, NAVY_HEX)
        cb.paragraphs[0].paragraph_format.space_before = Twips(20)
        cb.paragraphs[0].paragraph_format.space_after  = Twips(20)

    doc.add_paragraph()

    # ── CONNECTION MODEL EXPLAINER ─────────────────────────────────────────────
    exp_tbl = doc.add_table(rows=1, cols=2)
    exp_tbl.style = "Table Grid"
    mA, mB = exp_tbl.rows[0].cells

    for cell, badge, title, desc, flow in [
        (mA, "MODEL A · 4H HYBRID",
         "Hybrid PV + BESS Connection",
         "Standard hybrid grid licence. PV exports directly during daytime at DAM rates. "
         "BESS captures curtailed energy (~65% curtailment) and discharges at evening peak. "
         "Dual revenue: direct PV sales + BESS arbitrage.",
         "☀ PV 35% → Grid  +  🔋 65% curtailed → BESS  →  🌙 €183/MWh night"),
        (mB, "MODEL B · 8H FLEXIBLE",
         "Flexible Nighttime-Only Dispatch",
         "CERA flexible connection: zero daytime grid export. 100% of PV generation "
         "flows into BESS. All energy discharged overnight at DAM peak. Operationally "
         "simpler; eliminates midday price exposure; larger BESS required (8h duration).",
         "☀ 100% PV → BESS  →  🌙 100% night dispatch  →  €183/MWh"),
    ]:
        set_cell_bg(cell, "F8FAFC")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Twips(40)
        p.paragraph_format.space_after  = Twips(10)
        rb = p.add_run(badge + "\n")
        rb.bold = True; rb.font.size = Pt(8); rb.font.color.rgb = NAVY
        rt = p.add_run(title + "\n")
        rt.bold = True; rt.font.size = Pt(11); rt.font.color.rgb = NAVY
        rd = p.add_run(desc + "\n")
        rd.font.size = Pt(9); rd.font.color.rgb = GREY
        rf = p.add_run(flow)
        rf.bold = True; rf.font.size = Pt(8.5); rf.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        p.paragraph_format.space_after = Twips(60)

    # Callout box
    doc.add_paragraph()
    cal_tbl = doc.add_table(rows=1, cols=1)
    cal_tbl.style = "Table Grid"
    cc = cal_tbl.rows[0].cells[0]
    set_cell_bg(cc, "FFF8E6")
    cp = cc.paragraphs[0]
    cp.paragraph_format.space_before = Twips(40)
    cp.paragraph_format.space_after  = Twips(40)
    cp.add_run("⚡  Why nighttime dispatch?  ").bold = True
    cp.runs[-1].font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    cp.runs[-1].font.size = Pt(9)
    r2 = cp.add_run("TSOC DAM sample (134 trading days, Oct 2025–Feb 2026): ")
    r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    r3 = cp.add_run("57 of 134 days recorded zero or negative midday prices. ")
    r3.bold = True; r3.font.size = Pt(9); r3.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    r4 = cp.add_run("Evening peak averaged €182.99/MWh vs €101.13/MWh at midday — an 81% premium. "
                    "BESS-backed nighttime dispatch eliminates midday price risk entirely.")
    r4.font.size = Pt(9); r4.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)

    doc.add_paragraph()

    # ── MAIN TABLES ────────────────────────────────────────────────────────────

    def build_model_table(scenarios, results, model_label, model_sub):
        section_banner(f"MODEL {model_label}", model_sub)
        doc.add_paragraph()

        labels = [s["label"] for s in scenarios]
        tbl = doc.add_table(rows=0, cols=4)
        tbl.style = "Table Grid"
        tbl.columns[0].width = Inches(2.9)
        for i in range(1, 4):
            tbl.columns[i].width = Inches(1.4)

        def add_header_row(texts, bg=NAVY_HEX, fglist=None):
            row = tbl.add_row()
            for i, txt in enumerate(texts):
                c = row.cells[i]
                set_cell_bg(c, bg)
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Twips(30)
                p.paragraph_format.space_after  = Twips(30)
                rr = p.add_run(txt)
                rr.bold = True; rr.font.size = Pt(8.5)
                if fglist:
                    rr.font.color.rgb = fglist[i]
                else:
                    rr.font.color.rgb = GOLD if i == 0 else WHITE

        def add_section_row(label):
            row = tbl.add_row()
            mc = row.cells[0].merge(row.cells[3])
            set_cell_bg(mc, "EEF3FB")
            p = mc.paragraphs[0]
            p.paragraph_format.space_before = Twips(20)
            p.paragraph_format.space_after  = Twips(20)
            r = p.add_run(label.upper())
            r.bold = True; r.font.size = Pt(7.5); r.font.color.rgb = NAVY

        def add_data_row(label, vals, bold=False, bg=None, color=None, size=9.5):
            row = tbl.add_row()
            cells = row.cells
            if bg:
                for c in cells: set_cell_bg(c, bg)
            c0 = cells[0]
            p0 = c0.paragraphs[0]
            p0.paragraph_format.space_before = Twips(25)
            p0.paragraph_format.space_after  = Twips(25)
            r0 = p0.add_run(label)
            r0.bold = bold; r0.font.size = Pt(size)
            r0.font.color.rgb = color or (NAVY if bold else GREY)
            for i, v in enumerate(vals, 1):
                p = cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_before = Twips(25)
                p.paragraph_format.space_after  = Twips(25)
                r = p.add_run(str(v))
                r.bold = bold; r.font.size = Pt(size)
                r.font.color.rgb = color or (NAVY if bold else GREY)

        # Column headers
        add_header_row(["Metric"] + labels)

        # CAPEX
        add_section_row("Capital Expenditure")
        add_data_row("PV EPC — AE Solar Meteor TOPCon (€720/kWp)",
                     [eur(r["pv_epc"]) for r in results])
        bess_label = "BESS EPC — Linyang LFP 4h (all-in turnkey)" if model_label == "A" \
                     else "BESS EPC — Linyang LFP 8h (V4 internal × 1.40 margin)"
        add_data_row(bess_label,
                     [f"{eur(r['bess_epc'])}  ({eur(r['bess_eur_mwh']/1000, 0)}k/MWh)" for r in results])
        add_data_row("RTB licence acquisition (€380k/MW)",
                     [eur(r["rtb"]) for r in results])
        add_data_row("EAC grid connection terms (€80k/MW)",
                     [eur(r["eac"]) for r in results])
        add_data_row("MV cable to substation (1 km × €20k/km — site variable)",
                     [eur(r["cable"]) for r in results])
        add_data_row("Permitting & documentation",
                     [eur(r["permitting"]) for r in results])
        add_data_row("Total CAPEX",
                     [eur(r["capex"]) for r in results],
                     bold=True, bg="EEF3FB", color=NAVY, size=10)

        # Revenue
        add_section_row("Annual Revenue (Year 1)")
        add_data_row("PV production (AE Meteor · 2,150 kWh/kWp)",
                     [f"{r['pv_prod_mwh']:,.0f} MWh" for r in results])
        if model_label == "A":
            add_data_row("Direct PV sales (35% uncurtailed × €140.88/MWh)",
                         [eur(r["direct_rev"]) for r in results])
            add_data_row("BESS arbitrage (65% × 95% capture × RTE 86.32% × €182.99)",
                         [eur(r["bess_rev"]) for r in results])
        else:
            add_data_row("BESS nighttime discharge (RTE 86.32% × 100% production)",
                         [f"{r['discharged']:,.0f} MWh" for r in results])
            add_data_row("Gross revenue (DAM evening peak × €182.99/MWh)",
                         [eur(r["bess_rev"]) for r in results])
        add_data_row("Less: aggregator fee (10%)",
                     [f"−{eur(r['gross_rev'] * AGGREGATOR_FEE)}" for r in results])
        add_data_row("Net Annual Revenue",
                     [eur(r["net_rev"]) for r in results],
                     bold=True, bg="EEF6EE", color=GREEN, size=10)

        # OPEX / EBITDA
        add_section_row("Operating Costs & EBITDA")
        add_data_row("PV O&M (€15/kWp/yr · cleaning + inspection)",
                     [f"−{eur(r['pv_om'])}" for r in results])
        add_data_row("BESS LTSA (€1,740/MWh/yr · 97% availability · 15yr)",
                     [f"−{eur(r['bess_om'])}" for r in results])
        add_data_row("EBITDA",
                     [eur(r["ebitda"]) for r in results],
                     bold=True, bg="EEF6EE", color=GREEN, size=10)

        # Returns
        add_header_row(["Return Metrics"] + [""] * 3,
                       bg=NAVY_HEX,
                       fglist=[GOLD, WHITE, WHITE, WHITE])
        add_data_row("Unlevered payback",
                     [f"{r['upayback']:.1f} yr" for r in results], bold=True, color=NAVY)
        add_data_row("Equity required (35% of CAPEX)",
                     [eur(r["equity"]) for r in results])
        add_data_row("Senior debt (65% · 5.5% p.a. · 12yr annuity)",
                     [eur(r["debt"]) for r in results])
        add_data_row("Annual debt service",
                     [f"−{eur(r['svc'])}" for r in results])
        add_data_row("DSCR",
                     [f"{r['dscr']:.2f}×" for r in results], bold=True, color=GREEN)
        add_data_row("Net cash to equity (post 15% CIT)",
                     [eur(r["net_cash"]) for r in results], color=GREEN)
        add_data_row("Equity payback (levered)",
                     [f"{r['epayback']:.1f} yr" for r in results],
                     bold=True, bg="FFF9E6", color=RGBColor(0x9C, 0x7D, 0x22), size=10)

    ra = [calc(s, "A") for s in MODEL_A_SCENARIOS]
    rb = [calc(s, "B") for s in MODEL_B_SCENARIOS]

    build_model_table(MODEL_A_SCENARIOS, ra, "A — 4-Hour Hybrid Connection",
                      "PV day sales + BESS curtailment arbitrage · Linyang LFP 4h")
    doc.add_paragraph()
    build_model_table(MODEL_B_SCENARIOS, rb, "B — 8-Hour Flexible Connection",
                      "100% nighttime dispatch · zero daytime export · Linyang LFP 8h")
    doc.add_paragraph()

    # ── TECHNOLOGY TABLE ───────────────────────────────────────────────────────
    section_banner("Technology & EPC Capability",
                   "Lighthief — exclusive Linyang Energy distributor in Cyprus · 28 parks / 134 MW / 496.5 MWh pipeline")
    doc.add_paragraph()

    tech_tbl = doc.add_table(rows=0, cols=4)
    tech_tbl.style = "Table Grid"
    for w, cw in zip(tech_tbl.columns, [Inches(1.5), Inches(2.0), Inches(1.5), Inches(2.0)]):
        w.width = cw

    # headers
    hrow = tech_tbl.add_row()
    for i, (txt, bg) in enumerate([("Solar PV — AE Solar Meteor", NAVY_HEX),
                                   ("", NAVY_HEX),
                                   ("BESS — Linyang LFP", NAVY_HEX),
                                   ("", NAVY_HEX)]):
        c = hrow.cells[i]
        set_cell_bg(c, bg)
        if txt:
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Twips(30)
            p.paragraph_format.space_after  = Twips(30)
            r = p.add_run(txt.upper())
            r.bold = True; r.font.size = Pt(8); r.font.color.rgb = GOLD

    tech_pv = [
        ("Module", "AE Solar Meteor N-type TOPCon bifacial"),
        ("Wattage / efficiency", "620–630 Wp · η 22.9%"),
        ("Bifaciality / albedo", "80% · white ground 0.70"),
        ("Specific yield", "2,150 kWh/kWp · 25° south"),
        ("Degradation", "0.4%/yr (N-type — low)"),
        ("Product warranty", "30 years"),
        ("PV EPC", "€720/kWp all-in turnkey"),
    ]
    tech_bess = [
        ("Chemistry / cells", "LFP (EVE) · EN 50549-2 TÜV cert."),
        ("PCS models", "1.0 / 1.25 MW C-series"),
        ("RTE", "86.32% AC-AC (full system)"),
        ("Cycle life", "7,000 cycles @ 70% EOL DoD"),
        ("SoH Yr5 / 10 / 15", "85% / 79.6% / 70% guaranteed"),
        ("Base warranty", "5 years (in EPC price)"),
        ("LTSA", "€1,740/MWh/yr · 97% avail. · 15yr"),
        ("Incoterms", "CIF Limassol · 90d prod + 50d transit"),
        ("Pipeline", "28 parks · 134 MW · 496.5 MWh confirmed"),
    ]
    for i in range(max(len(tech_pv), len(tech_bess))):
        row = tech_tbl.add_row()
        if i < len(tech_pv):
            set_cell_bg(row.cells[0], "F8FAFC")
            set_cell_bg(row.cells[1], "FFFFFF")
            cell_text(row.cells[0], tech_pv[i][0], bold=True, color=GREY, size=9)
            cell_text(row.cells[1], tech_pv[i][1], color=GREY, size=9)
        if i < len(tech_bess):
            set_cell_bg(row.cells[2], "F8FAFC")
            set_cell_bg(row.cells[3], "FFFFFF")
            cell_text(row.cells[2], tech_bess[i][0], bold=True, color=GREY, size=9)
            cell_text(row.cells[3], tech_bess[i][1], color=GREY, size=9)

    doc.add_paragraph()

    # ── TIMELINE ───────────────────────────────────────────────────────────────
    section_banner("Indicative Project Timeline",
                   "From RTB acquisition to Provisional Acceptance Certificate (PAC)")
    doc.add_paragraph()

    tl_tbl = doc.add_table(rows=1, cols=5)
    tl_tbl.style = "Table Grid"
    tl_steps = [
        ("RTB / EPC Contract", "Month 0"),
        ("Linyang Production", "+90 days"),
        ("CIF Limassol", "+50 days"),
        ("Installation & Commissioning", "+60–90 days"),
        ("PAC — Grid Connected", "~Month 9"),
    ]
    for i, (lbl, val) in enumerate(tl_steps):
        c = tl_tbl.rows[0].cells[i]
        set_cell_bg(c, "EEF3FB")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Twips(40)
        p.paragraph_format.space_after  = Twips(40)
        r1 = p.add_run(lbl + "\n")
        r1.font.size = Pt(8); r1.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        r2 = p.add_run(val)
        r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = NAVY

    note = add_para("MV cable: illustrative 1 km × €20,000/km — adjust for site distance. "
                    "RTB + EAC connection terms apply for standard tier ≤8 MW. "
                    "Above 8 MW, higher EAC connection cost tiers apply. "
                    "All figures indicative; formal EPC proposal on request.",
                    color=RGBColor(0xAA, 0xBB, 0xCC), size=8)
    note.paragraph_format.space_before = Pt(4)

    doc.add_paragraph()

    # ── FOOTER ─────────────────────────────────────────────────────────────────
    ft_tbl = doc.add_table(rows=1, cols=2)
    ft_tbl.style = "Table Grid"
    fl, fr = ft_tbl.rows[0].cells
    set_cell_bg(fl, NAVY_HEX)
    set_cell_bg(fr, NAVY_HEX)
    fl.width = Inches(4.0)

    p_fl = fl.paragraphs[0]
    p_fl.paragraph_format.space_before = Twips(50)
    p_fl.paragraph_format.space_after  = Twips(50)
    for label, val in [("Contact", "Alexander Papacosta"),
                       ("Phone", "+357 99 164 158"),
                       ("Email", "alexander.papacosta@lighthief.com"),
                       ("Web", "lighthief.com")]:
        rl = p_fl.add_run(f"{label}: ")
        rl.font.size = Pt(8); rl.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)
        rv = p_fl.add_run(f"{val}    ")
        rv.bold = True; rv.font.size = Pt(9); rv.font.color.rgb = WHITE

    p_fr = fr.paragraphs[0]
    p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fr.paragraph_format.space_before = Twips(50)
    p_fr.paragraph_format.space_after  = Twips(50)
    rl2 = p_fr.add_run("Lighthief Cyprus Ltd · HE 477423\n"
                        "15 Agaritsis, Nektaria Court, Office 201\n"
                        "3045 Zakaki, Limassol, Cyprus")
    rl2.font.size = Pt(8); rl2.font.color.rgb = RGBColor(0x66, 0x77, 0x88)

    doc.save(str(DOCX_OUT))
    print(f"  DOCX → {DOCX_OUT}")


# ══════════════════════════════════════════════════════════════════════════════
# 3. GENERATE EXCEL WORKBOOK
# ══════════════════════════════════════════════════════════════════════════════

def make_xlsx(results_a, results_b):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XLImage

    wb = openpyxl.Workbook()

    # ── style helpers ─────────────────────────────────────────────────────────

    def fill(hex_c):
        return PatternFill(start_color=hex_c, end_color=hex_c, fill_type="solid")

    def thin_border(color="D0DDEF"):
        s = Side(style="thin", color=color)
        return Border(left=s, right=s, top=s, bottom=s)

    def hfont(sz=10, bold=True, color=WHITE_HEX):
        return Font(name="Calibri", bold=bold, size=sz, color=color)

    def bfont(sz=10, bold=False, color=GREY_HEX):
        return Font(name="Calibri", bold=bold, size=sz, color=color)

    def center():
        return Alignment(horizontal="center", vertical="center", wrap_text=True)

    def left():
        return Alignment(horizontal="left", vertical="center", indent=1)

    def right():
        return Alignment(horizontal="right", vertical="center")

    def hcell(ws, row, col, val, bg=NAVY_HEX, fg=WHITE_HEX, sz=10, bold=True, align="center"):
        c = ws.cell(row=row, column=col, value=val)
        c.font   = Font(name="Calibri", bold=bold, size=sz, color=fg)
        c.fill   = fill(bg)
        c.border = thin_border()
        c.alignment = center() if align == "center" else (left() if align == "left" else right())
        return c

    def dcell(ws, row, col, val, fmt=None, bold=False, color=GREY_HEX,
              bg=None, align="left"):
        c = ws.cell(row=row, column=col, value=val)
        c.font   = Font(name="Calibri", bold=bold, size=10, color=color)
        c.border = thin_border()
        c.alignment = center() if align == "center" else (left() if align == "left" else right())
        if fmt:
            c.number_format = fmt
        if bg:
            c.fill = fill(bg)
        return c

    def section_row(ws, row, label, ncols=7):
        ws.row_dimensions[row].height = 16
        c = ws.cell(row=row, column=1, value=label.upper())
        c.font      = Font(name="Calibri", bold=True, size=8.5, color=NAVY_HEX)
        c.fill      = fill(LGREY_HEX)
        c.alignment = left()
        c.border    = thin_border()
        for col in range(2, ncols + 1):
            cc = ws.cell(row=row, column=col)
            cc.fill   = fill(LGREY_HEX)
            cc.border = thin_border()

    def total_row(ws, row, label, vals, ncols=7, bg=LGREY_HEX, color=NAVY_HEX, fmt='€#,##0'):
        ws.row_dimensions[row].height = 16
        c = ws.cell(row=row, column=1, value=label)
        c.font = Font(name="Calibri", bold=True, size=10.5, color=color)
        c.fill = fill(bg); c.border = thin_border(); c.alignment = left()
        for i, v in enumerate(vals):
            col = 2 + i * 2
            cc = ws.cell(row=row, column=col, value=v)
            cc.font = Font(name="Calibri", bold=True, size=10.5, color=color)
            cc.fill = fill(bg); cc.border = thin_border(); cc.alignment = right()
            cc.number_format = fmt
            # blank spacer
            sp = ws.cell(row=row, column=col + 1)
            sp.fill = fill(bg); sp.border = thin_border()

    # ── sheet builder ─────────────────────────────────────────────────────────

    def build_sheet(ws, scenarios, results, model_char, model_title, model_desc):
        ws.sheet_view.showGridLines = False

        # column widths
        ws.column_dimensions["A"].width = 42
        for i, w in enumerate([14, 1, 14, 1, 14, 1], start=2):
            ws.column_dimensions[get_column_letter(i)].width = w

        row = 1
        # ── header band ──
        ws.row_dimensions[row].height = 60
        # logo cell
        if LOGO_PATH.exists():
            try:
                img = XLImage(str(LOGO_PATH))
                img.width, img.height = 220, 44
                ws.add_image(img, "A1")
            except Exception:
                pass
        for col in range(1, 8):
            ws.cell(row=row, column=col).fill = fill(NAVY_HEX)
            ws.cell(row=row, column=col).border = thin_border(NAVY_HEX)
        # right-side ref
        ref_cell = ws.cell(row=row, column=6, value="Ref: LH-CY-PVBESS-APR2026")
        ref_cell.font = Font(name="Calibri", size=8, color="8899AA"); ref_cell.fill = fill(NAVY_HEX)
        ref_cell.alignment = right()
        date_cell = ws.cell(row=row, column=7, value="April 2026 · Confidential")
        date_cell.font = Font(name="Calibri", bold=True, size=9, color="CCDDE0"); date_cell.fill = fill(NAVY_HEX)
        date_cell.alignment = right()

        row += 1
        ws.row_dimensions[row].height = 28
        c = ws.cell(row=row, column=1, value=f"Cyprus PV + Battery Energy Storage — Model {model_char}: {model_title}")
        c.font      = Font(name="Calibri", bold=True, size=14, color=GOLD_HEX)
        c.fill      = fill(NAVY_HEX); c.border = thin_border(NAVY_HEX); c.alignment = left()
        for col in range(2, 8):
            ws.cell(row=row, column=col).fill = fill(NAVY_HEX)

        row += 1
        ws.row_dimensions[row].height = 18
        c = ws.cell(row=row, column=1, value=model_desc)
        c.font      = Font(name="Calibri", size=9, color="7A8FA6", italic=True)
        c.fill      = fill(NAVY_HEX); c.border = thin_border(NAVY_HEX); c.alignment = left()
        for col in range(2, 8):
            ws.cell(row=row, column=col).fill = fill(NAVY_HEX)

        row += 1
        # ── column headers ──
        ws.row_dimensions[row].height = 22
        hcell(ws, row, 1, "Metric", bg=NAVY_HEX, fg=GOLD_HEX, align="left")
        scen_labels = [s["label"] for s in scenarios]
        for i, lbl in enumerate(scen_labels):
            col = 2 + i * 2
            hcell(ws, row, col, lbl, bg=NAVY_HEX, fg=WHITE_HEX)
            ws.cell(row=row, column=col+1).fill = fill(NAVY_HEX)  # spacer

        # ── CAPEX ──
        row += 1; section_row(ws, row, "Capital Expenditure")
        rows_capex = [
            ("PV EPC — AE Solar Meteor TOPCon (€720/kWp)", [r["pv_epc"] for r in results]),
            (("BESS EPC — Linyang LFP 4h (all-in)" if model_char == "A"
              else "BESS EPC — Linyang LFP 8h (V4 internal × 1.40)"),
             [r["bess_epc"] for r in results]),
            ("RTB licence acquisition (€380k/MW)", [r["rtb"] for r in results]),
            ("EAC grid connection terms (€80k/MW)", [r["eac"] for r in results]),
            ("MV cable to substation (1 km × €20k)", [r["cable"] for r in results]),
            ("Permitting & documentation", [r["permitting"] for r in results]),
        ]
        for label, vals in rows_capex:
            row += 1; ws.row_dimensions[row].height = 15
            dcell(ws, row, 1, label)
            for i, v in enumerate(vals):
                col = 2 + i * 2
                dcell(ws, row, col, v, fmt='€#,##0', align="right")
                ws.cell(row=row, column=col+1).border = thin_border()
        row += 1; total_row(ws, row, "Total CAPEX", [r["capex"] for r in results])

        # ── Revenue ──
        row += 1; section_row(ws, row, "Annual Revenue (Year 1)")
        rev_rows = [
            (f"PV production (2,150 kWh/kWp)",
             [f"{r['pv_prod_mwh']:,.0f} MWh" for r in results]),
        ]
        if model_char == "A":
            rev_rows += [
                ("Direct PV sales (35% uncurtailed × €140.88/MWh)", [r["direct_rev"] for r in results]),
                ("BESS arbitrage (65% × 95% cap × RTE 86.32% × €182.99)", [r["bess_rev"] for r in results]),
            ]
        else:
            rev_rows += [
                ("BESS nighttime discharge (RTE 86.32%)",
                 [f"{r['discharged']:,.0f} MWh" for r in results]),
                ("Gross revenue (DAM evening peak × €182.99/MWh)", [r["bess_rev"] for r in results]),
            ]
        rev_rows += [
            ("Less: aggregator fee (10%)", [-r["gross_rev"] * AGGREGATOR_FEE for r in results]),
        ]
        for label, vals in rev_rows:
            row += 1; ws.row_dimensions[row].height = 15
            dcell(ws, row, 1, label)
            for i, v in enumerate(vals):
                col = 2 + i * 2
                if isinstance(v, str):
                    dcell(ws, row, col, v, align="right")
                else:
                    dcell(ws, row, col, v, fmt='€#,##0', align="right")
                ws.cell(row=row, column=col+1).border = thin_border()
        row += 1
        total_row(ws, row, "Net Annual Revenue",
                  [r["net_rev"] for r in results], color=GREEN_HEX, bg=CALC_HEX)

        # ── OPEX / EBITDA ──
        row += 1; section_row(ws, row, "Operating Costs & EBITDA")
        for label, vals in [
            ("PV O&M (€15/kWp/yr)", [-r["pv_om"] for r in results]),
            ("BESS LTSA (€1,740/MWh/yr · 97% avail. · 15yr)", [-r["bess_om"] for r in results]),
        ]:
            row += 1; ws.row_dimensions[row].height = 15
            dcell(ws, row, 1, label)
            for i, v in enumerate(vals):
                col = 2 + i * 2
                dcell(ws, row, col, v, fmt='€#,##0', align="right")
                ws.cell(row=row, column=col+1).border = thin_border()
        row += 1
        total_row(ws, row, "EBITDA",
                  [r["ebitda"] for r in results], color=GREEN_HEX, bg=CALC_HEX)

        # ── Returns ──
        row += 1; section_row(ws, row, "Return Metrics")
        for label, vals, fmt in [
            ("Unlevered payback", [f"{r['upayback']:.1f} yr" for r in results], "@"),
            ("Equity required (35%)", [r["equity"] for r in results], '€#,##0'),
            ("Senior debt (65% · 5.5% · 12yr)", [r["debt"] for r in results], '€#,##0'),
            ("Annual debt service", [-r["svc"] for r in results], '€#,##0'),
            ("DSCR", [r["dscr"] for r in results], '0.00"×"'),
            ("Net cash to equity (post 15% CIT)", [r["net_cash"] for r in results], '€#,##0'),
            ("Equity payback (levered)", [f"{r['epayback']:.1f} yr" for r in results], "@"),
        ]:
            row += 1; ws.row_dimensions[row].height = 15
            is_bold = label in ("DSCR", "Equity payback (levered)")
            color = GREEN_HEX if label == "DSCR" else (GOLD_HEX if "payback" in label.lower() else GREY_HEX)
            dcell(ws, row, 1, label, bold=is_bold)
            bg_ = INPUT_HEX if "payback" in label.lower() and "equity" in label.lower() else None
            for i, v in enumerate(vals):
                col = 2 + i * 2
                dcell(ws, row, col, v, fmt=fmt, bold=is_bold,
                      color=color, bg=bg_, align="right")
                ws.cell(row=row, column=col+1).border = thin_border()

        # ── Assumptions block ──
        row += 2; ws.row_dimensions[row].height = 14
        section_row(ws, row, "Key Assumptions")
        assump = [
            ("PV EPC cost", "€720/kWp all-in turnkey (AE Solar Meteor TOPCon bifacial)"),
            ("PV specific yield", "2,150 kWh/kWp (25° south, white albedo 0.70, bifaciality 80%)"),
            ("BESS RTE", "86.32% AC-AC round-trip efficiency (Linyang confirmed)"),
            ("DAM evening peak", "€182.99/MWh (TSOC sample Oct 2025–Feb 2026, 134 days)"),
            ("DAM daytime average", f"€{DAM_DAYTIME}/MWh (06:00–17:00 average)" if model_char == "A" else "N/A (100% BESS dispatch)"),
            ("Curtailment (Model A)", "65% (Agios Theodoros / Galascope validated model)" if model_char == "A" else "N/A — 100% stored in BESS"),
            ("BESS capture rate (Model A)", "95% of curtailed energy captured" if model_char == "A" else "N/A"),
            ("Aggregator fee", "10% of gross revenue"),
            ("RTB acquisition", "€380k/MW"),
            ("EAC connection", "€80k/MW (standard tier ≤8 MW)"),
            ("MV cable", f"€20k/km · illustrative {MV_CABLE_KM} km — site-specific"),
            ("Finance", "65% LTV · 5.5% p.a. · 12yr annuity (factor 0.1160)"),
            ("Depreciation", "20 years straight-line on total CAPEX"),
            ("Cyprus CIT", "15% (from 1 Jan 2026, OECD Pillar II)"),
            ("BESS LTSA", "€1,740/MWh/yr · 97% availability · 15yr · Lighthief Cyprus"),
            ("PV O&M", "€15/kWp/yr · cleaning + inspection"),
        ]
        for i, (k, v) in enumerate(assump):
            row += 1; ws.row_dimensions[row].height = 14
            ws.cell(row=row, column=1, value=k).font = Font(name="Calibri", bold=True, size=9, color=GREY_HEX)
            ws.cell(row=row, column=1).border = thin_border()
            ws.cell(row=row, column=1).alignment = left()
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=7)
            vc = ws.cell(row=row, column=2, value=v)
            vc.font = Font(name="Calibri", size=9, color=GREY_HEX)
            vc.border = thin_border(); vc.alignment = left()

        # ── footer ──
        row += 2
        ws.row_dimensions[row].height = 20
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
        fc = ws.cell(row=row, column=1,
                     value="Lighthief Cyprus Ltd · HE 477423 · +357 99 164 158 · "
                           "alexander.papacosta@lighthief.com · lighthief.com · "
                           "Strictly Confidential — April 2026")
        fc.font = Font(name="Calibri", size=8, color="AABBCC")
        fc.fill = fill(NAVY_HEX); fc.alignment = center()

    # ── build sheets ──────────────────────────────────────────────────────────
    ws_a = wb.active
    ws_a.title = "Model A — 4h Hybrid"
    ra = [calc(s, "A") for s in MODEL_A_SCENARIOS]
    build_sheet(ws_a, MODEL_A_SCENARIOS, ra, "A", "4-Hour Hybrid Connection",
                "PV day sales + BESS curtailment arbitrage · AE Solar Meteor TOPCon · Linyang LFP 4h")

    ws_b = wb.create_sheet("Model B — 8h Flexible")
    rb = [calc(s, "B") for s in MODEL_B_SCENARIOS]
    build_sheet(ws_b, MODEL_B_SCENARIOS, rb, "B", "8-Hour Flexible Connection",
                "100% nighttime BESS dispatch · zero daytime export · V4 internal costs × 1.40 · Linyang LFP 8h")

    wb.save(str(XLSX_OUT))
    print(f"  XLSX -> {XLSX_OUT}")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("\nLighthief PV+BESS Flexible Connection Teaser Generator")
    print("=" * 56)

    print("\n[1/3] Generating logo PNG...")
    make_logo()

    ra = [calc(s, "A") for s in MODEL_A_SCENARIOS]
    rb = [calc(s, "B") for s in MODEL_B_SCENARIOS]

    print("\n[2/3] Generating Word document...")
    make_docx(ra, rb)

    print("\n[3/3] Generating Excel workbook...")
    make_xlsx(ra, rb)

    print("\nDone.")
    print(f"  DOCX: {DOCX_OUT}")
    print(f"  XLSX: {XLSX_OUT}")
    print(f"  Logo: {LOGO_PATH}")

    # Print key metrics summary
    print("\n--- Model A (4h Hybrid) ---")
    for s, r in zip(MODEL_A_SCENARIOS, ra):
        print(f"  {s['label']:22s}  CAPEX {eurf(r['capex'])}  "
              f"EBITDA {eurf(r['ebitda'])}  "
              f"Upayback {r['upayback']:.1f}yr  DSCR {r['dscr']:.2f}x  "
              f"Epayback {r['epayback']:.1f}yr")
    print("\n--- Model B (8h Flexible) ---")
    for s, r in zip(MODEL_B_SCENARIOS, rb):
        print(f"  {s['label']:22s}  CAPEX {eurf(r['capex'])}  "
              f"EBITDA {eurf(r['ebitda'])}  "
              f"Upayback {r['upayback']:.1f}yr  DSCR {r['dscr']:.2f}x  "
              f"Epayback {r['epayback']:.1f}yr")
