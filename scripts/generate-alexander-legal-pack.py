#!/usr/bin/env python3
"""Generate combined Word legal pack from Alexander/ folder documents."""

from __future__ import annotations

import html as html_lib
import os
import re
import subprocess
import sys
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
ALEXANDER = ROOT / "Alexander"
OUT_FILE = ALEXANDER / "Alexander-Legal-Pack-May2026.docx"
OUT_FILE_ALT = ALEXANDER / "Alexander-Legal-Pack-May2026-v10.docx"

LOI_SEARCH_DIRS = [
    ROOT / "docs" / "clients" / "group-order" / "Group2_Esperia_Energy" / "contracts",
    ROOT / "docs" / "clients" / "group-order" / "Group2_Esperia_Energy",
]

NAVY = RGBColor(0x1A, 0x36, 0x5D)
GOLD = RGBColor(0xC9, 0xA4, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x40, 0x40, 0x40)
NAVY_HEX = "1A365D"


def add_run(para, text, *, bold=False, italic=False, size_pt=10, color=BLACK, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = font
    run.font.color.rgb = color
    return run


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def restart_section_page_numbering(section):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), "1")


def add_page_number_field(paragraph):
    """Insert Word PAGE field."""
    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    text_run = OxmlElement("w:r")
    text_run.append(fld_text)
    r.addnext(text_run)
    r.addnext(fld_end)


def configure_section_footer(section, label: str, *, linked_previous: bool = False):
    section.footer.is_linked_to_previous = linked_previous
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, label + "  |  Page ", size_pt=8, color=GREY)
    add_page_number_field(p)


def setup_document_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def start_new_section(doc: Document, footer_label: str) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    restart_section_page_numbering(section)
    configure_section_footer(section, footer_label, linked_previous=False)


def add_doc_title(doc: Document, title: str, subtitle: str = "", ref: str = ""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title.upper(), bold=True, size_pt=14, color=NAVY)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, subtitle, italic=True, size_pt=11, color=GREY)
    if ref:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p3, ref, size_pt=9, color=GREY)
    doc.add_paragraph()


def add_confidential_banner(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "STRICTLY CONFIDENTIAL — LEGALLY PRIVILEGED — DRAFT FOR COUNSEL REVIEW", bold=True, size_pt=9, color=RGBColor(0x99, 0x00, 0x00))


def inline_to_text(element) -> str:
    if isinstance(element, NavigableString):
        return str(element)
    if not isinstance(element, Tag):
        return ""
    if element.name == "br":
        return "\n"
    parts = []
    for child in element.children:
        parts.append(inline_to_text(child))
    return "".join(parts)


def add_formatted_runs(paragraph, element, *, base_size=10, base_bold=False, base_italic=False):
    if isinstance(element, NavigableString):
        text = html_lib.unescape(str(element))
        if text:
            add_run(paragraph, text, bold=base_bold, italic=base_italic, size_pt=base_size)
        return
    if not isinstance(element, Tag):
        return
    if element.name in ("strong", "b"):
        for child in element.children:
            add_formatted_runs(paragraph, child, base_size=base_size, base_bold=True, base_italic=base_italic)
        return
    if element.name in ("em", "i"):
        for child in element.children:
            add_formatted_runs(paragraph, child, base_size=base_size, base_bold=base_bold, base_italic=True)
        return
    if element.name == "span":
        for child in element.children:
            add_formatted_runs(paragraph, child, base_size=base_size, base_bold=base_bold, base_italic=base_italic)
        return
    text = inline_to_text(element)
    if text:
        add_run(paragraph, html_lib.unescape(text), bold=base_bold, italic=base_italic, size_pt=base_size)


def add_html_table(doc: Document, table_tag: Tag, *, font_size=9):
    rows = table_tag.find_all("tr")
    if not rows:
        return
    col_count = max(len(r.find_all(["td", "th"])) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for ci in range(col_count):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            if ci < len(cells):
                src = cells[ci]
                p = cell.paragraphs[0]
                add_formatted_runs(p, src, base_size=font_size)
                classes = src.get("class") or []
                if "r" in classes or src.name == "th":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if "r" in classes else WD_ALIGN_PARAGRAPH.LEFT
                if src.name == "th" or ri == 0 and row.find("th"):
                    set_cell_bg(cell, NAVY_HEX)
                    for run in p.runs:
                        run.font.color.rgb = WHITE
                        run.bold = True
    doc.add_paragraph()


def render_html_fragment(doc: Document, html_fragment: str):
    soup = BeautifulSoup(f"<body>{html_fragment}</body>", "html.parser")
    body = soup.body
    if not body:
        return

    skip_classes = {"print-btn", "no-print", "footer", "sig-block", "sig-party", "sig-line"}

    for node in body.children:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                p = doc.add_paragraph()
                add_run(p, text, size_pt=10)
            continue
        if not isinstance(node, Tag):
            continue
        if node.name == "button":
            continue
        if node.get("class") and "footer" in node.get("class"):
            continue

        if node.name == "div":
            classes = node.get("class") or []
            if "confidential" in classes:
                add_confidential_banner(doc)
                continue
            if "header-bar" in classes:
                h1 = node.find("h1")
                sub = node.find(class_="subtitle")
                if h1:
                    add_doc_title(doc, inline_to_text(h1), inline_to_text(sub) if sub else "")
                continue
            if "parties-block" in classes:
                for box in node.find_all(class_="party-box"):
                    h3 = box.find("h3")
                    title = inline_to_text(h3).strip() if h3 else "Party"
                    p = doc.add_paragraph()
                    add_run(p, title, bold=True, size_pt=10, color=NAVY)
                    for line in box.find_all("p"):
                        if line.find("h3"):
                            continue
                        tp = doc.add_paragraph()
                        add_formatted_runs(tp, line, base_size=9)
                doc.add_paragraph()
                continue
            if "part-header" in classes:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                add_run(p, inline_to_text(node).strip(), bold=True, size_pt=11, color=WHITE)
                # simulate navy bar via shading on paragraph not easy — bold navy text instead
                for run in p.runs:
                    run.font.color.rgb = NAVY
                continue
            if "protection-box" in classes or "schedule" in classes or "resolution-box" in classes or "undertaking-box" in classes:
                for child in node.children:
                    if isinstance(child, Tag):
                        render_html_fragment(doc, str(child))
                continue
            if "sig-block" in node.get("class", []):
                continue
            # generic div — recurse children
            inner = "".join(str(c) for c in node.children if not (isinstance(c, Tag) and c.name == "button"))
            if inner.strip():
                render_html_fragment(doc, inner)
            continue

        if node.name == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            add_run(p, inline_to_text(node).strip(), bold=True, size_pt=12, color=GOLD)
            continue
        if node.name == "h3":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            add_run(p, inline_to_text(node).strip(), bold=True, size_pt=11, color=NAVY)
            continue
        if node.name == "p":
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_formatted_runs(p, node, base_size=10)
            continue
        if node.name == "table":
            add_html_table(doc, node)
            continue
        if node.name == "hr":
            doc.add_paragraph()
            continue


def split_commission_html(path: Path) -> dict[str, str]:
    text = path.read_text(encoding="utf-8")
    body_match = re.search(r"<body>(.*)</body>", text, re.DOTALL | re.IGNORECASE)
    body = body_match.group(1) if body_match else text

    markers = [
        ("header", r"<div class=\"confidential\">"),
        ("main", r"<!-- ═══ PARTIES ═══ -->"),
        ("schedule1", r"<!-- ═══ SCHEDULE 1 ═══ -->"),
        ("schedule2", r"<!-- ═══ SCHEDULE 2 — MARGIN FORMULA ═══ -->"),
        ("schedule3", r"<!-- ═══ SCHEDULE 3 — SCOPE OF SERVICES ═══ -->"),
        ("execution", r"<!-- ═══ SIGNATURES — AGREEMENT ═══ -->"),
        ("annex_a", r"<!-- ═══ ANNEX A — SHAREHOLDER RESOLUTION ═══ -->"),
        ("annex_b", r"<!-- ═══ ANNEX B — UBO PERSONAL UNDERTAKING ═══ -->"),
        ("end", r"<!-- ═══ FOOTER ═══ -->"),
    ]

    positions = []
    for key, pattern in markers:
        m = re.search(pattern, body)
        if m:
            positions.append((m.start(), key))

    positions.sort()
    sections: dict[str, str] = {}
    for i, (start, key) in enumerate(positions):
        if key == "end":
            break
        end = positions[i + 1][0] if i + 1 < len(positions) else len(body)
        chunk = body[start:end]
        if key == "header":
            sections.setdefault("main", "")
            sections["main"] = chunk + sections.get("main", "")
        elif key == "main":
            sections["main"] = sections.get("main", "") + chunk
        else:
            sections[key] = chunk

    # Main should stop before schedule 1 — already handled by split
    # Merge execution signatures into main or keep separate — user asked separate annexes; execution stays own section
    return sections


def add_signature_blocks_commission(doc: Document):
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run(p, "EXECUTION", bold=True, size_pt=12, color=GOLD)
    doc.add_paragraph()
    add_run(doc.add_paragraph(), "This Agreement has been executed as a Deed on the date first written above.", size_pt=10)

    for party, co, name in (
        ("LIGHTHIEF CYPRUS LTD", "HE 477423", "Alexander Papacosta — Director"),
        ("MOI OSTROV LTD", "HE 319483", "Alexander Papacosta — Director"),
    ):
        doc.add_paragraph()
        p = doc.add_paragraph()
        add_run(p, f"Signed for and on behalf of {party} ({co})", bold=True, size_pt=10)
        doc.add_paragraph()
        doc.add_paragraph()
        sig = doc.add_paragraph()
        add_run(sig, "______________________________", size_pt=10)
        add_run(doc.add_paragraph(), name, size_pt=10)
        add_run(doc.add_paragraph(), "Date: ______________________", size_pt=10)

    doc.add_paragraph()
    add_run(doc.add_paragraph(), "Witness:", bold=True, size_pt=10)
    add_run(doc.add_paragraph(), "Name: ______________________", size_pt=10)
    add_run(doc.add_paragraph(), "Address: ______________________", size_pt=10)
    add_run(doc.add_paragraph(), "Date: ______________________", size_pt=10)


def render_markdown(doc: Document, md_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            add_doc_title(doc, line[2:].strip())
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            add_run(p, line[3:].strip(), bold=True, size_pt=12, color=GOLD)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            add_run(p, line[4:].strip(), bold=True, size_pt=11, color=NAVY)
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[-| :]+\|", lines[i + 1]):
            table_lines = [line]
            i += 1
            i += 1  # skip separator
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            if rows:
                col_count = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=col_count)
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(col_count):
                        txt = row[ci] if ci < len(row) else ""
                        txt = re.sub(r"\*\*(.+?)\*\*", r"\1", txt)
                        cell = tbl.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_run(p, txt, bold=(ri == 0), size_pt=9, color=WHITE if ri == 0 else BLACK)
                        if ri == 0:
                            set_cell_bg(cell, NAVY_HEX)
                doc.add_paragraph()
            continue
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            text = line[2:].strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            add_run(p, text, size_pt=10)
            i += 1
            continue
        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            text = line.lstrip("> ").strip()
            add_run(p, text, italic=True, size_pt=10, color=GREY)
            i += 1
            continue
        if line.strip() == "---":
            i += 1
            continue
        # bold lines
        text = line.strip()
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bold = text.startswith("**") or (line.startswith("**") and line.endswith("**"))
        add_run(p, text.strip("*"), bold=bold, size_pt=10)
        i += 1


def add_cover_page(doc: Document, sections: list[tuple[str, str, str]]):
    add_confidential_banner(doc)
    add_doc_title(
        doc,
        "Alexander Papacosta — Legal Document Pack",
        "Lighthief Cyprus Ltd · Founder Alignment · May 2026",
        "Combined draft for counsel review — not for execution without legal sign-off",
    )

    p = doc.add_paragraph()
    add_run(
        p,
        "Page 1 of the legal content begins with the Executive Summary on the following page. "
        "Each numbered Part thereafter starts on a new page with its own footer.",
        size_pt=9,
        color=GREY,
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    add_run(p, "DOCUMENT INDEX", bold=True, size_pt=12, color=NAVY)
    doc.add_paragraph()

    index_rows = [("—", "LCY-SUMMARY-2026", "Executive Summary — economics, signing, document map")] + list(sections)

    tbl = doc.add_table(rows=len(index_rows) + 1, cols=4)
    tbl.style = "Table Grid"
    headers = ("Part", "Reference", "Document", "Pagination")
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        cell.text = ""
        add_run(cell.paragraphs[0], h, bold=True, size_pt=9, color=WHITE)
        set_cell_bg(cell, NAVY_HEX)

    for ri, (part, ref, title) in enumerate(index_rows, start=1):
        row = tbl.rows[ri]
        vals = (part, ref, title, "See footer" if part == "—" else "Restarts at page 1")
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = ""
            add_run(cell.paragraphs[0], val, size_pt=9)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(
        note,
        "Each part below begins on a new page with its own footer and page numbering. "
        "Print or PDF each section separately using the footer label, or print the full pack for counsel review.",
        size_pt=9,
        color=GREY,
    )


PACK_SECTIONS = [
    ("R1", "LCY-PARK-REG-2026-001", "Park & Pipeline Register (51 + LOI parks)"),
    ("R2", "LCY-FUTURE-BIZ-2026-001", "Future Business — Protection for Not-Yet-Existing Work"),
    ("L1", "LCY-LOI-GAL-B1-2026-R15", "LOI — Galascope Ltd (G1 + G2)"),
    ("L2", "LCY-LOI-ESP-PIPELINE-2026-R15", "LOI — Esperia Energy Pipeline (9 parks)"),
    ("1", "LCY-MOI-TOCA-2026-001", "Team Operational Consultancy Agreement — Main Agreement"),
    ("2", "LCY-MOI-TOCA-2026-001 · Sch. 1", "Schedule 1 & 1A — Existing + Pipeline LOI Parks"),
    ("3", "LCY-MOI-TOCA-2026-001 · Sch. 2", "Schedule 2 — Net Margin Calculation Example"),
    ("4", "LCY-MOI-TOCA-2026-001 · Sch. 3", "Schedule 3 — Scope of Operational Services"),
    ("5", "LCY-MOI-TOCA-2026-001 · Execution", "Execution — Consultancy Agreement Signatures"),
    ("6", "LCY-MOI-TOCA-2026-001 · Annex A", "Annex A — Written Shareholder Resolution (Moi Ostrov)"),
    ("7", "LCY-MOI-TOCA-2026-001 · Annex B", "Annex B — UBO Personal Undertaking"),
    ("8", "LCY-EMP-AP-2026-001", "MD Employment Agreement (€2k net → €5k gross)"),
    ("9", "LCY-SHR-EMP-2026-001", "Shareholder Resolution — MD Employment Approval"),
    ("10", "LCY-MD-EQ-2026-001", "MD Equity Participation Side Letter"),
    ("11", "LCY-OPT-A-2026-001", "Option A — Founder Runway & MD Economics"),
    ("12", "LCY-MOI-ADH-2026-001", "Deed of Group Adherence (template)"),
    ("13", "LCY-DIR-IND-2026-001", "Director & Officer Protections (D&O + indemnity)"),
    ("14", "LCY-SHR-DO-2026-001", "Shareholder Resolution — D&O Insurance & Indemnity"),
    ("15", "LCY-GOV-SIGN-2026-001", "Signing Matrix & Governance (who signs what)"),
    ("16", "LCY-BRD-2026-001", "Board Minutes — Phase 1 (Founder alignment pack)"),
    ("17", "LCY-BRD-2026-002", "Board Minutes — Phase 2 (Equity side letter / pre-investor)"),
]


def ensure_loi_docx_generated():
    """Regenerate Galascope + Esperia LOIs so pack embeds current quotes."""
    for script in ("generate-loi-galascope.py", "generate-loi-esperia-pipeline.py"):
        path = ROOT / "scripts" / script
        if path.exists():
            subprocess.run([sys.executable, str(path)], cwd=ROOT, check=False)


def resolve_loi_docx(filename: str) -> Path | None:
    candidates = [filename, filename.replace(".docx", "-UPDATED.docx")]
    for directory in LOI_SEARCH_DIRS:
        for name in candidates:
            p = directory / name
            if p.exists():
                return p
    return None


def append_docx_document(target: Document, source_path: Path | None, missing_hint: str):
    """Append all body elements from a source .docx (LOI) into target."""
    if source_path is None or not source_path.exists():
        p = target.add_paragraph()
        add_run(
            p,
            f"[Missing: {missing_hint}. Run: python scripts/generate-loi-galascope.py "
            f"and python scripts/generate-loi-esperia-pipeline.py]",
            italic=True,
            size_pt=10,
            color=GREY,
        )
        return
    src = Document(str(source_path))
    for element in src.element.body:
        target.element.body.append(element)


def build():
    doc = Document()
    setup_document_margins(doc)

    # Cover section — no page numbers
    configure_section_footer(doc.sections[0], "Alexander Legal Pack — Cover", linked_previous=False)
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.clear()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer_p, "Lighthief Cyprus Ltd · HE 477423 · solarfarms.cy · DRAFT v10 — May 2026", size_pt=8, color=GREY)

    ensure_loi_docx_generated()

    add_cover_page(doc, [(p, r, t) for p, r, t in PACK_SECTIONS])

    # Executive Summary — page 1 of legal content
    start_new_section(doc, "Executive Summary")
    render_markdown(doc, ALEXANDER / "PACK-SUMMARY-PAGE.md")

    # Part R1 — Park register
    start_new_section(doc, "Part R1 · LCY-PARK-REG · Park & Pipeline Register")
    render_markdown(doc, ALEXANDER / "SCHEDULE-1-AND-LOI-PARK-REGISTER.md")

    # Part R2 — Future business protection
    start_new_section(doc, "Part R2 · LCY-FUTURE-BIZ · Future Business Protection")
    render_markdown(doc, ALEXANDER / "FUTURE-BUSINESS-PROTECTION-NOTE.md")

    # Part L1 — Galascope LOI
    start_new_section(doc, "Part L1 · LCY-LOI-GAL-B1-2026-R15 · LOI Galascope Ltd")
    append_docx_document(
        doc,
        resolve_loi_docx("LOI-Galascope-Ltd-batch1-may2026.docx"),
        "LOI-Galascope-Ltd-batch1-may2026.docx",
    )

    # Part L2 — Esperia pipeline LOI
    start_new_section(doc, "Part L2 · LCY-LOI-ESP-PIPELINE · LOI Esperia Pipeline")
    append_docx_document(
        doc,
        resolve_loi_docx("LOI-Esperia-Energy-pipeline-may2026.docx"),
        "LOI-Esperia-Energy-pipeline-may2026.docx",
    )

    commission_path = ALEXANDER / "commission-agreement-lighthief-moiostrov.html"
    chunks = split_commission_html(commission_path)

    # Part 1 — Main agreement (through clause 12)
    start_new_section(doc, "Part 1 · LCY-MOI-TOCA-2026-001 · Main Agreement")
    render_html_fragment(doc, chunks.get("main", ""))

    # Part 2-4 — Schedules
    for idx, key, footer in (
        ("2", "schedule1", "Part 2 · LCY-MOI-TOCA-2026-001 · Schedule 1 & 1A"),
        ("3", "schedule2", "Part 3 · LCY-MOI-TOCA-2026-001 · Schedule 2"),
        ("4", "schedule3", "Part 4 · LCY-MOI-TOCA-2026-001 · Schedule 3"),
    ):
        start_new_section(doc, footer)
        render_html_fragment(doc, chunks.get(key, ""))

    # Part 5 — Execution
    start_new_section(doc, "Part 5 · LCY-MOI-TOCA-2026-001 · Execution")
    add_signature_blocks_commission(doc)

    # Part 6 — Annex A (signature block is in HTML only — do not duplicate)
    start_new_section(doc, "Part 6 · LCY-MOI-TOCA-2026-001 · Annex A")
    render_html_fragment(doc, chunks.get("annex_a", ""))

    # Part 7 — Annex B (signature block is in HTML only — do not duplicate)
    start_new_section(doc, "Part 7 · LCY-MOI-TOCA-2026-001 · Annex B")
    render_html_fragment(doc, chunks.get("annex_b", ""))

    # Part 8 — MD Employment Agreement
    start_new_section(doc, "Part 8 · LCY-EMP-AP-2026-001 · MD Employment Agreement")
    render_markdown(doc, ALEXANDER / "DRAFT-MD-EMPLOYMENT-AGREEMENT.md")
    doc.add_paragraph()
    add_run(doc.add_paragraph(), "SIGNATURES", bold=True, size_pt=12, color=GOLD)
    for label in (
        "Employer — Lighthief Cyprus Ltd (authorised by shareholder resolution)",
        "Employee — Alexander Papacosta (Cyprus ID: 1241318)",
    ):
        doc.add_paragraph()
        add_run(doc.add_paragraph(), label, bold=True, size_pt=10)
        add_run(doc.add_paragraph(), "Signature: ______________________   Date: ______________________", size_pt=10)

    # Part 9 — Shareholder resolution (MD employment)
    start_new_section(doc, "Part 9 · LCY-SHR-EMP-2026-001 · Shareholder Resolution (MD Employment)")
    render_markdown(doc, ALEXANDER / "DRAFT-SHAREHOLDER-RESOLUTION-MD-EMPLOYMENT.md")

    # Part 10 — Equity side letter
    start_new_section(doc, "Part 10 · LCY-MD-EQ-2026-001 · Equity Participation Side Letter")
    render_markdown(doc, ALEXANDER / "DRAFT-EQUITY-PARTICIPATION-SIDE-LETTER.md")
    doc.add_paragraph()
    add_run(doc.add_paragraph(), "SIGNATURES", bold=True, size_pt=12, color=GOLD)
    for label in (
        "Founder: Dr. Arkadiusz Sybaris for Lighthief International Ltd",
        "Company: Lighthief Cyprus Ltd — Director",
        "MD: Alexander Papacosta",
    ):
        doc.add_paragraph()
        add_run(doc.add_paragraph(), label, bold=True, size_pt=10)
        add_run(doc.add_paragraph(), "Signature: ______________________   Date: ______________________", size_pt=10)

    # Part 11 — Option A
    start_new_section(doc, "Part 11 · LCY-OPT-A-2026-001 · Founder Runway & MD Economics")
    render_markdown(doc, ALEXANDER / "DRAFT-OPTION-A-FOUNDER-RUNWAY-AND-MD-TERMS.md")
    doc.add_paragraph()
    add_run(doc.add_paragraph(), "SIGNATURES", bold=True, size_pt=12, color=GOLD)
    for label in ("Company — Lighthief Cyprus Ltd", "Founder", "MD — Alexander Papacosta"):
        doc.add_paragraph()
        add_run(doc.add_paragraph(), label, bold=True, size_pt=10)
        add_run(doc.add_paragraph(), "Signature: ______________________   Date: ______________________", size_pt=10)

    # Part 12 — Group adherence template
    start_new_section(doc, "Part 12 · LCY-MOI-ADH-2026-001 · Deed of Group Adherence")
    render_markdown(doc, ALEXANDER / "DRAFT-DEED-OF-GROUP-ADHERENCE.md")

    # Part 13 — Director & officer protections
    start_new_section(doc, "Part 13 · LCY-DIR-IND-2026-001 · Director & Officer Protections")
    render_markdown(doc, ALEXANDER / "DRAFT-DIRECTOR-AND-OFFICER-PROTECTIONS.md")

    # Part 14 — D&O shareholder resolution
    start_new_section(doc, "Part 14 · LCY-SHR-DO-2026-001 · Shareholder Resolution (D&O)")
    render_markdown(doc, ALEXANDER / "DRAFT-SHAREHOLDER-RESOLUTION-DO-INSURANCE.md")

    # Part 15 — Signing matrix & governance
    start_new_section(doc, "Part 15 · LCY-GOV-SIGN-2026-001 · Signing Matrix & Governance")
    render_markdown(doc, ALEXANDER / "DRAFT-SIGNING-MATRIX-AND-GOVERNANCE.md")

    # Part 16 — Board minutes Phase 1
    start_new_section(doc, "Part 16 · LCY-BRD-2026-001 · Board Minutes (Phase 1)")
    render_markdown(doc, ALEXANDER / "DRAFT-LCY-BOARD-MINUTES-PHASE1-2026-001.md")

    # Part 17 — Board minutes Phase 2
    start_new_section(doc, "Part 17 · LCY-BRD-2026-002 · Board Minutes (Phase 2)")
    render_markdown(doc, ALEXANDER / "DRAFT-LCY-BOARD-MINUTES-PHASE2-2026-002.md")

    try:
        doc.save(OUT_FILE)
        return OUT_FILE
    except PermissionError:
        doc.save(OUT_FILE_ALT)
        return OUT_FILE_ALT


if __name__ == "__main__":
    out = build()
    print(f"Generated: {out}")
