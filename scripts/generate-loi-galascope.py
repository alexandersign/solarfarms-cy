"""
Generate LOI — Galascope Ltd only (G1 + G2).

Output: docs/clients/group-order/Group2_Esperia_Energy/contracts/
        LOI-Galascope-Ltd-batch1-may2026.docx

Esperia pipeline parks are in generate-loi-esperia-pipeline.py (separate LOI).
"""

import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from loi_docx_common import (  # noqa: E402
    AMBER_HDR,
    AMBER_HEX,
    AMBER_TXT,
    CLIENT_REP_TITLE_GALASCOPE,
    GALASCOPE,
    GREY,
    client_rep_line,
    LIGHT_HEX,
    NAVY,
    NAVY_HEX,
    SCHEDULE_TABLE_COLS,
    SCHEDULE_TABLE_RC,
    WHITE,
    add_clause_1_3,
    add_clause_6_good_faith,
    add_clause_7_exclusivity,
    add_footer,
    add_header_bar,
    add_parties_table,
    add_price_mechanism_section,
    add_run,
    add_signatures,
    body,
    fmt_eur,
    galascope_totals,
    h1,
    lock_table_widths,
    new_document,
    set_cell_bg,
    tbl_hdr,
    tbl_row,
)
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

REF = "Ref: LCY-LOI-GAL-B1-2026-R15"
EPC_REF = "LCY-EPC-GAL-B1-2026"
GT = galascope_totals()


def build():
    doc = new_document()
    add_header_bar(doc, REF)

    tp = doc.add_paragraph()
    add_run(tp, "Letter of Intent", bold=True, size_pt=18, color=NAVY)
    sp = doc.add_paragraph()
    add_run(
        sp,
        "Galascope Ltd — BESS EPC Batch 1 (Galascope 1 + Galascope 2)",
        italic=True,
        size_pt=11,
        color=GREY,
    )

    add_parties_table(
        doc,
        "Galascope Ltd\nRegistration No. HE 303759\n"
        "Karaiskaki 6, City House, 3032 Limassol, Cyprus\n"
        + client_rep_line(CLIENT_REP_TITLE_GALASCOPE)
        + "(hereinafter \"the Client\")",
        client_label="Client",
    )

    h1(doc, "RECITALS")
    for letter, text in [
        ("A.", "Lighthief supplies and installs grid-connected BESS for hybrid PV projects in Cyprus."),
        (
            "B.",
            "The Client (Galascope Ltd) is developing two BESS parks at Famagusta (Galascope 1 and "
            "Galascope 2) as part of the Esperia Energy Group portfolio.",
        ),
        (
            "C.",
            f"The Client wishes to proceed with EPC ref. {EPC_REF} for Galascope 1 (5.0 MW / 20.06 MWh "
            f"effective) and Galascope 2 (2.5 MW / 10.03 MWh effective), combined indicative value "
            f"{fmt_eur(GT['val'])} (excl. VAT).",
        ),
        (
            "D.",
            "Other Esperia group pipeline parks are covered under a separate Letter of Intent with "
            "Esperia Energy Group (Ref: LCY-LOI-ESP-PIPELINE-2026-R15).",
        ),
    ]:
        rp = doc.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(rp, letter + "  ", bold=True, color=NAVY)
        add_run(rp, text)
    body(doc, "NOW THEREFORE, the Parties agree as follows:", space_after=8)

    h1(doc, "1.   COMMITMENT AND PURPOSE")
    body(
        doc,
        f"1.1  The Client confirms its intention to award the BESS EPC contract to Lighthief for "
        f"two (2) parks in Schedule 1 under {EPC_REF}: Galascope 1 and Galascope 2 "
        f"({GT['mw']:.2f} MW; effective MWh per Schedule A — 20.06 + 10.03). "
        f"Indicative value {fmt_eur(GT['val'])} (excl. VAT), confirmed May 2026.",
    )
    body(
        doc,
        "1.2  This LOI supports procurement, OEM planning, and lender review. Except where "
        "expressly binding below, obligations arise only when the EPC and companion documents are signed.",
    )
    add_clause_1_3(doc, reimbursement_cap=100_000)

    h1(doc, "2.   SCHEDULE 1 — PARKS AND INDICATIVE PRICING")

    mx_tbl = doc.add_table(rows=1, cols=4)
    for i, (val, lbl) in enumerate([
        ("2 Parks", "Galascope Batch 1"),
        (f"{GT['mw']:.2f} MW", "Total BESS Power"),
        (f"{sum(p['effective_mwh'] for p in GALASCOPE):.2f} MWh", "Total BESS Energy"),
        (fmt_eur(GT["val"]), "Indicative Value (ex. VAT)"),
    ]):
        cell = mx_tbl.rows[0].cells[i]
        set_cell_bg(cell, LIGHT_HEX)
        mp = cell.paragraphs[0]
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(mp, val + "\n", bold=True, size_pt=12, color=NAVY)
        add_run(mp, lbl, size_pt=7.5, color=GREY)
    lock_table_widths(mx_tbl, [3.5, 3.5, 3.5, 3.5])

    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_hdr(tbl, ["Park", "District", "MW", "MWh (eff.)", "Indicative Price (ex. VAT)"])
    for p in GALASCOPE:
        tbl_row(
            tbl,
            [p["name"], p["district"], f"{p['mw']:.2f}", f"{p['effective_mwh']:.2f}",
             fmt_eur(p["client_price"])],
            right_cols=SCHEDULE_TABLE_RC,
        )
    tbl_row(
        tbl,
        ["TOTAL — Galascope Batch 1", "", f"{GT['mw']:.2f}",
         f"{sum(p['effective_mwh'] for p in GALASCOPE):.2f}", fmt_eur(GT["val"])],
        right_cols=SCHEDULE_TABLE_RC,
        total=True,
    )
    lock_table_widths(tbl, SCHEDULE_TABLE_COLS)

    h1(doc, "3.   INDICATIVE PROGRAMME")
    body(doc, f"EPC target: {EPC_REF} — signing Q2 2026; target PAC 31 January 2027.")

    add_price_mechanism_section(doc)

    h1(doc, "5.   NEXT STEPS")
    body(
        doc,
        "5.1  The Parties will negotiate and sign the EPC package (EPC, LTSA, EMS addendum, and "
        "OEM warranties). Commercial terms for those documents are agreed in the EPC — not in "
        "this LOI.",
    )

    add_clause_6_good_faith(doc)

    add_clause_7_exclusivity(doc)

    h1(doc, "8.   CONFIDENTIALITY")
    body(
        doc,
        "8.1  (Binding) Each Party shall keep confidential this LOI and related commercial, "
        "technical, and pricing information, except to advisers, lenders, or as required "
        "by law. This survives for five (5) years after this LOI ends.",
    )

    h1(doc, "9.   TERM AND GOVERNING LAW")
    body(
        doc,
        "9.1  This LOI runs until the earlier of: EPC and companion documents signed, eighteen "
        "(18) months from signing, or written termination.",
    )
    body(doc, "9.2  Governed by Cyprus law. Courts of Cyprus have exclusive jurisdiction.")

    nb_tbl = doc.add_table(rows=1, cols=1)
    nb_cell = nb_tbl.rows[0].cells[0]
    set_cell_bg(nb_cell, AMBER_HEX)
    nbp = nb_cell.paragraphs[0]
    nbp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(nbp, "NATURE OF THIS DOCUMENT:  ", bold=True, size_pt=9, color=AMBER_HDR)
    add_run(
        nbp,
        "Non-binding except Clauses 1.3, 7, and 8. Covers Galascope 1 & 2 only. "
        "Binding commercial terms are in the executed EPC and companion agreements.",
        size_pt=9,
        color=AMBER_TXT,
    )
    lock_table_widths(nb_tbl, [14.0])

    add_signatures(doc, client_sig_name="Galascope Ltd")
    add_footer(doc, REF)
    return doc


def save_doc(doc, filename):
    base = os.path.join(
        os.path.dirname(__file__), "..",
        "docs", "clients", "group-order", "Group2_Esperia_Energy",
    )
    contracts_dir = os.path.join(base, "contracts")
    os.makedirs(contracts_dir, exist_ok=True)
    path = os.path.join(contracts_dir, filename)
    doc.save(path)
    print("Saved -> " + os.path.abspath(path))


if __name__ == "__main__":
    save_doc(build(), "LOI-Galascope-Ltd-batch1-may2026.docx")
