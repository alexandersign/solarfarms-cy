"""
Contract Generator v5.1 — Final Issue (May 2026)
=================================================
Generates final client contract packages from updated master template.

CHANGES vs v5.0:
1. §1A.1: Connection Terms longstop = 12 months (was [●])
2. §1A.3: Advance trigger = 30 days (was 7)
3. §6.1: Indicative Price + Confirmed Price mechanism (NEW)
4. §6.2/6.3: relabelled as Indicative Component A / B
5. §8.4.1(g): LD Value split (Comp A delivery / Comp B commissioning)
6. §8.4.6(vii): Upstream FM flow-through (NEW)
7. Schedule A: Price Basis Reference block (NEW)
8. Version history reference to "client lawyer" instead of named individual

Outputs:
- legal/templates/client_sales_v5.1.docx (master, with [●] placeholders)
- docs/clients/group-order/Group2_Esperia_Energy/contracts/
    EPC-Galascope-Esperia-batch1-may2026.docx (regenerated, v5.1)
- docs/clients/Individual_Spanercom/contracts/
    EPC-Spanercom-Anarita-may2026.docx (regenerated, v5.1)
- docs/clients/group-order/Group2_Esperia_Energy/contracts/
    LOI-Esperia-Energy-pipeline-may2026.docx (updated to match)
"""

import os, shutil
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

BASE = '/Volumes/T7 Grey/solinvest'
TEMPL = f'{BASE}/legal/templates'
GAL_OUT = f'{BASE}/docs/clients/group-order/Group2_Esperia_Energy/contracts'
SPA_OUT = f'{BASE}/docs/clients/Individual_Spanercom/contracts'
DATE = 'May 2026'

# Reference values (Jan 2026 baseline for price adjustment formula)
REF_LITHIUM_CNY = '155,000 CNY/tonne (Mysteel China battery-grade spot, January 2026 monthly average)'
REF_EUR_CNY = '8.18 CNY per EUR (January 2026 average)'

# ──────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────

def make_p(text, bold=False, after_pt=4):
    """Create a new w:p element with the given text."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    if bold:
        rpr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rpr.append(b)
        r.append(rpr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def set_p_text(para, text):
    """Replace all runs in a paragraph with a single new run containing text."""
    for el in list(para._p.findall(qn('w:r'))):
        para._p.remove(el)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    para._p.append(r)

def find_para_idx(doc, predicate):
    """Find first paragraph index matching predicate."""
    for i, p in enumerate(doc.paragraphs):
        if predicate(p.text):
            return i
    return None

def replace_text_in_doc(doc, old, new):
    """Replace text in all paragraphs and tables (handles multi-run paragraphs)."""
    for p in doc.paragraphs:
        if old in p.text and p.runs:
            # Try simple run-level replace first
            for r in p.runs:
                if old in r.text:
                    r.text = r.text.replace(old, new)
                    break
            else:
                # If split across runs, rewrite paragraph
                if old in p.text:
                    set_p_text(p, p.text.replace(old, new))
        elif old in p.text:
            # No runs — set directly
            set_p_text(p, p.text.replace(old, new))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        replace_text_in_doc_para(p, old, new)

def replace_text_in_doc_para(p, old, new):
    if old in p.text and p.runs:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
                return
        set_p_text(p, p.text.replace(old, new))

# ──────────────────────────────────────────────────────────────────────────
# STEP 1: Build v5.1 master template
# ──────────────────────────────────────────────────────────────────────────

def build_v51_template():
    """Read v5.0 template and apply structural changes; save as v5.1."""
    src = f'{TEMPL}/client_sales_v5.0.docx'
    dst = f'{TEMPL}/client_sales_v5.1.docx'
    shutil.copy(src, dst)

    doc = Document(dst)

    # 1. Update Document header — version, date, sanitize Anastasis reference
    for p in doc.paragraphs[:10]:
        if 'Document Reference: LCY-EPC-001' in p.text:
            new_header = (
                'Document Reference: LCY-EPC-001\n'
                'Version: 5.1\n'
                f'Date: {DATE}\n\n'
                'VERSION HISTORY\n'
                'v1.0 — Oct 2025: Initial draft\n'
                'v2.0 — Jan 2026: Linyang RFI confirmed terms\n'
                'v3.0 — 22 Feb 2026: Installation scope, delay LD framework, bankability\n'
                'v4.0 — 17 Mar 2026: Single counterparty warranty, title at PAC, unified warranty, performance guarantees\n'
                'v5.0 — 6 May 2026: Two-component price split, FAC removed (3-month DLP), tiered liability cap (10/50/uncapped), client lawyer review comments addressed\n'
                'v5.1 — May 2026: Indicative Price + Confirmed Price mechanism (raw material / FX adjustment), LD Value split (Comp A delivery / Comp B commissioning), upstream FM flow-through, 12-month Connection Terms longstop, 30-day advance payment trigger'
            )
            set_p_text(p, new_header)
            break

    # 2. §1A.1 — 12-month longstop
    for p in doc.paragraphs:
        if '1A.1' in p.text and 'Connection Terms' in p.text and '[●] months' in p.text:
            set_p_text(p, p.text.replace('[●] months', 'twelve (12) months'))
            break

    # 3. §1A.3 — 30-day advance trigger (only this one — keep §7.1(a) trigger same)
    for p in doc.paragraphs:
        if '1A.3' in p.text and 'seven (7) days' in p.text:
            set_p_text(p, p.text.replace('seven (7) days', 'thirty (30) days'))
            break

    # 4. Replace §6.1 (Contract Price) with Indicative + Confirmed mechanism
    # Find §6.1 paragraph
    sec61_idx = find_para_idx(doc, lambda t: t.startswith('6.1 The total contract price'))
    if sec61_idx is None:
        sec61_idx = find_para_idx(doc, lambda t: '6.1' in t and 'Contract Price' in t)
    if sec61_idx is not None:
        # Replace §6.1 text
        new_6_1 = (
            '6.1 Contract Price — Indicative and Confirmed\n'
            '(a) Indicative Contract Price: The Contract Price set out in Sections 6.2 and 6.3 is an Indicative Price '
            'based on Linyang Quotation LY202601271 (January 2026), exclusive of VAT. The Indicative Price is binding '
            'on the Contractor for the period between the Effective Date and the Confirmed Price Certificate Date '
            '(as defined below).\n'
            '(b) Confirmed Price Trigger: Within fourteen (14) days following the Client\'s receipt of Connection Terms '
            '(per Section 1A), the Contractor shall issue a Confirmed Price Certificate, fixing the final Contract Price '
            'to be applied to all subsequent payment milestones.\n'
            '(c) Adjustment Methodology: The Confirmed Price shall equal the Indicative Price, save where:\n'
            '  (i) the Linyang battery cell raw material reference index '
            f'(reference: {REF_LITHIUM_CNY}) has moved by more than ten percent (10%); or\n'
            '  (ii) the EUR/CNY exchange rate '
            f'(reference: {REF_EUR_CNY}) has moved by more than eight percent (8%); or\n'
            '  (iii) the equipment configuration in Schedule A is modified at the Client\'s request.\n'
            '(d) Adjustment Cap: Any upward adjustment to the Indicative Price shall not exceed five percent (5%) of '
            'the Indicative Contract Price. Where the calculated adjustment exceeds 5%, either Party may terminate this '
            'Agreement without penalty by written notice within fourteen (14) days of the Confirmed Price Certificate. '
            'Upon such termination, the Contractor shall refund any advance payments received within thirty (30) days, '
            'less reasonable and documented costs incurred prior to termination.\n'
            '(e) Pass-Through Limitation: The Contractor shall not pass through any cost increases that result from the '
            'Contractor\'s own delay or fault.\n'
            '(f) All milestone percentages in Section 7.1 are calculated on the Confirmed Contract Price (or, prior to '
            'issue of the Confirmed Price Certificate, on the Indicative Contract Price).'
        )
        set_p_text(doc.paragraphs[sec61_idx], new_6_1)

    # 5. Update §6.2 and §6.3 labels to "Indicative"
    for p in doc.paragraphs:
        if p.text.startswith('6.2 Component A — Equipment Supply Price'):
            set_p_text(p, '6.2 Indicative Component A — Equipment Supply Price')
        elif p.text.startswith('6.3 Component B — EPC Services Price'):
            set_p_text(p, '6.3 Indicative Component B — EPC Services Price')

    # 6. §8.4.1 — Update LD Value definition
    for p in doc.paragraphs:
        if '"LD Value"' in p.text and 'Full Park' in p.text and 'Affected Capacity' in p.text:
            new_ld = (
                '(g) "LD Value" — Calculated separately by delay type:\n'
                '  (i) Equipment Delivery Delay (delivery of equipment to Site delayed): '
                'Confirmed Component A Equipment Supply Price for Full Park; or '
                'Confirmed Component A Equipment Supply Price × Affected Capacity % for Partial Delay.\n'
                '  (ii) EPC Commissioning Delay (PAC not achieved after equipment delivered to Site for reasons within '
                'the Contractor\'s reasonable control): Confirmed Component B EPC Services Price for Full Park; or '
                'Confirmed Component B EPC Services Price × Affected Capacity % for Partial Delay.\n'
                'The combined Delay LDs across both categories shall not exceed the cap in Section 8.4.4.'
            )
            set_p_text(p, new_ld)
            break

    # 7. §8.4.6 — Add (vii) Upstream FM flow-through (insert AFTER the existing exclusions list)
    for i, p in enumerate(doc.paragraphs):
        if '8.4.6' in p.text and 'Exclusions' in p.text:
            # The exclusions list is in this paragraph
            # Append (vii) clause as a new paragraph immediately after
            existing = p.text
            # The current text ends with "...(vi) Client's late confirmation..." — append (vii)
            new_excl = existing.rstrip()
            if new_excl.endswith('.'):
                new_excl = new_excl[:-1]
            new_excl += '; (vii) Upstream Force Majeure declared by the OEM and notified to the Client per Section 8.4.7A.'
            set_p_text(p, new_excl)
            
            # Now insert a NEW paragraph for §8.4.7A immediately after
            anchor = doc.paragraphs[i + 1]._p
            new_clause = make_p(
                '8.4.7A Upstream Force Majeure (OEM Flow-Through). '
                'Where the OEM has formally notified the Contractor in writing of a Force Majeure event under the '
                'upstream supply contract, and such event directly causes a delay in delivery of equipment to Site, '
                'the Contractor may declare the same event as Force Majeure under this Agreement, provided the '
                'Contractor gives written notice to the Client within seven (7) Business Days of receiving the OEM\'s '
                'notification, including a copy of the OEM\'s Force Majeure notice. The Target PAC Date shall extend '
                'for the same duration as the OEM\'s Force Majeure period. No Delay LDs shall accrue during the '
                'Upstream Force Majeure period.'
            )
            anchor.addprevious(new_clause)
            break

    # 8. Schedule A — Add Price Basis Reference block (placeholder for filling per client)
    # The Schedule A heading has been customised in client docs; in master, it's a placeholder
    # We keep master template generic — actual reference values inserted at client doc generation

    doc.save(dst)
    print(f'  Saved master template: {dst}')
    return dst


# ──────────────────────────────────────────────────────────────────────────
# STEP 2: Render client EPC from v5.1 template
# ──────────────────────────────────────────────────────────────────────────

def fmt(n):
    return f'{n:,.2f}'

def render_epc(client_name, client_reg_line, client_address_block, parks, total_price,
               pac_date, ref, ltsa_ref, sig_name, sig_title, out_path):
    """Generate populated EPC v5.1 for a client."""
    doc = Document(f'{TEMPL}/client_sales_v5.1.docx')

    # Party block — Client
    replace_text_in_doc(doc, '[● End Customer Legal Name]', client_name)
    replace_text_in_doc(doc, 'a company incorporated under the laws of [●], with registered office at [●] ("Client").',
                        client_reg_line + ' ("Client").')

    # Document reference + date
    replace_text_in_doc(doc, 'Document Reference: LCY-EPC-001', f'Document Reference: {ref}')
    replace_text_in_doc(doc, '[●] 2026 ("Effective Date")', f'[●] {DATE} ("Effective Date")')
    replace_text_in_doc(doc, 'LCY-LTSA-001 v4.0', ltsa_ref)

    # Signature party line
    replace_text_in_doc(doc, 'For and on behalf of [● End Customer] ("Client"):',
                        f'For and on behalf of {client_name} ("Client"):')
    replace_text_in_doc(doc, '[● End Customer]', client_name)

    # Build Schedule A content (replace existing Schedule A from heading to before SCHEDULE B)
    sched_a_lines = [
        'SCHEDULE A — TECHNICAL SPECIFICATIONS, COMMERCIAL DETAILS AND PRICE BASIS',
        f'Reference: {ref} | Version: 5.1 | Date: {DATE}',
        '',
        '=' * 60,
        'PART 1 — PARK DETAILS',
        '=' * 60,
    ]
    for pk in parks:
        sched_a_lines += [
            '',
            f"Park Name:              {pk['name']}",
            f"Site / District:        {pk['district']} District, Cyprus",
            f"BESS Capacity:          {pk['mwh']} MWh / {pk['mw']} MW",
            f"Containers (BESS):      {pk['bess']} x Linyang Power Atlantic 5.015 MWh (20HC)",
            f"MV Skid:                {pk['mv_desc']}",
            f"PCS:                    {pk['pcs']}",
            f"Coastal Distance:       {pk['coastal']} km (C5 enclosure — full 5-year warranty)",
            '',
            f"  Indicative Component A (Equipment CIF):    EUR {fmt(pk['comp_a'])}",
            f"  Indicative Component B (EPC Services):     EUR {fmt(pk['comp_b'])}",
            f"  INDICATIVE CONTRACT PRICE (A+B, ex VAT):   EUR {fmt(pk['price'])}",
            '',
            f"  APG (70% × Indicative Component A):        EUR {fmt(pk['apg'])}",
            f"  Performance Bond (5% × Component A):       EUR {fmt(pk['pb'])}",
            '-' * 60,
        ]
    sched_a_lines += [
        '',
        '=' * 60,
        'PART 2 — COMBINED CONTRACT SUMMARY',
        '=' * 60,
        '',
        f"TOTAL INDICATIVE CONTRACT PRICE (ex VAT):  EUR {fmt(total_price)}",
        '',
        'PAYMENT MILESTONES (applied to Confirmed Contract Price):',
        f"  Advance (30%):     EUR {fmt(total_price * 0.30)}  — Within 30 days of payment trigger (§1A.3)",
        f"  Pre-Ship (55%):    EUR {fmt(total_price * 0.55)}  — On FAT + Linyang written confirmation",
        f"  PAC (10%):         EUR {fmt(total_price * 0.10)}  — System commissioned, grid-connected",
        f"  Retention (5%):    EUR {fmt(total_price * 0.05)}  — Released after 3-month DLP",
        f"  TOTAL:             EUR {fmt(total_price)}",
        '',
        '=' * 60,
        'PART 3 — KEY DATES',
        '=' * 60,
        '',
        'Production Start:       Q2 2026',
        'Factory Acceptance:     Q3 2026',
        'CIF Limassol:           August / September 2026',
        f'Target PAC Date:        {pac_date}',
        'Warranty Start:         On PAC',
        'Base Warranty End:      5 years from PAC (January 2032)',
        'DLP End:                3 months post-PAC (approx. April 2027)',
        'Retention Release:      On DLP expiry (approx. April 2027)',
        '',
        '=' * 60,
        'PART 4 — PRICE BASIS REFERENCE (per Section 6.1)',
        '=' * 60,
        '',
        'Quotation Source:       Linyang Quotation LY202601271 (January 2026)',
        f'Lithium Carbonate Index: {REF_LITHIUM_CNY}',
        f'EUR/CNY Reference:      {REF_EUR_CNY}',
        'Reference Date:         January 2026',
        '',
        'Adjustment Triggers (per §6.1(c)):',
        '  • Lithium carbonate index movement > 10% from reference value, OR',
        '  • EUR/CNY rate movement > 8% from reference value, OR',
        '  • Configuration change requested by Client',
        '',
        'Adjustment Cap: 5% of Indicative Contract Price (§6.1(d))',
        'Walk-away Right: Either Party may terminate without penalty if adjustment > 5%',
        '',
        '=' * 60,
        'PART 5 — BESS TECHNICAL SPECIFICATION',
        '=' * 60,
        '',
        'BESS: Linyang Power Atlantic 5.015 MWh LFP (EVE LF314, 12P416S)',
        '  Thermal: liquid cooling 45 kW; IP54; C5 enclosure (Cyprus standard)',
        '  Fire: FK5112 perfluoroketone + water; BMS: Modbus TCP / IEC 60870-5-104',
        '',
        'PCS: Kehua BCS1250K-C-HUD (1.25 MW); ≥98% efficiency; EN 50549-2 (TÜV D 115067 0077)',
        'Incoterms: CIF Limassol, Cyprus (Incoterms 2020); Port of Shipment: Shanghai',
        '',
        '=' * 60,
        'PART 6 — COMPANION DOCUMENTS (executed simultaneously)',
        '=' * 60,
        '',
        f'1. LTSA v4.0 (Ref: {ltsa_ref})',
        '2. OEM Direct Warranty Undertaking — Schedule C (LCY-OEM-DWU-001, one per park)',
        '   Linyang contacts: Conor Yang / Kamil Tyburski / Tomasz Wieckowski',
        '3. EMS Subscription Addendum — DISPERON (EUR 400/MWh/yr from PAC)',
        '4. APG from Linyang (CP: received before advance payment)',
        '5. Performance Bond from Linyang (5% Component A, valid to DLP end)',
        '6. Confirmed Price Certificate (issued within 14 days of Connection Terms receipt)',
    ]

    paras = doc.paragraphs
    sa_idx = sb_idx = None
    for i, p in enumerate(paras):
        if 'SCHEDULE A' in p.text and sa_idx is None:
            sa_idx = i
        if 'SCHEDULE B' in p.text and sa_idx is not None and sb_idx is None:
            sb_idx = i
            break

    # Clear existing Schedule A paragraphs
    for i in range(sa_idx, sb_idx):
        set_p_text(paras[i], '')

    # Set heading + insert all lines before SCHEDULE B
    set_p_text(paras[sa_idx], sched_a_lines[0])
    anchor = paras[sb_idx]._p
    for line in sched_a_lines[1:]:
        anchor.addprevious(make_p(line))

    # Pre-fill signature lines (if provided)
    if sig_name:
        for i, p in enumerate(paras):
            if f'For and on behalf of {client_name}' in p.text:
                # Look for Name: blank in next paragraphs
                for j in range(i + 1, min(i + 10, len(paras))):
                    if paras[j].text.strip() == 'Name: __________________________':
                        set_p_text(paras[j], f'Name:   {sig_name}')
                    elif paras[j].text.strip() == 'Title: __________________________':
                        set_p_text(paras[j], f'Title:   {sig_title}')
                break

    doc.save(out_path)
    print(f'  Saved: {os.path.basename(out_path)}')


# ──────────────────────────────────────────────────────────────────────────
# STEP 3: Update Esperia LOI for v5.1 alignment
# ──────────────────────────────────────────────────────────────────────────

def update_loi(loi_path):
    """Apply v5.1 changes to the Esperia pipeline LOI."""
    doc = Document(loi_path)

    repls = [
        # 1. Indicative Price language — header
        ('Letter of Intent - REVISED (Revision 1, May 2026)',
         'Letter of Intent - REVISED (Revision 2, May 2026)'),
        # 2. Add note re: Indicative Price + Confirmed Price mechanism
        ('Esperia Energy Group - BESS Portfolio EPC Pipeline Commitment',
         'Esperia Energy Group - BESS Portfolio EPC Pipeline Commitment\n'
         'KEY UPDATES IN REVISION 2:\n'
         '- All pipeline pricing is INDICATIVE — Confirmed Contract Price set per EPC §6.1 mechanism on receipt of Connection Terms\n'
         '- 12-month Connection Terms longstop per phase (per EPC §1A.1)\n'
         '- Advance payment trigger: 30 days from later of Effective Date and Connection Terms\n'
         '- Indicative Price subject to ±5% adjustment for raw material (lithium >10%) or FX (EUR/CNY >8%)'),
        # 3. Replace fixed pipeline pricing language (§5.1)
        ('All Pipeline Projects shall be priced at Group-Order Pricing rates based on Quotation LY202511281 or a successor quotation to be agreed in writing.',
         'All Pipeline Projects shall be priced at Group-Order Pricing rates based on Linyang Quotation LY202601271 (January 2026) or a successor quotation. Each EPC shall apply the Indicative Price + Confirmed Price mechanism set out in EPC v5.1 §6.1 — a Confirmed Price Certificate is issued within 14 days of Connection Terms receipt for that phase.'),
        # 4. Update payment timing (5.2): 7 days → 30 days
        ('30%  —  Advance payment upon EPC Agreement signature;',
         '30%  —  Advance payment within 30 days of payment trigger (per EPC §1A.3 — later of Effective Date and Connection Terms);'),
        # 5. Update LTSA rate reference
        ('Tier C or as separately agreed', 'Tier C — EUR 1,740/MWh/year'),
    ]

    for old, new in repls:
        replace_text_in_doc(doc, old, new)

    # Add 12-month longstop note to phase descriptions (§4.1)
    # Find each "Phase N" description paragraph and add longstop reference
    for p in doc.paragraphs:
        if 'EPC Agreements for the three Phase 1 parks' in p.text and 'longstop' not in p.text.lower():
            new_text = p.text.rstrip()
            if new_text.endswith('.'):
                new_text = new_text[:-1]
            new_text += '. Each EPC is conditional on Client obtaining grid Connection Terms within 12 months of Effective Date.'
            set_p_text(p, new_text)

    doc.save(loi_path)
    print(f'  Updated LOI: {os.path.basename(loi_path)}')


# ──────────────────────────────────────────────────────────────────────────
# STEP 4: Copy Technical Agreements to contracts folder
# ──────────────────────────────────────────────────────────────────────────

def copy_tech_agreements():
    src_dir = f'{BASE}/docs/clients/group-order/Group2_Esperia_Energy/technical'
    dst_dir = GAL_OUT
    files = [
        'Cyprus 2.5MW_10MWh Technical Agreement V2.1 (draft) 260429.docx',
        'Cypurs 5MW_20MWh Technical Agreement V2.1 (draft) 260429(1).docx',
    ]
    for fname in files:
        src = os.path.join(src_dir, fname)
        if os.path.exists(src):
            # Clean filename for output
            new_name = fname.replace('Cypurs', 'Cyprus').replace(' (draft) 260429(1)', ' (draft) 260429')
            dst = os.path.join(dst_dir, new_name)
            shutil.copy2(src, dst)
            print(f'  Copied: {new_name}')
        else:
            print(f'  WARN: source not found: {src}')


# ══════════════════════════════════════════════════════════════════════════
# RUN
# ══════════════════════════════════════════════════════════════════════════

print('=' * 70)
print('STEP 1: Build v5.1 master template')
print('=' * 70)
build_v51_template()

# Park data
gal_parks = [
    {'name': 'Galascope 1', 'district': 'Famagusta', 'mw': 5.0, 'mwh': 20, 'bess': 4, 'mv': 1,
     'mv_desc': '1 x T4 MV Skid (4 x BCS1250K = 5 MW, SL-5000)',
     'pcs': '4 x Kehua BCS1250K-C-HUD = 5.0 MW', 'coastal': '>5',
     'comp_a': 1848712.43, 'comp_b': 410187.57, 'price': 2258900,
     'apg': 1848712.43 * 0.70, 'pb': 1848712.43 * 0.05},
    {'name': 'Galascope 2', 'district': 'Famagusta', 'mw': 2.5, 'mwh': 10, 'bess': 2, 'mv': 1,
     'mv_desc': '1 x T2 MV Skid (2 x BCS1250K = 2.5 MW, SL-2500)',
     'pcs': '2 x Kehua BCS1250K-C-HUD = 2.5 MW', 'coastal': '>5',
     'comp_a': 974457.00, 'comp_b': 231843.00, 'price': 1206300,
     'apg': 974457.00 * 0.70, 'pb': 974457.00 * 0.05},
]
spa_parks = [
    {'name': 'Anarita 1', 'district': 'Paphos', 'mw': 5.0, 'mwh': 20, 'bess': 4, 'mv': 1,
     'mv_desc': '1 x T4 MV Skid (4 x BCS1250K = 5 MW, SL-5000)',
     'pcs': '4 x Kehua BCS1250K-C-HUD = 5.0 MW', 'coastal': '>5',
     'comp_a': 1848712.43, 'comp_b': 531287.57, 'price': 2380000,
     'apg': 1848712.43 * 0.70, 'pb': 1848712.43 * 0.05},
    {'name': 'Anarita 2', 'district': 'Paphos', 'mw': 5.0, 'mwh': 20, 'bess': 4, 'mv': 1,
     'mv_desc': '1 x T4 MV Skid (4 x BCS1250K = 5 MW, SL-5000)',
     'pcs': '4 x Kehua BCS1250K-C-HUD = 5.0 MW', 'coastal': '>5',
     'comp_a': 1848712.43, 'comp_b': 531287.57, 'price': 2380000,
     'apg': 1848712.43 * 0.70, 'pb': 1848712.43 * 0.05},
]

print()
print('=' * 70)
print('STEP 2: Generate Galascope EPC v5.1')
print('=' * 70)
gal_reg_line = (
    'a company incorporated under the laws of the Republic of Cyprus, '
    'with registered office at Karaiskaki 6, City House, 3032 Limassol, Cyprus, '
    'Registration No. HE 303759'
)
render_epc(
    client_name='Galascope Ltd',
    client_reg_line=gal_reg_line,
    client_address_block='Karaiskaki 6, City House, 3032 Limassol, Cyprus',
    parks=gal_parks,
    total_price=3465200,
    pac_date='31 January 2027',
    ref='LCY-EPC-GAL-B1-2026',
    ltsa_ref='LCY-LTSA-GAL-2026',
    sig_name='Ntinos Konstantinos',
    sig_title='Director',
    out_path=f'{GAL_OUT}/EPC-Galascope-Esperia-batch1-may2026.docx'
)

print()
print('=' * 70)
print('STEP 3: Generate Spanercom EPC v5.1')
print('=' * 70)
spa_reg_line = (
    'a company incorporated under the laws of the Republic of Cyprus, '
    'with registered office at Stasikratous 32, Charalambous Tower, Floor 6, '
    'Flat/Office 603, 1065 Nicosia, Cyprus, Registration No. HE 396638'
)
render_epc(
    client_name='Spanercom Ltd',
    client_reg_line=spa_reg_line,
    client_address_block='Stasikratous 32, Charalambous Tower, Floor 6, Flat/Office 603, 1065 Nicosia, Cyprus',
    parks=spa_parks,
    total_price=4760000,
    pac_date='31 January 2027',
    ref='LCY-EPC-SPA-2026',
    ltsa_ref='LCY-LTSA-SPA-2026',
    sig_name=None,  # blank — different client
    sig_title=None,
    out_path=f'{SPA_OUT}/EPC-Spanercom-Anarita-may2026.docx'
)

print()
print('=' * 70)
print('STEP 4: Update Esperia LOI')
print('=' * 70)
update_loi(f'{GAL_OUT}/LOI-Esperia-Energy-pipeline-may2026.docx')

print()
print('=' * 70)
print('STEP 5: Copy Technical Agreement drafts')
print('=' * 70)
copy_tech_agreements()

print()
print('=' * 70)
print('DONE')
print('=' * 70)
