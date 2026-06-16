"""Generate the branded Word RFP that accompanies the prefilled producer Excel forms.

Covers BOTH transformers (T1 main + T2 earthing/aux), full specification, Q1 2027 delivery,
and client requirements. Producer-agnostic master — issued with each supplier's own prefilled
questionnaire (T1) and the Lighthief T2 datasheet.

Output: .../HV Transformer/RFP/HESS-transformer-RFP-jun2026.docx
"""
from __future__ import annotations

import os
import sys
from datetime import date

sys.path.insert(0, os.path.dirname(__file__))
from loi_docx_common import (  # noqa: E402
    AMBER_HEX,
    AMBER_HDR,
    AMBER_TXT,
    GOLD,
    GREY,
    NAVY,
    NAVY_HEX,
    WHITE,
    add_footer,
    add_run,
    body,
    h1,
    lock_table_widths,
    new_document,
    set_cell_bg,
    tbl_hdr,
)
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

REF = "Ref: LCY-RFP-HESS-TRF-2026"
DOC_DATE = date.today().strftime("%d %B %Y")
DEADLINE = "Tuesday 17 June 2026, 18:00 (Cyprus / EEST)"

OUT_DIR = (
    r"L:\My Drive\LINYANG\BESS CLIENTS\Individual_60-120-standalone\HV Transformer\RFP"
)
OUT_NAME = "HESS-transformer-RFP-jun2026.docx"

T1_SPECS = [
    ("Quantity", "1 unit"),
    ("Rating", "63 MVA continuous (ONAN / ONAF)"),
    ("Voltage ratio", "132 / 33 kV"),
    ("Vector group", "YNd11 (132 kV star / 33 kV delta)"),
    ("Short-circuit impedance uk", "21% on HV base @ 75 °C, CMR"),
    ("Highest system voltage", "145 / 36 kV"),
    ("Insulation level — HV", "550 kV BIL / 230 kV AC (1 min)"),
    ("Insulation level — LV (36 kV class)", "170 kV BIL / 70 kV AC (1 min)  [corrects any 125/50 in legacy tables]"),
    ("Neutral (HVN)", "Power-frequency withstand 38 kV"),
    ("Through-fault withstand", "132 kV: 31.5 kA / 1 s   |   33 kV: 20 kA / 3 s"),
    ("Cooling", "ONAN / ONAF — 2 cooler banks @ 50% CMR each"),
    ("OLTC", "+12.5% / −18.75%, 25 × 1.25% steps; MR/ABB vacuum; local/remote/supervisory"),
    ("Load power factor", "0.9 lag (design)"),
    ("Temperature rise", "55 °C top oil / 60 °C winding"),
    ("Sound level", "≤ 60 dB(A) outdoor (≤ 80 worst case)"),
    ("Pollution / creepage", "Heavy — IEC 60815 Class III; composite ≥ 35 mm/kV"),
    ("Standards", "Tier 2 Ecodesign (EU 2019/1783); EN 60076 series; TSOC T14"),
    ("Monitoring", "Dual WTI (HV + LV hotspot); online DGA / transformer monitoring system"),
    ("Oil / bushings / OLTC", "EN 60296 oil; EN 60137 bushings; EN 60214 OLTC; silica-gel breather"),
    ("Mounting / transport", "Flat base; anti-vibration pads; EN 795 anchorage; impact recorders w/ GPS"),
    ("Losses / PEI", "To be stated by producer (no-load, load kW, PEI %)"),
    ("Installation", "Outdoor, oil-immersed; Plot 26 Psevdas, Larnaca, Cyprus"),
]

T2_SPECS = [
    ("Equipment", "Separate earthing & auxiliary transformer (not a winding on T1)"),
    ("Function", "Derive & solidly earth the 33 kV system neutral + station-service supply"),
    ("Earthing connection", "Zig-zag (interconnected-star), ZN — SOLID earthing, no NER"),
    ("Earth-fault withstand", "≥ 20 kA for 3 s (or ≥ 25 kA for 1 s) per EN 60076-5"),
    ("MV voltage", "33 kV (highest equipment 36 kV); insulation 170 BIL / 70 AC"),
    ("Auxiliary winding", "≥ 315 kVA @ 400/230 V, 3-ph 4-wire (minimum — upratable; quote 315/400/500 kVA)"),
    ("Cooling / oil", "ONAN; mineral oil EN 60296"),
    ("Standards", "EN/IEC 60076-1/-2/-3/-5; IEC 60076-6; EN 60289"),
    ("Note", "Full requirements in the attached T2 datasheet (Excel)"),
]


def add_header_bar(doc) -> None:
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_tbl.autofit = False
    lc = hdr_tbl.rows[0].cells[0]
    set_cell_bg(lc, NAVY_HEX)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(6)
    lp.paragraph_format.space_after = Pt(6)
    add_run(lp, "Lighthief", bold=True, size_pt=22, color=WHITE)
    add_run(lp, " Cyprus Ltd", bold=False, size_pt=11, color=WHITE)
    rc = hdr_tbl.rows[0].cells[1]
    set_cell_bg(rc, NAVY_HEX)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(6)
    rp.paragraph_format.space_after = Pt(6)
    for line, is_gold in [
        ("REQUEST FOR PROPOSAL", True),
        (REF, False),
        (f"Date: {DOC_DATE}", False),
        (f"Bid deadline: {DEADLINE}", False),
        ("STRICTLY CONFIDENTIAL", True),
    ]:
        add_run(rp, line + "\n", bold=is_gold, size_pt=8, color=GOLD if is_gold else WHITE)
    lock_table_widths(hdr_tbl, [9.0, 5.0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def spec_table(doc, rows) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_hdr(tbl, ["Parameter", "Requirement"])
    for i, (param, val) in enumerate(rows):
        row = tbl.rows[0] if i == 0 else tbl.add_row()
        row.cells[0].paragraphs[0].clear()
        add_run(row.cells[0].paragraphs[0], param, bold=True, size_pt=9)
        row.cells[1].paragraphs[0].clear()
        add_run(row.cells[1].paragraphs[0], val, size_pt=9)
    lock_table_widths(tbl, [5.2, 8.8])


def note_box(doc, label, text) -> None:
    nb = doc.add_table(rows=1, cols=1)
    cell = nb.rows[0].cells[0]
    set_cell_bg(cell, AMBER_HEX)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, label, bold=True, size_pt=9, color=AMBER_HDR)
    add_run(p, text, size_pt=9, color=AMBER_TXT)
    lock_table_widths(nb, [14.0])


def build():
    doc = new_document()
    add_header_bar(doc)

    tp = doc.add_paragraph()
    add_run(tp, "Request for Proposal — Power, Earthing & Auxiliary Transformers", bold=True, size_pt=16, color=NAVY)
    sp = doc.add_paragraph()
    add_run(sp, "HESS Psevdas BESS — 63 MVA 132/33 kV step-up + 33 kV earthing/auxiliary transformer", italic=True, size_pt=11, color=GREY)

    body(
        doc,
        "Lighthief Cyprus Ltd (EPC contractor), on behalf of H.E.S.S. Hybrid Energy Storage Systems Ltd "
        "(the Client, project SPV), invites your firm offer for the supply of the transformers described "
        "below for the Client's standalone battery energy storage facility at Plot 26, Psevdas, Larnaca, "
        "Cyprus (TSOC connection ref. ΔΣΜΚ/ΠΟΣ/320.7.11).",
        space_after=6,
    )
    note_box(
        doc,
        "HOW TO RESPOND:  ",
        "This RFP accompanies your transformer questionnaire, which Lighthief has pre-filled from the "
        "Client's requirements (the 'Required' column). Please complete the 'OFFER' column of that Excel "
        "for the main transformer (T1) and complete the attached earthing & auxiliary transformer datasheet "
        "(T2), and return both with a signed compliance statement, certificates, and your commercial offer.",
    )

    h1(doc, "1.   SCOPE — TWO TRANSFORMERS")
    body(
        doc,
        "Quote T1 and T2 as SEPARATE line items. A bid covering only the main transformer is non-compliant. "
        "The earthing & auxiliary transformer (T2) is a separate unit — the TSOC connection terms list it as a "
        "distinct item — not a winding on the main transformer.",
        space_after=6,
    )

    h2 = doc.add_paragraph()
    add_run(h2, "T1 — Main power transformer (63 MVA, 132/33 kV)", bold=True, size_pt=11, color=NAVY)
    spec_table(doc, T1_SPECS)

    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(10)
    add_run(h3, "T2 — Earthing & auxiliary transformer (33 kV / 400-230 V)", bold=True, size_pt=11, color=NAVY)
    spec_table(doc, T2_SPECS)

    h1(doc, "2.   PROGRAMME — Q1 2027 DELIVERY")
    body(doc, "2.1  Target manufacturing slot: January 2027.")
    body(
        doc,
        "2.2  Target on-site delivery: Q1 2027 (DAP Limassol Port / Plot 26 Psevdas), aligned with the "
        "Client's implementation programme and TSOC connection schedule. State your manufacturing lead time "
        "and total delivery time to Limassol Port against this target.",
    )

    h1(doc, "3.   COMPLIANCE & OPEN ITEMS")
    body(
        doc,
        "3.1  Return a signed clause-by-clause compliance statement (COMPLY / DEVIATE + note). Explicitly "
        "confirm: YNd11 (not Dyn11), uk 21%, 170/70 LV insulation, Tier 2 Ecodesign PEI, and that T2 is included. "
        "State no-load loss, load loss and PEI % for both units.",
    )
    body(doc, "3.2  Bid the following open items as instructed (do not assume):")
    for t in [
        "T2 auxiliary winding: bid 315 kVA minimum, confirm upratable, give price per kVA step (315/400/500).",
        "T1 132 kV CTs: quote 6-CT base PLUS a priced option for 9-CT (neutral CT for REF) — final fixed when "
        "the ISM bay protection single-line is issued.",
        "Seismic: design to CYS EN 1998-1 Zone II, agR 0.23 g, importance factor 1.4.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, t, size_pt=10)

    h1(doc, "4.   DELIVERY & INCOTERMS")
    for t in [
        "State ex-works manufacturing lead time (weeks from order/LOI).",
        "State total delivery time to Limassol Port, Cyprus (incl. transit + customs clearance).",
        "Price DAP Limassol Port (Incoterms 2020) as baseline; also quote DAP/DDP Plot 26 Psevdas (inland "
        "haulage + offloading).",
        "Transport provisions: impact recorders w/ GPS, lifting/jacking points, shipping oil/nitrogen state, "
        "abnormal-load responsibility, any port-handling exclusions.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, t, size_pt=10)

    h1(doc, "5.   CERTIFICATIONS (attach with bid)")
    for t in [
        "EU Declaration of Conformity + CE marking.",
        "EU Ecodesign Tier 2 (Reg. 548/2014 as amended by 2019/1783) — PEI declaration.",
        "Independent short-circuit withstand TYPE TEST to EN 60076-5 (KEMA / CESI / IPH) — calculation-only "
        "not accepted.",
        "Temperature-rise (EN 60076-2); impulse & AC withstand (EN 60076-3) at the 170/70 LV level; sound "
        "level (EN 60076-10).",
        "Routine test certificates per EN 60076-1; material certs (oil EN 60296, bushings EN 60137, OLTC EN 60214).",
        "ISO 9001 / 14001 / 45001; EU references for 132 kV / ≥63 MVA units; warranty & LTSA terms.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, t, size_pt=10)

    h1(doc, "6.   COMMERCIAL & SUBMISSION")
    for t in [
        "Prices in EUR, itemised: T1, T2, transport/Incoterm options, spares, supervision of erection & commissioning.",
        "Offer validity: minimum 90 days.",
        "Warranty: 24 months from commissioning or 36 months from delivery, whichever comes first "
        "(60 months from delivery preferred). Must cover BOTH T1 and T2. State terms explicitly.",
        "Payment terms: state proposed (final per executed EPC).",
        f"Submit to office@lighthief.com by the bid deadline: {DEADLINE}.",
        "Format: completed T1 questionnaire (OFFER column) + completed T2 datasheet + compliance statement + "
        "certificates + commercial offer.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, t, size_pt=10)
    note_box(
        doc,
        "WHY THIS DEADLINE:  ",
        "Bids must arrive by 17 June 2026 so Lighthief can evaluate all offers and present the consolidated "
        "solution to the Client at Intersolar Europe (23–25 June 2026). Late or incomplete bids may be excluded.",
    )

    h1(doc, "7.   EVALUATION CRITERIA")
    for t in [
        "Technical compliance (T1 and T2), incl. genuine Tier 2 Ecodesign & independent short-circuit type test.",
        "Total cost of ownership (price + capitalised losses / PEI).",
        "Delivery time to Limassol vs January 2027 slot / Q1 2027 target.",
        "Manufacturer track record & EU references; warranty / LTSA.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, t, size_pt=10)

    body(
        doc,
        "Contact for clarifications: Alexander Papacosta, Cyprus Director — office@lighthief.com — "
        "+357 99 164 158 — solarfarms.cy.",
        space_after=4,
    )

    add_footer(doc, REF)
    return doc


def save_doc(doc) -> str:
    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.join(OUT_DIR, OUT_NAME)
    try:
        doc.save(path)
    except PermissionError:
        path = os.path.join(OUT_DIR, "HESS-transformer-RFP-jun2026-rev2.docx")
        doc.save(path)
    return os.path.abspath(path)


if __name__ == "__main__":
    print("Saved -> " + save_doc(build()))
