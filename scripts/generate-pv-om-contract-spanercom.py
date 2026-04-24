"""
Generate PV O&M Contract — Spanercom Ltd (Anarita East & West Parks)
Output: docs/clients/Individual_Spanercom/contracts/
        PV-OM-Contract-Spanercom-Anarita-2x5MW-apr2026.docx

Parks:   2 × 5.01 MWp (Anarita East + West), Paphos District, Cyprus
Price:   €56,700 / MW / year  →  €283,500 / park / year  →  €567,000 / year total
Cleaning: 2 × per year (spring + autumn)
Guarantee: 99% Availability with performance penalty schedule
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import date

# ── Brand colours ─────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1A, 0x36, 0x5D)
GOLD      = RGBColor(0xC9, 0xA4, 0x32)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)
GREY      = RGBColor(0x40, 0x40, 0x40)
NAVY_HEX  = "1A365D"
GOLD_HEX  = "C9A432"
LIGHT_HEX = "EBF0F7"
ALT_HEX   = "F5F7FA"

# ── Contract data ─────────────────────────────────────────────────────────────
CONTRACT = {
    "ref":               "LH-OM-SPANERCOM-2026-001",
    "date":              "23 April 2026",
    "effective_date":    "23 April 2026",
    "commencement_date": "1 May 2026",
    "term_years":    2,
    "client": {
        "name":    "Spanercom Ltd",
        "contact": "Muminjon",
        "address": "Anarita, Paphos District, Republic of Cyprus",
    },
    "provider": {
        "name":    "Lighthief Cyprus Ltd",
        "contact": "Alexander Papacosta",
        "title":   "Cyprus Director",
        "address": "Republic of Cyprus",
        "reg":     "HE [·]",
    },
    "parks": [
        {"name": "Anarita East Park",  "dc_mwp": 5.01, "ac_mw": 5.0, "location": "Anarita, Paphos District"},
        {"name": "Anarita West Park",  "dc_mwp": 5.01, "ac_mw": 5.0, "location": "Anarita, Paphos District"},
    ],
    "price_per_park_yr": 28_200,
    "parks_count":     2,
    "cleanings_per_yr":2,
    "availability_guarantee_pct": 99.0,
    "vat_rate":        0.19,
    "callout_cap_per_park": 6,
    "callout_extra_fee":    400,
}

PRICE_PER_PARK        = CONTRACT["price_per_park_yr"]
PRICE_TOTAL           = PRICE_PER_PARK * CONTRACT["parks_count"]
VAT_RATE              = CONTRACT["vat_rate"]
VAT_PER_PARK          = PRICE_PER_PARK * VAT_RATE
VAT_TOTAL             = PRICE_TOTAL * VAT_RATE
PRICE_PER_PARK_INCL   = PRICE_PER_PARK + VAT_PER_PARK
PRICE_TOTAL_INCL      = PRICE_TOTAL + VAT_TOTAL
PRICE_MONTHLY_EX      = PRICE_TOTAL / 12
PRICE_MONTHLY_VAT     = PRICE_MONTHLY_EX * VAT_RATE
PRICE_MONTHLY_INCL    = PRICE_MONTHLY_EX + PRICE_MONTHLY_VAT

# ── Output path ───────────────────────────────────────────────────────────────
OUT_DIR  = os.path.join(os.path.dirname(__file__), "..", "docs", "clients",
                        "Individual_Spanercom", "contracts")
OUT_FILE = os.path.join(OUT_DIR, "PV-OM-Contract-Spanercom-Anarita-2x5MW-apr2026.docx")
os.makedirs(OUT_DIR, exist_ok=True)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{ side }")
            el.set(qn("w:val"),   val.get("val",   "single"))
            el.set(qn("w:sz"),    val.get("sz",    "4"))
            el.set(qn("w:space"), val.get("space", "0"))
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def fmt(p, text, bold=False, italic=False, size=11,
        color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    run = p.add_run(text)
    run.bold  = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    p.alignment = align
    return run

def add_paragraph(doc, text="", bold=False, italic=False, size=11,
                  color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        fmt(p, text, bold=bold, italic=italic, size=size, color=color, align=align)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    size  = 14 if level == 1 else 12
    color = GOLD
    fmt(p, text, bold=True, size=size, color=color)
    # underline thin rule via bottom border on the paragraph
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), GOLD_HEX if level == 1 else NAVY_HEX)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def add_bullet(doc, text, indent_level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + indent_level * 0.6)
    p.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size  = Pt(10.5)
        r.font.color.rgb = BLACK
        r2 = p.add_run(text)
        r2.font.size  = Pt(10.5)
        r2.font.color.rgb = BLACK
    else:
        r = p.add_run(text)
        r.font.size  = Pt(10.5)
        r.font.color.rgb = BLACK
    return p

def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size  = Pt(10.5)
        r.font.color.rgb = BLACK
        r2 = p.add_run(text)
        r2.font.size  = Pt(10.5)
        r2.font.color.rgb = BLACK
    else:
        r = p.add_run(text)
        r.font.size  = Pt(10.5)
        r.font.color.rgb = BLACK
    return p

def add_table_row(table, cells_data, header=False, alt=False):
    """cells_data: list of (text, width_cm, align, bold)"""
    row = table.add_row()
    for i, (text, width, align_str, bold) in enumerate(cells_data):
        cell = row.cells[i]
        cell.width = Cm(width)
        if header:
            set_cell_bg(cell, NAVY_HEX)
        elif alt:
            set_cell_bg(cell, ALT_HEX)
        p   = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold or header
        run.font.size  = Pt(9.5)
        run.font.color.rgb = WHITE if header else BLACK
        if align_str == "C":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align_str == "R":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

def set_col_widths(table, widths_cm):
    for i, w in enumerate(widths_cm):
        for row in table.rows:
            row.cells[i].width = Cm(w)

# ── Page setup ────────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.page_width   = Cm(21)
section.page_height  = Cm(29.7)
section.left_margin  = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin   = Cm(2.0)
section.bottom_margin= Cm(2.0)

# ── Default style ─────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name  = "Calibri"
style.font.size  = Pt(11)
style.font.color.rgb = BLACK

# ══════════════════════════════════════════════════════════════════════════════
# COVER / HEADER BLOCK
# ══════════════════════════════════════════════════════════════════════════════
# Navy header bar via a single-cell table
hdr_table = doc.add_table(rows=1, cols=1)
hdr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cell = hdr_table.rows[0].cells[0]
hdr_cell.width = Cm(16)
set_cell_bg(hdr_cell, NAVY_HEX)

p = hdr_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("LIGHTHIEF CYPRUS LTD")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = WHITE

p2 = hdr_cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(12)
r2 = p2.add_run("Photovoltaic Operation & Maintenance Services Agreement")
r2.bold = True; r2.font.size = Pt(12); r2.font.color.rgb = GOLD

doc.add_paragraph()

# Reference block
ref_table = doc.add_table(rows=4, cols=2)
ref_table.alignment = WD_TABLE_ALIGNMENT.LEFT
ref_data = [
    ("Contract Reference:",      CONTRACT["ref"]),
    ("Date of Signing:",         CONTRACT["date"]),
    ("Effective Date:",          CONTRACT["effective_date"]),
    ("Service Commencement:",    CONTRACT["commencement_date"]),
]
for i, (label, value) in enumerate(ref_data):
    row = ref_table.rows[i]
    c0, c1 = row.cells[0], row.cells[1]
    c0.width = Cm(5); c1.width = Cm(11)
    lp = c0.paragraphs[0]
    lr = lp.add_run(label)
    lr.bold = True; lr.font.size = Pt(10); lr.font.color.rgb = GREY
    vp = c1.paragraphs[0]
    vr = vp.add_run(value)
    vr.font.size = Pt(10); vr.font.color.rgb = BLACK

doc.add_paragraph()

# ── Parties block ─────────────────────────────────────────────────────────────
parties_table = doc.add_table(rows=1, cols=2)
parties_table.alignment = WD_TABLE_ALIGNMENT.LEFT

for i, (role, party_key) in enumerate([("SERVICE PROVIDER", "provider"), ("CLIENT", "client")]):
    cell = parties_table.rows[0].cells[i]
    cell.width = Cm(8)
    set_cell_bg(cell, LIGHT_HEX)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(role)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = NAVY

    party = CONTRACT[party_key]
    for line in [party["name"], party.get("contact",""), party["address"]]:
        if line:
            lp = cell.add_paragraph()
            lp.paragraph_format.space_after = Pt(1)
            lr = lp.add_run(line)
            lr.font.size = Pt(9.5); lr.font.color.rgb = BLACK
    ep = cell.add_paragraph()
    ep.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 1 — DEFINITIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 1 — Definitions", level=1)
add_paragraph(doc,
    "In this Agreement, the following terms shall have the meanings set out below:",
    size=10.5, space_after=6)

definitions = [
    ('"Agreement"',
     'this Photovoltaic Operation & Maintenance Services Agreement, including all Schedules and Annexes.'),
    ('"Effective Date"',
     f'the date of signing of this Agreement, being {CONTRACT["effective_date"]}. The Effective Date governs contractual obligations, confidentiality, and dispute resolution from the moment of execution.'),
    ('"Service Commencement Date"',
     f'the date on which O&M services, KPI measurement, and billing commence, being {CONTRACT["commencement_date"]}. All performance guarantees, Availability calculations, and Service Fees are measured from the Service Commencement Date.'),
    ('"Contract Year"',
     'each consecutive twelve (12) month period commencing on the Service Commencement Date.'),
    ('"Availability"',
     'the ratio (expressed as a percentage) of time during a measurement period in which a Park\'s inverter systems are operational and capable of producing power when in-plane irradiance exceeds 50 W/m\u00b2, as measured by the on-site meteorological station, excluding Excused Downtime. Calculated per IEC 61724-1 using SCADA data at 5-minute resolution.'),
    ('"Availability Guarantee"',
     f'{CONTRACT["availability_guarantee_pct"]}% measured on an annual rolling basis per Park, from the Service Commencement Date.'),
    ('"Curtailment"',
     'any reduction or interruption of energy export imposed by the Distribution System Operator (DSO / EAC), the Transmission System Operator (TSO), or any grid operator, including but not limited to: ripple control signals, export limitation instructions, net billing cap restrictions, or any forced reduction not attributable to the Service Provider or to the condition of the Park equipment.'),
    ('"Excused Downtime"',
     'downtime periods excluded from Availability and PR calculations, comprising: (i) Force Majeure events; (ii) Client-instructed shutdowns; (iii) Curtailment by DSO/TSO or any grid operator; (iv) grid outages beyond the Park substation; (v) any defect or failure in equipment not maintained under this Agreement; (vi) planned maintenance windows pre-agreed in writing; (vii) irradiance periods below 50 W/m\u00b2; (viii) sensor malfunction or SCADA data loss not attributable to the Service Provider.'),
    ('"Fault Response Time"',
     'the elapsed time from the Service Provider\'s receipt of a fault notification (via SCADA alarm, email, or telephone) to commencement of active remediation. Response times apply during Service Hours for minor and cosmetic faults, and on a 24/7 basis for critical and major faults.'),
    ('"Service Hours"',
     'Monday to Friday, 08:00–18:00 Cyprus local time (EET/EEST), excluding Cyprus public holidays. Emergency response for Critical and Major faults is provided 24/7 regardless of Service Hours.'),
    ('"Parks"',
     'collectively, the Anarita East Park (5.01 MWp / 5 MW AC) and the Anarita West Park (5.01 MWp / 5 MW AC), both located at Anarita, Paphos District, Republic of Cyprus.'),
    ('"Performance Ratio"',
     'the ratio of the measured specific energy yield (kWh/kWp) to the theoretically achievable yield based on in-plane irradiation, calculated per IEC 61724-1, excluding Excused Downtime periods from both numerator and denominator.'),
    ('"Preventive Maintenance"',
     'all scheduled maintenance activities carried out in accordance with the Maintenance Plan set out in Schedule A.'),
    ('"Corrective Maintenance"',
     'all unscheduled maintenance activities required to restore the Parks to their normal operating condition following a fault or failure.'),
    ('"SCADA"',
     'the supervisory control and data acquisition system monitoring the Parks. SCADA data constitutes the primary source of record for Availability and Performance Ratio measurement under this Agreement.'),
    ('"Service Fee"',
     'the annual remuneration payable to the Service Provider as set out in Article 9.'),
]
for term, defn in definitions:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(term + "  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
    r2 = p.add_run(defn)
    r2.font.size = Pt(10); r2.font.color.rgb = BLACK

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 2 — SCOPE OF SERVICES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 2 — Scope of Services", level=1)

add_heading(doc, "2.0  Service Package Overview & Selected Configuration", level=2)
add_paragraph(doc,
    "Lighthief Cyprus Ltd offers O&M services across four tiers. The table below "
    "shows the full service menu and the configuration selected under this Agreement "
    "(CUSTOM). Services marked with a dash are not included in the selected package "
    "and are available as billable add-ons under Article 10.",
    size=10.5, space_after=6)

# Package comparison table
pkg_cols = 6
pkg_table = doc.add_table(rows=1, cols=pkg_cols)
pkg_table.style = "Table Grid"
pkg_table.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_row(pkg_table, [
    ("Service",                     5.8, "L", True),
    ("ECO",                         1.5, "C", True),
    ("SILVER",                      1.5, "C", True),
    ("GOLD",                        1.5, "C", True),
    ("CUSTOM\n(This Agreement)",    2.5, "C", True),
    ("Notes",                       3.2, "L", True),
], header=True)

PKG_YES  = "Yes"
PKG_NO   = "-"
PKG_OPT  = "Add-on"

pkg_rows = [
    ("SCADA 24/7 online monitoring",
     PKG_YES, PKG_YES, PKG_YES, PKG_YES,
     "Alarms, faults, production anomalies"),
    ("Monthly performance reports",
     PKG_YES, PKG_YES, PKG_YES, PKG_YES,
     "12 reports/year per park"),
    ("Periodic inspection",
     "2x/yr", "4x/yr", "6x/yr", "2x/yr",
     "Visual + electrical check per visit"),
    ("IR thermography",
     PKG_NO, PKG_YES, PKG_YES, PKG_YES,
     "2x/year at each inspection visit"),
    ("Panel cleaning",
     PKG_NO, "1x/yr", "2x/yr", "2x/yr",
     "IEC TS 62788 compliant, per visit"),
    ("Vegetation / grass cutting",
     PKG_NO, "2x/yr", "4x/yr", "4x/yr",
     "Spring, early summer, late summer, autumn"),
    ("Corrective call-outs (included)",
     "0", "2/park", "4/park", "6/park",
     "On-site; remote resolutions unlimited"),
    ("Electrical safety inspections",
     PKG_NO, PKG_YES, PKG_YES, PKG_YES,
     "Annual IR test, earth continuity, IEC 62446-1"),
    ("Spare parts inventory management",
     PKG_NO, PKG_YES, PKG_YES, PKG_YES,
     "Consumables: fuses, SPDs, filters"),
    ("HSE compliance & site register",
     PKG_NO, PKG_YES, PKG_YES, PKG_YES,
     "Site HSE plan, toolbox talks, waste disposal"),
    ("Annual summary report",
     PKG_NO, PKG_YES, PKG_YES, PKG_YES,
     "Year-on-year comparison + next-year plan"),
    ("Transformer general audit",
     PKG_NO, PKG_NO, PKG_YES, PKG_NO,
     "Available as add-on (Article 10)"),
    ("Insurance management",
     PKG_NO, PKG_YES, PKG_YES, PKG_NO,
     "Not required — Client manages own cover"),
    ("PPA / energy contract management",
     PKG_NO, PKG_NO, PKG_YES, PKG_NO,
     "Not applicable — Net Billing regime"),
    ("Security firm management",
     PKG_NO, PKG_NO, PKG_YES, PKG_NO,
     "Not required — existing arrangement"),
    ("Warranty & claims administration",
     PKG_NO, PKG_NO, PKG_YES, PKG_OPT,
     "Available as add-on (Article 10.4)"),
    ("Full IV-curve testing (all strings)",
     PKG_NO, PKG_NO, PKG_YES, PKG_YES,
     "Every 2 years (once per initial Term)"),
]

GOLD_HEX_LIGHT = "FFF8E1"

for i, row in enumerate(pkg_rows):
    svc, eco, silver, gold, custom, notes = row
    is_custom_yes = custom not in (PKG_NO,)
    data = [
        (svc,    5.8, "L", False),
        (eco,    1.5, "C", False),
        (silver, 1.5, "C", False),
        (gold,   1.5, "C", False),
        (custom, 2.5, "C", True),
        (notes,  3.2, "L", False),
    ]
    trow = pkg_table.add_row()
    for j, (text, width, align_str, bold) in enumerate(data):
        cell = trow.cells[j]
        cell.width = Cm(width)
        # Highlight CUSTOM column in light gold if selected
        if j == 4 and text not in (PKG_NO,):
            set_cell_bg(cell, GOLD_HEX_LIGHT)
        elif i % 2 == 1:
            set_cell_bg(cell, ALT_HEX)
        p   = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(9)
        run.font.color.rgb = BLACK
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_str == "C" else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

# Pricing footer row
price_row = pkg_table.add_row()
price_cells = [
    ("Annual fee (ex. VAT) — 2 parks", 5.8, "L", True),
    ("On request", 1.5, "C", False),
    ("On request", 1.5, "C", False),
    ("On request", 1.5, "C", False),
    (f"€{PRICE_TOTAL:,.0f}", 2.5, "C", True),
    ("Total both parks / year", 3.2, "L", False),
]
for j, (text, width, align_str, bold) in enumerate(price_cells):
    cell = price_row.cells[j]
    cell.width = Cm(width)
    set_cell_bg(cell, NAVY_HEX)
    p   = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_str == "C" else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

add_paragraph(doc,
    "The detailed scope of each included service is set out in Articles 2.1–2.8 below. "
    "Services marked 'Add-on' are available under the Additional Works mechanism "
    "in Article 10 at agreed rates.",
    size=10, color=GREY, space_before=4, space_after=8)

add_heading(doc, "2.1  Preventive Maintenance", level=2)
add_paragraph(doc,
    "The Service Provider shall perform a minimum of two (2) full preventive maintenance "
    "visits per Park per Contract Year, scheduled in spring (March–April) and autumn "
    "(September–October), unless otherwise agreed in writing. Each visit shall include:",
    size=10.5, space_after=4)

prev_items = [
    "Full visual inspection of all PV modules, mounting structures, cables, connectors, and junction boxes.",
    "Thermal infrared (IR) thermography of all PV strings and modules to detect hot-spot defects and soiling anomalies.",
    "IV-curve testing on a representative sample (minimum 10% of strings per Park) using a calibrated IV curve tracer.",
    "Inspection and testing of all DC isolators, string fuses, and surge protection devices (SPDs).",
    "Inverter inspection: firmware version verification, fault log review, cooling system cleaning, filter replacement.",
    "MV/LV switchgear visual inspection, torque verification of busbar connections, and interlock functional testing.",
    "Meteorological station calibration check and irradiance sensor cleaning.",
    "Fence integrity, gate locks, CCTV, and perimeter security inspection.",
    "Earthing and lightning protection continuity check.",
    "Vegetation management: grass cutting and weed control within 0.5 m of all mounting structures and cable routes.",
    "Compilation of a written visit report with photographic evidence, fault register update, and recommended actions.",
]
for item in prev_items:
    add_bullet(doc, item)

add_heading(doc, "2.2  Panel Cleaning", level=2)
add_paragraph(doc,
    f"The Service Provider shall perform two (2) professional PV module cleaning sessions "
    f"per Park per Contract Year (once per preventive maintenance visit). Cleaning shall comply "
    f"with IEC TS 62788 and module manufacturer guidelines, using deionised or approved water "
    f"and soft-brush equipment. No abrasive agents or high-pressure jets above 40 bar shall be used. "
    f"Post-cleaning irradiance-normalised energy output shall be recorded for performance benchmarking.",
    size=10.5, space_after=6)

add_heading(doc, "2.3  Corrective Maintenance", level=2)
add_paragraph(doc,
    "The Service Provider shall provide corrective maintenance response as follows:",
    size=10.5, space_after=4)

response_items = [
    ("Critical fault (total Park shutdown, >50 kW loss): ",
     "on-site response within 4 hours of notification; restoration target 24 hours."),
    ("Major fault (>10% capacity loss, inverter failure, MV fault): ",
     "on-site response within 8 hours; restoration target 48 hours."),
    ("Minor fault (string failure, sensor fault, communication loss): ",
     "on-site response within 24 hours; resolution within 5 business days."),
    ("Cosmetic / administrative fault: ",
     "addressed at the next scheduled preventive maintenance visit, but in any event no later than 90 days from the date of identification."),
]
for bold_part, text_part in response_items:
    add_bullet(doc, text_part, bold_prefix=bold_part)

add_paragraph(doc,
    "Response times apply regardless of whether the call-out cap has been reached. "
    "A distinction is made between remote resolution and on-site attendance:",
    size=10.5, space_before=4, space_after=4)

callout_items = [
    ("Remote resolution: ",
     "fault diagnosed and resolved via SCADA, remote inverter access, or telephone guidance "
     "to the Client, without physical site attendance. Remote resolutions are unlimited and "
     "do not count against the call-out cap."),
    ("On-site call-out: ",
     f"unscheduled physical attendance at the Park not forming part of a scheduled preventive "
     f"maintenance visit. The Service Fee includes up to {CONTRACT['callout_cap_per_park']} on-site "
     f"corrective call-outs per Park per Contract Year. Each additional on-site call-out beyond "
     f"this cap is invoiced at €{CONTRACT['callout_extra_fee']:,} (ex. VAT) per attendance, "
     f"due within 14 days of invoice."),
]
for bold_part, text_part in callout_items:
    add_bullet(doc, text_part, bold_prefix=bold_part)
doc.add_paragraph()

add_heading(doc, "2.4  SCADA / Performance Monitoring", level=2)
monitoring_items = [
    "24/7 remote monitoring of both Parks via SCADA with automatic alarm notification to the Service Provider's operations centre.",
    "Monthly performance report issued within 10 business days of month-end, covering: energy yield (MWh), specific yield (kWh/kWp), Performance Ratio (PR), Availability, fault log, and cleaning records.",
    "Quarterly trend analysis identifying degradation, soiling losses, and inverter performance deviations.",
    "Annual performance summary report with year-on-year comparison and maintenance recommendations.",
    "Immediate notification to Client for any fault or event causing >5% instantaneous capacity loss.",
]
for item in monitoring_items:
    add_bullet(doc, item)

add_heading(doc, "2.5  Electrical Safety & Compliance Inspections", level=2)
compliance_items = [
    "Annual insulation resistance (IR) testing of all DC and AC cabling per IEC 62446-1.",
    "Annual earth continuity testing and earth fault loop impedance measurement.",
    "Full IV-curve testing of all strings per Park — performed every two (2) Contract Years (i.e. once during the initial Term), in addition to the sample testing carried out at each preventive maintenance visit.",
    "Maintenance of electrical safety records and asset register in compliance with Cyprus Electricity Authority (EAC) requirements.",
    "Coordination with the EAC for any grid-related inspections or connection tests.",
]
for item in compliance_items:
    add_bullet(doc, item)

add_heading(doc, "2.6  Vegetation & Civil Maintenance", level=2)
civil_items = [
    "Minimum four (4) grass-cutting and vegetation control sessions per Park per Contract Year (spring, early summer, late summer, autumn).",
    "Inspection of civil foundations, cable ducting, and drainage systems; reporting of defects to Client.",
    "Post-storm site inspection within 48 hours of any storm event with wind speeds exceeding 80 km/h.",
    "Minor civil repairs (sealing of cable duct penetrations, fence post re-setting) up to €500 per incident included; larger works quoted separately.",
]
for item in civil_items:
    add_bullet(doc, item)

add_heading(doc, "2.7  Spare Parts & Consumables", level=2)
add_paragraph(doc,
    "The Service Provider shall maintain an on-island spare parts inventory covering: "
    "string fuse elements, DC SPDs, AC SPDs, inverter air filters, connector caps, "
    "and cable termination materials. Replacement of these consumable items during "
    "Preventive or Corrective Maintenance visits is included in the Service Fee. "
    "Major components (inverters, transformers, modules) are excluded and quoted separately.",
    size=10.5, space_after=6)

add_heading(doc, "2.8  Health, Safety & Environmental (HSE)", level=2)
hse_items = [
    "All works to be performed in compliance with the Cyprus Safety and Health at Work Law (89(I)/1996) and all applicable EU Directives.",
    "Service Provider to maintain a site-specific HSE Plan and Method Statements for all planned activities.",
    "Toolbox talks to be documented prior to each site visit.",
    "Environmental management: collection and licensed disposal of all waste materials including failed modules, oils, and packaging.",
    "All Service Provider personnel and subcontractors attending either Park must carry Lighthief-issued site identification and log their entry and exit in the site visitor register maintained at each Park.",
]
for item in hse_items:
    add_bullet(doc, item)

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 3 — AVAILABILITY GUARANTEE & PERFORMANCE PENALTIES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 3 — Availability Guarantee & Performance Penalties", level=1)

add_heading(doc, "3.1  Availability Guarantee", level=2)
add_paragraph(doc,
    f"The Service Provider guarantees a minimum Availability of "
    f"{CONTRACT['availability_guarantee_pct']}% for each Park, measured annually per "
    f"Contract Year from the Service Commencement Date. Availability is calculated as:",
    size=10.5, space_after=4)

p = doc.add_paragraph()
p.paragraph_format.left_indent  = Cm(1.5)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("Availability (%) = [(Eligible Hours - Unavailable Hours) / Eligible Hours] x 100")
r.bold = True; r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.paragraph_format.left_indent = Cm(1.5)
p2.paragraph_format.space_after = Pt(6)
r2 = p2.add_run("Where: Eligible Hours = Total Hours - Excused Downtime Hours")
r2.italic = True; r2.font.size = Pt(10); r2.font.color.rgb = GREY

add_paragraph(doc,
    "Measurement methodology: Availability is determined from SCADA inverter status data "
    "at 5-minute resolution. In the event of SCADA data loss, inverter event logs constitute "
    "the secondary source, followed by grid meter import/export records as tertiary source. "
    "Disputed periods default to Excused Downtime where the cause cannot be unambiguously "
    "attributed to Service Provider failure. Any dispute on Availability measurement shall "
    "first be escalated to the Parties' technical representatives; if unresolved within "
    "14 days, an Independent Engineer (jointly appointed) shall provide a binding determination "
    "within 30 days.",
    size=10, color=GREY, space_before=0, space_after=6)

add_heading(doc, "3.2  Availability Penalty Schedule", level=2)
add_paragraph(doc,
    "Where the measured annual Availability falls below the guaranteed 99%, the following "
    "liquidated damages shall be payable by the Service Provider to the Client within "
    "30 days of the end of the relevant Contract Year:",
    size=10.5, space_after=6)

pen_table = doc.add_table(rows=1, cols=4)
pen_table.alignment = WD_TABLE_ALIGNMENT.LEFT
pen_table.style = "Table Grid"
add_table_row(pen_table, [
    ("Annual Availability Achieved", 4.5, "C", True),
    ("Penalty Rate", 3.5, "C", True),
    ("Basis", 3.5, "C", True),
    ("Annual Cap", 4.5, "C", True),
], header=True)

penalty_rows = [
    ("98.00% – 98.99%", "5% of annual Service Fee per Park", "Per shortfall band", "—"),
    ("97.00% – 97.99%", "10% of annual Service Fee per Park", "Per shortfall band", "—"),
    ("95.00% – 96.99%", "20% of annual Service Fee per Park", "Per shortfall band", "—"),
    ("< 95.00%",        "30% of annual Service Fee per Park", "Per shortfall band", "30% of annual Service Fee per Park"),
]
for i, row_data in enumerate(penalty_rows):
    add_table_row(pen_table, [
        (row_data[0], 4.5, "C", False),
        (row_data[1], 3.5, "C", False),
        (row_data[2], 3.5, "C", False),
        (row_data[3], 4.5, "C", False),
    ], alt=(i % 2 == 1))

add_paragraph(doc,
    "A grace band of 0.5 percentage points applies: no penalty is triggered unless "
    "Availability falls below 98.5% (i.e. the first penalty band activates at <98.5%, "
    "not <99%). The total annual liquidated damages under this Article, combined with "
    "any PR penalties under Article 3.3, shall not exceed 30% of the annual Service Fee "
    "for the relevant Park in any Contract Year. Liquidated damages represent the "
    "Client's sole financial remedy for Availability and PR shortfalls.",
    size=10, color=GREY, space_before=6, space_after=6)

add_heading(doc, "3.3  Performance Ratio (PR) Guarantee", level=2)
add_paragraph(doc,
    "The Service Provider guarantees a minimum annual average Performance Ratio (PR) of "
    "75% for each Park, calculated per IEC 61724-1 using in-plane irradiation from the "
    "on-site meteorological station, with all Excused Downtime periods excluded from both "
    "the measured and reference yield. The following conditions are expressly excluded "
    "from the PR guarantee and shall not give rise to any liability or penalty:",
    size=10.5, space_after=4)

pr_exclusions = [
    "Curtailment or export limitation imposed by EAC, the DSO, TSO, or any grid operator.",
    "Irradiance sensor drift, failure, or data gaps not attributable to Service Provider negligence.",
    "Grid outages, voltage fluctuations, or frequency deviations outside equipment tolerances.",
    "Design defects, module degradation beyond manufacturer-warranted rates, or soiling losses during periods between scheduled cleaning visits.",
    "Force majeure weather anomalies, including irradiance levels materially below the 10-year P50 baseline for the site.",
    "Any restriction imposed by a regulatory authority, EAC licence condition, or net billing scheme cap.",
]
for item in pr_exclusions:
    add_bullet(doc, item)

add_paragraph(doc,
    "Where the measured annual PR (net of excluded periods) falls below 75%, the Service "
    "Provider shall, at no additional cost:",
    size=10.5, space_before=6, space_after=4)

pr_items = [
    "Conduct a full root-cause analysis within 10 business days, identifying factors within and outside the Service Provider's control.",
    "Present a corrective action plan addressing controllable factors to the Client within 20 business days.",
    "Implement approved corrective actions within the agreed timeframe.",
    "If PR remains below 70% for two consecutive Contract Years due to factors demonstrably within the Service Provider's control, the Client may terminate this Agreement without penalty.",
]
for item in pr_items:
    add_bullet(doc, item)

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 4 — EXCLUDED SERVICES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 4 — Excluded Services", level=1)
add_paragraph(doc,
    "The following services are not included in the selected Custom package and are "
    "expressly excluded from this Agreement. Services marked (Add-on) may be commissioned "
    "via the Additional Works mechanism in Article 10. All other items below require "
    "a separate agreement or specialist contractor.",
    size=10.5, space_after=6)

add_heading(doc, "4.0  Package Exclusions (Custom vs GOLD tier)", level=2)
add_paragraph(doc,
    "The following services are available under the GOLD tier but have been excluded "
    "from the Custom package for the reasons stated:",
    size=10.5, space_after=6)

excl_pkg_table = doc.add_table(rows=1, cols=3)
excl_pkg_table.style = "Table Grid"
excl_pkg_table.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_row(excl_pkg_table, [
    ("Excluded Service",        7.0, "L", True),
    ("Available As",            3.0, "C", True),
    ("Reason for Exclusion",    6.0, "L", True),
], header=True)

excl_pkg_rows = [
    ("Transformer general audit",
     "Add-on (Art. 10)",
     "To be commissioned if/when required based on equipment condition"),
    ("Insurance management",
     "Not required",
     "Client maintains own property and liability insurance"),
    ("PPA / energy contract management",
     "Not applicable",
     "Parks operate under Cyprus Net Billing scheme, not merchant PPA"),
    ("Security firm management",
     "Not required",
     "Client has existing security arrangement in place"),
    ("Warranty & claims administration",
     "Add-on (Art. 10.4)",
     "Available as billable service; Service Provider to flag warranty defects"),
]
for i, (svc, avail, reason) in enumerate(excl_pkg_rows):
    add_table_row(excl_pkg_table, [
        (svc,    7.0, "L", False),
        (avail,  3.0, "C", False),
        (reason, 6.0, "L", False),
    ], alt=(i % 2 == 1))

doc.add_paragraph()

add_heading(doc, "4.1  General Exclusions (all tiers)", level=2)
add_paragraph(doc,
    "The following services are excluded from all service tiers and require separate "
    "quotation or specialist engagement:",
    size=10.5, space_after=4)

excluded = [
    "Replacement of PV modules (including warranty claim administration unless commissioned under Article 10.4).",
    "Replacement of string inverters, central inverters, or transformers.",
    "Major electrical works including MV cable replacement or substation upgrades.",
    "Civil engineering works beyond minor repairs (as defined in Article 2.6).",
    "Grid connection modifications or grid operator compliance testing.",
    "Cybersecurity or IT infrastructure services beyond SCADA maintenance under this Agreement.",
    "Works required as a result of vandalism, theft, or third-party damage (quoted separately).",
]
for item in excluded:
    add_bullet(doc, item)

add_heading(doc, "4.1  BESS Operation & Maintenance — Separate Agreement", level=2)
add_paragraph(doc,
    "Both Parks include Battery Energy Storage System (BESS) installations currently "
    "subject to a separate Long-Term Service Agreement (LTSA) between the Parties. "
    "BESS O&M services are expressly excluded from the scope of this PV O&M Agreement. "
    "For reference, the key service response times under the separate BESS LTSA are:",
    size=10.5, space_after=4)

bess_response_items = [
    ("Critical BESS fault (full system shutdown, fire suppression activation, BMS alarm): ",
     "on-site response within 2 hours; system isolation and safe-state confirmation within 4 hours."),
    ("Major BESS fault (>20% capacity loss, PCS failure, thermal event): ",
     "on-site response within 4 hours; restoration target 24 hours."),
    ("Minor BESS fault (single container fault, communication loss, BMS warning): ",
     "on-site response within 8 hours; resolution within 48 hours."),
    ("Planned BESS preventive maintenance: ",
     "2 visits per year per system, scheduled in coordination with PV preventive maintenance visits "
     "to minimise combined downtime across both Parks."),
]
for bold_part, text_part in bess_response_items:
    add_bullet(doc, text_part, bold_prefix=bold_part)

add_paragraph(doc,
    "Where a fault event affects both the PV system and the BESS simultaneously, "
    "the Service Provider shall coordinate response under both agreements and provide "
    "a single unified incident report to the Client within 24 hours.",
    size=10, color=GREY, space_before=4, space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 5 — CLIENT OBLIGATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 5 — Client Obligations", level=1)
add_paragraph(doc,
    "The Client shall throughout the Term of this Agreement:", size=10.5, space_after=4)
client_obs = [
    "Within 7 days of the date of signing this Agreement, hand over to the Service Provider — by means of a written handover protocol — all keys, access codes, gate remote controls, SCADA login credentials, and any other materials necessary for unrestricted 24/7 access to both Parks.",
    "Provide the Service Provider and its personnel with unrestricted access to both Parks at all reasonable times, with reasonable notice except in cases of emergency.",
    "Maintain valid property and public liability insurance for both Parks and provide evidence of cover on request.",
    "Ensure that both Parks are connected to the SCADA / monitoring system and that communication links remain operational; promptly notify the Service Provider of any communication outage.",
    "Notify the Service Provider promptly of any observed fault, alarm, or unusual operating condition.",
    "Not perform or authorise any maintenance, modification, or repair to the Parks without the Service Provider's prior written consent, except in cases of emergency.",
    "Pay all Service Fees in accordance with Article 9.",
    "Provide the Service Provider with copies of all existing warranties, manuals, as-built drawings, and grid connection documentation within 14 days of the Commencement Date.",
]
for item in client_obs:
    add_bullet(doc, item)

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 6 — REPORTING & RECORDS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 6 — Reporting & Records", level=1)
report_items = [
    ("Monthly Performance Report: ", "issued by the 10th of each calendar month covering the preceding month's data (see Schedule C below). Each monthly report constitutes the basis for issuance of the monthly invoice under Article 9."),
    ("Maintenance Visit Report: ",   "issued within 5 business days of each site visit, including photographs, findings, and recommendations."),
    ("Annual Summary Report: ",      "issued within 30 days of each Contract Year-end, including year-on-year performance comparison, asset condition assessment, and maintenance plan for the following year."),
    ("Ad-hoc Incident Reports: ",    "issued within 24 hours of any Critical fault event."),
]
for bold_part, text_part in report_items:
    add_bullet(doc, text_part, bold_prefix=bold_part)

add_paragraph(doc,
    "All reports shall be delivered electronically in PDF format. The Service Provider "
    "shall maintain a maintenance log, fault register, and asset condition record for "
    "each Park and make these available to the Client on request.",
    size=10.5, space_after=6)

add_heading(doc, "Schedule C — Monthly Performance Report Required Fields", level=2)
add_paragraph(doc,
    "Each monthly report shall contain the following fields for each Park:",
    size=10.5, space_after=6)

rpt_table = doc.add_table(rows=1, cols=3)
rpt_table.alignment = WD_TABLE_ALIGNMENT.LEFT
rpt_table.style = "Table Grid"
add_table_row(rpt_table, [
    ("Field", 5.5, "L", True),
    ("Unit / Format", 3.5, "C", True),
    ("Notes", 7.0, "L", True),
], header=True)

rpt_rows = [
    ("Reporting Period",         "Month / Year",       "e.g. April 2026"),
    ("Park Name",                "Text",               "Anarita East / West"),
    ("Energy Yield",             "MWh",                "Total AC export for the period"),
    ("Specific Yield",           "kWh/kWp",            "Yield per installed kWp"),
    ("Performance Ratio (PR)",   "%",                  "Per IEC 61724-1"),
    ("Availability",             "%",                  "Per Article 3 definition"),
    ("Irradiation (in-plane)",   "kWh/m2",             "From on-site met station"),
    ("Fault Events — Critical",  "Count",              "Including duration and resolution"),
    ("Fault Events — Major",     "Count",              "Including duration and resolution"),
    ("Fault Events — Minor",     "Count",              "Including resolution date"),
    ("Corrective Call-outs Used","Count (of cap)",     f"e.g. 2 of {CONTRACT['callout_cap_per_park']} included"),
    ("Cleaning Carried Out",     "Yes / No",           "Date(s) if performed"),
    ("Maintenance Visit",        "Yes / No / Scheduled","Date if performed"),
    ("Recommended Actions",      "Free text",          "Items for next visit or client decision"),
    ("Prepared By",              "Name + Date",        "Service Provider signatory"),
]
for i, (field, unit, notes) in enumerate(rpt_rows):
    add_table_row(rpt_table, [
        (field, 5.5, "L", True),
        (unit,  3.5, "C", False),
        (notes, 7.0, "L", False),
    ], alt=(i % 2 == 1))

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 7 — TERM & RENEWAL
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 7 — Term & Renewal", level=1)
add_paragraph(doc,
    f"This Agreement is executed on the Effective Date ({CONTRACT['effective_date']}) and "
    f"O&M services commence on the Service Commencement Date ({CONTRACT['commencement_date']}). "
    f"The initial Term runs for {CONTRACT['term_years']} (two) Contract Years from the "
    f"Service Commencement Date unless terminated earlier in accordance with Article 14. "
    f"Unless either Party gives written notice of non-renewal at least ninety (90) days "
    f"before the expiry of the initial Term or any renewal period, this Agreement shall "
    f"automatically renew for successive one (1) year periods on the same terms, subject "
    f"to the annual Service Fee adjustment set out in Article 9.4.",
    size=10.5, space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 8 — SERVICE LEVELS & KPIs
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 8 — Service Levels & Key Performance Indicators", level=1)

kpi_table = doc.add_table(rows=1, cols=3)
kpi_table.alignment = WD_TABLE_ALIGNMENT.LEFT
kpi_table.style = "Table Grid"
add_table_row(kpi_table, [
    ("KPI", 5.5, "L", True),
    ("Target", 4.5, "C", True),
    ("Measurement Basis", 6.0, "L", True),
], header=True)

kpi_rows = [
    ("Annual Availability (per Park)",              f"≥ {CONTRACT['availability_guarantee_pct']}%",  "Annual, per IEC 61724-1"),
    ("Annual Performance Ratio (per Park)",         "≥ 75%",           "Annual, per IEC 61724-1"),
    ("Critical Fault Response Time",                "≤ 4 hours",       "From notification receipt"),
    ("Major Fault Response Time",                   "≤ 8 hours",       "From notification receipt"),
    ("Minor Fault Response Time",                   "≤ 24 hours",      "From notification receipt"),
    ("Cosmetic Fault Resolution",                   "≤ 90 days",       "From date of identification"),
    ("Corrective Call-outs Included (per Park/yr)", f"{CONTRACT['callout_cap_per_park']} per year", f"Additional at €{CONTRACT['callout_extra_fee']:,}/call-out"),
    ("Preventive Maintenance Visits (per Park)",    "2 per year",      "Scheduled spring & autumn"),
    ("Panel Cleaning Sessions (per Park)",          "2 per year",      "Per preventive visit"),
    ("Monthly Report + Invoice",                    "By 10th of month","Electronic PDF + VAT invoice"),
    ("Grass Cutting & Vegetation Control",          "4 per year",      "Per Park"),
]
for i, (kpi, target, basis) in enumerate(kpi_rows):
    add_table_row(kpi_table, [
        (kpi,    5.5, "L", False),
        (target, 4.5, "C", True),
        (basis,  6.0, "L", False),
    ], alt=(i % 2 == 1))

doc.add_paragraph()

add_heading(doc, "8.1  Response Time Breach — Service Credits", level=2)
add_paragraph(doc,
    "Where the Service Provider fails to meet the contracted response times for Critical "
    "or Major faults, the following service credits shall be applied as a deduction from "
    "the next monthly invoice:",
    size=10.5, space_after=6)

sla_table = doc.add_table(rows=1, cols=4)
sla_table.alignment = WD_TABLE_ALIGNMENT.LEFT
sla_table.style = "Table Grid"
add_table_row(sla_table, [
    ("Fault Type",          4.0, "L", True),
    ("Contracted SLA",      3.0, "C", True),
    ("Credit per Hour Exceeded", 4.0, "C", True),
    ("Maximum Credit / Incident", 4.0, "C", True),
], header=True)

sla_credit_rows = [
    ("Critical fault",  "4 hours on-site", f"0.5% of monthly fee per hour", "5% of monthly fee"),
    ("Major fault",     "8 hours on-site", f"0.25% of monthly fee per hour", "3% of monthly fee"),
    ("Minor fault",     "24 hours on-site", "No credit — tracked only", "—"),
]
for i, row in enumerate(sla_credit_rows):
    add_table_row(sla_table, [
        (row[0], 4.0, "L", False),
        (row[1], 3.0, "C", False),
        (row[2], 4.0, "C", False),
        (row[3], 4.0, "C", False),
    ], alt=(i % 2 == 1))

add_paragraph(doc,
    "Service credits are the Client's sole remedy for response time breaches and do not "
    "constitute a finding of liability. Credits are not payable where delay is caused by "
    "Client-side access restriction, force majeure, or Excused Downtime.",
    size=10, color=GREY, space_before=4, space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 9 — SERVICE FEE & PAYMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 9 — Service Fee & Payment Terms", level=1)

add_heading(doc, "9.1  Annual Service Fee", level=2)
add_paragraph(doc,
    "The annual Service Fee is a fixed lump-sum per Park per Contract Year, "
    "regardless of installed capacity or energy production. The combined annual "
    f"fee for both Parks is €{28_200*2:,} (excluding VAT). "
    "The full cost breakdown including VAT is set out below:",
    size=10.5, space_after=6)

fee_table = doc.add_table(rows=1, cols=4)
fee_table.alignment = WD_TABLE_ALIGNMENT.LEFT
fee_table.style = "Table Grid"
add_table_row(fee_table, [
    ("Park",                       5.5, "L", True),
    ("Annual Fee (ex. VAT)",       3.5, "R", True),
    (f"VAT ({int(VAT_RATE*100)}%)",2.5, "R", True),
    ("Annual Fee (incl. VAT)",     3.5, "R", True),
], header=True)

fee_rows = [
    ("Anarita East Park",  f"€{PRICE_PER_PARK:,.0f}",  f"€{VAT_PER_PARK:,.0f}",  f"€{PRICE_PER_PARK_INCL:,.0f}"),
    ("Anarita West Park",  f"€{PRICE_PER_PARK:,.0f}",  f"€{VAT_PER_PARK:,.0f}",  f"€{PRICE_PER_PARK_INCL:,.0f}"),
    ("TOTAL",              f"€{PRICE_TOTAL:,.0f}",      f"€{VAT_TOTAL:,.0f}",      f"€{PRICE_TOTAL_INCL:,.0f}"),
]
for i, (park, ex, vat, incl) in enumerate(fee_rows):
    is_total = (i == 2)
    add_table_row(fee_table, [
        (park, 5.5, "L", is_total),
        (ex,   3.5, "R", is_total),
        (vat,  2.5, "R", is_total),
        (incl, 3.5, "R", is_total),
    ], alt=(i % 2 == 1))

add_paragraph(doc,
    f"VAT is applied at the rate of {int(VAT_RATE*100)}% in accordance with Cyprus VAT Law "
    f"(95(I)/2000) as in force at the time of invoicing. Should the applicable VAT rate change, "
    f"invoices shall reflect the rate in force at the invoice date.",
    size=10, color=GREY, space_before=4, space_after=6)

add_heading(doc, "9.2  Monthly Payment Schedule", level=2)
add_paragraph(doc,
    "The annual Service Fee shall be invoiced monthly in arrears. The Service Provider shall "
    "issue a VAT invoice by the 10th of each calendar month for services rendered in the "
    "preceding month. Each invoice shall be accompanied by the Monthly Performance Report "
    "for the relevant Park(s) as set out in Schedule C. Payment is due within 30 days of "
    "the invoice date.",
    size=10.5, space_after=4)

monthly_table = doc.add_table(rows=1, cols=3)
monthly_table.alignment = WD_TABLE_ALIGNMENT.LEFT
monthly_table.style = "Table Grid"
add_table_row(monthly_table, [
    ("Payment Item",            6.0, "L", True),
    ("Amount (ex. VAT)",        3.5, "R", True),
    ("Amount (incl. VAT)",      3.5, "R", True),
], header=True)

monthly_rows = [
    ("Monthly instalment — both Parks", f"€{PRICE_MONTHLY_EX:,.2f}", f"€{PRICE_MONTHLY_INCL:,.2f}"),
    ("Monthly instalment — East Park only", f"€{PRICE_MONTHLY_EX/2:,.2f}", f"€{PRICE_MONTHLY_INCL/2:,.2f}"),
    ("Monthly instalment — West Park only", f"€{PRICE_MONTHLY_EX/2:,.2f}", f"€{PRICE_MONTHLY_INCL/2:,.2f}"),
]
for i, (item, ex, incl) in enumerate(monthly_rows):
    add_table_row(monthly_table, [
        (item, 6.0, "L", i == 0),
        (ex,   3.5, "R", i == 0),
        (incl, 3.5, "R", i == 0),
    ], alt=(i % 2 == 1))

add_paragraph(doc,
    "Where the Commencement Date falls mid-month, the first invoice shall be pro-rated "
    "to the actual number of days of service in that month.",
    size=10, color=GREY, space_before=4, space_after=6)

add_heading(doc, "9.3  Late Payment", level=2)
add_paragraph(doc,
    "Overdue amounts shall accrue interest at 2% per annum above the European Central Bank "
    "base rate from the due date until the date of payment. The Service Provider reserves "
    "the right to suspend non-emergency services after 45 days of non-payment, with written "
    "notice of 10 business days.",
    size=10.5, space_after=6)

add_heading(doc, "9.4  Annual Fee Escalation", level=2)
add_paragraph(doc,
    "From the second Contract Year onwards, the Service Fee shall be adjusted annually "
    "on 1 January by the change in the Harmonised Index of Consumer Prices for Cyprus "
    "(HICP-CY), as published by Eurostat, for the preceding calendar year. "
    "The maximum annual escalation shall not exceed 4%. The indexation does not apply "
    "during the first Contract Year.",
    size=10.5, space_after=6)

add_heading(doc, "9.5  Additional Call-Out Charges", level=2)
add_paragraph(doc,
    f"Corrective call-outs beyond the included cap of {CONTRACT['callout_cap_per_park']} per Park "
    f"per Contract Year (as defined in Article 2.3) shall be invoiced separately at "
    f"€{CONTRACT['callout_extra_fee']:,} (ex. VAT) per call-out, due within 14 days of invoice.",
    size=10.5, space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 10 — ADDITIONAL & URGENT WORKS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 10 — Additional & Urgent Works", level=1)

add_heading(doc, "10.1  Additional Works — Order Procedure", level=2)
add_paragraph(doc,
    "Works outside the scope of this Agreement (see Article 4) may be commissioned by "
    "the Client via an email-confirmed written order using the format set out in "
    "Schedule D below. No additional works shall commence until the Service Provider "
    "has confirmed the order in writing to the Client's designated email address, "
    "together with the agreed cost and timeline.",
    size=10.5, space_after=6)

add_heading(doc, "10.2  Urgent Works — Emergency Carve-Out", level=2)
add_paragraph(doc,
    "Where an urgent and necessary intervention is required to prevent imminent risk to "
    "life, health, property, or continued operation of a Park — and prior approval cannot "
    "reasonably be obtained — the Service Provider may carry out such works without a "
    "pre-confirmed order. The Service Provider shall:",
    size=10.5, space_after=4)

urgent_items = [
    "Notify the Client by telephone and email within 4 hours of commencing the urgent works.",
    "Submit a written cost summary to the Client within 5 business days of completion.",
    "Invoice the cost of such works separately from the monthly Service Fee invoice; payment due within 14 days.",
]
for item in urgent_items:
    add_bullet(doc, item)

add_paragraph(doc,
    "The emergency carve-out does not apply to works that could reasonably have been "
    "anticipated and pre-approved during normal business hours.",
    size=10, color=GREY, space_before=4, space_after=6)

add_heading(doc, "Schedule D — Additional Works Order Form", level=2)
add_paragraph(doc,
    "The following format shall be used for all additional works orders under Article 10.1. "
    "Orders must be sent to the Service Provider's designated email address and confirmed "
    "in writing before works commence.",
    size=10.5, space_after=6)

ord_table = doc.add_table(rows=1, cols=2)
ord_table.alignment = WD_TABLE_ALIGNMENT.LEFT
ord_table.style = "Table Grid"
add_table_row(ord_table, [
    ("Field", 5.0, "L", True),
    ("To Be Completed By Client / Agreed By Both Parties", 11.0, "L", True),
], header=True)

ord_rows = [
    ("Order Reference",          "[Auto-assigned by Service Provider]"),
    ("Date of Order",            ""),
    ("Park(s) Affected",         "Anarita East  /  Anarita West  /  Both"),
    ("Description of Works",     ""),
    ("Reason / Trigger",         ""),
    ("Estimated Duration",       ""),
    ("Agreed Cost (ex. VAT)",    "€"),
    ("VAT Amount (19%)",         "€"),
    ("Total Cost (incl. VAT)",   "€"),
    ("Authorised By (Client)",   "Name: _________________  Signature: _________________  Date: ________"),
    ("Confirmed By (Contractor)","Name: _________________  Signature: _________________  Date: ________"),
    ("Target Completion Date",   ""),
    ("Completion Confirmed",     "Date: _______________  Sign-off: _______________"),
]
for i, (field, value) in enumerate(ord_rows):
    add_table_row(ord_table, [
        (field, 5.0, "L", True),
        (value, 11.0, "L", False),
    ], alt=(i % 2 == 1))

doc.add_paragraph()

add_heading(doc, "10.3  SCADA Data Ownership & Access Rights", level=2)
scada_items = [
    "All performance data generated by the Parks' SCADA and monitoring systems remains the exclusive property of the Client.",
    "The Service Provider is granted a licence to access, store, and process SCADA data solely for the purposes of performing services under this Agreement.",
    "The Service Provider shall not share, sell, or otherwise transfer SCADA data or derived analytics to any third party without the Client's prior written consent.",
    "On termination of this Agreement, the Service Provider shall deliver a complete export of all SCADA data collected during the Term to the Client within 14 days, in a standard format (CSV or equivalent), and shall permanently delete all retained copies.",
    "The Service Provider is responsible for the cybersecurity of any monitoring equipment, remote access credentials, or VPN connections under its direct management. Any security breach affecting SCADA access must be reported to the Client within 4 hours of discovery.",
]
for item in scada_items:
    add_bullet(doc, item)

doc.add_paragraph()

add_heading(doc, "10.4  Warranty Management (Optional Service)", level=2)
add_paragraph(doc,
    "Warranty claim administration on behalf of the Client (for PV modules, inverters, "
    "mounting structures, or other equipment covered by manufacturer warranties) is not "
    "included in the standard Service Fee. The Service Provider may provide warranty "
    "management as an additional service upon written request, charged at the hourly "
    "rates agreed in a separate Additional Works Order (Schedule D). "
    "The Service Provider shall in all cases: (i) promptly notify the Client of any "
    "observed defect that may give rise to a warranty claim; and (ii) preserve evidence "
    "and documentation required to support a claim.",
    size=10.5, space_after=6)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 11 — INSURANCE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 11 — Insurance", level=1)
add_paragraph(doc,
    "Throughout the Term, the Service Provider shall maintain the following minimum "
    "insurance cover and provide evidence to the Client on request:",
    size=10.5, space_after=4)

insurance_items = [
    ("Public Liability Insurance: ", "minimum €2,000,000 per occurrence."),
    ("Employers' Liability Insurance: ", "as required by Cyprus law."),
    ("Professional Indemnity Insurance: ", "minimum €500,000 per claim."),
    ("Motor Vehicle Insurance: ", "third-party liability for all vehicles used on-site."),
]
for bold_part, text_part in insurance_items:
    add_bullet(doc, text_part, bold_prefix=bold_part)

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 11 — LIABILITY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 12 — Liability & Indemnity", level=1)
liability_paras = [
    ("12.1", "Each Party shall indemnify the other against any third-party claims, losses, damages, "
             "and costs arising from its negligent acts or omissions in connection with this Agreement."),
    ("12.2", "The Service Provider's total aggregate liability to the Client under this Agreement "
             "in any Contract Year shall not exceed 100% of the annual Service Fee payable for that "
             "Contract Year, except in cases of fraud, wilful misconduct, or death or personal injury "
             "caused by negligence."),
    ("12.3", "Neither Party shall be liable to the other for any indirect, consequential, or "
             "special losses, including loss of revenue, loss of profit, or loss of production, "
             "howsoever arising."),
    ("12.4", "Liquidated damages payable under Article 3 represent the Client's sole financial "
             "remedy for Availability and Performance Ratio shortfalls."),
]
for num, text in liability_paras:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(num + "  ")
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 12 — FORCE MAJEURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 13 — Force Majeure", level=1)
add_paragraph(doc,
    "Neither Party shall be in breach of this Agreement, nor liable for any failure or "
    "delay in performance of its obligations, to the extent that such failure or delay "
    "results from a Force Majeure event, being any event beyond a Party's reasonable "
    "control, including acts of God, natural disasters, war, terrorism, government action, "
    "pandemic, grid outages caused by the national transmission operator, or civil unrest. "
    "The Party invoking Force Majeure shall notify the other Party within 48 hours and "
    "resume performance as soon as reasonably practicable.",
    size=10.5, space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 13 — TERMINATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 14 — Termination", level=1)
term_items = [
    ("14.1  Termination for Convenience: ",
     "Either Party may terminate this Agreement at the end of any Contract Year "
     "by giving at least ninety (90) days' written notice."),
    ("14.2  Termination for Cause: ",
     "Either Party may terminate this Agreement immediately by written notice if "
     "the other Party: (i) commits a material breach and fails to remedy it within "
     "30 days of written notice; (ii) becomes insolvent, enters administration, or "
     "ceases business; or (iii) commits fraud or wilful misconduct."),
    ("14.3  Consequences of Termination: ",
     "On termination, the Service Provider shall hand over all Park records, "
     "maintenance logs, spare parts, and SCADA access credentials to the Client "
     "within 14 days. Any fees accrued and unpaid as of the termination date remain "
     "payable. Prepaid fees relating to periods after the termination date shall be "
     "refunded on a pro-rata basis."),
]
for bold_part, text_part in term_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(bold_part)
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
    r2 = p.add_run(text_part)
    r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 14 — CONFIDENTIALITY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 15 — Confidentiality", level=1)
add_paragraph(doc,
    "Each Party undertakes to keep confidential all Confidential Information received "
    "from the other Party and not to disclose it to any third party without prior "
    "written consent, except as required by law or regulation. This obligation "
    "shall survive termination of this Agreement for a period of three (3) years.",
    size=10.5, space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 15 — GOVERNING LAW & DISPUTE RESOLUTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 16 — Governing Law & Dispute Resolution", level=1)
add_paragraph(doc,
    "This Agreement shall be governed by and construed in accordance with the laws "
    "of the Republic of Cyprus. Any dispute arising out of or in connection with this "
    "Agreement shall first be submitted to mediation administered by the Cyprus Arbitration "
    "& Mediation Centre (CAMC) in accordance with its Mediation Rules in force at the time "
    "of the dispute. If the dispute is not resolved within 45 days of the appointment of the "
    "mediator (or such longer period as the Parties may agree in writing), either Party may "
    "refer the dispute to the exclusive jurisdiction of the courts of the Republic of Cyprus.",
    size=10.5, space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
# ARTICLE 16 — GENERAL PROVISIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Article 17 — General Provisions", level=1)
general_items = [
    ("Entire Agreement: ", "This Agreement constitutes the entire agreement between the Parties relating to its subject matter and supersedes all prior negotiations, representations, and agreements."),
    ("Amendments: ", "No amendment to this Agreement shall be effective unless made in writing and signed by authorised representatives of both Parties."),
    ("Assignment: ", "Neither Party may assign its rights or obligations under this Agreement without the prior written consent of the other Party, not to be unreasonably withheld."),
    ("Severability: ", "If any provision of this Agreement is held invalid or unenforceable, the remaining provisions shall continue in full force and effect."),
    ("Notices: ", "All notices shall be in writing and delivered by email with read receipt or by registered post to the addresses set out on the cover page."),
    ("Waiver: ", "Failure to exercise or delay in exercising any right shall not constitute a waiver of that right."),
    ("Counterparts: ", "This Agreement may be executed in counterparts, each of which shall constitute an original, and together they shall form one agreement."),
]
for bold_part, text_part in general_items:
    add_bullet(doc, text_part, bold_prefix=bold_part)

# ══════════════════════════════════════════════════════════════════════════════
# SCHEDULES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "Schedule A — Annual Maintenance Plan", level=1)
add_paragraph(doc, "Minimum maintenance activities per Park per Contract Year:", size=10.5, space_after=6)

sched_table = doc.add_table(rows=1, cols=5)
sched_table.alignment = WD_TABLE_ALIGNMENT.LEFT
sched_table.style = "Table Grid"
add_table_row(sched_table, [
    ("Activity",              5.0, "L", True),
    ("Q1 (Jan–Mar)",          2.5, "C", True),
    ("Q2 (Apr–Jun)",          2.5, "C", True),
    ("Q3 (Jul–Sep)",          2.5, "C", True),
    ("Q4 (Oct–Dec)",          2.5, "C", True),
], header=True)

sched_rows = [
    ("Full preventive maintenance visit",        "",   "Yes", "",    "Yes"),
    ("PV module cleaning",                       "",   "Yes", "",    "Yes"),
    ("IR thermography scan",                     "",   "Yes", "",    "Yes"),
    ("IV-curve testing (sample, min 10% of strings)", "",   "Yes", "",    "Yes"),
    ("IV-curve testing (FULL — all strings)",    "",   "",    "",    "Year 2 only"),
    ("Grass cutting / vegetation control",       "Yes","Yes", "Yes", "Yes"),
    ("SCADA / monitoring system check",          "Yes","Yes", "Yes", "Yes"),
    ("Insulation resistance testing",            "",   "",    "Yes", ""),
    ("Earth continuity testing",                 "",   "",    "Yes", ""),
    ("Perimeter / security inspection",          "Yes","Yes", "Yes", "Yes"),
    ("Post-storm emergency inspection",          "On demand - within 48 hours of trigger event", "", "", ""),
]
for i, row_d in enumerate(sched_rows):
    if len(row_d) == 2:
        add_table_row(sched_table, [
            (row_d[0], 5.0, "L", False),
            (row_d[1], 10.0, "C", True),
        ], alt=(i % 2 == 1))
    else:
        add_table_row(sched_table, [
            (row_d[0], 5.0, "L", False),
            (row_d[1], 2.5, "C", False),
            (row_d[2], 2.5, "C", False),
            (row_d[3], 2.5, "C", False),
            (row_d[4], 2.5, "C", False),
        ], alt=(i % 2 == 1))

doc.add_paragraph()

add_heading(doc, "Schedule B — Parks Technical Summary", level=1)
park_table = doc.add_table(rows=1, cols=4)
park_table.alignment = WD_TABLE_ALIGNMENT.LEFT
park_table.style = "Table Grid"
add_table_row(park_table, [
    ("Parameter",             5.0, "L", True),
    ("Anarita East Park",     4.0, "L", True),
    ("Anarita West Park",     4.0, "L", True),
    ("Combined",              3.0, "C", True),
], header=True)

park_rows = [
    ("Location",                    "Anarita, Paphos District",   "Anarita, Paphos District",  "—"),
    ("Installed PV Capacity",       "5.01 MWp",                   "5.01 MWp",                  "10.02 MWp"),
    ("AC Export Capacity",          "5.0 MW",                     "5.0 MW",                    "10.0 MW"),
    ("Grid Connection",             "22 kV (EAC)",                "22 kV (EAC)",               "—"),
    ("Annual O&M Fee (ex. VAT)",    f"€{PRICE_PER_PARK:,.0f}",    f"€{PRICE_PER_PARK:,.0f}",   f"€{PRICE_TOTAL:,.0f}"),
    ("VAT (19%)",                   f"€{VAT_PER_PARK:,.0f}",      f"€{VAT_PER_PARK:,.0f}",     f"€{VAT_TOTAL:,.0f}"),
    ("Annual O&M Fee (incl. VAT)",  f"€{PRICE_PER_PARK_INCL:,.0f}", f"€{PRICE_PER_PARK_INCL:,.0f}", f"€{PRICE_TOTAL_INCL:,.0f}"),
    ("Monthly Fee (incl. VAT)",     f"€{PRICE_MONTHLY_INCL/2:,.2f}", f"€{PRICE_MONTHLY_INCL/2:,.2f}", f"€{PRICE_MONTHLY_INCL:,.2f}"),
    ("Cleaning Sessions/yr",        "2",                          "2",                         "4 total"),
    ("Availability Guarantee",      "99%",                        "99%",                       "99% per park"),
]
for i, row_d in enumerate(park_rows):
    add_table_row(park_table, [
        (row_d[0], 5.0, "L", True),
        (row_d[1], 4.0, "L", False),
        (row_d[2], 4.0, "L", False),
        (row_d[3], 3.0, "C", False),
    ], alt=(i % 2 == 1))

# ══════════════════════════════════════════════════════════════════════════════
# SIGNATURE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "Execution", level=1)
add_paragraph(doc,
    "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date "
    "first written above.",
    size=10.5, space_after=16)

sig_table = doc.add_table(rows=1, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT

for i, (role, party_key, extra) in enumerate([
    ("SERVICE PROVIDER", "provider", f"Title: {CONTRACT['provider']['title']}"),
    ("CLIENT",           "client",   "Title: Authorised Signatory"),
]):
    cell = sig_table.rows[0].cells[i]
    cell.width = Cm(8)

    def sig_para(cell, text, bold=False, color=BLACK, after=4):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(10.5); r.font.color.rgb = color
        return p

    sig_para(cell, role, bold=True, color=NAVY)
    sig_para(cell, CONTRACT[party_key]["name"], bold=True)
    sig_para(cell, "", after=32)
    sig_para(cell, "Signature: ___________________________")
    sig_para(cell, f"Name: {CONTRACT[party_key]['contact']}")
    sig_para(cell, extra)
    sig_para(cell, "Date: ___________________________")

# remove first blank paragraph from each cell
for cell in sig_table.rows[0].cells:
    first = cell.paragraphs[0]
    if not first.text:
        p_elem = first._p
        p_elem.getparent().remove(p_elem)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print("Saved: " + OUT_FILE)
