"""
Generate simple LOI — HESS Psevdas BESS EPC (HV transformer supply).

Client-facing only: no OEM / supplier names disclosed.
Lighthief Cyprus Ltd = EPC Contractor; HESS = Client.

Output:
  L:\\My Drive\\LINYANG\\BESS CLIENTS\\Individual_60-120-standalone\\HV Transformer\\
      LOI-HESS-transformer-jun2026.docx
"""

from __future__ import annotations

import os
import sys
from datetime import date

sys.path.insert(0, os.path.dirname(__file__))
from loi_docx_common import (  # noqa: E402
    AMBER_HDR,
    AMBER_HEX,
    AMBER_TXT,
    GREY,
    GOLD,
    NAVY,
    NAVY_HEX,
    WHITE,
    add_footer,
    add_parties_table,
    add_run,
    body,
    h1,
    lock_table_widths,
    new_document,
    set_cell_bg,
    tbl_hdr,
)
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

REF = "Ref: LCY-LOI-HESS-TRF-2026-R1"
DOC_DATE = date.today().strftime("%d %B %Y")

CLIENT_COMPANY = "H.E.S.S. Hybrid Energy Storage Systems Ltd"
CLIENT_SIGNATORY = "Sotiris Shiakallis"
CLIENT_SIGNATORY_ORG = "Extensive Proficient Services Ltd (EPS)"
CLIENT_SIGNATORY_TITLE = "Director"
CLIENT_SIGNATORY_MOBILE = "+357 96 555 209"
CLIENT_SIGNATORY_ADDRESS = (
    "Nicosia Business Centre (NBC), 33 Neas Egkomis Street,\n"
    "2409 Egkomi, Nicosia, Cyprus"
)

CLIENT_BLOCK = (
    f"{CLIENT_COMPANY}\n"
    "Plot 26, Psevdas Community, Larnaca District, Cyprus\n"
    "CERA storage licence: KEA14-2024\n\n"
    f"Represented by: {CLIENT_SIGNATORY}\n"
    f"{CLIENT_SIGNATORY_TITLE}, {CLIENT_SIGNATORY_ORG}\n"
    f"{CLIENT_SIGNATORY_ADDRESS}\n"
    f"Mobile: {CLIENT_SIGNATORY_MOBILE}\n"
    "(signing for and on behalf of the Client as SPV for the Psevdas standalone BESS "
    "and transformer supply)\n\n"
    '(hereinafter "the Client")'
)


def add_header_bar_dated(doc, ref_line: str) -> None:
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_tbl.autofit = False
    lc = hdr_tbl.rows[0].cells[0]
    set_cell_bg(lc, NAVY_HEX)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(6)
    lp.paragraph_format.space_after = Pt(6)
    add_run(lp, "Lighthief", bold=True, size_pt=22, color=WHITE)
    add_run(lp, " Cyprus Ltd", bold=False, size_pt=11, color=WHITE)
    rc = hdr_tbl.rows[0].cells[1]
    set_cell_bg(rc, NAVY_HEX)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(6)
    rp.paragraph_format.space_after = Pt(6)
    for line, is_gold in [
        ("LETTER OF INTENT", True),
        (ref_line, False),
        (f"Date: {DOC_DATE}", False),
        ("STRICTLY CONFIDENTIAL", True),
    ]:
        add_run(rp, line + "\n", bold=is_gold, size_pt=8, color=GOLD if is_gold else WHITE)
    lock_table_widths(hdr_tbl, [10.0, 4.0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    doc = new_document()
    add_header_bar_dated(doc, REF)

    tp = doc.add_paragraph()
    add_run(tp, "Letter of Intent", bold=True, size_pt=18, color=NAVY)
    sp = doc.add_paragraph()
    add_run(
        sp,
        "HESS Psevdas — BESS EPC (132/33 kV Power Transformer Supply)",
        italic=True,
        size_pt=11,
        color=GREY,
    )

    add_parties_table(doc, CLIENT_BLOCK, client_label="Client")

    h1(doc, "RECITALS")
    for letter, text in [
        (
            "A.",
            "The Client is developing a standalone BESS at Plot 26, Psevdas, Larnaca "
            "(Hybrid Energy Storage Systems Ltd — TSOC ref. ΔΣΜΚ/ΠΟΣ/320.7.11).",
        ),
        (
            "B.",
            "Lighthief Cyprus Ltd is the proposed EPC contractor for the BESS plant and will "
            "design, procure, deliver, and install the 132/33 kV step-up transformer and "
            "related MV equipment at the KYEA.",
        ),
        (
            "C.",
            "The Parties wish to record their mutual intent to proceed so that Lighthief may "
            "reserve a January 2027 manufacturing slot for the Schedule 1 transformer and "
            "finalise the wider BESS EPC offer.",
        ),
    ]:
        rp = doc.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(rp, letter + "  ", bold=True, color=NAVY)
        add_run(rp, text)
    body(doc, "NOW THEREFORE, the Parties agree as follows:", space_after=8)

    h1(doc, "1.   COMMITMENT")
    body(
        doc,
        "1.1  The Client confirms its intention to appoint Lighthief as EPC contractor for "
        "the Psevdas BESS and to procure the Schedule 1 transformer through Lighthief as part "
        "of that EPC scope.",
    )
    body(
        doc,
        "1.2  This LOI supports manufacturing slot reservation, technical finalisation, and "
        "EPC pricing. Except where expressly stated as binding below, binding obligations arise "
        "only when the executed EPC and companion agreements are signed.",
    )

    h1(doc, "2.   SCHEDULE 1 — TRANSFORMER SCOPE (EPC)")
    spec_tbl = doc.add_table(rows=1, cols=2)
    spec_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_hdr(spec_tbl, ["Parameter", "Requirement"])
    specs = [
        ("Quantity", "1 unit"),
        ("Rating", "63 MVA continuous (ONAN / ONAF)"),
        ("Voltage ratio", "132 / 33 kV"),
        ("Vector group", "YNd11 (132 kV star / 33 kV delta)"),
        ("Short-circuit impedance uk", "21% on HV base @ 75 °C, CMR"),
        ("Cooling", "ONAN / ONAF — 2 cooler banks @ 50% CMR each"),
        ("OLTC", "+12.5% / −18.75%, 25 × 1.25% steps; local, remote, supervisory"),
        ("Standards", "Tier 2 Ecodesign (EU 2019/1783); EN 60076 series"),
        ("Installation", "Outdoor, oil-immersed; DAP Plot 26 Psevdas, Larnaca, Cyprus"),
        ("Delivery target", "On-site Q1 2027 (manufacturing slot January 2027)"),
        (
            "Included in Lighthief EPC scope",
            "Main 63 MVA unit; earthing and auxiliary transformer per Client GTR / EN 60289 "
            "(zig-zag earthing winding — solidly earthed, ≥20 kA / 3 s; "
            "station-service auxiliary winding ≥315 kVA @ 400/230 V); "
            "impact recorders; online DGA / monitoring per Client GTR",
        ),
        ("LV insulation (36 kV class)", "170 kV BIL / 70 kV AC — per Client clarification Jun 2026"),
        (
            "Excluded",
            "132 kV substation / ISM bay works; surge arresters on transformer (KYEA AIS scope)",
        ),
    ]
    for i, (param, val) in enumerate(specs):
        if i == 0:
            row = spec_tbl.rows[0]
        else:
            row = spec_tbl.add_row()
        row.cells[0].paragraphs[0].clear()
        add_run(row.cells[0].paragraphs[0], param, bold=True, size_pt=9)
        row.cells[1].paragraphs[0].clear()
        add_run(row.cells[1].paragraphs[0], val, size_pt=9)
    lock_table_widths(spec_tbl, [4.5, 9.5])

    h1(doc, "3.   PROGRAMME")
    body(doc, "3.1  Target manufacturing slot for Schedule 1 equipment: January 2027.")
    body(
        doc,
        "3.2  Target on-site delivery (DAP Limassol / Psevdas): Q1 2027, aligned with the Client "
        "Q2 2027 implementation target and TSOC connection programme.",
    )

    h1(doc, "4.   COMMERCIAL")
    body(
        doc,
        "4.1  Transformer and EPC pricing will be fixed in the executed EPC agreement — not in "
        "this LOI. This LOI enables Lighthief to reserve manufacturing capacity and finalise the "
        "BESS EPC quotation.",
    )
    body(
        doc,
        "4.2  Payment terms, warranties, and performance obligations will follow the executed EPC "
        "package and Client General Technical Requirements.",
    )
    body(
        doc,
        "4.3  Lighthief will procure Schedule 1 equipment through its supply chain. Manufacturer "
        "identity and sub-supplier arrangements remain confidential to Lighthief and are not "
        "disclosed under this LOI.",
    )

    h1(doc, "5.   CONFIDENTIALITY")
    body(
        doc,
        "5.1  (Binding) Each Party shall keep this LOI and related commercial and technical "
        "information confidential, except to advisers, lenders, or as required by law, for five "
        "(5) years after this LOI ends.",
    )

    h1(doc, "6.   TERM AND GOVERNING LAW")
    body(
        doc,
        "6.1  This LOI runs until the earlier of: (a) executed EPC and companion agreements; "
        "(b) ninety (90) days from signing; or (c) written termination.",
    )
    body(doc, "6.2  Governed by the laws of Cyprus. Courts of Cyprus have exclusive jurisdiction.")

    nb_tbl = doc.add_table(rows=1, cols=1)
    nb_cell = nb_tbl.rows[0].cells[0]
    set_cell_bg(nb_cell, AMBER_HEX)
    nbp = nb_cell.paragraphs[0]
    nbp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(nbp, "NATURE OF THIS DOCUMENT:  ", bold=True, size_pt=9, color=AMBER_HDR)
    add_run(
        nbp,
        "Non-binding except Clause 5.1 (confidentiality). Records mutual intent for BESS EPC "
        "appointment and Schedule 1 transformer supply. Binding commercial terms are in the "
        "executed EPC only.",
        size_pt=9,
        color=AMBER_TXT,
    )
    lock_table_widths(nb_tbl, [14.0])

    add_signatures_hess(doc)
    add_footer(doc, REF)
    return doc


def add_signatures_hess(doc) -> None:
    from loi_docx_common import BLACK, h1, body, add_run, set_cell_bg, lock_table_widths

    h1(doc, "SIGNATURES")
    body(
        doc,
        "IN WITNESS WHEREOF, the authorised representatives of the Parties have executed "
        "this Letter of Intent as of the date first written above.",
        space_after=10,
    )
    sig_tbl = doc.add_table(rows=6, cols=2)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    sig_data = [
        ("For and on behalf of\nLighthief Cyprus Ltd", f"For and on behalf of\n{CLIENT_COMPANY}"),
        ("Name:   Alexander Papacosta", f"Name:   {CLIENT_SIGNATORY}"),
        ("Title:    Cyprus Director", f"Title:    {CLIENT_SIGNATORY_TITLE}, {CLIENT_SIGNATORY_ORG}"),
        (
            "",
            f"Signing capacity:   Director of {CLIENT_SIGNATORY_ORG}, authorised to sign "
            f"for and on behalf of {CLIENT_COMPANY} (Client SPV)",
        ),
        ("Signature:   ___________________________", "Signature:   ___________________________"),
        ("Date:   ___________________________", "Date:   ___________________________"),
    ]
    for row_i, (lt, rt) in enumerate(sig_data):
        lc_s = sig_tbl.rows[row_i].cells[0]
        rc_s = sig_tbl.rows[row_i].cells[1]
        if row_i == 0:
            set_cell_bg(lc_s, NAVY_HEX)
            set_cell_bg(rc_s, NAVY_HEX)
            for cell, txt in ((lc_s, lt), (rc_s, rt)):
                sp = cell.paragraphs[0]
                add_run(sp, txt, bold=True, size_pt=9, color=WHITE)
        else:
            for cell, txt in ((lc_s, lt), (rc_s, rt)):
                sp = cell.paragraphs[0]
                add_run(sp, txt, size_pt=10, color=BLACK)
    lock_table_widths(sig_tbl, [7.0, 7.0])


def save_doc(doc) -> str:
    out_dir = (
        r"L:\My Drive\LINYANG\BESS CLIENTS\Individual_60-120-standalone\HV Transformer"
    )
    os.makedirs(out_dir, exist_ok=True)
    filename = "LOI-HESS-transformer-jun2026.docx"
    path = os.path.join(out_dir, filename)
    try:
        doc.save(path)
    except PermissionError:
        path = os.path.join(out_dir, "LOI-HESS-transformer-jun2026-rev2.docx")
        doc.save(path)
    return os.path.abspath(path)


if __name__ == "__main__":
    saved = save_doc(build())
    print("Saved -> " + saved)
