"""
Generate LOI — Esperia Energy Group + Galascope (standalone portfolio LOI)
Output: docs/clients/group-order/Group2_Esperia_Energy/contracts/
        LOI-Esperia-Energy-pipeline-may2026.docx

Data source: lib/portfolio-data.ts + esperia-energy.md (SSOT)
Scope:       Galascope 1 (5 MW / 20 MWh) + Galascope 2 (2.5 MW / 10 MWh) + nine (9) pipeline parks.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Brand colours ─────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1A, 0x36, 0x5D)
GOLD      = RGBColor(0xC9, 0xA4, 0x32)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)
GREY      = RGBColor(0x40, 0x40, 0x40)
AMBER_TXT = RGBColor(0x78, 0x35, 0x0F)
AMBER_HDR = RGBColor(0x92, 0x40, 0x0E)
NAVY_HEX  = "1A365D"
LIGHT_HEX = "EBF0F7"
AMBER_HEX = "FEF3C7"

# ── Portfolio data (SSOT: lib/portfolio-data.ts + esperia-energy.md) ──────────
GALASCOPE = [
    {"name": "Galascope 1", "district": "Famagusta", "mw": 5.0,  "mwh": 20, "containers": 5, "client_price": 2_238_000},
    {"name": "Galascope 2", "district": "Famagusta", "mw": 2.5,  "mwh": 10, "containers": 3, "client_price": 1_206_300},
]

PIPELINE = {
    "galascope": GALASCOPE,
    "phase1": [
        {"name": "Esperia Energy (Famagusta)",      "district": "Famagusta", "mw": 6.5,  "mwh": 20,   "containers": 4,  "client_price": 2_316_815, "epc_target": "Q2 2026", "cod": "Q4 2026"},
        {"name": "Esperia Green Energy (Limassol)", "district": "Limassol",  "mw": 8.0,  "mwh": 60,   "containers": 12, "client_price": 5_644_044, "epc_target": "Q2 2026", "cod": "Q4 2026"},
        {"name": "Esperia Energy (Frenaros)",       "district": "Famagusta", "mw": 25.0, "mwh": 100,  "containers": 20, "client_price": 10_088_014,"epc_target": "Q2 2026", "cod": "Q4 2026"},
    ],
    "phase2": [
        {"name": "Esperia Green Energy (Famagusta 2)", "district": "Famagusta", "mw": 5.0, "mwh": 20, "containers": 4,  "client_price": 2_189_416, "epc_target": "Q3 2027", "cod": "Q1 2028"},
    ],
    "phase3": [
        {"name": "Esperia Energy (Tseri)",    "district": "Nicosia", "mw": 7.0,  "mwh": 20,   "containers": 4, "client_price": 2_382_035, "epc_target": "2027", "cod": "Q2 2028"},
        {"name": "Esperia Energy (Tseri 2-A)","district": "Nicosia", "mw": 2.5,  "mwh": 7.5,  "containers": 2, "client_price": 1_084_620, "epc_target": "2027", "cod": "Q2 2028"},
        {"name": "Esperia Energy (Tseri 2-B)","district": "Nicosia", "mw": 7.5,  "mwh": 25,   "containers": 5, "client_price": 2_912_993, "epc_target": "2027", "cod": "Q3 2028"},
        {"name": "Esperia Energy (Tseri 2-C)","district": "Nicosia", "mw": 6.0,  "mwh": 20,   "containers": 4, "client_price": 1_551_848, "epc_target": "2027", "cod": "Q3 2028"},
        {"name": "Esperia Energy (Tseri 3)",  "district": "Nicosia", "mw": 4.75, "mwh": 15,   "containers": 3, "client_price": 1_874_143, "epc_target": "2027", "cod": "Q4 2028"},
    ],
}

PIPELINE_ONLY = PIPELINE["phase1"] + PIPELINE["phase2"] + PIPELINE["phase3"]
ALL_PARKS     = GALASCOPE + PIPELINE_ONLY

GAL_MW   = sum(p["mw"]  for p in GALASCOPE)
GAL_MWH  = sum(p["mwh"] for p in GALASCOPE)
GAL_CONT = sum(p["containers"] for p in GALASCOPE)
GAL_VAL  = sum(p["client_price"] for p in GALASCOPE)

PIPE_MW   = sum(p["mw"]  for p in PIPELINE_ONLY)
PIPE_MWH  = sum(p["mwh"] for p in PIPELINE_ONLY)
PIPE_CONT = sum(p["containers"] for p in PIPELINE_ONLY)
PIPE_VAL  = sum(p["client_price"] for p in PIPELINE_ONLY)

TOTAL_MW   = GAL_MW + PIPE_MW
TOTAL_MWH  = GAL_MWH + PIPE_MWH
TOTAL_CONT = GAL_CONT + PIPE_CONT
TOTAL_VAL  = GAL_VAL + PIPE_VAL

PHASE_ORDER = ("galascope", "phase1", "phase2", "phase3")

PHASE_TOTALS = {
    "galascope": {
        "label": "Galascope Ltd \u2014 Batch 1 (confirmed pricing May 2026)",
        "mw": GAL_MW, "mwh": GAL_MWH, "containers": GAL_CONT, "val": GAL_VAL,
    },
    **{
        k: {
            "label": lbl,
            "mw":  sum(p["mw"]  for p in PIPELINE[k]),
            "mwh": sum(p["mwh"] for p in PIPELINE[k]),
            "containers": sum(p["containers"] for p in PIPELINE[k]),
            "val": sum(p["client_price"] for p in PIPELINE[k]),
        }
        for k, lbl in (
            ("phase1", "Phase 1 \u2014 2026 Delivery (Esperia pipeline)"),
            ("phase2", "Phase 2 \u2014 2027 Delivery"),
            ("phase3", "Phase 3 \u2014 2028 Delivery (Tseri Portfolio)"),
        )
    },
}

# ── Helpers ───────────────────────────────────────────────────────────────────
def fmt_eur(v):
    return f"\u20ac{v:,.0f}"

def cm_to_twips(cm):
    return int(cm * 567)

def lock_table_widths(table, col_widths_cm):
    tbl   = table._tbl
    tblPr = tbl.tblPr
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(cm_to_twips(w) for w in col_widths_cm)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    tblPr.append(lay)
    existing = tbl.find(qn("w:tblGrid"))
    if existing is not None:
        tbl.remove(existing)
    tblGrid = OxmlElement("w:tblGrid")
    for w_cm in col_widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(cm_to_twips(w_cm)))
        tblGrid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(col_widths_cm):
                break
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(cm_to_twips(col_widths_cm[i])))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_run(para, text, bold=False, italic=False, size_pt=10, color=BLACK, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name  = font
    run.font.size  = Pt(size_pt)
    run.font.color.rgb = color
    return run

def para(doc, space_before=0, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = align
    return p

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, text, bold=True, size_pt=12, color=GOLD)
    # underline rule via border — simple approach: just use the run
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    add_run(p, text, bold=True, size_pt=10, color=NAVY)
    return p

def body(doc, text, space_after=5, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, text)
    return p

def bullet(doc, text, indent_cm=0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(indent_cm)
    add_run(p, text, size_pt=10)
    return p

def tbl_hdr(table, headers):
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, NAVY_HEX)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, hdr, bold=True, size_pt=8, color=WHITE)

def tbl_row(table, values, right_cols=None, total=False, shade=False):
    right_cols = right_cols or set()
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        if total:
            set_cell_bg(cell, NAVY_HEX)
        elif shade:
            set_cell_bg(cell, LIGHT_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in right_cols else (
            WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT)
        color = WHITE if total else BLACK
        add_run(p, str(val), bold=total, size_pt=8.5, color=color)
    return row

# ── Document ──────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ═══════════════════════════════════════════════════════════════════════════════
# HEADER BAR
# ═══════════════════════════════════════════════════════════════════════════════
hdr_tbl = doc.add_table(rows=1, cols=2)
hdr_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr_tbl.autofit   = False

lc = hdr_tbl.rows[0].cells[0]
set_cell_bg(lc, NAVY_HEX)
lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
lp = lc.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
lp.paragraph_format.space_before = Pt(6)
lp.paragraph_format.space_after  = Pt(6)
add_run(lp, "Lighthief", bold=True,  size_pt=22, color=WHITE)
add_run(lp, " Cyprus Ltd", bold=False, size_pt=11, color=WHITE)

rc = hdr_tbl.rows[0].cells[1]
set_cell_bg(rc, NAVY_HEX)
rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
rp = rc.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rp.paragraph_format.space_before = Pt(6)
rp.paragraph_format.space_after  = Pt(6)
for line, is_gold in [
    ("LETTER OF INTENT", True),
    ("Ref: LCY-LOI-ESP-PIPELINE-2026-R4", False),
    ("Date: May 2026", False),
    ("STRICTLY CONFIDENTIAL", True),
]:
    add_run(rp, line + "\n", bold=is_gold, size_pt=8,
            color=GOLD if is_gold else WHITE)

lock_table_widths(hdr_tbl, [10.0, 4.0])
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT TITLE
# ═══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
tp.paragraph_format.space_after = Pt(2)
add_run(tp, "Letter of Intent", bold=True, size_pt=18, color=NAVY)

sp = doc.add_paragraph()
sp.paragraph_format.space_after = Pt(10)
add_run(sp, "Esperia Energy Group \u2014 Galascope & BESS Pipeline Commitment",
        italic=True, size_pt=11, color=GREY)

# ═══════════════════════════════════════════════════════════════════════════════
# PARTIES
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "PARTIES")

parties_tbl = doc.add_table(rows=2, cols=2)
parties_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
parties_tbl.autofit   = False
parties_data = [
    ("Contractor", (
        "Lighthief Cyprus Ltd\n"
        "Registration No. HE 477423\n"
        "28 October Ave 249, Lophitis Business Center 1,\n"
        "Office 201, 3035 Limassol, Cyprus\n"
        "office@lighthief.com | +357 77 77 00 50\n"
        "(hereinafter \u201cLighthief\u201d or the \u201cContractor\u201d)"
    )),
    ("Client", (
        "Esperia Energy Group\n"
        "represented by: Dino Constantinou (Owner)\n"
        "Famagusta, Limassol and Nicosia districts, Cyprus\n"
        "\n"
        "(hereinafter \u201cthe Client\u201d)"
    )),
]
for row_i, (label, text) in enumerate(parties_data):
    lc_p = parties_tbl.rows[row_i].cells[0]
    rc_p = parties_tbl.rows[row_i].cells[1]
    set_cell_bg(lc_p, NAVY_HEX)
    lc_p.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lpp = lc_p.paragraphs[0]
    lpp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lpp.paragraph_format.space_before = Pt(4)
    lpp.paragraph_format.space_after  = Pt(4)
    add_run(lpp, label, bold=True, size_pt=10, color=WHITE)
    rpp = rc_p.paragraphs[0]
    rpp.paragraph_format.space_before = Pt(4)
    rpp.paragraph_format.space_after  = Pt(4)
    add_run(rpp, text, size_pt=9, color=BLACK)

lock_table_widths(parties_tbl, [2.5, 11.5])
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ═══════════════════════════════════════════════════════════════════════════════
# RECITALS
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "RECITALS")
recitals = [
    ("A.", "Lighthief supplies and installs grid-connected BESS for hybrid PV projects in Cyprus."),
    ("B.", "The Client (Esperia Energy Group) develops BESS projects in Famagusta, Limassol, and Nicosia."),
    ("C.", "The Client wishes to commit to the eleven (11) parks in Schedule 1, including Galascope 1 and Galascope 2 (Galascope Ltd) and nine (9) further group parks."),
]
for letter, text in recitals:
    rp = doc.add_paragraph()
    rp.paragraph_format.space_after = Pt(4)
    rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(rp, letter + "  ", bold=True, size_pt=10, color=NAVY)
    add_run(rp, text, size_pt=10)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
body(doc, "NOW THEREFORE, the Parties agree as follows:", space_after=8)

# ═══════════════════════════════════════════════════════════════════════════════
# 1. COMMITMENT
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "1.   COMMITMENT AND PURPOSE")
body(doc,
    "1.1  The Client confirms its intention to award BESS EPC contracts to Lighthief for "
    "eleven (11) parks in Schedule 1: Galascope 1 (5.0 MW / 20 MWh) and Galascope 2 "
    "(2.5 MW / 10 MWh) with Galascope Ltd under EPC ref. LCY-EPC-GAL-B1-2026; and nine "
    f"(9) further Esperia group parks ({PIPE_MW:.1f} MW / {PIPE_MWH:.1f} MWh). Combined "
    f"indicative value {fmt_eur(TOTAL_VAL)} (excl. VAT), based on Linyang quotation "
    "LY202601271 and group-order pricing.")
body(doc,
    "1.2  This LOI supports Lighthief\u2019s procurement and OEM planning. Except where "
    "expressly stated as binding below, this LOI is not an obligation to sign any EPC; "
    "each EPC (Galascope batch and each pipeline phase) is binding only when signed.")
body(doc,
    "1.3  (Binding) If the Client withdraws from Schedule 1 without valid reason after "
    "Lighthief has reasonably relied on this LOI, the Client shall reimburse documented "
    "pre-contractual costs directly incurred.")

# ═══════════════════════════════════════════════════════════════════════════════
# 2. SCHEDULE 1
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "2.   SCHEDULE 1 — PARKS AND INDICATIVE PRICING")
body(doc,
    "Indicative prices only. Final terms are set in each executed EPC (Lighthief EPC "
    "template v5.1 or successor).")

# Metrics strip
mx_tbl = doc.add_table(rows=1, cols=4)
mx_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
mx_tbl.autofit   = False
for i, (val, lbl) in enumerate([
    ("11 Parks",         "Schedule 1 Total"),
    (f"{TOTAL_MW:.2f} MW",   "Total BESS Power"),
    (f"{TOTAL_MWH:.1f} MWh", "Total BESS Energy"),
    (fmt_eur(TOTAL_VAL), "Indicative Value (ex. VAT)"),
]):
    cell = mx_tbl.rows[0].cells[i]
    set_cell_bg(cell, LIGHT_HEX)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    mp = cell.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_before = Pt(6)
    mp.paragraph_format.space_after  = Pt(6)
    add_run(mp, val + "\n", bold=True, size_pt=12, color=NAVY)
    add_run(mp, lbl, size_pt=7.5, color=GREY)
lock_table_widths(mx_tbl, [3.5, 3.5, 3.5, 3.5])
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Pipeline table — 6 columns, 14.0 cm total (body = 16 cm)
_PC = [5.2, 1.7, 1.3, 1.3, 1.0, 3.5]
pipe_tbl = doc.add_table(rows=1, cols=6)
pipe_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
pipe_tbl.autofit   = False
tbl_hdr(pipe_tbl, ["Park / Project", "District", "MW", "MWh", "Cont.", "Indicative Price (ex. VAT)"])

RC = {2, 3, 4, 5}
for phase_key in PHASE_ORDER:
    parks = PIPELINE[phase_key]
    for p in parks:
        tbl_row(pipe_tbl, [
            p["name"], p["district"],
            f"{p['mw']:.2f}", f"{p['mwh']:.1f}",
            str(p["containers"]), fmt_eur(p["client_price"]),
        ], right_cols=RC)
    pt = PHASE_TOTALS[phase_key]
    tbl_row(pipe_tbl, [
        pt["label"] + " \u2014 Subtotal", "",
        f"{pt['mw']:.2f}", f"{pt['mwh']:.1f}",
        str(pt["containers"]), fmt_eur(pt["val"]),
    ], right_cols=RC, total=True)

tbl_row(pipe_tbl, [
    "TOTAL SCHEDULE 1  (11 parks: 2 Galascope + 9 pipeline)", "",
    f"{TOTAL_MW:.2f}", f"{TOTAL_MWH:.1f}",
    str(TOTAL_CONT), fmt_eur(TOTAL_VAL),
], right_cols=RC, total=True)
lock_table_widths(pipe_tbl, _PC)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. PROGRAMME (INDICATIVE)
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "3.   INDICATIVE PROGRAMME")
body(doc,
    "Target dates are indicative; each EPC sets binding milestones. Pipeline phases "
    "generally require EAC/DSO grid connection and satisfactory PAC of the prior phase.")

prog_tbl = doc.add_table(rows=1, cols=5)
prog_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
prog_tbl.autofit   = False
tbl_hdr(prog_tbl, ["Batch / Phase", "Parks", "MWh", "EPC target", "Indicative value (ex VAT)"])
prog_rows = [
    ("Galascope Ltd \u2014 LCY-EPC-GAL-B1-2026", "2", 30.0, "Signed / Q2 2026", PHASE_TOTALS["galascope"]["val"]),
    ("Phase 1 \u2014 2026", "3", 180.0, "Q2 2026", PHASE_TOTALS["phase1"]["val"]),
    ("Phase 2 \u2014 2027", "1", 20.0,  "Q3 2027", PHASE_TOTALS["phase2"]["val"]),
    ("Phase 3 \u2014 Tseri 2028", "5", 87.5, "2027", PHASE_TOTALS["phase3"]["val"]),
]
for label, n_parks, mwh, epc_tgt, val in prog_rows:
    tbl_row(prog_tbl, [label, str(n_parks), f"{mwh:.1f}", epc_tgt, fmt_eur(val)], right_cols={2, 3, 4})
lock_table_widths(prog_tbl, [5.5, 1.0, 1.2, 2.0, 3.3])
doc.add_paragraph().paragraph_format.space_after = Pt(6)
body(doc,
    "Galascope target PAC: 31 January 2027. Galascope EPC is with Galascope Ltd; companion "
    "LTSA, EMS addendum, and OEM warranty undertaking per park are intended with that EPC.")

# ═══════════════════════════════════════════════════════════════════════════════
# 4. COMMERCIAL SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "4.   COMMERCIAL SUMMARY")
body(doc,
    "4.1  Indicative payment structure for each EPC (detail in the EPC): 30% advance; "
    "55% pre-shipment; 10% on PAC; 5% retention released after a three (3) month defects "
    "liability period following PAC.")
body(doc,
    "4.2  Prices exclude VAT. Client scope includes civil works, planning, permits, and "
    "grid connection fees unless agreed otherwise in the EPC.")
body(doc,
    "4.3  Schedule 1 pricing is indicative for ninety (90) days from this LOI. Material "
    "changes will be notified at least thirty (30) days before a phase EPC is signed.")

# ═══════════════════════════════════════════════════════════════════════════════
# 5. COMPANION AGREEMENTS
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "5.   COMPANION AGREEMENTS (INTENT)")
body(doc,
    "5.1  LTSA and EMS: The Client intends to enter LTSA(s) and, where applicable, "
    "DISPERON EMS subscription (Lighthief EU BESS Ltd) with each commissioned park. "
    "Indicative EMS rate EUR 400/MWh/year from PAC (separate addendum). Terms per "
    "the executed LTSA and EMS documents.")
body(doc,
    "5.2  Galascope: Companion documents for the Galascope batch (LTSA, EMS addendum, "
    "OEM warranty undertaking) are intended to be signed with the Galascope EPC.")

# ═══════════════════════════════════════════════════════════════════════════════
# 6. NEGOTIATION
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "6.   GOOD FAITH AND PHASE EXIT")
body(doc,
    "6.1  The Parties will negotiate each phase EPC in good faith. If terms are not "
    "agreed within sixty (60) days after the relevant EPC target date in Section 3, "
    "either Party may end this LOI for that phase only, without liability except "
    "Clause 1.3 costs.")

# ═══════════════════════════════════════════════════════════════════════════════
# 7. EXCLUSIVITY (BINDING)
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "7.   EXCLUSIVITY")
body(doc,
    "7.1  (Binding) Until this LOI expires or all Schedule 1 EPCs are signed, and for "
    "six (6) months after each park\u2019s PAC, the Client will not appoint another EPC "
    "contractor for any Schedule 1 park without Lighthief\u2019s consent.")
body(doc,
    "7.2  This does not apply if Lighthief fails to agree an EPC for a phase within "
    "sixty (60) days after the target date, or if Lighthief materially breaches this "
    "LOI and does not remedy within thirty (30) days.")

# ═══════════════════════════════════════════════════════════════════════════════
# 8. CONFIDENTIALITY (BINDING)
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "8.   CONFIDENTIALITY")
body(doc,
    "8.1  (Binding) Each Party shall keep confidential this LOI and related commercial, "
    "technical, and pricing information, except to advisers, lenders, or as required "
    "by law. This survives for five (5) years after this LOI ends.")

# ═══════════════════════════════════════════════════════════════════════════════
# 9. TERM AND GOVERNING LAW
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "9.   TERM AND GOVERNING LAW")
body(doc,
    "9.1  This LOI runs until the earlier of: all Schedule 1 EPCs signed, thirty-six "
    "(36) months from signing, or termination of a phase under Section 6.")
body(doc,
    "9.2  This LOI is governed by Cyprus law. Cyprus courts have exclusive jurisdiction.")
body(doc,
    "9.3  Amendments require written agreement. This LOI and Schedule 1 are the full "
    "record for the portfolio commitment described here.")
body(doc,
    "Notices: Lighthief \u2014 office@lighthief.com (Alexander Papacosta). "
    "Client \u2014 Dino Constantinou (details to be inserted).")

# ═══════════════════════════════════════════════════════════════════════════════
# NON-BINDING NOTICE BOX
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph().paragraph_format.space_after = Pt(4)
nb_tbl = doc.add_table(rows=1, cols=1)
nb_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
nb_tbl.autofit   = False
nb_cell = nb_tbl.rows[0].cells[0]
set_cell_bg(nb_cell, AMBER_HEX)
nb_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
nbp = nb_cell.paragraphs[0]
nbp.paragraph_format.space_before = Pt(6)
nbp.paragraph_format.space_after  = Pt(6)
nbp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(nbp, "NATURE OF THIS DOCUMENT:  ", bold=True, size_pt=9, color=AMBER_HDR)
add_run(nbp,
    "Non-binding except Clauses 1.3 (reliance costs), 7 (exclusivity), and 8 "
    "(confidentiality). This LOI covers Galascope 1 & 2 and nine (9) pipeline parks in "
    "Schedule 1. Each EPC is binding only when signed. Legal and commercial detail "
    "(warranty, LDs, APG, price confirmation, etc.) is in the relevant EPC v5.1, not here.",
    size_pt=9, color=AMBER_TXT)
lock_table_widths(nb_tbl, [14.0])

# ═══════════════════════════════════════════════════════════════════════════════
# SIGNATURES
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "SIGNATURES")
body(doc,
    "IN WITNESS WHEREOF, the authorised representatives of the Parties have executed "
    "this Letter of Intent as of the date first written above.", space_after=10)

sig_tbl = doc.add_table(rows=5, cols=2)
sig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
sig_tbl.autofit   = False
sig_data = [
    ("For and on behalf of\nLighthief Cyprus Ltd", "For and on behalf of\nEsperia Energy Group"),
    ("Name:   Alexander Papacosta",          "Name:   ___________________________"),
    ("Title:    Managing Director",           "Title:    ___________________________"),
    ("Signature:   ___________________________", "Signature:   ___________________________"),
    ("Date:   ___________________________",  "Date:   ___________________________"),
]
for row_i, (lt, rt) in enumerate(sig_data):
    lc_s = sig_tbl.rows[row_i].cells[0]
    rc_s = sig_tbl.rows[row_i].cells[1]
    if row_i == 0:
        set_cell_bg(lc_s, NAVY_HEX)
        set_cell_bg(rc_s, NAVY_HEX)
        for cell, txt in ((lc_s, lt), (rc_s, rt)):
            sp2 = cell.paragraphs[0]
            sp2.paragraph_format.space_before = Pt(5)
            sp2.paragraph_format.space_after  = Pt(5)
            add_run(sp2, txt, bold=True, size_pt=9, color=WHITE)
    else:
        for cell, txt in ((lc_s, lt), (rc_s, rt)):
            sp2 = cell.paragraphs[0]
            sp2.paragraph_format.space_before = Pt(7)
            sp2.paragraph_format.space_after  = Pt(7)
            add_run(sp2, txt, size_pt=10, color=BLACK)
lock_table_widths(sig_tbl, [7.0, 7.0])

# ═══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph().paragraph_format.space_after = Pt(8)
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fp,
    "Lighthief Cyprus Ltd  \u2502  HE 477423  \u2502  28 October Ave 249, Lophitis Business Center 1, "
    "Office 201, 3035 Limassol, Cyprus\n"
    "office@lighthief.com  \u2502  +357 77 77 00 50  \u2502  solarfarms.cy  \u2502  "
    "Ref: LCY-LOI-ESP-PIPELINE-2026-R4",
    size_pt=7.5, color=GREY)

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
base         = os.path.join(os.path.dirname(__file__), "..",
                            "docs", "clients", "group-order", "Group2_Esperia_Energy")
contracts_dir = os.path.join(base, "contracts")
os.makedirs(contracts_dir, exist_ok=True)
filename = "LOI-Esperia-Energy-pipeline-may2026.docx"

p1 = os.path.join(contracts_dir, filename)
doc.save(p1)
print("Saved -> " + os.path.abspath(p1))

p2 = os.path.join(base, filename)
try:
    doc.save(p2)
    print("Saved -> " + os.path.abspath(p2))
except PermissionError:
    tmp = os.path.join(base, "LOI-Esperia-Pipeline-apr2026-UPDATED.docx")
    doc.save(tmp)
    print("Root file locked by Word - saved as -> " + os.path.abspath(tmp))
    print("Close the old file in Word, then rename.")
