"""
Standard PV O&M Contract Generator — Lighthief Cyprus Ltd
=========================================================
Usage:
    Edit the CLIENT and PARKS sections below, then run:
    python generate-standard-pv-om-contract.py

Output: pv-om/contracts/output/PV-OM-Contract-[CLIENT]-[DATE].docx

Reference: See pv-om/internal/ for cost model, pricing benchmarks, and LESA rates.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import date

# ══════════════════════════════════════════════════════════════════════════════
# ▶ CONFIGURE CONTRACT HERE — edit this section for each new client
# ══════════════════════════════════════════════════════════════════════════════

CLIENT = {
    "name":          "[CLIENT COMPANY NAME]",
    "contact":       "[CONTACT PERSON NAME]",
    "address":       "[CLIENT ADDRESS], Republic of Cyprus",
    "vat_number":    "[VAT NUMBER]",
}

PARKS = [
    {
        "name":     "[Park Name 1]",
        "dc_mwp":  0.0,          # Installed PV capacity in MWp
        "ac_mw":   0.0,          # AC export capacity in MW
        "location": "[Location, District], Cyprus",
    },
    # Add more parks here if needed:
    # {
    #     "name":     "[Park Name 2]",
    #     "dc_mwp":  0.0,
    #     "ac_mw":   0.0,
    #     "location": "[Location, District], Cyprus",
    # },
]

# Pricing — set the annual fee per park (ex. VAT)
# Refer to pv-om/internal/pricing-benchmarks.md for guidance
PRICE_PER_PARK_YR   = 0.0        # EUR per park per year, ex. VAT
VAT_RATE            = 0.19       # 19% Cyprus VAT

# Dates
CONTRACT_DATE       = "[DD Month YYYY]"           # Date of signing
COMMENCEMENT_DATE   = "[DD Month YYYY]"           # O&M service start date
TERM_YEARS          = 2                           # Initial term in years

# Contract reference (format: LH-OM-[CLIENT_SHORT]-[YEAR]-[SEQ])
CONTRACT_REF        = "LH-OM-[CLIENT]-[YEAR]-001"

# Service parameters
CALLOUT_CAP         = 6          # On-site corrective call-outs included per park/year
CALLOUT_EXTRA_WD    = 600        # Extra call-out rate — weekday (EUR ex. VAT)
CALLOUT_EXTRA_EVE   = 750        # Extra call-out rate — evening/Saturday (EUR ex. VAT)
CALLOUT_EXTRA_SUN   = 900        # Extra call-out rate — Sunday/Bank Holiday (EUR ex. VAT)
AVAILABILITY_PCT    = 99.0       # Availability guarantee %
PR_GUARANTEE_PCT    = 75.0       # Performance Ratio guarantee %

# LESA sub-contractor (update if different from Dimos)
LESA = {
    "name":  "Dimos Demosthenos",
    "rate":  1600,               # EUR per park per year retainer
}

# ══════════════════════════════════════════════════════════════════════════════
# DO NOT EDIT BELOW THIS LINE
# ══════════════════════════════════════════════════════════════════════════════

# Calculations
TOTAL_DC_MWP    = sum(p["dc_mwp"] for p in PARKS)
TOTAL_AC_MW     = sum(p["ac_mw"]  for p in PARKS)
PARKS_COUNT     = len(PARKS)
PRICE_TOTAL     = PRICE_PER_PARK_YR * PARKS_COUNT
VAT_PER_PARK    = PRICE_PER_PARK_YR * VAT_RATE
VAT_TOTAL       = PRICE_TOTAL * VAT_RATE
PRICE_PER_PARK_INCL  = PRICE_PER_PARK_YR + VAT_PER_PARK
PRICE_TOTAL_INCL     = PRICE_TOTAL + VAT_TOTAL
PRICE_MONTHLY_EX     = PRICE_TOTAL / 12
PRICE_MONTHLY_INCL   = PRICE_MONTHLY_EX * (1 + VAT_RATE)

# Output path
today = date.today()
client_slug = CLIENT["name"].replace(" ", "-").replace("[", "").replace("]", "")[:20]
OUT_DIR  = os.path.join(os.path.dirname(__file__), "output")
OUT_FILE = os.path.join(OUT_DIR, f"PV-OM-Contract-{client_slug}-{today.strftime('%b%Y')}.docx")
os.makedirs(OUT_DIR, exist_ok=True)

# ── Brand colours ─────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1A, 0x36, 0x5D)
GOLD      = RGBColor(0xC9, 0xA4, 0x32)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)
GREY      = RGBColor(0x40, 0x40, 0x40)
NAVY_HEX  = "1A365D"
GOLD_HEX  = "C9A432"
LIGHT_HEX = "EBF0F7"
ALT_HEX   = "F5F7FA"

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(14 if level == 1 else 12)
    r.font.color.rgb = GOLD
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), GOLD_HEX if level == 1 else NAVY_HEX)
    pBdr.append(bot); pPr.append(pBdr)
    return p

def add_paragraph(doc, text="", bold=False, italic=False, size=11,
                  color=BLACK, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.color.rgb = color
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix); r1.bold = True
        r1.font.size = Pt(10.5); r1.font.color.rgb = BLACK
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK

def add_table_row(table, cells_data, header=False, alt=False):
    row = table.add_row()
    for i, (text, width, align_str, bold) in enumerate(cells_data):
        cell = row.cells[i]; cell.width = Cm(width)
        if header: set_cell_bg(cell, NAVY_HEX)
        elif alt:  set_cell_bg(cell, ALT_HEX)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold or header; run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE if header else BLACK
        p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align_str == "C"
                       else WD_ALIGN_PARAGRAPH.RIGHT if align_str == "R"
                       else WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21); section.page_height = Cm(29.7)
section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── Header bar ────────────────────────────────────────────────────────────────
hdr = doc.add_table(rows=1, cols=1)
hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
hc = hdr.rows[0].cells[0]; set_cell_bg(hc, NAVY_HEX)
p = hc.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
r = p.add_run("LIGHTHIEF CYPRUS LTD")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = WHITE
p2 = hc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(12)
r2 = p2.add_run("Photovoltaic Operation & Maintenance Services Agreement")
r2.bold = True; r2.font.size = Pt(12); r2.font.color.rgb = GOLD
doc.add_paragraph()

# ── Reference block ───────────────────────────────────────────────────────────
ref_table = doc.add_table(rows=4, cols=2); ref_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for label, value in [
    ("Contract Reference:",   CONTRACT_REF),
    ("Date of Signing:",      CONTRACT_DATE),
    ("Effective Date:",       CONTRACT_DATE),
    ("Service Commencement:", COMMENCEMENT_DATE),
]:
    row = ref_table.add_row()
    c0, c1 = row.cells[0], row.cells[1]
    c0.width = Cm(5); c1.width = Cm(11)
    r0 = c0.paragraphs[0].add_run(label)
    r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = GREY
    r1 = c1.paragraphs[0].add_run(value)
    r1.font.size = Pt(10); r1.font.color.rgb = BLACK
# Remove auto-added first blank row
first = ref_table.rows[0]
if not any(c.text for c in first.cells):
    for cell in first.cells:
        cell._tc.getparent().remove(cell._tc)

doc.add_paragraph()

# ── Parties block ─────────────────────────────────────────────────────────────
pt = doc.add_table(rows=1, cols=2); pt.alignment = WD_TABLE_ALIGNMENT.LEFT
for role, party in [("SERVICE PROVIDER", {
        "name": "Lighthief Cyprus Ltd", "contact": "Alexander Papacosta",
        "address": "Republic of Cyprus", "reg": "HE 477423"}),
    ("CLIENT", CLIENT)]:
    cell = pt.rows[0].cells[0 if role == "SERVICE PROVIDER" else 1]
    cell.width = Cm(8); set_cell_bg(cell, LIGHT_HEX)
    p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(6)
    r = p.add_run(role); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = NAVY
    for line in [party["name"], party.get("contact",""), party["address"]]:
        if line:
            lp = cell.add_paragraph(); lp.paragraph_format.space_after = Pt(1)
            lr = lp.add_run(line); lr.font.size = Pt(9.5); lr.font.color.rgb = BLACK
    cell.add_paragraph().paragraph_format.space_after = Pt(6)
doc.add_paragraph()

# ── Art 1 Definitions ─────────────────────────────────────────────────────────
add_heading(doc, "Article 1 — Definitions", level=1)
add_paragraph(doc, "In this Agreement:", size=10.5, space_after=4)
defs = [
    ('"Effective Date"',        f'the date of signing, being {CONTRACT_DATE}.'),
    ('"Service Commencement Date"', f'the date O&M services, KPI measurement, and billing commence, being {COMMENCEMENT_DATE}.'),
    ('"Contract Year"',         'each 12-month period from the Service Commencement Date.'),
    ('"Availability"',          f'the % of time Park inverters are operational and capable of producing power when in-plane irradiance > 50 W/m2, per IEC 61724-1 (5-min SCADA data), excluding Excused Downtime.'),
    ('"Availability Guarantee"',f'{AVAILABILITY_PCT}% annually per Park.'),
    ('"Curtailment"',           'any export reduction imposed by DSO/EAC, TSO, or any grid operator (incl. ripple control, net billing caps) — classified as Excused Downtime.'),
    ('"Excused Downtime"',      'Force Majeure; Client shutdowns; Curtailment; grid outages beyond Park substation; equipment defects outside this Agreement; planned maintenance windows; irradiance < 50 W/m2; SCADA data loss not attributable to Service Provider.'),
    ('"Service Hours"',         'Mon–Fri 08:00–18:00 Cyprus time, excl. public holidays. Critical/Major faults: 24/7.'),
    ('"Parks"',                 ", ".join([f'{p["name"]} ({p["dc_mwp"]} MWp / {p["ac_mw"]} MW AC, {p["location"]})' for p in PARKS]) + "."),
    ('"Performance Ratio"',     f'Measured vs theoretical yield per IEC 61724-1, Excused Downtime excluded.'),
    ('"Service Fee"',           'Annual remuneration per Article 9.'),
]
for term, defn in defs:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(term + "  "); r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
    r2 = p.add_run(defn); r2.font.size = Pt(10); r2.font.color.rgb = BLACK

# ── Art 2 Scope ───────────────────────────────────────────────────────────────
add_heading(doc, "Article 2 — Scope of Services (Custom Package)", level=1)

# Package overview table
pkg_table = doc.add_table(rows=1, cols=5); pkg_table.style = "Table Grid"
pkg_table.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_row(pkg_table, [
    ("Service", 5.8, "L", True), ("ECO", 1.5, "C", True),
    ("SILVER", 1.5, "C", True), ("GOLD", 1.5, "C", True),
    ("CUSTOM (Selected)", 2.7, "C", True),
], header=True)

pkg_rows = [
    ("SCADA 24/7 monitoring + monthly reports", "Yes","Yes","Yes","Yes"),
    ("Periodic inspection", "2x/yr","4x/yr","6x/yr","2x/yr"),
    ("IR thermography", "-","Yes","Yes","Yes (2x/yr)"),
    ("Panel cleaning", "-","1x/yr","2x/yr","2x/yr"),
    ("Vegetation / grass cutting", "-","2x/yr","4x/yr","4x/yr"),
    ("Corrective call-outs (incl.)", "0","2/park","4/park",f"{CALLOUT_CAP}/park"),
    ("Electrical safety inspections (annual)", "-","Yes","Yes","Yes"),
    ("Spare parts inventory management", "-","Yes","Yes","Yes"),
    ("Full IV-curve testing (all strings)", "-","-","Yes","Every 2 yrs"),
    ("Transformer general audit", "-","-","Yes","Add-on"),
    ("Insurance management", "-","Yes","Yes","—"),
    ("PPA / energy contract management", "-","-","Yes","—"),
    ("Security firm management", "-","-","Yes","—"),
    ("Warranty & claims admin", "-","-","Yes","Add-on"),
]
GOLD_LIGHT = "FFF8E1"
for i, (svc, eco, silver, gold, custom) in enumerate(pkg_rows):
    trow = pkg_table.add_row()
    for j, (txt, w) in enumerate(zip([svc,eco,silver,gold,custom],[5.8,1.5,1.5,1.5,2.7])):
        cell = trow.cells[j]; cell.width = Cm(w)
        if j == 4 and txt not in ("-","—"): set_cell_bg(cell, GOLD_LIGHT)
        elif i % 2: set_cell_bg(cell, ALT_HEX)
        p = cell.paragraphs[0]; run = p.add_run(txt)
        run.bold = (j == 4 and txt not in ("-","—"))
        run.font.size = Pt(9); run.font.color.rgb = BLACK
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

# Price footer
price_row = pkg_table.add_row()
for j, (txt, w) in enumerate(zip([
    "Annual fee ex. VAT","On request","On request","On request",
    f"EUR {PRICE_PER_PARK_YR:,.0f}/park"], [5.8,1.5,1.5,1.5,2.7])):
    cell = price_row.cells[j]; cell.width = Cm(w); set_cell_bg(cell, NAVY_HEX)
    p = cell.paragraphs[0]; run = p.add_run(txt)
    run.bold = True; run.font.size = Pt(9); run.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

add_paragraph(doc, "Detailed scope of each included service is set out in Articles 2.1–2.8.",
              size=10, color=GREY, space_before=4, space_after=8)

for heading, items in [
    ("2.1  Preventive Maintenance (2x per year per Park)", [
        "Full visual inspection: PV modules, mounting structures, cables, connectors, junction boxes.",
        "IR thermography of all PV strings and modules (hot-spot detection).",
        "IV-curve testing — sample (minimum 10% of strings) at each visit.",
        "DC isolator, string fuse, and SPD inspection and testing.",
        "Inverter inspection: firmware check, fault log review, cooling system, filter replacement.",
        "MV/LV switchgear: visual inspection, torque check, interlock function test.",
        "Meteorological station calibration check and irradiance sensor cleaning.",
        "Fence, gate, CCTV, and perimeter security inspection.",
        "Earthing and lightning protection continuity check.",
        "Vegetation management within 0.5 m of all mounting structures.",
        "Written visit report with photographs, fault register update, and recommendations.",
    ]),
    ("2.2  Panel Cleaning (2x per year per Park)", [
        "IEC TS 62788 compliant, using deionised or approved water and soft-brush equipment.",
        "No abrasive agents or high-pressure jets above 40 bar.",
        "Post-cleaning irradiance-normalised energy output recorded for benchmarking.",
    ]),
    ("2.3  Corrective Maintenance", [
        f"Critical fault (total Park shutdown, >50 kW loss): on-site within 4 hours; restoration target 24 hours.",
        "Major fault (>10% capacity loss, inverter failure, MV fault): on-site within 8 hours; restoration target 48 hours.",
        "Minor fault (string failure, sensor, comms): on-site within 24 hours; resolution within 5 business days.",
        "Cosmetic fault: at next preventive visit, no later than 90 days.",
        f"Included on-site call-outs: {CALLOUT_CAP}/park/year. Remote resolutions unlimited.",
    ]),
    ("2.4  SCADA / Performance Monitoring", [
        "24/7 SCADA monitoring with automatic alarm notifications.",
        "Monthly performance report by 10th of following month (Schedule C).",
        "Quarterly trend analysis (degradation, soiling, inverter deviations).",
        "Annual performance summary with year-on-year comparison.",
        "Immediate notification for any fault causing >5% instantaneous capacity loss.",
    ]),
    ("2.5  Electrical Safety Inspections (Annual)", [
        "Insulation resistance testing (DC + AC cabling, IEC 62446-1).",
        "Earth continuity testing and earth fault loop impedance measurement.",
        "Full IV-curve testing of all strings — every 2 Contract Years.",
        "Asset register maintenance per EAC requirements.",
    ]),
    ("2.6  Vegetation & Civil Maintenance", [
        "4x grass cutting and vegetation control per Park per year.",
        "Civil inspection (foundations, cable ducting, drainage); defect reporting.",
        "Post-storm inspection within 48 hours of wind > 80 km/h.",
        "Minor civil repairs up to EUR 500/incident included.",
    ]),
    ("2.7  Spare Parts & Consumables", [
        "On-island inventory: string fuses, DC/AC SPDs, inverter filters, connector caps.",
        "Consumable replacement during preventive/corrective visits included.",
        "Major components (inverters, transformers, modules) excluded.",
    ]),
    ("2.8  Health, Safety & Environmental", [
        "Cyprus Safety and Health at Work Law (89(I)/1996) and EU Directives.",
        "Site-specific HSE Plan and Method Statements for all activities.",
        "Toolbox talks documented before each visit.",
        "Licensed disposal of all waste materials (modules, oils, packaging).",
        "All personnel to carry Lighthief-issued site ID and log entry/exit in site visitor register.",
    ]),
]:
    add_heading(doc, heading, level=2)
    for item in items:
        add_bullet(doc, item)

# ── Art 3 Availability ────────────────────────────────────────────────────────
add_heading(doc, "Article 3 — Availability Guarantee & Performance Penalties", level=1)
add_heading(doc, "3.1  Availability Guarantee", level=2)
add_paragraph(doc,
    f"Service Provider guarantees minimum Availability of {AVAILABILITY_PCT}% per Park annually. "
    "Availability = [(Eligible Hours - Unavailable Hours) / Eligible Hours] x 100. "
    "Eligible Hours = Total Hours - Excused Downtime Hours. "
    "Measured from SCADA data at 5-minute resolution; inverter logs and grid meter as secondary/tertiary sources. "
    "A grace band of 0.5 percentage points applies (first penalty band activates below 98.5%). "
    "Disputes unresolved within 14 days referred to jointly appointed Independent Engineer (binding, 30 days).",
    size=10.5, space_after=6)

add_heading(doc, "3.2  Availability Penalty Schedule", level=2)
pen_table = doc.add_table(rows=1, cols=3); pen_table.style = "Table Grid"
add_table_row(pen_table, [
    ("Annual Availability", 5.0, "C", True),
    ("Penalty", 5.0, "C", True),
    ("Cap", 6.0, "C", True),
], header=True)
for i, row in enumerate([
    ("98.5% – 98.99%", "5% of annual fee/park", "—"),
    ("97.0% – 98.49%", "10% of annual fee/park", "—"),
    ("95.0% – 96.99%", "20% of annual fee/park", "—"),
    ("< 95.0%",        "30% of annual fee/park", "30% of annual fee/park"),
]):
    add_table_row(pen_table, [(v, w, "C", False) for v, w in zip(row, [5.0,5.0,6.0])], alt=(i%2==1))

add_paragraph(doc,
    f"Total annual penalties (Availability + PR combined) capped at 30% of annual fee per Park. "
    "Liquidated damages are the Client's sole financial remedy for Availability/PR shortfalls.",
    size=10, color=GREY, space_before=4, space_after=6)

add_heading(doc, "3.3  Performance Ratio Guarantee", level=2)
add_paragraph(doc,
    f"Service Provider guarantees minimum annual PR of {PR_GUARANTEE_PCT}% per Park (IEC 61724-1, Excused Downtime excluded). "
    "Excluded from PR guarantee: Curtailment/EAC limitations; sensor drift/failure; grid anomalies; "
    "design defects; module degradation beyond warranted rates; force majeure weather; regulatory/net billing caps.",
    size=10.5, space_after=4)
for item in [
    "Root-cause analysis within 10 business days of year-end shortfall identification.",
    "Corrective action plan within 20 business days.",
    "If PR < 70% for two consecutive Contract Years due to Service Provider factors: Client may terminate without penalty.",
]:
    add_bullet(doc, item)

# ── Art 4 Excluded Services ───────────────────────────────────────────────────
add_heading(doc, "Article 4 — Excluded Services", level=1)
add_heading(doc, "4.0  Package Exclusions (Custom vs GOLD)", level=2)
excl_table = doc.add_table(rows=1, cols=3); excl_table.style = "Table Grid"
add_table_row(excl_table, [
    ("Excluded Service", 7.0, "L", True),
    ("Available As", 3.0, "C", True),
    ("Reason", 6.0, "L", True),
], header=True)
for i, row in enumerate([
    ("Transformer general audit", "Add-on (Art. 10)", "Commission as needed per equipment condition"),
    ("Insurance management", "Not required", "Client maintains own property/liability insurance"),
    ("PPA/energy contract management", "Not applicable", "Cyprus Net Billing scheme — no PPA"),
    ("Security management", "Not required", "Client's existing arrangement"),
    ("Warranty & claims admin", "Add-on (Art. 10.4)", "Available as billable service"),
]):
    add_table_row(excl_table, [(v, w, a, False) for v, w, a in zip(row, [7.0,3.0,6.0],["L","C","L"])], alt=(i%2==1))

doc.add_paragraph()
add_heading(doc, "4.1  LESA / MV Sub-Contractor", level=2)
add_paragraph(doc,
    f"MV network fault response is fulfilled via licensed LESA sub-contractor "
    f"{LESA['name']} (retainer: EUR {LESA['rate']:,}/park/year, included in Service Fee). "
    "Response for MV Urgent Calls: within 6 hours. "
    "BESS O&M (where applicable) is covered under a separate LTSA agreement.",
    size=10.5, space_after=6)

add_heading(doc, "4.2  General Exclusions", level=2)
for item in [
    "PV module replacement (warranty admin available as add-on under Art. 10.4).",
    "Inverter, central inverter, or transformer replacement.",
    "MV cable replacement or substation upgrades.",
    "Civil engineering works beyond minor repairs (Art. 2.6).",
    "Grid connection modifications or operator compliance testing.",
    "Cybersecurity or IT infrastructure beyond SCADA maintenance under this Agreement.",
    "Works caused by vandalism, theft, or third-party damage (quoted separately).",
]:
    add_bullet(doc, item)

# ── Art 5 Client Obligations ──────────────────────────────────────────────────
add_heading(doc, "Article 5 — Client Obligations", level=1)
add_paragraph(doc, "The Client shall:", size=10.5, space_after=4)
for item in [
    "Within 7 days of signing, hand over keys, access codes, SCADA credentials, and all access materials via written handover protocol.",
    "Provide unrestricted access to all Parks at reasonable times; immediate access for emergencies.",
    "Maintain valid property and public liability insurance for all Parks.",
    "Keep Parks connected to SCADA; notify Service Provider of any communication outage.",
    "Notify Service Provider promptly of any fault, alarm, or unusual operating condition.",
    "Not perform or authorise maintenance or modification without Service Provider's prior written consent.",
    "Pay all Service Fees per Article 9.",
    "Provide warranties, manuals, as-built drawings, and grid documentation within 14 days of commencement.",
]:
    add_bullet(doc, item)

# ── Art 6 Reporting ───────────────────────────────────────────────────────────
add_heading(doc, "Article 6 — Reporting & Records", level=1)
for bold, text in [
    ("Monthly Performance Report: ", "By 10th of each month; basis for invoice. Required fields: per Schedule C."),
    ("Visit Report: ", "Within 5 business days of each site visit."),
    ("Annual Summary: ", "Within 30 days of each Contract Year-end."),
    ("Incident Report: ", "Within 24 hours of any Critical fault."),
]:
    add_bullet(doc, text, bold_prefix=bold)

add_paragraph(doc, "All reports delivered electronically (PDF). Service Provider maintains maintenance log and fault register.",
              size=10.5, space_after=6)

add_heading(doc, "Schedule C — Monthly Report Required Fields", level=2)
rpt_table = doc.add_table(rows=1, cols=3); rpt_table.style = "Table Grid"
add_table_row(rpt_table, [("Field", 5.5, "L", True), ("Unit", 3.5, "C", True), ("Notes", 7.0, "L", True)], header=True)
rpt_fields = [
    ("Reporting Period", "Month/Year", ""),
    ("Park Name", "Text", ""),
    ("Energy Yield", "MWh", "Total AC export"),
    ("Specific Yield", "kWh/kWp", ""),
    ("Performance Ratio", "%", "IEC 61724-1"),
    ("Availability", "%", "Per Art. 3 definition"),
    ("Irradiation (in-plane)", "kWh/m2", "On-site met station"),
    ("Critical Faults", "Count", "Duration + resolution"),
    ("Major Faults", "Count", "Duration + resolution"),
    ("Minor Faults", "Count", "Resolution date"),
    (f"Call-outs used (of {CALLOUT_CAP} cap)", "Count", ""),
    ("Cleaning Carried Out", "Yes/No", "Date if performed"),
    ("Maintenance Visit", "Yes/No/Scheduled", "Date if performed"),
    ("Recommended Actions", "Free text", ""),
]
for i, (f, u, n) in enumerate(rpt_fields):
    add_table_row(rpt_table, [(f,5.5,"L",True),(u,3.5,"C",False),(n,7.0,"L",False)], alt=(i%2==1))
doc.add_paragraph()

# ── Art 7 Term ────────────────────────────────────────────────────────────────
add_heading(doc, "Article 7 — Term & Renewal", level=1)
add_paragraph(doc,
    f"Agreement executed {CONTRACT_DATE}. O&M services commence {COMMENCEMENT_DATE}. "
    f"Initial Term: {TERM_YEARS} Contract Year(s). Auto-renews for successive 1-year periods unless "
    "90 days' written notice of non-renewal given before expiry. Fee escalation per Article 9.4.",
    size=10.5, space_after=6)

# ── Art 8 KPIs ────────────────────────────────────────────────────────────────
add_heading(doc, "Article 8 — Service Levels & KPIs", level=1)
kpi_table = doc.add_table(rows=1, cols=3); kpi_table.style = "Table Grid"
add_table_row(kpi_table, [("KPI", 6.0, "L", True), ("Target", 4.0, "C", True), ("Measurement", 6.0, "L", True)], header=True)
kpi_rows = [
    ("Annual Availability", f"≥ {AVAILABILITY_PCT}%", "Annual per IEC 61724-1"),
    ("Annual Performance Ratio", f"≥ {PR_GUARANTEE_PCT}%", "Annual per IEC 61724-1"),
    ("Critical Fault Response", "≤ 4 hours", "From notification receipt"),
    ("Major Fault Response", "≤ 8 hours", "From notification receipt"),
    ("Minor Fault Response", "≤ 24 hours", "From notification receipt"),
    ("Cosmetic Fault Resolution", "≤ 90 days", "From identification"),
    (f"On-site call-outs included", f"{CALLOUT_CAP}/park/year", "Remote resolutions unlimited"),
    ("Preventive Maintenance Visits", "2/park/year", "Spring + Autumn"),
    ("Panel Cleaning", "2/park/year", "Per preventive visit"),
    ("Vegetation Control", "4/park/year", ""),
    ("Monthly Report + Invoice", "By 10th of month", "Electronic PDF + VAT invoice"),
]
for i, (k, t, m) in enumerate(kpi_rows):
    add_table_row(kpi_table, [(k,6.0,"L",False),(t,4.0,"C",True),(m,6.0,"L",False)], alt=(i%2==1))

doc.add_paragraph()
add_heading(doc, "8.1  Response Time Breach — Service Credits", level=2)
sla_table = doc.add_table(rows=1, cols=4); sla_table.style = "Table Grid"
add_table_row(sla_table, [
    ("Fault Type", 3.5, "L", True), ("SLA", 2.5, "C", True),
    ("Credit / Hour Exceeded", 4.5, "C", True), ("Max / Incident", 4.5, "C", True),
], header=True)
for i, row in enumerate([
    ("Critical", "4h on-site", "0.5% of monthly fee", "5% of monthly fee"),
    ("Major",    "8h on-site", "0.25% of monthly fee", "3% of monthly fee"),
    ("Minor",    "24h on-site", "No credit — tracked only", "—"),
]):
    add_table_row(sla_table, [(v,w,"C",False) for v,w in zip(row,[3.5,2.5,4.5,4.5])], alt=(i%2==1))

doc.add_paragraph()

# ── Art 9 Service Fee ─────────────────────────────────────────────────────────
add_heading(doc, "Article 9 — Service Fee & Payment Terms", level=1)
add_heading(doc, "9.1  Annual Service Fee (Fixed Lump Sum)", level=2)

fee_table = doc.add_table(rows=1, cols=4); fee_table.style = "Table Grid"
add_table_row(fee_table, [
    ("Park", 5.5, "L", True),
    ("Annual Fee (ex. VAT)", 3.5, "R", True),
    (f"VAT ({int(VAT_RATE*100)}%)", 2.5, "R", True),
    ("Annual Fee (incl. VAT)", 3.5, "R", True),
], header=True)
for i, park in enumerate(PARKS):
    add_table_row(fee_table, [
        (park["name"], 5.5, "L", False),
        (f"EUR {PRICE_PER_PARK_YR:,.0f}", 3.5, "R", False),
        (f"EUR {VAT_PER_PARK:,.0f}", 2.5, "R", False),
        (f"EUR {PRICE_PER_PARK_INCL:,.0f}", 3.5, "R", False),
    ], alt=(i%2==1))

add_table_row(fee_table, [
    ("TOTAL", 5.5, "L", True),
    (f"EUR {PRICE_TOTAL:,.0f}", 3.5, "R", True),
    (f"EUR {VAT_TOTAL:,.0f}", 2.5, "R", True),
    (f"EUR {PRICE_TOTAL_INCL:,.0f}", 3.5, "R", True),
])

add_paragraph(doc,
    f"VAT at {int(VAT_RATE*100)}% per Cyprus VAT Law 95(I)/2000. Rate subject to change per applicable law at invoice date.",
    size=10, color=GREY, space_before=4, space_after=6)

add_heading(doc, "9.2  Monthly Payment", level=2)
add_paragraph(doc,
    f"Invoiced monthly in arrears by 10th of each month with Monthly Performance Report. "
    f"Monthly instalment: EUR {PRICE_MONTHLY_EX:,.2f} ex. VAT / EUR {PRICE_MONTHLY_INCL:,.2f} incl. VAT. "
    "Payment due within 30 days of invoice. First invoice pro-rated if commencement is mid-month.",
    size=10.5, space_after=6)

add_heading(doc, "9.3  Late Payment", level=2)
add_paragraph(doc,
    "Overdue amounts: interest at 2% p.a. above ECB base rate. "
    "Service Provider may suspend non-emergency services after 45 days non-payment (10 business days' notice).",
    size=10.5, space_after=6)

add_heading(doc, "9.4  Annual Escalation", level=2)
add_paragraph(doc,
    "From Year 2: annual adjustment by Cyprus HICP-CY (Eurostat). Maximum 4%. No adjustment in Year 1.",
    size=10.5, space_after=6)

add_heading(doc, "9.5  Additional Call-Out Charges", level=2)
add_paragraph(doc,
    f"On-site corrective call-outs beyond {CALLOUT_CAP}/park/year invoiced separately:",
    size=10.5, space_after=4)
for bold, text in [
    ("Weekday (08:00–18:00): ", f"EUR {CALLOUT_EXTRA_WD} ex. VAT per visit."),
    ("Evening / Saturday: ",    f"EUR {CALLOUT_EXTRA_EVE} ex. VAT per visit."),
    ("Sunday / Bank Holiday: ", f"EUR {CALLOUT_EXTRA_SUN} ex. VAT per visit."),
]:
    add_bullet(doc, text, bold_prefix=bold)
add_paragraph(doc, "Due within 14 days of invoice.", size=10, color=GREY, space_before=4, space_after=6)

# ── Art 10 Additional Works ───────────────────────────────────────────────────
add_heading(doc, "Article 10 — Additional & Urgent Works", level=1)
add_heading(doc, "10.1  Additional Works — Order Procedure", level=2)
add_paragraph(doc,
    "Works outside this Agreement require email-confirmed written order per Schedule D before commencement.",
    size=10.5, space_after=6)

add_heading(doc, "10.2  Pre-Agreed Add-On Rate Schedule", level=2)
rate_table = doc.add_table(rows=1, cols=3); rate_table.style = "Table Grid"
add_table_row(rate_table, [
    ("Additional Service", 7.0, "L", True),
    ("Coverage", 3.0, "C", True),
    ("Rate (ex. VAT)", 6.0, "R", True),
], header=True)
addon_rows = [
    ("Extra panel cleaning session", "Both parks (same visit)", "EUR 5,500"),
    ("Extra panel cleaning session", "Single park only", "EUR 3,200"),
    ("Extra vegetation / grass cutting", "Both parks (same visit)", "EUR 2,800"),
    ("Extra vegetation / grass cutting", "Single park only", "EUR 1,700"),
    ("Combined cleaning + cutting (same day)", "Both parks", "EUR 7,500"),
    ("Transformer general audit", "Per substation", "EUR 1,500"),
    ("Extra on-site call-out — weekday", "Per visit", f"EUR {CALLOUT_EXTRA_WD}"),
    ("Extra on-site call-out — evening/Sat", "Per visit", f"EUR {CALLOUT_EXTRA_EVE}"),
    ("Extra on-site call-out — Sun/holiday", "Per visit", f"EUR {CALLOUT_EXTRA_SUN}"),
    ("Warranty claim administration", "Per claim", "EUR 150/hr (min EUR 300)"),
]
for i, (svc, cov, rate) in enumerate(addon_rows):
    add_table_row(rate_table, [(svc,7.0,"L",False),(cov,3.0,"C",False),(rate,6.0,"R",True)], alt=(i%2==1))

doc.add_paragraph()
add_heading(doc, "10.3  Urgent Works Carve-Out", level=2)
add_paragraph(doc,
    "Service Provider may perform urgent works without pre-approval where delay would risk life, health, property, "
    "or park continuity. Client notified within 4 hours; cost summary within 5 business days; "
    "invoiced separately within 14 days.",
    size=10.5, space_after=6)

add_heading(doc, "10.4  SCADA Data Ownership & Cybersecurity", level=2)
for item in [
    "All SCADA and performance data is the exclusive property of the Client.",
    "Service Provider licensed to process data solely for delivering services under this Agreement.",
    "No data sharing with third parties without prior written Client consent.",
    "On termination: full SCADA data export to Client within 14 days (CSV or equivalent); all copies deleted.",
    "Security breach affecting SCADA access must be reported to Client within 4 hours.",
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_heading(doc, "Schedule D — Additional Works Order Form", level=2)
ord_table = doc.add_table(rows=1, cols=2); ord_table.style = "Table Grid"
add_table_row(ord_table, [("Field", 5.0, "L", True), ("To Be Completed", 11.0, "L", True)], header=True)
for i, (f, v) in enumerate([
    ("Order Reference", "[Auto-assigned]"),
    ("Date of Order", ""),
    ("Park(s) Affected", ""),
    ("Description of Works", ""),
    ("Agreed Cost (ex. VAT)", "EUR"),
    ("VAT (19%)", "EUR"),
    ("Total (incl. VAT)", "EUR"),
    ("Authorised By Client", "Name: _______________  Signature: _______________  Date: ________"),
    ("Confirmed By Contractor", "Name: _______________  Signature: _______________  Date: ________"),
    ("Target Completion", ""),
]):
    add_table_row(ord_table, [(f,5.0,"L",True),(v,11.0,"L",False)], alt=(i%2==1))

doc.add_paragraph()

# ── Art 11–18 ─────────────────────────────────────────────────────────────────
add_heading(doc, "Article 11 — Insurance", level=1)
add_paragraph(doc, "Service Provider shall maintain throughout the Term:", size=10.5, space_after=4)
for bold, text in [
    ("Public Liability: ", "minimum EUR 2,000,000 per occurrence."),
    ("Employers' Liability: ", "as required by Cyprus law."),
    ("Professional Indemnity: ", "minimum EUR 500,000 per claim."),
    ("Motor Vehicle: ", "third-party and on-site liability for all vehicles, including subcontractors."),
]:
    add_bullet(doc, text, bold_prefix=bold)

add_heading(doc, "Article 12 — Liability & Indemnity", level=1)
for num, text in [
    ("12.1", "Each Party indemnifies the other against third-party claims arising from its negligent acts or omissions."),
    ("12.2", "Service Provider aggregate liability capped at 100% of annual Service Fee per Contract Year, except for fraud, wilful misconduct, or death/injury caused by negligence."),
    ("12.3", "Neither Party liable for indirect, consequential, or special losses."),
    ("12.4", "Liquidated damages under Article 3 are the Client's sole remedy for Availability and PR shortfalls."),
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(num + "  "); r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
    r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK

add_heading(doc, "Article 13 — Force Majeure", level=1)
add_paragraph(doc,
    "Neither Party in breach for failures caused by Force Majeure (acts of God, war, terrorism, pandemic, "
    "government action, grid outages by national operator, civil unrest). Notifying Party must advise within 48 hours "
    "and resume performance as soon as reasonably practicable.",
    size=10.5, space_after=6)

add_heading(doc, "Article 14 — Termination", level=1)
for bold, text in [
    ("14.1  Termination for Convenience: ",
     "90 days' written notice, effective at end of Contract Year."),
    ("14.2  Termination for Cause: ",
     "Immediate by written notice for: material breach unremedied within 30 days; insolvency; fraud/wilful misconduct."),
    ("14.3  Consequences: ",
     "Service Provider returns records, SCADA credentials, and spare parts within 14 days. "
     "Accrued fees remain payable. Prepaid fees for post-termination periods refunded pro-rata."),
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(bold); r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
    r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK

add_heading(doc, "Article 15 — Confidentiality", level=1)
add_paragraph(doc,
    "Each Party keeps confidential all Confidential Information received from the other Party. "
    "No disclosure to third parties without prior written consent, except as required by law. "
    "Obligation survives termination for 3 years.",
    size=10.5, space_after=6)

add_heading(doc, "Article 16 — Governing Law & Dispute Resolution", level=1)
add_paragraph(doc,
    "Governed by the laws of the Republic of Cyprus. Disputes first submitted to mediation administered by "
    "the Cyprus Arbitration & Mediation Centre (CAMC) per its Mediation Rules. "
    "If unresolved within 45 days of mediator appointment, referred to the courts of the Republic of Cyprus.",
    size=10.5, space_after=6)

add_heading(doc, "Article 17 — General Provisions", level=1)
for bold, text in [
    ("Entire Agreement: ", "Supersedes all prior negotiations and representations."),
    ("Amendments: ", "In writing, signed by authorised representatives of both Parties."),
    ("Assignment: ", "Neither Party may assign without prior written consent (not unreasonably withheld)."),
    ("Severability: ", "Invalid provisions severed; remaining provisions continue in full force."),
    ("Notices: ", "In writing, email with read receipt or registered post to cover page addresses."),
    ("Counterparts: ", "May be executed in counterparts, each an original."),
]:
    add_bullet(doc, text, bold_prefix=bold)

# ── Schedules ─────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Schedule A — Annual Maintenance Plan", level=1)
add_paragraph(doc, "Minimum activities per Park per Contract Year:", size=10.5, space_after=6)
sched_table = doc.add_table(rows=1, cols=5); sched_table.style = "Table Grid"
add_table_row(sched_table, [
    ("Activity", 5.0, "L", True), ("Q1", 2.5, "C", True),
    ("Q2", 2.5, "C", True), ("Q3", 2.5, "C", True), ("Q4", 2.5, "C", True),
], header=True)
for i, row in enumerate([
    ("Full preventive maintenance visit",          "", "Yes", "", "Yes"),
    ("PV module cleaning",                         "", "Yes", "", "Yes"),
    ("IR thermography scan",                       "", "Yes", "", "Yes"),
    ("IV-curve testing (sample, min 10% strings)", "", "Yes", "", "Yes"),
    ("IV-curve testing (FULL — all strings)",      "", "", "", "Year 2 only"),
    ("Grass cutting / vegetation control",         "Yes","Yes","Yes","Yes"),
    ("SCADA / monitoring check",                   "Yes","Yes","Yes","Yes"),
    ("Insulation resistance testing",              "", "", "Yes", ""),
    ("Earth continuity testing",                   "", "", "Yes", ""),
    ("Perimeter / security inspection",            "Yes","Yes","Yes","Yes"),
    ("Post-storm inspection",                      "On demand — within 48h of trigger event","","",""),
]):
    add_table_row(sched_table, [(row[0],5.0,"L",False),(row[1],2.5,"C",False),
        (row[2],2.5,"C",False),(row[3],2.5,"C",False),(row[4],2.5,"C",False)], alt=(i%2==1))

doc.add_paragraph()
add_heading(doc, "Schedule B — Parks Technical Summary", level=1)
park_table = doc.add_table(rows=1, cols=1+len(PARKS)); park_table.style = "Table Grid"
headers = [("Parameter", 5.0, "L", True)] + [(p["name"], 5.0, "L", True) for p in PARKS]
add_table_row(park_table, headers, header=True)
for i, (param, *values) in enumerate([
    ("Location",)               + tuple(p["location"] for p in PARKS),
    ("Installed PV Capacity",)  + tuple(f"{p['dc_mwp']} MWp" for p in PARKS),
    ("AC Export Capacity",)     + tuple(f"{p['ac_mw']} MW" for p in PARKS),
    ("Annual Fee (ex. VAT)",)   + tuple(f"EUR {PRICE_PER_PARK_YR:,.0f}" for _ in PARKS),
    ("VAT (19%)",)              + tuple(f"EUR {VAT_PER_PARK:,.0f}" for _ in PARKS),
    ("Annual Fee (incl. VAT)",) + tuple(f"EUR {PRICE_PER_PARK_INCL:,.0f}" for _ in PARKS),
    ("Monthly (incl. VAT)",)    + tuple(f"EUR {PRICE_MONTHLY_INCL/PARKS_COUNT:,.2f}" for _ in PARKS),
    ("Cleaning Sessions/yr",)   + tuple("2" for _ in PARKS),
    ("Availability Guarantee",) + tuple(f"{AVAILABILITY_PCT}%" for _ in PARKS),
]):
    row_data = [(param, 5.0, "L", True)] + [(v, 5.0, "L", False) for v in values]
    add_table_row(park_table, row_data, alt=(i%2==1))

# ── Signature ─────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Execution", level=1)
add_paragraph(doc, "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written above.",
              size=10.5, space_after=16)
sig_table = doc.add_table(rows=1, cols=2); sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (role, name, title) in enumerate([
    ("SERVICE PROVIDER", "Alexander Papacosta", "Cyprus Director, Lighthief Cyprus Ltd"),
    ("CLIENT", CLIENT["contact"], "Authorised Signatory"),
]):
    cell = sig_table.rows[0].cells[i]; cell.width = Cm(8)
    for text, bold, color, after in [
        (role, True, NAVY, 4), (name, True, BLACK, 32),
        ("Signature: ___________________________", False, BLACK, 4),
        (f"Name: {name}", False, BLACK, 4),
        (f"Title: {title}", False, BLACK, 4),
        ("Date: ___________________________", False, BLACK, 4),
    ]:
        p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(10.5); r.font.color.rgb = color

# Remove blank first paragraph in sig cells
for cell in sig_table.rows[0].cells:
    first = cell.paragraphs[0]
    if not first.text:
        first._p.getparent().remove(first._p)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print("Saved: " + OUT_FILE)
print(f"Client: {CLIENT['name']}")
print(f"Parks: {PARKS_COUNT} x {TOTAL_AC_MW} MW AC ({TOTAL_DC_MWP} MWp)")
print(f"Annual fee: EUR {PRICE_TOTAL:,.0f} ex. VAT / EUR {PRICE_TOTAL_INCL:,.0f} incl. VAT")
print(f"Monthly: EUR {PRICE_MONTHLY_EX:,.2f} ex. VAT / EUR {PRICE_MONTHLY_INCL:,.2f} incl. VAT")
